"""
SOZO Clinical Handbook Generator — clones PD template and applies condition-specific replacements.
Usage: from handbooks.base_generator import build_handbook
       build_handbook(cdata)  # cdata = one condition dict from condition_data.py
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_PROJECT_ROOT = Path(__file__).resolve().parent.parent

# Ensure src/ is importable for shared helpers
sys.path.insert(0, str(_PROJECT_ROOT / "src"))
from sozo_generator.docx.legacy_helpers import (
    para_replace,
    para_set,
    cell_write,
    replace_table,
    apply_global_replacements,
)

TEMPLATE = _PROJECT_ROOT / "templates" / "gold_standard" / "Clinical_Handbook.docx"
OUTPUT_ROOT = Path("outputs/documents")


# ─── global text replacements ────────────────────────────────────────────────

def _apply_global_replacements(doc, c):
    """Replace every PD-specific phrase with condition-specific equivalents."""
    reps = [
        # Cover / headings
        ("PARKINSON'S DISEASE",                c["cover_title"]),
        ("Parkinson's Disease (PD)",            c["name"]),
        ("Parkinson's disease (PD)",            c["name"]),
        ("Parkinson's Disease",                 c["name"]),
        ("Parkinson's disease",                 c["name"]),
        ("Parkinson's",                         c["short"]),
        (" (PD)",                               f" ({c['abbr']})"),
        ("(PD) ",                               f"({c['abbr']}) "),
        (" PD ",                                f" {c['abbr']} "),
        (" PD.",                                f" {c['abbr']}."),
        (" PD,",                                f" {c['abbr']},"),
        ("for PD ",                             f"for {c['abbr']} "),
        ("for PD.",                             f"for {c['abbr']}."),
        ("for PD,",                             f"for {c['abbr']},"),
        ("for PD\n",                            f"for {c['abbr']}\n"),
        # Medication timing
        ("last levodopa dose, ON/OFF state",    c["med_check_item"]),
        ("medication state (ON/OFF, timing of last dose)", c["session_med_doc"]),
        # Exam step heading (Heading 3)
        ("4.1 Motor Examination",               c["exam_step_heading"]),
        ("Motor examination completed",         c["exam_checklist_item"]),
        # PRS domain label
        ("Motor Symptoms Domain",               c["prs_domain_label"]),
        # About text (just the PD-specific part)
        ("for Parkinson's Disease (PD) using",  f"for {c['name']} using"),
    ]
    apply_global_replacements(doc, reps)


# ─── paragraph-level targeted replacements ───────────────────────────────────

def _apply_para_replacements(doc, c):
    """Replace full paragraph text for paragraphs identified by index or content."""
    paras = doc.paragraphs

    # [26] TPS off-label callout body
    para_set(paras[26], c["tps_offlabel_body"])

    # [79] MANDATORY callout body
    para_set(paras[79], c["mandatory_consent_body"])

    # [110] Motor symptoms domain header
    para_set(paras[110], c["prs_domain_header_line"])

    # [132] Motor exam description
    para_set(paras[132], c["exam_step_desc"])

    # [283] False classification callout body
    para_set(paras[283], c["false_class_body"])


# ─── table-level replacements ────────────────────────────────────────────────

def _apply_table_replacements(doc, c):
    tables = doc.tables

    # Table 1: Available Modalities (5 rows, 4 cols)
    replace_table(tables[1], c["modality_table"])

    # Table 7: Off-Label Disclosure Summary (5 rows, 4 cols)
    replace_table(tables[7], c["offlabel_table"])

    # Table 9: SOZO PRS (9 rows, 2 cols)
    replace_table(tables[9], c["prs_table"])

    # Table 11: Clinical Exam (7 rows, 3 cols)
    replace_table(tables[11], c["exam_table"])

    # Table 13: Cognitive/Mood Screening (8 rows, 3 cols)
    replace_table(tables[13], c["screening_table"])

    # Table 14: Phenotype Identification (8 rows, 2 cols)
    replace_table(tables[14], c["phenotype_table"])

    # Table 23: Task Pairing (6 rows, 3 cols)
    replace_table(tables[23], c["task_pairing_table"])

    # Table 24: Response Domains (5 rows, 2 cols)
    replace_table(tables[24], c["response_domains_table"])

    # Table 28: Inclusion Criteria (7 rows, 2 cols)
    replace_table(tables[28], c["inclusion_table"])

    # Table 29: Absolute Exclusion (8 rows, 2 cols)
    replace_table(tables[29], c["exclusion_table"])

    # Table 30: Conditional Criteria (9 rows, 2 cols)
    replace_table(tables[30], c["conditional_table"])

    # Table 31: Phenotype-to-Protocol (9 rows, 4 cols)
    replace_table(tables[31], c["protocol_table"])


# ─── main entry point ────────────────────────────────────────────────────────

def build_handbook(c: dict):
    """Build handbook for one condition. c = condition data dict."""
    doc = Document(str(TEMPLATE))

    _apply_global_replacements(doc, c)
    _apply_para_replacements(doc, c)
    _apply_table_replacements(doc, c)

    slug = c["slug"]
    out_path = OUTPUT_ROOT / slug / "fellow" / f"SOZO_Clinical_Handbook_{slug}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  [OK] {out_path}")
    return out_path
