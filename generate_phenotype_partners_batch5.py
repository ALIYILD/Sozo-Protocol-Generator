# DEPRECATED: This script is superseded by the canonical generation pipeline.
# Use instead: GenerationService.generate(condition="...", tier="...", doc_type="...")
# Or CLI: PYTHONPATH=src python -m sozo_generator.cli.main build condition --condition <slug> --tier <tier> --doc-type <type>
# See docs/MIGRATION_PLAN.md for details.

from pathlib import Path
from docx import Document
from docx.shared import RGBColor

_PROJECT_ROOT = Path(__file__).resolve().parent

C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)


def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full: return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None
    size = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except Exception:
        color = None
    for r in para.runs: r.text = ""
    if fr:
        fr.text = replaced; fr.font.bold = bold
        if size: fr.font.size = size
        if color: fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs: r.text = ""
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type:
            run.font.color.rgb = C_BLACK


def _replace_table(table, rows):
    for ri in range(len(table.rows)):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells: _cell_write(cell, "")
            continue
        is_hdr = (ri == 0)
        for ci, cell in enumerate(row.cells):
            _cell_write(cell, str(rows[ri][ci]) if ci < len(rows[ri]) else "", bold=is_hdr, white=is_hdr)


def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs):
            _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs):
                        _para_replace(para, old, new)


def build(c):
    TEMPLATE = _PROJECT_ROOT / "templates" / "gold_standard" / "Phenotype_Classification.docx"
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs
    tables = doc.tables
    for old, new in [
        ("Parkinson's Disease (PD)", c["full"]),
        ("Parkinson's Disease", c["full"]),
        ("Parkinson's", c["short"]),
        (" (PD)", f" ({c['abbr']})"),
        (" PD ", f" {c['abbr']} "),
        (" PD.", f" {c['abbr']}."),
        (" PD,", f" {c['abbr']},"),
        ("for PD", f"for {c['abbr']}"),
        ("H&Y Stage", c["severity_field"]),
        ("Disease Duration", c["duration_field"]),
        ("PD Phenotype", f"{c['abbr']} Phenotype"),
        ("Hoehn & Yahr", "disease severity rating"),
        ("levodopa", c.get("med_name", "primary medication")),
    ]:
        _global_replace(doc, old, new)
    _para_set(paras[1], f"{c['short']} Phenotype Classification & Protocol Mapping")
    _para_set(paras[7], f"STEP 1 \u2014 {c['short']} Phenotype Determination")
    _para_set(paras[11], f"1.2 {c['short']} Phenotype Classification")
    _para_set(paras[14], f"ASSIGNED PHENOTYPE: {c['phenotype_codes']}")
    _para_set(paras[19], c["g1"])
    _para_set(paras[21], c["g2"])
    _para_set(paras[23], c["g3"])
    _para_set(paras[27], c["tps_ol"])
    _para_set(paras[40], c["sel_opt"])
    _replace_table(tables[0], [
        ["Field", "Details"],
        ["Patient Name", ""],
        ["File Number", ""],
        [c["severity_field"], ""],
        ["Date", ""],
        ["Treating Doctor", ""],
        [c["duration_field"], "___ years / ___ months"],
    ])
    for ti, key in [(1, "dc"), (2, "ph"), (3, "nw"), (5, "t5"), (6, "t6"), (7, "t7"), (9, "t9"), (12, "tavns"), (13, "ces"), (14, "seq")]:
        _replace_table(tables[ti], c[key])
    _cell_write(tables[15].rows[1].cells[0], c["resp"])
    _cell_write(tables[15].rows[1].cells[1], "\u25a1 Completed | Score: ___ / baseline: ___")
    _cell_write(tables[18].rows[1].cells[0], f"{c['abbr']} Phenotype")
    out = Path("outputs/documents") / c["slug"] / "partners" / f"Phenotype_Classification_Partners_{c['slug']}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  [OK] {out}")
    return out


OCD = {
    "slug": "ocd", "short": "OCD", "abbr": "OCD", "full": "Obsessive-Compulsive Disorder (OCD)",
    "severity_field": "Y-BOCS / OCI-R Severity", "duration_field": "OCD Duration",
    "med_name": "SSRI/clomipramine",
    "phenotype_codes": "\u25a1 CO  \u25a1 SY  \u25a1 HA  \u25a1 HO  \u25a1 PO  \u25a1 TR  \u25a1 TI",
    "g1": "3.1 Contamination / Symmetry (CO / SY)", "g2": "3.2 Harm / Hoarding (HA / HO)",
    "g3": "3.3 Pure-O / TR / Tic-Related (PO / TR / TI)",
    "tps_ol": "TPS is OFF-LABEL for OCD. Requires Doctor authorisation and off-label consent.",
    "sel_opt": "SELECTED OPTION: \u25a1 A (Contamination/Symmetry)  \u25a1 B (Harm/Hoarding)  \u25a1 C (TR-OCD)  \u25a1 D (Mixed)",
    "resp": "Reassess SOZO PRS + Y-BOCS",
    "dc": [["Data Point", "Completed", "Source"], ["Y-BOCS / OCI-R severity score", "\u25a1", "Rating Scales"], ["OCD symptom dimension profiling (contamination/harm/symmetry)", "\u25a1", "Clinical Interview"], ["Insight level (BIQ, BABS)", "\u25a1", "Insight Scales"], ["Tic comorbidity screening (YGTSS)", "\u25a1", "YGTSS / Interview"], ["Hoarding assessment (HRS, SI-R)", "\u25a1", "Hoarding Scales"], ["Treatment history (SRI trials, ERP)", "\u25a1", "Medical History"], ["Anxiety and depression comorbidity", "\u25a1", "PHQ-9 / GAD-7"], ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"]],
    "ph": [["Dominant Features", "Assign Phenotype", "Code"], ["Contamination fears, washing, cleaning compulsions dominant", "Contamination (CO)", "CO"], ["Need for symmetry, order, just-right feeling, arranging", "Symmetry (SY)", "SY"], ["Harm, aggression, religious, sexual obsessions dominant", "Harm (HA)", "HA"], ["Excessive acquisition, difficulty discarding, clutter dominant", "Hoarding (HO)", "HO"], ["Intrusive thoughts without overt compulsions, mental rituals", "Pure-O (PO)", "PO"], ["\u22652 SSRI failures, ERP non-response, chronic refractory course", "Treatment-Resistant (TR)", "TR"], ["OCD with comorbid tic disorder or Tourette syndrome", "Tic-Related (TI)", "TI"]],
    "nw": [["Phenotype", "Primary Network", "Secondary Network", "Tertiary"], ["CO", "Fronto-striatal (OFC-caudate)", "SN (Insula)", "CEN"], ["SY", "Fronto-striatal (SMA-putamen)", "CEN", "SN"], ["HA", "Limbic (OFC/amygdala)", "Fronto-striatal", "CEN"], ["HO", "OFC-ventral striatum", "CEN (DLPFC)", "Limbic"], ["PO", "Fronto-striatal", "DMN", "CEN"], ["TR", "Fronto-striatal + Limbic", "CEN bilateral", "SN"], ["TI", "SMN (SMA/M1)", "Fronto-striatal", "CEN"]],
    "t5": [["Goal", "Anode", "Cathode"], ["Pre-SMA inhibition \u2014 OCD drive reduction (CO/SY)", "Left DLPFC (anodal)", "Pre-SMA (cathodal)"], ["OFC downregulation (CO/HA)", "Left DLPFC", "Right OFC cathodal (Fp2)"], ["Bilateral DLPFC enhancement (TR)", "Left DLPFC", "Right supraorbital"]],
    "t6": [["Goal", "Anode", "Cathode"], ["SMA/M1 inhibition for tic-OCD (TI)", "Right DLPFC", "SMA cathodal"], ["DLPFC executive strengthening (PO)", "Left DLPFC", "Right supraorbital"]],
    "t7": [["Goal", "Anode", "Cathode"], ["OFC cathodal / DLPFC anodal (HA/HO)", "Left DLPFC", "Right OFC (Fp2) cathodal"], ["Bilateral DLPFC TR augmentation", "Bilateral DLPFC", "Extracephalic"]],
    "t9": [["Clinical Target", "TPS Protocol", "Phenotype"], ["Pre-SMA cathodal (OCD drive)", "FT1", "CO, SY"], ["OFC downregulation", "FT3", "CO, HA, HO"], ["Caudate-fronto-striatal circuit", "FT4", "SY, TR"], ["SMA tic suppression", "FT5", "TI"], ["Bilateral DLPFC TR-OCD", "FT7", "TR"], ["Amygdala fear circuit (HA)", "FT8", "HA, PO"], ["Ventral striatal (reward/hoarding)", "FT9", "HO"], ["OFC + DLPFC combined", "Multiple", "TR / severe"], ["Multi-network distributed OCD", "Multiple", "MX / complex"]],
    "tavns": [["Domain", "Details"], ["Indications", "Anxiety comorbidity, sleep disruption, autonomic dysregulation in OCD"], ["Protocol", "20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"], ["Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"]],
    "ces": [["Domain", "Details"], ["Indications", "Anxiety, insomnia, TR-OCD adjunct, all phenotypes"], ["Protocol", "20\u201360 min daily, Alpha-Stim AID, especially TR and anxiety-comorbid"], ["Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"]],
    "seq": [["Option", "Sequence", "Phenotype"], ["OPTION A \u2014 Contamination/Symmetry", "Wk 1\u20134: Pre-SMA cathodal + L-DLPFC anodal tDCS (5\u00d7/wk) \u2192 Wk 2\u20136: TPS Pre-SMA/OFC", "CO / SY"], ["OPTION B \u2014 Harm/Hoarding", "Wk 1\u20134: OFC cathodal tDCS + ERP sessions \u2192 Wk 3\u20136: TPS OFC/Limbic", "HA / HO"], ["OPTION C \u2014 Treatment-Resistant", "Wk 1\u20132: CES stabilization \u2192 Wk 2\u20136: Bilateral DLPFC tDCS + TPS OFC", "TR"], ["OPTION D \u2014 Tic-Related/Mixed", "Sequential: TPS SMA (tic) \u2192 TPS OFC (OCD) \u2192 Maintenance DLPFC tDCS", "TI / MX"]],
}

MS = {
    "slug": "ms", "short": "MS", "abbr": "MS", "full": "Multiple Sclerosis (MS)",
    "severity_field": "EDSS Score", "duration_field": "Disease Duration",
    "med_name": "disease-modifying therapy",
    "phenotype_codes": "\u25a1 MO  \u25a1 FA  \u25a1 CG  \u25a1 AT  \u25a1 PA  \u25a1 DE  \u25a1 WA",
    "g1": "3.1 Motor / Walking (MO / WA / AT)", "g2": "3.2 Fatigue / Cognitive (FA / CG)",
    "g3": "3.3 Pain / Depression (PA / DE)",
    "tps_ol": "TPS is OFF-LABEL for MS. Requires Doctor authorisation and off-label consent.",
    "sel_opt": "SELECTED OPTION: \u25a1 A (Motor)  \u25a1 B (Fatigue/Cognitive)  \u25a1 C (Pain/Mood)  \u25a1 D (Mixed)",
    "resp": "Reassess SOZO PRS + EDSS / MFIS / MSWS",
    "dc": [["Data Point", "Completed", "Source"], ["EDSS staging and functional systems", "\u25a1", "Neurological Exam"], ["Motor function (9HPT, T25FW, MSWS-12)", "\u25a1", "Motor Assessment"], ["Fatigue severity (MFIS, FSS)", "\u25a1", "Fatigue Scales"], ["Cognitive assessment (BICAMS, SDMT)", "\u25a1", "Neuropsychological"], ["Pain assessment (NRS, neuropathic type)", "\u25a1", "Rating Scales"], ["Depression screening (PHQ-9, HADS)", "\u25a1", "Rating Scales"], ["Balance and ataxia (BERG, DGI)", "\u25a1", "Physiotherapy"], ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"]],
    "ph": [["Dominant Features", "Assign Phenotype", "Code"], ["Motor weakness, spasticity, pyramidal signs dominant", "Motor (MO)", "MO"], ["Fatigue as primary disabling symptom, MS-related fatigue", "Fatigue (FA)", "FA"], ["Cognitive impairment: processing speed, memory, attention", "Cognitive (CG)", "CG"], ["Cerebellar ataxia, tremor, coordination deficit dominant", "Ataxia (AT)", "AT"], ["Central neuropathic pain, Lhermitte's, dysesthetic pain", "Pain (PA)", "PA"], ["MS-related depression, emotional lability, pseudo-bulbar", "Depression (DE)", "DE"], ["Walking impairment, gait disturbance, spasticity dominant", "Walking (WA)", "WA"]],
    "nw": [["Phenotype", "Primary Network", "Secondary Network", "Tertiary"], ["MO", "SMN (bilateral M1)", "SMA", "Corticospinal tract"], ["FA", "CEN (L-DLPFC)", "SN", "DMN"], ["CG", "CEN (bilateral DLPFC)", "DAN", "DMN"], ["AT", "Cerebellar network", "SMN", "DAN"], ["PA", "SN (Insula/dACC)", "SMN", "CEN (L-DLPFC)"], ["DE", "CEN (L-DLPFC)", "DMN (vmPFC)", "Limbic"], ["WA", "SMN (SMA/M1-LL)", "Cerebellar", "Brainstem locomotor"]],
    "t5": [["Goal", "Anode", "Cathode"], ["Bilateral M1 motor excitability (MO/WA)", "Bilateral M1", "Extracephalic"], ["L-DLPFC fatigue/cognitive (FA/CG)", "Left DLPFC (F3)", "Right supraorbital"], ["Cerebellar modulation for ataxia (AT)", "Cerebellar (Oz-based)", "Extracephalic"]],
    "t6": [["Goal", "Anode", "Cathode"], ["SN/DLPFC pain modulation (PA)", "Left DLPFC", "Right DLPFC cathodal"], ["L-DLPFC antidepressant (DE)", "Left DLPFC", "Right supraorbital"]],
    "t7": [["Goal", "Anode", "Cathode"], ["Bilateral DLPFC cognition (CG)", "Bilateral DLPFC", "Extracephalic"], ["SMA gait network (WA)", "SMA", "Contralateral mastoid"]],
    "t9": [["Clinical Target", "TPS Protocol", "Phenotype"], ["Bilateral M1 motor network", "FT1", "MO, WA"], ["L-DLPFC fatigue circuit", "FT3", "FA, CG"], ["Cerebellar-cortical ataxia", "FT4", "AT"], ["SN pain modulation", "FT5", "PA"], ["Bilateral DLPFC cognitive", "FT7", "CG"], ["Limbic-DLPFC depression", "FT8", "DE"], ["SMA walking network", "FT9", "WA"], ["Cerebellar + M1 combined", "Multiple", "AT / MO"], ["Multi-network distributed MS", "Multiple", "MX / progressive"]],
    "tavns": [["Domain", "Details"], ["Indications", "Fatigue, autonomic dysfunction, depression, pain in MS"], ["Protocol", "20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"], ["Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"]],
    "ces": [["Domain", "Details"], ["Indications", "Depression, pain, insomnia, fatigue in MS \u2014 adjunct to all phenotypes"], ["Protocol", "20\u201360 min daily, Alpha-Stim AID, especially DE, PA, FA phenotypes"], ["Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"]],
    "seq": [["Option", "Sequence", "Phenotype"], ["OPTION A \u2014 Motor/Walking", "Wk 1\u20134: Bilateral M1 tDCS during physio (5\u00d7/wk) \u2192 Wk 2\u20136: TPS M1/SMA", "MO / WA / AT"], ["OPTION B \u2014 Fatigue/Cognitive", "Wk 1\u20134: L-DLPFC tDCS + Daily taVNS \u2192 Wk 3\u20136: TPS DLPFC/cerebellar", "FA / CG"], ["OPTION C \u2014 Pain/Depression", "Wk 1\u20132: CES stabilization \u2192 Wk 2\u20136: L-DLPFC tDCS + TPS SN/DLPFC", "PA / DE"], ["OPTION D \u2014 Mixed/Progressive", "Sequential: TPS (highest burden) \u2192 Maintenance bilateral tDCS + taVNS", "MX / advanced"]],
}

if __name__ == "__main__":
    print("Generating Partners Phenotype Classification documents...")
    for condition in [OCD, MS]:
        print(f"\nBuilding: {condition['full']}")
        build(condition)
    print("\nDone.")
