"""
Batch 3 — Partners Phenotype Classification DOCX generator
Conditions: Post-Stroke Rehabilitation (stroke) + Traumatic Brain Injury (tbi)
"""

from pathlib import Path
from docx import Document
from docx.shared import RGBColor

C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)


def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None
    size = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except Exception:
        color = None
    for r in para.runs:
        r.text = ""
    if fr:
        fr.text = replaced
        fr.font.bold = bold
        if size:
            fr.font.size = size
        if color:
            fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type:
            run.font.color.rgb = C_BLACK


def _replace_table(table, rows):
    n_tmpl = len(table.rows)
    for ri in range(n_tmpl):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells:
                _cell_write(cell, "")
            continue
        is_hdr = (ri == 0)
        for ci, cell in enumerate(row.cells):
            text = str(rows[ri][ci]) if ci < len(rows[ri]) else ""
            _cell_write(cell, text, bold=is_hdr, white=is_hdr)


def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs):
            _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs):
                        _para_replace(para, old, new)


def build_phenotype_classification(c):
    TEMPLATE = Path(
        r"C:/Users/yildi/OneDrive/Desktop/Parkinson D/Partners/Assessments/PD_Phenotype_Classification_Partners.docx"
    )
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs
    tables = doc.tables

    for old, new in [
        ("Parkinson's Disease (PD)", c["full"]),
        ("Parkinson's Disease", c["full"]),
        ("Parkinson's", c["short"]),
        (" (PD)", f" ({c['abbr']})"),
        (" PD ", f" {c['abbr']} "),
        (" PD.", f" {c['abbr']}."),
        (" PD,", f" {c['abbr']},"),
        ("for PD", f"for {c['abbr']}"),
        ("H&Y Stage", c["severity_field"]),
        ("Disease Duration", c["duration_field"]),
        ("PD Phenotype", f"{c['abbr']} Phenotype"),
        ("Hoehn & Yahr", "disease severity rating"),
        ("levodopa", c.get("med_name", "primary medication")),
    ]:
        _global_replace(doc, old, new)

    _para_set(paras[1], f"{c['short']} Phenotype Classification & Protocol Mapping")
    _para_set(paras[7], f"STEP 1 \u2014 {c['short']} Phenotype Determination")
    _para_set(paras[11], f"1.2 {c['short']} Phenotype Classification")
    _para_set(paras[14], f"ASSIGNED PHENOTYPE: {c['phenotype_codes']}")
    _para_set(paras[19], c["group1_label"])
    _para_set(paras[21], c["group2_label"])
    _para_set(paras[23], c["group3_label"])
    _para_set(paras[27], c["tps_offlabel"])
    _para_set(paras[40], c["selected_option"])

    _replace_table(
        tables[0],
        [
            ["Field", "Details"],
            ["Patient Name", ""],
            ["File Number", ""],
            [c["severity_field"], ""],
            ["Date", ""],
            ["Treating Doctor", ""],
            [c["duration_field"], "___ years / ___ months"],
        ],
    )
    _replace_table(tables[1], c["data_collection_table"])
    _replace_table(tables[2], c["phenotype_table"])
    _replace_table(tables[3], c["network_table"])
    _replace_table(tables[5], c["tdcs_group1"])
    _replace_table(tables[6], c["tdcs_group2"])
    _replace_table(tables[7], c["tdcs_group3"])
    _replace_table(tables[9], c["tps_table"])
    _replace_table(tables[12], c["tavns_table"])
    _replace_table(tables[13], c["ces_table"])
    _replace_table(tables[14], c["sequencing_table"])
    _cell_write(tables[15].rows[1].cells[0], c["response_domain"])
    _cell_write(
        tables[15].rows[1].cells[1], "\u25a1 Completed | Score: ___ / baseline: ___"
    )
    _cell_write(tables[18].rows[1].cells[0], f"{c['abbr']} Phenotype")

    slug = c["slug"]
    out = (
        Path("outputs/documents")
        / slug
        / "partners"
        / f"Phenotype_Classification_Partners_{slug}.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  [OK] {out}")
    return out


# ── Condition definitions ─────────────────────────────────────────────────────

STROKE = {
    "slug": "stroke",
    "short": "Stroke",
    "abbr": "CVA",
    "full": "Post-Stroke Rehabilitation (CVA)",
    "severity_field": "NIHSS / Stroke Severity",
    "duration_field": "Time Since Stroke",
    "med_name": "anticoagulant/antiplatelet",
    "phenotype_codes": "\u25a1 MU  \u25a1 ML  \u25a1 AP  \u25a1 NE  \u25a1 CG  \u25a1 SP  \u25a1 DE",
    "group1_label": "3.1 Motor Phenotypes (MU / ML / SP)",
    "group2_label": "3.2 Language / Neglect (AP / NE)",
    "group3_label": "3.3 Cognitive / Depression (CG / DE)",
    "tps_offlabel": "TPS is OFF-LABEL for Post-Stroke Rehabilitation. Requires Doctor authorisation and off-label consent.",
    "selected_option": "SELECTED OPTION: \u25a1 A (Motor UL)  \u25a1 B (Language)  \u25a1 C (Cognitive)  \u25a1 D (Mixed)",
    "response_domain": "Reassess SOZO PRS + FMA / NIHSS / mRS",
    "data_collection_table": [
        ["Data Point", "Completed", "Source"],
        ["NIHSS / mRS / Fugl-Meyer Assessment", "\u25a1", "Neurological Exam"],
        ["Upper limb motor function (FMA-UL, grip)", "\u25a1", "Motor Assessment"],
        ["Lower limb / gait assessment (FAC, 10MWT)", "\u25a1", "Physiotherapy"],
        ["Language / aphasia screening (WAB, BNT)", "\u25a1", "Speech Pathology"],
        ["Neglect assessment (CBS, line bisection)", "\u25a1", "Neuropsychological"],
        ["Post-stroke depression (PHQ-9, HADS)", "\u25a1", "Rating Scales"],
        ["Cognitive screening post-stroke (MoCA)", "\u25a1", "MoCA / Interview"],
        ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"],
    ],
    "phenotype_table": [
        ["Dominant Features", "Assign Phenotype", "Code"],
        ["Upper limb paresis / plegia, fine motor deficit dominant", "Motor-UL (MU)", "MU"],
        ["Lower limb weakness, gait impairment, balance deficit dominant", "Motor-LL (ML)", "ML"],
        ["Expressive or receptive aphasia, language deficit dominant", "Aphasia (AP)", "AP"],
        ["Unilateral spatial neglect, hemispatial inattention dominant", "Neglect (NE)", "NE"],
        ["Post-stroke cognitive impairment, executive/memory deficit", "Cognitive (CG)", "CG"],
        ["Spasticity, increased tone, flexor/extensor pattern", "Spasticity (SP)", "SP"],
        ["Post-stroke depression, emotional lability, apathy", "Depression (DE)", "DE"],
    ],
    "network_table": [
        ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
        ["MU", "SMN (ipsilesional M1)", "SMA", "Contralesional M1"],
        ["ML", "SMN (SMA/M1-LL)", "Cerebellar", "Brainstem locomotor"],
        ["AP", "L-temporal (Broca/Wernicke)", "CEN (L-DLPFC)", "SMA"],
        ["NE", "DAN (R-parietal)", "SN", "CEN (R-DLPFC)"],
        ["CG", "CEN (L-DLPFC)", "DMN", "DAN"],
        ["SP", "SMN (ipsilesional M1)", "SMA", "Cerebellar"],
        ["DE", "Limbic (OFC/ACC)", "CEN (L-DLPFC)", "SN"],
    ],
    "tdcs_group1": [
        ["Goal", "Anode", "Cathode"],
        ["Ipsilesional M1 excitation (MU)", "Ipsilesional M1", "Contralesional M1"],
        ["SMA activation for gait/ML", "Ipsilesional SMA", "Contralesional shoulder"],
        ["Bilateral M1 rebalancing (SP)", "Ipsilesional M1", "Contralesional M1 cathodal"],
    ],
    "tdcs_group2": [
        ["Goal", "Anode", "Cathode"],
        ["Broca's area stimulation for aphasia", "L-IFG / Broca's (F7)", "Right supraorbital"],
        ["Right parietal inhibition for neglect", "Left parietal (P3)", "Right parietal cathodal"],
    ],
    "tdcs_group3": [
        ["Goal", "Anode", "Cathode"],
        ["L-DLPFC cognitive rehabilitation (CG)", "Left DLPFC (F3)", "Right supraorbital"],
        ["Limbic modulation for depression (DE)", "Left DLPFC", "Right DLPFC cathodal"],
    ],
    "tps_table": [
        ["Clinical Target", "TPS Protocol", "Phenotype"],
        ["Ipsilesional M1 cortical excitability", "FT1", "MU, SP"],
        ["SMA-M1 motor network (gait)", "FT3", "ML, MU"],
        ["Broca's / Wernicke's language area", "FT4", "AP"],
        ["Right parietal neglect network", "FT5", "NE"],
        ["L-DLPFC cognitive network", "FT7", "CG"],
        ["Limbic-prefrontal depression circuit", "FT8", "DE"],
        ["Cerebellar-cortical balance circuit", "FT9", "ML, SP"],
        ["Bilateral M1 ipsi/contra rebalancing", "Multiple", "SP / advanced"],
        ["Multi-network post-stroke", "Multiple", "MX / complex"],
    ],
    "tavns_table": [
        ["Domain", "Details"],
        ["Indications", "Post-stroke depression, autonomic dysregulation, aphasia recovery (VNS-paired therapy)"],
        ["Protocol", "20\u201330 min during rehab tasks, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
        ["\u25a1 Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "ces_table": [
        ["Domain", "Details"],
        ["Indications", "Post-stroke depression, anxiety, sleep disruption"],
        ["Protocol", "20\u201360 min daily, Alpha-Stim AID, especially DE and CG phenotypes"],
        ["\u25a1 Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "sequencing_table": [
        ["Option", "Sequence", "Phenotype"],
        ["OPTION A \u2014 Motor UL", "Wk 1\u20134: Ipsilesional M1 tDCS during OT (5\u00d7/wk) \u2192 Wk 2\u20136: TPS M1/SMA", "MU / SP"],
        ["OPTION B \u2014 Language", "Wk 1\u20134: Broca\u2019s tDCS during speech therapy \u2192 Wk 3\u20136: TPS language areas", "AP"],
        ["OPTION C \u2014 Cognitive", "Wk 1\u20132: CES/taVNS stabilization \u2192 Wk 2\u20136: L-DLPFC tDCS + TPS CEN", "CG / DE"],
        ["OPTION D \u2014 Mixed", "Sequential: TPS (highest deficit network) \u2192 Maintenance ipsilesional tDCS", "NE / MX"],
    ],
}

TBI = {
    "slug": "tbi",
    "short": "TBI",
    "abbr": "TBI",
    "full": "Traumatic Brain Injury (TBI)",
    "severity_field": "GCS / TBI Severity Rating",
    "duration_field": "Time Since Injury",
    "med_name": "neuroprotective agent",
    "phenotype_codes": "\u25a1 CG  \u25a1 EX  \u25a1 HA  \u25a1 BE  \u25a1 VE  \u25a1 FA  \u25a1 ME",
    "group1_label": "3.1 Cognitive / Executive Phenotypes (CG / EX / ME)",
    "group2_label": "3.2 Headache / Vestibular (HA / VE)",
    "group3_label": "3.3 Behavioral / Fatigue (BE / FA)",
    "tps_offlabel": "TPS is OFF-LABEL for TBI. Requires Doctor authorisation and off-label consent.",
    "selected_option": "SELECTED OPTION: \u25a1 A (Cognitive)  \u25a1 B (Headache)  \u25a1 C (Behavioral)  \u25a1 D (Mixed)",
    "response_domain": "Reassess SOZO PRS + COGSTATE / RPQ",
    "data_collection_table": [
        ["Data Point", "Completed", "Source"],
        ["TBI severity (GCS, LOC duration, PTA)", "\u25a1", "Medical History"],
        ["Cognitive profile (COGSTATE, MoCA, digit span)", "\u25a1", "Neuropsychological"],
        ["Post-concussion symptoms (RPQ, PCSI)", "\u25a1", "Rating Scales"],
        ["Headache characteristics (MIDAS, frequency/severity)", "\u25a1", "Headache Diary"],
        ["Behavioral / emotional dysregulation (NPI, BRIEF)", "\u25a1", "NPI / BRIEF"],
        ["Vestibular / balance assessment (BESS, DHI)", "\u25a1", "Physiotherapy"],
        ["Fatigue severity (FSS, MFI)", "\u25a1", "Rating Scales"],
        ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"],
    ],
    "phenotype_table": [
        ["Dominant Features", "Assign Phenotype", "Code"],
        ["Cognitive impairment: processing speed, attention, memory", "Cognitive (CG)", "CG"],
        ["Executive dysfunction: planning, inhibition, set-shifting", "Executive (EX)", "EX"],
        ["Post-traumatic headache dominant, migraine-like features", "Headache (HA)", "HA"],
        ["Behavioral dysregulation, impulsivity, aggression, disinhibition", "Behavioral (BE)", "BE"],
        ["Vestibular dysfunction, dizziness, balance impairment", "Vestibular (VE)", "VE"],
        ["Fatigue dominant: cognitive and/or physical exhaustion", "Fatigue (FA)", "FA"],
        ["Memory impairment: episodic, prospective memory deficit", "Memory (ME)", "ME"],
    ],
    "network_table": [
        ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
        ["CG", "CEN (bilateral DLPFC)", "DAN", "DMN"],
        ["EX", "CEN (L-DLPFC/R-IFG)", "Fronto-striatal", "SN"],
        ["HA", "SN (Insula/dACC)", "Thalamic", "SMN (trigeminal)"],
        ["BE", "CEN (OFC/DLPFC)", "Limbic (amygdala)", "SN"],
        ["VE", "Cerebellar-vestibular", "SMN", "DAN"],
        ["FA", "CEN (L-DLPFC)", "SN", "DMN"],
        ["ME", "DMN (hippocampal)", "CEN", "DAN"],
    ],
    "tdcs_group1": [
        ["Goal", "Anode", "Cathode"],
        ["Bilateral DLPFC cognitive restoration (CG/EX)", "Left DLPFC (F3)", "Right supraorbital"],
        ["Right IFG executive inhibition", "Right IFG (F8)", "Left supraorbital"],
        ["Hippocampal memory support (ME)", "Left temporal (T3)", "Right supraorbital"],
    ],
    "tdcs_group2": [
        ["Goal", "Anode", "Cathode"],
        ["Thalamic/SN modulation for headache (HA)", "Left DLPFC", "Right DLPFC cathodal"],
        ["Cerebellar modulation for vestibular (VE)", "Cerebellum (Oz-based)", "Extracephalic"],
    ],
    "tdcs_group3": [
        ["Goal", "Anode", "Cathode"],
        ["OFC/DLPFC behavioral circuit (BE)", "Left DLPFC (F3)", "Right supraorbital"],
        ["Frontal fatigue network (FA)", "Bilateral DLPFC", "Extracephalic"],
    ],
    "tps_table": [
        ["Clinical Target", "TPS Protocol", "Phenotype"],
        ["Bilateral DLPFC CEN restoration", "FT1", "CG, EX"],
        ["L-DLPFC executive network", "FT3", "EX, BE"],
        ["OFC behavioral regulation circuit", "FT4", "BE"],
        ["Hippocampal memory consolidation", "FT5", "ME"],
        ["Thalamic-SN headache modulation", "FT7", "HA"],
        ["Cerebellar-vestibular network", "FT8", "VE"],
        ["Frontal fatigue circuit", "FT9", "FA"],
        ["Bilateral frontal post-TBI", "Multiple", "CG / severe"],
        ["Multi-network distributed TBI", "Multiple", "MX / complex TBI"],
    ],
    "tavns_table": [
        ["Domain", "Details"],
        ["Indications", "Fatigue, sleep disruption, autonomic dysregulation post-TBI"],
        ["Protocol", "20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
        ["\u25a1 Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "ces_table": [
        ["Domain", "Details"],
        ["Indications", "Headache, anxiety, sleep disruption, behavioral dysregulation in TBI"],
        ["Protocol", "20\u201360 min daily, Alpha-Stim AID, especially HA, BE, FA phenotypes"],
        ["\u25a1 Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "sequencing_table": [
        ["Option", "Sequence", "Phenotype"],
        ["OPTION A \u2014 Cognitive", "Wk 1\u20134: Bilateral DLPFC tDCS (5\u00d7/wk) \u2192 Wk 2\u20136: TPS CEN/memory", "CG / EX / ME"],
        ["OPTION B \u2014 Headache", "Wk 1\u20132: CES stabilization \u2192 Wk 2\u20136: tDCS DLPFC + TPS SN/thalamic", "HA / VE"],
        ["OPTION C \u2014 Behavioral", "Wk 1\u20134: L-DLPFC tDCS + Daily taVNS \u2192 Wk 3\u20136: TPS OFC/DLPFC", "BE / FA"],
        ["OPTION D \u2014 Mixed", "Sequential: TPS (highest burden network) \u2192 Maintenance bilateral DLPFC", "MX / complex"],
    ],
}


if __name__ == "__main__":
    import os

    os.chdir(r"C:/Users/yildi/Sozo-Protocol-Generator")
    print("Generating Partners Phenotype Classification documents — Batch 3\n")

    results = []
    for condition in [STROKE, TBI]:
        print(f"Building: {condition['full']} ...")
        out = build_phenotype_classification(condition)
        results.append(out)

    print("\nVerification:")
    for p in results:
        size = Path(p).stat().st_size
        print(f"  {p}  ({size:,} bytes)")

    print("\nDone.")
