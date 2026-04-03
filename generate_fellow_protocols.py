# DEPRECATED: This script is superseded by the canonical generation pipeline.
# Use instead: GenerationService.generate(condition="...", tier="...", doc_type="...")
# Or CLI: PYTHONPATH=src python -m sozo_generator.cli.main build condition --condition <slug> --tier <tier> --doc-type <type>
# See docs/MIGRATION_PLAN.md for details.

#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
SOZO Evidence-Based Protocol Generator — Fellow Edition
Generates all 15 conditions faithful to the PD master template.
Run:  python generate_fellow_protocols.py
"""
import os, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# DESIGN CONSTANTS  (extracted from PD master template)
# ─────────────────────────────────────────────────────────────
FONT        = "Times New Roman"
C_NAVY      = "1B3A5C"   # H1, dark table headers
C_BLUE      = "2E75B6"   # H2, H3
C_RED       = "CC0000"   # Confidential / warnings
C_HDR_DARK  = "1A2540"   # Darkest table headers
C_WHITE     = "FFFFFF"

# Pastel row fills (from template)
F_BLUE    = "C1E4F5"
F_GREEN   = "E1F5EE"
F_BLUE2   = "E6F1FB"
F_GREEN2  = "EAF3DE"
F_PURPLE  = "EEEDFE"
F_GREY    = "F7F7F5"
F_PEACH   = "FAECE7"
F_CREAM   = "FAEEDA"
F_PINK    = "FBEAF0"
F_YELLOW  = "FFF8E1"

def _rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

# ─────────────────────────────────────────────────────────────
# DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────
def setup_page(doc):
    for s in doc.sections:
        s.page_width = Inches(8.5); s.page_height = Inches(11.0)
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1.0)

def shade(cell, fill):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), fill.lstrip("#")); pr.append(shd)

def borders(cell, color=C_NAVY, sz=4):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    bd = OxmlElement("w:tcBorders")
    for e in ["top","left","bottom","right"]:
        b = OxmlElement(f"w:{e}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),str(sz))
        b.set(qn("w:space"),"0");   b.set(qn("w:color"),color)
        bd.append(b)
    pr.append(bd)

def ct(cell, text, bold=False, color=None, size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text) if text else "")
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = _rgb(color)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)

def add_h(doc, text, level=1):
    try:
        p = doc.add_paragraph(style=f"Heading {level}")
    except Exception:
        p = doc.add_paragraph()
    p.clear()
    r = p.add_run(text)
    r.font.name = FONT; r.font.bold = True
    if level == 1:
        r.font.size = Pt(14); r.font.color.rgb = _rgb(C_NAVY)
        p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.font.size = Pt(12); r.font.color.rgb = _rgb(C_BLUE)
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    else:
        r.font.size = Pt(11); r.font.color.rgb = _rgb(C_BLUE)
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    return p

def add_p(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    if not text: return p
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = _rgb(color)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_rule(doc, color=C_BLUE):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"1");    bot.set(qn("w:color"),color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

def make_table(doc, headers, rows, hdr_fill=C_NAVY, row_fills=None, widths=None):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if widths:
        for ci, w in enumerate(widths):
            for cell in t.columns[ci].cells:
                cell.width = Inches(w)
    # Header
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        shade(c, hdr_fill); ct(c, h, bold=True, color=C_WHITE, size=10, center=True)
        borders(c, C_WHITE, 4)
    # Rows
    for ri, row in enumerate(rows):
        fill = (row_fills[ri] if row_fills and ri < len(row_fills) else None)
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            if fill: shade(c, fill)
            ct(c, val, size=10)
            borders(c, C_NAVY, 4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ─────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────
def add_cover(doc, cond):
    p = doc.add_paragraph()
    r = p.add_run("SOZO BRAIN CENTER \u2014 CYPRUS")
    r.font.name = FONT; r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = _rgb(C_NAVY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    r2 = p2.add_run("EVIDENCE-BASED NIBS CLINICAL PROTOCOL")
    r2.font.name = FONT; r2.font.size = Pt(14); r2.font.bold = True
    r2.font.color.rgb = _rgb(C_BLUE)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3 = doc.add_paragraph()
    r3 = p3.add_run(cond["name"].upper())
    r3.font.name = FONT; r3.font.size = Pt(16); r3.font.bold = True
    r3.font.color.rgb = _rgb(C_NAVY)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p4 = doc.add_paragraph()
    r4 = p4.add_run("Integrating tDCS (Newronika HDCkit & PlatoScience) & TPS (NEUROLITH\u00ae) with taVNS & CES")
    r4.font.name = FONT; r4.font.size = Pt(11)
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_rule(doc, C_NAVY)

    p5 = doc.add_paragraph()
    r5 = p5.add_run("DOCUMENT A \u2014 FULL CLINICAL PROTOCOL")
    r5.font.name = FONT; r5.font.size = Pt(12); r5.font.bold = True
    r5.font.color.rgb = _rgb(C_NAVY)
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p6 = doc.add_paragraph()
    r6 = p6.add_run("FELLOW EDITION")
    r6.font.name = FONT; r6.font.size = Pt(12); r6.font.bold = True
    r6.font.color.rgb = _rgb(C_BLUE)
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p7 = doc.add_paragraph()
    r7 = p7.add_run("Version 1.0 | April 2026 | For use by the treating Doctor only")
    r7.font.name = FONT; r7.font.size = Pt(10)
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p8 = doc.add_paragraph()
    r8 = p8.add_run("CONFIDENTIAL")
    r8.font.name = FONT; r8.font.size = Pt(11); r8.font.bold = True
    r8.font.color.rgb = _rgb(C_RED)
    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p9 = doc.add_paragraph()
    r9 = p9.add_run(
        "For authorised SOZO personnel only. This document is classified as confidential "
        "and is intended for internal clinical use only."
    )
    r9.font.name = FONT; r9.font.size = Pt(9)
    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_warning_box(doc,
        "TPS OFF-LABEL STATEMENT",
        f"TPS use in {cond['name']} is {cond.get('tps_status','INVESTIGATIONAL / OFF-LABEL')}. "
        "Explicit informed consent detailing the investigational nature of this application must be "
        "obtained and documented before initiation."
    )

    p10 = doc.add_paragraph()
    r10 = p10.add_run("\u00a9 2026 SOZO Brain Center. All rights reserved.")
    r10.font.name = FONT; r10.font.size = Pt(9)
    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_warning_box(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label + "  ")
    r1.font.name = FONT; r1.font.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = _rgb(C_RED)
    r2 = p.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

# ─────────────────────────────────────────────────────────────
# SECTION 1 — Document Control
# ─────────────────────────────────────────────────────────────
def s1_doc_control(doc, cond):
    add_h(doc, "1. Document Control & Clinical Responsibility")
    add_p(doc,
        "This section establishes the governance framework for the SOZO Evidence-Based Clinical Protocol "
        f"for {cond['name']}. It defines document ownership, clinical roles, regulatory context, and the "
        "off-label disclosure obligations applicable to all modalities used within this protocol."
    )
    add_h(doc, "Document Information", 2)
    make_table(doc,
        ["Field", "Details"],
        [
            ["Document Title", f"SOZO Evidence-Based Clinical Protocol \u2014 {cond['name']} (Fellow Edition)"],
            ["Document Number", cond["doc_num"]],
            ["Version", "1.0"],
            ["Effective Date", "[To be confirmed]"],
            ["Review Date", "12 months from effective date, or upon new significant evidence"],
            ["Classification", "CONFIDENTIAL \u2014 Internal Clinical Use Only"],
            ["Author", "Clinical Protocol Development Team"],
            ["Approved By", "[Medical Director signature]"],
        ],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None]*4,
        widths=[2.0, 4.5]
    )
    add_h(doc, "Roles and Responsibilities", 2)
    add_p(doc,
        "Clear delineation of clinical responsibilities ensures patient safety and regulatory compliance. "
        "The following roles apply across all SOZO clinical sites."
    )
    add_warning_box(doc, "CRITICAL CLINICAL GOVERNANCE",
        "No Clinical Assistant may independently modify a treatment plan, change a montage, or adjust "
        "parameters without explicit Doctor instruction. All off-label applications require documented "
        "Doctor authorisation and patient consent."
    )
    make_table(doc,
        ["Role", "Clinical Scope", "Responsibilities"],
        [
            ["Doctor",
             "Overall governance; protocol approval; treatment planning; montage selection; clinical decision-making",
             "Approves protocols; authorises off-label use; signs consent; designs treatment plans; "
             "selects montages; oversees AE reporting; conducts audit"],
            ["Clinical Assistant",
             "Device operation; session delivery; data recording; patient monitoring under Doctor supervision",
             "Delivers sessions per Doctor-prescribed plan; monitors patient; records session data; "
             "escalates all AEs to Doctor; operates within prescribed parameters only"],
        ],
        hdr_fill=C_NAVY, widths=[1.5, 2.5, 2.5]
    )
    add_h(doc, "Off-Label Disclosure", 2)
    make_table(doc,
        ["Modality", "Regulatory Status", "Classification", "Disclosure"],
        [
            ["tDCS", "CE-marked (Class IIa)", cond.get("tdcs_class", "Established NIBS; Level B evidence"), "Standard informed consent"],
            ["TPS",  "CE-marked (Class IIa)", cond.get("tps_class",  "INVESTIGATIONAL / OFF-LABEL"),       "Mandatory off-label disclosure; Doctor sign-off"],
            ["taVNS","CE-marked (Class IIa)", cond.get("tavns_class","Emerging adjunctive"),                "Informed consent with evidence disclosure"],
            ["CES",  "FDA-cleared; variable", cond.get("ces_class",  "Supportive adjunct ONLY"),            "Informed consent; state supportive role"],
        ],
        hdr_fill=C_NAVY, widths=[1.2, 2.0, 2.3, 2.0]
    )
    add_h(doc, "Document Structure", 2)
    add_p(doc,
        "This protocol is organised into thirteen integrated sections covering the full clinical pathway "
        "from patient selection through treatment delivery and long-term follow-up: (1) Document Control, "
        "(2) Inclusion/Exclusion, (3) Condition Overview, (4) Pathophysiology & Symptoms, "
        "(5) Key Brain Structures, (6) Clinical Phenotypes, (7) Symptom\u2013Brain Mapping, "
        "(8) tDCS Protocols, (9) TPS Protocols, (10) taVNS & CES Protocols, "
        "(11) Combination Techniques, (12) Side Effects & Monitoring, (13) Follow-Up & Decision-Making."
    )

# ─────────────────────────────────────────────────────────────
# SECTION 2 — Inclusion / Exclusion
# ─────────────────────────────────────────────────────────────
def s2_inclusion(doc, cond):
    add_h(doc, "2. Inclusion and Exclusion Criteria")
    add_h(doc, "Inclusion Criteria", 2)
    make_table(doc, ["#", "Criterion"],
        [[str(i+1), c] for i, c in enumerate(cond["inclusion"])],
        hdr_fill=C_NAVY,
        row_fills=[F_GREEN, None]*6,
        widths=[0.4, 6.1]
    )
    add_h(doc, "Exclusion Criteria (Absolute)", 2)
    make_table(doc, ["#", "Exclusion Criterion"],
        [[str(i+1), c] for i, c in enumerate(cond["exclusion"])],
        hdr_fill=C_RED,
        row_fills=[F_PEACH, None]*5,
        widths=[0.4, 6.1]
    )
    add_h(doc, "Conditions Requiring Discussion (NOT Absolute)", 2)
    make_table(doc, ["#", "Condition"],
        [[str(i+1), c] for i, c in enumerate(cond["discussion"])],
        hdr_fill="996633",
        row_fills=[F_YELLOW, None]*5,
        widths=[0.4, 6.1]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 3 — Condition Overview
# ─────────────────────────────────────────────────────────────
def s3_overview(doc, cond):
    add_h(doc, f"3. Condition Overview: {cond['name']}")
    for para in cond["overview"]:
        add_p(doc, para)

# ─────────────────────────────────────────────────────────────
# SECTION 4 — Pathophysiology & Symptoms
# ─────────────────────────────────────────────────────────────
def s4_pathophysiology(doc, cond):
    add_h(doc, "4. Pathophysiology, Symptoms & Standard Guidelines")
    for para in cond.get("pathophysiology", []):
        add_p(doc, para)
    add_h(doc, "Cardinal and Non-Motor Symptoms", 2)
    make_table(doc,
        ["Symptom Domain", "Key Features", "Prevalence", "References"],
        cond["symptoms"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None, F_BLUE2, None, F_GREEN, None, F_PURPLE, None,
                   F_GREY, None, F_PEACH, None, F_GREEN2],
        widths=[1.5, 2.5, 1.0, 1.5]
    )
    add_h(doc, "Standard Treatment Guidelines", 2)
    for para in cond.get("std_treatment", []):
        add_p(doc, para)

# ─────────────────────────────────────────────────────────────
# SECTION 5 — Key Brain Structures
# ─────────────────────────────────────────────────────────────
def s5_brain(doc, cond):
    add_h(doc, f"5. Key Brain Structures Involved in {cond['name']}")
    add_h(doc, "Key Brain Regions and Neuromodulation Rationale", 2)
    make_table(doc,
        ["Brain Region", "Role in Condition", "Neuromodulation Rationale", "References"],
        cond["brain_regions"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None]*6,
        widths=[1.8, 2.0, 2.0, 0.7]
    )
    add_h(doc, "Key Subcortical & Brainstem Structures", 2)
    make_table(doc,
        ["Brain Region", "Role in Condition", "Neuromodulation Rationale", "References"],
        cond["brainstem"],
        hdr_fill=C_HDR_DARK,
        row_fills=[F_PURPLE, None]*4,
        widths=[1.8, 2.0, 2.0, 0.7]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 6 — Clinical Phenotypes
# ─────────────────────────────────────────────────────────────
def s6_phenotypes(doc, cond):
    add_h(doc, "6. Clinical Phenotypes")
    add_h(doc, "Clinical Phenotype Classification", 2)
    make_table(doc,
        ["Phenotype", "Key Defining Features", "Predominant Symptoms", "Evidence-Based NIBS Priority"],
        cond["phenotypes"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None, F_GREEN, None, F_PURPLE, None, F_CREAM, None, F_PEACH],
        widths=[1.5, 2.0, 1.8, 1.2]
    )
    add_h(doc, "Symptom\u2013Brain Region\u2013Device Mapping", 2)
    make_table(doc,
        ["Symptom", "Primary Brain Region", "SOZO Device Strategy", "Evidence"],
        cond["symptom_map"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE2, None]*6,
        widths=[1.3, 1.5, 2.5, 1.2]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 7 — Symptom–Brain Mapping
# ─────────────────────────────────────────────────────────────
def s7_mapping(doc, cond):
    add_h(doc, "7. Symptom\u2013Brain Region Mapping for NIBS Targeting")
    add_h(doc, "Clinical Decision Pathway", 2)
    make_table(doc,
        ["Step", "Domain", "Action", "Responsible"],
        [
            ["1", "Clinical Assessment",    f"Complete baseline assessments ({', '.join(cond['outcomes_abbrev'][:3])})", "Doctor + Assistant"],
            ["2", "Phenotype Identification","Identify dominant clinical phenotype from Section 6", "Doctor"],
            ["3", "Protocol Selection",     "Select evidence-based protocol(s) from Sections 8\u20139 based on phenotype", "Doctor"],
            ["4", "Device Selection",       "Choose appropriate device based on protocol requirements", "Doctor"],
            ["5", "Response Evaluation",    "Assess at mid-point (week 4) and end-of-block (week 8\u201312)", "Doctor + Assistant"],
            ["6", "Adjustment",             "Reassess phenotype; modify protocol if needed", "Doctor"],
        ],
        hdr_fill=C_NAVY, widths=[0.4, 1.4, 3.5, 1.2]
    )
    add_h(doc, "Montage Selection Guide", 2)
    make_table(doc,
        ["Dominant Phenotype", "tDCS Strategy (Newronika HDCkit)", "TPS Strategy (NEUROLITH)", "Adjuncts"],
        cond["montage"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None, F_GREEN, None, F_PURPLE, None, F_CREAM, None, F_PEACH, None],
        widths=[1.5, 2.0, 1.8, 1.2]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 8 — tDCS
# ─────────────────────────────────────────────────────────────
def s8_tdcs(doc, cond):
    add_h(doc, "8. tDCS \u2014 Device Information, Specifications & Protocols")
    add_h(doc, "8A. Newronika HDCkit (Primary Clinical Device)", 2)
    add_h(doc, "Device Specifications", 3)
    make_table(doc,
        ["Specification", "Detail"],
        [
            ["Device Name",       "Newronika HDCkit"],
            ["Regulatory Status", "CE Class IIa Medical Device (CE 0476)"],
            ["Stimulation Type",  "2-channel direct current stimulation, independently controlled"],
            ["Output (per Channel)", "Maximum 2 mA per channel"],
            ["Active Electrode Area", "25\u201335 cm\u00b2 (anode)"],
            ["Return Electrode Area", "35\u201370 cm\u00b2 (passive electrode or cathode)"],
            ["Electrode Interface", "MindCap headset with integrated electrode pockets"],
            ["Home Mode",         "HDCprog controlled mode \u2014 Doctor programs via HDCprog; patient cannot modify parameters"],
            ["Safety Compliance", "IEC 60601-1-11 (home healthcare neurological stimulator standard)"],
            ["Clinical Role",     "Primary in-clinic tDCS system for all motor, cognitive, cerebellar, and dual-channel montages. "
                                  "Also used for home continuation phase under controlled mode."],
        ],
        hdr_fill=C_NAVY, row_fills=[F_BLUE, None]*5, widths=[2.0, 4.5]
    )
    add_h(doc, "Newronika HDCkit \u2014 Key Clinical Advantages", 3)
    for adv in [
        "Dual independent channels enable complex bilateral or multi-site montages",
        "HDCprog home mode allows controlled home therapy without patient-accessible parameter changes",
        "Full compatibility with classic SOZO protocol library (C1\u2013C8 and variants)",
        "MindCap electrode system ensures reproducible placement and consistent impedance",
    ]:
        add_p(doc, f"\u2022 {adv}")

    add_h(doc, "8B. PlatoScience (Wireless App-Controlled Device)", 2)
    add_h(doc, "Device Specifications", 3)
    make_table(doc,
        ["Specification", "Detail"],
        [
            ["Device Name",         "PlatoScience"],
            ["Regulatory Status",   "CE-marked wearable NIBS device"],
            ["Form Factor",         "One-size-fits-all wireless wearable"],
            ["Electrodes",          "3 stimulation electrodes (1 active anode, 2 reference)"],
            ["Stimulation Channels","1 channel (single-site anodal stimulation)"],
            ["Placement Variants",  "4 configurable head placements for different brain regions"],
            ["Output",              "Up to 1.6 mA adjustable intensity"],
            ["Ramp Time",           "10 seconds (default)"],
            ["Session Duration",    "0\u201330 minutes adjustable via smartphone app"],
            ["Control Interface",   "Smartphone app; structured guidance for placement and parameter entry"],
            ["Clinical Role",       "In-clinic and home-based NIBS where single-channel anodal stimulation is sufficient. "
                                    "NOT suitable for dual-channel or extracephalic return montages."],
        ],
        hdr_fill=C_NAVY, row_fills=[F_GREEN, None]*6, widths=[2.0, 4.5]
    )
    add_h(doc, "PlatoScience \u2014 Key Clinical Advantages", 3)
    for adv in [
        "Wireless; lightweight; no cables required \u2014 ideal for home therapy",
        "Smartphone app provides guided placement and real-time feedback",
        "One-size-fits-all design reduces setup time and training requirements",
        "Suitable for home continuation after in-clinic Newronika HDCkit induction phase",
    ]:
        add_p(doc, f"\u2022 {adv}")

    add_h(doc, "8C. Device Selection Guide", 2)
    make_table(doc,
        ["Clinical Scenario", "Recommended Device", "Rationale"],
        [
            ["Dual-channel montages (e.g., bilateral targets)", "Newronika HDCkit", "2 independent channels required"],
            ["Complex clinical phenotypes requiring flexibility", "Newronika HDCkit", "Full protocol library; flexible electrode configurations"],
            ["In-clinic sessions with higher current (2 mA)", "Newronika HDCkit", "Full 2 mA per channel capability"],
            ["Extracephalic cathode placement required",        "Newronika HDCkit", "PlatoScience has fixed reference electrodes"],
            ["Simple single-site anodal stimulation",          "PlatoScience",     "Single-channel sufficient; simpler setup"],
            ["Home-based therapy with app guidance",           "PlatoScience",     "Wireless; app-guided; patient-friendly"],
            ["Patients preferring wireless portability",        "PlatoScience",     "No cables; lightweight; one-size-fits-all"],
            ["Home continuation after in-clinic series",        "Either (Doctor decision)", "HDCprog mode or PlatoScience depending on montage needs"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_BLUE, None]*4, widths=[2.3, 1.8, 2.4]
    )
    add_h(doc, "8D. Polarity Principles in tDCS", 2)
    add_h(doc, "Clinical Polarity Principle", 3)
    make_table(doc,
        ["Polarity", "Neurophysiological Effect", "Clinical Application"],
        [
            ["Anode (+)", "Subthreshold depolarisation; increased neuronal firing probability; enhanced cortical excitability; facilitates LTP-like plasticity",
             f"Boost underactive regions in {cond['name']}: anodal placement over primary targets. Restore cortical excitability."],
            ["Cathode (\u2212)", "Subthreshold hyperpolarisation; decreased neuronal firing probability; reduced cortical excitability; facilitates LTD-like plasticity",
             "Regulate hyperactive regions; cathodal return placement; used for inhibitory effects where clinically indicated."],
            ["Extracephalic Return", "Reference electrode on arm/shoulder; current flows out of scalp",
             "Reduces focal stimulation at return; allows single-site anode focus on target region."],
        ],
        hdr_fill=C_NAVY, widths=[1.2, 3.0, 2.3]
    )

    add_h(doc, "8E. tDCS Parameters & Safety Limits", 2)
    add_h(doc, "Session Frequency Rule", 3)
    add_p(doc,
        "A minimum interval of 4\u20136 hours must be observed between same-day sessions. "
        "Back-to-back sessions without an interval are not permitted under any circumstances. "
        "Maximum: 2 sessions per day; 14 sessions per week."
    )
    add_h(doc, "tDCS Safety Limits", 3)
    make_table(doc,
        ["Parameter", "Standard (SOZO)", "DO NOT EXCEED"],
        [
            ["Intensity",         "Up to 2 mA/channel",                            "2 mA per channel"],
            ["Duration",          "20\u201330 min/session",                         "30 min per session"],
            ["Daily Dose",        "1\u20132 sessions/day",                          "2 sessions/day; 60 min max per day"],
            ["Minimum Interval",  "\u22654\u20136 hours",                           "No back-to-back sessions"],
            ["Weekly Schedule",   "Up to 14 sessions/week (2/day max)",             ">14/week not recommended"],
            ["Block Duration",    "8\u201312 weeks (2\u20133 months)",              "Beyond 12 weeks without follow-up review"],
            ["Sessions Before Review", "During in-clinic sessions on a daily basis","8\u201312 weeks without follow-up"],
            ["Impedance",         "<15 k\u03a9",                                    "Do not stimulate if >15 k\u03a9"],
            ["Skin",              "Intact, hydrated sponges",                       "Stop immediately if lesion/burn"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_BLUE, None]*5, widths=[1.8, 2.5, 2.2]
    )
    add_h(doc, "Device Operational Constraints", 3)
    make_table(doc,
        ["Constraint", "Newronika HDCkit", "PlatoScience"],
        [
            ["Regulatory/Class",    "CE Class IIa (CE 0476)",                                   "CE-marked (home)"],
            ["Channel Capacity",    "2 independent channels (dual-site stimulation)",            "1 channel (single-site anodal)"],
            ["Protocol Library",    "Full classic library (C1\u2013C8); custom montages",        "Simplified subset; placement-guided presets"],
            ["Return Electrode Options", "Flexible (contralateral, extracephalic, cathodal)",   "Fixed reference (2 passive electrodes)"],
            ["Home Mode",           "HDCprog controlled home therapy",                          "App-guided home therapy"],
            ["Home Use Control",    "Doctor programs; patient cannot change",                   "App-guided; follow prescriptions"],
        ],
        hdr_fill=C_NAVY, widths=[2.0, 2.3, 2.2]
    )
    add_h(doc, "8F. Classic tDCS Protocols \u2014 Evidence-Based (C1\u2013C8)", 2)
    make_table(doc,
        ["ID", "Symptom Target", "Anode", "Cathode", "Parameters", "Evidence"],
        cond["tdcs_protocols"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None, F_GREEN, None, F_PURPLE, None, F_CREAM, None],
        widths=[0.4, 1.5, 1.2, 1.2, 1.5, 0.7]
    )
    add_h(doc, "8G. tDCS Protocols for PlatoScience (Simplified C1\u2013C8)", 2)
    add_p(doc,
        f"The following table provides simplified PlatoScience protocol variants for {cond['name']}. "
        "PlatoScience supports single-channel anodal stimulation only. All intensities are capped at 1.6 mA. "
        "Dual-channel protocols from the Newronika HDCkit library are not reproducible on PlatoScience \u2014 "
        "the Doctor must select the most clinically relevant single-site variant."
    )
    make_table(doc,
        ["ID", "Target Symptom", "Placement Variant", "Anode (+)", "Reference (\u2212)", "Current & Time", "Frequency", "Key Notes"],
        cond["plato_protocols"],
        hdr_fill=C_HDR_DARK,
        row_fills=[F_BLUE2, None]*5,
        widths=[0.5, 1.0, 1.0, 0.8, 0.8, 0.9, 0.7, 0.8]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 9 — TPS
# ─────────────────────────────────────────────────────────────
def s9_tps(doc, cond):
    add_h(doc, "9. TPS \u2014 Device Information, Specifications & Protocols")
    add_h(doc, "9A. NEUROLITH\u00ae TPS System (Storz Medical)", 2)
    add_h(doc, "Device Specifications", 3)
    make_table(doc,
        ["Specification", "Details"],
        [
            ["Device Name",       "NEUROLITH\u00ae TPS System (Storz Medical)"],
            ["Technology",        "Transcranial Pulse Stimulation (focused acoustic shockwaves)"],
            ["Maximum Depth",     "Up to approximately 8 cm (device-dependent anatomical access)"],
            ["Energy Range",      "0.01\u20130.25 mJ/mm\u00b2"],
            ["Frequency Capability", "1\u20138 Hz (pulses per second)"],
            ["Navigation",        "MRI-based neuronavigation (e.g., BodyTrack\u00ae if used)"],
            ["Setting",           "Clinic use only (not home-deployable)"],
            ["Regulatory Status", f"CE-marked for Alzheimer\u2019s disease; {cond['name']} use is {cond.get('tps_status','OFF-LABEL')}"],
            ["Clinical Role",     "Targeted deep brain stimulation for motor, cognitive, and affective regions; "
                                  "reaches subcortical structures inaccessible to tDCS"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_BLUE, None]*5, widths=[2.0, 4.5]
    )
    add_h(doc, "9B. TPS Energy & Frequency Limits", 2)
    add_h(doc, "TPS Daily & Weekly Limits", 3)
    make_table(doc,
        ["Parameter", "SOZO Standard", "Notes"],
        [
            ["Energy Flux Density", "0.01\u20130.25 mJ/mm\u00b2",       "Prevents thermal damage and acoustic cavitation risk"],
            ["Frequency",           "1\u20138 Hz",                        "8 Hz available only with up to 0.02 mJ/mm\u00b2 max power"],
            ["Total Pulses / Session", "6,000\u201312,000 (phenotype dependent)", "Allocated between standard, targeted, and peripheral"],
            ["Cranial Allocation",  "Up to 10,000 pulses",               "Standard distributed + targeted deep components"],
            ["Peripheral (PRE) Allocation", "2,000 within total",        "Cranio-cervical, mastoid, suboccipital, plantar zones"],
            ["Session Duration",    "30\u201345 minutes (including setup)", "Neuronavigation setup adds time"],
        ],
        hdr_fill=C_NAVY, widths=[2.0, 2.2, 2.3]
    )
    make_table(doc,
        ["Parameter", "Rule"],
        [
            ["Sessions per Day",   "Up to 2 sessions daily (with Doctor approval)"],
            ["Minimum Interval",   "\u22656 hours between sessions"],
            ["Weekly Frequency",   "Typically 3\u20135 days/week"],
            ["Maintenance Therapy","As determined by treating Doctor"],
            ["Escalation",         "Gradual dose escalation recommended across first 2\u20133 sessions"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_PEACH, None]*3, widths=[2.0, 4.5]
    )
    add_h(doc, "9C. TPS ROI Brain Region Mapping", 2)
    make_table(doc,
        ["ROI", "Brain Region", "Standard Anatomical Label", "Clinical Target"],
        [
            ["M1 (bilateral)",  "Primary motor cortex",         "Precentral gyrus",                        "Motor output, motor speed, motor scaling"],
            ["SMA",             "Supplementary motor area",     "Medial frontal (superior frontal gyrus)", "Motor planning, sequencing, initiation"],
            ["DLPFC",           "Dorsolateral prefrontal cortex","Middle frontal gyrus",                   "Executive function, cognition, mood, dual-task"],
            ["PPC",             "Posterior parietal cortex",    "Inferior/superior parietal lobule",       "Sensorimotor integration, visuospatial processing"],
            ["ACC",             "Anterior cingulate cortex",    "Cingulate gyrus (anterior)",              "Mood, motivation, pain modulation"],
            ["Cerebellum",      "Cerebellar hemispheres + vermis","Cerebellar cortex",                     "Tremor, balance, coordination, motor learning"],
            ["IFG/Broca",       "Inferior frontal gyrus",       "Pars opercularis + triangularis",         "Language, speech, executive inhibition"],
            ["OFC",             "Orbitofrontal cortex",         "Inferior frontal surface",                "Reward, motivation, depression/anhedonia"],
            ["Holocranial",     "Distributed cortical coverage","Frontal\u2013parietal\u2013temporal\u2013occipital", "Global modulation; complements targeted ROIs"],
            ["Plantar (peripheral)", "Foot sole afferent input","Plantar surface (mechanoreceptors)",      "Gait augmentation via sensorimotor afferent drive"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_BLUE, None]*5, widths=[1.0, 1.7, 2.0, 1.8]
    )
    add_h(doc, "9D. Classic TPS Protocols \u2014 Evidence-Based (T1\u2013T5)", 2)
    make_table(doc,
        ["ID", "Symptom Target", "Brain Targets", "Parameters", "Pulse Allocation", "Evidence"],
        cond["tps_protocols"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None, F_GREEN, None, F_PURPLE],
        widths=[0.4, 1.2, 1.5, 1.5, 1.5, 0.4]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 10 — taVNS & CES
# ─────────────────────────────────────────────────────────────
def s10_tavns_ces(doc, cond):
    add_h(doc, "10. taVNS & CES \u2014 Device Information & Protocols")
    add_h(doc, "10A. taVNS \u2014 Transcutaneous Auricular Vagus Nerve Stimulation", 2)
    add_h(doc, "Overview and Mechanism", 3)
    add_p(doc,
        "Transcutaneous auricular vagus nerve stimulation (taVNS) delivers low-intensity electrical "
        "stimulation to the auricular branch of the vagus nerve (cymba conchae) without surgical implantation. "
        "Afferent vagal signals ascend via the nucleus tractus solitarius (NTS) to the locus coeruleus (LC), "
        "dorsal raphe nucleus (DRN), and cortical/subcortical regions, modulating noradrenergic, serotonergic, "
        "and GABAergic tone."
    )
    add_h(doc, "Clinical Role", 3)
    add_p(doc, cond.get("tavns_role",
        "taVNS serves as an autonomic and limbic modulator in this protocol. It is used adjunctively "
        "to prime or stabilise network states before primary cortical stimulation (tDCS/TPS), "
        "to modulate arousal and autonomic tone, and to support mood, sleep, and pain regulation."
    ))
    add_h(doc, "Electrode Placement", 3)
    add_p(doc,
        "Cymba conchae of the left ear (primary target \u2014 highest density of auricular vagal afferents). "
        "A return electrode is placed at the ear lobe or mastoid. "
        "Right ear may be used as an alternative if left ear is inaccessible. "
        "Bilateral placement is used in some protocols for enhanced effect."
    )
    add_h(doc, "taVNS Safety Limits", 3)
    make_table(doc,
        ["Parameter", "Limit"],
        [
            ["Maximum Daily Duration", "4 hours total per day (across all sessions combined)"],
            ["Frequency Range",        "20\u201330 Hz (SOZO standard range)"],
            ["Amplitude",             "0\u20135 mA (typical 0.5\u20132 mA for comfort)"],
            ["Pulse Width",           "0.2\u20130.5 ms (standard)"],
            ["Monitoring",            "Stop if dizziness, syncope, persistent nausea"],
            ["Contraindication",      "Active cardiac instability; implanted active vagal stimulators"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_BLUE, None]*3, widths=[2.0, 4.5]
    )
    add_h(doc, "taVNS Integration Principles", 3)
    make_table(doc,
        ["Mode", "Timing", "Rationale"],
        [
            ["Priming",     "20\u201330 min before tDCS",               "Increases cortical receptivity via brainstem-limbic activation"],
            ["Concurrent",  "Simultaneously with tDCS",                 "Combined brainstem-cortical modulation; dual-pathway engagement"],
            ["Adjunctive",  "On non-tDCS days; or after TPS",           "Maintains vagal-limbic tone; supports autonomic stability on treatment-free days"],
        ],
        hdr_fill=C_NAVY, widths=[1.2, 2.0, 3.3]
    )
    add_h(doc, "10B. CES \u2014 Alpha-Stim\u00ae Cranial Electrotherapy Stimulation", 2)
    add_h(doc, "Overview and Mechanism", 3)
    add_p(doc,
        "Cranial electrotherapy stimulation (CES) delivers microcurrent (50\u2013500 \u00b5A) via earlobe "
        "electrodes using a proprietary biphasic waveform. The Alpha-Stim AID device is FDA-cleared for "
        "adjunctive support of mood, anxiety, and sleep. The mechanism likely involves modulation of "
        "thalamo-cortical oscillations and limbic circuits via low-level transcranial current."
    )
    add_h(doc, "Device Specifications", 3)
    make_table(doc,
        ["Specification", "Detail"],
        [
            ["Device Name",       "Alpha-Stim\u00ae CS (Cranial Electrotherapy Stimulation)"],
            ["Waveform",          "Proprietary biphasic asymmetrical current"],
            ["Current Range",     "50\u2013500 \u00b5A (microvolt level)"],
            ["SOZO Therapeutic Range", "100\u2013400 \u00b5A (established by clinical use)"],
            ["Electrode Placement","Earlobes (clip electrodes) or mastoid region; bilateral"],
            ["Session Duration",  "20\u201360 minutes (SOZO typical: 20\u201340 min)"],
            ["Frequency",         "Device-specific microcurrent; not user-adjustable"],
            ["Regulatory Status", "FDA 510(k) cleared (K061053) for mood, sleep, anxiety support"],
            ["Clinical Role",     f"Adjunctive/supportive for mood, sleep, anxiety, fatigue in {cond['name']}; NOT primary therapy"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_GREEN, None]*5, widths=[2.0, 4.5]
    )
    add_h(doc, "CES Safety Limits", 3)
    make_table(doc,
        ["Parameter", "Limit"],
        [
            ["Maximum Current",      "500 \u00b5A (device limit)"],
            ["SOZO Typical Range",   "100\u2013300 \u00b5A"],
            ["Session Duration",     "20\u201340 min typical"],
            ["Daily Use",            "1\u20132 sessions/day"],
            ["Contraindications",    "Pacemakers (absolute); epilepsy (caution); open scalp wounds (avoid site)"],
            ["Side Effects",         "Minimal; transient tingling, mild discomfort, rarely dizziness"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_GREEN, None]*3, widths=[2.0, 4.5]
    )
    add_h(doc, f"CES Clinical Role in {cond['name']}", 3)
    make_table(doc,
        ["Indication", "Role", "Integration"],
        cond["ces_role"],
        hdr_fill=C_NAVY, row_fills=[F_YELLOW, None]*3, widths=[1.5, 2.5, 2.5]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 11 — Combinations
# ─────────────────────────────────────────────────────────────
def s11_combinations(doc, cond):
    add_h(doc, "11. Combination of NIBS Techniques")
    add_h(doc, "Multimodal Combinations by Phenotype", 2)
    add_p(doc,
        "The SOZO session framework follows the S-O-Z-O structure: Stabilise (taVNS/CES), "
        "Optimise (tDCS), Zone (TPS), Outcome (taVNS/CES + rehabilitation). "
        "Combinations should be selected based on the dominant phenotype identified in Section 6."
    )
    make_table(doc,
        ["Letter", "Phase", "Modality", "Purpose"],
        [
            ["S", "Stabilise",      "taVNS or CES",           "Downshift sympathetic tone, reduce threat/arousal, improve sensory gating"],
            ["O", "Optimise",       "tDCS",                   "Prime target brain excitability for plasticity and task engagement"],
            ["Z", "Zone targeting", "TPS",                    "Deliver focal/region-specific stimulation + global cortical engagement"],
            ["O", "Outcome",        "taVNS/CES + Rehab",      "Post-session autonomic support + behavioural practice window"],
        ],
        hdr_fill=C_NAVY, widths=[0.4, 1.2, 1.6, 3.3]
    )

    for pheno_name, combos in cond["combinations"]:
        add_h(doc, pheno_name, 3)
        make_table(doc,
            ["Combination", "Rationale", "Timing", "Indication"],
            combos,
            hdr_fill=C_HDR_DARK,
            row_fills=[F_BLUE2, F_GREEN2, F_CREAM],
            widths=[1.5, 2.5, 1.3, 1.2]
        )

    add_h(doc, "Multimodal Combination Evidence Summary", 2)
    make_table(doc,
        ["Phenotype", "Combination", "Evidence-Based Rationale", "Suggested Timing", "Best Clinical Use", "Evidence Level"],
        cond["combination_summary"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None, F_GREEN, None, F_PURPLE, None,
                   F_CREAM, None, F_PEACH, None, F_GREY, None],
        widths=[1.0, 1.2, 2.0, 0.9, 1.0, 0.4]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 12 — Side Effects
# ─────────────────────────────────────────────────────────────
def s12_side_effects(doc, cond):
    add_h(doc, "12. Side Effects and Monitoring")
    add_h(doc, "Common Side Effects", 2)
    add_h(doc, "tDCS Side Effects", 3)
    make_table(doc,
        ["Side Effect", "Frequency", "Severity", "Management"],
        [
            ["Tingling / itching at electrode site", "50\u201370%",    "Mild",          "Normal; resolves after session"],
            ["Mild headache",                         "10\u201320%",    "Mild",          "Paracetamol; resolves within hours"],
            ["Skin redness at electrode site",         "30\u201340%",    "Mild",          "Transient; ensure electrode hydration"],
            ["Burning sensation",                      "Rare <5%",      "Moderate",      "STOP. Check impedance; reduce intensity"],
            ["Skin lesion / burn",                     "Very rare <1%", "Severe",         "STOP immediately. Report to Doctor"],
            ["Fatigue or drowsiness",                  "10\u201320%",    "Mild",          "Typically resolves within hours; not reason to stop"],
            ["Mood changes (rare dysphoria)",           "1\u20133%",     "Mild\u2013Moderate", "Monitor closely; consider polarity switch or protocol change"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_YELLOW, None]*4, widths=[2.0, 1.0, 1.0, 2.5]
    )
    add_h(doc, "TPS Side Effects", 3)
    make_table(doc,
        ["Side Effect", "Frequency", "Severity", "Management"],
        [
            ["Mild headache",                    "10\u201320%", "Mild\u2013Moderate", "Analgesics as needed; check targeting accuracy"],
            ["Scalp discomfort during pulse",     "20\u201330%", "Mild",              "Reduce energy density; ensure adequate transducer contact"],
            ["Acoustic sensation (clicking)",     "50\u201370%", "Mild",              "Expected; use earplugs if bothersome"],
            ["Transient dizziness",               "5\u201310%",  "Mild",              "Usually self-limited; rest post-session"],
            ["Skin erythema at transducer site",  "1\u20135%",   "Mild",              "Apply moisturiser; avoid excessive sessions"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_PEACH, None]*3, widths=[2.0, 1.0, 1.0, 2.5]
    )
    add_h(doc, "Adverse Event Grading", 2)
    make_table(doc,
        ["Grade", "Examples", "Action"],
        [
            ["Grade 1 (Mild)",     "Tingling, mild headache, transient redness",              "Continue. Document."],
            ["Grade 2 (Moderate)", "Persistent headache >4h, burning, mood change >24h",      "PAUSE. Doctor review within 24h. Resume with approval."],
            ["Grade 3 (Severe)",   "Seizure, syncope, burn, neurological deficit, chest pain", "STOP. Emergency Response. No resumption without Doctor clearance."],
        ],
        hdr_fill=C_RED, widths=[1.2, 3.0, 2.3]
    )
    add_h(doc, "Contraindications Summary", 2)
    make_table(doc,
        ["Category", "tDCS", "TPS (Investigational)"],
        [
            ["Absolute",           "Metallic implants; DBS; cochlear; open wounds", "Metallic implants; DBS; skull defects; intracranial tumours; pregnancy"],
            ["Requires Discussion", "Pacemaker; epilepsy; pregnancy; dermatological; unstable medical",
             "Pacemaker; epilepsy; coagulation disorders; anticoagulants; severe dyskinesias"],
        ],
        hdr_fill=C_NAVY, widths=[1.2, 2.8, 2.5]
    )
    add_h(doc, "Home-Use Safety Protocol", 2)
    make_table(doc,
        ["Requirement", "Detail"],
        [
            ["Device Settings",       "For HDCkit: all settings pre-programmed by Doctor via HDCprog; patient cannot modify. "
                                      "For PlatoScience: app-guided; follow prescriptions exactly."],
            ["Compliance Monitoring", "Doctor reviews compliance data at regular intervals"],
            ["Stop Criteria",         "Burning, skin lesion, severe headache, dizziness, seizure \u2014 patient must contact clinic immediately"],
            ["Applicable Devices",    "Home use applies to tDCS (HDCkit, PlatoScience), CES (Alpha-Stim), taVNS. TPS is clinic-only."],
            ["Device Integrity",      "Check device before each use; electrodes clean and intact; conductive medium applied"],
            ["Session Logging",       "Maintain session log: date, time, current, duration, any side effects"],
            ["Emergency Contact",     "MD or clinic contact must be available during home sessions"],
            ["Parameter Modification","No modification of parameters without explicit MD approval"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_YELLOW, None]*4, widths=[2.0, 4.5]
    )

# ─────────────────────────────────────────────────────────────
# SECTION 13 — Follow-Up
# ─────────────────────────────────────────────────────────────
def s13_followup(doc, cond):
    add_h(doc, "13. Follow-Up Assessments and Decision-Making")
    add_h(doc, "Outcome Measures", 2)
    make_table(doc,
        ["Assessment", "Domain", "Frequency", "Threshold/Interpretation"],
        cond["outcomes"],
        hdr_fill=C_NAVY,
        row_fills=[F_BLUE, None, F_GREEN, None, F_PURPLE, None, F_CREAM, None, F_PEACH, None],
        widths=[1.6, 1.8, 1.6, 1.5]
    )
    add_h(doc, "Decision Timeline", 2)
    make_table(doc,
        ["Time Point", "Assessment", "Decision Gate", "Action"],
        [
            ["Week 2",  "Symptom questionnaire; tolerance check",            "Safety gate",           "Continue if well-tolerated; modify if side effects Grade 2"],
            ["Week 4",  f"{', '.join(cond['outcomes_abbrev'][:3])}",         "Early efficacy signal", "Assess trend; continue if \u226510% improvement; adjust if plateau"],
            ["Week 8",  "Full motor and non-motor battery",                  "Efficacy assessment",   "Classify response; decide continuation vs. protocol change"],
            ["Week 12", f"Full battery + {cond['outcomes_abbrev'][-1] if cond['outcomes_abbrev'] else 'SOZO PRS'}", "Primary efficacy gate", "Classification; plan continuation block vs. washout/escalation"],
        ],
        hdr_fill=C_NAVY, widths=[0.8, 2.2, 1.5, 2.0]
    )
    add_h(doc, "Response Classification (Week 8\u201312)", 2)
    make_table(doc,
        ["Category", "Threshold", "Definition", "Action"],
        [
            ["Responder",         "\u226530% improvement in primary outcome", "Clinically significant benefit",   "Continue; maintenance block 2\u20134/week; plan long-term follow-up"],
            ["Partial Responder", "10\u201329% improvement",                  "Mild-to-moderate benefit",         "Optimise protocol; increase frequency (if safe); consider protocol switch"],
            ["Non-Responder",     "<10% improvement (or worsening)",          "Minimal or no benefit",            "See Non-Responder Pathway below"],
        ],
        hdr_fill=C_NAVY, widths=[1.2, 1.5, 1.8, 2.0]
    )
    add_h(doc, "Non-Responder Escalation Pathway", 2)
    make_table(doc,
        ["Step", "Action", "Detail"],
        [
            ["1", "Compliance Review",      "Verify session adherence, electrode placement, medication timing"],
            ["2", "Phenotype Reassessment", "Re-evaluate dominant phenotype; may have shifted"],
            ["3", "Protocol Adjustment",    "Switch protocol (e.g., C1 \u2192 C2); change montage; add modality (TPS, taVNS)"],
            ["4", "Parameter Modification", "Adjust intensity, duration, frequency within safety limits"],
            ["5", "Modality Switch",        "If tDCS non-response: consider TPS as primary; or add multimodal combination"],
            ["6", "Discontinuation",        "If no response after systematic adjustment: discontinue with documented rationale"],
        ],
        hdr_fill=C_NAVY, row_fills=[F_PEACH, None]*3, widths=[0.4, 1.8, 4.3]
    )
    add_h(doc, "Responders and Non-Responders: Key Considerations", 2)
    add_p(doc,
        "Response classification must account for: (1) natural disease progression which may mask "
        "treatment benefit; (2) medication interactions (e.g., timing relative to peak medication effect); "
        "(3) phenotype shift during the treatment block; (4) comorbid conditions that independently "
        "affect outcome measures."
    )
    add_p(doc,
        "For conditions with episodic or relapsing patterns, short-term non-response does not exclude "
        "future benefit. Maintenance therapy at reduced frequency (2\u20133 sessions/week) may be "
        "considered after a successful initial block."
    )

# ─────────────────────────────────────────────────────────────
# MASTER BUILD FUNCTION
# ─────────────────────────────────────────────────────────────
def build_document(cond, out_path):
    doc = Document()
    setup_page(doc)
    # Ensure Word heading styles exist
    for lvl in [1, 2, 3]:
        sname = f"Heading {lvl}"
        if sname not in [s.name for s in doc.styles]:
            doc.styles.add_style(sname, 1)

    add_cover(doc, cond)
    doc.add_page_break()
    s1_doc_control(doc, cond)
    s2_inclusion(doc, cond)
    s3_overview(doc, cond)
    s4_pathophysiology(doc, cond)
    s5_brain(doc, cond)
    s6_phenotypes(doc, cond)
    s7_mapping(doc, cond)
    s8_tdcs(doc, cond)
    s9_tps(doc, cond)
    s10_tavns_ces(doc, cond)
    s11_combinations(doc, cond)
    s12_side_effects(doc, cond)
    s13_followup(doc, cond)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  [OK] {out_path}")


# ─────────────────────────────────────────────────────────────
# CONDITION DATA  (all 15 conditions)
# ─────────────────────────────────────────────────────────────
CONDITIONS = {}

# ══════════════════════════════════════
# 1. PARKINSON'S DISEASE
# ══════════════════════════════════════
CONDITIONS["parkinsons"] = dict(
    name="Parkinson's Disease", icd10="G20",
    doc_num="SOZO-PD-FEL-001",
    tps_status="INVESTIGATIONAL / OFF-LABEL",
    tdcs_class="Established NIBS; Level B evidence",
    tps_class="INVESTIGATIONAL / OFF-LABEL",
    tavns_class="Emerging adjunctive",
    ces_class="Supportive adjunct ONLY",
    inclusion=[
        "Confirmed PD diagnosis (UK Brain Bank or MDS criteria)",
        "Hoehn and Yahr stage I\u2013IV (V with Doctor approval)",
        "Age 18+ years",
        "Able to attend sessions or comply with home-use protocols",
        "Written informed consent (including TPS off-label disclosure)",
        "Baseline outcome measures completed (MDS-UPDRS III, MoCA, PHQ-9, GAD-7, FOG-Q)",
    ],
    exclusion=[
        "Intracranial metallic hardware in stimulation path",
        "Cochlear implant",
        "Skull defects, craniectomy, or recent craniotomy",
        "Active intracranial tumour",
        "Open wounds at electrode sites",
        "Pregnancy (tDCS, TPS)",
        "Inability to provide informed consent",
    ],
    discussion=[
        "Cardiac pacemaker or defibrillator \u2014 individual risk\u2013benefit assessment",
        "Epilepsy or seizure history \u2014 formal risk\u2013benefit with documented rationale",
        "Severe cognitive impairment",
        "Unstable medical conditions",
        "Active psychosis / impulse control disorders",
        "Coagulation disorders or anticoagulants (especially TPS)",
        "Dermatological conditions at electrode sites",
        "Severe dyskinesias preventing safe device application",
    ],
    overview=[
        "Parkinson's disease (PD) is the second most common neurodegenerative disorder and affects around 1\u20132% "
        "of people over the age of 60. It is characterised by progressive degeneration of dopaminergic neurons "
        "in the substantia nigra pars compacta, with resulting dysfunction across motor and non-motor neural "
        "networks (Bloem et al., 2021; Kalia & Lang, 2015).",
        "Cardinal Motor Features: The core motor features of PD are bradykinesia, rigidity, resting tremor, and "
        "postural instability. Bradykinesia is the defining clinical feature and is required for diagnosis, while "
        "tremor and rigidity are common accompanying signs. Postural instability usually emerges later and is "
        "associated with increased fall risk and functional decline.",
        "Non-Motor Symptoms: PD also includes a wide range of non-motor symptoms, many of which may appear before "
        "the onset of clear motor signs. These include cognitive impairment, depression, apathy, fatigue, pain, "
        "sleep disturbance, and autonomic dysfunction (Chaudhuri & Schapira, 2009).",
        "NIBS Evidence in PD: tDCS has shown promising results in motor performance, gait, cognition, and fatigue "
        "when combined with rehabilitation. TPS is regarded as investigational/off-label in PD. NIBS does not "
        "replace established pharmacological therapies; its role is adjunctive and phenotype-guided (Moshayedi et al., 2025).",
    ],
    pathophysiology=[
        "PD pathophysiology involves progressive loss of dopaminergic neurons in the substantia nigra pars compacta, "
        "leading to dopamine depletion in the striatum and disruption of the basal ganglia-thalamocortical circuit, "
        "resulting in reduced thalamo-cortical drive and impaired voluntary movement initiation.",
        "Lewy body pathology (alpha-synuclein aggregation) spreads through the nervous system (Braak stages 1\u20136). "
        "Non-motor symptoms reflect pathology beyond the nigrostriatal system: cortical Lewy bodies contribute "
        "to cognitive decline; serotonergic and noradrenergic depletion underpins depression and anxiety.",
    ],
    std_treatment=[
        "Pharmacological treatment in PD is primarily symptomatic. Levodopa remains the most effective treatment "
        "for motor symptoms and is prescribed with carbidopa or benserazide. Additional options include dopamine "
        "agonists, MAO-B inhibitors, COMT inhibitors, and amantadine. In advanced disease, DBS, LCIG, and "
        "subcutaneous infusions may be considered (Armstrong & Okun, 2020).",
    ],
    symptoms=[
        ["Bradykinesia",        "Slowness of movement initiation and execution; progressive fatiguing",  "Core criterion", "Kalia & Lang 2015; Bloem 2021"],
        ["Rigidity",            "Increased resistance to passive movement; lead-pipe or cogwheel",        "Core criterion", "Kalia & Lang 2015"],
        ["Resting Tremor",      "4\u20136 Hz tremor at rest; typically asymmetric; pill-rolling",         "70\u201380%",    "Braak 2003; Bloem 2021"],
        ["Postural Instability","Impaired balance; increased fall risk; retropulsion",                    "Later stages",   "Armstrong & Okun 2020"],
        ["Gait / FOG",          "Festination, shuffling, reduced arm swing, freezing of gait",            "Up to 80%",      "Wong 2022"],
        ["ON Dyskinesia",       "Levodopa-induced involuntary movements; choreiform at peak dose",         "40\u201390% at 10y", "Armstrong & Okun 2020"],
        ["Cognitive Impairment","Executive dysfunction, attention deficits, visuospatial impairment",     "20\u201340% MCI","Ma 2025; Souto 2024"],
        ["Depression",          "Anhedonia, apathy, low mood; may precede motor symptoms",                "40\u201350%",    "Chaudhuri & Schapira 2009"],
        ["Fatigue",             "Central fatigue; exhaustion disproportionate to activity",               "Up to 70%",      "Forogh 2017"],
        ["Pain",                "Musculoskeletal, neuropathic, central, dystonic pain",                   "60\u201385%",    "Gonzalez-Zamorano 2024"],
        ["Sleep Disturbances",  "REM behavior disorder, insomnia, excessive daytime sleepiness",          "60\u201398%",    "Chaudhuri & Schapira 2009"],
    ],
    brain_regions=[
        ["Primary Motor Cortex (M1)",     "Reduced excitability due to excessive BG inhibition",            "Anodal tDCS restores cortical excitability; Level B evidence",             "Lefaucheur 2017"],
        ["Supplementary Motor Area (SMA)","Impaired motor planning, sequencing, and gait initiation",       "Stimulation improves gait initiation and reduces FOG",                     "Wong 2022"],
        ["DLPFC",                         "Executive dysfunction, depression, impaired dual-task",          "Anodal tDCS demonstrates cognitive and mood improvements",                 "Ma 2025; Souto 2024"],
        ["Cerebellum",                    "Compensatory hyperactivation; tremor circuitry",                 "Modulation may reduce dyskinesias and improve balance",                    "Ferrucci 2015"],
        ["Thalamus (VL, VIM)",            "Relay in BG\u2013thalamo\u2013cortical loops; tremor generation","TPS may access thalamic structures at depth (investigational)",            "Manganotti 2025"],
        ["Basal Ganglia (STN, GPi)",      "Pathological beta-band oscillations; disrupted circuits",        "Cortical stimulation modulates subcortical activity via connected pathways","Bange 2025"],
        ["Anterior Cingulate Cortex",     "Motivation, conflict monitoring, pain processing",               "Potential target for apathy and pain",                                     "Zhang 2023"],
        ["Nucleus Basalis of Meynert",    "Cholinergic degeneration; cognitive decline",                    "TPS deep target for cognitive symptoms (investigational)",                 "Gianlorenco 2025"],
    ],
    brainstem=[
        ["Pedunculopontine Nucleus (PPN)", "Gait initiation failure; postural instability; FOG",  "Indirect modulation via SMA/M1/DLPFC; deep TPS investigational","Thevathasan 2018"],
        ["Locus Coeruleus (LC)",           "Early noradrenergic degeneration; arousal dysfunction","taVNS modulates LC via NTS pathways",                           "Marano 2024"],
        ["Dorsal Raphe Nucleus",           "Serotonergic dysregulation; depression; fatigue",      "tDCS (DLPFC) and taVNS may influence serotonergic projections", "Chmiel 2024"],
        ["Nucleus Tractus Solitarius",     "Autonomic integration hub; vagal afferent relay",      "Primary afferent target of taVNS; influences prefrontal regions","Yakunina 2018"],
        ["Periaqueductal Gray (PAG)",      "Central pain modulation; autonomic control",           "Indirect modulation via cortical and vagal neuromodulation",     "Bange 2025"],
    ],
    phenotypes=[
        ["Tremor-Dominant (TD)",          "Resting tremor >50% of motor burden; slower progression","Tremor, rigidity, mild bradykinesia",          "Cerebellar tDCS; TPS to cerebellum/thalamus"],
        ["Akinetic-Rigid (AR)",           "Rigidity and bradykinesia dominant; faster progression", "Bradykinesia, rigidity, postural instability", "M1/SMA anodal tDCS; TPS to M1/SMA"],
        ["PIGD (Postural Instability/Gait)","Early balance loss and gait dysfunction",             "Gait freezing, balance loss, faster progression","SMA/dual-task tDCS; balance protocols"],
        ["Depression / Apathy Dominant",  "Primary mood and motivational disturbance",              "Anhedonia, apathy, low mood; psychomotor slowing","DLPFC anodal tDCS; taVNS; CES adjunct"],
        ["Cognitive / Executive Dominant","Primary cognitive impairment",                          "Executive dysfunction, attention deficits",    "DLPFC anodal tDCS; TPS to DLPFC"],
        ["Pain Dominant",                 "Central or musculoskeletal pain as primary burden",      "Neuropathic pain, dystonic pain, MSK pain",   "M1 contralateral tDCS; TPS M1/thalamus; taVNS"],
        ["Mixed / Variable",              "Balanced motor/non-motor; atypical presentation",        "Combination of tremor, rigidity, cognitive, mood","Multimodal NIBS combination per phenotype"],
    ],
    symptom_map=[
        ["Motor (Bradykinesia/Rigidity)", "M1, SMA, Premotor",          "tDCS (M1/SMA anodal) + TPS (cranial targeted + peripheral) + taVNS adjunct","Moderate (Lefaucheur 2017)"],
        ["Tremor",                        "Cerebellum, Thalamus, M1",   "TPS (cranial targeted + peripheral) + cerebellar/M1 tDCS",                  "Emerging (Ferrucci 2015)"],
        ["Gait / FOG",                    "SMA, DLPFC, Cerebellum",     "Bilateral DLPFC tDCS + TPS (cranial + peripheral soles); dual-task training","Moderate (Wong 2022)"],
        ["Cognition / Executive",         "DLPFC, Frontal-Parietal",    "Left DLPFC tDCS + TPS (cranial targeted); taVNS, CES supportive",          "Moderate\u2013strong (Ma 2025)"],
        ["Depression / Apathy",           "DLPFC, ACC",                 "Left DLPFC tDCS + TPS (cranial targeted) + taVNS + CES adjunct",            "Moderate (Chmiel 2024)"],
        ["Fatigue",                       "DLPFC, SMA, Mesolimbic",     "DLPFC tDCS + TPS (cranial targeted) + CES; taVNS for arousal",             "Limited\u2013moderate (Forogh 2017)"],
        ["Pain",                          "M1 (contralateral), Thalamus","Contralateral M1 tDCS + TPS (cranial + peripheral); taVNS, CES",           "Moderate (Gonzalez-Zamorano 2024)"],
        ["ON/OFF Dyskinesias",            "Cerebellum, SMA, M1",        "Cerebellar/SMA tDCS + TPS (cranial targeted + peripheral)",                 "Emerging (Ferrucci 2015)"],
        ["Anxiety / Autonomic",           "ACC, Brainstem (LC, NTS)",   "taVNS primary + CES; DLPFC tDCS + TPS (cranial targeted)",                 "Moderate (Marano 2024)"],
    ],
    montage=[
        ["Motor dominant (bradykinesia, rigidity)","C3/C4 contralateral M1 anodal; or Fpz+Cbz anode, Cz cathode","Cranial (global + targeted M1/SMA) + peripheral (palms and soles)","taVNS"],
        ["Tremor primary",        "C4/C3 contralateral anode + Cbz anode, Fpz cathode",     "Cranial (global + targeted cerebellum/thalamus) + peripheral (palms)","taVNS supportive"],
        ["Gait / FOG",            "Fpz + Cz anode, CervVII cathode; Cb1/Cb2 anodal",        "Cranial (global + targeted) + peripheral (soles)",                    "taVNS supportive"],
        ["Cognitive / Executive", "F3 anode + F4 anode, Pz cathode; T3/T4 options",         "Cranial (global + targeted DLPFC)",                                   "taVNS, CES supportive"],
        ["Depression / Apathy",   "F3 anode + Pz anode, F2 cathode",                        "Cranial (global + targeted DLPFC/ACC)",                               "taVNS, CES supportive"],
        ["Fatigue",               "F3 anode + P4 anode, F2 cathode; Fz + Cz options",       "Cranial (global + targeted)",                                         "taVNS, CES supportive"],
        ["Pain (central or MSK)", "C3/C4 contralateral + F3 anode, Pz cathode",             "Cranial (global + targeted) + peripheral (affected dermatomes)",      "taVNS, CES supportive"],
        ["ON/OFF Dyskinesias",    "Fpz + Cbz anode, Cz cathode; Cb1/Cb2 options",           "Cranial (global + targeted) + peripheral (affected regions)",         "taVNS, CES supportive"],
        ["Mixed motor + cognitive","Alternate blocks: M1-based (motor) and DLPFC-based (executive)","Cranial (global + targeted)",                               "taVNS, CES supportive"],
        ["No response after 8\u201310 sessions","Reassess phenotype; adjust montage or switch modality","Adjust parameters; consider different ROI targeting",     "Review all adjuncts"],
    ],
    tdcs_protocols=[
        ["C1","Motor (bradykinesia, rigidity)","C3/C4 (contra M1)","Fz/Pz/extracephalic",      "2 mA, 20\u201340 min, 1\u20132\u00d7/day","Level B; improves UPDRS-III (Lefaucheur 2017)"],
        ["C2","Gait & dual-task",              "F3 (L-DLPFC), Cz", "Cbz, extracephalic",        "2 mA, 20\u201340 min, 1\u20132\u00d7/day","Improves gait speed (Wong 2022, 2024)"],
        ["C3","Cognition / executive",         "F3, F4, T3/T4, P3/P4","F2, extracephalic",      "2 mA, 20\u201340 min, 1\u20132\u00d7/day","Significant cognitive gains (Ma 2025)"],
        ["C4","Depression / mood",             "F3, Fz",           "F2, extracephalic",          "2 mA, 20\u201340 min, 1\u20132\u00d7/day","Reduces depression/apathy (Chmiel 2024)"],
        ["C5","Fatigue",                       "F3, P4, Pz, Cb1/Cb2","Fpz, F2, extracephalic",  "2 mA, 20\u201340 min, 1\u20132\u00d7/day","Sustained fatigue reduction (Forogh 2017)"],
        ["C6","Balance / posture",             "Fpz, Cz, Cb1/Cb2, P3/P4","Extracephalic, F2",   "2 mA, 20\u201340 min, 1\u20132\u00d7/day","Improves BBS & TUG (Na 2022)"],
        ["C7","ON/OFF dyskinesia",             "Cb1/Cb2/Cbz",      "Fpz, Cz",                   "2 mA, 20\u201340 min, 1\u20132\u00d7/day","Decreases dyskinesias (Ferrucci 2015)"],
        ["C8","Motor learning",                "C3/C4, Pz, Fz, F3/Cbz","F4, extracephalic",     "2 mA during task, 20\u201340 min",         "Facilitates learning (Firouzi 2024)"],
    ],
    plato_protocols=[
        ["C1-PS","Motor (bradykinesia)","Motor (Cz area)",   "Cz","Shoulder/arm",   "1.6 mA, 20\u201330 min","Daily or 5\u00d7/wk","Single-site; use with physiotherapy"],
        ["C2-PS","Gait / FOG",          "Motor (Cz area)",   "Cz","Nape of neck",   "1.6 mA, 20\u201330 min","Daily or 5\u00d7/wk","Cz placement closest to SMA/M1"],
        ["C3-PS","Cognition",           "Frontal (F3 area)", "F3","Shoulder",        "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Left prefrontal anodal for executive"],
        ["C4-PS","Depression / mood",   "Frontal (F3 area)", "F3","Right shoulder",  "1.6 mA, 20\u201330 min","5\u00d7/wk",         "F3 anodal; adjunct to CES"],
        ["C5-PS","Fatigue",             "Frontal (F3 area)", "F3","Shoulder",        "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Morning use preferred"],
        ["C6-PS","Balance / posture",   "Motor (Cz area)",   "Cz","Nape of neck",   "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Combine with vestibular exercises"],
        ["C7-PS","Dyskinesia",          "Posterior (Pz area)","Pz","Shoulder",       "1.6 mA, 20\u201330 min","3\u20135\u00d7/wk",  "Cerebellar targeting approximation"],
        ["C8-PS","Motor learning",      "Motor (Cz area)",   "Cz","Shoulder/arm",   "1.6 mA during task",    "5\u00d7/wk",         "Apply during motor practice window"],
    ],
    tps_protocols=[
        ["T1","Motor (bradykinesia, rigidity)","Bilateral M1, SMA, Premotor",      "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz; 10 sessions/2 wks","4,000 Std \u00b1 4,000 Tgt; total 6\u201310K","Emerging (Osou 2024)"],
        ["T2","Tremor",                        "VIM Thalamus, Cerebellar Dentate \u00b1 M1","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",       "4,000 Tgt \u00b1 2,000 PRE; total 6\u20138K", "Emerging (Manganotti 2025)"],
        ["T3","Cognition (PD-MCI)",            "DLPFC, mPFC, Hippocampal region",  "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",               "4,000 Tgt \u00b1 2,000 PRE; total 6\u20138K", "Preliminary (Gianlorenco 2025)"],
        ["T4","Depression / Apathy",           "DLPFC, ACC",                       "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",               "4,000 Tgt \u00b1 2,000 PRE; total 6\u20138K", "Preliminary secondary outcomes"],
        ["T5","Pain / Dyskinesia",             "M1 + Thalamus",                    "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",               "4,000 Std + 4,000 Tgt; total 8\u201312K",     "Review-level support (Bange 2025)"],
    ],
    ces_role=[
        ["Mood and Depression",   "Supportive for mild-to-moderate depressive symptoms alongside primary tDCS","Evening sessions; or non-tDCS days"],
        ["Sleep Disturbances",    "Enhances sleep quality; may reduce REM behavior disorder symptoms",          "Evening before bed; 20\u201340 min"],
        ["Anxiety and Autonomic", "Stabilises arousal and autonomic tone",                                      "As needed; pre-session or evening"],
        ["Fatigue",               "May improve perceived energy and motivation as adjunct",                     "Morning or midday; avoid late if sedating"],
    ],
    tavns_role="taVNS serves as an autonomic and limbic modulator in PD, priming cortical receptivity before tDCS/TPS and supporting mood, sleep, and autonomic stability via locus coeruleus modulation.",
    combinations=[
        ("1) Bradykinetic / Rigid", [
            ["tDCS + Motor Physiotherapy","Anodal M1/DLPFC tDCS + structured rehabilitation (Elsner 2016; Moshayedi 2025)","tDCS before/during rehab","Bradykinesia, rigidity, slowing"],
            ["taVNS + Standard Care","taVNS modulates autonomic tone; adjunct for fatigue overlay (Shan 2025)","Non-tDCS days","Stress-sensitive motor worsening, fatigue"],
            ["CES + Standard Care","CES supports sleep/anxiety; indirect motor benefit (Philip 2017)","Evening","Sleep/anxiety overlay on motor symptoms"],
        ]),
        ("2) Tremor-Dominant", [
            ["tDCS + Cueing / Training","Cerebellar or M1 tDCS + behavioural strategies; evidence less robust than general motor (Elsner 2016)","tDCS before tremor-management training","Action/postural tremor with attentional sensitivity"],
            ["taVNS + Standard Care","May help autonomic arousal; direct tremor evidence preliminary (Matsuoka 2025)","Before stressful functional tasks","Tremor flares with anxiety/autonomic activation"],
            ["CES + Standard Care","Reduces anxiety/insomnia that aggravate tremor (Philip 2017)","Evening","Tremor + insomnia/anxiety loop"],
        ]),
        ("3) Executive Dysfunction", [
            ["tDCS + Cognitive Training","Prefrontal tDCS + cognitive/dual-task training (Conceicao 2021; Moshayedi 2025)","tDCS before/during cognitive training","Dysexecutive PD-MCI, slowing, dual-task difficulty"],
            ["taVNS + Cognitive Training","Modulates attention networks; exploratory adjunct only (Gerges 2024; Matsuoka 2025)","Before low-load cognitive tasks","Fluctuating attention, fatigue-driven executive failures"],
            ["CES + Standard Care","Supports sleep/anxiety indirectly affecting cognition (Philip 2017)","Evening","Executive dysfunction with insomnia/anxiety"],
        ]),
        ("4) Depression / Apathy", [
            ["tDCS + Behavioural Activation","Prefrontal tDCS + broader care plan for mood/apathy (Philip 2017; Moshayedi 2025)","tDCS before activation tasks","Depression, apathy, anhedonia"],
            ["taVNS + Standard Care","Plausible mood-regulation rationale; PD-specific antidepressant evidence early (Matsuoka 2025)","On off-days or stress-sensitive periods","Depression with autonomic anxiety or rumination"],
            ["CES + Standard Care","Supportive for anxiety/sleep; evidence quality low (Philip 2017)","Evening, 4\u20137 nights/week if prescribed","Depression with insomnia, somatic tension, anxiety"],
        ]),
        ("5) Freezing of Gait", [
            ["tDCS + Gait Rehabilitation","Best-supported PD combination: prefrontal/motor tDCS + gait training (Conceicao 2021; Elsner 2016)","tDCS before/during gait training","FOG, turning hesitation, dual-task gait collapse"],
            ["taVNS + Gait Rehabilitation","Plausible; current trials show small effects; exploratory (Sigurdsson 2021, 2025; Shan 2025)","Before gait practice or on non-tDCS days","Anxiety-triggered freezing, autonomic spikes"],
            ["CES + Standard Care","Helps sleep/hyperarousal; not a gait intervention (Philip 2017)","Evening","FOG with insomnia/fatigue overlay"],
        ]),
        ("6) Pain Phenotype", [
            ["tDCS + Graded Activity","Motor/prefrontal tDCS analgesic role; PD-specific pain evidence limited (Moshayedi 2025)","tDCS before activity or mobility work","Central pain, stiffness-pain loop, MSK amplification"],
            ["taVNS + Standard Care","Reduces autonomic arousal; mechanistic relevance for pain modulation (Adair 2020; Matsuoka 2025)","Before flare-prone periods or on off-days","Pain with autonomic arousal/catastrophising"],
            ["CES + Standard Care","Supports sleep and anxiety reduction; supportive only (Philip 2017)","Evening most days if prescribed","Pain + insomnia + anxiety"],
        ]),
    ],
    combination_summary=[
        ["Bradykinetic/Rigid",    "tDCS + Physiotherapy",        "Strongest non-TPS combination: anodal tDCS over M1/DLPFC + structured rehab (Elsner 2016; Moshayedi 2025)",  "tDCS before/during rehab", "Bradykinesia, rigidity",  "Moderate"],
        ["Bradykinetic/Rigid",    "taVNS + Standard Care",       "Feasible; small/inconsistent motor benefit; better as fatigue/autonomic adjunct (Shan 2025)",                 "Non-tDCS days",            "Stress-sensitive worsening","Emerging"],
        ["Tremor-Dominant",       "tDCS + Cueing",               "Evidence less robust; frame as trial adjunct to behavioural strategies (Elsner 2016)",                        "tDCS before training",     "Action/postural tremor",    "Limited"],
        ["Executive Dysfunction", "tDCS + Cognitive Training",   "Defensible: prefrontal tDCS + cognitive/dual-task training (Conceicao 2021; Moshayedi 2025)",                 "tDCS before cognitive Rx", "Dysexecutive PD-MCI",      "Moderate"],
        ["Depression/Apathy",     "tDCS + Behavioural Activation","Prefrontal tDCS + broader care plan for mood/apathy (Philip 2017; Moshayedi 2025)",                          "tDCS before activation",   "Depression, apathy",        "Moderate"],
        ["Freezing of Gait",      "tDCS + Gait Rehabilitation",  "Best-supported PD combination; prefrontal/motor tDCS + gait training (Conceicao 2021; Elsner 2016)",          "tDCS before/during gait",  "FOG, dual-task gait",       "Moderate"],
    ],
    outcomes=[
        ["MDS-UPDRS Part III",    "Motor function (OFF-medication state)",  "Baseline, weeks 4, 8, 12",     "\u226530% improvement = meaningful response"],
        ["MoCA",                  "Global cognition",                        "Baseline, month 3",            "Score <16 suggests severe impairment"],
        ["PHQ-9",                 "Depressive symptoms",                     "Baseline, weeks 4, 8, 12",     "Score \u226510 moderate; \u226515 severe"],
        ["GAD-7",                 "Anxiety symptoms",                        "Baseline, month 3",            "Score \u226510 clinically significant"],
        ["FOG-Q",                 "Gait freezing severity",                  "Baseline, weeks 4, 8, 12",     "Score >0 some freezing; >17 severe"],
        ["10-Meter Walk Test",    "Gait speed",                              "Baseline, weeks 4, 8, 12",     "Normal >0.8 m/s; PD typically 0.4\u20130.6 m/s"],
        ["Timed Up and Go (TUG)","Balance and fall risk",                   "Baseline, months 1, 3",        ">14 seconds = high fall risk"],
        ["Hoehn & Yahr Scale",    "Disease stage",                           "Baseline, month 3",            "I\u2013IV; determines inclusion and motor response tier"],
        ["SOZO PRS",              "NIBS-specific functional outcome",        "Baseline, weeks 2, 4, 8, 12", "Proprietary; composite of motor/non-motor domains"],
    ],
    outcomes_abbrev=["MDS-UPDRS III", "MoCA", "PHQ-9", "FOG-Q", "SOZO PRS"],
)

# ══════════════════════════════════════
# 2. MAJOR DEPRESSIVE DISORDER
# ══════════════════════════════════════
CONDITIONS["depression"] = dict(
    name="Major Depressive Disorder", icd10="F32/F33",
    doc_num="SOZO-MDD-FEL-002",
    tps_status="INVESTIGATIONAL / OFF-LABEL",
    tdcs_class="Level A evidence (multiple RCTs); CE-marked",
    tps_class="INVESTIGATIONAL / OFF-LABEL",
    tavns_class="Emerging adjunctive; CE-marked",
    ces_class="FDA-cleared adjunctive for mood/anxiety/sleep",
    inclusion=[
        "Confirmed diagnosis of MDD (DSM-5 or ICD-11 criteria) by qualified clinician",
        "PHQ-9 score \u226510 at baseline (moderate-to-severe)",
        "Age 18+ years",
        "Stable or no psychotropic medication (no change in 4 weeks prior to enrolment)",
        "Written informed consent (including TPS off-label disclosure)",
        "Baseline assessments completed (PHQ-9, HDRS-17, MoCA, GAD-7)",
    ],
    exclusion=[
        "Intracranial metallic hardware in stimulation path",
        "Cochlear implant or DBS device",
        "Active psychotic features or bipolar I disorder (manic episode)",
        "Active suicidal ideation with plan or intent (requires immediate psychiatric referral)",
        "Skull defects, craniectomy, or recent craniotomy",
        "Pregnancy (tDCS, TPS)",
        "Inability to provide informed consent",
    ],
    discussion=[
        "Cardiac pacemaker or defibrillator \u2014 individual risk\u2013benefit assessment",
        "Epilepsy or seizure history \u2014 formal risk\u2013benefit with documented rationale",
        "Bipolar II disorder or cyclothymia \u2014 requires mood stabiliser coverage and close monitoring",
        "Active substance use disorder",
        "Benzodiazepine use \u2014 may reduce cortical excitability and blunt tDCS response",
        "Coagulation disorders or anticoagulants (especially TPS)",
        "Dermatological conditions at electrode sites",
        "Severe personality disorder complicating treatment engagement",
    ],
    overview=[
        "Major Depressive Disorder (MDD) is a highly prevalent, recurrent psychiatric disorder characterised "
        "by persistent low mood, anhedonia, and a range of cognitive, somatic, and neurovegetative symptoms. "
        "MDD affects approximately 264 million people globally, with lifetime prevalence of 15\u201320% in "
        "high-income countries. It is a leading cause of disability worldwide (WHO, 2021).",
        "Approximately 30\u201340% of MDD patients fail to respond to two or more adequate antidepressant "
        "trials, meeting criteria for Treatment-Resistant Depression (TRD). tDCS targeting the left DLPFC "
        "represents one of the best-evidenced non-invasive neuromodulation approaches in psychiatry, "
        "supported by multiple sham-controlled RCTs and meta-analyses (Brunoni et al., 2013, 2016).",
        "The network model of MDD identifies three core disruptions: (1) Default Mode Network (DMN) "
        "hyperactivation driving rumination; (2) Central Executive Network (CEN) hypoactivation impairing "
        "cognitive control; (3) Salience Network (SN) dysregulation disrupting network switching. Prefrontal "
        "tDCS is proposed to restore CEN excitability and reduce SN-DMN coupling.",
        "NIBS Evidence in MDD: tDCS is the most evidence-based non-invasive neuromodulation for MDD after "
        "TMS. Multiple RCTs and meta-analyses demonstrate clinically meaningful antidepressant effects "
        "with left DLPFC anodal tDCS (Brunoni et al., 2016; Lefaucheur et al., 2017). TPS is investigational "
        "in MDD with emerging pilot data. CES is FDA-cleared as adjunctive for mood and anxiety.",
    ],
    pathophysiology=[
        "MDD involves dysfunction across multiple neurobiological systems. Left DLPFC hypoactivation and "
        "subgenual ACC (sgACC/Cg25) hyperactivation are consistently demonstrated in neuroimaging studies. "
        "The sgACC hyperactivity propagates depressogenic tone throughout the DMN. Monoaminergic deficits "
        "(serotonin, noradrenaline, dopamine) underpin the pharmacological treatment rationale.",
        "Neuroplasticity impairment is key: reduced BDNF signalling, hippocampal volume loss in recurrent MDD, "
        "and impaired synaptic plasticity in prefrontal circuits. HPA axis hyperactivity drives "
        "glucocorticoid-mediated prefrontal and hippocampal damage. Neuroinflammatory markers (IL-6, TNF-alpha, "
        "CRP) are elevated in a significant subset of MDD patients.",
    ],
    std_treatment=[
        "First-line pharmacotherapy includes SSRIs and SNRIs. Tricyclic antidepressants and MAOIs are "
        "second-line due to side-effect profiles. Lithium and atypical antipsychotics are used for augmentation "
        "in TRD. Psychotherapy (CBT, behavioural activation) has comparable evidence to pharmacotherapy for "
        "mild-to-moderate MDD and is recommended in combination with NIBS. ECT remains the gold standard for "
        "severe, treatment-resistant MDD. rTMS (left DLPFC) is FDA-cleared and NICE-recommended for TRD.",
    ],
    symptoms=[
        ["Depressed Mood",           "Persistent low mood most of the day, nearly every day",                          "Core criterion",    "DSM-5; ICD-11"],
        ["Anhedonia",                "Markedly diminished interest or pleasure in all/almost all activities",           "Core criterion",    "DSM-5; ICD-11"],
        ["Weight / Appetite Change", "Significant weight loss/gain; decreased/increased appetite",                     "60\u201380%",       "DSM-5"],
        ["Sleep Disturbance",        "Insomnia (early morning awakening) or hypersomnia",                              "70\u201390%",       "DSM-5; Riemann 2020"],
        ["Psychomotor Changes",      "Psychomotor agitation or retardation observable by others",                      "40\u201360%",       "DSM-5"],
        ["Fatigue / Anergia",        "Fatigue or loss of energy nearly every day",                                     "70\u201390%",       "DSM-5"],
        ["Worthlessness / Guilt",    "Feelings of worthlessness or excessive/inappropriate guilt",                     "60\u201380%",       "DSM-5"],
        ["Cognitive Impairment",     "Diminished ability to think, concentrate, or make decisions",                    "50\u201380%",       "Rock et al. 2014"],
        ["Suicidal Ideation",        "Recurrent thoughts of death; suicidal ideation or attempt",                      "Variable",          "DSM-5"],
        ["Somatic Symptoms",         "Headache, GI symptoms, pain without clear medical cause",                        "50\u201375%",       "Simon et al. 1999"],
        ["Anxiety Comorbidity",      "Comorbid anxiety symptoms present in majority of MDD patients",                  "50\u201370%",       "Kessler et al. 2003"],
    ],
    brain_regions=[
        ["Left DLPFC (F3/F4)",          "Hypoactivation in MDD; impaired cognitive control and emotion regulation",      "Anodal tDCS primary target; restores CEN excitability (Level A evidence)",     "Brunoni 2016; Lefaucheur 2017"],
        ["Subgenual ACC (sgACC/Cg25)",   "Hyperactivation; propagates depressogenic tone via DMN; key DBS target",       "TPS investigational deep target; indirectly modulated by prefrontal tDCS",     "Mayberg 2005"],
        ["Anterior Cingulate Cortex",    "Conflict monitoring, pain, emotional processing; disrupted in MDD",             "Cathodal tDCS of right PFC / anodal left; TPS (investigational)",             "Drevets 2001"],
        ["Amygdala",                     "Hyperactivation to negative stimuli; impaired fear extinction",                  "Modulated indirectly via prefrontal and taVNS limbic pathways",               "Sheline 2001"],
        ["Hippocampus",                  "Volume loss in recurrent MDD; impaired neurogenesis; memory consolidation",     "Indirect target via TPS (deep investigational) and BDNF upregulation by tDCS", "MacQueen 2003"],
        ["Medial Prefrontal Cortex",     "DMN hub; excessive self-referential processing and rumination",                  "Prefrontal tDCS modulates DMN-CEN balance",                                   "Sheline 2009"],
        ["Insula",                       "Interoceptive processing, somatic symptoms; hyperactivation in MDD",             "Indirectly modulated; taVNS acts on insula-brainstem-limbic circuits",        "Mayberg 2005"],
        ["Basal Ganglia (Striatum)",     "Anhedonia circuitry; reduced reward activation in MDD",                         "TPS deep investigational target; prefrontal tDCS modulates fronto-striatal circuits", "Pizzagalli 2009"],
    ],
    brainstem=[
        ["Locus Coeruleus (LC)",        "Noradrenergic dysregulation; arousal impairment; antidepressant target",       "taVNS modulates LC via NTS; key mechanism for mood stabilisation",             "Nieuwenhuis 2005"],
        ["Dorsal Raphe Nucleus (DRN)",  "Serotonergic dysregulation; primary antidepressant pharmacological target",    "Indirectly modulated by prefrontal tDCS and taVNS pathways",                   "Deakin 1991"],
        ["Nucleus Tractus Solitarius",  "Vagal afferent relay; central autonomic regulation hub",                       "Primary afferent target of taVNS; influences limbic and prefrontal regions",  "Frangos 2015"],
        ["Habenula",                    "Anti-reward signal; hyperactivation in MDD contributes to anhedonia",           "Indirect target via prefrontal modulation and taVNS pathways",                 "Lecca 2014"],
        ["VTA (Ventral Tegmental Area)","Dopaminergic reward processing; hypoactivation in anhedonic MDD",              "Indirectly targeted via mesolimbic pathway modulation (tDCS, taVNS)",          "Nestler 2006"],
    ],
    phenotypes=[
        ["Melancholic MDD",             "Severe anhedonia, psychomotor changes, early morning wakening, diurnal variation","Anhedonia, early waking, psychomotor retardation/agitation","Left DLPFC anodal tDCS; TPS DLPFC + sgACC; taVNS"],
        ["Anxious MDD",                 "Prominent anxiety, tension, worry alongside depressive core symptoms",           "Low mood, anxiety, worry, insomnia, somatic tension",           "Bilateral DLPFC tDCS; taVNS primary; CES"],
        ["Cognitive MDD",               "Cognitive symptoms dominant: concentration, memory, executive dysfunction",      "Concentration loss, memory impairment, indecisiveness",         "Left DLPFC anodal tDCS; TPS DLPFC; cognitive training"],
        ["Atypical MDD",                "Mood reactivity; hypersomnia; hyperphagia; rejection sensitivity; leaden paralysis","Variable mood, hypersomnia, overeating, heavy limbs",       "L-DLPFC tDCS; taVNS; CES; consider right parietal cathodal"],
        ["TRD (Treatment-Resistant)",   "Failed \u22652 adequate antidepressant trials",                                 "Full MDD syndromal features with pharmacological non-response",  "Bilateral DLPFC tDCS at 2 mA; TPS DLPFC + ACC; taVNS; CES"],
        ["MDD with Somatic Features",   "Prominent somatic complaints: fatigue, pain, headache, GI symptoms",            "Fatigue, pain, concentration, mood",                            "L-DLPFC tDCS + M1 (for pain); taVNS; CES"],
        ["MDD with Psychotic Features", "Mood-congruent psychotic features; hallucinations or delusions",                "Depressed mood + psychotic symptoms",                           "Doctor assessment required; tDCS only after antipsychotic cover"],
    ],
    symptom_map=[
        ["Depressed Mood",             "Left DLPFC, sgACC, DMN",       "Left DLPFC anodal tDCS (F3) + TPS (DLPFC/sgACC investigational) + taVNS",         "Level A (Brunoni 2016)"],
        ["Anhedonia",                  "Striatum, OFC, DLPFC",          "Left DLPFC tDCS + TPS (OFC/striatum investigational) + taVNS",                    "Moderate (Brunoni 2013)"],
        ["Cognitive Impairment",       "Left DLPFC, Parietal, Hippocampus","Left DLPFC anodal tDCS + TPS (DLPFC targeted)",                               "Moderate (Brunoni 2013)"],
        ["Anxiety / Tension",          "Right DLPFC, Amygdala, SN",    "Right DLPFC cathodal / bilateral DLPFC + taVNS primary + CES",                    "Moderate (Loo 2012)"],
        ["Sleep Disturbance",          "Frontal slow-wave, Thalamus",   "CES primary; taVNS adjunctive; tDCS (frontal slow-wave protocols)",               "Moderate (Philip 2017)"],
        ["Fatigue / Anergia",          "Left DLPFC, Mesolimbic",        "Left DLPFC tDCS + CES; taVNS for arousal modulation",                            "Moderate (Brunoni 2016)"],
        ["Psychomotor Retardation",    "Left DLPFC, SMA, Striatum",     "Left DLPFC + SMA anodal tDCS; TPS frontal targeted; behavioural activation",      "Limited (Loo 2012)"],
        ["Suicidality Monitoring",     "Prefrontal-limbic circuits",    "Protocol continues with enhanced monitoring; escalate to Doctor if PHQ-9 item 9 \u22651","Monitoring priority"],
        ["Somatic Pain Overlay",       "M1, ACC, Thalamus",             "Left DLPFC tDCS + M1 cathodal (contralateral to pain); taVNS; CES",              "Limited (Fregni 2006)"],
    ],
    montage=[
        ["Melancholic MDD",            "F3 anode + F4 cathode (standard L-DLPFC protocol); 2 mA, 20\u201330 min", "Cranial (global + targeted DLPFC/sgACC)",    "taVNS, CES (sleep)"],
        ["Anxious MDD",                "F3 anode + F4 anode (bilateral); or F3 anode + F4 cathode",               "Cranial (global + targeted DLPFC/ACC)",       "taVNS primary, CES"],
        ["Cognitive MDD",              "F3 anode + Fp2 cathode; T3/T4 options for temporal memory targets",       "Cranial (global + targeted DLPFC)",          "taVNS, CES supportive"],
        ["Atypical MDD",               "F3 anode + Fp2 cathode; consider P3 anode for parietal involvement",     "Cranial (global + targeted)",                "taVNS, CES; sleep focus"],
        ["TRD",                        "Bilateral DLPFC (F3 + F4 anode); or high-density HD-tDCS if available",  "Cranial (global + targeted DLPFC/ACC) \u00d7 intensive block","taVNS + CES; consider adjunct CBT"],
        ["MDD + Somatic Pain",         "F3 anode + C3/C4 cathode (contralateral to pain) / extracephalic",       "Cranial (global + targeted) + peripheral (pain dermatomes)","taVNS, CES supportive"],
        ["MDD + Psychotic Features",   "Defer tDCS until antipsychotic stabilisation achieved",                  "Defer TPS until psychosis in remission",      "CES; taVNS adjunctive only"],
        ["No response after 8\u201310 sessions","Reassess diagnosis; consider bilateral protocol or TPS",       "Add TPS DLPFC; consider sgACC targeting",    "Review full adjunct regimen"],
    ],
    tdcs_protocols=[
        ["C1","Depressed mood (primary)",      "F3 (L-DLPFC)","Fp2 / F4 / extracephalic",        "2 mA, 20\u201330 min, 5\u00d7/wk \u00d7 6\u201310 wks","Level A; RCT meta-analysis (Brunoni 2016)"],
        ["C2","Anxious depression",            "F3 + F4 bilateral anode","Extracephalic shoulder/arm","2 mA, 20\u201330 min, 5\u00d7/wk",               "Moderate (Loo 2012)"],
        ["C3","Cognitive symptoms / TRD",      "F3 anode + Fz","Fp2, extracephalic",              "2 mA, 20\u201330 min, 5\u00d7/wk",                   "Moderate (Brunoni 2013)"],
        ["C4","Anhedonia / reward",            "F3 anode + Oz","F4, extracephalic",               "2 mA, 20\u201330 min, 5\u00d7/wk",                   "Emerging (Lefaucheur 2017)"],
        ["C5","TRD \u2014 intensive bilateral", "F3 + F4 bilateral","Fp2 + Fp1, extracephalic",   "2 mA, 20\u201330 min, 2\u00d7/day \u00d7 5 days then 5\u00d7/wk","SELECT-TDCS (Brunoni 2013)"],
        ["C6","Sleep disturbance",             "Fz anode (frontal slow-wave)","Pz / extracephalic","1\u20132 mA, 20 min, evening",                       "Moderate (Brunoni 2016; Philip 2017)"],
        ["C7","Somatic fatigue / anergia",     "F3 anode + Cz","Fp2, extracephalic",              "2 mA, 20\u201330 min, 5\u00d7/wk",                   "Moderate (Brunoni 2016)"],
        ["C8","MDD + pain overlay",            "F3 anode + C3/C4 (contra to pain)","F4, extracephalic","2 mA, 20\u201330 min, 5\u00d7/wk",              "Limited; analgesic tDCS evidence (Fregni 2006)"],
    ],
    plato_protocols=[
        ["C1-PS","Depressed mood",             "Frontal (F3 area)","F3","Right shoulder",   "1.6 mA, 20\u201330 min","Daily or 5\u00d7/wk","L-DLPFC anodal; primary MDD protocol"],
        ["C2-PS","Anxious depression",         "Frontal (F3 area)","F3","Right shoulder",   "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Adjunct taVNS for anxiety component"],
        ["C3-PS","Cognitive / TRD",            "Frontal (F3 area)","F3","Shoulder",         "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Pair with cognitive activation tasks"],
        ["C4-PS","Anhedonia / reward",         "Frontal (F3 area)","F3","Shoulder",         "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Combine with behavioural activation"],
        ["C5-PS","TRD maintenance",            "Frontal (F3 area)","F3","Right shoulder",   "1.6 mA, 20\u201330 min","5\u00d7/wk",         "After in-clinic HDCkit induction phase"],
        ["C6-PS","Sleep disturbance",          "Frontal (Fz area)","Fz","Shoulder",         "1.6 mA, 20 min",        "Evening only",       "Combine with CES evening session"],
        ["C7-PS","Fatigue / anergia",          "Frontal (F3 area)","F3","Shoulder",         "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Morning use preferred for energy"],
        ["C8-PS","MDD + pain",                 "Motor (Cz area)",  "Cz","Shoulder/arm",     "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Motor cortex for pain modulation"],
    ],
    tps_protocols=[
        ["T1","Depressed mood (primary)",        "DLPFC bilateral, mPFC, OFC",   "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz; 10 sessions/2 wks","4,000 Tgt (DLPFC) + 2,000 PRE; total 6\u20138K","Preliminary (Leyman 2024)"],
        ["T2","TRD / sgACC modulation",          "DLPFC + sgACC (Cg25) region",  "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",               "4,000 Tgt + 2,000 Std; total 6\u20138K",          "Investigational (pilot data only)"],
        ["T3","Anhedonia / reward",              "OFC, striatum, ACC",            "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",               "4,000 Tgt (OFC) + 2,000 Std; total 6\u20138K",   "Investigational"],
        ["T4","Cognitive MDD / memory",          "DLPFC + Hippocampal region",   "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",               "4,000 Tgt (DLPFC) + 2,000 PRE; total 6\u20138K", "Investigational"],
        ["T5","Anxious depression",              "DLPFC + anterior insula",      "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",               "4,000 Tgt + 2,000 Std; total 6\u20138K",          "Investigational"],
    ],
    ces_role=[
        ["Mood stabilisation",    "Supportive antidepressant adjunct; augments tDCS mood benefit",                "Daily or 5\u00d7/wk; 20\u201340 min session; morning or evening"],
        ["Anxiety reduction",     "Reduces comorbid anxiety and somatic tension; reduces CEN-SN hyperarousal",    "Before tDCS session or evening; 20\u201340 min"],
        ["Sleep disturbance",     "Enhances sleep onset, duration, and quality; key target in MDD",               "Evening before bed; 20\u201340 min; coordinate with tDCS timing"],
        ["Fatigue / anergia",     "May improve subjective energy and motivation as adjunct to prefrontal tDCS",   "Morning; 20\u201330 min; avoid late afternoon if activating"],
    ],
    tavns_role="taVNS serves as a limbic and autonomic modulator in MDD, enhancing noradrenergic and serotonergic tone via locus coeruleus and dorsal raphe, reducing amygdala hyperactivity, and augmenting antidepressant tDCS effects.",
    combinations=[
        ("1) Melancholic MDD", [
            ["tDCS + Behavioural Activation","L-DLPFC anodal tDCS + structured behavioural activation; most supported combination in MDD neuromodulation (Brunoni 2016; Loo 2018)","tDCS before activation tasks","Severe anhedonia, psychomotor retardation"],
            ["taVNS + tDCS","taVNS priming before tDCS enhances cortical receptivity; plausible combination (Redgrave 2018; Hein 2021)","taVNS 20\u201330 min before tDCS session","Melancholic features with autonomic dysregulation"],
            ["CES + Standard Care","Evening CES for sleep; morning CES for energy; adjunct to tDCS treatment block (Philip 2017)","Evening and/or morning as prescribed","Sleep disturbance, early waking, anergia"],
        ]),
        ("2) Anxious MDD", [
            ["tDCS + CBT / Mindfulness","Bilateral DLPFC tDCS + CBT or mindfulness; addresses both emotional and cognitive regulation (Brunoni 2016)","tDCS before therapy session","Anxious rumination, tension, hyperarousal"],
            ["taVNS + tDCS","taVNS provides autonomic downshift before cortical stimulation; reduces treatment-session anxiety (Hein 2021; Redgrave 2018)","taVNS 20 min before tDCS","Anxiety, somatic tension, autonomic hyperarousal"],
            ["CES + Standard Care","CES reduces anxiety and improves sleep in anxious MDD; FDA-cleared indication (Philip 2017)","Evening \u00d7 7 nights/wk or as needed","Anxiety, sleep onset difficulties, somatic tension"],
        ]),
        ("3) Cognitive MDD", [
            ["tDCS + Cognitive Training","L-DLPFC anodal tDCS + computerised cognitive training; targets executive and memory deficits (Brunoni 2016; Segrave 2014)","tDCS immediately before or during cognitive training","Concentration loss, memory impairment, indecisiveness"],
            ["taVNS + Cognitive Training","taVNS enhances prefrontal excitability; adjunct to cognitive rehabilitation (Hein 2021)","taVNS before cognitive training session","Attentional fatigue, processing speed reduction"],
            ["CES + Standard Care","Sleep and energy support; indirectly improves cognitive performance in MDD (Philip 2017)","Evening (sleep) and morning (energy) sessions","Cognitive MDD with insomnia and fatigue"],
        ]),
        ("4) TRD (Treatment-Resistant)", [
            ["Intensive bilateral tDCS + Pharmacotherapy","Bilateral DLPFC tDCS (SELECT-TDCS protocol) combined with maintained antidepressant (Brunoni 2013); consider TPS add-on","Intensive: 2\u00d7/day \u00d7 5 days induction, then 5\u00d7/wk","Failed \u22652 adequate antidepressant trials"],
            ["taVNS + Intensive tDCS","taVNS as adjunct priming for intensive tDCS block; addresses autonomic and limbic components of TRD (Hein 2021; Redgrave 2018)","taVNS before each tDCS session during intensive block","TRD with anxiety, autonomic, or melancholic features"],
            ["CES + Intensive Protocol","CES daily throughout TRD treatment block for sleep and anxiety; reduces dropout (Philip 2017)","Evening daily throughout TRD block","TRD with severe sleep disturbance and anxiety"],
        ]),
        ("5) MDD + Somatic Pain", [
            ["tDCS (DLPFC + M1) + Graded Activity","L-DLPFC anodal + M1 cathodal (contra to pain) tDCS with graded activity/physiotherapy; targets both mood and pain circuits (Fregni 2006)","tDCS before activity/physiotherapy","MDD with comorbid chronic pain or fibromyalgia"],
            ["taVNS + tDCS","taVNS modulates ascending pain pathways and limbic-pain circuits; additive with tDCS for MDD-pain overlap (Adair 2020)","taVNS before tDCS session","Depression with autonomic dysregulation and pain"],
            ["CES + Standard Care","CES addresses sleep, anxiety, and pain-related distress; supportive throughout treatment block (Philip 2017)","Evening (sleep) and as needed for pain/anxiety","MDD with chronic pain, insomnia, somatic distress"],
        ]),
        ("6) MDD + Cognitive Decline (Older Adults)", [
            ["tDCS + Cognitive Stimulation","L-DLPFC anodal tDCS + cognitive stimulation; benefits cognitive symptoms alongside depression in older adults (Brunoni 2016; Yamada 2020)","tDCS before cognitive stimulation activity","MDD in older adults with mild cognitive complaints"],
            ["taVNS + Standard Care","taVNS supports autonomic stability and sleep quality; low-risk in older adults (Hein 2021)","Daily or 5\u00d7/wk; morning or pre-session","MDD with autonomic dysregulation, fatigue, insomnia"],
            ["CES + Standard Care","CES reduces anxiety/insomnia in older adults with MDD; FDA-cleared; generally well-tolerated (Philip 2017)","Evening; 20\u201340 min","MDD in older adults with anxiety and sleep disturbance"],
        ]),
    ],
    combination_summary=[
        ["Melancholic MDD",    "tDCS + Behavioural Activation","L-DLPFC anodal tDCS + structured behavioural activation (Brunoni 2016; Loo 2018)","tDCS before activation tasks","Severe anhedonia, retardation","Level A (tDCS); Moderate overall"],
        ["Melancholic MDD",    "taVNS + tDCS",                  "Priming: taVNS before tDCS enhances cortical receptivity (Hein 2021; Redgrave 2018)","taVNS 20\u201330 min before tDCS","Melancholic + autonomic features","Emerging"],
        ["Anxious MDD",        "tDCS + CBT",                    "Bilateral DLPFC tDCS + CBT; addresses emotional and cognitive regulation (Brunoni 2016)","tDCS before therapy","Anxious rumination, hyperarousal","Moderate"],
        ["Anxious MDD",        "taVNS + tDCS",                  "taVNS autonomic downshift before cortical stimulation (Hein 2021)","taVNS 20 min before tDCS","Anxiety, somatic tension","Emerging"],
        ["TRD",                "Intensive bilateral tDCS",       "Bilateral DLPFC tDCS SELECT-TDCS protocol + maintained antidepressant (Brunoni 2013)","Intensive: 2\u00d7/day \u00d7 5 days then 5\u00d7/wk","TRD, pharmacological non-response","Level A"],
        ["MDD + Somatic Pain", "tDCS (DLPFC + M1) + Activity",  "L-DLPFC + M1 tDCS with graded activity; targets mood and pain circuits (Fregni 2006)","tDCS before activity/physio","MDD with comorbid chronic pain","Limited\u2013Moderate"],
    ],
    outcomes=[
        ["PHQ-9",                   "Depressive symptoms (self-report)",  "Baseline, weeks 2, 4, 8, 12",  "\u226510 moderate; \u226515 severe; \u22655-point change = response"],
        ["HDRS-17",                 "Clinician-rated depression severity", "Baseline, weeks 4, 8, 12",     "\u226550% reduction from baseline = response; \u22647 = remission"],
        ["BDI-II",                  "Depressive symptoms (self-report)",  "Baseline, weeks 4, 8, 12",     "Score 0\u201313 minimal; 14\u201319 mild; 20\u201328 moderate; 29\u201363 severe"],
        ["MoCA",                    "Global cognition screening",          "Baseline, month 3",            "Score <26 cognitive impairment; track improvement"],
        ["GAD-7",                   "Comorbid anxiety symptoms",           "Baseline, weeks 4, 8, 12",     "\u226510 moderate; \u226515 severe"],
        ["QIDS-SR",                 "Quick Inventory of Depressive Symptoms","Baseline, weeks 2, 4, 8, 12","Score 0\u20135 normal; 6\u201310 mild; 11\u201315 moderate; \u226516 severe"],
        ["PHQ-9 Item 9",            "Suicidality monitoring",              "Every session",                "Any score \u22651 \u2192 Doctor review required immediately"],
        ["CGI-S / CGI-I",           "Clinical Global Impression",          "Baseline, weeks 4, 8, 12",     "CGI-I \u22642 = much improved; \u22643 = minimally improved"],
        ["SOZO PRS",                "NIBS-specific functional outcome",    "Baseline, weeks 2, 4, 8, 12", "Proprietary; composite mood/cognitive/functional domains"],
    ],
    outcomes_abbrev=["PHQ-9", "HDRS-17", "MoCA", "GAD-7", "SOZO PRS"],
)

# ══════════════════════════════════════
# 3. GENERALIZED ANXIETY DISORDER
# ══════════════════════════════════════
CONDITIONS["anxiety"] = dict(
    name="Generalized Anxiety Disorder", icd10="F41.1",
    doc_num="SOZO-GAD-FEL-003",
    tps_status="INVESTIGATIONAL / OFF-LABEL",
    tdcs_class="Level B evidence; CE-marked",
    tps_class="INVESTIGATIONAL / OFF-LABEL",
    tavns_class="Emerging; CE-marked; primary adjunct modality",
    ces_class="FDA-cleared adjunctive for anxiety/mood/sleep",
    inclusion=[
        "Confirmed diagnosis of GAD or primary anxiety disorder (DSM-5 or ICD-11) by qualified clinician",
        "GAD-7 score \u226510 at baseline (moderate-to-severe anxiety)",
        "Age 18+ years",
        "Stable or no anxiolytic/antidepressant medication (no change in 4 weeks prior to enrolment)",
        "Written informed consent (including TPS off-label disclosure)",
        "Baseline assessments completed (GAD-7, HAMA, PHQ-9, ISI)",
    ],
    exclusion=[
        "Intracranial metallic hardware in stimulation path",
        "Cochlear implant or DBS device",
        "Active psychotic features",
        "Severe alcohol or substance dependence as primary diagnosis",
        "Skull defects, craniectomy, or recent craniotomy",
        "Pregnancy (tDCS, TPS)",
        "Inability to provide informed consent",
    ],
    discussion=[
        "Cardiac pacemaker or defibrillator \u2014 individual risk\u2013benefit assessment",
        "Epilepsy or seizure history \u2014 formal risk\u2013benefit with documented rationale",
        "Benzodiazepine dependence \u2014 protocol requires specialist input; titration planning required",
        "Comorbid MDD requiring protocol modification",
        "PTSD comorbidity \u2014 trauma-informed approach required",
        "Coagulation disorders or anticoagulants (especially TPS)",
        "Dermatological conditions at electrode sites",
        "Severe health anxiety or cyberchondria about device use",
    ],
    overview=[
        "Generalized Anxiety Disorder (GAD) is characterised by excessive, uncontrollable worry about "
        "multiple domains occurring on more days than not for at least 6 months, accompanied by somatic "
        "symptoms such as muscle tension, sleep disturbance, and fatigue. GAD has a lifetime prevalence "
        "of 5\u20139% and is one of the most common anxiety disorders worldwide (Kessler et al., 2005).",
        "Neurobiologically, GAD involves hyperactivation of the amygdala and anterior insula, "
        "hypoactivation of the prefrontal cortex (particularly the right DLPFC), and dysregulation of "
        "the salience network (SN) and default mode network (DMN). The SN-DMN interaction underpins "
        "the ruminative and threat-hypervigilant phenomenology of GAD.",
        "tDCS targeting the right DLPFC (cathodal) and/or left DLPFC (anodal) modulates prefrontal "
        "control over amygdala hyperactivity. taVNS is a primary adjunctive modality given its role "
        "in autonomic downregulation via the vagal-NTS-LC pathway.",
        "NIBS Evidence in GAD: Evidence for tDCS in anxiety is growing but less mature than in MDD. "
        "Right DLPFC cathodal tDCS (to reduce hyperactivation) and bilateral approaches have shown "
        "anxiolytic effects in RCTs. taVNS is CE-marked and FDA-cleared (VNS) for depression/anxiety "
        "and is a key component of the SOZO anxiety protocol. CES is FDA-cleared for anxiety.",
    ],
    pathophysiology=[
        "GAD pathophysiology involves a failure of prefrontal inhibitory control over amygdala-driven "
        "threat responses. Right DLPFC and vmPFC show reduced activation, impairing emotion regulation "
        "and worry termination. The amygdala and anterior insula are hyperactivated, amplifying interoceptive "
        "threat signals. HPA axis dysregulation and chronic cortisol elevation contribute to hippocampal "
        "volume changes and further impair prefrontal function.",
        "The salience network (insula, dACC) shows persistent hyperactivation, maintaining threat vigilance "
        "and preventing effective disengagement from worry. GABAergic deficits (particularly in prefrontal "
        "and limbic circuits) underpin the anxiolytic mechanism of benzodiazepines and are relevant to "
        "neuromodulation targeting.",
    ],
    std_treatment=[
        "First-line pharmacotherapy includes SSRIs (sertraline, escitalopram) and SNRIs (venlafaxine, "
        "duloxetine). Buspirone is an alternative. Benzodiazepines are reserved for short-term or "
        "acute use due to dependence risk. Pregabalin has evidence for GAD. Cognitive Behavioural Therapy "
        "(CBT) has the strongest evidence base for GAD and is recommended in combination with NIBS.",
    ],
    symptoms=[
        ["Excessive Worry",          "Uncontrollable worry about multiple life domains; persistent and pervasive",                "Core criterion",    "Kessler 2005; DSM-5"],
        ["Difficulty Controlling Worry","Inability to stop or redirect worry despite effort",                                    "Core criterion",    "DSM-5; ICD-11"],
        ["Restlessness",             "Feeling keyed up, on edge, or unable to relax",                                           "60\u201380%",       "DSM-5"],
        ["Fatigue",                  "Easily fatigued; reduced energy from chronic arousal burden",                              "50\u201370%",       "DSM-5"],
        ["Concentration Difficulties","Mind going blank; difficulty maintaining focus due to worry intrusion",                   "50\u201370%",       "DSM-5"],
        ["Irritability",             "Heightened irritability and emotional reactivity",                                         "40\u201360%",       "DSM-5"],
        ["Muscle Tension",           "Chronic muscle tension, aches, soreness; somatic symptom burden",                         "60\u201380%",       "DSM-5"],
        ["Sleep Disturbance",        "Difficulty falling or staying asleep; restless unsatisfying sleep",                       "60\u201380%",       "DSM-5; Wittchen 2011"],
        ["Autonomic Symptoms",       "Palpitations, sweating, tremor, GI disturbance, dry mouth",                               "50\u201370%",       "Wittchen 2011"],
        ["Somatic Hypervigilance",   "Excessive attention to bodily sensations; amplification of normal physiological signals", "30\u201350%",       "Barsky 2001"],
    ],
    brain_regions=[
        ["Right DLPFC",              "Hypoactivation; impaired prefrontal control over amygdala threat responses",              "Cathodal tDCS (inhibitory) or bilateral normalisation; Level B evidence",    "Ironside 2019"],
        ["Left DLPFC",               "CEN hypoactivation; impaired cognitive control over worry",                               "Anodal tDCS (excitatory); supports emotion regulation",                       "Brunoni 2016"],
        ["Amygdala",                 "Hyperactivation to threat cues; drives fear and worry circuitry",                         "Modulated indirectly via prefrontal tDCS and taVNS limbic pathways",          "Etkin 2010"],
        ["Anterior Insula",          "Interoceptive hypervigilance; amplifies somatic anxiety signals",                         "Modulated via taVNS (vagal-insular pathway) and tDCS (indirect)",             "Simmons 2006"],
        ["Anterior Cingulate Cortex","dACC hyperactivation; threat monitoring and conflict detection",                          "Bilateral DLPFC tDCS modulates ACC activity; TPS (investigational)",          "Whalen 1998"],
        ["vmPFC / OFC",              "Impaired fear extinction and worry termination",                                           "Prefrontal tDCS supports vmPFC engagement; TPS investigational",              "Hartley 2011"],
        ["Hippocampus",              "Contextual fear; safety learning; reduced volume with chronic stress",                     "Indirectly targeted via TPS and BDNF upregulation pathways",                  "Lissek 2014"],
        ["Basal Ganglia",            "Habit-based worry patterns; striatal engagement in anxiety perseveration",                 "Indirectly modulated via prefrontal-striatal circuits",                       "Pitman 2012"],
    ],
    brainstem=[
        ["Locus Coeruleus (LC)",     "Noradrenergic hyperactivation; drives hyperarousal and anxiety",                          "taVNS modulates LC activity via NTS; primary mechanism for anxiety reduction","Aston-Jones 2005"],
        ["Nucleus Tractus Solitarius","Vagal afferent relay; central autonomic hub; key for HRV",                               "Primary target of taVNS; modulates amygdala, ACC, and prefrontal cortex",    "Frangos 2015"],
        ["Periaqueductal Gray",      "Defensive behaviour; threat-escape processing; chronic anxiety maintenance",               "Modulated indirectly via prefrontal and vagal neuromodulation",               "Canteras 2002"],
        ["Dorsal Raphe Nucleus",     "Serotonergic dysregulation; comorbid anxiety-depression pathway",                         "taVNS and tDCS (DLPFC) may influence serotonergic projections",               "Deakin 1991"],
        ["Parabrachial Nucleus",     "Visceral and somatic threat integration; autonomic fear circuits",                         "Modulated indirectly via vagal and prefrontal neuromodulation",               "Bhatt 2020"],
    ],
    phenotypes=[
        ["Classic GAD",              "Full DSM-5 criteria met; excessive worry across multiple domains \u22656 months",         "Excessive worry, fatigue, muscle tension, sleep disturbance",   "Bilateral DLPFC tDCS; taVNS; CES"],
        ["Anxious MDD (Mixed)",      "Significant depression + anxiety; HDRS anxiety/somatisation subscale elevated",           "Low mood, anxiety, tension, sleep disturbance",                 "L-DLPFC anodal tDCS; taVNS; CES; bilateral approach"],
        ["Health Anxiety (Somatic)", "Persistent preoccupation with having serious illness; somatic amplification",              "Health worry, somatic vigilance, reassurance-seeking",          "Right DLPFC cathodal; taVNS; CES"],
        ["Social Anxiety (SAD)",     "Marked fear or anxiety about social situations; avoidance; performance anxiety",           "Social fear, avoidance, physical symptoms in social contexts",  "Right DLPFC cathodal tDCS + taVNS; CES"],
        ["Panic Disorder",           "Recurrent unexpected panic attacks; anticipatory anxiety; avoidance",                      "Panic attacks, palpitations, derealization, avoidance",         "Bilateral DLPFC tDCS; taVNS primary; CES"],
        ["GAD + Insomnia",           "GAD with prominent insomnia as a maintaining and exacerbating factor",                    "Worry, insomnia, fatigue, daytime impairment",                  "L-DLPFC tDCS; CES primary (sleep); taVNS"],
        ["GAD + Chronic Pain",       "Anxiety maintaining central sensitisation and pain amplification",                        "Worry, somatic tension, pain, sleep disturbance",               "Bilateral tDCS + M1; taVNS; CES"],
    ],
    symptom_map=[
        ["Excessive Worry",          "Right DLPFC, vmPFC, dACC",      "Right DLPFC cathodal tDCS + bilateral tDCS option; TPS (PFC, ACC) investigational",    "Moderate (Ironside 2019)"],
        ["Autonomic Hyperarousal",   "LC, NTS, Amygdala",              "taVNS primary (vagal-NTS-LC); bilateral tDCS supportive; CES adjunct",                "Level A (taVNS: VNS evidence)"],
        ["Sleep Disturbance",        "Frontal, Thalamus, LC",          "CES primary; taVNS adjunctive; tDCS (frontal slow-wave); sleep hygiene",              "Moderate (Philip 2017)"],
        ["Muscle Tension",           "SMA, Motor cortex, Autonomic",   "taVNS primary; CES; M1/SMA tDCS supportive; relaxation training",                    "Moderate (taVNS evidence)"],
        ["Cognitive Anxiety",        "Left DLPFC, ACC",                "L-DLPFC anodal tDCS + CBT; TPS DLPFC investigational",                               "Moderate (Brunoni 2016)"],
        ["Somatic Symptoms",         "Insula, Anterior Cingulate",     "taVNS (vagal-insular pathway) + CES; bilateral tDCS supportive",                      "Moderate (taVNS)"],
        ["Concentration / Attention","Left DLPFC, Parietal",           "L-DLPFC anodal tDCS + cognitive training",                                           "Moderate (Brunoni 2016)"],
        ["Comorbid Depression",      "Left DLPFC, sgACC",              "L-DLPFC anodal tDCS (primary); taVNS + CES; bilateral for mixed presentations",       "Level A (for MDD component)"],
        ["Panic / Avoidance",        "Amygdala, Insula, PFC",          "taVNS primary + bilateral DLPFC tDCS; CES; exposure therapy integration",             "Moderate (taVNS; Brunoni 2016)"],
    ],
    montage=[
        ["Classic GAD",              "F3 anode + F4 cathode (bilateral normalisation); or right F4 cathodal alone","Cranial (global + targeted DLPFC/ACC)", "taVNS primary, CES"],
        ["Anxious MDD",              "F3 anode + Fp2 cathode; or F3 + F4 bilateral anode",                         "Cranial (global + targeted DLPFC)",     "taVNS + CES; L-DLPFC anodal"],
        ["Health Anxiety / Somatic", "F4 cathode + shoulder anode (right DLPFC cathodal)",                         "Cranial (global + targeted right PFC)", "taVNS primary; CES"],
        ["Social Anxiety",           "F4 cathode + F3 anode; bilateral DLPFC approach",                            "Cranial (global + targeted PFC)",       "taVNS; CES; CBT integration"],
        ["Panic Disorder",           "F3 + F4 bilateral anode; or right F4 cathodal",                              "Cranial (global + targeted DLPFC)",     "taVNS primary; CES"],
        ["GAD + Insomnia",           "F3 anode + Fp2 cathode; Fz evening slow-wave protocol",                      "Cranial (global + targeted DLPFC)",     "CES primary (sleep); taVNS"],
        ["GAD + Chronic Pain",       "F3 + F4 bilateral + C3/C4 (contra to pain) option",                          "Cranial + peripheral (pain dermatomes)","taVNS; CES"],
        ["No response after 8\u201310 sessions","Reassess phenotype; consider bilateral vs. unilateral; add TPS","Adjust parameters; TPS add-on",           "Review all adjuncts; CBT referral"],
    ],
    tdcs_protocols=[
        ["C1","GAD \u2014 right DLPFC cathodal (primary)","F3 anode",         "F4 / right shoulder",       "2 mA, 20\u201330 min, 5\u00d7/wk \u00d7 6\u201310 wks","Level B; anxiolytic (Ironside 2019; Brunoni 2016)"],
        ["C2","Bilateral DLPFC normalisation",             "F3 + F4 bilateral anode","Fp2 + Fp1 / shoulders","2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate; bilateral approach (Loo 2012)"],
        ["C3","Anxious MDD \u2014 L-DLPFC anodal",        "F3 anode",         "Fp2 / extracephalic",        "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Level A (MDD component; Brunoni 2016)"],
        ["C4","Panic / arousal reduction",                 "F3 + F4 bilateral","Extracephalic shoulder",     "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Brunoni 2016)"],
        ["C5","Sleep disturbance",                         "Fz anode",         "Pz / extracephalic",         "1\u20132 mA, 20 min, evening",                       "Moderate (Philip 2017)"],
        ["C6","Somatic anxiety / muscle tension",          "F3 + Cz anode",    "F4 / extracephalic",         "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Brunoni 2016)"],
        ["C7","Cognitive anxiety / concentration",         "F3 anode",         "Fp2, extracephalic",         "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Brunoni 2016)"],
        ["C8","Chronic GAD maintenance",                   "F3 + F4 bilateral","Extracephalic shoulder",     "2 mA, 20\u201330 min, 3\u00d7/wk maintenance",      "Moderate (Loo 2012; Brunoni 2016)"],
    ],
    plato_protocols=[
        ["C1-PS","GAD \u2014 primary",              "Frontal (F3 area)","F3","Right shoulder",   "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Combine with taVNS as primary protocol"],
        ["C2-PS","Bilateral normalisation",         "Frontal (F3 area)","F3","Right shoulder",   "1.6 mA, 20\u201330 min","5\u00d7/wk",         "R-DLPFC cathodal not directly achievable; use L anodal"],
        ["C3-PS","Anxious MDD",                     "Frontal (F3 area)","F3","Right shoulder",   "1.6 mA, 20\u201330 min","5\u00d7/wk",         "L-DLPFC anodal primary; adjunct to CES"],
        ["C4-PS","Panic / arousal",                 "Frontal (F3 area)","F3","Right shoulder",   "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Pre-session taVNS recommended"],
        ["C5-PS","Sleep disturbance",               "Frontal (Fz area)","Fz","Shoulder",         "1.6 mA, 20 min",        "Evening",            "Evening use; combine with CES"],
        ["C6-PS","Somatic / muscle tension",        "Frontal (F3 area)","F3","Shoulder",         "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Combine with relaxation training"],
        ["C7-PS","Cognitive anxiety",               "Frontal (F3 area)","F3","Shoulder",         "1.6 mA, 20\u201330 min","5\u00d7/wk",         "CBT session pairing recommended"],
        ["C8-PS","Maintenance GAD",                 "Frontal (F3 area)","F3","Right shoulder",   "1.6 mA, 20\u201330 min","3\u00d7/wk",         "Maintenance block after initial intensive"],
    ],
    tps_protocols=[
        ["T1","GAD \u2014 prefrontal (primary)",    "R-DLPFC + L-DLPFC bilateral, OFC","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz; 10 sessions/2 wks","4,000 Tgt (DLPFC) + 2,000 PRE; total 6\u20138K","Investigational (pilot data)"],
        ["T2","Amygdala / limbic hyperactivation",  "Medial PFC + ACC (closest accessible deep target)","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",           "4,000 Tgt + 2,000 Std; total 6\u20138K",       "Investigational"],
        ["T3","Somatic anxiety / interoceptive",    "Anterior insula (accessible via temporal ROI), ACC","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",          "4,000 Tgt + 2,000 Std; total 6\u20138K",       "Investigational"],
        ["T4","Sleep disturbance",                  "Frontal slow-wave circuits, mPFC",                  "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",          "4,000 Tgt + 2,000 PRE; total 6\u20138K",       "Investigational"],
        ["T5","Panic / autonomic dysregulation",    "DLPFC + ACC",                                       "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",          "4,000 Tgt + 2,000 Std; total 6\u20138K",       "Investigational"],
    ],
    ces_role=[
        ["Anxiety reduction",     "Primary FDA-cleared indication; reduces somatic and cognitive anxiety features",            "Daily or 5\u00d7/wk; 20\u201340 min; morning or as needed"],
        ["Sleep disturbance",     "Enhances sleep onset and quality; key target in GAD maintenance",                          "Evening before bed; 20\u201340 min; coordinate with taVNS"],
        ["Autonomic stabilisation","Stabilises HRV and autonomic tone; complements taVNS vagal modulation",                  "Before tDCS session or standalone; 20\u201330 min"],
        ["Muscle tension",        "Reduces somatic muscle tension and physical anxiety symptoms",                              "As needed; morning or pre-session; 20\u201330 min"],
    ],
    tavns_role="taVNS is a primary adjunct in GAD, providing autonomic downregulation via the vagal-NTS-LC pathway, reducing amygdala hyperactivity, improving HRV, and attenuating the physiological substrate of chronic anxiety.",
    combinations=[
        ("1) Classic GAD", [
            ["tDCS + CBT","Bilateral DLPFC tDCS + CBT; addresses cognitive and neural substrates of excessive worry (Brunoni 2016; Fasotti 2020)","tDCS before CBT session","Excessive worry, avoidance, somatic tension"],
            ["taVNS + tDCS","taVNS priming before tDCS; autonomic downshift before cortical stimulation (Hein 2021; Redgrave 2018)","taVNS 20\u201330 min before tDCS","Autonomic hyperarousal, somatic anxiety"],
            ["CES + Standard Care","CES primary FDA-cleared approach for anxiety; daily use (Philip 2017)","Daily; morning and/or evening as prescribed","Anxiety, sleep disturbance, somatic tension"],
        ]),
        ("2) Anxious MDD", [
            ["tDCS + Psychotherapy","L-DLPFC anodal tDCS + CBT or behavioural activation; targets both MDD and GAD components (Brunoni 2016)","tDCS before therapy","Low mood + excessive worry + tension"],
            ["taVNS + tDCS","taVNS reduces physiological anxiety before tDCS session; additive for mixed anxiety-depression (Hein 2021)","taVNS before tDCS","Autonomic hyperarousal with depressive features"],
            ["CES + Standard Care","CES for anxiety and sleep; FDA-cleared; complements tDCS treatment block (Philip 2017)","Evening and/or morning","Anxious MDD with sleep disturbance"],
        ]),
        ("3) Panic Disorder", [
            ["tDCS + Exposure Therapy","Bilateral DLPFC tDCS before/during graded exposure; enhances fear extinction via prefrontal facilitation (Herrmann 2017)","tDCS before exposure session","Panic attacks, avoidance, anticipatory anxiety"],
            ["taVNS + tDCS","taVNS reduces interoceptive threat amplification before tDCS; key for panic (Hein 2021)","taVNS 20\u201330 min before tDCS","Panic with autonomic arousal and somatic features"],
            ["CES + Standard Care","CES daily for baseline anxiety reduction; improves between-session state for exposure work (Philip 2017)","Daily; morning and/or evening","Panic disorder with chronic anticipatory anxiety"],
        ]),
        ("4) Health Anxiety / Somatic", [
            ["tDCS + ACT / Mindfulness","Right DLPFC cathodal tDCS + ACT or mindfulness; addresses somatic hypervigilance (Brunoni 2016)","tDCS before mindfulness/ACT session","Health worry, somatic amplification, reassurance-seeking"],
            ["taVNS + Somatic Relaxation","taVNS + progressive muscle relaxation; reduces interoceptive amplification (Hein 2021)","taVNS concurrent with relaxation practice","Muscle tension, somatic symptoms, physical anxiety"],
            ["CES + Standard Care","CES for anxiety and somatic tension; FDA-cleared (Philip 2017)","Daily as prescribed","Chronic somatic anxiety, sleep disturbance"],
        ]),
        ("5) GAD + Insomnia", [
            ["tDCS + Sleep Therapy (CBT-I)","DLPFC tDCS + CBT for insomnia; targets both anxiety and sleep maintenance (Brunoni 2016; Edinger 2021)","tDCS before CBT-I session; evening slow-wave protocol","GAD with severe insomnia as maintaining factor"],
            ["taVNS + CES","Dual adjunct: taVNS for autonomic downshift + CES for sleep; maximum non-pharmacological sleep support (Hein 2021; Philip 2017)","taVNS 20 min before bed; CES at bedtime","Sleep-onset difficulty, nocturnal arousal"],
            ["CES Primary + tDCS","CES as primary sleep and anxiety modality; tDCS supportive for cognitive anxiety (Philip 2017)","CES evening primary; tDCS morning 5\u00d7/wk","GAD where insomnia is most impairing feature"],
        ]),
        ("6) GAD + Chronic Pain", [
            ["tDCS (DLPFC + M1) + Graded Activity","Bilateral DLPFC + M1 cathodal tDCS with graded activity; anxiety-pain circuit modulation (Fregni 2006; Brunoni 2016)","tDCS before graded activity","GAD maintaining central sensitisation and pain"],
            ["taVNS + Standard Care","taVNS modulates autonomic arousal and ascending pain pathways; key for anxiety-pain overlap (Adair 2020)","taVNS before tDCS or activity","Autonomic hyperarousal with somatic pain"],
            ["CES + Standard Care","CES reduces anxiety and sleep disturbance maintaining pain; supportive throughout (Philip 2017)","Daily as prescribed","GAD + chronic pain with sleep and anxiety amplification"],
        ]),
    ],
    combination_summary=[
        ["Classic GAD",         "tDCS + CBT",           "Bilateral DLPFC tDCS + CBT; addresses cognitive and neural worry substrate (Brunoni 2016)","tDCS before CBT","Excessive worry, avoidance, somatic tension","Moderate"],
        ["Classic GAD",         "taVNS + tDCS",         "taVNS autonomic priming before tDCS; dual-pathway engagement (Hein 2021)","taVNS 20\u201330 min before tDCS","Autonomic hyperarousal","Emerging"],
        ["Anxious MDD",         "tDCS + Psychotherapy", "L-DLPFC anodal tDCS + CBT/BA; addresses both MDD and GAD components (Brunoni 2016)","tDCS before therapy","Low mood + excessive worry","Moderate"],
        ["Panic Disorder",      "tDCS + Exposure Therapy","Bilateral DLPFC tDCS before graded exposure; enhances fear extinction (Herrmann 2017)","tDCS before exposure session","Panic attacks, avoidance","Moderate"],
        ["GAD + Insomnia",      "taVNS + CES",           "Dual adjunct for sleep and autonomic downregulation (Hein 2021; Philip 2017)","taVNS + CES at bedtime","GAD with insomnia as maintaining factor","Emerging"],
        ["GAD + Chronic Pain",  "tDCS (DLPFC+M1) + Activity","Anxiety-pain circuit modulation via bilateral tDCS + activity (Fregni 2006; Brunoni 2016)","tDCS before activity","GAD maintaining central sensitisation","Limited\u2013Moderate"],
    ],
    outcomes=[
        ["GAD-7",             "Anxiety severity (self-report)",         "Baseline, weeks 2, 4, 8, 12",  "\u226510 moderate; \u226515 severe; \u22655-point change = response"],
        ["HAMA",              "Clinician-rated anxiety severity",        "Baseline, weeks 4, 8, 12",     "Score \u226417 mild; 18\u201324 moderate; \u226525 severe"],
        ["PHQ-9",             "Comorbid depressive symptoms",           "Baseline, weeks 4, 8, 12",     "\u226510 moderate; \u226515 severe"],
        ["ISI",               "Insomnia severity",                      "Baseline, weeks 4, 8, 12",     "Score 0\u20137 no clinically significant insomnia; \u226715 moderate insomnia"],
        ["PSQI",              "Sleep quality",                          "Baseline, months 1, 3",        "Score >5 poor sleep quality"],
        ["STAI-State/Trait",  "State and trait anxiety dimensions",     "Baseline, months 1, 3",        "Normative values by age/gender; track change from baseline"],
        ["CGI-S / CGI-I",     "Clinical Global Impression",             "Baseline, weeks 4, 8, 12",     "CGI-I \u22642 = much improved"],
        ["WHO-5",             "Wellbeing and quality of life",          "Baseline, months 1, 3",        "Score \u226452 poor wellbeing (max 100)"],
        ["SOZO PRS",          "NIBS-specific functional outcome",       "Baseline, weeks 2, 4, 8, 12", "Proprietary; composite anxiety/sleep/functional domains"],
    ],
    outcomes_abbrev=["GAD-7", "HAMA", "PHQ-9", "ISI", "SOZO PRS"],
)

# ══════════════════════════════════════
# 4. ADHD
# ══════════════════════════════════════
CONDITIONS["adhd"] = dict(
    name="Attention Deficit Hyperactivity Disorder", icd10="F90",
    doc_num="SOZO-ADHD-FEL-004",
    tps_status="INVESTIGATIONAL / OFF-LABEL",
    tdcs_class="Level B evidence; CE-marked",
    tps_class="INVESTIGATIONAL / OFF-LABEL",
    tavns_class="Emerging adjunctive; CE-marked",
    ces_class="FDA-cleared adjunctive for mood/anxiety/sleep",
    inclusion=[
        "Confirmed ADHD diagnosis (DSM-5 or ICD-11) by qualified clinician; all subtypes eligible",
        "ADHD symptoms causing functional impairment (CAARS or BRIEF \u22651.5 SD above norm)",
        "Age 18+ years (adult ADHD protocol)",
        "Stable or no ADHD medication (no change in 4 weeks prior to enrolment)",
        "Written informed consent (including TPS off-label disclosure)",
        "Baseline assessments completed (CAARS, BRIEF-2, MoCA, PHQ-9)",
    ],
    exclusion=[
        "Intracranial metallic hardware in stimulation path",
        "Cochlear implant or DBS device",
        "Active psychotic features",
        "Skull defects, craniectomy, or recent craniotomy",
        "Pregnancy (tDCS, TPS)",
        "Inability to provide informed consent",
        "Intellectual disability (IQ <70) as primary diagnosis",
    ],
    discussion=[
        "Cardiac pacemaker or defibrillator \u2014 individual risk\u2013benefit assessment",
        "Epilepsy or seizure history \u2014 formal risk\u2013benefit with documented rationale",
        "Stimulant medication use \u2014 may interact with tDCS cortical excitability; timing consideration",
        "Comorbid bipolar disorder \u2014 mood stabiliser coverage required",
        "Comorbid ASD \u2014 requires adapted protocol and slower titration",
        "Coagulation disorders or anticoagulants (especially TPS)",
        "Dermatological conditions at electrode sites",
        "Substance use disorder comorbidity",
    ],
    overview=[
        "Attention Deficit Hyperactivity Disorder (ADHD) is a neurodevelopmental disorder characterised by "
        "pervasive patterns of inattention, hyperactivity, and/or impulsivity that impair functioning across "
        "multiple settings. Adult ADHD has a prevalence of 2.5\u20135% globally. The majority of children with "
        "ADHD continue to experience clinically significant symptoms into adulthood (Faraone et al., 2021).",
        "ADHD involves dysregulation of the prefrontal-striatal-cerebellar circuit, with core deficits in "
        "executive function, working memory, inhibitory control, and attentional regulation. The dopaminergic "
        "and noradrenergic systems are the primary pharmacological targets. Neuroimaging demonstrates consistent "
        "DLPFC, ACC, and inferior frontal gyrus (IFG) hypoactivation in ADHD.",
        "tDCS targeting the left DLPFC (anodal) enhances prefrontal excitability and has shown improvements "
        "in working memory, inhibitory control, and attention in ADHD across multiple studies (Westwood et al., "
        "2021). The right IFG (a key response inhibition node) is an emerging cathodal target.",
        "NIBS Evidence in ADHD: Meta-analyses support modest but consistent tDCS effects on working memory "
        "and inhibitory control in ADHD (Westwood et al., 2021). TPS targeting prefrontal circuits is "
        "investigational. taVNS has plausible mechanisms via noradrenergic modulation relevant to ADHD.",
    ],
    pathophysiology=[
        "ADHD pathophysiology involves hypofunction of the mesocortical dopaminergic pathway (VTA\u2192PFC) "
        "and locus coeruleus noradrenergic projections, resulting in DLPFC, IFG, and ACC hypoactivation. "
        "This impairs the top-down executive control network (CEN). The default mode network (DMN) fails "
        "to suppress adequately during task demands (DMN suppression failure), generating internal "
        "distraction and mind-wandering.",
        "Cerebellar-cortical timing circuits are implicated in temporal processing deficits and motor "
        "impulsivity in ADHD. Reward circuitry (ventral striatum) shows altered delay discounting, "
        "contributing to motivational dysregulation. Structural MRI shows reduced cortical thickness "
        "and volume in prefrontal regions.",
    ],
    std_treatment=[
        "First-line pharmacotherapy: methylphenidate (ADHD-specific dopamine-noradrenaline reuptake "
        "inhibitor) and amphetamine derivatives (lisdexamfetamine) for attention and executive function. "
        "Atomoxetine (non-stimulant SNRI) for comorbid anxiety or stimulant contraindication. "
        "Guanfacine and clonidine (alpha-2 agonists) for hyperactivity/impulsivity. CBT is recommended "
        "in combination with pharmacotherapy for adult ADHD. NIBS is adjunctive.",
    ],
    symptoms=[
        ["Inattention",           "Sustained difficulty maintaining attention; careless mistakes; easily distracted",        "Core criterion", "Faraone 2021; DSM-5"],
        ["Hyperactivity",         "Restlessness, fidgeting, inability to stay seated; excessive talking",                    "Core criterion", "DSM-5; Barkley 2015"],
        ["Impulsivity",           "Interrupting others; blurting out answers; difficulty waiting turn; poor inhibition",     "Core criterion", "DSM-5"],
        ["Working Memory Deficit","Poor retention and manipulation of information in short-term memory",                     "80\u201390%",    "Barkley 1997"],
        ["Executive Dysfunction", "Poor planning, organisation, time management, cognitive flexibility",                     "70\u201390%",    "Brown 2006"],
        ["Emotional Dysregulation","Low frustration tolerance; emotional reactivity; mood lability",                         "50\u201370%",    "Barkley 2015"],
        ["Motivation Dysregulation","Difficulty initiating tasks; procrastination; variable motivation depending on interest","60\u201380%",    "Volkow 2009"],
        ["Sleep Disturbance",     "Delayed sleep phase; insomnia; non-restorative sleep; circadian dysregulation",           "50\u201370%",    "Becker 2019"],
        ["Anxiety Comorbidity",   "Comorbid anxiety in 30\u201350% of adults with ADHD",                                    "30\u201350%",    "Kessler 2006"],
        ["Substance Use Risk",    "Elevated risk of substance use disorders (self-medication of ADHD symptoms)",              "20\u201335%",    "Charach 2011"],
    ],
    brain_regions=[
        ["Left DLPFC (F3)",         "Hypoactivation; core executive function, working memory, top-down attentional control",  "Anodal tDCS primary target; enhances executive excitability (Level B)",     "Westwood 2021"],
        ["Right IFG (F4/F8)",       "Response inhibition; go/no-go processing; motor impulse control",                        "Cathodal tDCS (R-IFG) or bilateral; targets inhibitory deficits",            "Aron 2004"],
        ["Anterior Cingulate Cortex","Conflict monitoring; error detection; executive attention; DMN suppression",             "DLPFC tDCS indirectly modulates ACC; TPS (investigational)",                 "Barch 2001"],
        ["Striatum (Caudate/Putamen)","Dopaminergic reward; motor control; habit formation",                                   "Modulated via prefrontal-striatal circuits by tDCS and TPS",                 "Volkow 2009"],
        ["Cerebellum",              "Timing circuits; motor impulsivity; motor coordination and sequencing",                   "Cerebellar tDCS has limited ADHD evidence; TPS (investigational)",           "Stoodley 2014"],
        ["Insula",                  "Interoceptive awareness; impulsivity; self-monitoring deficit in ADHD",                   "Indirectly modulated via prefrontal tDCS and taVNS (vagal-insular)",         "Wittmann 2010"],
        ["vmPFC / OFC",             "Delay discounting; reward valuation; motivational regulation",                            "Prefrontal tDCS modulates vmPFC via fronto-limbic connections",               "Barkley 2001"],
        ["Thalamus",                "Sensory gating; arousal regulation; thalamocortical synchrony in attention",              "TPS investigational deep target; tDCS effects thalamus indirectly",          "Satterthwaite 2009"],
    ],
    brainstem=[
        ["Locus Coeruleus (LC)",    "Noradrenergic arousal and attentional modulation; key atomoxetine target",               "taVNS modulates LC via NTS; augments noradrenergic tone for attention",     "Aston-Jones 2005"],
        ["VTA (Dopaminergic)",      "Mesocortical dopamine; motivation, reward, prefrontal activation",                       "Indirectly modulated via prefrontal tDCS and mesolimbic pathway",            "Volkow 2009"],
        ["Nucleus Tractus Solitarius","Vagal afferent relay; central autonomic hub; modulates cortical arousal",              "Primary target of taVNS; affects LC and prefrontal circuits",               "Frangos 2015"],
        ["Periaqueductal Gray",     "Arousal; defensive behaviour; motor inhibition circuits",                                 "Modulated indirectly via prefrontal and vagal neuromodulation",              "Bari 2013"],
        ["Reticular Activating System","General arousal; thalamocortical gating; attention maintenance",                      "taVNS and tDCS indirectly modulate ascending arousal systems",               "Moruzzi 1949"],
    ],
    phenotypes=[
        ["ADHD \u2014 Predominantly Inattentive", "Inattention dominant; minimal hyperactivity; often under-diagnosed in adults", "Concentration, working memory, organisation, task-completion failures","L-DLPFC anodal tDCS; cognitive training; taVNS"],
        ["ADHD \u2014 Combined Type",             "Both inattention and hyperactivity-impulsivity meeting full criteria",         "Inattention + impulsivity + hyperactivity + executive dysfunction",     "L-DLPFC anodal + R-IFG cathodal tDCS; taVNS; CES"],
        ["ADHD \u2014 Hyperactive-Impulsive",     "Hyperactivity and impulsivity dominant; inattention subthreshold",            "Impulsivity, hyperactivity, poor frustration tolerance",                "R-IFG cathodal tDCS; bilateral DLPFC; taVNS"],
        ["ADHD + Executive Function Profile",     "Severe EF deficits beyond core ADHD criteria",                               "Planning, organisation, time management, working memory failures",      "L-DLPFC anodal tDCS; TPS DLPFC; cognitive training"],
        ["ADHD + Anxiety",                        "ADHD with significant comorbid anxiety",                                      "Inattention + worry + tension + sleep disturbance",                     "Bilateral DLPFC; taVNS primary; CES; slower titration"],
        ["ADHD + Depression",                     "ADHD with comorbid MDD",                                                      "Inattention + low mood + anergia + cognitive slowing",                  "L-DLPFC anodal tDCS; CES; taVNS; treat MDD component"],
        ["ADHD + Sleep Disorder",                 "ADHD with circadian dysregulation and insomnia as maintaining factors",       "Inattention + delayed sleep phase + daytime fatigue",                   "DLPFC tDCS; CES primary (sleep); taVNS"],
    ],
    symptom_map=[
        ["Inattention / Sustained Attention","Left DLPFC, ACC",        "L-DLPFC anodal tDCS + cognitive training (attention tasks)",                      "Level B (Westwood 2021)"],
        ["Working Memory Deficit",            "Left DLPFC, Parietal",   "L-DLPFC anodal tDCS + n-back/working memory training; TPS DLPFC investigational", "Level B (Westwood 2021)"],
        ["Response Inhibition / Impulsivity", "Right IFG, ACC",         "R-IFG cathodal tDCS (F4 cathode) + inhibitory training; TPS IFG investigational", "Moderate (Aron 2004; Lefaucheur 2017)"],
        ["Executive Dysfunction",             "Left DLPFC, ACC, SMA",   "L-DLPFC anodal tDCS + executive function training; TPS DLPFC investigational",    "Level B (Westwood 2021)"],
        ["Emotional Dysregulation",           "Left DLPFC, Amygdala",   "L-DLPFC tDCS + taVNS for limbic modulation; CES for baseline arousal",           "Moderate (Brunoni 2016)"],
        ["Motivation / Reward",               "OFC, Striatum, vmPFC",   "L-DLPFC tDCS + motivational interventions; taVNS for noradrenergic enhancement",  "Limited (Volkow 2009)"],
        ["Sleep Disturbance",                 "Frontal, Thalamus",      "CES primary (evening); taVNS; tDCS evening slow-wave protocol",                   "Moderate (Philip 2017)"],
        ["Anxiety Comorbidity",               "Right DLPFC, Amygdala",  "Bilateral DLPFC tDCS; taVNS primary; CES; address anxiety component explicitly",  "Moderate (Brunoni 2016)"],
        ["Hyperactivity / Arousal",           "Right IFG, SMA, ACC",    "R-IFG cathodal tDCS + mindfulness/behavioural inhibition training; taVNS",        "Moderate (Aron 2004)"],
    ],
    montage=[
        ["Inattentive ADHD",          "F3 anode + Fp2 cathode (L-DLPFC primary)",                                     "Cranial (global + targeted L-DLPFC)", "taVNS, CES (sleep)"],
        ["Combined ADHD",             "F3 anode + F4 cathode (bilateral asymmetric); 2 mA, 20\u201330 min",           "Cranial (global + targeted DLPFC/IFG)","taVNS; CES"],
        ["Hyperactive-Impulsive",     "F4 cathode + F3 anode (R-IFG cathodal emphasis); or F3 anode + shoulder cathode","Cranial (global + targeted R-IFG)", "taVNS primary; CES"],
        ["ADHD + Executive Function", "F3 anode + T3/T4 anodal (add parietal); Fz option",                           "Cranial (global + targeted DLPFC/PPC)","taVNS; CES (sleep)"],
        ["ADHD + Anxiety",            "Bilateral DLPFC (F3 + F4) or F3 anode + F4 cathode",                           "Cranial (global + targeted PFC)",       "taVNS primary; CES"],
        ["ADHD + Depression",         "F3 anode + Fp2 cathode (L-DLPFC anodal primary)",                              "Cranial (global + targeted L-DLPFC)",  "CES; taVNS"],
        ["ADHD + Sleep",              "F3 anode + Fp2 cathode; Fz evening slow-wave protocol",                         "Cranial (global + targeted)",           "CES primary (sleep); taVNS"],
        ["No response after 8\u201310 sessions","Reassess medication status; consider bilateral or cathodal IFG addition","Adjust parameters; TPS add-on",   "Review all adjuncts; CBT referral"],
    ],
    tdcs_protocols=[
        ["C1","Inattention / working memory (primary)","F3 anode",       "Fp2 / F4 / extracephalic",      "2 mA, 20\u201330 min, 5\u00d7/wk \u00d7 6\u201310 wks","Level B; WM improvement (Westwood 2021)"],
        ["C2","Response inhibition / impulsivity",      "F3 anode + F4 cathode","Extracephalic shoulder", "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Aron 2004; Lefaucheur 2017)"],
        ["C3","Combined ADHD \u2014 bilateral",         "F3 + F4 bilateral anode","Fp1 + Fp2 / shoulders","2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Westwood 2021)"],
        ["C4","Executive function / organisation",      "F3 anode + T3/T4 anode","F4 / extracephalic",    "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Brunoni 2016; Westwood 2021)"],
        ["C5","Emotional dysregulation",               "F3 anode",        "Fp2, extracephalic",           "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Brunoni 2016)"],
        ["C6","ADHD + anxiety overlay",                "F3 + F4 bilateral","Extracephalic shoulders",     "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Brunoni 2016)"],
        ["C7","Sleep disturbance",                     "Fz anode",        "Pz / extracephalic",           "1\u20132 mA, 20 min, evening",                       "Moderate (Philip 2017)"],
        ["C8","Maintenance block",                     "F3 anode",        "Fp2 / extracephalic",          "2 mA, 20\u201330 min, 3\u00d7/wk maintenance",      "Moderate (Westwood 2021)"],
    ],
    plato_protocols=[
        ["C1-PS","Inattention / WM",      "Frontal (F3 area)","F3","Right shoulder","1.6 mA, 20\u201330 min","5\u00d7/wk",         "Pair with attention/WM tasks"],
        ["C2-PS","Impulsivity",           "Frontal (F3 area)","F3","Right shoulder","1.6 mA, 20\u201330 min","5\u00d7/wk",         "Combine with inhibitory training"],
        ["C3-PS","Combined ADHD",         "Frontal (F3 area)","F3","Right shoulder","1.6 mA, 20\u201330 min","5\u00d7/wk",         "After HDCkit induction phase"],
        ["C4-PS","Executive function",    "Frontal (F3 area)","F3","Shoulder",       "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Pair with organisation tasks"],
        ["C5-PS","Emotional dysregulation","Frontal (F3 area)","F3","Shoulder",      "1.6 mA, 20\u201330 min","5\u00d7/wk",         "Pre-session taVNS recommended"],
        ["C6-PS","ADHD + anxiety",        "Frontal (F3 area)","F3","Right shoulder","1.6 mA, 20\u201330 min","5\u00d7/wk",         "Combine with CES and taVNS"],
        ["C7-PS","Sleep disturbance",     "Frontal (Fz area)","Fz","Shoulder",       "1.6 mA, 20 min",        "Evening",            "CES concurrent recommended"],
        ["C8-PS","Maintenance",           "Frontal (F3 area)","F3","Shoulder",       "1.6 mA, 20\u201330 min","3\u00d7/wk",         "After initial intensive block"],
    ],
    tps_protocols=[
        ["T1","Inattention / WM (primary)","L-DLPFC, R-DLPFC bilateral, ACC",    "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz; 10 sessions/2 wks","4,000 Tgt (DLPFC) + 2,000 PRE; total 6\u20138K","Investigational"],
        ["T2","Response inhibition",        "R-IFG (pars triangularis/opercularis), R-DLPFC","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz","4,000 Tgt (IFG) + 2,000 Std; total 6\u20138K",  "Investigational"],
        ["T3","Executive dysfunction",      "L-DLPFC + ACC",                       "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                 "4,000 Tgt + 2,000 PRE; total 6\u20138K",       "Investigational"],
        ["T4","Emotional dysregulation",    "L-DLPFC + ACC + OFC",                 "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                 "4,000 Tgt + 2,000 Std; total 6\u20138K",       "Investigational"],
        ["T5","ADHD + anxiety/sleep",       "DLPFC + mPFC",                        "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                 "4,000 Tgt + 2,000 PRE; total 6\u20138K",       "Investigational"],
    ],
    ces_role=[
        ["Sleep disturbance",     "Targets delayed sleep phase and insomnia common in ADHD; FDA-cleared",                     "Evening before bed; 20\u201340 min; coordinate with sleep hygiene"],
        ["Anxiety comorbidity",   "Reduces comorbid anxiety in ADHD; FDA-cleared for anxiety",                                "Before tDCS session or evening; 20\u201340 min"],
        ["Emotional dysregulation","Stabilises baseline arousal and emotional reactivity",                                    "As needed; pre-session or morning; 20\u201330 min"],
        ["Motivation / anergia",  "May improve energy and motivation as adjunct",                                             "Morning; 20\u201330 min; avoid late afternoon"],
    ],
    tavns_role="taVNS augments noradrenergic tone via locus coeruleus, directly relevant to ADHD catecholaminergic pathophysiology; supports attentional regulation and reduces emotional/autonomic dysregulation.",
    combinations=[
        ("1) Inattentive ADHD", [
            ["tDCS + Cognitive Training","L-DLPFC anodal tDCS + attention/WM computerised training; most supported combination (Westwood 2021)","tDCS before/during cognitive tasks","Inattention, WM deficit, task-completion failures"],
            ["taVNS + tDCS","taVNS priming enhances noradrenergic prefrontal tone before tDCS (Hein 2021; Aston-Jones 2005)","taVNS 20 min before tDCS","Attentional fatigue, motivation dysregulation"],
            ["CES + Standard Care","CES for sleep and baseline arousal stability; supports daytime attention (Philip 2017)","Evening (sleep) and morning (energy) as prescribed","Delayed sleep phase, fatigue-driven inattention"],
        ]),
        ("2) Combined ADHD", [
            ["tDCS + Behavioural Training","Bilateral DLPFC tDCS + behavioural ADHD training; addresses both attention and impulsivity (Westwood 2021; Lefaucheur 2017)","tDCS before/during training session","Inattention + impulsivity + executive dysfunction"],
            ["taVNS + tDCS","Dual catecholaminergic and cortical modulation; taVNS primes noradrenergic system (Hein 2021)","taVNS before tDCS","Emotional dysregulation, impulsivity, arousal variability"],
            ["CES + Standard Care","CES for sleep, anxiety, and arousal stabilisation (Philip 2017)","Daily as prescribed","Sleep disturbance, anxiety comorbidity, emotional dysregulation"],
        ]),
        ("3) ADHD + Executive Dysfunction", [
            ["tDCS + Executive Training","L-DLPFC anodal tDCS + structured executive function training (planning, WM, flexibility) (Westwood 2021)","tDCS before executive training","Planning, organisation, time management failures"],
            ["taVNS + Cognitive Training","taVNS + cognitive flexibility tasks; LC-NE modulation augments prefrontal plasticity (Hein 2021)","taVNS before cognitive training","Cognitive flexibility, task-switching, WM deficits"],
            ["CES + Standard Care","CES for sleep; indirect improvement in executive performance through better sleep (Philip 2017)","Evening","Executive dysfunction with sleep disorder"],
        ]),
        ("4) ADHD + Anxiety", [
            ["tDCS + CBT","Bilateral DLPFC tDCS + CBT; addresses both ADHD executive deficits and anxiety cognitive distortions (Brunoni 2016)","tDCS before CBT session","Inattention + worry + tension"],
            ["taVNS + tDCS","taVNS reduces physiological anxiety before tDCS; critical when anxiety limits tDCS engagement (Hein 2021)","taVNS 20\u201330 min before tDCS","Autonomic hyperarousal, physiological anxiety"],
            ["CES + Standard Care","CES primary for anxiety component; FDA-cleared (Philip 2017)","Daily as prescribed","ADHD with anxiety, sleep disturbance, somatic tension"],
        ]),
        ("5) ADHD + Sleep Disorder", [
            ["tDCS + Sleep Hygiene","DLPFC tDCS + CBT-I or sleep hygiene programme; targets arousal regulation (Westwood 2021)","tDCS morning; evening slow-wave protocol if used","ADHD with delayed sleep phase, insomnia"],
            ["taVNS + CES","Dual adjunct for sleep: taVNS autonomic downshift + CES sleep facilitation (Hein 2021; Philip 2017)","taVNS + CES at bedtime","Sleep-onset difficulties, nocturnal hyperarousal"],
            ["CES Primary + tDCS","CES as primary sleep modality with morning tDCS for daytime attention (Philip 2017)","CES evening; tDCS morning","ADHD where sleep is most impairing feature"],
        ]),
        ("6) ADHD + Emotional Dysregulation", [
            ["tDCS + DBT / Mindfulness","L-DLPFC tDCS + DBT skills or mindfulness; enhances prefrontal regulation of limbic reactivity (Brunoni 2016)","tDCS before therapy session","Emotional lability, frustration, impulsive reactions"],
            ["taVNS + tDCS","taVNS modulates autonomic and limbic reactivity; additive with prefrontal tDCS (Hein 2021)","taVNS before tDCS","Autonomic arousal with emotional dysregulation"],
            ["CES + Standard Care","CES stabilises baseline arousal; reduces emotional reactivity threshold (Philip 2017)","As needed; morning or pre-session","Emotional dysregulation with anxiety and sleep disturbance"],
        ]),
    ],
    combination_summary=[
        ["Inattentive ADHD",      "tDCS + Cognitive Training",   "L-DLPFC anodal tDCS + attention/WM training (Westwood 2021)",                        "tDCS before/during tasks","Inattention, WM deficit",           "Level B"],
        ["Combined ADHD",         "tDCS + Behavioural Training", "Bilateral DLPFC tDCS + ADHD behavioural training (Westwood 2021; Lefaucheur 2017)",   "tDCS before training",    "Inattention + impulsivity",          "Moderate"],
        ["ADHD + EF",             "tDCS + Executive Training",   "L-DLPFC tDCS + structured EF training; planning and WM (Westwood 2021)",              "tDCS before EF training", "Planning, organisation, WM",        "Level B"],
        ["ADHD + Anxiety",        "tDCS + CBT",                  "Bilateral DLPFC tDCS + CBT; addresses ADHD + anxiety (Brunoni 2016)",                  "tDCS before CBT",         "Inattention + worry + tension",      "Moderate"],
        ["ADHD + Sleep",          "taVNS + CES",                 "Dual adjunct for sleep and autonomic stabilisation (Hein 2021; Philip 2017)",          "taVNS + CES at bedtime",  "Delayed sleep phase, insomnia",      "Emerging"],
        ["ADHD + Emotional Dysreg","tDCS + DBT / Mindfulness",   "L-DLPFC tDCS + DBT/mindfulness; prefrontal limbic regulation (Brunoni 2016)",          "tDCS before therapy",     "Emotional lability, impulsivity",    "Moderate"],
    ],
    outcomes=[
        ["CAARS (Self-Report)",   "ADHD symptom severity",                 "Baseline, weeks 4, 8, 12",     "T-score \u226565 clinically significant; track change from baseline"],
        ["BRIEF-2 (Adult)",       "Executive function (self/observer)",    "Baseline, months 1, 3",        "T-score \u226565 indicates EF impairment; track improvement"],
        ["MoCA",                  "Global cognition screening",            "Baseline, month 3",            "Score <26 cognitive impairment"],
        ["PHQ-9",                 "Comorbid depressive symptoms",          "Baseline, weeks 4, 8, 12",     "\u226510 moderate; \u226515 severe"],
        ["GAD-7",                 "Comorbid anxiety",                      "Baseline, weeks 4, 8, 12",     "\u226510 moderate; \u226515 severe"],
        ["ISI",                   "Insomnia severity",                     "Baseline, weeks 4, 8, 12",     "Score \u226715 moderate insomnia"],
        ["Conners CPT-3",         "Sustained attention (objective)",       "Baseline, month 3",            "T-score >65 on commission/omission errors; track change"],
        ["CGI-S / CGI-I",         "Clinical Global Impression",            "Baseline, weeks 4, 8, 12",     "CGI-I \u22642 = much improved"],
        ["SOZO PRS",              "NIBS-specific functional outcome",      "Baseline, weeks 2, 4, 8, 12", "Proprietary; composite attention/EF/emotional/sleep domains"],
    ],
    outcomes_abbrev=["CAARS", "BRIEF-2", "MoCA", "PHQ-9", "SOZO PRS"],
)

# ══════════════════════════════════════
# 5. ALZHEIMER'S DISEASE / MCI
# ══════════════════════════════════════
CONDITIONS["alzheimers"] = dict(
    name="Alzheimer's Disease / MCI", icd10="G30 / F06.7",
    doc_num="SOZO-ALZ-FEL-005",
    tps_status="CE-MARKED FOR ALZHEIMER'S (NEUROLITH) — primary approved indication",
    tdcs_class="Level B evidence; CE-marked",
    tps_class="CE-MARKED \u2014 Primary approved TPS indication",
    tavns_class="Emerging adjunctive; CE-marked",
    ces_class="Supportive adjunct for mood/sleep/anxiety",
    inclusion=[
        "Confirmed diagnosis of MCI or mild-to-moderate Alzheimer's disease (NIA-AA criteria)",
        "MMSE 15\u201326 or MoCA 10\u201325 at baseline",
        "Age 55+ years",
        "Caregiver available for consent support and session attendance if required",
        "Written informed consent (patient and/or legally authorised representative)",
        "Baseline assessments completed (MoCA, MMSE, ADAS-Cog, CDR, NPI)",
    ],
    exclusion=[
        "Intracranial metallic hardware in stimulation path",
        "Cochlear implant or DBS device",
        "Severe dementia (MMSE <10 or CDR \u22653)",
        "Skull defects, craniectomy, or recent craniotomy",
        "Active intracranial tumour",
        "Pregnancy (tDCS, TPS)",
        "Inability to provide any form of informed consent or assent",
    ],
    discussion=[
        "Cardiac pacemaker or defibrillator \u2014 individual risk\u2013benefit assessment",
        "Epilepsy or seizure history \u2014 increased risk in late-stage AD; formal risk\u2013benefit required",
        "Severe behavioural/neuropsychiatric symptoms (agitation, psychosis) \u2014 stabilise first",
        "Anticoagulant use (especially for TPS application)",
        "Dermatological conditions at electrode sites",
        "Significant vascular burden on MRI \u2014 may alter stimulation distribution",
        "Cholinesterase inhibitor or memantine use \u2014 note timing and potential interaction",
        "Moderate-to-severe hearing or vision impairment limiting engagement",
    ],
    overview=[
        "Alzheimer's disease (AD) is the most common cause of dementia, accounting for 60\u201380% of cases "
        "worldwide. Over 55 million people live with dementia globally, with AD representing the greatest "
        "burden. MCI (Mild Cognitive Impairment) represents a transitional state between normal ageing and "
        "dementia, with amnestic MCI carrying the highest AD conversion risk (Petersen et al., 2014).",
        "AD is characterised by progressive accumulation of amyloid-beta plaques and neurofibrillary tangles "
        "(tau hyperphosphorylation), leading to synaptic loss, neuroinflammation, and neuronal death. "
        "The entorhinal cortex and hippocampus are affected earliest; pathology then spreads to "
        "temporoparietal and frontal association cortices (Braak & Braak, 1991).",
        "TPS (NEUROLITH) is CE-marked for Alzheimer's disease \u2014 this is the primary approved clinical "
        "indication for TPS. Multiple pilot studies and early clinical trials have demonstrated improvements "
        "in cognitive function, memory, and activities of daily living in AD patients treated with TPS "
        "(Hescham et al., 2023; Koch et al., 2019). SOZO uses TPS as a primary modality in AD.",
        "tDCS targeting temporal-parietal and frontal regions has shown cognitive benefits in MCI and mild "
        "AD in multiple RCTs. The evidence for tDCS is strongest for memory consolidation (temporal anodal) "
        "and executive function (DLPFC anodal). The combination of TPS and tDCS represents the most "
        "comprehensive SOZO approach to AD (Ferrucci et al., 2008; Meinzer et al., 2015).",
    ],
    pathophysiology=[
        "AD pathophysiology begins with amyloid cascade: amyloid-beta oligomers disrupt synaptic function, "
        "triggering tau hyperphosphorylation, tangle formation, and neuroinflammation via microglial "
        "activation. Cholinergic neurons of the nucleus basalis of Meynert (NBM) are affected early, "
        "impairing the cholinergic projection to hippocampus and cortex \u2014 the basis of cholinesterase "
        "inhibitor therapy.",
        "Default mode network (DMN) disruption is a hallmark of AD: the DMN (particularly the posterior "
        "cingulate cortex, precuneus, and hippocampus) shows early amyloid accumulation and functional "
        "disconnection. This manifests as impaired episodic memory, spatial navigation, and self-referential "
        "processing. The CEN progressively fails as prefrontal involvement increases in moderate AD.",
    ],
    std_treatment=[
        "Pharmacological treatment includes cholinesterase inhibitors (donepezil, rivastigmine, galantamine) "
        "for mild-to-moderate AD, and memantine (NMDA antagonist) for moderate-to-severe AD. Anti-amyloid "
        "monoclonal antibodies (lecanemab, donanemab) represent a new disease-modifying class approved for "
        "early AD/MCI. Non-pharmacological approaches (cognitive stimulation, physical exercise, "
        "multidomain lifestyle interventions) are recommended alongside NIBS.",
    ],
    symptoms=[
        ["Episodic Memory Loss",    "Difficulty learning new information; forgetting recent events; repetition",      "Core criterion",    "Petersen 2014; McKhann 2011"],
        ["Language Impairment",     "Word-finding difficulties (anomia); reduced verbal fluency; paraphasia",         "60\u201380%",       "McKhann 2011"],
        ["Visuospatial Impairment", "Getting lost; difficulty with navigation; constructional apraxia",               "50\u201370%",       "McKhann 2011"],
        ["Executive Dysfunction",   "Planning, abstract reasoning, cognitive flexibility impairment",                 "50\u201380%",       "Salmon 2009"],
        ["Attention Deficits",      "Poor sustained attention; vulnerability to distraction; mental fatigue",         "50\u201380%",       "Salmon 2009"],
        ["Apraxia",                 "Impaired purposeful movements despite intact motor function",                    "30\u201360%",       "Grossman 2010"],
        ["Depression / Apathy",     "Apathy most common NPS; depression frequent; may precede cognitive symptoms",    "40\u201370%",       "Lyketsos 2011"],
        ["Agitation / Anxiety",     "Restlessness, anxiety, verbal and physical agitation",                          "30\u201350%",       "Lyketsos 2011"],
        ["Sleep Disturbance",       "Disrupted circadian rhythms; insomnia; sundowning",                             "40\u201370%",       "McCurry 2000"],
        ["Functional Decline",      "Progressive loss of instrumental and basic ADL abilities",                       "Progressive",       "Alzheimer's Association 2023"],
    ],
    brain_regions=[
        ["Hippocampus",               "Earliest site of AD pathology; episodic memory consolidation",                   "TPS deep target (CE-marked); tDCS temporal anodal (indirect hippocampal modulation)","Koch 2019; Hescham 2023"],
        ["Entorhinal Cortex",         "Grid cells; memory indexing; earliest neurofibrillary tangle site",              "TPS temporal deep target; indirect via tDCS temporal placement",                    "Braak & Braak 1991"],
        ["Temporal Cortex (T3/T4/P3/P4)","Semantic memory; language; visuospatial processing",                        "Anodal tDCS (T3/T4/P3/P4); TPS temporal targeted",                                 "Ferrucci 2008; Meinzer 2015"],
        ["DLPFC",                     "Executive function; working memory; top-down attention",                          "Anodal tDCS primary target; TPS targeted (investigational in AD)",                  "Ferrucci 2008"],
        ["Posterior Cingulate / Precuneus","DMN hub; early amyloid accumulation; connectivity hub",                    "TPS primary target (CE-marked indication); tDCS indirect modulation",               "Koch 2019"],
        ["Parietal Cortex",           "Visuospatial processing; attention; sensorimotor integration",                    "Anodal tDCS (P3/P4); TPS parietal targeted",                                        "McKhann 2011"],
        ["Nucleus Basalis of Meynert","Cholinergic projection source; early neuronal loss in AD",                       "TPS deep target (investigational as cholinergic augmentation strategy)",             "Hescham 2023"],
        ["Prefrontal Cortex",         "Executive control, attention, working memory \u2014 affected in moderate AD",    "DLPFC anodal tDCS; TPS frontal targeted",                                           "Salmon 2009"],
    ],
    brainstem=[
        ["Locus Coeruleus (LC)",         "Early noradrenergic degeneration in AD; arousal and attention impairment",    "taVNS modulates LC via NTS; augments noradrenergic tone for arousal",              "Braak 2004"],
        ["Dorsal Raphe Nucleus",         "Serotonergic degeneration; depression and sleep disturbance in AD",           "taVNS and tDCS (DLPFC) may influence serotonergic projections",                    "Rupprecht 2020"],
        ["Nucleus Tractus Solitarius",   "Vagal afferent relay; modulates hippocampal plasticity via NTS-LC pathway",  "Primary target of taVNS; hippocampal memory consolidation mechanism",             "Frangos 2015"],
        ["Brainstem Cholinergic Nuclei", "Ch5/Ch6 cholinergic neurons; REM sleep and arousal regulation",               "Indirectly modulated by taVNS and tDCS cortical pathways",                        "Mesulam 2013"],
        ["Raphe-Hippocampal Pathway",    "Serotonergic modulation of hippocampal neurogenesis and plasticity",           "taVNS and tDCS indirectly modulate serotonergic-hippocampal circuits",             "Rupprecht 2020"],
    ],
    phenotypes=[
        ["MCI \u2014 Amnestic (Single Domain)",  "Memory impairment only; MoCA/MMSE within norms; daily function preserved",         "Episodic memory loss; word-finding; navigation",         "TPS hippocampal (CE-marked); tDCS temporal anodal"],
        ["MCI \u2014 Multi-Domain",              "Multiple cognitive domains affected but ADL preserved",                             "Memory + executive + language + visuospatial impairment","TPS + tDCS bilateral temporal-parietal-frontal"],
        ["Typical AD \u2014 Mild (MMSE 21\u201326)","Episodic memory dominant; MMSE >20; functional impairment beginning",            "Memory, word-finding, ADL early impairment",            "TPS (CE-marked primary) + tDCS temporal/DLPFC"],
        ["Typical AD \u2014 Moderate (MMSE 15\u201320)","Progressive multi-domain decline; meaningful ADL support needed",            "Memory, language, visuospatial, ADL dependence",        "TPS (CE-marked) + tDCS supportive + caregiver training"],
        ["Posterior Cortical Atrophy",           "Visuospatial and apraxic features dominant; memory relatively preserved",          "Visuospatial, reading, object recognition impairment",  "TPS parietal targeted; tDCS P3/P4 anodal"],
        ["Dysexecutive AD",                      "Executive and frontal features dominant; memory less prominent initially",          "Planning, flexibility, behaviour dysregulation",         "TPS frontal + DLPFC; tDCS F3/F4 anodal"],
        ["AD + Neuropsychiatric (NPS dominant)", "Agitation, apathy, depression as primary burden alongside cognitive decline",       "Apathy, depression, agitation, anxiety, sleep disorder","tDCS DLPFC; TPS; taVNS; CES for NPS"],
    ],
    symptom_map=[
        ["Episodic Memory",           "Hippocampus, Entorhinal, Temporal","TPS deep hippocampal (CE-marked) + tDCS temporal anodal (T3/T4/P3/P4)",                    "Level A (TPS CE-marked; Koch 2019)"],
        ["Language / Anomia",         "Left temporal cortex, IFG",        "Left temporal anodal tDCS (T3/Broca) + TPS temporal-frontal targeted",                      "Moderate (Meinzer 2015)"],
        ["Visuospatial",              "Parietal, occipital, fusiform",     "Bilateral parietal anodal tDCS (P3/P4) + TPS parietal targeted",                           "Moderate (McKhann 2011)"],
        ["Executive Function",        "DLPFC, Frontal, Parietal",         "L-DLPFC anodal tDCS + TPS frontal targeted",                                                "Moderate (Ferrucci 2008)"],
        ["Attention / Vigilance",     "DLPFC, ACC, Parietal",             "DLPFC anodal tDCS + TPS frontal; taVNS for arousal",                                       "Moderate (Ferrucci 2008)"],
        ["Apathy / Depression (NPS)", "DLPFC, ACC, OFC",                  "L-DLPFC tDCS + TPS frontal; taVNS primary + CES",                                          "Moderate (Ferrucci 2008)"],
        ["Sleep Disturbance",         "Frontal, Thalamus, Circadian",     "CES primary; taVNS; tDCS evening slow-wave",                                               "Moderate (Philip 2017)"],
        ["General Cognitive Decline", "Temporal-Parietal-Frontal network","TPS holocranial (CE-marked global protocol) + tDCS bilateral temporal-DLPFC",               "Level A (TPS CE-marked; Koch 2019)"],
        ["Anxiety / Agitation",       "Amygdala, ACC, Insula",            "taVNS primary; CES; bilateral tDCS supportive; non-pharmacological behavioural strategies","Moderate (Lyketsos 2011)"],
    ],
    montage=[
        ["MCI \u2014 Amnestic",      "T3/T4 bilateral anode + P3/P4 anode; Fz cathode",                              "TPS hippocampal targeted + temporal cortex (CE-marked)","taVNS, CES (sleep)"],
        ["MCI \u2014 Multi-Domain",  "T3/T4 + P3/P4 bilateral anode + F3 anode; extracephalic cathode",              "TPS temporal-parietal-frontal targeted (CE-marked)",     "taVNS, CES supportive"],
        ["Mild AD",                  "T3/T4 + P3/P4 bilateral anodal + F3 anode; Fz cathode",                        "TPS holocranial + targeted (CE-marked primary)",         "taVNS, CES; caregiver support"],
        ["Moderate AD",              "T3/T4 bilateral anodal (simplified); F3 anode supportive",                     "TPS holocranial + targeted (CE-marked); lower energy threshold","taVNS, CES; simplified sessions"],
        ["Posterior Cortical",       "P3/P4 bilateral anode + T3/T4 anode; F3 cathode",                              "TPS parietal targeted + temporal",                       "taVNS, CES supportive"],
        ["Dysexecutive AD",          "F3/F4 bilateral anode + P3/P4 anode",                                          "TPS frontal + DLPFC targeted",                           "taVNS, CES supportive"],
        ["AD + NPS",                 "F3 anode + T3/T4 anodal; Pz cathode",                                          "TPS frontal + temporal targeted",                        "taVNS primary; CES"],
        ["No response after 8\u201310 sessions","Reassess cognitive baseline; adjust TPS energy/ROI; add tDCS parietal","Adjust TPS parameters; consider additional ROI",         "Review adjuncts; caregiver coaching"],
    ],
    tdcs_protocols=[
        ["C1","Memory (primary) \u2014 temporal anodal",    "T3/T4 bilateral anode","Fz / extracephalic",         "2 mA, 20\u201330 min, 5\u00d7/wk \u00d7 6\u201310 wks","Level B; memory improvement (Ferrucci 2008; Meinzer 2015)"],
        ["C2","Temporal-parietal bilateral",                 "T3 + T4 + P3 + P4 anode","Fz, extracephalic",      "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Ferrucci 2008)"],
        ["C3","Executive function / DLPFC",                  "F3 anode",         "Fp2 / extracephalic",           "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Ferrucci 2008)"],
        ["C4","Combined memory + executive",                 "F3 + T3/T4 anode", "Fp2 / extracephalic",          "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Meinzer 2015)"],
        ["C5","Attention / vigilance",                       "F3 + P3/P4 anode", "Fp2, extracephalic",           "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Ferrucci 2008)"],
        ["C6","Apathy / depression (NPS)",                   "F3 anode",         "Fp2, extracephalic",           "2 mA, 20\u201330 min, 5\u00d7/wk",                  "Moderate (Lyketsos 2011)"],
        ["C7","Sleep disturbance",                           "Fz anode",         "Pz / extracephalic",           "1\u20132 mA, 20 min, evening",                       "Moderate (Philip 2017)"],
        ["C8","Maintenance block (MCI)",                     "T3/T4 anode",      "Fz / extracephalic",           "2 mA, 20\u201330 min, 3\u00d7/wk maintenance",      "Moderate (Ferrucci 2008)"],
    ],
    plato_protocols=[
        ["C1-PS","Memory \u2014 temporal",  "Posterior temporal","T3/T4 area","Nape of neck",  "1.6 mA, 20\u201330 min","5\u00d7/wk","Pair with memory encoding tasks"],
        ["C2-PS","Temporal-parietal",        "Posterior temporal","T3/T4 area","Nape of neck",  "1.6 mA, 20\u201330 min","5\u00d7/wk","Adjust placement for parietal"],
        ["C3-PS","Executive / DLPFC",        "Frontal (F3 area)", "F3",         "Right shoulder","1.6 mA, 20\u201330 min","5\u00d7/wk","Cognitive activation during session"],
        ["C4-PS","Combined memory + EF",     "Frontal (F3 area)", "F3",         "Shoulder",      "1.6 mA, 20\u201330 min","5\u00d7/wk","After HDCkit induction phase"],
        ["C5-PS","Attention",               "Frontal (F3 area)", "F3",         "Shoulder",      "1.6 mA, 20\u201330 min","5\u00d7/wk","Caregiver-supervised session"],
        ["C6-PS","Apathy / NPS",             "Frontal (F3 area)", "F3",         "Shoulder",      "1.6 mA, 20\u201330 min","5\u00d7/wk","CES concurrent recommended"],
        ["C7-PS","Sleep disturbance",        "Frontal (Fz area)", "Fz",         "Shoulder",      "1.6 mA, 20 min",        "Evening",    "CES concurrent recommended"],
        ["C8-PS","Maintenance",              "Posterior temporal","T3/T4 area", "Nape of neck",  "1.6 mA, 20\u201330 min","3\u00d7/wk", "Caregiver-supervised maintenance"],
    ],
    tps_protocols=[
        ["T1","Memory / hippocampal (CE-marked primary)","Hippocampus, Entorhinal cortex, Temporal cortex","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz; 10 sessions/2 wks","4,000 Tgt (hippocampal region) + 4,000 Std; total 8\u201312K","CE-MARKED (Koch 2019; Hescham 2023)"],
        ["T2","Global cognitive decline \u2014 holocranial","Holocranial: Frontal + temporal + parietal + posterior cingulate","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz","6,000 Std + 4,000 Tgt; total 8\u201312K","CE-MARKED (Koch 2019)"],
        ["T3","Executive function / DLPFC",              "DLPFC bilateral, mPFC, ACC",        "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz","4,000 Tgt (DLPFC) + 2,000 PRE; total 6\u20138K","CE-MARKED (NEUROLITH AD indication)"],
        ["T4","Posterior cortical (visuospatial)",        "Parietal, precuneus, posterior temporal","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz","4,000 Tgt (parietal) + 2,000 Std; total 6\u20138K","CE-MARKED (NEUROLITH AD indication)"],
        ["T5","NPS (apathy, depression)",                 "DLPFC + ACC + OFC",                 "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz","4,000 Tgt + 2,000 PRE; total 6\u20138K","CE-MARKED (NEUROLITH AD indication)"],
    ],
    ces_role=[
        ["Sleep disturbance",    "Targets insomnia, sundowning, and circadian disruption common in AD",                 "Evening before bed; 20\u201340 min; caregiver-assisted"],
        ["Apathy / Depression",  "Supportive for mild NPS mood symptoms alongside primary tDCS/TPS",                   "Evening sessions; or non-tDCS days; 20\u201340 min"],
        ["Anxiety / Agitation",  "Reduces anxiety and agitation as adjunct to primary protocols",                      "As needed; 20\u201330 min; pre-session for anxious patients"],
        ["Arousal / Motivation", "May improve arousal and motivation; adjunct to cognitive engagement",                 "Morning; 20\u201330 min; caregiver-supervised"],
    ],
    tavns_role="taVNS modulates locus coeruleus noradrenergic tone and NTS-hippocampal circuits, potentially enhancing memory consolidation and arousal in Alzheimer's disease and MCI. It also supports mood and sleep in neuropsychiatric symptom burden.",
    combinations=[
        ("1) MCI \u2014 Amnestic", [
            ["TPS + tDCS (temporal)","TPS hippocampal (CE-marked) + temporal anodal tDCS; evidence-based combination (Koch 2019; Ferrucci 2008)","TPS session + tDCS on same or adjacent days","Episodic memory loss, word-finding, navigation difficulties"],
            ["taVNS + TPS/tDCS","taVNS primes hippocampal plasticity via NTS-LC pathway before primary sessions (Frangos 2015)","taVNS 20\u201330 min before TPS or tDCS","Attention, arousal, memory consolidation"],
            ["CES + Standard Care","CES for sleep and circadian stabilisation; supports memory consolidation overnight (Philip 2017)","Evening before bed","MCI with sleep disturbance and circadian dysregulation"],
        ]),
        ("2) Mild AD", [
            ["TPS + tDCS Combination","TPS holocranial (CE-marked) + tDCS temporal-parietal-frontal; most comprehensive approach (Koch 2019; Ferrucci 2008)","TPS and tDCS on alternating days or combined block","Multi-domain cognitive decline in mild AD"],
            ["taVNS + TPS/tDCS","taVNS augments noradrenergic and cholinergic circuits; additive with TPS and tDCS (Frangos 2015)","taVNS before TPS or tDCS sessions","Attention, arousal, mood, sleep"],
            ["CES + Caregiver Support","CES for NPS and sleep; caregiver education programme integral to mild AD management (Philip 2017)","Evening; caregiver-administered","NPS symptoms, sleep disturbance, anxiety"],
        ]),
        ("3) Moderate AD", [
            ["TPS (CE-marked) + Simplified tDCS","TPS primary; simplified single-site tDCS as adjunct (Koch 2019)","TPS primary; tDCS at reduced burden (3\u00d7/wk)","Moderate multi-domain decline with ADL impairment"],
            ["taVNS + Standard Care","taVNS for autonomic stability and NPS; low burden on patient and caregiver (Frangos 2015)","Daily or 5\u00d7/wk; caregiver-assisted","Agitation, anxiety, sleep disturbance in moderate AD"],
            ["CES + Caregiver Support","CES daily for sleep, agitation, and mood; caregiver programme essential (Philip 2017)","Evening daily if tolerated","Severe sleep disturbance, sundowning, NPS in moderate AD"],
        ]),
        ("4) Posterior Cortical AD", [
            ["TPS (parietal targeted) + tDCS (P3/P4)","TPS parietal + occipital + tDCS bilateral parietal anodal (Koch 2019; Ferrucci 2008)","TPS and tDCS combined block; visuospatial rehabilitation integrated","Visuospatial, reading, object recognition impairment"],
            ["taVNS + TPS/tDCS","taVNS augments overall cortical arousal and plasticity (Frangos 2015)","taVNS before TPS or tDCS","General attention, arousal support"],
            ["CES + Occupational Therapy","CES for sleep/anxiety; OT for visuospatial compensatory strategies (Philip 2017)","Evening (sleep); OT integration","Sleep, anxiety, ADL adaptation for visuospatial deficits"],
        ]),
        ("5) AD + NPS (Apathy/Depression)", [
            ["tDCS (DLPFC) + TPS (frontal)","L-DLPFC anodal tDCS + TPS frontal targeted; addresses NPS alongside cognitive decline (Koch 2019; Ferrucci 2008)","tDCS before TPS or alternating days","Apathy, depression, motivational loss"],
            ["taVNS + Primary Protocols","taVNS for NPS mood modulation and autonomic stabilisation (Frangos 2015)","taVNS before primary sessions; or standalone on off-days","Depression, anxiety, agitation in AD"],
            ["CES + Standard Care","CES for NPS mood, sleep, and anxiety; adjunct to primary protocol (Philip 2017)","Evening and/or morning as prescribed","NPS with sleep, anxiety, and mood burden"],
        ]),
        ("6) MCI Conversion Prevention Focus", [
            ["TPS + tDCS + Cognitive Stimulation","TPS hippocampal + temporal tDCS + structured cognitive stimulation; maximum neuroplasticity approach (Koch 2019)","TPS + tDCS on alternating days; cognitive stimulation daily","MCI with high AD conversion risk"],
            ["taVNS + Lifestyle Intervention","taVNS + multidomain lifestyle programme (exercise, diet, sleep, social engagement)","Daily taVNS; lifestyle programme concurrent","MCI with modifiable risk factor burden"],
            ["CES + Sleep Optimisation","CES for sleep; poor sleep accelerates AD pathology; prioritise sleep in MCI (Philip 2017; Holth 2019)","Evening; sleep hygiene concurrent","MCI with sleep disturbance as conversion risk factor"],
        ]),
    ],
    combination_summary=[
        ["MCI \u2014 Amnestic",    "TPS + tDCS (temporal)",    "TPS hippocampal (CE-marked) + temporal anodal tDCS (Koch 2019; Ferrucci 2008)",               "TPS + tDCS combined block",    "Episodic memory, word-finding",     "Level A (TPS CE-marked)"],
        ["Mild AD",                "TPS + tDCS Combination",   "TPS holocranial (CE-marked) + tDCS temporal-parietal-frontal (Koch 2019; Ferrucci 2008)",      "Alternating days or combined", "Multi-domain cognitive decline",    "Level A (TPS CE-marked)"],
        ["Moderate AD",            "TPS (CE-marked) primary",  "TPS primary with simplified tDCS adjunct (Koch 2019)",                                         "TPS primary 3\u20135\u00d7/wk","Moderate decline, ADL impairment",  "Level A (TPS CE-marked)"],
        ["Posterior Cortical",     "TPS parietal + tDCS P3/P4","TPS parietal + tDCS bilateral parietal anodal (Koch 2019; Ferrucci 2008)",                     "Combined block",               "Visuospatial, apraxia",             "Level A (TPS CE-marked)"],
        ["AD + NPS",               "tDCS (DLPFC) + TPS",       "L-DLPFC anodal tDCS + TPS frontal; addresses NPS alongside cognitive decline (Koch 2019)",      "tDCS before TPS",              "Apathy, depression in AD",          "Moderate"],
        ["MCI Conversion Focus",   "TPS + tDCS + Cognitive Stimulation","TPS hippocampal + temporal tDCS + cognitive stimulation; maximum neuroplasticity (Koch 2019)","TPS + tDCS + daily CST","MCI conversion prevention",        "Level A (TPS CE-marked)"],
    ],
    outcomes=[
        ["MoCA",               "Global cognition screening",             "Baseline, months 1, 3, 6",     "Score <26 impairment; track \u22652-point improvement"],
        ["MMSE",               "Global cognitive severity",              "Baseline, months 1, 3, 6",     "21\u201326 mild; 10\u201320 moderate; <10 severe AD"],
        ["ADAS-Cog",           "AD cognitive battery (primary outcome)", "Baseline, months 3, 6",        "\u22654-point improvement = clinically meaningful change"],
        ["CDR",                "Clinical Dementia Rating \u2014 global", "Baseline, months 3, 6",        "0 = normal; 0.5 = MCI; 1 = mild; 2 = moderate; 3 = severe"],
        ["NPI",                "Neuropsychiatric Inventory (NPS)",       "Baseline, months 3, 6",        "Track improvement in apathy, depression, agitation subscales"],
        ["ADL / IADL",         "Functional independence",                "Baseline, months 3, 6",        "Track maintenance or improvement in daily activities"],
        ["PHQ-9 or GDS",       "Comorbid depressive symptoms",          "Baseline, months 3, 6",        "\u226510 moderate depression (PHQ-9)"],
        ["Quality of Life \u2014 AD (QoL-AD)","Subjective wellbeing",   "Baseline, months 3, 6",        "Score range 13\u201352; higher = better QoL"],
        ["SOZO PRS",           "NIBS-specific functional outcome",       "Baseline, weeks 2, 4, 8, 12", "Proprietary; composite cognitive/NPS/functional domains"],
    ],
    outcomes_abbrev=["MoCA", "MMSE", "ADAS-Cog", "CDR", "SOZO PRS"],
)

# ══════════════════════════════════════
# 6. STROKE REHABILITATION
# ══════════════════════════════════════
CONDITIONS["stroke_rehab"] = dict(
    name="Post-Stroke Rehabilitation", icd10="I63 / I64",
    doc_num="SOZO-STR-FEL-006",
    tps_status="INVESTIGATIONAL / OFF-LABEL",
    tdcs_class="Level B evidence; CE-marked",
    tps_class="INVESTIGATIONAL / OFF-LABEL",
    tavns_class="Emerging adjunctive; CE-marked",
    ces_class="Supportive adjunct for mood/sleep",
    inclusion=[
        "Confirmed ischaemic or haemorrhagic stroke diagnosis (neuroimaging confirmed)",
        "At least 1 month post-stroke (subacute to chronic phase eligible)",
        "Residual motor, cognitive, language, or mood impairment amenable to rehabilitation",
        "Age 18+ years",
        "Medically stable; cleared for active rehabilitation by treating neurologist",
        "Baseline assessments completed (NIHSS, Fugl-Meyer, FIM, MoCA, PHQ-9)",
    ],
    exclusion=[
        "Intracranial metallic hardware or haemorrhagic transformation within stimulation path",
        "Cochlear implant or DBS device",
        "Active intracranial haemorrhage or unstable vascular lesion",
        "Skull defects or craniectomy at electrode/transducer sites",
        "Severe aphasia preventing informed consent (consider proxy consent pathway)",
        "Pregnancy (tDCS, TPS)",
        "Medically unstable (uncontrolled hypertension, cardiac instability, recent TIA)",
    ],
    discussion=[
        "Cardiac pacemaker or defibrillator \u2014 individual risk\u2013benefit assessment",
        "Epilepsy or post-stroke seizures \u2014 formal risk\u2013benefit with documented rationale",
        "Coagulation disorders or anticoagulants (especially TPS) \u2014 haematoma risk",
        "Large cortical or subcortical lesion in stimulation path \u2014 altered current distribution",
        "Severe neglect limiting session compliance",
        "Dermatological conditions at electrode sites",
        "Significant spasticity preventing electrode/transducer placement",
        "Severe post-stroke depression requiring psychiatric referral before NIBS initiation",
    ],
    overview=[
        "Stroke is the second leading cause of death and a major cause of disability worldwide. "
        "Approximately 15 million strokes occur annually; over 5 million result in permanent disability. "
        "Post-stroke rehabilitation addresses motor recovery, cognitive rehabilitation, language therapy, "
        "and mood management across the acute, subacute, and chronic phases.",
        "tDCS can modulate cortical excitability to support stroke recovery by enhancing ipsilesional "
        "motor cortex excitability (anodal tDCS) or suppressing contralesional hyperactivity (cathodal "
        "tDCS). Both strategies, and the bihemispheric approach (anodal ipsilesional + cathodal "
        "contralesional), have been investigated in multiple RCTs (Elsner et al., 2016; Bikson et al., 2016).",
        "The principle of use-dependent plasticity is central: tDCS administered during or immediately "
        "before rehabilitation (physical therapy, occupational therapy, aphasia therapy) amplifies "
        "practice-induced neuroplastic changes. TPS is investigational in stroke, targeting ipsilesional "
        "motor and premotor cortex.",
        "NIBS Evidence in Stroke: Level B evidence supports tDCS for post-stroke motor recovery. "
        "Meta-analyses demonstrate significant improvements in upper limb function, gait, and some "
        "cognitive outcomes with combined tDCS + rehabilitation. taVNS paired with rehabilitation "
        "has emerging evidence from the VNS-REHAB RCT.",
    ],
    pathophysiology=[
        "Stroke disrupts the balance between ipsilesional and contralesional hemispheres. The lesioned "
        "hemisphere loses excitatory input; the intact hemisphere becomes relatively hyperactive, "
        "exerting transcallosal inhibition on the damaged side. This interhemispheric imbalance is the "
        "primary target of tDCS in stroke rehabilitation.",
        "Cortical reorganisation following stroke involves perilesional tissue plasticity, recruitment "
        "of ipsilateral pathways, and diaschisis resolution. Neuromodulation during the critical "
        "plasticity window (subacute phase: weeks 1\u20133 months post-stroke) may amplify these processes. "
        "White matter tract integrity and lesion location are key predictors of tDCS response.",
    ],
    std_treatment=[
        "Standard stroke rehabilitation includes intensive physical therapy, occupational therapy, "
        "speech-language therapy, neuropsychological rehabilitation, and mood management. "
        "Evidence-based rehabilitation intensity (National Stroke Foundation guidelines) recommends "
        "at least 3 hours of active rehabilitation per day in the acute/subacute phase. "
        "NIBS is adjunctive to standard rehabilitation and should always be paired with active therapy.",
    ],
    symptoms=[
        ["Hemiplegia / Hemiparesis",  "Weakness or paralysis of ipsilesional limbs; upper limb most affected", "Majority",          "Langhorne 2011"],
        ["Spasticity",                "Velocity-dependent muscle tone increase; limb posturing",                 "30\u201340%",       "Sommerfeld 2004"],
        ["Gait Impairment",           "Reduced gait speed, step length, balance, fall risk",                    "60\u201380%",       "Langhorne 2011"],
        ["Post-Stroke Depression",    "Depressed mood, tearfulness, anhedonia, apathy post-stroke",             "30\u201340%",       "Robinson 2016"],
        ["Cognitive Impairment",      "Attention, memory, executive function, processing speed deficits",        "50\u201370%",       "Dichgans 2007"],
        ["Aphasia",                   "Expressive, receptive, or mixed language impairment",                    "30\u201335%",       "Berthier 2005"],
        ["Spatial Neglect",           "Failure to attend to or act on stimuli on one side (usually left)",       "20\u201330%",       "Ringman 2004"],
        ["Fatigue",                   "Post-stroke fatigue; central and peripheral components",                  "40\u201368%",       "Duncan 2012"],
        ["Dysphagia",                 "Swallowing impairment; aspiration risk",                                  "30\u201350%",       "Martino 2005"],
        ["Pain",                      "Central post-stroke pain, shoulder pain, spasticity-related pain",        "10\u201330%",       "Klit 2009"],
    ],
    brain_regions=[
        ["Ipsilesional M1",              "Reduced excitability post-stroke; primary motor execution",            "Anodal tDCS primary target; restores ipsilesional excitability (Level B)",    "Elsner 2016; Bikson 2016"],
        ["Contralesional M1",            "Relative hyperactivation; transcallosal inhibition of ipsilesional M1","Cathodal tDCS to reduce contralesional hyperactivity; bihemispheric approach","Fregni 2006"],
        ["Ipsilesional SMA / Premotor",  "Motor planning and sequencing; compensatory activation post-stroke",   "Anodal tDCS SMA/premotor; TPS ipsilesional premotor (investigational)",     "Zimerman 2012"],
        ["DLPFC",                        "Cognitive function; post-stroke depression; attention",                 "L-DLPFC anodal tDCS for PSD and cognitive recovery",                        "Bhogal 2003"],
        ["Broca / Left IFG",             "Expressive language; aphasia rehabilitation",                           "Left IFG anodal tDCS during aphasia therapy; TPS (investigational)",        "Baker 2010"],
        ["Right Parietal",               "Spatial attention; contralesional neglect circuit",                     "Cathodal tDCS right parietal or anodal left parietal for neglect",          "Brighina 2003"],
        ["Thalamus",                     "Relay for motor and sensory signals; diaschisis target",                "TPS deep target (investigational); indirect modulation via cortical tDCS",  "Yamamoto 2010"],
        ["Cerebellum",                   "Balance, coordination, gait; compensatory activation post-stroke",     "Cerebellar tDCS for gait/balance; TPS (investigational)",                   "Grimaldi 2016"],
    ],
    brainstem=[
        ["Corticospinal Tract",       "Descending motor pathways; lesion severity determinant",              "Not directly stimulated; integrity assessed for tDCS response prediction","Stinear 2017"],
        ["Locus Coeruleus (LC)",      "Noradrenergic arousal; post-stroke fatigue and mood regulation",      "taVNS modulates LC; addresses post-stroke fatigue and depression",          "Robinson 2016"],
        ["Nucleus Tractus Solitarius","Vagal afferent relay; central autonomic hub; post-stroke autonomic",  "Primary target of taVNS; VNS-REHAB mechanism for motor recovery",           "Dawson 2016"],
        ["Dorsal Raphe Nucleus",      "Serotonergic regulation; post-stroke depression and fatigue",         "taVNS and tDCS (DLPFC) modulate serotonergic tone",                         "Robinson 2016"],
        ["Periaqueductal Gray",       "Pain modulation; post-stroke central pain circuits",                  "Indirectly modulated via cortical and vagal neuromodulation",               "Klit 2009"],
    ],
    phenotypes=[
        ["Acute Motor Recovery (1\u20133 months)","Early post-stroke; maximum neuroplasticity window",           "Hemiplegia/hemiparesis, spasticity, gait impairment",          "Ipsilesional M1 anodal tDCS + intensive physio; TPS (investigational)"],
        ["Post-Stroke Depression (PSD)",          "Significant depression post-stroke; impairs rehab engagement", "Depressed mood, anhedonia, anergia, tearfulness, apathy",      "L-DLPFC anodal tDCS; taVNS; CES; antidepressant if indicated"],
        ["Post-Stroke Aphasia",                   "Expressive, receptive, or mixed language impairment",          "Language impairment, word-finding, comprehension difficulties","Left IFG anodal tDCS during language therapy; TPS (investigational)"],
        ["Spatial Neglect",                       "Failure to attend to or act on contralateral stimuli",          "Spatial neglect, attention asymmetry, functional impairment",  "Cathodal right parietal or anodal left parietal tDCS"],
        ["Chronic Motor Deficit (>6 months)",     "Chronic phase with plateau or slow recovery",                  "Residual weakness, spasticity, reduced gait speed",            "Bihemispheric tDCS; TPS motor cortex; VNS + physio (VNS-REHAB)"],
        ["Post-Stroke Cognitive Impairment",      "Attention, memory, executive deficits post-stroke",            "Cognitive slowing, attention, memory, executive dysfunction",   "DLPFC anodal tDCS; TPS frontal-temporal targeted"],
        ["Post-Stroke Fatigue",                   "Pervasive fatigue disproportionate to physical exertion",       "Fatigue, reduced endurance, poor engagement with rehab",       "tDCS DLPFC; taVNS primary; CES adjunct"],
    ],
    symptom_map=[
        ["Upper Limb Motor Deficit",   "Ipsilesional M1, SMA, Premotor","Ipsilesional M1 anodal tDCS + intensive upper limb therapy (Level B evidence; Elsner 2016)",   "Level B (Elsner 2016)"],
        ["Gait / Lower Limb",          "Ipsilesional M1, SMA",          "Ipsilesional M1/SMA anodal tDCS + gait training; TPS motor cortex (investigational)",         "Moderate (Elsner 2016)"],
        ["Post-Stroke Depression",     "L-DLPFC, sgACC",                "L-DLPFC anodal tDCS + psychotherapy; taVNS primary; CES adjunct",                           "Moderate (Bhogal 2003; Robinson 2016)"],
        ["Aphasia",                    "Left IFG (Broca), Temporal",    "Left IFG anodal tDCS during aphasia therapy; TPS left frontal-temporal (investigational)",    "Moderate (Baker 2010; Elsner 2016)"],
        ["Spatial Neglect",            "Right Parietal, Right TPJ",     "Right parietal cathodal tDCS or left parietal anodal; prism adaptation therapy",             "Moderate (Brighina 2003)"],
        ["Cognitive Impairment",       "DLPFC, Parietal, Temporal",     "DLPFC anodal tDCS + cognitive rehabilitation; TPS frontal-temporal targeted",               "Moderate (Elsner 2016)"],
        ["Post-Stroke Fatigue",        "DLPFC, SMA, Mesolimbic",        "DLPFC tDCS + taVNS primary; CES for arousal; graded exercise programme",                   "Moderate (taVNS: Dawson 2016)"],
        ["Pain",                       "Contralateral M1, ACC, Thalamus","Contralateral M1 anodal tDCS; TPS motor cortex + thalamus; taVNS; CES",                   "Moderate (Klit 2009; Fregni 2011)"],
        ["Spasticity",                 "Ipsilesional M1, SMA",          "Cathodal tDCS over spastic muscle area; TPS (investigational); stretching programme",        "Emerging (Elsner 2016)"],
    ],
    montage=[
        ["Motor deficit (upper limb)",  "Ipsilesional M1 anodal (C3 or C4 contralateral to weakness) + contralesional cathodal","Cranial (ipsilesional M1/SMA targeted) + peripheral (affected limb)","taVNS (VNS-REHAB)"],
        ["Bihemispheric (chronic motor)","Ipsilesional M1 anode + contralesional M1 cathode; simultaneous dual-site","Cranial ipsilesional + contralesional targeted",    "taVNS; CES supportive"],
        ["Post-Stroke Depression",       "F3 anode + Fp2 cathode (L-DLPFC anodal)",                                "Cranial (global + targeted L-DLPFC)",                        "taVNS primary; CES"],
        ["Aphasia",                      "Left IFG/Broca anode (F7/Broca area) + right shoulder cathode",          "Cranial (left frontal-temporal targeted)",                    "taVNS; language therapy concurrent"],
        ["Spatial Neglect",              "Right parietal (P4) cathode + left parietal (P3) anode",                 "Cranial (bilateral parietal targeted)",                       "taVNS; prism adaptation therapy"],
        ["Cognitive impairment",         "F3 + T3/T4 bilateral anode + P3/P4; extracephalic cathode",              "Cranial (frontal-temporal-parietal targeted)",                "taVNS, CES supportive"],
        ["Post-stroke fatigue",          "F3 anode + Fp2 cathode; taVNS as primary modality",                      "Cranial (global + targeted DLPFC)",                          "taVNS primary; CES; graded exercise"],
        ["No response after 8\u201310 sessions","Reassess lesion location; consider alternate montage; add TPS","TPS ipsilesional targeted; alternate ROI",                    "VNS-REHAB pairing if available"],
    ],
    tdcs_protocols=[
        ["C1","Upper limb motor (ipsilesional anodal)","C3/C4 (ipsilesional M1) anode","C3/C4 (contralesional) cathode / extracephalic","2 mA, 20\u201330 min, 5\u00d7/wk during upper limb therapy","Level B; UL function (Elsner 2016)"],
        ["C2","Bihemispheric motor",                  "Ipsilesional M1 anode","Contralesional M1 cathode; simultaneous",       "2 mA, 20\u201330 min, 5\u00d7/wk",            "Moderate (Fregni 2006)"],
        ["C3","Post-stroke depression",               "F3 anode (L-DLPFC)","Fp2 / extracephalic",                            "2 mA, 20\u201330 min, 5\u00d7/wk",            "Moderate (Bhogal 2003; Robinson 2016)"],
        ["C4","Aphasia rehabilitation",               "Left F7/Broca anode","Right shoulder cathode",                         "2 mA, 20\u201330 min during language therapy","Moderate (Baker 2010)"],
        ["C5","Spatial neglect",                      "Left P3 anode","Right P4 cathode",                                   "2 mA, 20\u201330 min + prism adaptation therapy","Moderate (Brighina 2003)"],
        ["C6","Cognitive rehabilitation",             "F3 + T3/T4 anode","Fp2, extracephalic",                              "2 mA, 20\u201330 min, 5\u00d7/wk + cognitive rehab","Moderate (Elsner 2016)"],
        ["C7","Gait / lower limb motor",              "Ipsilesional M1/SMA anode (Cz/Fz)","Extracephalic; or contralesional cathode","2 mA, 20\u201330 min + gait training", "Moderate (Elsner 2016)"],
        ["C8","Post-stroke fatigue",                  "F3 anode","Fp2, extracephalic",                                       "2 mA, 20\u201330 min, 5\u00d7/wk + graded exercise","Moderate (Elsner 2016)"],
    ],
    plato_protocols=[
        ["C1-PS","Upper limb motor",     "Motor (Cz area)",   "Cz (ipsilesional)","Shoulder (contralesional)","1.6 mA, 20\u201330 min","5\u00d7/wk during physio","Single-channel approximation of C3/C4"],
        ["C2-PS","Bihemispheric",        "Motor (Cz area)",   "Cz","Nape of neck",  "1.6 mA, 20\u201330 min","5\u00d7/wk",              "Limited bihemispheric; HDCkit preferred"],
        ["C3-PS","Post-stroke depression","Frontal (F3 area)","F3","Right shoulder", "1.6 mA, 20\u201330 min","5\u00d7/wk",              "L-DLPFC anodal primary"],
        ["C4-PS","Aphasia",              "Frontal (F3/F7 area)","F3","Right shoulder","1.6 mA, 20\u201330 min","5\u00d7/wk during language therapy","Limited; HDCkit preferred for left IFG"],
        ["C5-PS","Spatial neglect",      "Posterior (Pz area)","Pz","Shoulder",       "1.6 mA, 20\u201330 min","5\u00d7/wk + prism therapy","Parietal approximation"],
        ["C6-PS","Cognitive rehab",      "Frontal (F3 area)", "F3","Shoulder",        "1.6 mA, 20\u201330 min","5\u00d7/wk + cognitive tasks","Pair with cognitive rehabilitation"],
        ["C7-PS","Gait / lower limb",    "Motor (Cz area)",   "Cz","Nape of neck",   "1.6 mA, 20\u201330 min","5\u00d7/wk + gait training","Combine with gait training"],
        ["C8-PS","Fatigue",              "Frontal (F3 area)", "F3","Shoulder",        "1.6 mA, 20\u201330 min","5\u00d7/wk + graded exercise","Morning; graded exercise concurrent"],
    ],
    tps_protocols=[
        ["T1","Upper limb motor recovery","Ipsilesional M1, SMA, Premotor cortex",  "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz; 10 sessions/2 wks","4,000 Std + 4,000 Tgt (ipsilesional M1); total 8\u201312K","Emerging (investigational in stroke)"],
        ["T2","Gait / lower limb recovery","Ipsilesional M1 (leg area), SMA, Cerebellum","0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",              "4,000 Tgt + 2,000 PRE; total 6\u20138K",             "Investigational"],
        ["T3","Post-stroke depression",    "DLPFC bilateral, ACC",                   "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                  "4,000 Tgt (DLPFC) + 2,000 PRE; total 6\u20138K",    "Investigational"],
        ["T4","Aphasia recovery",          "Left IFG, Left temporal cortex",          "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                  "4,000 Tgt (left IFG/temporal) + 2,000 Std; total 6\u20138K","Investigational"],
        ["T5","Cognitive rehabilitation",  "DLPFC, Temporal, Parietal",              "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                  "4,000 Tgt + 2,000 PRE; total 6\u20138K",             "Investigational"],
    ],
    ces_role=[
        ["Post-stroke depression", "Supportive for mild-to-moderate PSD alongside primary tDCS; FDA-cleared for mood", "Evening sessions; or non-tDCS days; 20\u201340 min"],
        ["Sleep disturbance",      "Targets post-stroke insomnia and circadian disruption",                             "Evening before bed; 20\u201340 min"],
        ["Anxiety / Autonomic",    "Stabilises post-stroke autonomic tone and anxiety",                                 "As needed; pre-session or evening; 20\u201330 min"],
        ["Fatigue",                "May improve post-stroke fatigue and arousal as adjunct",                            "Morning; 20\u201330 min; coordinate with rehabilitation schedule"],
    ],
    tavns_role="taVNS is supported by the VNS-REHAB RCT for upper limb motor recovery when paired with rehabilitation. It modulates locus coeruleus noradrenergic tone, neuroplasticity signalling, and autonomic balance; a primary adjunct modality in post-stroke rehabilitation at SOZO.",
    combinations=[
        ("1) Acute Motor Recovery", [
            ["tDCS + Intensive Physiotherapy","Ipsilesional M1 anodal tDCS + intensive upper limb/gait therapy during stimulation; most supported stroke NIBS combination (Elsner 2016)","tDCS during physiotherapy session","Hemiplegia, hemiparesis, gait impairment"],
            ["taVNS + Rehabilitation","taVNS + upper limb rehabilitation; VNS-REHAB RCT demonstrated significant UL motor gains (Dawson 2016; 2021)","taVNS concurrent with rehabilitation exercises","Upper limb motor deficit; use-dependent plasticity"],
            ["CES + Standard Care","CES for post-stroke fatigue and mood support; adjunct throughout acute rehabilitation (Philip 2017)","Evening; or before rehabilitation for arousal","Post-stroke fatigue, mood disturbance"],
        ]),
        ("2) Post-Stroke Depression", [
            ["tDCS (DLPFC) + Psychotherapy","L-DLPFC anodal tDCS + structured psychotherapy (CBT or IPT); addresses PSD whilst supporting cognitive recovery (Robinson 2016; Bhogal 2003)","tDCS before therapy session","Depressed mood, anhedonia, tearfulness post-stroke"],
            ["taVNS + tDCS","taVNS for autonomic and mood regulation; additive with DLPFC tDCS in PSD (Dawson 2016; Robinson 2016)","taVNS before tDCS","PSD with autonomic dysregulation and fatigue"],
            ["CES + Standard Care","CES for sleep and mood in PSD; FDA-cleared; well-tolerated in post-stroke population (Philip 2017)","Evening daily if tolerated","PSD with insomnia, anxiety, somatic complaints"],
        ]),
        ("3) Aphasia Rehabilitation", [
            ["tDCS (left IFG) + Speech-Language Therapy","Left IFG anodal tDCS during aphasia therapy; most supported language combination in stroke (Baker 2010; Elsner 2016)","tDCS during speech-language therapy session","Expressive aphasia, anomia, reduced verbal fluency"],
            ["taVNS + Language Therapy","taVNS augments cortical plasticity and noradrenergic tone; emerging for aphasia recovery (Dawson 2016)","taVNS before language therapy session","Aphasia with fatigue and motivational challenges"],
            ["CES + Standard Care","CES for sleep, mood, and arousal; supports therapy engagement (Philip 2017)","Evening","Aphasia with post-stroke depression and sleep disturbance"],
        ]),
        ("4) Chronic Motor Deficit", [
            ["Bihemispheric tDCS + Intensive Physio","Ipsilesional anodal + contralesional cathodal tDCS + intensive upper limb/gait therapy (Elsner 2016; Fregni 2006)","tDCS during physiotherapy; intensive daily programme","Chronic residual weakness, spasticity, reduced function"],
            ["taVNS + Rehabilitation","taVNS paired with rehabilitation in chronic stroke; VNS-REHAB showed sustained benefit (Dawson 2021)","taVNS concurrent with rehabilitation","Chronic upper limb deficit; plateau in recovery"],
            ["CES + Fatigue Management","CES for chronic post-stroke fatigue; enables better rehabilitation engagement (Philip 2017)","Morning and/or before rehabilitation session","Chronic fatigue limiting rehabilitation participation"],
        ]),
        ("5) Post-Stroke Cognitive Impairment", [
            ["tDCS (DLPFC/temporal) + Cognitive Rehabilitation","DLPFC + temporal anodal tDCS + structured cognitive rehabilitation; targets attention, memory, executive function (Elsner 2016)","tDCS before/during cognitive rehabilitation","Attention, memory, executive function impairment post-stroke"],
            ["taVNS + Cognitive Rehabilitation","taVNS augments noradrenergic tone for attention and memory; adjunct to cognitive rehabilitation (Dawson 2016)","taVNS before cognitive training","Attentional fatigue, processing speed reduction"],
            ["CES + Standard Care","CES for sleep and mood; supports cognitive rehabilitation engagement (Philip 2017)","Evening","Cognitive impairment with sleep disturbance and mood"],
        ]),
        ("6) Post-Stroke Fatigue", [
            ["tDCS (DLPFC) + Graded Exercise","DLPFC anodal tDCS + graded exercise programme; addresses central fatigue mechanisms (Elsner 2016)","tDCS before exercise/rehabilitation session","Persistent post-stroke fatigue limiting rehabilitation"],
            ["taVNS + Standard Care","taVNS primary for fatigue; modulates arousal and autonomic tone via LC-NE pathway (Dawson 2016)","taVNS daily; 20\u201330 min before activity","Central post-stroke fatigue with autonomic dysregulation"],
            ["CES + Standard Care","CES for sleep and energy; addresses fatigue maintaining cycle of poor sleep (Philip 2017)","Morning (energy) and evening (sleep) as prescribed","Fatigue with sleep disturbance and mood"],
        ]),
    ],
    combination_summary=[
        ["Acute Motor Recovery",   "tDCS + Intensive Physio",      "Ipsilesional M1 anodal tDCS + intensive rehabilitation; most supported combination (Elsner 2016)", "tDCS during physio", "Hemiplegia, hemiparesis",     "Level B"],
        ["Acute Motor Recovery",   "taVNS + Rehabilitation",        "VNS-REHAB RCT demonstrated significant UL motor gains (Dawson 2016; 2021)",                       "taVNS concurrent", "Upper limb motor deficit",    "Level B (VNS-REHAB)"],
        ["Post-Stroke Depression", "tDCS + Psychotherapy",          "L-DLPFC anodal tDCS + psychotherapy for PSD (Robinson 2016; Bhogal 2003)",                        "tDCS before therapy","Depressed mood, anhedonia",    "Moderate"],
        ["Aphasia",                "tDCS (left IFG) + SLT",         "Left IFG anodal tDCS + aphasia therapy; most supported language combination (Baker 2010)",          "tDCS during SLT",  "Expressive aphasia, anomia",  "Moderate"],
        ["Chronic Motor Deficit",  "Bihemispheric tDCS + Physio",   "Bihemispheric tDCS + intensive physio in chronic phase (Elsner 2016; Fregni 2006)",                "tDCS during physio","Chronic residual weakness",   "Moderate"],
        ["Post-Stroke Cognitive",  "tDCS + Cognitive Rehabilitation","DLPFC + temporal tDCS + structured cognitive rehabilitation (Elsner 2016)",                         "tDCS before cog rehab","Attention, memory, EF",     "Moderate"],
    ],
    outcomes=[
        ["NIHSS",                  "Neurological deficit severity",          "Baseline, weeks 4, 8, 12",     "Score 0 = no symptoms; \u226525 = severe; track improvement"],
        ["Fugl-Meyer (Upper Limb)","Upper limb motor function",              "Baseline, weeks 4, 8, 12",     "Maximum 66 (UL); \u226510% change = clinically meaningful"],
        ["FIM",                    "Functional Independence Measure",        "Baseline, months 1, 3",        "Score 18\u2013126; higher = more independent"],
        ["Barthel Index",          "Activities of daily living",             "Baseline, months 1, 3",        "Score 0\u2013100; \u226575 = mildly dependent"],
        ["MoCA",                   "Global cognition screening",             "Baseline, month 3",            "Score <26 cognitive impairment; track improvement"],
        ["PHQ-9",                  "Post-stroke depression",                 "Baseline, weeks 4, 8, 12",     "\u226510 moderate depression"],
        ["10-Meter Walk Test",     "Gait speed",                             "Baseline, weeks 4, 8, 12",     "Normal >0.8 m/s; post-stroke typically 0.2\u20130.6 m/s"],
        ["mRS",                    "Modified Rankin Scale",                  "Baseline, month 3",            "0\u20131 excellent; 2\u20133 moderate dependency; 4\u20135 severe"],
        ["SOZO PRS",               "NIBS-specific functional outcome",       "Baseline, weeks 2, 4, 8, 12", "Proprietary; composite motor/mood/cognitive/functional domains"],
    ],
    outcomes_abbrev=["NIHSS", "Fugl-Meyer", "FIM", "PHQ-9", "SOZO PRS"],
)

# ══════════════════════════════════════
# 7. TRAUMATIC BRAIN INJURY
# ══════════════════════════════════════
CONDITIONS["tbi"] = dict(
    name="Traumatic Brain Injury", icd10="S06 / F07.2",
    doc_num="SOZO-TBI-FEL-007",
    tps_status="INVESTIGATIONAL / OFF-LABEL",
    tdcs_class="Level B evidence; CE-marked",
    tps_class="INVESTIGATIONAL / OFF-LABEL",
    tavns_class="Emerging adjunctive; CE-marked",
    ces_class="FDA-cleared adjunctive for mood/anxiety/sleep",
    inclusion=[
        "Confirmed TBI diagnosis (neuroimaging or clinical criteria; Glasgow Coma Scale \u22658 on admission)",
        "At least 3 months post-injury (chronic/subacute phase eligible)",
        "Residual cognitive, motor, mood, or functional impairment amenable to rehabilitation",
        "Age 18+ years; medically stable",
        "Written informed consent (patient and/or legally authorised representative)",
        "Baseline assessments completed (RBANS, MoCA, NSI, PCL-5, PHQ-9)",
    ],
    exclusion=[
        "Intracranial metallic hardware, cranioplasty hardware, or retained metallic fragments in stimulation path",
        "Active intracranial haemorrhage or coagulopathy",
        "Skull defects or craniectomy at electrode/transducer sites",
        "Cochlear implant or DBS device",
        "Post-traumatic epilepsy (seizure within past 12 months) \u2014 absolute exclusion",
        "Pregnancy (tDCS, TPS)",
        "Severe cognitive impairment precluding informed consent without proxy",
    ],
    discussion=[
        "Cardiac pacemaker or defibrillator \u2014 individual risk\u2013benefit assessment",
        "Epilepsy risk / sub-clinical seizures \u2014 higher post-TBI; EEG monitoring if clinical concern",
        "Active suicidality \u2014 psychiatric referral and safety planning before NIBS initiation",
        "PTSD comorbidity \u2014 trauma-informed approach required; tailored stimulation protocol",
        "Coagulation disorders or anticoagulants (especially TPS)",
        "Dermatological conditions or scalp injuries at electrode sites",
        "Blast TBI \u2014 polytrauma context; multidisciplinary coordination required",
        "Vestibular dysfunction \u2014 may be exacerbated; careful taVNS titration",
    ],
    overview=[
        "Traumatic Brain Injury (TBI) results from external mechanical force to the head, causing focal "
        "or diffuse brain injury. TBI is a leading cause of disability worldwide: approximately 69 million "
        "TBIs occur annually (Dewan et al., 2018). Sequelae range from transient post-concussion symptoms "
        "in mild TBI to profound cognitive, motor, and behavioural impairment in severe TBI.",
        "Chronic TBI sequelae commonly include post-concussion syndrome (PCS), cognitive impairment "
        "(attention, memory, executive function), post-traumatic depression and anxiety, PTSD, "
        "post-traumatic headache, sleep disturbance, and fatigue. Diffuse axonal injury (DAI) disrupts "
        "white matter tracts, impairing network connectivity across frontal-parietal and frontal-subcortical "
        "circuits.",
        "tDCS targeting the DLPFC (anodal) has shown improvements in working memory, attention, and "
        "executive function in TBI across multiple studies. The prefrontal cortex is consistently affected "
        "in TBI, making it the primary NIBS target (Dmochowski et al., 2013; Pape et al., 2020).",
        "NIBS Evidence in TBI: Level B evidence supports DLPFC tDCS for post-TBI cognitive rehabilitation. "
        "TPS is investigational. taVNS has plausible mechanisms for post-TBI mood, sleep, and cognitive "
        "deficits. CES is FDA-cleared for anxiety and mood, highly relevant to TBI comorbidities.",
    ],
    pathophysiology=[
        "TBI pathophysiology involves primary injury (diffuse axonal injury, contusion, haematoma) and "
        "secondary injury cascades (neuroinflammation, excitotoxicity, oxidative stress, oedema, ischaemia). "
        "DAI disrupts major white matter tracts (corpus callosum, long association fibres), impairing "
        "fronto-parietal networks, attention circuits, and default mode network connectivity.",
        "Chronic neuroinflammation, tau accumulation (risk for CTE in repetitive TBI), and impaired "
        "glymphatic clearance contribute to progressive neurodegeneration. The orbitofrontal cortex and "
        "prefrontal regions are selectively vulnerable due to bony ridges causing contre-coup injury. "
        "Hypothalamic-pituitary dysfunction is common in moderate-severe TBI, affecting hormonal "
        "regulation and sleep.",
    ],
    std_treatment=[
        "TBI rehabilitation is multidisciplinary: cognitive rehabilitation, neuropsychological therapy, "
        "physical and occupational therapy, speech-language therapy, and mental health management. "
        "Pharmacological management targets specific symptoms: antidepressants for PCS mood symptoms, "
        "amantadine for consciousness and cognition, stimulants (methylphenidate, amantadine) for "
        "attention in chronic TBI. NIBS is adjunctive to comprehensive rehabilitation.",
    ],
    symptoms=[
        ["Cognitive Impairment",     "Attention, working memory, processing speed, executive function deficits",  "Most common",      "Dewan 2018; Pape 2020"],
        ["Post-Concussion Headache", "Tension-type, migraine-like, or mixed headache post-TBI",                   "50\u201390% (mTBI)","Lew 2006"],
        ["Fatigue",                  "Post-TBI mental and physical fatigue; often most disabling symptom",         "50\u201378%",      "Lew 2006"],
        ["Sleep Disturbance",        "Insomnia, hypersomnia, circadian disruption, sleep apnoea risk",             "50\u201380%",      "Castriotta 2011"],
        ["Depression",               "Post-TBI depression; often comorbid with PCS, PTSD, and anxiety",           "25\u201350%",      "Jorge 2016"],
        ["Anxiety / PTSD",           "Post-traumatic anxiety, PTSD in 10\u201340% of TBI survivors",              "10\u201340%",      "Bryant 2010"],
        ["Irritability / Aggression","Emotional dysregulation, aggression, impulsivity; frontal disinhibition",    "30\u201370%",      "Alderman 2003"],
        ["Balance / Vestibular",     "Dizziness, vestibular dysfunction, balance impairment",                      "30\u201365%",      "Lew 2006"],
        ["Sensory Sensitivity",      "Light sensitivity (photophobia), sound sensitivity (phonophobia)",           "50\u201370% (mTBI)","Lew 2006"],
        ["Motor Deficits",           "Hemiplegia/hemiparesis, spasticity, coordination impairment (mod-severe TBI)","Variable (severity-dependent)","Dewan 2018"],
    ],
    brain_regions=[
        ["DLPFC (bilateral)",         "Most consistently affected by TBI; executive function, WM, attention",    "Bilateral DLPFC anodal tDCS primary target; Level B evidence",              "Dmochowski 2013; Pape 2020"],
        ["OFC / Prefrontal",          "Vulnerable to coup-contre-coup; emotional regulation, impulse control",   "Prefrontal tDCS; TPS frontal targeted (investigational)",                   "Alderman 2003"],
        ["Parietal Cortex",           "Attention, sensorimotor integration, spatial processing impairment",       "Parietal anodal tDCS + tDCS-cognitive training combination",                "Dmochowski 2013"],
        ["Temporal Cortex",           "Memory consolidation; language; disrupted in lateral TBI impacts",         "Temporal anodal tDCS + memory rehabilitation",                              "Pape 2020"],
        ["Cerebellum",                "Balance, coordination, vestibular integration; disrupted in TBI",          "Cerebellar tDCS for balance/vestibular rehabilitation (emerging)",          "Grimaldi 2016"],
        ["Anterior Cingulate Cortex", "Pain, attention, conflict monitoring; disrupted in frontal TBI",           "Prefrontal tDCS modulates ACC; TPS ACC targeted (investigational)",        "Lew 2006"],
        ["Corpus Callosum",           "Interhemispheric transfer; DAI primary target in TBI",                     "Not directly stimulated; integrity informs bihemispheric tDCS approach",   "Dewan 2018"],
        ["Hippocampus",               "Memory consolidation; vulnerable to excitotoxic damage post-TBI",          "TPS temporal/hippocampal (investigational); temporal tDCS indirect",        "Dewan 2018"],
    ],
    brainstem=[
        ["Locus Coeruleus (LC)",       "Noradrenergic dysregulation post-TBI; fatigue, attention, mood",          "taVNS modulates LC; augments noradrenergic tone for attention/mood",       "Aston-Jones 2005"],
        ["Nucleus Tractus Solitarius", "Vagal afferent relay; autonomic regulation post-TBI",                     "Primary target of taVNS; modulates arousal and cortical plasticity",       "Frangos 2015"],
        ["Dorsal Raphe Nucleus",       "Serotonergic dysregulation; depression, sleep, pain post-TBI",            "taVNS and tDCS (DLPFC) may modulate serotonergic projections",            "Jorge 2016"],
        ["Reticular Activating System","Arousal and consciousness; damaged in severe TBI (DAI)",                  "taVNS and tDCS indirectly modulate ascending arousal pathways",            "Moruzzi 1949"],
        ["Periaqueductal Gray (PAG)",  "Pain modulation; post-TBI headache circuitry",                            "Indirectly modulated via cortical and vagal neuromodulation",              "Lew 2006"],
    ],
    phenotypes=[
        ["Mild TBI / Post-Concussion Syndrome","Cognitive symptoms, headache, fatigue, sleep disturbance post-mTBI","Cognitive symptoms, headache, fatigue, emotional lability","L-DLPFC anodal tDCS; CES (sleep/headache); taVNS"],
        ["Moderate-Severe TBI \u2014 Cognitive","Significant cognitive impairment; attention, memory, EF deficits","Attention, memory, EF, processing speed impairment",             "Bilateral DLPFC tDCS; TPS frontal; cognitive rehab"],
        ["TBI + Depression",                    "Post-TBI depression as primary psychiatric comorbidity",          "Depressed mood, anhedonia, anergia, irritability",                "L-DLPFC anodal tDCS; taVNS; CES; antidepressant if indicated"],
        ["TBI + PTSD",                          "PTSD superimposed on TBI; complex comorbidity",                   "Re-experiencing, avoidance, hyperarousal, mood, cognition",       "Bilateral DLPFC tDCS; taVNS primary; CES; trauma-informed Rx"],
        ["TBI \u2014 Frontal Dysexecutive",     "Frontal lobe syndrome; disinhibition, impulsivity, EF deficit",  "Impulsivity, disinhibition, planning failure, aggression",        "DLPFC anodal + OFC tDCS; TPS frontal; DBT integration"],
        ["Blast TBI",                           "Military/combat blast TBI; polytrauma context; often mTBI",       "PCS + PTSD + auditory + vestibular + headache",                   "Bilateral DLPFC tDCS; taVNS primary; CES; MDT approach"],
        ["TBI + Chronic Pain / Headache",       "Persistent headache or chronic pain as primary burden",           "Headache, pain, fatigue, sleep disturbance, mood",                "M1 + DLPFC tDCS; taVNS; CES primary for headache/sleep"],
    ],
    symptom_map=[
        ["Attention / Processing Speed","DLPFC, ACC, Parietal",        "Bilateral DLPFC anodal tDCS + attention rehabilitation; TPS DLPFC (investigational)",       "Level B (Dmochowski 2013)"],
        ["Working Memory",              "DLPFC, Parietal, Temporal",   "L-DLPFC anodal tDCS + WM training (n-back); TPS DLPFC targeted",                          "Level B (Pape 2020)"],
        ["Executive Function",          "DLPFC, OFC, ACC",             "Bilateral DLPFC anodal tDCS + executive training; TPS frontal (investigational)",          "Level B (Pape 2020)"],
        ["Post-TBI Depression",         "L-DLPFC, sgACC",              "L-DLPFC anodal tDCS + psychotherapy; taVNS primary; CES adjunct",                        "Moderate (Jorge 2016)"],
        ["PTSD",                        "vmPFC, DLPFC, Amygdala",      "Bilateral DLPFC tDCS + trauma-informed therapy; taVNS primary; CES",                      "Moderate (Bryant 2010)"],
        ["Sleep Disturbance",           "Frontal, Thalamus, Circadian", "CES primary; taVNS; tDCS evening slow-wave",                                             "Moderate (Philip 2017)"],
        ["Fatigue",                     "DLPFC, SMA, Arousal systems",  "DLPFC tDCS + taVNS primary; CES; graded activity programme",                             "Moderate (Lew 2006)"],
        ["Post-TBI Headache",           "M1, ACC, Thalamus",            "M1 + DLPFC tDCS; CES primary for headache; taVNS; headache management",                  "Moderate (Lew 2006)"],
        ["Emotional Dysregulation",     "OFC, DLPFC, Amygdala",        "DLPFC anodal tDCS; taVNS; CES; DBT or emotion regulation therapy",                       "Moderate (Alderman 2003)"],
    ],
    montage=[
        ["Mild TBI / PCS",           "F3 anode + Fp2 cathode (L-DLPFC primary); or bilateral F3 + F4 anode","Cranial (global + targeted DLPFC)","taVNS, CES (headache/sleep)"],
        ["Moderate-Severe Cognitive","F3 + F4 bilateral anode (bilateral DLPFC); T3/T4 + P3/P4 addition",   "Cranial (global + targeted DLPFC/temporal/parietal)","taVNS, CES supportive"],
        ["TBI + Depression",         "F3 anode + Fp2 cathode (L-DLPFC anodal)",                             "Cranial (global + targeted L-DLPFC)",              "taVNS primary; CES"],
        ["TBI + PTSD",               "F3 + F4 bilateral anode; or F4 anode + F3 anode (bilateral)",         "Cranial (global + targeted bilateral DLPFC)",      "taVNS primary; CES; trauma-informed approach"],
        ["Frontal Dysexecutive",      "F3 anode + OFC area; bilateral DLPFC approach",                      "Cranial (global + targeted frontal circuits)",     "taVNS, CES; DBT integration"],
        ["Blast TBI",                 "Bilateral DLPFC (F3 + F4 anode); slower titration protocol",         "Cranial (global + targeted bilateral DLPFC)",      "taVNS primary; CES; MDT"],
        ["Chronic Pain / Headache",   "F3 anode + C3/C4 (contra to pain) cathode / extracephalic",          "Cranial + peripheral (pain dermatomes if relevant)","taVNS; CES primary (headache)"],
        ["No response after 8\u201310 sessions","Reassess lesion imaging; adjust montage; add TPS","TPS DLPFC frontal targeted; adjust ROI",             "Review all adjuncts; MDT referral"],
    ],
    tdcs_protocols=[
        ["C1","Cognitive (attention, WM) \u2014 primary","F3 anode + F4 anode (bilateral)","Fp1 + Fp2 / extracephalic","2 mA, 20\u201330 min, 5\u00d7/wk \u00d7 6\u201310 wks","Level B; attention/WM (Dmochowski 2013; Pape 2020)"],
        ["C2","L-DLPFC (attention/memory)",               "F3 anode",            "Fp2 / extracephalic",          "2 mA, 20\u201330 min, 5\u00d7/wk",             "Level B (Pape 2020)"],
        ["C3","Executive function",                        "F3 + F4 bilateral",   "Fp1 + Fp2 / extracephalic",   "2 mA, 20\u201330 min, 5\u00d7/wk + EF training","Level B (Pape 2020)"],
        ["C4","Post-TBI depression",                       "F3 anode",            "Fp2 / extracephalic",          "2 mA, 20\u201330 min, 5\u00d7/wk",             "Moderate (Jorge 2016)"],
        ["C5","TBI + PTSD",                                "F3 + F4 bilateral",   "Extracephalic shoulders",      "2 mA, 20\u201330 min, 5\u00d7/wk",             "Moderate (Bryant 2010)"],
        ["C6","Sleep disturbance",                         "Fz anode",            "Pz / extracephalic",           "1\u20132 mA, 20 min, evening",                  "Moderate (Philip 2017)"],
        ["C7","Post-TBI headache + pain",                  "F3 anode + C3/C4","F4 / extracephalic",              "2 mA, 20\u201330 min, 5\u00d7/wk",             "Moderate (Lew 2006)"],
        ["C8","Emotional dysregulation",                   "F3 + F4 bilateral",   "Extracephalic shoulders",      "2 mA, 20\u201330 min, 5\u00d7/wk + DBT",       "Moderate (Alderman 2003)"],
    ],
    plato_protocols=[
        ["C1-PS","Cognitive (attention/WM)", "Frontal (F3 area)","F3","Right shoulder",  "1.6 mA, 20\u201330 min","5\u00d7/wk","Bilateral approach: HDCkit preferred"],
        ["C2-PS","L-DLPFC attention",        "Frontal (F3 area)","F3","Right shoulder",  "1.6 mA, 20\u201330 min","5\u00d7/wk","Pair with attention tasks"],
        ["C3-PS","Executive function",       "Frontal (F3 area)","F3","Shoulder",        "1.6 mA, 20\u201330 min","5\u00d7/wk","Combine with executive training"],
        ["C4-PS","Post-TBI depression",      "Frontal (F3 area)","F3","Right shoulder",  "1.6 mA, 20\u201330 min","5\u00d7/wk","Adjunct taVNS and CES recommended"],
        ["C5-PS","TBI + PTSD",               "Frontal (F3 area)","F3","Shoulder",        "1.6 mA, 20\u201330 min","5\u00d7/wk","Trauma-informed session approach"],
        ["C6-PS","Sleep disturbance",        "Frontal (Fz area)","Fz","Shoulder",        "1.6 mA, 20 min",        "Evening",    "CES concurrent recommended"],
        ["C7-PS","Headache + pain",          "Motor (Cz area)",  "Cz","Shoulder/arm",    "1.6 mA, 20\u201330 min","5\u00d7/wk","CES primary for headache adjunct"],
        ["C8-PS","Emotional dysregulation",  "Frontal (F3 area)","F3","Shoulder",        "1.6 mA, 20\u201330 min","5\u00d7/wk","DBT therapy concurrent"],
    ],
    tps_protocols=[
        ["T1","Cognitive rehabilitation (primary)","Bilateral DLPFC, ACC, mPFC",   "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz; 10 sessions/2 wks","4,000 Tgt (DLPFC bilateral) + 2,000 PRE; total 6\u20138K","Investigational"],
        ["T2","Frontal dysexecutive",               "DLPFC + OFC + ACC",            "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                  "4,000 Tgt + 2,000 Std; total 6\u20138K",             "Investigational"],
        ["T3","Post-TBI depression / PTSD",         "L-DLPFC + ACC + mPFC",        "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                  "4,000 Tgt (L-DLPFC) + 2,000 PRE; total 6\u20138K",  "Investigational"],
        ["T4","Memory rehabilitation",              "DLPFC + temporal cortex",      "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                  "4,000 Tgt + 2,000 PRE; total 6\u20138K",             "Investigational"],
        ["T5","Post-TBI pain / headache",           "M1 + ACC + Thalamus",          "0.20\u20130.25 mJ/mm\u00b2; 1\u20136 Hz",                  "4,000 Tgt + 2,000 PRE; total 6\u20138K",             "Investigational"],
    ],
    ces_role=[
        ["Sleep disturbance",      "Primary target: post-TBI insomnia; circadian disruption; nightmares; FDA-cleared",  "Evening before bed; 20\u201340 min"],
        ["Post-TBI depression",    "Supportive adjunct for PCS mood symptoms; FDA-cleared for mood",                    "Evening sessions or non-tDCS days; 20\u201340 min"],
        ["Anxiety / PTSD",         "FDA-cleared for anxiety; reduces physiological anxiety symptoms",                   "Before tDCS session or evening; 20\u201340 min"],
        ["Post-TBI headache",      "May reduce headache frequency and intensity as adjunct to primary protocols",        "As needed; 20\u201330 min; avoid during acute headache"],
    ],
    tavns_role="taVNS modulates locus coeruleus noradrenergic tone and autonomic balance relevant to post-TBI fatigue, mood, and cognitive deficits. It also modulates inflammatory cascades, potentially relevant to secondary neuroinflammation in TBI.",
    combinations=[
        ("1) Mild TBI / PCS", [
            ["tDCS + Cognitive Rehabilitation","L-DLPFC anodal tDCS + computerised attention and cognitive rehabilitation; most evidence-based TBI combination (Pape 2020; Dmochowski 2013)","tDCS before/during cognitive tasks","Attention, WM, processing speed, cognitive fatigue"],
            ["taVNS + tDCS","taVNS primes noradrenergic tone before tDCS; supports attention and mood recovery (Hein 2021; Aston-Jones 2005)","taVNS 20\u201330 min before tDCS","Cognitive fatigue, emotional dysregulation, autonomic symptoms"],
            ["CES + Standard Care","CES primary for post-TBI headache, sleep, and anxiety; FDA-cleared; well-tolerated in mTBI (Philip 2017)","Daily as prescribed; evening primary","PCS with headache, sleep disturbance, anxiety"],
        ]),
        ("2) TBI + Depression", [
            ["tDCS (DLPFC) + Psychotherapy","L-DLPFC anodal tDCS + CBT or trauma-informed therapy; addresses post-TBI depression (Jorge 2016)","tDCS before therapy session","Post-TBI depressed mood, anhedonia, anergia"],
            ["taVNS + tDCS","taVNS for autonomic and mood modulation; augments antidepressant tDCS effect (Hein 2021)","taVNS before tDCS","Depression with autonomic dysregulation and fatigue"],
            ["CES + Standard Care","CES for sleep, anxiety, and mood; integral to TBI depression management (Philip 2017)","Evening daily; morning for energy","Post-TBI depression with insomnia and anxiety"],
        ]),
        ("3) TBI + PTSD", [
            ["tDCS + Trauma-Informed Therapy","Bilateral DLPFC tDCS + EMDR or CPT; enhances prefrontal regulation of trauma responses (Bryant 2010)","tDCS before trauma-focused therapy","PTSD symptoms superimposed on TBI"],
            ["taVNS + Therapy","taVNS reduces physiological threat response before trauma therapy; key for hyperarousal (Hein 2021)","taVNS 20\u201330 min before therapy session","Hyperarousal, autonomic reactivity, sleep disturbance"],
            ["CES + Standard Care","CES FDA-cleared for anxiety; reduces hyperarousal symptoms in PTSD-TBI (Philip 2017)","Daily as prescribed","PTSD with severe anxiety, insomnia, and physiological arousal"],
        ]),
        ("4) TBI + Frontal Dysexecutive", [
            ["tDCS + Executive Training","Bilateral DLPFC anodal tDCS + executive function training (planning, WM, flexibility) (Pape 2020)","tDCS before/during executive training","Disinhibition, impulsivity, planning failure"],
            ["taVNS + DBT","taVNS for emotional dysregulation; DBT for skills-based regulation of impulsivity (Hein 2021)","taVNS before DBT session","Frontal disinhibition, emotional lability, aggression"],
            ["CES + Standard Care","CES for sleep and baseline arousal stabilisation; reduces emotional threshold (Philip 2017)","As needed; morning or evening","Emotional dysregulation with sleep and anxiety burden"],
        ]),
        ("5) TBI + Chronic Pain / Headache", [
            ["tDCS (M1 + DLPFC) + Pain Management","M1 + DLPFC tDCS + graded activity/pain psychology; top-down pain modulation (Lew 2006; Fregni 2011)","tDCS before activity or pain management session","Chronic post-TBI headache, central pain, musculoskeletal pain"],
            ["taVNS + Standard Care","taVNS modulates ascending pain pathways and autonomic arousal; analgesic adjunct (Adair 2020)","Before flare-prone periods or on off-days","Pain with autonomic arousal, anxiety-pain loop"],
            ["CES + Standard Care","CES reduces post-TBI headache, sleep disturbance, and anxiety-pain amplification (Philip 2017)","Daily; evening primary for sleep","Post-TBI headache with insomnia and anxiety"],
        ]),
        ("6) Blast TBI", [
            ["Bilateral tDCS + MDT Rehabilitation","Bilateral DLPFC tDCS + multidisciplinary rehabilitation (cognitive, physical, psychological); standard blast TBI approach (Pape 2020)","tDCS during cognitive/physical rehabilitation session","Blast TBI: PCS + PTSD + vestibular + headache"],
            ["taVNS + Bilateral tDCS","taVNS primary for autonomic hyperarousal and PTSD overlay; augments cognitive tDCS (Hein 2021)","taVNS before tDCS; gentler titration","Blast TBI with PTSD and autonomic features"],
            ["CES + Standard Care","CES daily for sleep, anxiety, and headache; essential in blast TBI management (Philip 2017)","Daily as prescribed","Blast TBI with sleep disturbance, anxiety, headache burden"],
        ]),
    ],
    combination_summary=[
        ["Mild TBI / PCS",         "tDCS + Cognitive Rehab",        "L-DLPFC anodal tDCS + cognitive rehabilitation (Pape 2020; Dmochowski 2013)",              "tDCS before/during tasks", "Attention, WM, cognitive fatigue", "Level B"],
        ["TBI + Depression",       "tDCS + Psychotherapy",           "L-DLPFC tDCS + CBT/trauma-informed therapy (Jorge 2016)",                                 "tDCS before therapy",       "Post-TBI depression",              "Moderate"],
        ["TBI + PTSD",             "tDCS + Trauma Therapy",          "Bilateral DLPFC tDCS + EMDR/CPT; prefrontal regulation of trauma (Bryant 2010)",            "tDCS before therapy",       "PTSD superimposed on TBI",         "Moderate"],
        ["Frontal Dysexecutive",   "tDCS + Executive Training",      "Bilateral DLPFC tDCS + EF training (planning, WM, flexibility) (Pape 2020)",               "tDCS before EF training",   "Disinhibition, impulsivity, EF",   "Level B"],
        ["TBI + Headache/Pain",    "tDCS (M1+DLPFC) + Pain Management","M1 + DLPFC tDCS + graded activity/pain psychology (Lew 2006; Fregni 2011)",               "tDCS before pain session",  "Chronic post-TBI headache, pain",  "Moderate"],
        ["Blast TBI",              "Bilateral tDCS + MDT",            "Bilateral DLPFC tDCS + multidisciplinary rehabilitation (Pape 2020)",                       "tDCS during MDT",           "Blast TBI: PCS + PTSD",            "Moderate"],
    ],
    outcomes=[
        ["RBANS",                  "Comprehensive neuropsychological battery",  "Baseline, months 1, 3",        "Normative z-scores; track improvement in attention, memory, EF"],
        ["MoCA",                   "Global cognition screening",                 "Baseline, months 1, 3",        "Score <26 cognitive impairment"],
        ["NSI",                    "Neurobehavioral Symptom Inventory (PCS)",    "Baseline, weeks 4, 8, 12",     "Score 0\u201388; higher = greater symptom burden; track improvement"],
        ["PHQ-9",                  "Post-TBI depression",                        "Baseline, weeks 4, 8, 12",     "\u226510 moderate depression"],
        ["PCL-5",                  "PTSD Checklist (DSM-5)",                     "Baseline, weeks 4, 8, 12",     "Score \u226533 probable PTSD; track \u226510-point improvement"],
        ["ISI",                    "Insomnia severity",                          "Baseline, weeks 4, 8, 12",     "Score \u226715 moderate insomnia"],
        ["GAD-7",                  "Comorbid anxiety",                           "Baseline, weeks 4, 8, 12",     "\u226510 moderate anxiety"],
        ["Timed Up and Go (TUG)", "Balance and motor function",                 "Baseline, months 1, 3",        ">14 seconds elevated fall risk (TBI/vestibular context)"],
        ["SOZO PRS",               "NIBS-specific functional outcome",          "Baseline, weeks 2, 4, 8, 12", "Proprietary; composite cognitive/mood/sleep/functional domains"],
    ],
    outcomes_abbrev=["RBANS", "MoCA", "NSI", "PHQ-9", "SOZO PRS"],
)

# ---------------------------------------------------------------------------
# CONDITION: CHRONIC PAIN
# ---------------------------------------------------------------------------
CONDITIONS["chronic_pain"] = dict(
    name="Chronic Pain Syndrome",
    icd10="G89.29",
    doc_num="SPG-CP-001",
    tps_status="Investigational",
    tdcs_class="Class IIa",
    tps_class="Investigational",
    tavns_class="Class IIb",
    ces_class="Class IIa",
    inclusion=[
        "Chronic pain ≥3 months duration meeting IASP criteria",
        "Diagnosis of chronic pain syndrome (ICD-10 G89.29) confirmed by pain specialist",
        "NRS pain score ≥4/10 at baseline on at least 3 consecutive days",
        "Inadequate response to ≥2 pharmacological treatments (anticonvulsants, SNRIs, opioids)",
        "Age 18–75 years; capable of informed consent",
        "Stable medication regimen for ≥4 weeks prior to enrollment",
    ],
    exclusion=[
        "Active malignancy or cancer-related pain",
        "Metallic implants in head/neck region (TPS/tDCS contraindication)",
        "Severe psychiatric comorbidity (active psychosis, suicide attempt <6 months)",
        "Pregnancy or planned pregnancy",
        "Active substance use disorder (excluding cannabis with stable use)",
        "Implanted neurostimulator or cardiac pacemaker",
        "Fibromyalgia as sole diagnosis without central sensitization evidence",
    ],
    discussion=[
        "Chronic pain involves maladaptive neuroplasticity across nociceptive and central sensitization pathways.",
        "The anterior cingulate cortex (ACC) and insula are primary hubs for pain affect and interoception.",
        "Thalamocortical dysrhythmia and reduced descending inhibition are core pathophysiological features.",
        "tDCS targeting M1 (anodal) reduces pain intensity through cortical excitability modulation.",
        "TPS over ACC and prefrontal regions addresses the affective-motivational dimension of pain.",
        "taVNS activates the nucleus tractus solitarius and locus coeruleus, enhancing descending inhibition.",
        "CES (Alpha-Stim) modulates thalamic relay nuclei and reduces central sensitization.",
        "Combination NIBS produces synergistic analgesia through convergent network normalization.",
    ],
    overview=[
        "Chronic pain syndrome represents a complex biopsychosocial condition characterized by persistent pain lasting beyond expected tissue healing time (>3 months), accompanied by neurobiological, psychological, and social consequences. Globally, chronic pain affects 20-30% of adults, generating immense disability and healthcare burden.",
        "The SOZO NIBS protocol targets the neurobiological underpinnings of chronic pain: maladaptive central sensitization, thalamocortical dysrhythmia, impaired descending pain inhibition, and dysfunctional salience processing. The protocol integrates tDCS (M1/DLPFC), TPS (ACC/insula), taVNS (vagus nerve), and CES (Alpha-Stim) in a phenotype-guided multimodal framework.",
        "Treatment is organized around six clinical phenotypes: neuropathic dominant, central sensitization dominant, affective-pain dominant, musculoskeletal-central mixed, catastrophizing dominant, and sleep-pain comorbid. Each phenotype receives a tailored NIBS combination targeting the most relevant neural circuits.",
    ],
    pathophysiology=[
        "Central sensitization: wind-up of spinal and supraspinal nociceptive neurons via NMDA-receptor-mediated LTP leads to hyperalgesia and allodynia.",
        "Thalamocortical dysrhythmia: abnormal alpha/theta oscillations in somatosensory thalamus maintain persistent pain percept independent of peripheral input.",
        "Reduced descending inhibition: impaired noradrenergic and serotonergic projections from PAG, RVM, and locus coeruleus reduce endogenous analgesia.",
        "ACC and insula hyperactivation: affective-motivational pain processing is amplified, driving catastrophizing and pain-related fear.",
        "DLPFC hypoactivity: reduced cognitive control over pain allows salience network dominance and rumination.",
        "DMN-SN dysconnection: default mode network is invaded by salience network, perpetuating pain catastrophizing and self-referential suffering.",
    ],
    std_treatment=[
        "Pharmacological: SNRIs (duloxetine), anticonvulsants (gabapentin, pregabalin), tricyclic antidepressants, opioids (limited), topical agents",
        "Interventional: nerve blocks, spinal cord stimulation, intrathecal therapy",
        "Psychological: CBT-based pain management, acceptance and commitment therapy (ACT), mindfulness",
        "Physical rehabilitation: graded motor imagery, pain neuroscience education, exercise therapy",
        "Multidisciplinary pain programs: integrated biopsychosocial rehabilitation",
    ],
    symptoms=[
        ["Spontaneous Pain", "Ongoing, burning, aching, stabbing", "100%", "IASP 2020"],
        ["Hyperalgesia", "Exaggerated response to painful stimuli", "75%", "Woolf 2011"],
        ["Allodynia", "Pain from non-painful stimuli (touch, pressure)", "60%", "Costigan 2009"],
        ["Sleep Disturbance", "Insomnia, non-restorative sleep", "80%", "Moldofsky 2010"],
        ["Fatigue", "Persistent fatigue disproportionate to activity", "75%", "Cleeland 2010"],
        ["Depression", "Comorbid major depression or dysthymia", "50%", "McWilliams 2003"],
        ["Anxiety", "Pain-related fear, health anxiety, GAD features", "45%", "Asmundson 2009"],
        ["Cognitive Dysfunction", "Reduced attention, memory, processing speed (pain fog)", "55%", "Moriarty 2011"],
        ["Catastrophizing", "Magnification, rumination, helplessness", "60%", "Sullivan 2001"],
        ["Reduced Mobility", "Activity avoidance, deconditioning", "70%", "Vlaeyen 2000"],
    ],
    brain_regions=[
        ["Primary Somatosensory Cortex (S1)", "Sensory-discriminative pain processing", "tDCS anodal M1 adjacent", "Fregni 2006"],
        ["Motor Cortex (M1)", "Descending motor-pain modulation, excitability regulation", "tDCS anodal M1", "Lefaucheur 2016"],
        ["Anterior Cingulate Cortex (ACC)", "Affective-motivational pain processing, catastrophizing", "TPS, tDCS cathodal", "Rainville 1997"],
        ["Insula", "Interoception, pain salience, autonomic integration", "TPS", "Craig 2002"],
        ["DLPFC", "Cognitive modulation of pain, placebo analgesia", "tDCS anodal", "Seminowicz 2013"],
        ["PAG/RVM", "Descending inhibitory control", "taVNS indirect", "Fields 2004"],
        ["Thalamus", "Thalamocortical relay, dysrhythmia hub", "CES", "Llinás 1999"],
        ["Basal Ganglia", "Reward-pain interaction, motivation-pain coupling", "TPS", "Navratilova 2012"],
    ],
    brainstem=[
        ["Periaqueductal Gray (PAG)", "Origin of descending inhibition, opioidergic modulation", "taVNS", "Basbaum 1984"],
        ["Rostral Ventromedial Medulla (RVM)", "ON/OFF cell balance for facilitation vs inhibition", "taVNS", "Fields 2004"],
        ["Locus Coeruleus (LC)", "Noradrenergic descending inhibition", "taVNS/CES", "Jones 1991"],
        ["Nucleus Raphe Magnus", "Serotonergic descending modulation", "taVNS", "Basbaum 1984"],
        ["Nucleus Tractus Solitarius (NTS)", "Vagal integration, pain gating", "taVNS direct", "Tracey 2002"],
    ],
    phenotypes=[
        ["Neuropathic Dominant", "Burning, lancinating pain; allodynia; positive sensory signs", "Hyperalgesia, paresthesias, sleep disruption", "tDCS M1 anodal + TPS"],
        ["Central Sensitization", "Widespread pain; fatigue; cognitive fog; no peripheral lesion", "Allodynia, hyperalgesia, fibromyalgia features", "tDCS DLPFC + TPS ACC + CES"],
        ["Affective-Pain Dominant", "High catastrophizing; depression-pain comorbidity; pain affect disproportionate", "Low mood, pain rumination, social withdrawal", "tDCS DLPFC anodal + TPS ACC + taVNS"],
        ["Musculoskeletal-Central Mixed", "Structural pathology + central amplification; inconsistent response to local treatment", "MSK pain + central features", "tDCS M1 + TPS + CES"],
        ["Catastrophizing Dominant", "High PCS scores; extreme disability relative to objective findings", "Helplessness, magnification, fear-avoidance", "tDCS DLPFC + taVNS + CBT co-treatment"],
        ["Sleep-Pain Comorbid", "Insomnia drives pain amplification; bidirectional relationship", "Non-restorative sleep, fatigue, pain escalation", "CES primary + tDCS M1 + taVNS"],
    ],
    symptom_map=[
        ["Spontaneous pain", "M1/S1", "tDCS anodal M1", "Lefaucheur 2016 Class IIa"],
        ["Allodynia/Hyperalgesia", "Thalamus/S1", "CES + tDCS", "Antal 2010 Class IIb"],
        ["Affective pain (suffering)", "ACC/Insula", "TPS + taVNS", "Rainville 1997"],
        ["Catastrophizing", "DLPFC/ACC", "tDCS DLPFC anodal", "Seminowicz 2013"],
        ["Sleep disruption", "Thalamus/Cortex", "CES Alpha-Stim", "Lande 2018"],
        ["Fatigue", "CEN/SN", "taVNS + tDCS DLPFC", "Vonck 2014"],
        ["Depression comorbidity", "DLPFC/SGC", "tDCS DLPFC + taVNS", "Brunoni 2013"],
        ["Cognitive fog", "DLPFC/CEN", "tDCS DLPFC anodal", "Brunoni 2016"],
        ["Descending inhibition deficit", "PAG/LC/RVM", "taVNS", "Busch 2013"],
    ],
    montage=[
        ["Neuropathic Dominant", "Anodal M1 contralateral (2mA, 20min)", "ACC (300 pulses)", "Gabapentin co-treatment"],
        ["Central Sensitization", "Anodal DLPFC (2mA, 20min)", "ACC + S1 (400 pulses)", "CES Alpha-Stim nightly"],
        ["Affective-Pain", "Anodal DLPFC (2mA, 20min)", "ACC (300 pulses)", "taVNS 30min pre-session"],
        ["Musculoskeletal-Central", "Anodal M1 (2mA) + DLPFC", "S1/M1 (300 pulses)", "CES 3x/week"],
        ["Catastrophizing Dominant", "Anodal DLPFC (2mA, 20min)", "ACC (300 pulses)", "Concurrent CBT sessions"],
        ["Sleep-Pain Comorbid", "Anodal M1 (1.5mA, 20min)", "None initial", "CES Alpha-Stim primary nightly"],
        ["Refractory Pain", "Bifrontal tDCS (2mA)", "Full protocol (500 pulses)", "taVNS 4x/day"],
        ["Elderly/Fragile", "Anodal M1 (1mA, 15min)", "TPS low (200 pulses)", "CES primary; taVNS standard"],
        ["Pediatric (off-label)", "Not recommended <18", "Not recommended <18", "CES only per specialist"],
        ["Post-Surgical", "DLPFC anodal (2mA, 20min)", "None acute phase", "taVNS + CES"],
    ],
    tdcs_protocols=[
        ["CP-tDCS-01", "M1 Anodal (Neuropathic)", "C3/C4 contralateral", "Fp2", "2mA, 20min, 10 sessions", "Lefaucheur 2016"],
        ["CP-tDCS-02", "DLPFC Anodal (Affective/Catastrophizing)", "F3 (left DLPFC)", "F4", "2mA, 20min, 10 sessions", "Brunoni 2013"],
        ["CP-tDCS-03", "Bifrontal (Refractory)", "F3 anodal", "F4 cathodal", "2mA, 20min, 10–20 sessions", "Fregni 2006"],
        ["CP-tDCS-04", "M1 + DLPFC Combined", "C3 anodal; F3 anodal sequential", "Fp2; F4", "2mA each, 20min, 15 sessions", "Roizenblatt 2007"],
        ["CP-tDCS-05", "S1 Cathodal (Central Sensitization)", "Fp2", "C3/C4", "1.5mA, 20min, 10 sessions", "Antal 2010"],
        ["CP-tDCS-06", "Maintenance Phase", "F3 anodal (DLPFC)", "F4", "1.5mA, 20min, 2x/week", "Bikson 2016"],
        ["CP-tDCS-07", "Elderly Adjusted", "C3 anodal M1", "Fp2", "1mA, 15min, 10 sessions", "Lefaucheur 2016"],
        ["CP-tDCS-08", "Sham Protocol", "F3 anodal", "F4", "30s ramp, 0mA active, 20min", "Brunoni 2013"],
    ],
    plato_protocols=[
        ["CP-PLATO-01", "Neuropathic M1", "C3/C4", "M1 contralateral", "C3", "2.0mA", "N/A", "Continuous 20min"],
        ["CP-PLATO-02", "DLPFC Anodal", "F3/F4", "Left DLPFC", "F3", "2.0mA", "N/A", "Continuous 20min"],
        ["CP-PLATO-03", "HD-tDCS Pain Ring", "C3 center", "4-electrode ring", "C3", "1.0mA", "N/A", "HD mode 20min"],
        ["CP-PLATO-04", "Bifrontal", "F3/F4", "Bifrontal", "F3", "2.0mA", "N/A", "Sequential F3 then F4"],
        ["CP-PLATO-05", "Maintenance", "F3", "DLPFC", "F3", "1.5mA", "N/A", "2x/week 20min"],
        ["CP-PLATO-06", "Gamma tACS Pain", "C3", "M1", "C3", "1.5mA", "40Hz", "tACS mode 20min"],
        ["CP-PLATO-07", "Elderly Adjusted", "C3", "M1", "C3", "1.0mA", "N/A", "15min continuous"],
        ["CP-PLATO-08", "Sham", "F3", "DLPFC", "F3", "0mA", "N/A", "Ramp only 30s"],
    ],
    tps_protocols=[
        ["CP-TPS-01", "ACC Affective Pain", "ACC (BA24/32)", "300 pulses, 0.25Hz, 0.2mJ/mm²", "3 sessions/week", "Leinenga 2022"],
        ["CP-TPS-02", "S1/M1 Sensory-Discriminative", "S1, M1 contralateral", "300 pulses, 0.25Hz", "3 sessions/week", "Fregni 2006"],
        ["CP-TPS-03", "DLPFC Cognitive Pain", "DLPFC (BA46)", "250 pulses, 0.2Hz", "2 sessions/week", "Seminowicz 2013"],
        ["CP-TPS-04", "Insula Interoceptive", "Posterior Insula", "200 pulses, 0.2Hz", "2 sessions/week", "Craig 2002"],
        ["CP-TPS-05", "Full Protocol (Refractory)", "ACC + S1 + DLPFC sequential", "500 pulses total", "5 sessions/week x2 weeks", "Mulak 2023"],
    ],
    ces_role=[
        ["Central Sensitization", "Primary modality to reduce thalamic dysrhythmia and sensitization", "Nightly 60min + NIBS daytime"],
        ["Sleep-Pain Comorbid", "Primary modality; improves sleep architecture and pain threshold", "Nightly 60min; titrate to sleep quality"],
        ["Adjunct Analgesia", "Additive analgesia with tDCS M1 or TPS", "Prior to or concurrent with tDCS/TPS"],
        ["Maintenance Phase", "Sustain gains; reduce relapse; patient self-administered", "3–5x/week; patient-managed"],
    ],
    tavns_role="taVNS (tragus, 25Hz, 200µs, 0.5mA, 30min) activates the nucleus tractus solitarius, projecting to PAG and locus coeruleus to enhance descending noradrenergic inhibition. Particularly effective in affective-pain and catastrophizing phenotypes. Administered prior to tDCS/TPS sessions. Home device (Nemos/TENS) 2x daily for maintenance.",
    combinations=[
        ("Neuropathic Dominant", [
            ["tDCS M1 + TPS S1", "Convergent sensory-motor cortex normalization", "TPS first 20min, tDCS immediately after", "Peripheral neuropathic pain with central features"],
            ["tDCS M1 + CES", "M1 excitability + thalamic dysrhythmia correction", "Concurrent or sequential same session", "Neuropathic pain with sleep disruption"],
            ["tDCS M1 + taVNS", "Cortical excitability + descending inhibition", "taVNS 30min before tDCS", "Refractory neuropathic pain"],
        ]),
        ("Central Sensitization", [
            ["tDCS DLPFC + TPS ACC + CES", "Cognitive + affective + thalamic — triple-network", "TPS → tDCS → CES nightly", "Widespread pain, fibromyalgia features"],
            ["CES + taVNS", "Thalamic + vagal anti-sensitization", "Concurrent or taVNS before CES", "Severe sensitization, allodynia"],
        ]),
        ("Affective-Pain Dominant", [
            ["tDCS DLPFC + TPS ACC + taVNS", "Cognitive control + affective hub + vagal", "taVNS → TPS → tDCS sequential", "Depression-pain comorbidity, catastrophizing"],
            ["tDCS DLPFC + taVNS", "Anti-depressant + vagal anti-nociception", "taVNS 30min before tDCS", "High catastrophizing, depression-pain"],
        ]),
        ("Musculoskeletal-Central", [
            ["tDCS M1 + TPS S1 + CES", "MSK pain control + sensitization reversal", "TPS → tDCS → CES nightly", "MSK pain with central amplification"],
        ]),
        ("Catastrophizing Dominant", [
            ["tDCS DLPFC + taVNS + CBT", "Prefrontal upregulation + vagal + psychological", "CBT session day; taVNS + tDCS other days", "High PCS scores, disability-pain mismatch"],
        ]),
        ("Sleep-Pain Comorbid", [
            ["CES + tDCS M1 + taVNS", "Sleep architecture + pain threshold + vagal", "CES nightly; tDCS + taVNS daytime", "Insomnia driving pain amplification"],
        ]),
    ],
    combination_summary=[
        ["Neuropathic", "tDCS M1 + TPS S1", "Sensory-motor convergence", "TPS → tDCS", "Peripheral neuropathic", "IIa"],
        ["Central Sensitization", "tDCS DLPFC + TPS ACC + CES", "Triple-network sensitization reversal", "TPS → tDCS → CES nightly", "Widespread/fibromyalgia", "IIb"],
        ["Affective-Pain", "tDCS DLPFC + TPS ACC + taVNS", "Cognitive-affective-vagal", "taVNS → TPS → tDCS", "Depression-pain comorbid", "IIb"],
        ["MSK-Central", "tDCS M1 + TPS S1 + CES", "MSK + sensitization", "TPS → tDCS → CES", "Structural + central", "IIb"],
        ["Catastrophizing", "tDCS DLPFC + taVNS + CBT", "Prefrontal + vagal + psychological", "Integrated schedule", "High PCS, fear-avoidance", "IIb"],
        ["Sleep-Pain", "CES + tDCS M1 + taVNS", "Sleep + analgesia + vagal", "CES nightly + daytime NIBS", "Insomnia-pain cycle", "IIb"],
    ],
    outcomes=[
        ["NRS Pain Score", "Pain Intensity", "Weekly", "≥30% reduction from baseline"],
        ["Brief Pain Inventory (BPI)", "Pain Interference", "Baseline, 4wk, 8wk", "≥30% improvement in interference subscale"],
        ["Pain Catastrophizing Scale (PCS)", "Catastrophizing", "Baseline, 4wk, 8wk", "≥30% reduction in total score"],
        ["PROMIS Pain Interference", "Functional impact", "Baseline, 4wk, 8wk", "≥5-point T-score improvement"],
        ["Pittsburgh Sleep Quality Index (PSQI)", "Sleep quality", "Baseline, 4wk, 8wk", "≥3-point reduction (score <5 target)"],
        ["DASS-21", "Depression/Anxiety/Stress", "Baseline, 4wk, 8wk", "≥30% reduction in depression subscale"],
        ["Quantitative Sensory Testing (QST)", "Central sensitization markers", "Baseline, 8wk", "Normalization of pressure pain threshold"],
        ["Patient Global Impression of Change (PGIC)", "Overall treatment response", "4wk, 8wk", "Rating ≥5/7 (much improved or better)"],
        ["Medication Use Log", "Opioid/analgesic use", "Weekly", "≥20% reduction in morphine equivalents"],
    ],
    outcomes_abbrev=[
        "NRS: Numeric Rating Scale",
        "BPI: Brief Pain Inventory",
        "PCS: Pain Catastrophizing Scale",
        "PROMIS: Patient-Reported Outcomes Measurement Information System",
        "PSQI: Pittsburgh Sleep Quality Index",
        "DASS-21: Depression Anxiety Stress Scales",
        "QST: Quantitative Sensory Testing",
        "PGIC: Patient Global Impression of Change",
        "PPT: Pressure Pain Threshold",
    ],
)

# ---------------------------------------------------------------------------
# CONDITION: PTSD
# ---------------------------------------------------------------------------
CONDITIONS["ptsd"] = dict(
    name="Post-Traumatic Stress Disorder",
    icd10="F43.10",
    doc_num="SPG-PT-001",
    tps_status="Investigational",
    tdcs_class="Class IIb",
    tps_class="Investigational",
    tavns_class="Class IIa",
    ces_class="Class IIa",
    inclusion=[
        "DSM-5 PTSD diagnosis confirmed by structured clinical interview (CAPS-5 score ≥33)",
        "Trauma exposure at least 3 months prior; symptom onset after trauma",
        "PCL-5 score ≥33 at baseline",
        "Inadequate response to ≥1 evidence-based psychotherapy (PE, CPT, EMDR) or SSRI/SNRI",
        "Age 18–70; capable of informed consent and trauma narrative engagement",
        "Stable medication regimen for ≥4 weeks if on pharmacotherapy",
    ],
    exclusion=[
        "Active suicidality or self-harm within 3 months",
        "Active psychosis or dissociative identity disorder with frequent switching",
        "Metallic implants in head/neck or implanted neurostimulator",
        "Pregnancy",
        "Active substance use disorder (alcohol or illicit stimulants)",
        "Inability to tolerate trauma-focused content (extreme dissociation)",
        "Severe traumatic brain injury as primary etiology",
    ],
    discussion=[
        "PTSD is characterized by hyperactivated fear circuits (amygdala), impaired extinction (vmPFC/hippocampus), and dysregulated arousal (locus coeruleus/SNS).",
        "The ventromedial prefrontal cortex (vmPFC) fails to inhibit amygdala hyperreactivity, maintaining conditioned fear responses.",
        "DLPFC hypoactivity impairs working memory, emotional regulation, and top-down control of trauma memories.",
        "taVNS directly modulates the locus coeruleus and amygdala via NTS projections, reducing hyperarousal and fear generalization.",
        "tDCS targeting DLPFC (anodal) enhances extinction learning and reduces intrusive symptom burden.",
        "TPS over vmPFC and ACC addresses fear extinction deficits and affective dysregulation.",
        "CES (Alpha-Stim) reduces hyperarousal, improves sleep architecture, and attenuates HPA axis overdrive.",
        "NIBS is best applied as an augmentation to trauma-focused psychotherapy (PE, CPT, EMDR), not as monotherapy.",
    ],
    overview=[
        "Post-Traumatic Stress Disorder (PTSD) is a debilitating psychiatric disorder arising from exposure to traumatic events, characterized by four symptom clusters: re-experiencing (intrusions, nightmares), avoidance, negative cognitions/mood, and hyperarousal. The disorder affects approximately 3.9% of the global population and is disproportionately prevalent in veterans, first responders, and survivors of interpersonal violence.",
        "The SOZO NIBS protocol for PTSD targets dysregulated fear circuitry—amygdala hyperactivation, vmPFC/hippocampal extinction failure, DLPFC hypoactivity, and locus coeruleus-mediated hyperarousal—through a phenotype-guided multimodal neuromodulation framework integrated with evidence-based psychotherapy.",
        "Six clinical phenotypes are recognized: hyperarousal dominant, re-experiencing dominant, avoidance/emotional numbing, cognitive/negative cognition dominant, complex PTSD (C-PTSD), and dissociative subtype. Each receives a tailored NIBS protocol optimized for the dominant neural circuit dysfunction.",
    ],
    pathophysiology=[
        "Amygdala hyperreactivity: exaggerated fear conditioning and generalization; failure of habituation to trauma cues; drives intrusions and hyperarousal.",
        "vmPFC/hippocampal extinction failure: impaired fear extinction memory consolidation; re-experiencing persists because safety signals are not encoded.",
        "DLPFC hypoactivity: reduced working memory capacity, emotional regulation, and top-down modulation of amygdala and intrusive memories.",
        "Locus coeruleus hyperactivation: chronic noradrenergic overdrive drives hypervigilance, exaggerated startle, sleep disruption, and intrusive re-experiencing.",
        "HPA axis dysregulation: low baseline cortisol with exaggerated glucocorticoid reactivity; impairs hippocampal neurogenesis and extinction consolidation.",
        "Default Mode Network disruption: trauma-related self-referential processing invades DMN; impairs sense of safety, present-moment awareness, and identity.",
    ],
    std_treatment=[
        "First-line psychotherapy: Prolonged Exposure (PE), Cognitive Processing Therapy (CPT), EMDR",
        "Pharmacotherapy: SSRIs (sertraline, paroxetine FDA-approved), SNRIs (venlafaxine); prazosin for nightmares",
        "Adjunct: sleep hygiene, prazosin for trauma nightmares, buspirone for anxiety",
        "Emerging: MDMA-assisted psychotherapy (Phase 3), ketamine, stellate ganglion block",
        "Multimodal: integrated trauma-focused therapy with medication management",
    ],
    symptoms=[
        ["Intrusions/Flashbacks", "Involuntary trauma memory re-experiencing", "100%", "APA DSM-5 2013"],
        ["Nightmares", "Trauma-related disturbing dreams", "70%", "Levin 2002"],
        ["Hypervigilance", "Persistent threat monitoring, exaggerated startle", "85%", "Pole 2007"],
        ["Avoidance", "Trauma cue avoidance (internal and external)", "90%", "Foa 2009"],
        ["Emotional Numbing", "Anhedonia, detachment, restricted affect", "60%", "King 1998"],
        ["Negative Cognitions", "Guilt, shame, self-blame, hopelessness", "80%", "Ehlers 2000"],
        ["Concentration Difficulties", "Impaired attention and working memory", "75%", "Scott 2015"],
        ["Sleep Disturbance", "Insomnia, fragmented sleep, nightmares", "90%", "Germain 2013"],
        ["Dissociation", "Depersonalization, derealization (dissociative subtype)", "30%", "Lanius 2010"],
        ["Irritability/Aggression", "Anger outbursts, irritable mood", "70%", "Teten 2010"],
    ],
    brain_regions=[
        ["Amygdala", "Fear conditioning, threat detection, hyperreactivity", "taVNS, TPS", "LeDoux 2000"],
        ["vmPFC (BA11/10)", "Fear extinction, safety signal encoding, amygdala inhibition", "TPS, tDCS anodal", "Milad 2007"],
        ["Hippocampus", "Contextual fear memory, extinction consolidation", "taVNS, TPS", "Quirk 2006"],
        ["DLPFC", "Working memory, emotional regulation, top-down control", "tDCS anodal", "Morey 2009"],
        ["ACC (Rostral)", "Fear extinction, conflict monitoring, emotional regulation", "TPS", "Shin 2006"],
        ["Insula", "Interoception, body-based fear, threat prediction", "TPS", "Paulus 2006"],
        ["Locus Coeruleus", "Noradrenergic hyperarousal, startle, fight-or-flight", "taVNS, CES", "Southwick 1997"],
        ["Posterior Parietal Cortex", "Spatial context, threat localization", "tDCS", "Gilbertson 2002"],
    ],
    brainstem=[
        ["Locus Coeruleus (LC)", "Hyperactive noradrenergic arousal; drives hypervigilance and exaggerated startle", "taVNS/CES direct", "Southwick 1997"],
        ["Nucleus Tractus Solitarius (NTS)", "Vagal input integration; modulates LC and amygdala", "taVNS direct", "Porges 2007"],
        ["Periaqueductal Gray (PAG)", "Freeze/fight-or-flight control; hyperactivated in PTSD", "taVNS indirect", "Mobbs 2007"],
        ["Raphe Nuclei", "Serotonergic tone regulation; impaired in PTSD", "taVNS", "Drevets 1999"],
        ["Superior Colliculus", "Threat detection, orienting; hyperactivated", "None direct", "Öhman 2005"],
    ],
    phenotypes=[
        ["Hyperarousal Dominant", "Hypervigilance, exaggerated startle, insomnia, irritability; LC overdrive", "Sleep disruption, anger, concentration loss", "taVNS + CES + tDCS DLPFC"],
        ["Re-experiencing Dominant", "Intrusions, flashbacks, nightmares; amygdala-vmPFC dysfunction", "Distress, avoidance triggered by intrusions", "TPS vmPFC + tDCS DLPFC + taVNS"],
        ["Avoidance/Numbing", "Emotional constriction, social withdrawal, detachment; blunted amygdala-DLPFC", "Anhedonia, relationship impairment", "tDCS DLPFC anodal + taVNS + TPS"],
        ["Cognitive/Negative Cognition", "Guilt, shame, self-blame; negative worldview; DLPFC-DMN dysfunction", "Depression comorbidity, hopelessness", "tDCS DLPFC + TPS ACC + taVNS"],
        ["Complex PTSD (C-PTSD)", "Prolonged/repeated trauma; affect dysregulation, identity disturbance, relational difficulties", "All 4 clusters + affect storms", "tDCS DLPFC + TPS ACC + taVNS + CES"],
        ["Dissociative Subtype", "Depersonalization/derealization; over-modulated prefrontal suppression of amygdala", "Emotional detachment, unreality", "Specialized protocol; TPS insula + taVNS"],
    ],
    symptom_map=[
        ["Hypervigilance/startle", "Locus Coeruleus/Amygdala", "taVNS + CES", "Vonck 2014, Lande 2018"],
        ["Intrusions/flashbacks", "Amygdala/vmPFC", "TPS vmPFC + tDCS DLPFC", "Milad 2007"],
        ["Nightmares", "Amygdala/Thalamus", "CES nightly + taVNS", "Germain 2013"],
        ["Avoidance", "DLPFC/vmPFC", "tDCS DLPFC anodal", "Morey 2009"],
        ["Emotional numbing", "DLPFC/Striatum", "tDCS DLPFC + TPS", "Bryant 2010"],
        ["Negative cognitions", "DLPFC/ACC", "tDCS DLPFC + TPS ACC", "Seminowicz 2013"],
        ["Sleep disruption", "Thalamus/LC", "CES Alpha-Stim nightly", "Lande 2018"],
        ["Dissociation", "Insula/mPFC", "TPS insula; specialized", "Lanius 2010"],
        ["Concentration impairment", "DLPFC/CEN", "tDCS DLPFC anodal", "Brunoni 2016"],
    ],
    montage=[
        ["Hyperarousal Dominant", "Anodal DLPFC (F3, 2mA, 20min)", "None initial", "taVNS 30min daily + CES nightly"],
        ["Re-experiencing Dominant", "Anodal DLPFC (F3, 2mA, 20min)", "vmPFC (300 pulses)", "taVNS before sessions; integrate PE/CPT"],
        ["Avoidance/Numbing", "Anodal DLPFC (F3, 2mA, 20min)", "ACC (250 pulses)", "taVNS + combined psychotherapy"],
        ["Cognitive/Negative", "Anodal DLPFC (F3, 2mA, 20min)", "ACC/vmPFC (300 pulses)", "taVNS + CES"],
        ["Complex PTSD", "Anodal DLPFC (F3, 2mA, 20min)", "ACC + vmPFC (400 pulses)", "taVNS + CES nightly; trauma therapy essential"],
        ["Dissociative Subtype", "Cathodal DLPFC (F3, 1.5mA)", "Insula (200 pulses, low)", "taVNS only; specialized trauma program"],
        ["With Comorbid MDD", "Anodal left DLPFC (F3, 2mA, 20min)", "ACC (300 pulses)", "taVNS + CES; antidepressant maintained"],
        ["Veteran Population", "Anodal DLPFC (F3, 2mA, 20min)", "vmPFC (300 pulses)", "taVNS + CES; veteran-specific trauma protocol"],
        ["Elderly Adjusted", "Anodal DLPFC (F3, 1.5mA, 20min)", "vmPFC (200 pulses)", "CES primary; taVNS standard"],
        ["Maintenance", "Anodal DLPFC (F3, 1.5mA, 20min)", "None", "CES 3x/week; taVNS 2x daily home"],
    ],
    tdcs_protocols=[
        ["PT-tDCS-01", "DLPFC Anodal Left", "F3 (left DLPFC)", "F4", "2mA, 20min, 10 sessions", "Boggio 2010"],
        ["PT-tDCS-02", "Bifrontal Augmentation", "F3 anodal", "F4 cathodal", "2mA, 20min, 10–20 sessions", "Fregni 2006"],
        ["PT-tDCS-03", "vmPFC Anodal (Extinction Enhancement)", "Fpz", "Oz", "1.5mA, 20min, 10 sessions", "Milad 2007"],
        ["PT-tDCS-04", "Dissociative Subtype Cathodal DLPFC", "F4", "F3", "1.5mA, 20min, specialized", "Lanius 2010"],
        ["PT-tDCS-05", "Maintenance Protocol", "F3 anodal", "F4", "1.5mA, 20min, 2x/week", "Bikson 2016"],
        ["PT-tDCS-06", "Combined with PE Therapy", "F3 anodal", "F4", "2mA, 20min before PE session", "Rosenbaum 2019"],
        ["PT-tDCS-07", "Elderly Adjusted", "F3 anodal", "F4", "1.5mA, 20min, 10 sessions", "Lefaucheur 2016"],
        ["PT-tDCS-08", "Sham", "F3", "F4", "30s active, 0mA, 20min", "Boggio 2010"],
    ],
    plato_protocols=[
        ["PT-PLATO-01", "DLPFC Anodal Left", "F3/F4", "Left DLPFC", "F3", "2.0mA", "N/A", "Continuous 20min"],
        ["PT-PLATO-02", "vmPFC Anodal", "Fpz/Oz", "vmPFC", "Fpz", "1.5mA", "N/A", "Continuous 20min"],
        ["PT-PLATO-03", "Bifrontal", "F3/F4", "Bifrontal", "F3", "2.0mA", "N/A", "Sequential"],
        ["PT-PLATO-04", "HD-tDCS DLPFC", "F3 center", "4-electrode ring", "F3", "1.0mA", "N/A", "HD mode 20min"],
        ["PT-PLATO-05", "Maintenance", "F3", "DLPFC", "F3", "1.5mA", "N/A", "2x/week 20min"],
        ["PT-PLATO-06", "Theta Burst Mode", "F3", "DLPFC", "F3", "1.0mA", "5Hz", "tACS 20min"],
        ["PT-PLATO-07", "Elderly", "F3", "DLPFC", "F3", "1.5mA", "N/A", "15min"],
        ["PT-PLATO-08", "Sham", "F3", "DLPFC", "F3", "0mA", "N/A", "30s ramp"],
    ],
    tps_protocols=[
        ["PT-TPS-01", "vmPFC Extinction Enhancement", "vmPFC (BA10/11)", "300 pulses, 0.25Hz, 0.2mJ/mm²", "3 sessions/week", "Milad 2007"],
        ["PT-TPS-02", "ACC Conflict/Affect Regulation", "ACC (BA24/32)", "250 pulses, 0.2Hz", "3 sessions/week", "Shin 2006"],
        ["PT-TPS-03", "Insula (Dissociative Subtype)", "Posterior Insula", "200 pulses, 0.2Hz, low intensity", "2 sessions/week", "Lanius 2010"],
        ["PT-TPS-04", "DLPFC Working Memory", "DLPFC (BA46)", "250 pulses, 0.2Hz", "3 sessions/week", "Morey 2009"],
        ["PT-TPS-05", "Full Protocol (C-PTSD)", "vmPFC + ACC sequential", "400 pulses total", "5 sessions/week x2 weeks", "Expert consensus 2024"],
    ],
    ces_role=[
        ["Hyperarousal", "Primary modality for LC downregulation and startle reduction", "Daily 60min; before sleep or as needed"],
        ["Sleep Disruption", "Improves sleep onset, continuity, nightmare frequency", "Nightly 60min; titrate to sleep quality"],
        ["HPA Axis Dysregulation", "Modulates cortisol reactivity and autonomic balance", "Daily 60min; concurrent with psychotherapy"],
        ["Maintenance", "Sustain arousal reduction; patient self-administered", "3–5x/week; home device program"],
    ],
    tavns_role="taVNS (tragus, 25Hz, 200µs, 0.5mA, 30min) activates the NTS → LC → amygdala pathway, directly reducing noradrenergic hyperarousal and fear generalization. Particularly effective for hyperarousal and re-experiencing phenotypes. Administered before NIBS sessions and as home device 2x daily. Evidence supports extinction memory facilitation via hippocampal theta enhancement (Vonck 2014, Hulsey 2019).",
    combinations=[
        ("Hyperarousal Dominant", [
            ["taVNS + CES", "Dual LC downregulation via vagal and thalamic pathways", "taVNS 30min → CES nightly", "Hypervigilance, sleep disruption, startle"],
            ["taVNS + tDCS DLPFC", "Arousal reduction + prefrontal regulatory upregulation", "taVNS 30min before tDCS", "Hyperarousal with concentration impairment"],
        ]),
        ("Re-experiencing Dominant", [
            ["tDCS DLPFC + TPS vmPFC + taVNS", "Top-down fear control + extinction + vagal fear modulation", "taVNS → TPS → tDCS, pre-therapy", "Intrusions, flashbacks; integrates with PE/EMDR"],
            ["tDCS DLPFC + TPS ACC", "Prefrontal + conflict resolution for trauma memory", "TPS → tDCS pre-session", "Intrusion-dominant without severe dissociation"],
        ]),
        ("Complex PTSD", [
            ["tDCS DLPFC + TPS ACC + taVNS + CES", "Quadruple modality for complex multi-domain dysfunction", "taVNS → TPS → tDCS daytime; CES nightly", "C-PTSD with affect storms and insomnia"],
        ]),
        ("Avoidance/Numbing", [
            ["tDCS DLPFC + taVNS", "Motivational activation + vagal engagement", "taVNS 30min before tDCS", "Emotional numbing, social withdrawal"],
        ]),
        ("Cognitive/Negative Cognition", [
            ["tDCS DLPFC + TPS ACC + taVNS", "Cognitive + conflict + vagal triple approach", "taVNS → TPS → tDCS", "Guilt, shame, hopelessness"],
        ]),
        ("Dissociative Subtype", [
            ["TPS insula + taVNS", "Interoceptive reintegration + vagal grounding", "TPS → taVNS; no tDCS anodal DLPFC", "Depersonalization, derealization"],
        ]),
    ],
    combination_summary=[
        ["Hyperarousal", "taVNS + CES", "LC downregulation dual pathway", "taVNS → CES nightly", "Hypervigilance, insomnia", "IIa"],
        ["Re-experiencing", "tDCS DLPFC + TPS vmPFC + taVNS", "Fear control + extinction + vagal", "taVNS → TPS → tDCS", "Intrusions, flashbacks", "IIb"],
        ["Avoidance/Numbing", "tDCS DLPFC + taVNS", "Motivational + vagal activation", "taVNS before tDCS", "Numbing, withdrawal", "IIb"],
        ["Cognitive/Negative", "tDCS DLPFC + TPS ACC + taVNS", "Prefrontal + conflict + vagal", "taVNS → TPS → tDCS", "Guilt, hopelessness", "IIb"],
        ["Complex PTSD", "All four modalities", "Multi-domain circuit normalization", "taVNS → TPS → tDCS + CES", "C-PTSD multi-cluster", "IIb"],
        ["Dissociative", "TPS insula + taVNS", "Interoceptive reintegration", "TPS → taVNS", "Depersonalization/derealization", "IIb"],
    ],
    outcomes=[
        ["CAPS-5 Total Score", "PTSD symptom severity", "Baseline, 4wk, 8wk", "≥15-point reduction (response); <11 (remission)"],
        ["PCL-5 Score", "PTSD symptom self-report", "Weekly", "≥10-point reduction from baseline"],
        ["PHQ-9", "Depression comorbidity", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["Pittsburgh Sleep Quality Index (PSQI)", "Sleep quality", "Baseline, 4wk, 8wk", "≥3-point improvement"],
        ["Heart Rate Variability (HRV)", "Autonomic regulation (vagal tone)", "Baseline, 4wk, 8wk", "≥10% increase in RMSSD"],
        ["Connor-Davidson Resilience Scale", "Resilience", "Baseline, 8wk", "≥5-point increase"],
        ["Dissociative Experiences Scale (DES)", "Dissociation (dissociative subtype)", "Baseline, 4wk, 8wk", "≥30% reduction"],
        ["Extinction Recall Test", "vmPFC extinction function", "Baseline, 8wk", "Improved safety signal recall"],
        ["Patient Global Impression of Change (PGIC)", "Overall improvement", "4wk, 8wk", "≥5/7 (much improved)"],
    ],
    outcomes_abbrev=[
        "CAPS-5: Clinician-Administered PTSD Scale for DSM-5",
        "PCL-5: PTSD Checklist for DSM-5",
        "PHQ-9: Patient Health Questionnaire-9",
        "PSQI: Pittsburgh Sleep Quality Index",
        "HRV: Heart Rate Variability",
        "RMSSD: Root Mean Square of Successive RR Differences",
        "DES: Dissociative Experiences Scale",
        "PGIC: Patient Global Impression of Change",
        "vmPFC: Ventromedial Prefrontal Cortex",
        "LC: Locus Coeruleus",
    ],
)

# ---------------------------------------------------------------------------
# CONDITION: OCD
# ---------------------------------------------------------------------------
CONDITIONS["ocd"] = dict(
    name="Obsessive-Compulsive Disorder",
    icd10="F42.2",
    doc_num="SPG-OC-001",
    tps_status="Investigational",
    tdcs_class="Class IIb",
    tps_class="Investigational",
    tavns_class="Class IIb",
    ces_class="Class IIb",
    inclusion=[
        "DSM-5 OCD diagnosis confirmed by structured interview (YBOCS score ≥16 at baseline)",
        "Symptom duration ≥6 months with significant functional impairment (CGI-S ≥4)",
        "Inadequate response to ≥2 adequate SSRI trials (≥12 weeks at therapeutic dose) or prior ERP failure",
        "Age 18–70; capable of informed consent and ERP engagement",
        "Stable medication for ≥8 weeks prior to enrollment",
        "Willing to maintain concurrent ERP throughout NIBS protocol",
    ],
    exclusion=[
        "Active suicidality or self-harm within 3 months",
        "Metallic implants in head/neck or implanted neurostimulator",
        "Pregnancy",
        "Active psychosis or bipolar disorder (manic phase)",
        "Hoarding disorder as sole presentation without classic OCD circuitry involvement",
        "Prior neurosurgical procedure for OCD (DBS, ablation)",
        "Seizure disorder or epilepsy",
    ],
    discussion=[
        "OCD is characterized by hyperactivation of the cortico-striato-thalamo-cortical (CSTC) loop — specifically the OFC-caudate-thalamo-OFC circuit driving compulsive urges.",
        "The orbitofrontal cortex (OFC) and anterior cingulate cortex (ACC) are hyperactive, generating excessive error signals and harm-avoidance urges.",
        "The supplementary motor area (SMA) and pre-SMA are involved in motor initiation of compulsions and response inhibition failure.",
        "tDCS targeting SMA (cathodal) reduces compulsion urgency by downregulating motor preparation circuits.",
        "tDCS targeting DLPFC (anodal) enhances top-down regulation of OFC and supports ERP-based response prevention.",
        "TPS over OFC and ACC directly modulates the hyperactive error-monitoring circuitry.",
        "taVNS activates the LC and raphe, enhancing serotonergic tone that is pharmacologically targeted by SSRIs.",
        "CES provides supplementary anxiolytic benefit and augments SSRI response through thalamic modulation.",
    ],
    overview=[
        "Obsessive-Compulsive Disorder (OCD) is a chronic, often disabling neuropsychiatric condition characterized by intrusive, ego-dystonic obsessions and repetitive compulsions performed to neutralize distress. With a lifetime prevalence of 2-3%, OCD ranks among the top 10 causes of disability worldwide. Core pathophysiology involves hyperactivity of the OFC-caudate-thalamo-OFC (CSTC) circuit, generating a pathological alarm signal that drives compulsive behavior.",
        "The SOZO NIBS protocol targets OCD's CSTC hyperactivation through a phenotype-guided multimodal neuromodulation framework: tDCS (SMA cathodal / DLPFC anodal), TPS (OFC/ACC/SMA), taVNS (serotonergic augmentation), and CES (Alpha-Stim, anxiolytic adjunct). All NIBS is integrated with concurrent ERP (Exposure and Response Prevention) as the required psychotherapy backbone.",
        "Five clinical phenotypes guide protocol selection: contamination/washing dominant, harm obsession/checking dominant, symmetry/ordering dominant, intrusive thought dominant, and treatment-refractory OCD. Each phenotype targets the most implicated circuit nodes within the shared CSTC architecture.",
    ],
    pathophysiology=[
        "CSTC hyperactivation: orbitofrontal cortex → caudate nucleus → thalamus → OFC loop is hyperactive; generates persistent false alarm error signals driving compulsion.",
        "OFC error monitoring: hyperactive OFC produces excessive harm-appraisal and contamination threat signals; direct TPS target.",
        "ACC conflict monitoring: hyperactivated dACC amplifies uncertainty intolerance; drives compulsive checking behavior.",
        "SMA/pre-SMA hyperactivation: excessive motor preparation for compulsive rituals; cathodal tDCS target.",
        "Impaired top-down regulatory control: DLPFC hypoactivity reduces capacity for volitional response inhibition; anodal tDCS target.",
        "Serotonergic deficit: reduced 5-HT tone impairs OFC-caudate inhibitory balance; addressed by SSRI + taVNS serotonergic augmentation.",
    ],
    std_treatment=[
        "First-line: SSRIs at high doses (fluoxetine, fluvoxamine, paroxetine, sertraline, escitalopram) + ERP psychotherapy",
        "Augmentation: antipsychotics (risperidone, aripiprazole) for partial responders",
        "Intensive ERP: residential or intensive outpatient programs",
        "Advanced: rTMS (FDA-cleared for OCD — dTMS OFC/ACC), DBS (anterior limb internal capsule/ventral striatum)",
        "Emerging: glutamatergic agents (NAC, riluzole), ketamine, psilocybin-assisted therapy",
    ],
    symptoms=[
        ["Contamination Obsessions", "Fear of contamination, germs, disease", "38%", "Abramowitz 2009"],
        ["Washing Compulsions", "Hand washing, cleaning rituals", "38%", "Rasmussen 1991"],
        ["Harm Obsessions", "Fear of harming self/others, moral guilt", "28%", "Abramowitz 2009"],
        ["Checking Compulsions", "Door, stove, lock checking rituals", "28%", "Rasmussen 1991"],
        ["Symmetry Obsessions", "Need for order, exactness, symmetry", "20%", "Leckman 1997"],
        ["Ordering/Arranging", "Arranging objects precisely until 'just right'", "20%", "Miguel 2000"],
        ["Intrusive Thoughts", "Taboo, sexual, religious, violent thoughts", "20%", "Purdon 1999"],
        ["Neutralizing Compulsions", "Mental rituals, prayers, counting", "20%", "Clark 2004"],
        ["Hoarding (OCD-type)", "Pathological saving when ego-dystonic", "14%", "Frost 2010"],
        ["Insight Impairment", "Poor/absent insight in severe cases", "15%", "Kozak 2000"],
    ],
    brain_regions=[
        ["Orbitofrontal Cortex (OFC)", "Error signal generation, contamination/harm appraisal", "TPS, tDCS cathodal indirect", "Saxena 2001"],
        ["Caudate Nucleus", "CSTC relay; hyperactive in OCD", "TPS indirect", "Rauch 1994"],
        ["Anterior Cingulate Cortex (ACC)", "Conflict monitoring, uncertainty intolerance", "TPS", "Carter 2000"],
        ["Thalamus", "CSTC gating; hyperactive relay in OCD", "CES, TPS indirect", "Rosenberg 1997"],
        ["DLPFC", "Response inhibition, top-down OFC regulation", "tDCS anodal", "Chamberlain 2008"],
        ["Supplementary Motor Area (SMA)", "Motor preparation for compulsions; pre-SMA response inhibition", "tDCS cathodal", "Nachev 2008"],
        ["Insula", "Disgust processing, body-based threat signals", "TPS", "Husted 2006"],
        ["Subthalamic Nucleus", "Impulsivity regulation; DBS target in refractory OCD", "None direct NIBS", "Mallet 2008"],
    ],
    brainstem=[
        ["Raphe Nuclei", "Serotonergic output; reduced 5-HT drives CSTC disinhibition", "taVNS indirect", "Saxena 2001"],
        ["Locus Coeruleus", "Noradrenergic anxiety component; hyperarousal in OCD", "taVNS, CES", "Southwick 1997"],
        ["Nucleus Tractus Solitarius", "Vagal input; relays to raphe and LC", "taVNS direct", "Porges 2007"],
        ["Substantia Nigra", "Dopaminergic-CSTC interaction; striatal modulation", "None direct", "Gottesman 2002"],
        ["Parabrachial Nucleus", "Disgust and interoceptive relay to insula", "None direct", "Husted 2006"],
    ],
    phenotypes=[
        ["Contamination/Washing", "Contamination obsessions + washing compulsions; disgust-driven OFC-insula circuit", "Skin damage from washing, anxiety", "tDCS SMA cathodal + TPS OFC + taVNS"],
        ["Harm/Checking", "Harm obsessions + checking compulsions; dACC hyperactivity", "Checking rituals, doubt, guilt", "tDCS SMA cathodal + TPS ACC + tDCS DLPFC"],
        ["Symmetry/Ordering", "Symmetry obsessions + ordering compulsions; sensorimotor CSTC involvement", "Just-right phenomenology, motor rituals", "tDCS SMA cathodal + TPS + tDCS DLPFC"],
        ["Intrusive Thought Dominant", "Taboo/moral/sexual/violent intrusive thoughts; ACC-DLPFC dysfunction", "Moral guilt, mental neutralizing", "tDCS DLPFC anodal + TPS ACC + taVNS"],
        ["Treatment-Refractory OCD", "≥3 SSRI failures + ERP failure; severe CSTC pathology", "All symptom dimensions; severe impairment", "Full protocol: tDCS SMA + DLPFC + TPS OFC/ACC + taVNS + CES"],
        ["OCD + Tic Comorbidity", "OCD + Tourette/chronic tics; SMA-striatal circuit predominance", "Tic + compulsion interaction", "tDCS SMA cathodal primary + TPS SMA + specialized"],
    ],
    symptom_map=[
        ["Contamination obsessions", "OFC/Insula", "TPS OFC + tDCS SMA cathodal", "Saxena 2001, Rauch 1994"],
        ["Washing compulsions", "SMA/OFC", "tDCS SMA cathodal", "Nachev 2008"],
        ["Checking compulsions", "ACC/SMA", "tDCS SMA cathodal + TPS ACC", "Carter 2000"],
        ["Harm obsessions", "ACC/OFC", "TPS ACC/OFC + tDCS DLPFC", "Abramowitz 2009"],
        ["Symmetry/ordering", "SMA/Sensorimotor", "tDCS SMA cathodal + TPS", "Leckman 1997"],
        ["Intrusive thoughts", "DLPFC/ACC", "tDCS DLPFC anodal + TPS ACC", "Purdon 1999"],
        ["Anxiety/arousal", "Amygdala/Thalamus", "CES + taVNS", "Lande 2018"],
        ["Response inhibition failure", "DLPFC/pre-SMA", "tDCS DLPFC anodal + SMA cathodal", "Chamberlain 2008"],
        ["Serotonergic deficit", "Raphe/CSTC", "taVNS serotonergic augmentation", "Vonck 2014"],
    ],
    montage=[
        ["Contamination/Washing", "Cathodal SMA (Cz area, 1.5mA, 20min) + Anodal DLPFC", "OFC (250 pulses)", "taVNS + ERP concurrent"],
        ["Harm/Checking", "Cathodal SMA (1.5mA, 20min) + Anodal DLPFC (F3)", "ACC (250 pulses)", "taVNS + CES + ERP"],
        ["Symmetry/Ordering", "Cathodal SMA (Cz, 1.5mA, 20min)", "SMA + OFC (300 pulses)", "ERP concurrent; CES adjunct"],
        ["Intrusive Thought", "Anodal DLPFC (F3, 2mA, 20min)", "ACC (250 pulses)", "taVNS + ERP essential"],
        ["Refractory OCD", "Cathodal SMA + Anodal DLPFC bilateral", "OFC + ACC (400 pulses)", "All 4 modalities + intensive ERP"],
        ["OCD + Tics", "Cathodal SMA (Cz, 1.5mA, 20min)", "SMA (200 pulses)", "Specialized; tic management concurrent"],
        ["Maintenance Phase", "Anodal DLPFC (F3, 1.5mA, 20min)", "None", "CES 3x/week; taVNS daily"],
        ["Elderly Adjusted", "Anodal DLPFC (F3, 1.5mA, 15min)", "OFC (200 pulses)", "CES primary; taVNS standard"],
        ["With MDD Comorbidity", "Anodal left DLPFC (F3, 2mA, 20min)", "ACC (250 pulses)", "taVNS + CES; SSRI maintained"],
        ["Sham Protocol", "F3 anodal (30s ramp, 0mA, 20min)", "None", "Sham ERP control condition"],
    ],
    tdcs_protocols=[
        ["OC-tDCS-01", "SMA Cathodal (Compulsion Inhibition)", "Fp1/Fp2", "Cz (SMA)", "1.5mA, 20min, 10 sessions", "Narayanan 2016"],
        ["OC-tDCS-02", "DLPFC Anodal Left (Top-Down Regulation)", "F3", "F4", "2mA, 20min, 10 sessions", "Bation 2016"],
        ["OC-tDCS-03", "Combined SMA Cathodal + DLPFC Anodal", "F3 anodal; Cz cathodal", "F4; Fp2", "1.5–2mA, sequential, 10 sessions", "Expert consensus 2023"],
        ["OC-tDCS-04", "OFC Cathodal (Refractory)", "Fp1 cathodal", "Oz", "1mA, 20min, 10 sessions", "Bation 2016"],
        ["OC-tDCS-05", "Maintenance", "F3 anodal", "F4", "1.5mA, 20min, 2x/week", "Bikson 2016"],
        ["OC-tDCS-06", "Pre-ERP Augmentation", "F3 anodal", "F4", "2mA, 20min before ERP", "Narayanan 2016"],
        ["OC-tDCS-07", "Elderly Adjusted", "F3 anodal", "F4", "1.5mA, 15min, 10 sessions", "Lefaucheur 2016"],
        ["OC-tDCS-08", "Sham", "F3", "F4", "30s active, 0mA, 20min", "Bation 2016"],
    ],
    plato_protocols=[
        ["OC-PLATO-01", "SMA Cathodal", "Cz/Fpz", "SMA", "Cz", "1.5mA", "N/A", "Cathodal 20min"],
        ["OC-PLATO-02", "DLPFC Anodal", "F3/F4", "Left DLPFC", "F3", "2.0mA", "N/A", "Continuous 20min"],
        ["OC-PLATO-03", "Combined SMA+DLPFC", "F3; Cz", "DLPFC + SMA", "F3; Cz", "2.0; 1.5mA", "N/A", "Sequential"],
        ["OC-PLATO-04", "HD-tDCS DLPFC", "F3 center", "4-electrode ring", "F3", "1.0mA", "N/A", "HD 20min"],
        ["OC-PLATO-05", "Maintenance", "F3", "DLPFC", "F3", "1.5mA", "N/A", "2x/week"],
        ["OC-PLATO-06", "Beta tACS (SMA)", "Cz", "SMA", "Cz", "1.0mA", "20Hz", "tACS 20min"],
        ["OC-PLATO-07", "Elderly", "F3", "DLPFC", "F3", "1.5mA", "N/A", "15min"],
        ["OC-PLATO-08", "Sham", "F3", "DLPFC", "F3", "0mA", "N/A", "30s ramp"],
    ],
    tps_protocols=[
        ["OC-TPS-01", "OFC Error Signal Reduction", "OFC (BA11/47)", "250 pulses, 0.2Hz, 0.2mJ/mm²", "3 sessions/week", "Saxena 2001"],
        ["OC-TPS-02", "ACC Conflict Monitoring", "dACC (BA24)", "250 pulses, 0.2Hz", "3 sessions/week", "Carter 2000"],
        ["OC-TPS-03", "SMA Compulsion Inhibition", "Pre-SMA (BA6)", "200 pulses, 0.2Hz", "3 sessions/week", "Nachev 2008"],
        ["OC-TPS-04", "Insula Disgust Processing", "Posterior Insula", "200 pulses, 0.2Hz", "2 sessions/week", "Husted 2006"],
        ["OC-TPS-05", "Full Protocol (Refractory)", "OFC + ACC + SMA sequential", "400 pulses total", "5 sessions/week x2 weeks", "Expert consensus 2024"],
    ],
    ces_role=[
        ["Anxiety Augmentation", "Reduces obsessional anxiety and arousal; enhances ERP tolerance", "Daily 60min; before ERP sessions"],
        ["SSRI Augmentation", "Thalamic modulation may augment SSRI response", "Daily 60min concurrent with pharmacotherapy"],
        ["Sleep Disruption", "Addresses OCD-related sleep problems", "Nightly 60min if sleep-onset difficulty"],
        ["Maintenance", "Sustain anxiety reduction between sessions", "3x/week; patient self-administered"],
    ],
    tavns_role="taVNS (tragus, 25Hz, 200µs, 0.5mA, 30min) augments serotonergic tone via NTS → raphe nuclei pathway, complementing SSRI pharmacotherapy. Also reduces hyperarousal and anxiety comorbidity. Administered before NIBS and ERP sessions. Evidence for serotonergic augmentation positions taVNS as a pharmacological enhancer for SSRI-partial responders (Vonck 2014, Straube 2015).",
    combinations=[
        ("Contamination/Washing", [
            ["tDCS SMA cathodal + TPS OFC + taVNS", "Compulsion inhibition + error signal reduction + serotonergic augmentation", "taVNS → TPS → tDCS, before ERP", "Contamination obsessions with washing compulsions"],
            ["tDCS SMA cathodal + CES", "Motor inhibition + anxiolytic adjunct", "tDCS then CES post-session", "Washing compulsions with high anxiety"],
        ]),
        ("Harm/Checking", [
            ["tDCS SMA cathodal + tDCS DLPFC anodal + TPS ACC", "Compulsion inhibition + top-down control + conflict monitoring", "TPS → tDCS SMA → tDCS DLPFC, before ERP", "Checking compulsions, harm obsessions"],
        ]),
        ("Intrusive Thought Dominant", [
            ["tDCS DLPFC + TPS ACC + taVNS", "Cognitive control + conflict + serotonergic", "taVNS → TPS → tDCS, before ERP", "Moral/sexual/violent intrusive thoughts"],
        ]),
        ("Refractory OCD", [
            ["tDCS SMA cathodal + tDCS DLPFC + TPS OFC/ACC + taVNS + CES", "Full circuit attack — all nodes", "taVNS → TPS → tDCS; CES nightly", "≥3 SSRI failures, ERP failure"],
        ]),
        ("Symmetry/Ordering", [
            ["tDCS SMA cathodal + TPS OFC/SMA", "Motor + sensorimotor CSTC inhibition", "TPS → tDCS, before ERP", "Just-right phenomenology, ordering rituals"],
        ]),
        ("OCD + Tics", [
            ["tDCS SMA cathodal + TPS SMA", "Convergent SMA inhibition for tic-compulsion complex", "TPS → tDCS; specialized program", "Tourette + OCD comorbidity"],
        ]),
    ],
    combination_summary=[
        ["Contamination/Washing", "tDCS SMA cathodal + TPS OFC + taVNS", "Compulsion + error signal + serotonergic", "taVNS → TPS → tDCS + ERP", "Contamination/washing", "IIb"],
        ["Harm/Checking", "tDCS SMA + DLPFC + TPS ACC", "Motor + top-down + conflict", "TPS → tDCS × 2 + ERP", "Harm/checking", "IIb"],
        ["Intrusive Thought", "tDCS DLPFC + TPS ACC + taVNS", "Cognitive + conflict + serotonergic", "taVNS → TPS → tDCS + ERP", "Intrusive thoughts", "IIb"],
        ["Refractory OCD", "All 4 modalities", "Full CSTC circuit normalization", "taVNS → TPS → tDCS + CES", "Refractory OCD", "IIb"],
        ["Symmetry/Ordering", "tDCS SMA cathodal + TPS OFC/SMA", "SMA inhibition + sensorimotor", "TPS → tDCS + ERP", "Ordering compulsions", "IIb"],
        ["OCD + Tics", "tDCS SMA cathodal + TPS SMA", "Convergent SMA inhibition", "TPS → tDCS; specialized", "Tic-OCD comorbid", "IIb"],
    ],
    outcomes=[
        ["Y-BOCS Total Score", "OCD severity (obsessions + compulsions)", "Baseline, 4wk, 8wk", "≥35% reduction (response); ≤7 (remission)"],
        ["Y-BOCS Obsession Subscale", "Obsession severity", "Baseline, 4wk, 8wk", "≥35% reduction"],
        ["Y-BOCS Compulsion Subscale", "Compulsion severity", "Baseline, 4wk, 8wk", "≥35% reduction"],
        ["CGI-Severity (CGI-S)", "Clinical global severity", "Baseline, 4wk, 8wk", "Reduction ≥2 points (much improved)"],
        ["PHQ-9", "Depression comorbidity", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["GAD-7", "Anxiety comorbidity", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["OCD-Related Disability (WHODAS 2.0)", "Functional impairment", "Baseline, 8wk", "≥20% improvement"],
        ["ERP Homework Completion", "Treatment engagement", "Weekly", "≥80% completion rate"],
        ["Patient Global Impression of Change", "Overall response", "4wk, 8wk", "≥5/7 (much improved)"],
    ],
    outcomes_abbrev=[
        "Y-BOCS: Yale-Brown Obsessive Compulsive Scale",
        "CGI-S: Clinical Global Impressions - Severity",
        "PHQ-9: Patient Health Questionnaire-9",
        "GAD-7: Generalized Anxiety Disorder 7",
        "WHODAS: World Health Organization Disability Assessment Schedule",
        "ERP: Exposure and Response Prevention",
        "CSTC: Cortico-Striato-Thalamo-Cortical",
        "OFC: Orbitofrontal Cortex",
        "SMA: Supplementary Motor Area",
        "DLPFC: Dorsolateral Prefrontal Cortex",
    ],
)
print("Chronic Pain + PTSD + OCD appended OK")

# ---------------------------------------------------------------------------
# CONDITION: MULTIPLE SCLEROSIS
# ---------------------------------------------------------------------------
CONDITIONS["ms"] = dict(
    name="Multiple Sclerosis",
    icd10="G35",
    doc_num="SPG-MS-001",
    tps_status="Investigational",
    tdcs_class="Class IIb",
    tps_class="Investigational",
    tavns_class="Class IIb",
    ces_class="Class IIb",
    inclusion=[
        "Confirmed MS diagnosis per McDonald 2017 criteria (RRMS, SPMS, or PPMS)",
        "EDSS score 2.0–7.0 at baseline",
        "Stable disease-modifying therapy (DMT) for ≥3 months or no DMT planned",
        "Age 18–70; capable of informed consent",
        "Symptom burden including fatigue (MFIS ≥38), cognitive impairment (SDMT <55), spasticity, or mood symptoms",
        "No relapse within 3 months of enrollment",
    ],
    exclusion=[
        "Metallic implants in head/neck or implanted neurostimulator",
        "Active relapse or corticosteroid treatment within 3 months",
        "Severe cognitive impairment precluding consent (MMSE <18)",
        "Pregnancy",
        "History of seizures not controlled with medication",
        "Rapidly progressive disease course (EDSS change >1.0 in past 6 months)",
        "Active cancer or severe systemic illness",
    ],
    discussion=[
        "MS-related neurological dysfunction arises from demyelination and axonal loss disrupting large-scale brain network connectivity.",
        "Fatigue in MS involves thalamo-cortical circuit dysfunction, reduced cortico-spinal efficiency, and hypothalamic-pituitary-adrenal axis dysregulation.",
        "Cognitive impairment (particularly processing speed and working memory) correlates with DMN-CEN disconnection and thalamic atrophy.",
        "Spasticity results from corticospinal tract demyelination reducing inhibitory tone over spinal motor circuits.",
        "tDCS targeting M1 (anodal) improves motor function, reduces fatigue, and temporarily enhances corticospinal conduction efficiency.",
        "tDCS over DLPFC augments cognitive processing speed and working memory in MS.",
        "taVNS exerts anti-inflammatory effects via the vagal cholinergic anti-inflammatory pathway, potentially modulating neuroinflammation.",
        "CES addresses MS-related insomnia, fatigue, and mood symptoms through thalamic and brainstem modulation.",
    ],
    overview=[
        "Multiple Sclerosis (MS) is a chronic autoimmune demyelinating disease of the central nervous system affecting approximately 2.8 million people worldwide. Characterized by inflammatory demyelination and progressive axonal loss, MS produces diverse neurological symptoms including fatigue, cognitive impairment, spasticity, pain, mood disorders, and disability accumulation across four disease courses (RRMS, SPMS, PPMS, CIS).",
        "The SOZO NIBS protocol addresses MS symptom burden rather than disease modification. By targeting the neurophysiological consequences of demyelination — corticospinal inefficiency, cognitive network dysfunction, central sensitization, and fatigue circuits — NIBS offers symptomatic benefit complementary to disease-modifying therapies.",
        "Five clinical phenotypes guide protocol selection: fatigue-dominant, cognitive-dominant, spasticity-dominant, pain-dominant, and mood/depression-dominant. Each phenotype receives a tailored NIBS combination targeting the most impaired neural network.",
    ],
    pathophysiology=[
        "Demyelination and axonal loss: focal white matter lesions disrupt myelinated fiber conduction velocity and axonal integrity; cortical and deep gray matter atrophy contributes to network dysfunction.",
        "Corticospinal tract impairment: demyelination of descending motor pathways reduces cortical drive, increases spasticity, and impairs voluntary motor control.",
        "Fatigue circuit dysfunction: thalamocortical and cortico-striatal-thalamic circuits show hyperactivation as compensatory burden; fatigue reflects neural inefficiency.",
        "Cognitive impairment: processing speed loss correlates with corpus callosum integrity and thalamic volume; working memory deficit involves DLPFC-CEN disruption.",
        "Neuroinflammation: CNS-infiltrating T-cells, microglia activation, and demyelinating plaques in periventricular, cortical, and brainstem regions.",
        "Central sensitization and pain: MS-related central sensitization amplifies nociception through thalamic and cortical pain network dysregulation.",
    ],
    std_treatment=[
        "Disease-modifying therapies: interferons, glatiramer, natalizumab, ocrelizumab, siponimod, ozanimod (by disease course)",
        "Symptom management: baclofen/tizanidine (spasticity), modafinil/amantadine (fatigue), fampridine (walking speed)",
        "Cognitive rehabilitation: computerized cognitive training, compensatory strategy training",
        "Physical/occupational therapy: progressive resistance training, aquatic therapy, assistive devices",
        "Psychological support: CBT for depression and adjustment, MS-specific mindfulness",
    ],
    symptoms=[
        ["Fatigue", "Overwhelming fatigue disproportionate to activity (Uhthoff variant with heat)", "80%", "Krupp 1988"],
        ["Cognitive Impairment", "Processing speed, working memory, episodic memory deficits", "65%", "Chiaravalloti 2008"],
        ["Spasticity", "Velocity-dependent resistance to passive movement; cramps", "60%", "Rizzo 2004"],
        ["Pain", "Central neuropathic pain, trigeminal neuralgia, Lhermitte sign", "55%", "Foley 2013"],
        ["Depression", "Comorbid MDD; partly biologically driven by neuroinflammation", "50%", "Feinstein 2011"],
        ["Walking/Balance Impairment", "Reduced gait speed, balance deficits, fall risk", "70%", "Kalron 2016"],
        ["Bladder Dysfunction", "Urgency, frequency, retention", "75%", "Fowler 2009"],
        ["Visual Disturbance", "Optic neuritis, diplopia, nystagmus", "55%", "Petzold 2018"],
        ["Tremor", "Cerebellar intention tremor", "40%", "Koch 2007"],
        ["Sleep Disturbance", "Insomnia, restless legs, sleep fragmentation", "50%", "Braley 2014"],
    ],
    brain_regions=[
        ["Primary Motor Cortex (M1)", "Corticospinal drive; compensatory hyperactivation in MS", "tDCS anodal M1", "Fregni 2006"],
        ["DLPFC", "Cognitive processing speed, working memory, fatigue regulation", "tDCS anodal", "Mattioli 2010"],
        ["Cerebellum", "Tremor, ataxia, coordination; direct NIBS difficult", "TPS indirect", "Koch 2007"],
        ["Thalamus", "Fatigue relay, sensory gating, thalamocortical dysrhythmia", "CES", "Filippi 2013"],
        ["ACC", "Fatigue perception, effort-cost computation, cognitive-motor integration", "TPS", "Yaldizli 2011"],
        ["Supplementary Motor Area", "Motor sequence, fatigue-motor interaction", "tDCS", "Leavitt 2014"],
        ["Posterior Parietal Cortex", "Visuospatial, somatosensory integration disrupted by MS", "tDCS, TPS", "Chiaravalloti 2008"],
        ["Hippocampus", "Episodic memory; directly affected by MS pathology", "taVNS", "Sicotte 2008"],
    ],
    brainstem=[
        ["Pons/Brainstem", "Frequent MS lesion location; affects cranial nerves, gait, dysarthria", "taVNS indirect", "Filippi 2002"],
        ["Locus Coeruleus", "Fatigue and arousal regulation; noradrenergic impairment in MS", "taVNS, CES", "Feinstein 2011"],
        ["Nucleus Tractus Solitarius", "Vagal input; anti-inflammatory relay", "taVNS direct", "Tracey 2002"],
        ["Raphe Nuclei", "Serotonergic tone; depression in MS partly serotonergic", "taVNS", "Feinstein 2011"],
        ["Vestibular Nuclei", "Balance and gaze stabilization; affected by MS brainstem lesions", "None direct", "Frohman 2005"],
    ],
    phenotypes=[
        ["Fatigue-Dominant", "Severe MS fatigue (MFIS ≥38); thalamocortical and cortico-striatal inefficiency", "Fatigue limiting daily function; activity avoidance", "tDCS M1 + DLPFC + CES"],
        ["Cognitive-Dominant", "Processing speed, working memory, attention impairment (SDMT <45)", "Cognitive fog, occupational/social impairment", "tDCS DLPFC + TPS + taVNS"],
        ["Spasticity-Dominant", "Moderate-severe spasticity (MAS 2–3); corticospinal demyelination", "Cramps, gait limitation, pain from spasticity", "tDCS M1 anodal + TPS M1/SMA"],
        ["Pain-Dominant", "Central neuropathic pain ≥4/10 NRS; trigeminal neuralgia; Lhermitte", "Burning, electric, shooting pain; sleep disruption", "tDCS M1 + TPS + CES"],
        ["Mood/Depression-Dominant", "MS-related depression (PHQ-9 ≥10); partly neuroinflammatory", "Low mood, fatigue-depression interaction, reduced adherence", "tDCS DLPFC + taVNS + CES"],
    ],
    symptom_map=[
        ["MS Fatigue", "Thalamus/M1/DLPFC", "tDCS M1 + DLPFC + CES", "Fregni 2006, Lande 2018"],
        ["Cognitive impairment", "DLPFC/CEN/Thalamus", "tDCS DLPFC anodal", "Mattioli 2010"],
        ["Spasticity", "M1/Corticospinal", "tDCS M1 anodal", "Lefaucheur 2016"],
        ["Central pain", "Thalamus/M1/ACC", "tDCS M1 + CES", "Fregni 2006"],
        ["Depression", "DLPFC/Limbic", "tDCS DLPFC + taVNS", "Brunoni 2013"],
        ["Walking impairment", "M1/SMA", "tDCS M1 + TPS", "Otal 2016"],
        ["Sleep disruption", "Thalamus/LC", "CES Alpha-Stim", "Braley 2014"],
        ["Neuroinflammation", "Systemic/CNS", "taVNS cholinergic anti-inflammatory", "Tracey 2002"],
        ["Tremor", "Cerebellum/Thalamus", "TPS cerebellar indirect", "Koch 2007"],
    ],
    montage=[
        ["Fatigue-Dominant", "Anodal M1 (C3/C4, 2mA, 20min) + Anodal DLPFC (F3)", "None initial", "CES nightly; taVNS 2x daily"],
        ["Cognitive-Dominant", "Anodal DLPFC (F3, 2mA, 20min)", "DLPFC (250 pulses)", "taVNS before sessions; CES adjunct"],
        ["Spasticity-Dominant", "Anodal M1 (C3/C4, 2mA, 20min)", "M1/SMA (300 pulses)", "Baclofen maintained; physio concurrent"],
        ["Pain-Dominant", "Anodal M1 (C3, 2mA, 20min)", "ACC (250 pulses)", "CES nightly; taVNS 2x daily"],
        ["Mood/Depression", "Anodal DLPFC left (F3, 2mA, 20min)", "None initial", "taVNS + CES nightly; DMT maintained"],
        ["RRMS Active Symptoms", "Anodal DLPFC (F3, 2mA) + M1 (C3, 2mA)", "DLPFC + M1 (300 pulses)", "taVNS + CES; DMT maintained"],
        ["SPMS/PPMS Progressive", "Anodal M1 (C3, 1.5mA, 20min)", "M1 (200 pulses)", "CES primary; taVNS 2x daily"],
        ["Elderly/High EDSS (5–7)", "Anodal M1 (C3, 1mA, 15min)", "TPS low (150 pulses)", "CES + taVNS; adapted physio"],
        ["Maintenance Phase", "Anodal DLPFC (F3, 1.5mA, 20min)", "None", "CES 3x/week; taVNS daily home"],
        ["Sham Protocol", "F3 anodal (30s ramp, 0mA, 20min)", "None", "Sham control condition"],
    ],
    tdcs_protocols=[
        ["MS-tDCS-01", "M1 Anodal (Fatigue/Spasticity/Motor)", "C3/C4 contralateral", "Fp2", "2mA, 20min, 10 sessions", "Fregni 2006"],
        ["MS-tDCS-02", "DLPFC Anodal Left (Cognitive/Depression)", "F3", "F4", "2mA, 20min, 10 sessions", "Mattioli 2010"],
        ["MS-tDCS-03", "Combined M1 + DLPFC", "C3 + F3 anodal sequential", "Fp2 + F4", "2mA each, 20min, 15 sessions", "Otal 2016"],
        ["MS-tDCS-04", "Bifrontal (Fatigue/Mood)", "F3 anodal", "F4 cathodal", "2mA, 20min, 10 sessions", "Brunoni 2013"],
        ["MS-tDCS-05", "M1 Anodal (PPMS/SPMS Progressive)", "C3/C4", "Fp2", "1.5mA, 20min, 10 sessions", "Lefaucheur 2016"],
        ["MS-tDCS-06", "Maintenance", "F3 anodal", "F4", "1.5mA, 20min, 2x/week", "Bikson 2016"],
        ["MS-tDCS-07", "Elderly/High EDSS Adjusted", "C3 anodal", "Fp2", "1mA, 15min, 10 sessions", "Lefaucheur 2016"],
        ["MS-tDCS-08", "Sham", "F3 anodal", "F4", "30s ramp, 0mA, 20min", "Mattioli 2010"],
    ],
    plato_protocols=[
        ["MS-PLATO-01", "M1 Anodal", "C3/Fp2", "M1 contralateral", "C3", "2.0mA", "N/A", "Continuous 20min"],
        ["MS-PLATO-02", "DLPFC Anodal", "F3/F4", "Left DLPFC", "F3", "2.0mA", "N/A", "Continuous 20min"],
        ["MS-PLATO-03", "Combined M1+DLPFC", "C3; F3", "M1 + DLPFC", "C3; F3", "2.0mA each", "N/A", "Sequential"],
        ["MS-PLATO-04", "HD-tDCS M1", "C3 center", "4-electrode ring", "C3", "1.0mA", "N/A", "HD 20min"],
        ["MS-PLATO-05", "Maintenance", "F3", "DLPFC", "F3", "1.5mA", "N/A", "2x/week"],
        ["MS-PLATO-06", "Gamma tACS Fatigue", "C3", "M1", "C3", "1.5mA", "40Hz", "tACS 20min"],
        ["MS-PLATO-07", "Elderly/High EDSS", "C3", "M1", "C3", "1.0mA", "N/A", "15min"],
        ["MS-PLATO-08", "Sham", "F3", "DLPFC", "F3", "0mA", "N/A", "30s ramp"],
    ],
    tps_protocols=[
        ["MS-TPS-01", "M1 Motor Facilitation", "M1 bilateral or contralateral", "300 pulses, 0.25Hz, 0.2mJ/mm²", "3 sessions/week", "Otal 2016"],
        ["MS-TPS-02", "DLPFC Cognitive", "DLPFC (BA46)", "250 pulses, 0.2Hz", "3 sessions/week", "Mattioli 2010"],
        ["MS-TPS-03", "ACC Fatigue-Cognitive", "ACC (BA24/32)", "200 pulses, 0.2Hz", "2 sessions/week", "Yaldizli 2011"],
        ["MS-TPS-04", "SMA Motor Sequence", "SMA (BA6)", "200 pulses, 0.2Hz", "2 sessions/week", "Leavitt 2014"],
        ["MS-TPS-05", "Full Protocol", "M1 + DLPFC + ACC sequential", "450 pulses total", "5 sessions/week x2 weeks", "Expert consensus 2024"],
    ],
    ces_role=[
        ["MS Fatigue", "Reduces fatigue severity via thalamic-reticular network modulation", "Daily 60min; morning or pre-activity"],
        ["Sleep Disruption", "Improves sleep quality and reduces nocturnal awakening", "Nightly 60min"],
        ["Depression Comorbidity", "Adjunct anxiolytic-antidepressant effect in MS depression", "Daily 60min concurrent with pharmacotherapy"],
        ["Maintenance", "Long-term fatigue and mood maintenance", "3–5x/week; patient self-administered home device"],
    ],
    tavns_role="taVNS (tragus, 25Hz, 200µs, 0.5mA, 30min) activates the cholinergic anti-inflammatory pathway via NTS → vagal efferents → spleen, potentially modulating CNS neuroinflammation. Also addresses MS fatigue and depression through LC-noradrenergic and raphe-serotonergic activation. Administered before NIBS sessions and as home device (2x daily). While disease modification is not expected, symptomatic anti-inflammatory benefits are documented in autoimmune conditions (Tracey 2002, Bonaz 2016).",
    combinations=[
        ("Fatigue-Dominant", [
            ["tDCS M1 + tDCS DLPFC + CES", "Motor efficiency + cognitive fatigue + thalamic dysrhythmia", "tDCS sequential → CES nightly", "MS fatigue affecting motor and cognitive function"],
            ["tDCS DLPFC + taVNS + CES", "Cognitive fatigue + anti-inflammatory + thalamic", "taVNS before tDCS; CES nightly", "Cognitive-fatigue dominant with neuroinflammation concern"],
        ]),
        ("Cognitive-Dominant", [
            ["tDCS DLPFC + TPS DLPFC + taVNS", "Convergent prefrontal enhancement + serotonergic/noradrenergic", "taVNS → TPS → tDCS", "Processing speed and working memory deficits"],
        ]),
        ("Spasticity-Dominant", [
            ["tDCS M1 + TPS M1/SMA", "Convergent motor cortex facilitation for spasticity", "TPS → tDCS immediately after", "Spasticity with motor impairment"],
        ]),
        ("Pain-Dominant", [
            ["tDCS M1 + TPS ACC + CES", "Pain threshold elevation + affective pain + thalamic", "TPS → tDCS; CES nightly", "Central neuropathic pain in MS"],
        ]),
        ("Mood/Depression", [
            ["tDCS DLPFC + taVNS + CES", "Antidepressant + vagal-serotonergic + anxiolytic", "taVNS before tDCS; CES nightly", "MS-related depression"],
        ]),
        ("PPMS/SPMS Progressive", [
            ["tDCS M1 + CES + taVNS", "Motor maintenance + fatigue + anti-inflammatory", "tDCS + taVNS daytime; CES nightly", "Progressive MS — symptom maintenance focus"],
        ]),
    ],
    combination_summary=[
        ["Fatigue-Dominant", "tDCS M1 + DLPFC + CES", "Motor + cognitive fatigue + thalamic", "tDCS sequential → CES nightly", "MS fatigue", "IIb"],
        ["Cognitive-Dominant", "tDCS DLPFC + TPS + taVNS", "Prefrontal + serotonergic", "taVNS → TPS → tDCS", "Processing speed deficit", "IIb"],
        ["Spasticity-Dominant", "tDCS M1 + TPS M1/SMA", "Convergent motor facilitation", "TPS → tDCS", "Spasticity/motor", "IIb"],
        ["Pain-Dominant", "tDCS M1 + TPS ACC + CES", "Pain + affective + thalamic", "TPS → tDCS + CES", "Central neuropathic pain", "IIb"],
        ["Mood/Depression", "tDCS DLPFC + taVNS + CES", "Antidepressant triple", "taVNS → tDCS + CES", "MS depression", "IIb"],
        ["Progressive MS", "tDCS M1 + CES + taVNS", "Motor + fatigue + anti-inflammatory", "tDCS + taVNS + CES nightly", "SPMS/PPMS", "IIb"],
    ],
    outcomes=[
        ["Modified Fatigue Impact Scale (MFIS)", "MS-specific fatigue", "Baseline, 4wk, 8wk", "≥5-point reduction (MCID)"],
        ["Symbol Digit Modalities Test (SDMT)", "Processing speed (cognitive primary)", "Baseline, 4wk, 8wk", "≥4-point improvement (MCID)"],
        ["Modified Ashworth Scale (MAS)", "Spasticity", "Baseline, 4wk, 8wk", "≥1-point reduction"],
        ["Timed 25-Foot Walk (T25FW)", "Walking speed", "Baseline, 4wk, 8wk", "≥20% improvement"],
        ["NRS Pain Score", "Central pain intensity", "Weekly", "≥30% reduction"],
        ["PHQ-9", "Depression", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["Pittsburgh Sleep Quality Index (PSQI)", "Sleep quality", "Baseline, 4wk, 8wk", "≥3-point improvement"],
        ["9-Hole Peg Test (9HPT)", "Upper limb dexterity", "Baseline, 4wk, 8wk", "≥20% improvement in dominant hand"],
        ["Patient Global Impression of Change", "Overall response", "4wk, 8wk", "≥5/7 (much improved)"],
    ],
    outcomes_abbrev=[
        "MFIS: Modified Fatigue Impact Scale",
        "SDMT: Symbol Digit Modalities Test",
        "MAS: Modified Ashworth Scale",
        "T25FW: Timed 25-Foot Walk",
        "9HPT: 9-Hole Peg Test",
        "EDSS: Expanded Disability Status Scale",
        "PSQI: Pittsburgh Sleep Quality Index",
        "PHQ-9: Patient Health Questionnaire-9",
        "RRMS/SPMS/PPMS: Relapsing-Remitting/Secondary Progressive/Primary Progressive MS",
        "DMT: Disease-Modifying Therapy",
    ],
)

# ---------------------------------------------------------------------------
# CONDITION: AUTISM SPECTRUM DISORDER (ASD)
# ---------------------------------------------------------------------------
CONDITIONS["asd"] = dict(
    name="Autism Spectrum Disorder",
    icd10="F84.0",
    doc_num="SPG-AU-001",
    tps_status="Investigational",
    tdcs_class="Class IIb",
    tps_class="Investigational",
    tavns_class="Class IIb",
    ces_class="Class IIb",
    inclusion=[
        "DSM-5 ASD diagnosis confirmed by ADOS-2 and ADI-R",
        "Age 12–45 years; Level 1 or Level 2 ASD (not requiring substantial support for communication)",
        "Baseline ADOS-2 Social Affect score ≥8 or RRB score ≥3",
        "Capable of informed consent (or assent with guardian consent)",
        "Target symptoms including social cognition deficits, repetitive behaviors, anxiety, or executive dysfunction",
        "Stable behavioral/pharmacological intervention for ≥4 weeks prior to enrollment",
    ],
    exclusion=[
        "Level 3 ASD (requiring very substantial support); nonverbal without AAC",
        "Active seizure disorder not controlled with medication (10–30% ASD comorbidity)",
        "Metallic implants in head/neck or implanted neurostimulator",
        "Active psychiatric crisis (psychosis, suicidality)",
        "Pregnancy",
        "Severe intellectual disability (IQ <50) precluding consent/assent and participation",
        "Recent medication change (<4 weeks) that could confound outcomes",
    ],
    discussion=[
        "ASD involves atypical development of large-scale brain networks: reduced long-range connectivity (social brain) and enhanced local connectivity (repetitive behavior networks).",
        "The social brain network — mPFC, TPJ, STS, amygdala — shows hypoconnectivity and atypical activation during social information processing.",
        "Mirror neuron system dysfunction in IFG and inferior parietal lobule contributes to social imitation and empathy deficits.",
        "Frontal lobe executive dysfunction (DLPFC hypoactivation) drives cognitive inflexibility, repetitive behaviors, and poor set-shifting.",
        "tDCS targeting DLPFC (anodal) improves cognitive flexibility, social cognition, and response inhibition.",
        "TPS over mPFC and DLPFC targets social brain network nodes and executive control regions.",
        "taVNS modulates autonomic nervous system dysregulation (common in ASD) and may improve social engagement via polyvagal mechanisms.",
        "CES addresses comorbid anxiety and sleep disruption, which are highly prevalent and functionally impairing in ASD.",
    ],
    overview=[
        "Autism Spectrum Disorder (ASD) is a neurodevelopmental condition characterized by persistent deficits in social communication and interaction, and restricted/repetitive patterns of behavior, interests, or activities. With a global prevalence of approximately 1-2%, ASD spans a wide spectrum of functional abilities and support needs. Core neurobiological features include atypical large-scale brain network organization: hypoconnectivity of long-range social brain networks and hyperconnectivity of local sensorimotor networks.",
        "The SOZO NIBS protocol for ASD targets core symptom domains through network-specific neuromodulation: social cognition (mPFC/TPJ/STS), executive function (DLPFC), repetitive behaviors (SMA/striatum), anxiety (taVNS/CES), and sensory processing. All NIBS is embedded within a behavioral/therapeutic framework (social skills training, CBT, ABA).",
        "Four clinical phenotypes guide protocol selection: social cognition and communication dominant, executive dysfunction and rigidity dominant, anxiety and sensory-overload dominant, and repetitive behavior dominant. The protocol is designed for adolescents and adults (12+) with Level 1 or Level 2 ASD capable of engaging with behavioral therapy.",
    ],
    pathophysiology=[
        "Atypical large-scale network organization: reduced long-range cortico-cortical connectivity between frontal and posterior social brain regions; overconnected local networks underlying repetitive behaviors.",
        "Social brain network hypoactivation: mPFC, TPJ, and superior temporal sulcus show reduced activity during social cognition tasks; mirror neuron system (IFG/IPL) shows atypical activation.",
        "DLPFC executive dysfunction: reduced activation during set-shifting, inhibitory control, and cognitive flexibility tasks; drives rigid, inflexible behavior.",
        "Amygdala hyperreactivity: exaggerated arousal response to social stimuli; contributes to social anxiety and avoidance in ASD.",
        "Autonomic dysregulation: reduced parasympathetic tone (low HRV); heightened sympathetic arousal; impaired social engagement via polyvagal system.",
        "Sensory processing differences: atypical sensory gating and integration in S1, S2, and sensory association areas; drives sensory over/under-reactivity.",
    ],
    std_treatment=[
        "Behavioral: Applied Behavior Analysis (ABA), Naturalistic Developmental Behavioral Interventions (NDBI)",
        "Communication: speech-language therapy, AAC (augmentative and alternative communication)",
        "Social skills: social skills training groups, social cognition interventions",
        "Pharmacological (comorbidities): risperidone/aripiprazole (irritability, aggression; FDA-approved), SSRIs (anxiety, repetitive behaviors), stimulants (ADHD comorbidity)",
        "Psychological: CBT for anxiety, Cognitive Behavioral Social Skills Training (CBSST)",
    ],
    symptoms=[
        ["Social Communication Deficits", "Reduced reciprocal social interaction, nonverbal communication difficulties", "100% (defining)", "APA DSM-5 2013"],
        ["Restricted Interests", "Intense, narrow special interests; rigid adherence to routines", "100% (defining)", "APA DSM-5 2013"],
        ["Repetitive Behaviors", "Stereotypies, echolalia, self-stimulatory behaviors (stimming)", "100% (defining)", "APA DSM-5 2013"],
        ["Anxiety", "Comorbid anxiety disorders (GAD, social anxiety, specific phobias)", "40–50%", "White 2009"],
        ["ADHD Features", "Inattention, hyperactivity, impulsivity (comorbid ADHD 50–70%)", "50–70%", "Simonoff 2008"],
        ["Sensory Reactivity", "Hyper/hyposensitivity across sensory modalities", "90%", "Marco 2011"],
        ["Executive Dysfunction", "Rigid thinking, poor set-shifting, cognitive inflexibility", "70%", "Hill 2004"],
        ["Sleep Disturbance", "Insomnia, irregular sleep-wake rhythm (60–80%)", "60–80%", "Mannion 2014"],
        ["Irritability/Aggression", "Meltdowns, self-injurious behavior (in more affected individuals)", "25%", "Sukhodolsky 2008"],
        ["Intellectual Disability", "Comorbid ID (~30%); does not preclude NIBS at Level 1–2", "30%", "Baio 2018"],
    ],
    brain_regions=[
        ["mPFC (BA10/11)", "Social cognition, mentalizing, theory of mind; hypoactive in ASD", "tDCS anodal, TPS", "Castelli 2002"],
        ["Temporoparietal Junction (TPJ)", "Theory of mind, perspective-taking; hypoconnected", "TPS, tDCS", "Saxe 2003"],
        ["Superior Temporal Sulcus (STS)", "Biological motion, face processing, social perception", "TPS", "Pelphrey 2005"],
        ["DLPFC", "Executive function, cognitive flexibility, working memory", "tDCS anodal", "Hill 2004"],
        ["Inferior Frontal Gyrus (IFG)", "Mirror neuron system, language, social imitation", "tDCS anodal, TPS", "Iacoboni 2006"],
        ["Amygdala", "Social-emotional arousal, threat detection; hyperreactive in ASD", "taVNS", "Baron-Cohen 2000"],
        ["Supplementary Motor Area", "Repetitive motor behaviors, stereotypies", "tDCS cathodal", "Mostofsky 2009"],
        ["Inferior Parietal Lobule (IPL)", "Mirror system, imitation, somatosensory social", "TPS", "Iacoboni 2006"],
    ],
    brainstem=[
        ["Locus Coeruleus", "Arousal, alertness, noradrenergic tone; dysregulated in ASD", "taVNS, CES", "Mehler 2009"],
        ["Nucleus Tractus Solitarius", "Vagal integration; autonomic-social interface (polyvagal)", "taVNS direct", "Porges 2007"],
        ["Raphe Nuclei", "Serotonergic tone; serotonin synthesis atypical in ASD", "taVNS", "Chugani 1997"],
        ["Superior Colliculus", "Visual orienting, joint attention; atypical in ASD", "None direct", "Senju 2005"],
        ["Cerebellum (Dentate)", "Social prediction errors, sensorimotor integration; affected in ASD", "TPS indirect", "Fatemi 2012"],
    ],
    phenotypes=[
        ["Social Cognition Dominant", "Core social-communication deficits; mPFC/TPJ/STS hypoactivation", "Limited reciprocity, eye contact, ToM deficits", "tDCS DLPFC + TPS mPFC/STS + taVNS"],
        ["Executive Dysfunction/Rigidity", "Cognitive inflexibility, perseveration, rigid routines; DLPFC deficit", "Meltdowns at transitions, restricted problem-solving", "tDCS DLPFC anodal + TPS DLPFC"],
        ["Anxiety/Sensory-Overload", "Comorbid anxiety (40%); sensory hypersensitivity driving overwhelm", "Meltdowns, avoidance, sensory distress", "taVNS + CES + tDCS DLPFC"],
        ["Repetitive Behavior Dominant", "Stereotypies, compulsions, restricted interests; SMA/striatal hyperactivation", "Motor and verbal stereotypies, rigid routines", "tDCS SMA cathodal + TPS SMA + behavioral"],
    ],
    symptom_map=[
        ["Social cognition deficits", "mPFC/TPJ/STS", "tDCS DLPFC + TPS mPFC", "Fecteau 2011, Schneider 2017"],
        ["Theory of mind deficit", "TPJ/mPFC", "TPS TPJ + tDCS DLPFC", "Saxe 2003"],
        ["Executive dysfunction", "DLPFC", "tDCS DLPFC anodal", "Hill 2004"],
        ["Repetitive behaviors", "SMA/Striatum", "tDCS SMA cathodal", "Mostofsky 2009"],
        ["Anxiety", "Amygdala/DLPFC", "taVNS + CES + tDCS DLPFC", "White 2009"],
        ["Sleep disruption", "Thalamus/LC", "CES Alpha-Stim nightly", "Mannion 2014"],
        ["Sensory reactivity", "S1/S2/STS", "tDCS cathodal S1 (investigational)", "Marco 2011"],
        ["Autonomic dysregulation", "Brainstem/Vagal", "taVNS polyvagal", "Porges 2007"],
        ["Irritability", "Amygdala/OFC", "taVNS + CES", "Sukhodolsky 2008"],
    ],
    montage=[
        ["Social Cognition Dominant", "Anodal DLPFC (F3, 1.5mA, 20min)", "mPFC/STS (250 pulses)", "taVNS 30min before; social skills concurrent"],
        ["Executive Dysfunction", "Anodal DLPFC (F3, 1.5mA, 20min)", "DLPFC (250 pulses)", "Behavioral flexibility training concurrent"],
        ["Anxiety/Sensory-Overload", "Anodal DLPFC (F3, 1.5mA, 20min)", "None initial", "taVNS primary + CES nightly; sensory desensitization"],
        ["Repetitive Behavior", "Cathodal SMA (Cz, 1.5mA, 20min)", "SMA (200 pulses)", "Behavioral support concurrent; habit reversal"],
        ["Combined Social + Executive", "Anodal DLPFC (F3, 1.5mA) + TPS", "mPFC + DLPFC (300 pulses)", "taVNS before; comprehensive behavioral program"],
        ["Comorbid ADHD", "Anodal DLPFC (F3, 1.5mA, 20min)", "DLPFC (200 pulses)", "Stimulant maintained if on; structured routine"],
        ["Adolescent (12–17)", "Anodal DLPFC (F3, 1mA, 15min)", "None initial (TPS later)", "taVNS + CES; parent-involved therapy"],
        ["Maintenance Phase", "Anodal DLPFC (F3, 1.5mA, 20min)", "None", "CES 3x/week; taVNS daily home"],
        ["High Anxiety/Meltdowns", "Anodal DLPFC (F3, 1.5mA, 20min)", "None", "taVNS primary; CES nightly; crisis plan"],
        ["Sham Protocol", "F3 anodal (30s ramp, 0mA, 20min)", "None", "Sham control condition"],
    ],
    tdcs_protocols=[
        ["AU-tDCS-01", "DLPFC Anodal (Social Cognition/Executive)", "F3", "F4", "1.5mA, 20min, 10 sessions", "Schneider 2017"],
        ["AU-tDCS-02", "SMA Cathodal (Repetitive Behaviors)", "Fp1", "Cz (SMA)", "1.5mA, 20min, 10 sessions", "Amatachaya 2014"],
        ["AU-tDCS-03", "Combined DLPFC + SMA", "F3 anodal; Cz cathodal", "F4; Fpz", "1.5mA each, 15 sessions", "Expert consensus 2023"],
        ["AU-tDCS-04", "mPFC Anodal (Social Brain)", "Fpz anodal", "Oz", "1mA, 20min, 10 sessions", "Fecteau 2011"],
        ["AU-tDCS-05", "Maintenance", "F3 anodal", "F4", "1.5mA, 20min, 2x/week", "Bikson 2016"],
        ["AU-tDCS-06", "Adolescent Adjusted", "F3 anodal", "F4", "1mA, 15min, 10 sessions", "Schneider 2017"],
        ["AU-tDCS-07", "Cathodal S1 (Sensory Reactivity — Research)", "C3/C4 cathodal", "Fpz", "1mA, 15min, investigational", "Marco 2011"],
        ["AU-tDCS-08", "Sham", "F3", "F4", "30s ramp, 0mA, 20min", "Schneider 2017"],
    ],
    plato_protocols=[
        ["AU-PLATO-01", "DLPFC Anodal", "F3/F4", "Left DLPFC", "F3", "1.5mA", "N/A", "Continuous 20min"],
        ["AU-PLATO-02", "SMA Cathodal", "Cz/Fpz", "SMA", "Cz", "1.5mA", "N/A", "Cathodal 20min"],
        ["AU-PLATO-03", "Combined DLPFC+SMA", "F3; Cz", "DLPFC + SMA", "F3; Cz", "1.5mA each", "N/A", "Sequential"],
        ["AU-PLATO-04", "mPFC Anodal", "Fpz/Oz", "mPFC", "Fpz", "1.0mA", "N/A", "Continuous 20min"],
        ["AU-PLATO-05", "Maintenance", "F3", "DLPFC", "F3", "1.5mA", "N/A", "2x/week"],
        ["AU-PLATO-06", "Theta tACS Social", "F3", "DLPFC", "F3", "1.0mA", "6Hz", "tACS 20min"],
        ["AU-PLATO-07", "Adolescent", "F3", "DLPFC", "F3", "1.0mA", "N/A", "15min"],
        ["AU-PLATO-08", "Sham", "F3", "DLPFC", "F3", "0mA", "N/A", "30s ramp"],
    ],
    tps_protocols=[
        ["AU-TPS-01", "mPFC Social Brain", "mPFC (BA10/11)", "250 pulses, 0.2Hz, 0.2mJ/mm²", "3 sessions/week", "Castelli 2002"],
        ["AU-TPS-02", "DLPFC Executive", "DLPFC (BA46)", "250 pulses, 0.2Hz", "3 sessions/week", "Hill 2004"],
        ["AU-TPS-03", "STS Social Perception", "Superior Temporal Sulcus", "200 pulses, 0.2Hz", "2 sessions/week", "Pelphrey 2005"],
        ["AU-TPS-04", "SMA Repetitive Behavior", "SMA (BA6)", "200 pulses, 0.2Hz", "2 sessions/week", "Mostofsky 2009"],
        ["AU-TPS-05", "Full Social Brain Protocol", "mPFC + DLPFC + STS sequential", "400 pulses total", "5 sessions/week x2 weeks", "Expert consensus 2024"],
    ],
    ces_role=[
        ["Anxiety Comorbidity", "Primary modality for comorbid anxiety in ASD; reduces arousal and reactivity", "Daily 60min; before stressful activities"],
        ["Sleep Disruption", "Improves sleep onset and continuity; highly prevalent in ASD", "Nightly 60min; titrate to sleep quality"],
        ["Sensory Overload Recovery", "Reduces post-meltdown hyperarousal and sensory distress", "As needed after sensory overload events"],
        ["Maintenance", "Long-term anxiety and arousal management", "3–5x/week; patient/caregiver managed"],
    ],
    tavns_role="taVNS (tragus, 25Hz, 200µs, 0.5mA, 30min) targets the polyvagal system: activating NTS → vagal efferents → amygdala/LC pathway to enhance parasympathetic tone (measured by HRV), reduce social anxiety, and support the social engagement system. Low HRV is a documented biomarker of autonomic dysregulation in ASD. Administered before NIBS sessions and as home device (2x daily). Social engagement facilitation via polyvagal normalization is the primary mechanism (Porges 2007, Patriquin 2019).",
    combinations=[
        ("Social Cognition Dominant", [
            ["tDCS DLPFC + TPS mPFC/STS + taVNS", "DLPFC executive + social brain nodes + polyvagal engagement", "taVNS → TPS → tDCS; pre-social skills session", "Social communication deficits, ToM impairment"],
        ]),
        ("Executive Dysfunction/Rigidity", [
            ["tDCS DLPFC + TPS DLPFC", "Convergent prefrontal excitability enhancement", "TPS → tDCS; pre-behavioral flexibility training", "Cognitive rigidity, perseveration"],
        ]),
        ("Anxiety/Sensory-Overload", [
            ["taVNS + CES + tDCS DLPFC", "Polyvagal + thalamic + prefrontal anxiety management", "taVNS → tDCS; CES nightly", "Comorbid anxiety, sensory meltdowns"],
        ]),
        ("Repetitive Behavior Dominant", [
            ["tDCS SMA cathodal + TPS SMA + behavioral", "SMA inhibition — motor and behavioral level", "TPS → tDCS; habit reversal concurrent", "Stereotypies, repetitive motor behaviors"],
        ]),
        ("Comprehensive ASD Protocol", [
            ["tDCS DLPFC + TPS mPFC/STS + taVNS + CES", "Social + executive + polyvagal + anxiolytic", "taVNS → TPS → tDCS + CES nightly", "Multi-domain ASD with anxiety"],
        ]),
        ("Adolescent Protocol", [
            ["tDCS DLPFC (low) + taVNS + CES", "Conservative: DLPFC + polyvagal + anxiolytic", "taVNS before tDCS; CES nightly", "Age 12–17 ASD with anxiety/executive deficit"],
        ]),
    ],
    combination_summary=[
        ["Social Cognition", "tDCS DLPFC + TPS mPFC/STS + taVNS", "Social brain + polyvagal", "taVNS → TPS → tDCS + social skills", "Social-communication deficits", "IIb"],
        ["Executive/Rigidity", "tDCS DLPFC + TPS DLPFC", "Convergent prefrontal", "TPS → tDCS + behavioral", "Cognitive inflexibility", "IIb"],
        ["Anxiety/Sensory", "taVNS + CES + tDCS DLPFC", "Polyvagal + thalamic + prefrontal", "taVNS → tDCS + CES nightly", "Anxiety, sensory overload", "IIb"],
        ["Repetitive Behaviors", "tDCS SMA cathodal + TPS SMA", "SMA inhibition convergent", "TPS → tDCS + behavioral", "Stereotypies, rituals", "IIb"],
        ["Comprehensive", "All four modalities", "Multi-domain network normalization", "taVNS → TPS → tDCS + CES", "Multi-domain ASD", "IIb"],
        ["Adolescent", "tDCS DLPFC (low) + taVNS + CES", "Conservative comprehensive", "taVNS → tDCS + CES", "Age 12–17 ASD", "IIb"],
    ],
    outcomes=[
        ["ADOS-2 Social Affect Score", "Social communication severity", "Baseline, 8wk, 16wk", "≥3-point reduction"],
        ["Social Responsiveness Scale-2 (SRS-2)", "Social responsiveness (parent/self-report)", "Baseline, 4wk, 8wk", "≥10-point T-score reduction"],
        ["BRIEF-2 Executive Composite", "Executive function", "Baseline, 4wk, 8wk", "≥10-point reduction in T-score"],
        ["Repetitive Behavior Scale-Revised (RBS-R)", "Repetitive behavior frequency/severity", "Baseline, 4wk, 8wk", "≥15% reduction in total score"],
        ["SCARED (Anxiety)", "Comorbid anxiety", "Baseline, 4wk, 8wk", "≥10-point reduction"],
        ["Pittsburgh Sleep Quality Index (PSQI)", "Sleep quality", "Baseline, 4wk, 8wk", "≥3-point improvement"],
        ["Heart Rate Variability (RMSSD)", "Autonomic/polyvagal regulation", "Baseline, 4wk, 8wk", "≥10% increase"],
        ["Vineland Adaptive Behavior Scales-3", "Adaptive functioning", "Baseline, 8wk", "≥5-point increase in composite"],
        ["Caregiver Stress Index (CSI)", "Caregiver burden", "Baseline, 8wk", "≥20% reduction"],
    ],
    outcomes_abbrev=[
        "ADOS-2: Autism Diagnostic Observation Schedule-2",
        "SRS-2: Social Responsiveness Scale-2",
        "BRIEF-2: Behavior Rating Inventory of Executive Function-2",
        "RBS-R: Repetitive Behavior Scale-Revised",
        "SCARED: Screen for Child Anxiety Related Disorders",
        "PSQI: Pittsburgh Sleep Quality Index",
        "HRV/RMSSD: Heart Rate Variability / Root Mean Square of Successive Differences",
        "ToM: Theory of Mind",
        "mPFC: Medial Prefrontal Cortex",
        "STS: Superior Temporal Sulcus",
    ],
)

# ---------------------------------------------------------------------------
# CONDITION: LONG COVID (Post-Acute Sequelae of SARS-CoV-2)
# ---------------------------------------------------------------------------
CONDITIONS["long_covid"] = dict(
    name="Long COVID (Post-Acute Sequelae of SARS-CoV-2)",
    icd10="U09.9",
    doc_num="SPG-LC-001",
    tps_status="Investigational",
    tdcs_class="Class IIb",
    tps_class="Investigational",
    tavns_class="Class IIb",
    ces_class="Class IIb",
    inclusion=[
        "WHO-defined Long COVID: symptoms persisting ≥12 weeks post-SARS-CoV-2 infection not explained by alternative diagnosis",
        "Symptom burden including cognitive impairment, fatigue, autonomic dysfunction, or mood symptoms",
        "Cognitive domain: MoCA ≤25 or MFIS ≥38 at baseline",
        "Age 18–65; capable of informed consent",
        "Confirmed prior SARS-CoV-2 infection (PCR, antigen, or serology)",
        "Stable symptoms for ≥4 weeks (not in acute post-COVID relapse)",
    ],
    exclusion=[
        "Active SARS-CoV-2 infection (acute phase)",
        "Metallic implants in head/neck or implanted neurostimulator",
        "Pregnancy",
        "Pre-existing neurological or psychiatric diagnosis that adequately explains symptoms",
        "Severe post-exertional malaise precluding protocol participation (energy envelope too limited)",
        "Active myocarditis or pericarditis (POTS screening required)",
        "EDSS >6 equivalent disability from Long COVID",
    ],
    discussion=[
        "Long COVID involves multisystem pathology including neuroinflammation, cerebrovascular microangiopathy, autonomic nervous system dysfunction (POTS), mast cell activation, and mitochondrial impairment.",
        "Cognitive impairment (brain fog) reflects frontal-network disruption, thalamic microstructural damage, and neuroinflammatory cytokine effects on synaptic function.",
        "Autonomic dysfunction (POTS, orthostatic intolerance) is a dominant mechanism of fatigue and exercise intolerance; taVNS targets the autonomic-cardiovascular interface.",
        "tDCS targeting DLPFC addresses cognitive fog through direct prefrontal excitability enhancement and network normalization.",
        "TPS over the frontal and thalamic projections targets network-level cognitive recovery.",
        "taVNS offers anti-inflammatory, autonomic regulatory, and direct cognitive-fatigue benefits in Long COVID.",
        "CES addresses the insomnia, anxiety, and thalamic dysrhythmia components of Long COVID symptom burden.",
        "Pacing and energy conservation are essential co-treatment principles; NIBS intensity must be calibrated to individual energy envelope.",
    ],
    overview=[
        "Long COVID (Post-Acute Sequelae of SARS-CoV-2, PASC) is defined by the WHO as symptoms persisting for ≥12 weeks following SARS-CoV-2 infection, not explained by an alternative diagnosis. Affecting an estimated 10-30% of COVID-19 survivors, Long COVID represents a major global health burden with neurological symptoms — cognitive impairment, fatigue, autonomic dysfunction, headache, and mood disorders — among the most prevalent and disabling.",
        "The SOZO NIBS protocol for Long COVID targets the neurological and autonomic mechanisms underlying the most common symptom domains: prefrontal-thalamic cognitive network disruption (brain fog), autonomic dysregulation (POTS/HRV impairment), neuroinflammatory fatigue, and insomnia/mood comorbidity.",
        "Four clinical phenotypes are identified: cognitive-dominant (brain fog), autonomic-dominant (POTS/fatigue), neuropsychiatric (anxiety/depression/insomnia), and post-exertional malaise (PEM)-dominant. Energy envelope management is mandatory across all phenotypes.",
    ],
    pathophysiology=[
        "Neuroinflammation: activated microglia, astrogliosis, and inflammatory cytokine (IL-6, TNF-alpha) effects on synaptic function; white matter microstructural damage on advanced MRI.",
        "Thalamic microangiopathy: COVID-19-related cerebrovascular endotheliopathy and microclot formation disrupt thalamo-cortical relay networks, causing cognitive fog.",
        "Autonomic nervous system dysfunction: reduced HRV, POTS, sympathetic overdrive, and impaired baroreceptor reflex; primary driver of fatigue and exercise intolerance.",
        "Mitochondrial dysfunction: impaired cellular energy production amplifies fatigue, cognitive impairment, and post-exertional malaise.",
        "Reactivated viral reservoirs: persistent spike protein or viral RNA in CNS/gut may maintain ongoing neuroinflammatory signaling.",
        "DLPFC-CEN hypoactivation: direct neuroinflammatory and vascular effects reduce frontal network efficiency, producing cognitive fog and executive dysfunction.",
    ],
    std_treatment=[
        "Symptom management: pacing/activity management (energy envelope), graded exercise (POTS subtypes only — not PEM-dominant)",
        "Autonomic: compression garments, salt/fluid loading, beta-blockers or ivabradine (POTS)",
        "Cognitive rehabilitation: cognitive training, compensatory strategies, occupational therapy",
        "Pharmacological: low-dose naltrexone (LDN, anti-inflammatory), antihistamines (mast cell activation), SSRIs for mood/autonomic",
        "Research interventions: Paxlovid extended (viral reservoir), BC007 (autoantibodies), hyperbaric oxygen therapy",
    ],
    symptoms=[
        ["Cognitive Impairment (Brain Fog)", "Concentration, memory, word-finding, mental fatigue", "50–70%", "Taquet 2021"],
        ["Fatigue", "Debilitating, disproportionate fatigue; post-exertional malaise", "58–70%", "Davis 2021"],
        ["Post-Exertional Malaise (PEM)", "Symptom worsening ≥24h after minimal physical/mental exertion", "40–60%", "Davis 2021"],
        ["Dyspnea", "Breathlessness persisting beyond acute illness", "25%", "Nalbandian 2021"],
        ["Sleep Disturbance", "Insomnia, hypersomnia, unrefreshing sleep", "40%", "Huang 2021"],
        ["Headache", "Persistent new-onset or worsened headache", "30%", "Nalbandian 2021"],
        ["Anxiety/Depression", "Comorbid mood disorders (partly neuroinflammatory)", "30%", "Taquet 2021"],
        ["Autonomic Dysfunction (POTS)", "Orthostatic tachycardia, palpitations, dizziness", "20–30%", "Blitshteyn 2021"],
        ["Sensory Symptoms", "Paresthesias, dysesthesias, anosmia, parosmia", "30%", "Nalbandian 2021"],
        ["Pain", "Myalgia, arthralgia, headache; central sensitization features", "35%", "Davis 2021"],
    ],
    brain_regions=[
        ["DLPFC", "Cognitive control, working memory; hypoactivated in brain fog", "tDCS anodal", "Taquet 2021"],
        ["Thalamus", "Cognitive relay; microangiopathy-damaged; thalamocortical dysrhythmia", "CES, TPS", "Qin 2021"],
        ["ACC", "Fatigue perception, cognitive effort, conflict; hyperactivated in Long COVID", "TPS", "Yaldizli 2011"],
        ["Default Mode Network", "Self-referential processing; disrupted DMN-CEN anticorrelation", "tDCS DLPFC", "Taquet 2021"],
        ["Hippocampus", "Episodic memory; neuroinflammatory vulnerability", "taVNS", "Heneka 2020"],
        ["Insula", "Interoception, autonomic integration, dyspnea; hyperactivated", "TPS", "Craig 2002"],
        ["Brainstem (NTS/LC)", "Autonomic regulation; COVID-19 direct brainstem pathology", "taVNS", "Guedj 2021"],
        ["White Matter Networks", "CEN and DMN connectivity reduced by microangiopathy", "tDCS DLPFC indirect", "Lu 2022"],
    ],
    brainstem=[
        ["Locus Coeruleus", "Autonomic arousal, fatigue, noradrenergic tone; COVID-19 direct target", "taVNS, CES", "Guedj 2021"],
        ["Nucleus Tractus Solitarius", "Vagal integration, autonomic relay; COVID-19 brainstem pathology", "taVNS direct", "Tracey 2002"],
        ["Dorsal Vagal Complex", "Autonomic-inflammatory interface; key Long COVID target", "taVNS direct", "Bonaz 2016"],
        ["Raphe Nuclei", "Serotonergic tone; anxiety/mood in Long COVID", "taVNS", "Feinstein 2011"],
        ["Nucleus Ambiguus", "Vagal cardiac efferents; HRV generation; POTS-relevant", "taVNS indirect", "Porges 2007"],
    ],
    phenotypes=[
        ["Cognitive-Dominant (Brain Fog)", "Processing speed, memory, word-finding deficits; DLPFC-thalamic network disruption", "Cognitive impairment, work/study incapacity", "tDCS DLPFC + TPS + CES"],
        ["Autonomic-Dominant (POTS/Fatigue)", "Orthostatic intolerance, tachycardia, fatigue; ANS dysregulation primary", "Exercise intolerance, palpitations, dizziness", "taVNS primary + CES + tDCS DLPFC low"],
        ["Neuropsychiatric (Anxiety/Depression/Insomnia)", "Mood disorders, anxiety, insomnia; neuroinflammatory-serotonergic", "Depression, panic, sleep disruption", "tDCS DLPFC + taVNS + CES"],
        ["PEM-Dominant", "Post-exertional malaise >24h with minimal exertion; mitochondrial/metabolic", "Activity restriction, severe fatigue crashes", "CES primary + taVNS; tDCS only if tolerated"],
    ],
    symptom_map=[
        ["Brain fog/cognitive impairment", "DLPFC/Thalamus", "tDCS DLPFC + CES", "Taquet 2021"],
        ["Fatigue", "ACC/Thalamus/Brainstem", "CES + taVNS", "Davis 2021"],
        ["Autonomic dysfunction/POTS", "Brainstem/NTS/LC", "taVNS primary", "Blitshteyn 2021"],
        ["Anxiety/Depression", "DLPFC/Limbic", "tDCS DLPFC + taVNS + CES", "Taquet 2021"],
        ["Sleep disruption", "Thalamus/LC", "CES nightly", "Huang 2021"],
        ["Headache", "Thalamus/Cortical", "TPS investigational; CES", "Nalbandian 2021"],
        ["Neuroinflammation", "CNS-wide", "taVNS anti-inflammatory", "Bonaz 2016"],
        ["PEM", "Systemic/Mitochondrial", "CES only; strict pacing", "Davis 2021"],
        ["Sensory symptoms", "Insula/Thalamus", "TPS insula + CES", "Nalbandian 2021"],
    ],
    montage=[
        ["Cognitive-Dominant", "Anodal DLPFC (F3, 1.5–2mA, 20min)", "DLPFC + Thalamic (250 pulses)", "CES nightly; pacing essential"],
        ["Autonomic-Dominant", "Anodal DLPFC (F3, 1mA, 15min) — low dose", "None initial", "taVNS primary 2x daily; CES nightly; POTS management concurrent"],
        ["Neuropsychiatric", "Anodal DLPFC (F3, 1.5mA, 20min)", "None initial", "taVNS + CES nightly; SSRI if prescribed maintained"],
        ["PEM-Dominant", "Avoid tDCS if causes PEM — start CES only", "None", "CES primary; taVNS only (ultra-low dose); strict energy management"],
        ["Cognitive + Autonomic Combined", "Anodal DLPFC (F3, 1mA, 15min)", "DLPFC (150 pulses max)", "taVNS before tDCS; CES nightly"],
        ["Maintenance Phase", "Anodal DLPFC (F3, 1mA, 15min)", "None", "CES 3x/week; taVNS daily home"],
        ["Severe PEM (Research)", "Not recommended — defer", "Not recommended", "CES only; monitor energy envelope weekly"],
        ["Elderly Adjusted", "Anodal DLPFC (F3, 1mA, 15min)", "TPS very low (100 pulses)", "CES + taVNS; conservative approach"],
        ["Headache Subtype", "Anodal DLPFC (F3, 1.5mA, 20min)", "None initial", "CES; TPS future consideration"],
        ["Sham Protocol", "F3 anodal (30s ramp, 0mA, 20min)", "None", "Sham control condition"],
    ],
    tdcs_protocols=[
        ["LC-tDCS-01", "DLPFC Anodal (Brain Fog/Cognitive)", "F3", "F4", "1.5–2mA, 20min, 10 sessions", "Expert consensus 2023"],
        ["LC-tDCS-02", "DLPFC Low Dose (Autonomic/PEM-sensitive)", "F3", "F4", "1mA, 15min, 10 sessions", "Pacing-adapted protocol"],
        ["LC-tDCS-03", "Bifrontal (Neuropsychiatric)", "F3 anodal", "F4 cathodal", "1.5mA, 20min, 10 sessions", "Brunoni 2013"],
        ["LC-tDCS-04", "PEM Monitoring Protocol", "F3", "F4", "Start 0.5mA, escalate if tolerated", "Energy envelope-adaptive"],
        ["LC-tDCS-05", "Maintenance", "F3 anodal", "F4", "1mA, 15min, 2x/week", "Bikson 2016"],
        ["LC-tDCS-06", "Combined DLPFC + CES adjunct", "F3 anodal", "F4", "1.5mA, 20min + CES concurrent", "Expert consensus 2024"],
        ["LC-tDCS-07", "Elderly Adjusted", "F3 anodal", "F4", "1mA, 15min, 10 sessions", "Lefaucheur 2016"],
        ["LC-tDCS-08", "Sham", "F3", "F4", "30s ramp, 0mA, 20min", "Sham control"],
    ],
    plato_protocols=[
        ["LC-PLATO-01", "DLPFC Anodal Standard", "F3/F4", "Left DLPFC", "F3", "1.5mA", "N/A", "Continuous 20min"],
        ["LC-PLATO-02", "DLPFC Low Dose", "F3/F4", "Left DLPFC", "F3", "1.0mA", "N/A", "15min continuous"],
        ["LC-PLATO-03", "Bifrontal", "F3/F4", "Bifrontal", "F3", "1.5mA", "N/A", "Sequential"],
        ["LC-PLATO-04", "HD-tDCS DLPFC", "F3 center", "4-electrode ring", "F3", "0.5mA", "N/A", "HD 15min — low dose"],
        ["LC-PLATO-05", "Maintenance", "F3", "DLPFC", "F3", "1.0mA", "N/A", "2x/week 15min"],
        ["LC-PLATO-06", "Gamma tACS Cognitive", "F3", "DLPFC", "F3", "1.0mA", "40Hz", "tACS 15min"],
        ["LC-PLATO-07", "Elderly/PEM-Sensitive", "F3", "DLPFC", "F3", "0.5mA", "N/A", "10min"],
        ["LC-PLATO-08", "Sham", "F3", "DLPFC", "F3", "0mA", "N/A", "30s ramp"],
    ],
    tps_protocols=[
        ["LC-TPS-01", "DLPFC Cognitive Network", "DLPFC (BA46)", "200 pulses, 0.2Hz, 0.15mJ/mm²", "2 sessions/week", "Expert consensus 2023"],
        ["LC-TPS-02", "ACC Fatigue Hub", "ACC (BA24/32)", "150 pulses, 0.15Hz", "2 sessions/week", "Yaldizli 2011"],
        ["LC-TPS-03", "Insula Autonomic/Sensory", "Posterior Insula", "150 pulses, 0.15Hz", "1 session/week — cautious", "Craig 2002"],
        ["LC-TPS-04", "Thalamic Projection (Indirect)", "Frontal projections to thalamus", "200 pulses, 0.2Hz", "2 sessions/week", "Expert consensus 2024"],
        ["LC-TPS-05", "PEM-Adapted Protocol", "DLPFC only — minimal exposure", "100 pulses, 0.15Hz", "1 session/week — monitor PEM", "Expert consensus 2024"],
    ],
    ces_role=[
        ["Primary Fatigue/Sleep", "Primary NIBS modality for fatigue and insomnia in Long COVID; safest in PEM-dominant", "Daily 60min; nightly for sleep"],
        ["PEM-Dominant", "Only recommended NIBS in severe PEM; monitor energy response", "Start 20min; escalate if tolerated"],
        ["Neuropsychiatric", "Anxiolytic and sleep-normalizing adjunct", "Daily 60min; before bedtime"],
        ["Maintenance", "Long-term symptom management; patient self-administered", "3–5x/week; home device program"],
    ],
    tavns_role="taVNS (tragus, 25Hz, 200µs, 0.5mA, 30min) is a cornerstone of Long COVID NIBS treatment through three mechanisms: (1) Cholinergic anti-inflammatory pathway activation — reduces neuroinflammatory cytokine burden; (2) Autonomic normalization — increases HRV, reduces sympathetic overdrive, directly addresses POTS pathophysiology; (3) Cognitive-fatigue modulation via LC-noradrenergic and raphe-serotonergic activation. Twice-daily home device use (Nemos/TENS tragus) is the recommended maintenance regimen. PEM monitoring is essential — if taVNS worsens PEM, reduce session duration to 15min or less (Bonaz 2016, Tracey 2002, Clancy 2022).",
    combinations=[
        ("Cognitive-Dominant", [
            ["tDCS DLPFC + TPS DLPFC + CES", "Convergent prefrontal + thalamic cognitive network", "TPS → tDCS; CES nightly; pacing", "Brain fog, cognitive impairment"],
            ["tDCS DLPFC + taVNS + CES", "Prefrontal + anti-inflammatory + thalamic", "taVNS before tDCS; CES nightly", "Brain fog with neuroinflammation/fatigue"],
        ]),
        ("Autonomic-Dominant", [
            ["taVNS + CES", "Dual autonomic normalization; safe in POTS", "taVNS 2x daily; CES nightly; POTS concurrent", "POTS, orthostatic intolerance, fatigue"],
            ["taVNS + tDCS DLPFC (low)", "Autonomic + minimal cognitive support", "taVNS before low-dose tDCS", "POTS with cognitive fog — energy-adapted"],
        ]),
        ("Neuropsychiatric", [
            ["tDCS DLPFC + taVNS + CES", "Antidepressant + vagal + anxiolytic-thalamic", "taVNS → tDCS + CES nightly", "Anxiety/depression/insomnia in Long COVID"],
        ]),
        ("PEM-Dominant", [
            ["CES only", "Safest modality; no exertional demand", "CES nightly 60min; strict energy management", "Severe PEM — defer other modalities"],
            ["CES + taVNS (ultra-low)", "Minimal dual anti-inflammatory + sleep", "CES nightly; taVNS 15min once daily max", "PEM with neuroinflammation"],
        ]),
        ("Headache Subtype", [
            ["tDCS DLPFC + CES", "Prefrontal + thalamic headache modulation", "tDCS daytime; CES nightly", "Persistent headache subtype"],
        ]),
        ("Comprehensive Stable Long COVID", [
            ["tDCS DLPFC + TPS + taVNS + CES", "Full neurological Long COVID protocol", "taVNS → TPS → tDCS + CES nightly", "Stable multi-domain Long COVID — not PEM-dominant"],
        ]),
    ],
    combination_summary=[
        ["Cognitive-Dominant", "tDCS DLPFC + TPS + CES", "Prefrontal + thalamic cognitive", "TPS → tDCS + CES nightly", "Brain fog", "IIb"],
        ["Autonomic-Dominant", "taVNS + CES", "Dual autonomic normalization", "taVNS 2x daily + CES nightly", "POTS/fatigue", "IIb"],
        ["Neuropsychiatric", "tDCS DLPFC + taVNS + CES", "Antidepressant + vagal + anxiolytic", "taVNS → tDCS + CES nightly", "Anxiety/depression/insomnia", "IIb"],
        ["PEM-Dominant", "CES only (± taVNS ultra-low)", "Safest; minimal exertion demand", "CES nightly; strict pacing", "Severe PEM", "IIb"],
        ["Headache Subtype", "tDCS DLPFC + CES", "Prefrontal + thalamic headache", "tDCS daytime + CES nightly", "Persistent headache", "IIb"],
        ["Comprehensive Stable", "All four modalities", "Multi-domain Long COVID", "taVNS → TPS → tDCS + CES", "Stable multi-domain", "IIb"],
    ],
    outcomes=[
        ["MoCA Score", "Cognitive screening", "Baseline, 4wk, 8wk", "≥2-point improvement"],
        ["MFIS Total Score", "Fatigue severity", "Baseline, 4wk, 8wk", "≥5-point reduction (MCID)"],
        ["Heart Rate Variability (RMSSD)", "Autonomic/vagal tone", "Baseline, 4wk, 8wk", "≥10% increase"],
        ["Orthostatic Heart Rate Test", "POTS assessment (10-min stand test)", "Baseline, 4wk, 8wk", "HR increase <30bpm at 10min stand"],
        ["PHQ-9", "Depression", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["Pittsburgh Sleep Quality Index (PSQI)", "Sleep quality", "Baseline, 4wk, 8wk", "≥3-point improvement"],
        ["Post-Exertional Malaise (DSQ-PEM)", "PEM frequency/severity", "Baseline, 4wk, 8wk", "≥30% reduction in PEM frequency"],
        ["PROMIS Cognitive Function", "Self-reported cognitive function", "Baseline, 4wk, 8wk", "≥5-point T-score improvement"],
        ["Patient Global Impression of Change", "Overall improvement", "4wk, 8wk", "≥5/7 (much improved)"],
    ],
    outcomes_abbrev=[
        "MoCA: Montreal Cognitive Assessment",
        "MFIS: Modified Fatigue Impact Scale",
        "HRV/RMSSD: Heart Rate Variability / Root Mean Square of Successive RR Differences",
        "POTS: Postural Orthostatic Tachycardia Syndrome",
        "PHQ-9: Patient Health Questionnaire-9",
        "PSQI: Pittsburgh Sleep Quality Index",
        "DSQ-PEM: DePaul Symptom Questionnaire — Post-Exertional Malaise",
        "PROMIS: Patient-Reported Outcomes Measurement Information System",
        "PEM: Post-Exertional Malaise",
        "ANS: Autonomic Nervous System",
    ],
)
print("MS + ASD + Long COVID appended OK")

# ---------------------------------------------------------------------------
# CONDITION: TINNITUS
# ---------------------------------------------------------------------------
CONDITIONS["tinnitus"] = dict(
    name="Chronic Tinnitus",
    icd10="H93.19",
    doc_num="SPG-TN-001",
    tps_status="Investigational",
    tdcs_class="Class IIb",
    tps_class="Investigational",
    tavns_class="Class IIb",
    ces_class="Class IIb",
    inclusion=[
        "Chronic tinnitus ≥6 months duration (unilateral or bilateral; tonal or noise-like)",
        "THI (Tinnitus Handicap Inventory) score ≥38 (moderate-severe) at baseline",
        "Audiological assessment confirming tinnitus etiology and excluding treatable causes",
        "Age 18–75; capable of informed consent",
        "Inadequate response to ≥1 conventional treatment (sound therapy, CBT, hearing aids)",
        "Stable audiological status for ≥3 months (no acute hearing loss)",
    ],
    exclusion=[
        "Pulsatile tinnitus (vascular/structural etiology requiring separate workup)",
        "Active Meniere disease with acute vertigo attacks",
        "Metallic implants in head/neck or implanted neurostimulator (cochlear implant — evaluate carefully)",
        "Active otological infection or post-surgical healing",
        "Pregnancy",
        "Active psychiatric comorbidity requiring immediate intervention (psychosis, active suicidality)",
        "Severe hyperacusis requiring noise-restricted environment (assess TPS compatibility)",
    ],
    discussion=[
        "Chronic tinnitus involves maladaptive central auditory network reorganization following peripheral deafferentation or acoustic trauma.",
        "The primary auditory cortex (A1) shows increased spontaneous activity and decreased lateral inhibition, generating a phantom auditory percept.",
        "The parahippocampal area and amygdala are recruited into the tinnitus network, driving the affective distress and attention capture dimensions.",
        "DLPFC hypoactivity impairs top-down suppression of the auditory phantom via the salience and executive networks.",
        "tDCS targeting auditory cortex (cathodal) reduces A1 hyperexcitability; tDCS over DLPFC (anodal) strengthens top-down suppression.",
        "TPS directly targets A1 and DLPFC to normalize the tinnitus network node hyperactivation.",
        "taVNS combined with auditory tone presentation (VNS-tone pairing) facilitates auditory cortex reorganization through neuromodulatory plasticity.",
        "CES addresses tinnitus-related insomnia and anxiety through thalamic and brainstem auditory relay modulation.",
    ],
    overview=[
        "Chronic tinnitus — the perception of sound without an external acoustic source — affects approximately 15% of the global population, with 2-3% experiencing severe, debilitating symptoms. The neurobiological substrate involves maladaptive central auditory plasticity: increased spontaneous firing in primary auditory cortex (A1), reduced lateral inhibition, and pathological coupling between the auditory network and distress-generating regions (amygdala, parahippocampal gyrus, prefrontal cortex).",
        "The SOZO NIBS protocol targets the tinnitus neural network through a phenotype-guided multimodal approach: tDCS (A1 cathodal / DLPFC anodal), TPS (A1/DLPFC), taVNS (auditory cortex plasticity/neuromodulation), and CES (Alpha-Stim, anxiety/sleep). The tinnitus network — not just the auditory cortex — is the target.",
        "Five clinical phenotypes guide protocol selection: audiological/peripheral (hearing loss dominant), central/network (no hearing loss; central reorganization dominant), distress-dominant (high THI, anxiety, depression), sleep-dominant (tinnitus-insomnia cycle), and hyperacusis comorbid. Each phenotype receives a tailored NIBS combination.",
    ],
    pathophysiology=[
        "Central auditory reorganization: peripheral deafferentation triggers tonotopic map reorganization in A1; neurons adjacent to deafferented frequencies expand territory, generating spurious signals perceived as tinnitus.",
        "Increased spontaneous neural activity: A1 and secondary auditory cortex (A2) show elevated baseline firing rates; reduced lateral inhibition allows phantom percept to dominate.",
        "Tinnitus distress network: amygdala, parahippocampal gyrus, and anterior cingulate are recruited into the tinnitus network, generating affective suffering and attention capture.",
        "DLPFC top-down suppression failure: prefrontal inhibition of auditory attention is impaired, allowing tinnitus percept to dominate attentional resources.",
        "Thalamocortical dysrhythmia: delta band power increases and gamma power decreases in auditory thalamus (MGN), producing a pathological oscillatory signature.",
        "Default mode network coupling: DMN incorporates tinnitus-related activity, causing persistent awareness even during mind-wandering states.",
    ],
    std_treatment=[
        "Sound therapy: broadband noise, tinnitus masking, notched music therapy",
        "Cognitive behavioral therapy (CBT) for tinnitus distress: gold-standard psychological intervention",
        "Tinnitus retraining therapy (TRT): habituation-based counseling + sound therapy",
        "Hearing aids: amplification-based tinnitus management for hearing loss subtype",
        "rTMS: 1Hz low-frequency rTMS over left temporoparietal cortex (investigational)",
        "Pharmacological: no FDA-approved drug; off-label melatonin, clonazepam, acamprosate (limited evidence)",
    ],
    symptoms=[
        ["Tinnitus Percept", "Persistent sound (ringing, buzzing, hissing, roaring) without external source", "100%", "Jastreboff 2004"],
        ["Tinnitus Distress", "Emotional distress, annoyance, suffering related to tinnitus", "80%", "Andersson 2002"],
        ["Sleep Disturbance", "Difficulty falling/staying asleep due to tinnitus", "60%", "Crummer 2004"],
        ["Anxiety", "Tinnitus-related anxiety; health anxiety; GAD features", "45%", "Andersson 2004"],
        ["Depression", "Comorbid depressive disorder", "30%", "Sullivan 1988"],
        ["Concentration Impairment", "Difficulty concentrating; tinnitus captures attention", "70%", "Andersson 2002"],
        ["Hyperacusis", "Sound sensitivity; discomfort from everyday sounds", "40%", "Andersson 2002"],
        ["Avoidance Behavior", "Avoiding quiet/social situations; compensatory behaviors", "55%", "Jastreboff 2004"],
        ["Hearing Loss", "Sensorineural hearing loss as frequent comorbidity/trigger", "70%", "Axelsson 1996"],
        ["Catastrophizing", "Tinnitus catastrophizing scale; excessive negative appraisal", "50%", "Andersson 2002"],
    ],
    brain_regions=[
        ["Primary Auditory Cortex (A1)", "Tinnitus generator; hyperactive spontaneous firing; tonotopic reorganization", "tDCS cathodal, TPS", "Eggermont 2004"],
        ["Secondary Auditory Cortex (A2)", "Auditory association; tinnitus percept elaboration", "TPS", "Lockwood 2002"],
        ["DLPFC", "Top-down suppression of auditory attention; hypoactive in tinnitus", "tDCS anodal", "Schecklmann 2011"],
        ["Parahippocampal Gyrus", "Tinnitus memory and emotional tagging; tinnitus network node", "TPS", "Lockwood 2002"],
        ["Amygdala", "Tinnitus distress, threat appraisal, affective suffering", "taVNS, TPS", "Rauschecker 2010"],
        ["ACC", "Tinnitus distress monitoring, attention capture", "TPS", "Landgrebe 2009"],
        ["Medial Geniculate Nucleus", "Auditory thalamic relay; thalamocortical dysrhythmia source", "CES", "Llinás 1999"],
        ["Nucleus Accumbens", "Reward-gating of auditory attention; tinnitus gating failure", "TPS indirect", "Rauschecker 2010"],
    ],
    brainstem=[
        ["Cochlear Nucleus", "First central auditory relay; compensatory hyperactivity post-deafferentation", "None direct NIBS", "Moller 2000"],
        ["Inferior Colliculus (IC)", "Auditory midbrain; gain control; hyperactive in tinnitus", "None direct", "Salvi 2000"],
        ["Medial Geniculate Nucleus", "Thalamic auditory relay; thalamocortical dysrhythmia generator", "CES indirect", "Llinás 1999"],
        ["Locus Coeruleus", "Noradrenergic gating of auditory attention; anxiety-arousal in tinnitus", "taVNS, CES", "Brozoski 2007"],
        ["Nucleus Tractus Solitarius", "Vagal relay; taVNS mechanism for auditory plasticity", "taVNS direct", "Engineer 2011"],
    ],
    phenotypes=[
        ["Audiological/Peripheral Dominant", "Tinnitus with significant SNHL; peripheral deafferentation primary; A1 reorganization", "Hearing loss + tinnitus; audiological management primary", "tDCS A1 cathodal + DLPFC anodal + CES"],
        ["Central/Network Dominant", "Tinnitus without significant hearing loss; central reorganization primary", "High tinnitus loudness/distress despite normal audiogram", "tDCS A1 cathodal + TPS A1/DLPFC + taVNS"],
        ["Distress-Dominant", "High THI (≥56); anxiety, depression, catastrophizing; amygdala-network activation", "Severe distress, life quality impact, depression", "tDCS DLPFC + TPS ACC/amygdala + taVNS + CES"],
        ["Sleep-Dominant", "Tinnitus-insomnia cycle; hyperarousal maintaining tinnitus awareness at night", "Severe sleep disruption; fatigue; daytime cognitive impairment", "CES primary + tDCS DLPFC + taVNS"],
        ["Hyperacusis Comorbid", "Tinnitus + sound sensitivity; auditory gain dysregulation", "Avoidance of sound environments; reduced TPS tolerance", "tDCS A1 cathodal + CES; TPS with caution"],
    ],
    symptom_map=[
        ["Tinnitus loudness", "A1/A2", "tDCS A1 cathodal + TPS A1", "Eggermont 2004"],
        ["Tinnitus distress", "Amygdala/ACC/Parahippocampal", "TPS ACC + taVNS + tDCS DLPFC", "Rauschecker 2010"],
        ["Attention capture", "DLPFC/SN", "tDCS DLPFC anodal", "Schecklmann 2011"],
        ["Anxiety", "Amygdala/DLPFC", "taVNS + CES + tDCS DLPFC", "Andersson 2004"],
        ["Sleep disruption", "Thalamus/LC", "CES primary", "Lande 2018"],
        ["Depression comorbidity", "DLPFC/Limbic", "tDCS DLPFC + taVNS", "Brunoni 2013"],
        ["Hyperacusis", "A1/A2 gain", "tDCS A1 cathodal; CES", "Andersson 2002"],
        ["Concentration impairment", "DLPFC/CEN", "tDCS DLPFC anodal", "Schecklmann 2011"],
        ["Thalamocortical dysrhythmia", "MGN/A1", "CES indirect", "Llinas 1999"],
    ],
    montage=[
        ["Audiological/Peripheral", "Cathodal A1 (T3/T4, 1.5mA, 20min) + Anodal DLPFC (F3)", "A1 (200 pulses)", "Hearing aid; CES nightly; sound therapy concurrent"],
        ["Central/Network", "Cathodal A1 (T3/T4, 1.5mA, 20min) + Anodal DLPFC (F3)", "A1 + DLPFC (300 pulses)", "taVNS-tone pairing; TRT concurrent"],
        ["Distress-Dominant", "Anodal DLPFC (F3, 2mA, 20min)", "ACC (250 pulses)", "taVNS + CES nightly; CBT concurrent essential"],
        ["Sleep-Dominant", "Anodal DLPFC (F3, 1.5mA, 20min)", "None initial", "CES primary nightly; taVNS before bed"],
        ["Hyperacusis Comorbid", "Cathodal A1 (T3/T4, 1mA, 15min)", "None (TPS limited by hyperacusis)", "CES primary; sound desensitization concurrent"],
        ["Maintenance Phase", "Anodal DLPFC (F3, 1.5mA, 20min)", "None", "CES 3x/week; taVNS daily home"],
        ["Elderly Adjusted", "Cathodal A1 (T3, 1mA, 15min)", "A1 (150 pulses)", "CES primary; hearing aid"],
        ["High THI Refractory", "Cathodal A1 + Anodal DLPFC (2mA, 20min each)", "A1 + DLPFC (400 pulses)", "All 4 modalities; intensive CBT"],
        ["Unilateral Tinnitus", "Cathodal A1 ipsilateral (T3 or T4)", "A1 ipsilateral (200 pulses)", "Lateralized protocol; taVNS standard"],
        ["Sham Protocol", "F3 anodal (30s ramp, 0mA, 20min)", "None", "Sham control condition"],
    ],
    tdcs_protocols=[
        ["TN-tDCS-01", "A1 Cathodal (Tinnitus Percept Reduction)", "Fp1 or Fp2", "T3 or T4 (A1 ipsilateral)", "1.5mA, 20min, 10 sessions", "Fregni 2006"],
        ["TN-tDCS-02", "DLPFC Anodal (Top-Down Suppression)", "F3", "F4", "2mA, 20min, 10 sessions", "Schecklmann 2011"],
        ["TN-tDCS-03", "A1 Cathodal + DLPFC Anodal Combined", "F3 anodal; T3 cathodal", "F4; Fpz", "1.5–2mA, 15 sessions", "Expert consensus 2023"],
        ["TN-tDCS-04", "Bifrontal (Distress/Depression)", "F3 anodal", "F4 cathodal", "2mA, 20min, 10 sessions", "Brunoni 2013"],
        ["TN-tDCS-05", "Maintenance", "F3 anodal", "F4", "1.5mA, 20min, 2x/week", "Bikson 2016"],
        ["TN-tDCS-06", "Low Intensity (Hyperacusis)", "T3 cathodal", "Fpz", "1mA, 15min, 10 sessions", "Adjusted"],
        ["TN-tDCS-07", "Elderly Adjusted", "T3 cathodal", "Fp1", "1mA, 15min, 10 sessions", "Lefaucheur 2016"],
        ["TN-tDCS-08", "Sham", "F3", "F4", "30s ramp, 0mA, 20min", "Schecklmann 2011"],
    ],
    plato_protocols=[
        ["TN-PLATO-01", "A1 Cathodal", "T3/Fpz", "A1 ipsilateral", "T3", "1.5mA", "N/A", "Cathodal 20min"],
        ["TN-PLATO-02", "DLPFC Anodal", "F3/F4", "Left DLPFC", "F3", "2.0mA", "N/A", "Continuous 20min"],
        ["TN-PLATO-03", "Combined A1+DLPFC", "T3; F3", "A1 + DLPFC", "T3 cathodal; F3 anodal", "1.5; 2.0mA", "N/A", "Sequential"],
        ["TN-PLATO-04", "HD-tDCS A1", "T3 center", "4-electrode ring", "T3", "0.75mA", "N/A", "HD cathodal 20min"],
        ["TN-PLATO-05", "Maintenance", "F3", "DLPFC", "F3", "1.5mA", "N/A", "2x/week"],
        ["TN-PLATO-06", "Alpha tACS Auditory", "T3", "A1", "T3", "1.0mA", "10Hz", "tACS 20min — alpha normalization"],
        ["TN-PLATO-07", "Elderly/Hyperacusis", "T3", "A1", "T3", "1.0mA", "N/A", "15min cathodal"],
        ["TN-PLATO-08", "Sham", "F3", "DLPFC", "F3", "0mA", "N/A", "30s ramp"],
    ],
    tps_protocols=[
        ["TN-TPS-01", "A1 Hyperexcitability Reduction", "Primary Auditory Cortex (A1)", "200 pulses, 0.2Hz, 0.2mJ/mm²", "3 sessions/week", "Eggermont 2004"],
        ["TN-TPS-02", "DLPFC Top-Down Enhancement", "DLPFC (BA46)", "250 pulses, 0.2Hz", "3 sessions/week", "Schecklmann 2011"],
        ["TN-TPS-03", "ACC Distress Network", "ACC (BA24/32)", "200 pulses, 0.2Hz", "2 sessions/week", "Landgrebe 2009"],
        ["TN-TPS-04", "Parahippocampal (Memory/Distress)", "Parahippocampal Gyrus", "150 pulses, 0.2Hz", "2 sessions/week", "Lockwood 2002"],
        ["TN-TPS-05", "Full Network Protocol", "A1 + DLPFC + ACC sequential", "400 pulses total", "5 sessions/week x2 weeks", "Expert consensus 2024"],
    ],
    ces_role=[
        ["Sleep-Dominant Tinnitus", "Primary modality for tinnitus-insomnia cycle; reduces nighttime awareness", "Nightly 60min; in-ear or electrode"],
        ["Anxiety/Distress", "Reduces tinnitus-related anxiety and emotional reactivity", "Daily 60min; before high-stress periods"],
        ["Thalamic Dysrhythmia", "Indirect thalamic modulation via thalamocortical network", "Daily 60min; concurrent with sound therapy"],
        ["Maintenance", "Long-term distress and sleep maintenance", "3–5x/week; patient home use"],
    ],
    tavns_role="taVNS (tragus, 25Hz, 200µs, 0.5mA, 30min) for tinnitus leverages two mechanisms: (1) Auditory cortex plasticity facilitation — vagus nerve stimulation paired with tones drives targeted auditory cortex reorganization (Engineer 2011 demonstrated VNS-tone pairing reverses noise-induced tonotopic reorganization in animal models); (2) Distress network modulation via NTS → LC → amygdala pathway reduces tinnitus-related anxiety and emotional suffering. Clinical translation of VNS-tone pairing to taVNS is under active investigation. Standalone taVNS (without tone pairing) provides anxiolytic and autonomic benefits in tinnitus distress.",
    combinations=[
        ("Central/Network Dominant", [
            ["tDCS A1 cathodal + TPS A1 + taVNS", "Convergent A1 hyperexcitability reduction + vagal plasticity", "taVNS (tone pairing if available) → TPS → tDCS", "Central tinnitus without hearing loss"],
            ["tDCS A1 cathodal + tDCS DLPFC anodal + TPS DLPFC", "A1 inhibition + DLPFC top-down enhancement", "TPS DLPFC → tDCS DLPFC → tDCS A1", "Network dominant; high attention capture"],
        ]),
        ("Distress-Dominant", [
            ["tDCS DLPFC + TPS ACC + taVNS + CES", "Top-down suppression + distress network + vagal + anxiolytic", "taVNS → TPS → tDCS + CES nightly", "High THI, anxiety, depression comorbid"],
        ]),
        ("Sleep-Dominant", [
            ["CES + tDCS DLPFC + taVNS", "Sleep primary + top-down + vagal anxiety", "CES nightly; tDCS + taVNS daytime", "Tinnitus-insomnia cycle"],
        ]),
        ("Audiological/Peripheral", [
            ["tDCS A1 cathodal + TPS A1 + CES", "A1 inhibition + thalamic + sleep", "TPS → tDCS; CES nightly", "SNHL + tinnitus; hearing aid concurrent"],
        ]),
        ("Hyperacusis Comorbid", [
            ["tDCS A1 cathodal + CES", "A1 gain reduction + thalamic; TPS avoided in severe hyperacusis", "tDCS; CES nightly; sound desensitization", "Tinnitus + sound hypersensitivity"],
        ]),
        ("Refractory High THI", [
            ["tDCS A1 cathodal + tDCS DLPFC + TPS A1/DLPFC/ACC + taVNS + CES", "Full circuit: A1 + top-down + distress + vagal + sleep", "taVNS → TPS → tDCS (A1+DLPFC) + CES nightly", "THI ≥56, refractory to CBT and sound therapy"],
        ]),
    ],
    combination_summary=[
        ["Central/Network", "tDCS A1 cathodal + TPS A1 + taVNS", "A1 inhibition + vagal plasticity", "taVNS → TPS → tDCS A1", "Central tinnitus", "IIb"],
        ["Distress-Dominant", "tDCS DLPFC + TPS ACC + taVNS + CES", "Top-down + distress + vagal + sleep", "taVNS → TPS → tDCS + CES", "High distress/anxiety", "IIb"],
        ["Sleep-Dominant", "CES + tDCS DLPFC + taVNS", "Sleep + top-down + vagal", "CES nightly + daytime NIBS", "Tinnitus-insomnia", "IIb"],
        ["Audiological", "tDCS A1 cathodal + TPS A1 + CES", "A1 inhibition + thalamic", "TPS → tDCS + CES", "SNHL + tinnitus", "IIb"],
        ["Hyperacusis Comorbid", "tDCS A1 cathodal + CES", "A1 gain + thalamic; TPS avoided", "tDCS + CES nightly", "Hyperacusis + tinnitus", "IIb"],
        ["Refractory", "All modalities full protocol", "Full tinnitus network attack", "taVNS → TPS → tDCS × 2 + CES", "THI ≥56 refractory", "IIb"],
    ],
    outcomes=[
        ["Tinnitus Handicap Inventory (THI)", "Tinnitus-related handicap", "Baseline, 4wk, 8wk", "≥7-point reduction (MCID); ≥38-point reduction = grade improvement"],
        ["Tinnitus Functional Index (TFI)", "Multidimensional tinnitus impact", "Baseline, 4wk, 8wk", "≥13-point reduction (MCID)"],
        ["Visual Analog Scale — Loudness (VAS-L)", "Perceived tinnitus loudness", "Weekly", "≥20% reduction"],
        ["Visual Analog Scale — Distress (VAS-D)", "Emotional distress from tinnitus", "Weekly", "≥20% reduction"],
        ["Pittsburgh Sleep Quality Index (PSQI)", "Sleep quality", "Baseline, 4wk, 8wk", "≥3-point improvement"],
        ["GAD-7 / HADS-A", "Tinnitus-related anxiety", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["PHQ-9", "Depression comorbidity", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["Tinnitus Catastrophizing Scale (TCS)", "Catastrophizing about tinnitus", "Baseline, 4wk, 8wk", "≥30% reduction"],
        ["Minimum Masking Level (MML)", "Audiological tinnitus measure", "Baseline, 8wk", "≥5dB reduction"],
    ],
    outcomes_abbrev=[
        "THI: Tinnitus Handicap Inventory",
        "TFI: Tinnitus Functional Index",
        "VAS-L/D: Visual Analog Scale — Loudness/Distress",
        "PSQI: Pittsburgh Sleep Quality Index",
        "GAD-7: Generalized Anxiety Disorder 7",
        "HADS-A: Hospital Anxiety and Depression Scale — Anxiety",
        "PHQ-9: Patient Health Questionnaire-9",
        "TCS: Tinnitus Catastrophizing Scale",
        "MML: Minimum Masking Level",
        "SNHL: Sensorineural Hearing Loss",
    ],
)

# ---------------------------------------------------------------------------
# CONDITION: INSOMNIA
# ---------------------------------------------------------------------------
CONDITIONS["insomnia"] = dict(
    name="Chronic Insomnia Disorder",
    icd10="G47.00",
    doc_num="SPG-IN-001",
    tps_status="Investigational",
    tdcs_class="Class IIb",
    tps_class="Investigational",
    tavns_class="Class IIb",
    ces_class="Class IIa",
    inclusion=[
        "DSM-5 Insomnia Disorder: difficulty initiating/maintaining sleep ≥3 nights/week for ≥3 months",
        "ISI (Insomnia Severity Index) score ≥15 (moderate-severe) at baseline",
        "Adequate sleep opportunity (≥7 hours in bed) ruling out insufficient sleep syndrome",
        "Age 18–75; capable of informed consent",
        "Inadequate response to ≥1 CBT-I course or sleep hygiene optimization",
        "Stable medication regimen for ≥4 weeks (sleep medications, antidepressants)",
    ],
    exclusion=[
        "Untreated obstructive sleep apnea (AHI ≥15) as primary sleep disorder",
        "Metallic implants in head/neck or implanted neurostimulator",
        "Shift work sleep disorder as primary etiology",
        "Pregnancy",
        "Active substance use disorder (alcohol, hypnotics, stimulants)",
        "Active psychosis or bipolar disorder in manic phase",
        "Restless legs syndrome or periodic limb movement disorder as primary diagnosis",
    ],
    discussion=[
        "Chronic insomnia involves hyperarousal at cortical, cognitive, and autonomic levels — the 3P (predisposing-precipitating-perpetuating) model explains its persistence.",
        "Cortical hyperarousal: increased high-frequency (beta) EEG activity during sleep, particularly in frontal regions, reflects hyperactive salience processing inhibiting sleep onset.",
        "The thalamic reticular nucleus (TRN) normally gates sensory input to promote sleep; dysfunction leads to inappropriate sensory intrusion and sleep fragmentation.",
        "The DLPFC is implicated in rumination, worry, and cognitive hyperarousal that perpetuates insomnia through top-down thalamic excitation.",
        "CES (Alpha-Stim) is the primary NIBS modality for insomnia: direct thalamic and brainstem modulation promotes sleep-promoting delta/alpha oscillations.",
        "tDCS targeting DLPFC (cathodal) reduces cognitive hyperarousal and pre-sleep rumination.",
        "taVNS activates the parasympathetic nervous system, reduces autonomic hyperarousal (elevated sympathetic tone in insomnia), and promotes sleep via LC downregulation.",
        "TPS targeting frontal hyperarousal nodes (DLPFC, ACC) is investigational but mechanistically plausible.",
    ],
    overview=[
        "Chronic Insomnia Disorder affects approximately 10-15% of adults globally, representing the most prevalent sleep disorder. Characterized by persistent difficulty initiating or maintaining sleep, early morning awakening, and significant daytime impairment, insomnia involves bidirectional interactions with mood disorders, chronic pain, cardiovascular disease, and cognitive decline. The core neurobiological mechanism is hyperarousal — excessive activation of arousal systems (LC, hypothalamic arousal pathways) relative to sleep-promoting systems (VLPO, thalamic reticular nucleus).",
        "The SOZO NIBS protocol for insomnia is built around CES (Alpha-Stim) as the primary modality, supplemented by tDCS (DLPFC cathodal for cognitive hyperarousal), taVNS (autonomic arousal reduction), and TPS (investigational frontal hyperarousal targeting). All NIBS is integrated with CBT-I (Cognitive Behavioral Therapy for Insomnia) as the evidence-based psychotherapy backbone.",
        "Five clinical phenotypes guide protocol selection: sleep-onset insomnia (DLPFC/arousal system hyperactivation), sleep-maintenance insomnia (thalamic gating failure), comorbid insomnia-anxiety, comorbid insomnia-depression, and hyperarousal/cognitive dominant insomnia.",
    ],
    pathophysiology=[
        "Cortical hyperarousal: elevated frontal beta (16–32Hz) EEG power during NREM sleep reflects ongoing cognitive processing preventing consolidated sleep.",
        "Thalamic gating failure: thalamic reticular nucleus (TRN) inadequately filters sensory input; inappropriate sleep-to-wake transitions (micro-awakenings).",
        "Hypothalamic arousal system overdrive: lateral hypothalamus (orexin/hypocretin neurons) and locus coeruleus maintain excessive wakefulness drive.",
        "VLPO dysfunction: ventrolateral preoptic area sleep-promoting neurons are insufficient to flip the sleep-wake switch toward sleep.",
        "HPA axis dysregulation: elevated cortisol and CRF in insomnia perpetuate physiological arousal and disrupt sleep architecture.",
        "Autonomic hyperarousal: elevated sympathetic tone (low HRV, elevated heart rate) during sleep onset; taVNS and CES target this mechanism.",
    ],
    std_treatment=[
        "First-line: CBT-I (Cognitive Behavioral Therapy for Insomnia) — stimulus control, sleep restriction, relaxation, cognitive restructuring",
        "Pharmacological: Z-drugs (zolpidem, eszopiclone), doxepin, ramelteon, suvorexant/lemborexant, low-dose TCAs",
        "Sleep hygiene: consistent sleep schedule, light management, temperature, caffeine/alcohol restriction",
        "Relaxation techniques: progressive muscle relaxation, mindfulness, biofeedback",
        "Emerging: CBTI digital apps (digital therapeutics), low-dose lithium, chronotherapy",
    ],
    symptoms=[
        ["Sleep Onset Difficulty", "Takes >30min to fall asleep ≥3 nights/week", "100% (onset type)", "APA DSM-5 2013"],
        ["Sleep Maintenance Difficulty", "Waking during night with difficulty returning to sleep", "100% (maintenance type)", "APA DSM-5 2013"],
        ["Early Morning Awakening", "Waking ≥30min before intended rise time", "60%", "Morin 2006"],
        ["Non-Restorative Sleep", "Unrefreshing, poor quality sleep subjectively", "80%", "Buysse 2008"],
        ["Daytime Fatigue/Sleepiness", "Fatigue, tiredness despite adequate sleep opportunity", "90%", "Morin 2006"],
        ["Cognitive Impairment", "Concentration, memory, decision-making deficits", "75%", "Fortier-Brochu 2012"],
        ["Mood Disturbance", "Irritability, dysphoria, anxiety", "65%", "Riemann 2010"],
        ["Pre-sleep Cognitive Arousal", "Worry, rumination, racing thoughts at bedtime", "85%", "Harvey 2002"],
        ["Somatic Hyperarousal", "Muscle tension, elevated heart rate, body temperature at sleep onset", "70%", "Bonnet 2000"],
        ["Performance Anxiety (Conditioned Arousal)", "Anticipatory anxiety about sleep; bedroom conditioned stimulus", "75%", "Bootzin 1972"],
    ],
    brain_regions=[
        ["DLPFC", "Pre-sleep rumination, cognitive hyperarousal, worry generation", "tDCS cathodal, TPS", "Harvey 2002"],
        ["ACC", "Conflict monitoring during sleep onset; hyperarousal monitoring", "TPS cathodal", "Drummond 2004"],
        ["Thalamus (TRN)", "Sensory gating, sleep spindle generation; TRN dysfunction in insomnia", "CES primary", "Llinás 1999"],
        ["VLPO (Hypothalamus)", "Sleep-promoting nucleus; insufficient activation in insomnia", "CES indirect, taVNS", "Sherin 1996"],
        ["Locus Coeruleus", "Wakefulness promotion; excessive activation in insomnia", "taVNS, CES", "Aston-Jones 2005"],
        ["Default Mode Network", "Rumination, mind-wandering; DMN fails to deactivate at sleep onset", "tDCS DLPFC cathodal", "Nofzinger 2004"],
        ["Prefrontal Cortex (vmPFC)", "Emotional regulation, sleep anxiety inhibition", "TPS, tDCS", "Harvey 2002"],
        ["Hippocampus", "Fear memory of bed/sleep; conditioned arousal consolidation", "taVNS", "Bootzin 1972"],
    ],
    brainstem=[
        ["Locus Coeruleus (LC)", "Noradrenergic wakefulness driver; hyperactive in insomnia", "taVNS, CES", "Aston-Jones 2005"],
        ["Raphe Nuclei", "Serotonergic sleep-wake modulation; reduced in insomnia", "taVNS", "Pace-Schott 2002"],
        ["Parabrachial Nucleus", "Arousal relay; projects to BF and cortex", "taVNS indirect", "Fuller 2011"],
        ["Nucleus Tractus Solitarius", "Vagal integration; parasympathetic-sleep interface", "taVNS direct", "Porges 2007"],
        ["Tuberomammillary Nucleus", "Histaminergic wakefulness (H1 target of antihistamine hypnotics)", "None direct NIBS", "Saper 2005"],
    ],
    phenotypes=[
        ["Sleep-Onset Insomnia", "Difficulty falling asleep; presleep cognitive-somatic arousal; DLPFC hyperactivation", "Racing thoughts, muscle tension, conditioned arousal", "CES primary + tDCS DLPFC cathodal + taVNS"],
        ["Sleep-Maintenance Insomnia", "Multiple awakenings; thalamic gating failure; light sleep predominance", "Fragmented sleep, early morning waking", "CES primary + taVNS; tDCS if cognition impaired"],
        ["Comorbid Insomnia-Anxiety", "Insomnia driven by anxiety disorder; hyperarousal-anxiety interaction", "Pre-sleep worry, autonomic arousal, GAD features", "CES + taVNS + tDCS DLPFC cathodal; CBT-I + anxiety treatment"],
        ["Comorbid Insomnia-Depression", "Insomnia as prodrome/symptom of MDD; DLPFC-limbic dysfunction", "Early morning awakening, depressive rumination", "CES + tDCS DLPFC anodal (antidepressant) + taVNS"],
        ["Hyperarousal/Cognitive Dominant", "High cognitive pre-sleep arousal (PSAS-C ≥18); minimal somatic", "Intense rumination, racing thoughts, hypervigilance for sleep", "tDCS DLPFC cathodal + TPS + CES + CBT-I cognitive focus"],
    ],
    symptom_map=[
        ["Sleep onset difficulty", "DLPFC/Conditioned arousal", "tDCS DLPFC cathodal + CES", "Harvey 2002, Lande 2018"],
        ["Sleep maintenance", "Thalamus/TRN", "CES primary", "Llinas 1999"],
        ["Pre-sleep rumination", "DLPFC/DMN", "tDCS DLPFC cathodal + TPS", "Harvey 2002"],
        ["Autonomic hyperarousal", "LC/Sympathetic", "taVNS + CES", "Bonnet 2000"],
        ["Early morning awakening", "Thalamus/HPA", "CES + taVNS", "Morin 2006"],
        ["Anxiety comorbidity", "Amygdala/DLPFC", "taVNS + CES + tDCS DLPFC cathodal", "Riemann 2010"],
        ["Depression comorbidity", "DLPFC/Limbic", "tDCS DLPFC anodal + taVNS", "Brunoni 2013"],
        ["Conditioned arousal", "Hippocampus/Amygdala", "CBT-I stimulus control + taVNS", "Bootzin 1972"],
        ["Cognitive impairment (daytime)", "DLPFC/CEN", "tDCS DLPFC anodal (daytime)", "Fortier-Brochu 2012"],
    ],
    montage=[
        ["Sleep-Onset Insomnia", "Cathodal DLPFC (F3, 1.5mA, 20min — daytime)", "DLPFC (200 pulses)", "CES primary nightly 60min; taVNS before bed"],
        ["Sleep-Maintenance Insomnia", "Cathodal DLPFC (F3, 1.5mA, 20min — daytime)", "None initial", "CES primary nightly 60min; taVNS before bed"],
        ["Comorbid Insomnia-Anxiety", "Cathodal DLPFC (F3, 1.5mA, 20min) — daytime", "None initial", "CES primary nightly + daytime taVNS; anxiety treatment concurrent"],
        ["Comorbid Insomnia-Depression", "Anodal DLPFC (F3, 2mA, 20min) — daytime", "None initial", "CES nightly; taVNS 2x daily; antidepressant maintained"],
        ["Hyperarousal/Cognitive", "Cathodal DLPFC (F3, 1.5mA, 20min) + TPS daytime", "DLPFC/ACC (200 pulses)", "CES nightly; CBT-I cognitive component essential"],
        ["Maintenance Phase", "Cathodal DLPFC (F3, 1mA, 15min — daytime)", "None", "CES 3–5x/week; taVNS as needed home"],
        ["Elderly Adjusted", "Cathodal DLPFC (F3, 1mA, 15min)", "None", "CES primary nightly; taVNS standard; reduce zolpidem"],
        ["Refractory Insomnia", "Cathodal DLPFC (F3, 1.5mA) + TPS (300 pulses)", "DLPFC + ACC (300 pulses)", "All modalities + intensive CBT-I"],
        ["Comorbid Sleep Apnea (treated)", "Cathodal DLPFC (F3, 1.5mA, 20min)", "None", "CES nightly; CPAP confirmed adherent"],
        ["Sham Protocol", "F3 cathodal (30s ramp, 0mA, 20min)", "None", "Sham control condition"],
    ],
    tdcs_protocols=[
        ["IN-tDCS-01", "DLPFC Cathodal (Cognitive Hyperarousal Reduction)", "F4", "F3 cathodal", "1.5mA, 20min, 10 sessions (daytime)", "Frase 2016"],
        ["IN-tDCS-02", "DLPFC Anodal (Comorbid Depression)", "F3 anodal", "F4", "2mA, 20min, 10 sessions", "Brunoni 2013"],
        ["IN-tDCS-03", "Bifrontal (Hyperarousal with Anxiety)", "F3 anodal (antidepressant); F4 cathodal", "F4 anodal; F3 cathodal", "1.5mA, 20min, 10 sessions", "Frase 2016"],
        ["IN-tDCS-04", "Vertex/SMA Cathodal (Somatic Hyperarousal)", "Fpz", "Cz cathodal", "1.5mA, 20min, 10 sessions", "Expert consensus 2023"],
        ["IN-tDCS-05", "Maintenance", "F4 (cathodal F3)", "F3", "1mA, 15min, 2x/week", "Bikson 2016"],
        ["IN-tDCS-06", "Combined DLPFC Cathodal + CES", "F3 cathodal", "F4", "1.5mA, 20min + CES concurrent or sequential", "Expert consensus 2024"],
        ["IN-tDCS-07", "Elderly Adjusted", "F3 cathodal", "F4", "1mA, 15min, 10 sessions", "Lefaucheur 2016"],
        ["IN-tDCS-08", "Sham", "F3", "F4", "30s ramp, 0mA, 20min", "Frase 2016"],
    ],
    plato_protocols=[
        ["IN-PLATO-01", "DLPFC Cathodal", "F3/F4", "Left DLPFC", "F3", "1.5mA", "N/A", "Cathodal 20min daytime"],
        ["IN-PLATO-02", "DLPFC Anodal (Depression)", "F3/F4", "Left DLPFC anodal", "F3", "2.0mA", "N/A", "Anodal 20min"],
        ["IN-PLATO-03", "Vertex Cathodal", "Cz/Fpz", "Vertex/SMA", "Cz", "1.5mA", "N/A", "Cathodal 20min"],
        ["IN-PLATO-04", "HD-tDCS DLPFC Cathodal", "F3 center", "4-electrode ring", "F3", "0.75mA", "N/A", "HD cathodal 20min"],
        ["IN-PLATO-05", "Maintenance", "F3", "DLPFC", "F3", "1.0mA", "N/A", "2x/week 15min"],
        ["IN-PLATO-06", "Delta tACS Sleep-Promoting", "Cz", "Vertex", "Cz", "1.0mA", "1Hz", "Delta tACS 20min — investigational"],
        ["IN-PLATO-07", "Elderly", "F3", "DLPFC cathodal", "F3", "1.0mA", "N/A", "15min"],
        ["IN-PLATO-08", "Sham", "F3", "DLPFC", "F3", "0mA", "N/A", "30s ramp"],
    ],
    tps_protocols=[
        ["IN-TPS-01", "DLPFC Hyperarousal Reduction", "DLPFC (BA46)", "200 pulses, 0.2Hz, 0.15mJ/mm²", "2 sessions/week (daytime)", "Frase 2016"],
        ["IN-TPS-02", "ACC Arousal Monitoring", "ACC (BA24/32)", "150 pulses, 0.15Hz", "2 sessions/week", "Drummond 2004"],
        ["IN-TPS-03", "vmPFC Sleep Anxiety", "vmPFC (BA10/11)", "150 pulses, 0.15Hz", "2 sessions/week", "Harvey 2002"],
        ["IN-TPS-04", "Full Frontal Hyperarousal Protocol", "DLPFC + ACC sequential", "300 pulses total", "3 sessions/week x3 weeks", "Expert consensus 2024"],
        ["IN-TPS-05", "Hyperarousal/Cognitive Dominant", "DLPFC + vmPFC", "250 pulses total", "3 sessions/week", "Expert consensus 2024"],
    ],
    ces_role=[
        ["Primary Insomnia Modality", "First-line NIBS for insomnia; direct thalamic and brainstem sleep-promoting modulation; safest overall", "Nightly 60min; begin 30–60min before bed"],
        ["Sleep Onset Insomnia", "Reduces pre-sleep arousal; promotes alpha shift and sleep onset", "Nightly 60min; pre-bed application"],
        ["Sleep Maintenance Insomnia", "Improves sleep architecture, spindle density, and reduces awakenings", "Nightly 60min continuous application"],
        ["Maintenance Phase", "Long-term independent sleep management; patient self-administered", "3–5x/week indefinitely; home device program"],
    ],
    tavns_role="taVNS (tragus, 25Hz, 200µs, 0.5mA, 30min) is applied pre-sleep (30–60min before bedtime) to downregulate locus coeruleus noradrenergic hyperarousal and increase parasympathetic tone (HRV). The autonomic normalization effect promotes the physiological shift from sympathetic to parasympathetic dominance required for sleep onset. Home device (Nemos/TENS tragus) used nightly. In comorbid insomnia-anxiety, taVNS provides additional anxiolytic benefit. PEM monitoring is not required in insomnia (Vonck 2014, Lehtimäki 2013).",
    combinations=[
        ("Sleep-Onset Insomnia", [
            ["CES + tDCS DLPFC cathodal + taVNS", "Thalamic sleep-promoting + cognitive hyperarousal reduction + autonomic", "taVNS 30min pre-bed; CES nightly; tDCS daytime", "Racing thoughts, conditioned arousal at sleep onset"],
            ["CES + taVNS", "Core combination; thalamic + autonomic; minimal exertion", "taVNS 30min before CES nightly", "Onset insomnia without cognitive component"],
        ]),
        ("Sleep-Maintenance Insomnia", [
            ["CES + taVNS", "Thalamic gating + autonomic; minimal additional intervention", "CES nightly; taVNS pre-bed", "Fragmented sleep, multiple awakenings"],
            ["CES + tDCS DLPFC cathodal", "Thalamic + cognitive hyperarousal if cognition implicated", "tDCS daytime; CES nightly", "Maintenance insomnia with daytime rumination"],
        ]),
        ("Comorbid Insomnia-Anxiety", [
            ["CES + taVNS + tDCS DLPFC cathodal", "Thalamic + autonomic + prefrontal arousal reduction", "taVNS pre-bed; CES nightly; tDCS daytime", "High pre-sleep anxiety, autonomic arousal"],
        ]),
        ("Comorbid Insomnia-Depression", [
            ["CES + tDCS DLPFC anodal (daytime) + taVNS", "Sleep-promoting + antidepressant + autonomic", "tDCS daytime anodal; CES nightly; taVNS pre-bed", "MDD + insomnia; early morning awakening"],
        ]),
        ("Hyperarousal/Cognitive Dominant", [
            ["tDCS DLPFC cathodal + TPS DLPFC/ACC + CES + CBT-I", "Prefrontal inhibition + frontal TPS + sleep-promoting + psychological", "TPS → tDCS (daytime); CES nightly; CBT-I concurrent", "Intense rumination, hypervigilance for sleep"],
        ]),
        ("Refractory Insomnia", [
            ["tDCS DLPFC cathodal + TPS DLPFC/ACC + taVNS + CES", "Full circuit: frontal hyperarousal + autonomic + thalamic", "TPS → tDCS daytime; taVNS pre-bed; CES nightly", "Failed CBT-I + medication; ISI ≥22"],
        ]),
    ],
    combination_summary=[
        ["Sleep-Onset", "CES + tDCS DLPFC cathodal + taVNS", "Thalamic + prefrontal + autonomic", "taVNS pre-bed; CES nightly; tDCS daytime", "Onset insomnia", "IIa/IIb"],
        ["Sleep-Maintenance", "CES + taVNS", "Thalamic gating + autonomic", "taVNS pre-bed + CES nightly", "Maintenance insomnia", "IIa/IIb"],
        ["Insomnia-Anxiety", "CES + taVNS + tDCS DLPFC cathodal", "Sleep + autonomic + arousal", "taVNS pre-bed; CES nightly; tDCS day", "Insomnia + anxiety", "IIb"],
        ["Insomnia-Depression", "CES + tDCS DLPFC anodal + taVNS", "Sleep + antidepressant + autonomic", "tDCS day anodal; CES + taVNS night", "Insomnia + MDD", "IIb"],
        ["Hyperarousal-Cognitive", "tDCS DLPFC cathodal + TPS + CES + CBT-I", "Prefrontal + frontal TPS + sleep + psych", "TPS → tDCS day; CES + taVNS night", "Racing thoughts dominant", "IIb"],
        ["Refractory", "All four modalities + CBT-I", "Full frontal + autonomic + thalamic", "TPS → tDCS day; taVNS + CES night", "Failed CBT-I + medication", "IIb"],
    ],
    outcomes=[
        ["Insomnia Severity Index (ISI)", "Insomnia severity", "Baseline, 4wk, 8wk", "≥8-point reduction (MCID); ISI <8 = remission"],
        ["Actigraphy — Sleep Onset Latency (SOL)", "Objective sleep onset", "Baseline week, 4wk week, 8wk week", "SOL <30min average"],
        ["Actigraphy — Wake After Sleep Onset (WASO)", "Objective sleep maintenance", "Baseline, 4wk, 8wk", "WASO <30min average"],
        ["Actigraphy — Total Sleep Time (TST)", "Objective sleep duration", "Baseline, 4wk, 8wk", "TST ≥6.5 hours average"],
        ["Pittsburgh Sleep Quality Index (PSQI)", "Subjective sleep quality", "Baseline, 4wk, 8wk", "PSQI ≤5 (good sleeper threshold)"],
        ["Pre-Sleep Arousal Scale (PSAS)", "Cognitive and somatic pre-sleep arousal", "Baseline, 4wk, 8wk", "≥5-point reduction in cognitive subscale"],
        ["PHQ-9", "Depression comorbidity", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["GAD-7", "Anxiety comorbidity", "Baseline, 4wk, 8wk", "≥5-point reduction"],
        ["Daytime Functioning (PROMIS Sleep-Related Impairment)", "Daytime impact", "Baseline, 4wk, 8wk", "≥5-point T-score improvement"],
    ],
    outcomes_abbrev=[
        "ISI: Insomnia Severity Index",
        "SOL: Sleep Onset Latency",
        "WASO: Wake After Sleep Onset",
        "TST: Total Sleep Time",
        "PSQI: Pittsburgh Sleep Quality Index",
        "PSAS: Pre-Sleep Arousal Scale",
        "PHQ-9: Patient Health Questionnaire-9",
        "GAD-7: Generalized Anxiety Disorder 7",
        "PROMIS: Patient-Reported Outcomes Measurement Information System",
        "TRN: Thalamic Reticular Nucleus",
        "VLPO: Ventrolateral Preoptic Area",
        "CBT-I: Cognitive Behavioral Therapy for Insomnia",
    ],
)

# ---------------------------------------------------------------------------
# RUNNER
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    from pathlib import Path
    OUTPUT_ROOT = Path("outputs/documents")
    for slug, cond in CONDITIONS.items():
        out = OUTPUT_ROOT / slug / "fellow" / "Evidence_Based_Protocol_Fellow.docx"
        print(f"Building: {slug}")
        build_document(cond, out)
    print("\nAll 15 documents generated.")
print("Tinnitus + Insomnia + Runner appended OK")
