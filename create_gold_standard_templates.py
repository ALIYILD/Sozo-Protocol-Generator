#!/usr/bin/env python3
"""
Create Gold Standard DOCX templates for all 8 document types.
These templates define the section structure that the generator will fill.
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent / "src"))

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _add_para(doc, text, italic=False, bold=False):
    p = doc.add_paragraph(text)
    if italic or bold:
        for run in p.runs:
            run.italic = italic
            run.bold = bold
    return p


def create_evidence_based_protocol():
    """Full Evidence-Based Protocol template — the most comprehensive document."""
    doc = docx.Document()

    _add_heading(doc, "Document Control & Clinical Responsibility", 1)
    _add_para(doc, "Organization: SOZO Brain Center")
    _add_para(doc, "Condition: [CONDITION NAME]")
    _add_para(doc, "ICD-10: [ICD-10]")
    _add_para(doc, "Version: 1.0")
    _add_para(doc, "Tier: [TIER]")
    _add_para(doc, "GOVERNANCE RULE: This document is for authorized SOZO personnel only.")

    _add_heading(doc, "Inclusion & Exclusion Criteria", 1)
    _add_para(doc, "[CONDITION NAME] inclusion and exclusion criteria for neuromodulation treatment.")

    _add_heading(doc, "Clinical Overview", 1)
    _add_para(doc, "Clinical overview of [CONDITION NAME] including epidemiology, burden, and treatment landscape.")

    _add_heading(doc, "Pathophysiology", 1)
    _add_para(doc, "Detailed pathophysiology and neurobiological mechanisms underlying [CONDITION NAME].")

    _add_heading(doc, "Key Brain Structures & Neuroanatomy", 1)
    _add_para(doc, "Brain regions implicated in [CONDITION NAME].")

    _add_heading(doc, "Functional Network Involvement", 1)
    _add_para(doc, "FNON network dysfunction profiles for [CONDITION NAME].")

    _add_heading(doc, "Clinical Phenotypes", 1)
    _add_para(doc, "Phenotype classification guide for [CONDITION NAME].")

    _add_heading(doc, "Symptom-Network-Modality Mapping", 1)
    _add_para(doc, "Mapping of symptoms to dysfunctional networks and treatment modalities.")

    _add_heading(doc, "tDCS Protocols", 1)
    _add_para(doc, "Transcranial direct current stimulation protocols for [CONDITION NAME].")

    _add_heading(doc, "TPS Protocols", 1)
    _add_para(doc, "Transcranial pulse stimulation protocols. NOTE: TPS is OFF-LABEL for conditions other than Alzheimer's disease.")

    _add_heading(doc, "taVNS & CES Protocols", 1)
    _add_para(doc, "Transcutaneous auricular vagus nerve stimulation and cranial electrotherapy stimulation protocols.")

    _add_heading(doc, "Multimodal Combination Protocols", 1)
    _add_para(doc, "SOZO S-O-Z-O sequencing framework for multimodal treatment.")

    _add_heading(doc, "Safety, Side Effects & Monitoring", 1)
    _add_para(doc, "Contraindications, adverse events, and monitoring requirements.")

    _add_heading(doc, "Assessment Tools & Outcome Measures", 1)
    _add_para(doc, "Validated assessment scales for baseline and follow-up measurement.")

    _add_heading(doc, "Responder Criteria & Follow-Up", 1)
    _add_para(doc, "Response classification and non-responder pathway.")

    _add_heading(doc, "Evidence Gaps & Review Flags", 1)
    _add_para(doc, "Known evidence gaps requiring ongoing literature monitoring.")

    _add_heading(doc, "References", 1)
    _add_para(doc, "All citations with PubMed IDs (PMIDs).")

    return doc


def create_clinical_handbook():
    """Clinical Handbook template — 8-stage patient journey guide."""
    doc = docx.Document()

    _add_heading(doc, "Document Control & Clinical Responsibility", 1)
    _add_para(doc, "Organization: SOZO Brain Center")

    _add_heading(doc, "Introduction & Available Modalities", 1)
    _add_para(doc, "Clinical overview of [CONDITION NAME] treatment at SOZO Brain Center.")

    _add_heading(doc, "Functional Network Involvement", 1)
    _add_para(doc, "FNON framework application for [CONDITION NAME].")

    _add_heading(doc, "Clinical Phenotypes", 1)
    _add_para(doc, "Phenotype-based treatment selection.")

    _add_heading(doc, "Stimulation Protocols", 1)
    _add_para(doc, "All available protocols for [CONDITION NAME].")

    _add_heading(doc, "Safety & Contraindications", 1)
    _add_para(doc, "Safety considerations and contraindication screening.")

    _add_heading(doc, "Assessment Tools", 1)
    _add_para(doc, "Validated outcome measures.")

    _add_heading(doc, "Responder Criteria", 1)
    _add_para(doc, "Response evaluation and non-responder pathway.")

    _add_heading(doc, "Governance Rules", 1)
    _add_para(doc, "Clinical governance and documentation requirements.")

    _add_heading(doc, "References", 1)
    _add_para(doc, "All citations with PMIDs.")

    return doc


def create_clinical_exam():
    """Clinical Examination Checklist template."""
    doc = docx.Document()

    _add_heading(doc, "Document Control & Clinical Responsibility", 1)
    _add_para(doc, "Organization: SOZO Brain Center")

    _add_heading(doc, "Patient Information", 1)
    _add_para(doc, "Patient identification and session metadata.")

    _add_heading(doc, "Clinical Overview", 1)
    _add_para(doc, "Brief clinical overview of [CONDITION NAME].")

    _add_heading(doc, "Assessment Tools & Screening", 1)
    _add_para(doc, "Validated assessment scales for clinical examination.")

    _add_heading(doc, "Clinical Phenotypes", 1)
    _add_para(doc, "Phenotype identification checklist.")

    _add_heading(doc, "Functional Network Assessment", 1)
    _add_para(doc, "Network dysfunction screening.")

    return doc


def create_phenotype_classification():
    """Phenotype Classification template."""
    doc = docx.Document()

    _add_heading(doc, "Document Control & Clinical Responsibility", 1)
    _add_para(doc, "Organization: SOZO Brain Center")

    _add_heading(doc, "Clinical Phenotypes", 1)
    _add_para(doc, "Phenotype classification guide for [CONDITION NAME].")

    _add_heading(doc, "Functional Network Involvement", 1)
    _add_para(doc, "Network profiles per phenotype.")

    _add_heading(doc, "Stimulation Protocols", 1)
    _add_para(doc, "Phenotype-specific protocol selection.")

    _add_heading(doc, "References", 1)
    _add_para(doc, "All citations with PMIDs.")

    return doc


def create_responder_tracking():
    """Responder Tracking template."""
    doc = docx.Document()

    _add_heading(doc, "Document Control & Clinical Responsibility", 1)
    _add_para(doc, "Organization: SOZO Brain Center")

    _add_heading(doc, "Responder Criteria", 1)
    _add_para(doc, "SOZO response definitions for [CONDITION NAME].")

    _add_heading(doc, "Assessment Tools", 1)
    _add_para(doc, "Outcome measures for response tracking.")

    _add_heading(doc, "Evidence Gaps", 1)
    _add_para(doc, "Known evidence gaps in response classification.")

    _add_heading(doc, "References", 1)
    _add_para(doc, "All citations with PMIDs.")

    return doc


def create_psych_intake():
    """Psychological Intake & PRS Baseline template."""
    doc = docx.Document()

    _add_heading(doc, "Document Control & Clinical Responsibility", 1)
    _add_para(doc, "Organization: SOZO Brain Center")

    _add_heading(doc, "Patient Information", 1)
    _add_para(doc, "Patient identification and session metadata.")

    _add_heading(doc, "Clinical Overview", 1)
    _add_para(doc, "Brief overview of [CONDITION NAME] for intake context.")

    _add_heading(doc, "Assessment Tools & Screening", 1)
    _add_para(doc, "Baseline assessment scales.")

    _add_heading(doc, "Clinical Phenotypes", 1)
    _add_para(doc, "Phenotype identification at intake.")

    return doc


def create_network_assessment():
    """6-Network Bedside Assessment template (Partners only)."""
    doc = docx.Document()

    _add_heading(doc, "Document Control & Clinical Responsibility", 1)
    _add_para(doc, "Organization: SOZO Brain Center. PARTNERS TIER ONLY.")

    _add_heading(doc, "Patient Information", 1)
    _add_para(doc, "Patient identification and session metadata.")

    _add_heading(doc, "Functional Network Assessment", 1)
    _add_para(doc, "6-Network scoring for [CONDITION NAME].")

    _add_heading(doc, "Clinical Overview", 1)
    _add_para(doc, "Network involvement in [CONDITION NAME].")

    _add_heading(doc, "Clinical Phenotypes", 1)
    _add_para(doc, "Phenotype-network mapping.")

    return doc


def create_all_in_one_protocol():
    """All-in-One Protocol template."""
    doc = docx.Document()

    _add_heading(doc, "Document Control & Clinical Responsibility", 1)
    _add_para(doc, "Organization: SOZO Brain Center")

    _add_heading(doc, "Stimulation Protocols", 1)
    _add_para(doc, "All available protocols for [CONDITION NAME].")

    _add_heading(doc, "tDCS Protocols", 2)
    _add_para(doc, "tDCS-specific protocols.")

    _add_heading(doc, "TPS Protocols", 2)
    _add_para(doc, "TPS-specific protocols. OFF-LABEL for most conditions.")

    _add_heading(doc, "taVNS & CES Protocols", 2)
    _add_para(doc, "taVNS and CES protocols.")

    _add_heading(doc, "Safety & Contraindications", 1)
    _add_para(doc, "Safety considerations.")

    _add_heading(doc, "Inclusion & Exclusion Criteria", 1)
    _add_para(doc, "Patient eligibility criteria.")

    _add_heading(doc, "References", 1)
    _add_para(doc, "All citations with PMIDs.")

    return doc


TEMPLATES = {
    "Evidence_Based_Protocol": create_evidence_based_protocol,
    "Clinical_Handbook": create_clinical_handbook,
    "Clinical_Examination_Checklist": create_clinical_exam,
    "Phenotype_Classification": create_phenotype_classification,
    "Responder_Tracking": create_responder_tracking,
    "Psychological_Intake_PRS": create_psych_intake,
    "6Network_Bedside_Assessment": create_network_assessment,
    "All_In_One_Protocol": create_all_in_one_protocol,
}


def main():
    templates_dir = Path("templates/gold_standard/")
    templates_dir.mkdir(parents=True, exist_ok=True)

    print("Creating Gold Standard templates...")
    for name, builder_fn in TEMPLATES.items():
        doc = builder_fn()
        path = templates_dir / f"{name}.docx"
        doc.save(str(path))
        section_count = len([p for p in doc.paragraphs if p.style.name.startswith("Heading")])
        print(f"  {name:45s} | {section_count:2d} sections | {path}")

    print(f"\nCreated {len(TEMPLATES)} templates in {templates_dir}/")


if __name__ == "__main__":
    main()
