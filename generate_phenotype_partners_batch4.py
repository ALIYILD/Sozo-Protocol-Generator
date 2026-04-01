"""
Generate Partners Tier Phenotype Classification & Protocol Mapping DOCX files
Batch 4: Chronic Pain / Fibromyalgia + PTSD
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)


def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None
    size = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except:
        color = None
    for r in para.runs:
        r.text = ""
    if fr:
        fr.text = replaced
        fr.font.bold = bold
        if size:
            fr.font.size = size
        if color:
            fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type:
            run.font.color.rgb = C_BLACK


def _replace_table(table, rows):
    n_tmpl = len(table.rows)
    for ri in range(n_tmpl):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells:
                _cell_write(cell, "")
            continue
        is_hdr = (ri == 0)
        for ci, cell in enumerate(row.cells):
            text = str(rows[ri][ci]) if ci < len(rows[ri]) else ""
            _cell_write(cell, text, bold=is_hdr, white=is_hdr)


def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs):
            _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs):
                        _para_replace(para, old, new)


def build_phenotype_classification(c):
    TEMPLATE = Path(r"C:/Users/yildi/OneDrive/Desktop/Parkinson D/Partners/Assessments/PD_Phenotype_Classification_Partners.docx")
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs
    tables = doc.tables

    for old, new in [
        ("Parkinson's Disease (PD)", c["full"]),
        ("Parkinson's Disease", c["full"]),
        ("Parkinson's", c["short"]),
        (" (PD)", f" ({c['abbr']})"),
        (" PD ", f" {c['abbr']} "),
        (" PD.", f" {c['abbr']}."),
        (" PD,", f" {c['abbr']},"),
        ("for PD", f"for {c['abbr']}"),
        ("H&Y Stage", c["severity_field"]),
        ("Disease Duration", c["duration_field"]),
        ("PD Phenotype", f"{c['abbr']} Phenotype"),
        ("Hoehn & Yahr", "disease severity rating"),
        ("levodopa", c.get("med_name", "primary medication")),
    ]:
        _global_replace(doc, old, new)

    _para_set(paras[1], f"{c['short']} Phenotype Classification & Protocol Mapping")
    _para_set(paras[7], f"STEP 1 \u2014 {c['short']} Phenotype Determination")
    _para_set(paras[11], f"1.2 {c['short']} Phenotype Classification")
    _para_set(paras[14], f"ASSIGNED PHENOTYPE: {c['phenotype_codes']}")
    _para_set(paras[19], c["group1_label"])
    _para_set(paras[21], c["group2_label"])
    _para_set(paras[23], c["group3_label"])
    _para_set(paras[27], c["tps_offlabel"])
    _para_set(paras[40], c["selected_option"])

    _replace_table(tables[0], [
        ["Field", "Details"],
        ["Patient Name", ""],
        ["File Number", ""],
        [c["severity_field"], ""],
        ["Date", ""],
        ["Treating Doctor", ""],
        [c["duration_field"], "___ years / ___ months"],
    ])

    for ti, key in [
        (1, "data_collection_table"),
        (2, "phenotype_table"),
        (3, "network_table"),
        (5, "tdcs_group1"),
        (6, "tdcs_group2"),
        (7, "tdcs_group3"),
        (9, "tps_table"),
        (12, "tavns_table"),
        (13, "ces_table"),
        (14, "sequencing_table"),
    ]:
        _replace_table(tables[ti], c[key])

    _cell_write(tables[15].rows[1].cells[0], c["response_domain"])
    _cell_write(tables[15].rows[1].cells[1], "\u25a1 Completed | Score: ___ / baseline: ___")
    _cell_write(tables[18].rows[1].cells[0], f"{c['abbr']} Phenotype")

    slug = c["slug"]
    out = Path("outputs/documents") / slug / "partners" / f"Phenotype_Classification_Partners_{slug}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  [OK] {out}")
    return out


# ── Condition definitions ──────────────────────────────────────────────────────

CHRONIC_PAIN = {
    "slug": "chronic_pain",
    "short": "Chronic Pain",
    "abbr": "CP",
    "full": "Chronic Pain / Fibromyalgia (CP)",
    "severity_field": "NRS / BPI Pain Severity",
    "duration_field": "Pain Duration",
    "med_name": "analgesic/pain medication",
    "phenotype_codes": "\u25a1 NO  \u25a1 NE  \u25a1 NP  \u25a1 FM  \u25a1 CR  \u25a1 MS  \u25a1 HA",
    "group1_label": "3.1 Nociceptive / MSK Phenotypes (NO / MS)",
    "group2_label": "3.2 Neuropathic / CRPS (NE / CR)",
    "group3_label": "3.3 Nociplastic / Fibromyalgia / Headache (NP / FM / HA)",
    "tps_offlabel": "TPS is OFF-LABEL for Chronic Pain. Requires Doctor authorisation and off-label consent.",
    "selected_option": "SELECTED OPTION: \u25a1 A (Nociceptive)  \u25a1 B (Neuropathic)  \u25a1 C (Fibromyalgia)  \u25a1 D (Mixed)",
    "response_domain": "Reassess SOZO PRS + NRS / BPI / FIQ",
    "data_collection_table": [
        ["Data Point", "Completed", "Source"],
        ["Pain intensity NRS / BPI severity", "\u25a1", "Rating Scales"],
        ["Pain type profiling (nociceptive vs neuropathic vs nociplastic)", "\u25a1", "DN4 / PainDETECT"],
        ["Fibromyalgia criteria (ACR 2010: WPI + SS)", "\u25a1", "ACR Criteria"],
        ["CRPS diagnostic criteria (Budapest)", "\u25a1", "Clinical Assessment"],
        ["Comorbid depression and anxiety (PHQ-9, GAD-7)", "\u25a1", "Rating Scales"],
        ["Sleep quality (ISI, PSQI)", "\u25a1", "ISI / PSQI"],
        ["Functional disability (PDI, PROMIS)", "\u25a1", "Rating Scales"],
        ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"],
    ],
    "phenotype_table": [
        ["Dominant Features", "Assign Phenotype", "Code"],
        ["Peripheral nociceptive source, inflammatory, tissue injury dominant", "Nociceptive (NO)", "NO"],
        ["Burning, shooting, allodynia, nerve distribution pain", "Neuropathic (NE)", "NE"],
        ["Widespread pain, central sensitization, no clear peripheral source", "Nociplastic (NP)", "NP"],
        ["ACR fibromyalgia criteria met, fatigue, sleep, cognitive fog", "Fibromyalgia (FM)", "FM"],
        ["CRPS: disproportionate pain, autonomic-trophic changes", "CRPS (CR)", "CR"],
        ["Musculoskeletal pain: joint, muscle, spine dominant", "MSK (MS)", "MS"],
        ["Chronic headache / migraine dominant pain presentation", "Headache (HA)", "HA"],
    ],
    "network_table": [
        ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
        ["NO", "SMN (contralateral M1)", "SN (Insula)", "CEN"],
        ["NE", "SN (Insula/dACC)", "SMN (contralateral M1)", "Limbic"],
        ["NP", "SN (Insula/dACC)", "DMN", "CEN (L-DLPFC)"],
        ["FM", "SN (Insula)", "DMN", "CEN (L-DLPFC)"],
        ["CR", "SN (Insula)", "SMN", "Limbic (amygdala)"],
        ["MS", "SMN (contralateral M1)", "SN", "CEN"],
        ["HA", "SN (dACC/Insula)", "Thalamic", "SMN (trigeminal)"],
    ],
    "tdcs_group1": [
        ["Goal", "Anode", "Cathode"],
        ["Contralateral M1 analgesic (NO/MS)", "Contralateral M1", "Ipsilateral supraorbital"],
        ["L-DLPFC descending inhibition", "Left DLPFC (F3)", "Right supraorbital"],
        ["Bilateral M1 widespread pain (FM/NP)", "Bilateral M1", "Extracephalic"],
    ],
    "tdcs_group2": [
        ["Goal", "Anode", "Cathode"],
        ["SN/Insula downregulation (NE/CR)", "Left DLPFC", "Right DLPFC cathodal"],
        ["M1 + DLPFC combined (NE)", "Left M1 / DLPFC", "Right supraorbital"],
    ],
    "tdcs_group3": [
        ["Goal", "Anode", "Cathode"],
        ["Thalamic modulation for headache (HA)", "Left DLPFC", "Right DLPFC cathodal"],
        ["Fibromyalgia central sensitization (FM/NP)", "Left M1", "Left DLPFC (bi-frontal)"],
    ],
    "tps_table": [
        ["Clinical Target", "TPS Protocol", "Phenotype"],
        ["Contralateral M1 analgesic effect", "FT1", "NO, MS, NE"],
        ["Bilateral M1 central pain modulation", "FT3", "FM, NP"],
        ["Insula/SN central sensitization", "FT4", "NP, FM, CR"],
        ["L-DLPFC descending pain inhibition", "FT5", "All phenotypes"],
        ["Thalamic relay modulation", "FT7", "HA, NE"],
        ["CRPS autonomic-motor circuit", "FT8", "CR"],
        ["Bilateral pain neuromatrix", "FT9", "FM / severe"],
        ["dACC pain affect circuit", "Multiple", "NP, HA"],
        ["Multi-network distributed pain", "Multiple", "MX / refractory"],
    ],
    "tavns_table": [
        ["Domain", "Details"],
        ["Indications", "Fibromyalgia, autonomic dysregulation, sleep disruption, CRPS"],
        ["Protocol", "20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
        ["Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "ces_table": [
        ["Domain", "Details"],
        ["Indications", "Anxiety-pain comorbidity, insomnia, fibromyalgia, all phenotypes"],
        ["Protocol", "20\u201360 min daily, Alpha-Stim AID, especially FM, NP, HA phenotypes"],
        ["Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "sequencing_table": [
        ["Option", "Sequence", "Phenotype"],
        ["OPTION A \u2014 Nociceptive/MSK", "Wk 1\u20134: Contralateral M1 tDCS (5\u00d7/wk) \u2192 Wk 2\u20136: TPS M1/DLPFC", "NO / MS"],
        ["OPTION B \u2014 Neuropathic", "Wk 1\u20134: M1+DLPFC tDCS + Daily taVNS \u2192 Wk 3\u20136: TPS SN/M1", "NE / CR"],
        ["OPTION C \u2014 Fibromyalgia", "Wk 1\u20132: CES stabilization \u2192 Wk 2\u20136: Bilateral M1 tDCS + TPS SN", "FM / NP"],
        ["OPTION D \u2014 Mixed/Headache", "Sequential: TPS (highest burden) \u2192 Maintenance DLPFC tDCS + CES", "HA / MX"],
    ],
}

PTSD = {
    "slug": "ptsd",
    "short": "PTSD",
    "abbr": "PTSD",
    "full": "Post-Traumatic Stress Disorder (PTSD)",
    "severity_field": "PCL-5 / CAPS Severity",
    "duration_field": "PTSD Duration",
    "med_name": "SSRI/prazosin",
    "phenotype_codes": "\u25a1 RE  \u25a1 AV  \u25a1 HY  \u25a1 DI  \u25a1 CX  \u25a1 TB  \u25a1 DE",
    "group1_label": "3.1 Hyperarousal / Re-experiencing (HY / RE)",
    "group2_label": "3.2 Avoidance / Dissociative (AV / DI)",
    "group3_label": "3.3 Complex / TBI-Comorbid / Depression (CX / TB / DE)",
    "tps_offlabel": "TPS is OFF-LABEL for PTSD. Requires Doctor authorisation and off-label consent.",
    "selected_option": "SELECTED OPTION: \u25a1 A (Hyperarousal)  \u25a1 B (Avoidance)  \u25a1 C (Complex)  \u25a1 D (Mixed)",
    "response_domain": "Reassess SOZO PRS + PCL-5 / CAPS",
    "data_collection_table": [
        ["Data Point", "Completed", "Source"],
        ["PCL-5 / CAPS-5 severity score", "\u25a1", "Rating Scales"],
        ["Hyperarousal cluster symptoms (sleep, startle, irritability)", "\u25a1", "Clinical Interview"],
        ["Re-experiencing symptoms (flashbacks, nightmares)", "\u25a1", "Clinical Interview"],
        ["Avoidance and numbing severity", "\u25a1", "PCL-5 / Interview"],
        ["Dissociative features (IES-R, MID)", "\u25a1", "Dissociation Scales"],
        ["Complex trauma history and childhood adversity", "\u25a1", "Clinical History"],
        ["TBI comorbidity screening", "\u25a1", "Medical History"],
        ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"],
    ],
    "phenotype_table": [
        ["Dominant Features", "Assign Phenotype", "Code"],
        ["Intrusive memories, flashbacks, nightmares dominant", "Re-experiencing (RE)", "RE"],
        ["Persistent avoidance of trauma reminders, emotional numbing", "Avoidance (AV)", "AV"],
        ["Hypervigilance, exaggerated startle, sleep disruption dominant", "Hyperarousal (HY)", "HY"],
        ["Depersonalization, derealization, dissociative episodes", "Dissociative (DI)", "DI"],
        ["Complex/developmental trauma, chronic course, affect dysregulation", "Complex (CX)", "CX"],
        ["PTSD with concurrent traumatic brain injury", "TBI-Comorbid (TB)", "TB"],
        ["PTSD with major depressive episode dominant", "Depression-Comorbid (DE)", "DE"],
    ],
    "network_table": [
        ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
        ["RE", "Limbic (amygdala/hippocampus)", "SN (dACC)", "DMN"],
        ["AV", "Limbic (BLA-vmPFC)", "CEN (R-DLPFC)", "DMN"],
        ["HY", "SN (Insula/dACC)", "Limbic (amygdala)", "CEN"],
        ["DI", "DMN (mPFC)", "SN", "Limbic"],
        ["CX", "Limbic + SN", "CEN bilateral", "DMN"],
        ["TB", "CEN (bilateral DLPFC)", "Limbic", "SN"],
        ["DE", "CEN (L-DLPFC)", "DMN (vmPFC)", "Limbic"],
    ],
    "tdcs_group1": [
        ["Goal", "Anode", "Cathode"],
        ["R-DLPFC inhibition of fear network (HY/RE)", "Left DLPFC", "Right DLPFC cathodal"],
        ["vmPFC extinction facilitation (AV)", "vmPFC (Fpz)", "Right mastoid"],
        ["Bilateral DLPFC fear regulation (CX)", "Left DLPFC", "Right supraorbital"],
    ],
    "tdcs_group2": [
        ["Goal", "Anode", "Cathode"],
        ["vmPFC-hippocampal extinction (AV/DI)", "vmPFC (Fpz)", "Right shoulder"],
        ["Reduce amygdala hyperreactivity via DLPFC", "Left DLPFC", "Right DLPFC cathodal"],
    ],
    "tdcs_group3": [
        ["Goal", "Anode", "Cathode"],
        ["Bilateral DLPFC TBI-PTSD (TB)", "Bilateral DLPFC", "Extracephalic"],
        ["L-DLPFC depression comorbidity (DE)", "Left DLPFC (F3)", "Right supraorbital"],
    ],
    "tps_table": [
        ["Clinical Target", "TPS Protocol", "Phenotype"],
        ["R-DLPFC fear inhibition circuit", "FT1", "HY, RE"],
        ["vmPFC extinction facilitation", "FT3", "AV, DI"],
        ["Limbic-amygdala fear network", "FT4", "RE, HY"],
        ["SN/dACC hyperarousal reduction", "FT5", "HY, CX"],
        ["Hippocampal memory reconsolidation", "FT7", "RE, AV"],
        ["Bilateral DLPFC (TBI comorbid)", "FT8", "TB"],
        ["L-DLPFC antidepressant (DE)", "FT9", "DE"],
        ["vmPFC + hippocampal combined", "Multiple", "AV / DI"],
        ["Multi-network complex PTSD", "Multiple", "CX / severe"],
    ],
    "tavns_table": [
        ["Domain", "Details"],
        ["Indications", "Hyperarousal, sleep disruption, autonomic dysregulation, all PTSD phenotypes"],
        ["Protocol", "20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
        ["Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "ces_table": [
        ["Domain", "Details"],
        ["Indications", "Anxiety, insomnia, hyperarousal \u2014 first-line adjunct for all PTSD phenotypes"],
        ["Protocol", "20\u201360 min daily, Alpha-Stim AID, especially HY, RE, IN comorbidity"],
        ["Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "sequencing_table": [
        ["Option", "Sequence", "Phenotype"],
        ["OPTION A \u2014 Hyperarousal", "Wk 1\u20132: CES/taVNS stabilization \u2192 Wk 2\u20136: R-DLPFC cathodal tDCS + TPS SN", "HY / RE"],
        ["OPTION B \u2014 Avoidance", "Wk 1\u20134: vmPFC tDCS during trauma processing \u2192 Wk 3\u20136: TPS vmPFC/hippocampus", "AV / DI"],
        ["OPTION C \u2014 Complex", "Wk 1\u20134: Bilateral DLPFC tDCS + Daily taVNS + CES \u2192 Wk 4\u20136: TPS Limbic", "CX / DE"],
        ["OPTION D \u2014 TBI-Comorbid", "Sequential: TPS bilateral DLPFC \u2192 Maintenance L-DLPFC tDCS + taVNS", "TB / MX"],
    ],
}


# ── Main ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import os
    os.chdir(r"C:/Users/yildi/Sozo-Protocol-Generator")

    print("Building Phenotype Classification Partners — Batch 4")
    print("=" * 60)

    results = []
    for condition in [CHRONIC_PAIN, PTSD]:
        print(f"\nGenerating: {condition['full']}")
        out = build_phenotype_classification(condition)
        results.append(out)

    print("\n" + "=" * 60)
    print("Verification:")
    all_ok = True
    for path in results:
        exists = Path(path).exists()
        size = Path(path).stat().st_size if exists else 0
        status = "PASS" if exists and size > 0 else "FAIL"
        print(f"  [{status}] {path}  ({size:,} bytes)")
        if status != "PASS":
            all_ok = False

    print("\n" + ("All files generated successfully." if all_ok else "ERROR: Some files missing!"))
