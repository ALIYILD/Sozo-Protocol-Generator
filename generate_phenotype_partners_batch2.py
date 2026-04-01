"""
Generate Partners Tier Phenotype Classification & Protocol Mapping DOCX files
for ADHD and Alzheimer's Disease / MCI.
"""

from pathlib import Path
from docx import Document
from docx.shared import RGBColor

TEMPLATE = Path("C:/Users/yildi/OneDrive/Desktop/Parkinson D/Partners/Assessments/PD_Phenotype_Classification_Partners.docx")

C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)


def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None
    size = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except:
        color = None
    for r in para.runs:
        r.text = ""
    if fr:
        fr.text = replaced
        fr.font.bold = bold
        if size:
            fr.font.size = size
        if color:
            fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type:
            run.font.color.rgb = C_BLACK


def _replace_table(table, rows):
    n_tmpl = len(table.rows)
    for ri in range(n_tmpl):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells:
                _cell_write(cell, "")
            continue
        is_hdr = (ri == 0)
        for ci, cell in enumerate(row.cells):
            text = str(rows[ri][ci]) if ci < len(rows[ri]) else ""
            _cell_write(cell, text, bold=is_hdr, white=is_hdr)


def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs):
            _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs):
                        _para_replace(para, old, new)


def build_phenotype_classification(c):
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs
    tables = doc.tables

    # Global replacements
    for old, new in [
        ("Parkinson's Disease (PD)", c["full"]),
        ("Parkinson's Disease", c["full"]),
        ("Parkinson's", c["short"]),
        (" (PD)", f" ({c['abbr']})"),
        (" PD ", f" {c['abbr']} "),
        (" PD.", f" {c['abbr']}."),
        (" PD,", f" {c['abbr']},"),
        ("for PD", f"for {c['abbr']}"),
        ("H&Y Stage", c["severity_field"]),
        ("Disease Duration", c["duration_field"]),
        ("PD Phenotype", f"{c['abbr']} Phenotype"),
        ("Hoehn & Yahr", "disease severity rating"),
        ("levodopa", c.get("med_name", "primary medication")),
    ]:
        _global_replace(doc, old, new)

    # Paragraph replacements
    _para_set(paras[1], f"{c['short']} Phenotype Classification & Protocol Mapping")
    _para_set(paras[7], f"STEP 1 \u2014 {c['short']} Phenotype Determination")
    _para_set(paras[11], f"1.2 {c['short']} Phenotype Classification")
    _para_set(paras[14], f"ASSIGNED PHENOTYPE: {c['phenotype_codes']}")
    _para_set(paras[19], c["group1_label"])
    _para_set(paras[21], c["group2_label"])
    _para_set(paras[23], c["group3_label"])
    _para_set(paras[27], c["tps_offlabel"])
    _para_set(paras[40], c["selected_option"])

    # Table[0]: patient info
    _replace_table(tables[0], [
        ["Field", "Details"],
        ["Patient Name", ""],
        ["File Number", ""],
        [c["severity_field"], ""],
        ["Date", ""],
        ["Treating Doctor", ""],
        [c["duration_field"], "___ years / ___ months"],
    ])

    # Key tables
    _replace_table(tables[1], c["data_collection_table"])
    _replace_table(tables[2], c["phenotype_table"])
    _replace_table(tables[3], c["network_table"])
    _replace_table(tables[5], c["tdcs_group1"])
    _replace_table(tables[6], c["tdcs_group2"])
    _replace_table(tables[7], c["tdcs_group3"])
    _replace_table(tables[9], c["tps_table"])
    _replace_table(tables[12], c["tavns_table"])
    _replace_table(tables[13], c["ces_table"])
    _replace_table(tables[14], c["sequencing_table"])

    # Table[15] row[1] response evaluation
    _cell_write(tables[15].rows[1].cells[0], c["response_domain"])
    _cell_write(tables[15].rows[1].cells[1], "\u25a1 Completed | Score: ___ / baseline: ___")

    # Table[18] row[1][0] phenotype label
    _cell_write(tables[18].rows[1].cells[0], f"{c['abbr']} Phenotype")

    slug = c["slug"]
    out = Path("outputs/documents") / slug / "partners" / f"Phenotype_Classification_Partners_{slug}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  [OK] {out}")
    return out


# ── Condition definitions ──────────────────────────────────────────────────────

CONDITIONS = [
    {
        "slug": "adhd",
        "short": "ADHD",
        "abbr": "ADHD",
        "full": "Attention Deficit Hyperactivity Disorder (ADHD)",
        "severity_field": "ADHD Rating Scale / Conners Severity",
        "duration_field": "Symptom Duration",
        "med_name": "stimulant medication",
        "phenotype_codes": "\u25a1 IA  \u25a1 HI  \u25a1 CO  \u25a1 EX  \u25a1 EM  \u25a1 SC  \u25a1 WM",
        "group1_label": "3.1 Inattentive / Executive Phenotypes (IA / EX / WM)",
        "group2_label": "3.2 Hyperactive / Emotional Phenotypes (HI / EM)",
        "group3_label": "3.3 Combined / SCT Phenotypes (CO / SC)",
        "tps_offlabel": "TPS is OFF-LABEL for ADHD. Requires Doctor authorisation and off-label consent.",
        "selected_option": "SELECTED OPTION: \u25a1 A (Inattentive)  \u25a1 B (Hyperactive)  \u25a1 C (Combined)  \u25a1 D (Mixed)",
        "response_domain": "Reassess SOZO PRS + ADHD-RS / Conners",
        "data_collection_table": [
            ["Data Point", "Completed", "Source"],
            ["ADHD Rating Scale (ADHD-RS-5) severity", "\u25a1", "Rating Scales"],
            ["Inattentive vs hyperactive-impulsive feature ratio", "\u25a1", "Clinical Interview"],
            ["Executive function profile (BRIEF-A, cognitive tasks)", "\u25a1", "BRIEF-A / Interview"],
            ["Emotional dysregulation screening (DEARS, CAARS)", "\u25a1", "CAARS / Interview"],
            ["Working memory assessment (digit span, WAIS)", "\u25a1", "Neuropsychological"],
            ["Sleep quality and ADHD-sleep interaction", "\u25a1", "ISI / Interview"],
            ["Stimulant medication timing and response", "\u25a1", "Medical History"],
            ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"],
        ],
        "phenotype_table": [
            ["Dominant Features", "Assign Phenotype", "Code"],
            ["Inattention, distractibility, forgetfulness, slow processing dominant", "Inattentive (IA)", "IA"],
            ["Hyperactivity, impulsivity, motor restlessness dominant", "Hyperactive (HI)", "HI"],
            ["Both inattention and hyperactivity-impulsivity equally prominent", "Combined (CO)", "CO"],
            ["Executive dysfunction: planning, initiation, inhibition foregrounded", "Executive (EX)", "EX"],
            ["Emotional dysregulation, frustration intolerance, mood lability", "Emotional (EM)", "EM"],
            ["Sluggish cognitive tempo: slow, daydreamy, mentally foggy", "SCT (SC)", "SC"],
            ["Prominent working memory deficit as primary presentation", "Working Memory (WM)", "WM"],
        ],
        "network_table": [
            ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
            ["IA", "CEN (R-DLPFC)", "DAN (parietal)", "DMN (suppression)"],
            ["HI", "CEN (R-IFG/DLPFC)", "SMN (SMA)", "SN (Insula)"],
            ["CO", "CEN bilateral", "SN", "SMN"],
            ["EX", "CEN (R-DLPFC)", "Fronto-striatal", "DAN"],
            ["EM", "SN (dACC)", "Limbic (amygdala)", "CEN"],
            ["SC", "CEN (L-DLPFC)", "DMN", "DAN"],
            ["WM", "CEN bilateral (DLPFC)", "DAN", "Fronto-parietal"],
        ],
        "tdcs_group1": [
            ["Goal", "Anode", "Cathode"],
            ["Enhance right DLPFC excitability (IA/EX)", "Right DLPFC (F4)", "Left supraorbital"],
            ["Bilateral DLPFC working memory (WM/CO)", "Right DLPFC", "Left DLPFC (bi-cephalic)"],
            ["Strengthen right IFG inhibitory control", "Right IFG (F8)", "Left supraorbital"],
        ],
        "tdcs_group2": [
            ["Goal", "Anode", "Cathode"],
            ["Reduce SMA hyperexcitability (HI)", "Right DLPFC", "SMA cathodal"],
            ["Limbic-SN emotional regulation (EM)", "Left DLPFC", "Right DLPFC cathodal"],
        ],
        "tdcs_group3": [
            ["Goal", "Anode", "Cathode"],
            ["Bilateral frontal executive boost (CO/SC)", "Bilateral DLPFC", "Extracephalic"],
            ["DMN suppression enhancement (IA/SC)", "Right DLPFC", "Pz (posterior)"],
        ],
        "tps_table": [
            ["Clinical Target", "TPS Protocol", "Phenotype"],
            ["R-DLPFC excitability (attention network)", "FT1", "IA, EX"],
            ["Bilateral CEN normalization", "FT3", "CO, WM"],
            ["Fronto-striatal inhibitory circuit", "FT4", "HI, CO"],
            ["Right IFG response inhibition", "FT5", "HI, EM"],
            ["SN/dACC emotional regulation", "FT7", "EM"],
            ["Parietal DAN attention modulation", "FT8", "IA, WM"],
            ["SMA hyperactivity suppression", "FT9", "HI"],
            ["Bilateral prefrontal executive", "Multiple", "EX, SC"],
            ["Multi-network distributed ADHD", "Multiple", "CO / severe"],
        ],
        "tavns_table": [
            ["Domain", "Details"],
            ["Indications", "Emotional dysregulation, sleep disruption, autonomic hyperarousal in ADHD"],
            ["Protocol", "20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
            ["Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
        ],
        "ces_table": [
            ["Domain", "Details"],
            ["Indications", "Anxiety comorbidity, insomnia, emotional lability in ADHD"],
            ["Protocol", "20\u201360 min daily, Alpha-Stim AID, adjunct to EM and HI phenotypes"],
            ["Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
        ],
        "sequencing_table": [
            ["Option", "Sequence", "Phenotype"],
            ["OPTION A \u2014 Inattentive", "Wk 1\u20134: R-DLPFC tDCS (5\u00d7/wk) \u2192 Wk 2\u20136: TPS R-DLPFC/DAN", "IA / EX"],
            ["OPTION B \u2014 Hyperactive", "Wk 1\u20134: R-IFG tDCS + Daily CES \u2192 Wk 3\u20136: TPS IFG/SMA", "HI / EM"],
            ["OPTION C \u2014 Combined", "Wk 1\u20132: taVNS/CES stabilization \u2192 Wk 2\u20136: Bilateral DLPFC tDCS + TPS", "CO / WM"],
            ["OPTION D \u2014 SCT/Mixed", "Sequential: TPS (highest burden) \u2192 Maintenance bilateral DLPFC tDCS", "SC / MX"],
        ],
    },
    {
        "slug": "alzheimers",
        "short": "Alzheimer's",
        "abbr": "AD",
        "full": "Alzheimer's Disease / MCI (AD)",
        "severity_field": "CDR Stage / MoCA Score",
        "duration_field": "Cognitive Decline Duration",
        "med_name": "cholinesterase inhibitor",
        "phenotype_codes": "\u25a1 AM  \u25a1 MD  \u25a1 EO  \u25a1 LO  \u25a1 PC  \u25a1 LG  \u25a1 AP",
        "group1_label": "3.1 Amnestic / Multi-Domain (AM / MD / LO)",
        "group2_label": "3.2 Language / Posterior Cortical (LG / PC)",
        "group3_label": "3.3 Early-Onset / Apathy Dominant (EO / AP)",
        "tps_offlabel": "TPS is OFF-LABEL for AD/MCI. Requires Doctor authorisation and off-label consent.",
        "selected_option": "SELECTED OPTION: \u25a1 A (Amnestic)  \u25a1 B (Language/PCA)  \u25a1 C (Executive)  \u25a1 D (Mixed)",
        "response_domain": "Reassess SOZO PRS + MoCA / ADAS-Cog",
        "data_collection_table": [
            ["Data Point", "Completed", "Source"],
            ["MoCA / MMSE / ADAS-Cog score", "\u25a1", "Neuropsychological"],
            ["CDR staging and functional domains", "\u25a1", "CDR Assessment"],
            ["Episodic memory profile (RAVLT, story recall)", "\u25a1", "Neuropsychological"],
            ["Language assessment (naming, fluency)", "\u25a1", "Neuropsychological"],
            ["Visuospatial / constructional ability", "\u25a1", "Clock Draw / Tests"],
            ["Apathy and neuropsychiatric inventory (NPI)", "\u25a1", "NPI / Carer Report"],
            ["Neuroimaging findings (MRI pattern, FDG-PET)", "\u25a1", "Radiology Report"],
            ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"],
        ],
        "phenotype_table": [
            ["Dominant Features", "Assign Phenotype", "Code"],
            ["Prominent episodic memory impairment, hippocampal pattern", "Amnestic (AM)", "AM"],
            ["Multiple cognitive domains impaired, executive + memory", "Multi-Domain (MD)", "MD"],
            ["Onset before age 65, often atypical presentation", "Early-Onset (EO)", "EO"],
            ["Typical late-onset amnestic pattern, age >65", "Late-Onset (LO)", "LO"],
            ["Visuospatial, constructional, occipital predominance", "Posterior Cortical (PC)", "PC"],
            ["Progressive naming difficulty, anomia, semantic impairment", "Logopenic (LG)", "LG"],
            ["Profound apathy, amotivation, reduced initiation", "Apathy-Dominant (AP)", "AP"],
        ],
        "network_table": [
            ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
            ["AM", "DMN (medial temporal)", "CEN (L-DLPFC)", "DAN"],
            ["MD", "CEN (L-DLPFC)", "DMN", "DAN (parietal)"],
            ["EO", "CEN bilateral", "DMN", "Limbic"],
            ["LO", "DMN (hippocampal)", "CEN", "DAN"],
            ["PC", "DAN (parietal-occipital)", "DMN", "CEN"],
            ["LG", "Temporal (L-STG/MTG)", "CEN (L-DLPFC)", "DMN"],
            ["AP", "Limbic (anterior cingulate)", "CEN (L-DLPFC)", "SN"],
        ],
        "tdcs_group1": [
            ["Goal", "Anode", "Cathode"],
            ["Enhance L-DLPFC encoding support (AM/MD)", "Left DLPFC (F3)", "Right supraorbital"],
            ["Temporal lobe memory circuit (AM)", "Left temporal (T3/P3)", "Right supraorbital"],
            ["Bilateral DLPFC executive (MD/EO)", "Bilateral DLPFC", "Extracephalic"],
        ],
        "tdcs_group2": [
            ["Goal", "Anode", "Cathode"],
            ["Left temporal language (LG)", "Left temporal (T3)", "Right supraorbital"],
            ["Parietal visuospatial (PC)", "Left parietal (P3)", "Right supraorbital"],
        ],
        "tdcs_group3": [
            ["Goal", "Anode", "Cathode"],
            ["Bilateral DLPFC early-onset (EO)", "Left DLPFC", "Right DLPFC (bi-cephalic)"],
            ["Prefrontal apathy circuit (AP)", "Left DLPFC", "Right supraorbital"],
        ],
        "tps_table": [
            ["Clinical Target", "TPS Protocol", "Phenotype"],
            ["L-DLPFC memory consolidation", "FT1", "AM, MD, LO"],
            ["Bilateral DLPFC executive network", "FT3", "EO, MD"],
            ["Hippocampal-entorhinal circuit", "FT4", "AM, LO"],
            ["Parietal attention network", "FT5", "PC, MD"],
            ["Left temporal language areas", "FT7", "LG"],
            ["Anterior cingulate apathy circuit", "FT8", "AP"],
            ["Default mode network hubs", "FT9", "AM, LO"],
            ["Bilateral temporal-parietal", "Multiple", "PC, LG"],
            ["Multi-network distributed AD", "Multiple", "MD / advanced"],
        ],
        "tavns_table": [
            ["Domain", "Details"],
            ["Indications", "Apathy, sleep disruption, autonomic dysfunction in AD/MCI"],
            ["Protocol", "20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
            ["Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
        ],
        "ces_table": [
            ["Domain", "Details"],
            ["Indications", "Sleep disruption, anxiety, agitation in AD/MCI"],
            ["Protocol", "20\u201360 min daily, Alpha-Stim AID, especially AP and agitated phenotypes"],
            ["Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
        ],
        "sequencing_table": [
            ["Option", "Sequence", "Phenotype"],
            ["OPTION A \u2014 Amnestic", "Wk 1\u20134: L-DLPFC + temporal tDCS (5\u00d7/wk) \u2192 Wk 2\u20136: TPS hippocampal/DLPFC", "AM / LO"],
            ["OPTION B \u2014 Language/PCA", "Wk 1\u20134: L-temporal tDCS + cognitive tasks \u2192 Wk 3\u20136: TPS temporal/parietal", "LG / PC"],
            ["OPTION C \u2014 Executive", "Wk 1\u20132: CES stabilization \u2192 Wk 2\u20136: Bilateral DLPFC tDCS + TPS CEN", "EO / MD"],
            ["OPTION D \u2014 Apathy/Mixed", "Sequential: TPS (highest burden) \u2192 Maintenance L-DLPFC tDCS + taVNS", "AP / MX"],
        ],
    },
]


if __name__ == "__main__":
    print("Generating Partners Phenotype Classification documents (Batch 2)...")
    for c in CONDITIONS:
        print(f"\nBuilding: {c['full']}")
        build_phenotype_classification(c)
    print("\nDone.")
