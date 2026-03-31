"""
SOZO QA — Template conformity checker.
Verifies generated documents contain required sections and SOZO branding.
"""
import logging
from pathlib import Path

from ..core.enums import DocumentType, Tier
from ..template.gold_standard import get_required_sections

logger = logging.getLogger(__name__)

# Mapping from DocumentType value string to DocumentType enum for lookup
_DOC_TYPE_MAP: dict[str, DocumentType] = {dt.value: dt for dt in DocumentType}

# Mapping from Tier value string to Tier enum
_TIER_MAP: dict[str, Tier] = {t.value: t for t in Tier}


def _normalise_heading(text: str) -> str:
    """Lower-case, strip and collapse whitespace for loose heading matching."""
    return " ".join(text.lower().split())


class TemplateConformityChecker:
    """Checks that generated .docx files contain the required sections and branding."""

    def check_document(
        self,
        docx_path: Path,
        condition_slug: str,
        doc_type: str,
        tier: str,
    ) -> dict:
        """
        Check a single document for template conformity.

        Returns a dict with keys:
            file, has_header, has_footer, has_title,
            required_sections_found, required_sections_missing,
            has_tables, paragraph_count, passed
        """
        result: dict = {
            "file": str(docx_path),
            "has_header": False,
            "has_footer": False,
            "has_title": False,
            "required_sections_found": [],
            "required_sections_missing": [],
            "has_tables": False,
            "paragraph_count": 0,
            "passed": False,
        }

        if not docx_path.exists():
            logger.warning(
                "Template conformity: file not found: %s", docx_path
            )
            result["required_sections_missing"] = self._required_sections(
                doc_type, tier
            )
            return result

        try:
            from docx import Document  # python-docx

            doc = Document(str(docx_path))
        except Exception as exc:
            logger.warning(
                "Template conformity: could not open %s — %s", docx_path, exc
            )
            result["required_sections_missing"] = self._required_sections(
                doc_type, tier
            )
            return result

        # --- Header / footer ---
        try:
            section0 = doc.sections[0]
            header_paras = section0.header.paragraphs
            result["has_header"] = any(
                p.text.strip() for p in header_paras
            )
            footer_paras = section0.footer.paragraphs
            result["has_footer"] = any(
                p.text.strip() for p in footer_paras
            )
        except Exception as exc:
            logger.warning(
                "Template conformity: error reading header/footer in %s — %s",
                docx_path,
                exc,
            )

        # --- Paragraphs ---
        paragraphs = doc.paragraphs
        result["paragraph_count"] = len(paragraphs)

        # Collect heading texts for section matching
        heading_texts: list[str] = []
        for para in paragraphs:
            style_name = para.style.name.lower() if para.style else ""
            if "heading" in style_name or "title" in style_name:
                heading_texts.append(_normalise_heading(para.text))

        # Title: any Heading 1 or Title style paragraph
        result["has_title"] = bool(heading_texts)

        # --- Tables ---
        result["has_tables"] = len(doc.tables) > 0

        # --- Required sections ---
        required = self._required_sections(doc_type, tier)
        found: list[str] = []
        missing: list[str] = []

        for section_id in required:
            # Convert snake_case section id to space-separated for loose matching
            normalised_id = _normalise_heading(section_id.replace("_", " "))
            matched = any(
                normalised_id in heading or heading in normalised_id
                for heading in heading_texts
            )
            if matched:
                found.append(section_id)
            else:
                missing.append(section_id)

        result["required_sections_found"] = found
        result["required_sections_missing"] = missing

        # Passed: has title, no missing required sections
        result["passed"] = result["has_title"] and len(missing) == 0

        if missing:
            logger.warning(
                "Template conformity: %s missing sections: %s",
                docx_path.name,
                missing,
            )

        return result

    # ------------------------------------------------------------------
    # Batch helpers
    # ------------------------------------------------------------------

    def check_condition_documents(
        self, condition_slug: str, output_dir: Path
    ) -> list[dict]:
        """Check all documents for a condition in output_dir."""
        from ..core.enums import DocumentType as DT, Tier as T

        results: list[dict] = []

        for doc_type_enum in DT:
            for tier_enum in T:
                filename = (
                    f"{condition_slug}_{doc_type_enum.value}_{tier_enum.value}.docx"
                )
                docx_path = output_dir / filename
                result = self.check_document(
                    docx_path,
                    condition_slug,
                    doc_type_enum.value,
                    tier_enum.value,
                )
                results.append(result)

        return results

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _required_sections(self, doc_type: str, tier: str) -> list[str]:
        """Return required section IDs for a given doc_type / tier pair."""
        dt_enum = _DOC_TYPE_MAP.get(doc_type)
        tier_enum = _TIER_MAP.get(tier)
        if dt_enum is None or tier_enum is None:
            logger.warning(
                "Template conformity: unknown doc_type '%s' or tier '%s'",
                doc_type,
                tier,
            )
            return []
        return get_required_sections(dt_enum, tier_enum)
