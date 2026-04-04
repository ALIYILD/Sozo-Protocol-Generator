"""
Template Profile Builder — creates TemplateProfile from uploaded DOCX files.

Uses the existing TemplateParser for basic structure extraction, then enriches
with formatting analysis, table schema detection, figure placement detection,
tone profiling, and placeholder mapping.

Usage:
    from sozo_generator.template_profiles.builder import build_template_profile

    profile = build_template_profile("path/to/template.docx", name="My Handbook")
"""
from __future__ import annotations

import hashlib
import logging
import re
import statistics
import uuid
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt

from .models import (
    FigureSpec,
    FormattingProfile,
    TableSpec,
    TemplateProfile,
    TemplateSectionSpec,
    ToneProfile,
)

logger = logging.getLogger(__name__)

# ── Condition name patterns for detection ──────────────────────────────────
_CONDITION_PATTERNS = {
    "parkinsons": [r"parkinson", r"\bPD\b"],
    "depression": [r"depress", r"\bMDD\b", r"major depressive"],
    "anxiety": [r"anxiety", r"\bGAD\b", r"generaliz.d anxiety"],
    "adhd": [r"\bADHD\b", r"attention.deficit"],
    "alzheimers": [r"alzheimer", r"\bMCI\b", r"mild cognitive"],
    "stroke_rehab": [r"stroke", r"post.stroke", r"cerebrovascular"],
    "tbi": [r"traumatic brain", r"\bTBI\b"],
    "chronic_pain": [r"chronic pain", r"fibromyalgia"],
    "ptsd": [r"\bPTSD\b", r"post.traumatic stress"],
    "ocd": [r"\bOCD\b", r"obsessive.compulsive"],
    "ms": [r"multiple sclerosis", r"\bMS\b"],
    "asd": [r"autism", r"\bASD\b", r"spectrum disorder"],
    "long_covid": [r"long.covid", r"post.covid", r"brain fog"],
    "tinnitus": [r"tinnitus"],
    "insomnia": [r"insomnia", r"sleep disorder"],
}

# ── Doc type keywords for inference ────────────────────────────────────────
_DOC_TYPE_KEYWORDS = {
    "handbook": ["handbook", "stage_1", "stage_2", "patient journey", "modalities", "governance"],
    "evidence_based_protocol": ["pathophysiology", "inclusion", "exclusion", "protocol", "evidence"],
    "clinical_exam": ["examination", "checklist", "assessment", "screening"],
    "phenotype_classification": ["phenotype", "classification", "protocol mapping"],
    "responder_tracking": ["responder", "tracking", "response criteria"],
    "psych_intake": ["psychological", "intake", "prs baseline"],
    "network_assessment": ["network", "bedside assessment", "dmn", "cen", "salience"],
    "all_in_one_protocol": ["all-in-one", "compendium", "tps", "tdcs", "multimodal"],
}

# ── Visual type hints ─────────────────────────────────────────────────────
_VISUAL_KEYWORDS = {
    "brain_map": ["brain map", "brain region", "target map"],
    "eeg_topomap": ["eeg", "10-20", "10-10", "electrode", "montage", "topomap"],
    "network_diagram": ["network diagram", "fnon", "connectivity"],
    "mri_target_view": ["mri", "sagittal", "stimulation target"],
    "protocol_panel": ["protocol panel", "montage panel"],
    "treatment_timeline": ["timeline", "session sequence", "sozo sequence"],
    "dose_response": ["dose response", "responder", "tracking chart"],
    "symptom_flow": ["symptom flow", "symptom network"],
}


def build_template_profile(
    template_path: str | Path,
    name: str = "",
    profile_id: str = "",
) -> TemplateProfile:
    """Build a TemplateProfile from a DOCX file.

    This is the main entry point for template ingestion.

    Args:
        template_path: Path to the .docx file
        name: Human-readable profile name (auto-generated if empty)
        profile_id: Profile ID (auto-generated if empty)

    Returns:
        TemplateProfile with full structure, formatting, and tone analysis
    """
    path = Path(template_path)
    if not path.exists():
        raise FileNotFoundError(f"Template not found: {path}")

    doc = Document(str(path))

    if not profile_id:
        profile_id = f"tp-{uuid.uuid4().hex[:12]}"
    if not name:
        name = path.stem.replace("_", " ").replace("-", " ")

    logger.info(f"Building template profile from: {path.name}")

    # 1. Extract sections
    sections = _extract_sections(doc)

    # 2. Extract tables
    tables = _extract_tables(doc, sections)

    # 3. Extract figure placements
    figures = _extract_figures(doc, sections)

    # 4. Extract formatting profile
    formatting = _extract_formatting(doc)

    # 5. Build tone profile
    tone = _build_tone_profile(doc, sections)

    # 6. Detect placeholders
    placeholders = _detect_placeholders(doc)

    # 7. Detect condition
    source_condition = _detect_condition(doc)

    # 8. Infer doc type
    inferred_type, template_type = _infer_doc_type(sections)

    # 9. Detect tier
    tier = _detect_tier(doc)

    # 10. Compute fingerprint
    fingerprint = _compute_fingerprint(path)

    # 11. Build heading hierarchy
    hierarchy = {}
    for s in sections:
        level = s.heading_level
        hierarchy[level] = hierarchy.get(level, 0) + 1

    profile = TemplateProfile(
        profile_id=profile_id,
        name=name,
        source_filename=path.name,
        template_type=template_type,
        inferred_doc_type=inferred_type,
        tier=tier,
        source_condition=source_condition,
        section_map=sections,
        heading_hierarchy=hierarchy,
        table_specs=tables,
        figure_specs=figures,
        formatting_profile=formatting,
        tone_profile=tone,
        placeholder_map=placeholders,
        template_fingerprint=fingerprint,
    )

    logger.info(
        f"Template profile built: {profile.total_sections} sections, "
        f"{profile.total_tables} tables, {profile.total_figures} figures"
    )
    return profile


# ── Section Extraction ─────────────────────────────────────────────────────

def _extract_sections(doc: Document) -> list[TemplateSectionSpec]:
    """Extract section hierarchy from DOCX headings."""
    sections = []
    current_section = None
    word_buffer = []
    para_count = 0

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        text = para.text.strip()

        # Detect headings
        is_heading = False
        heading_level = 0
        if "heading" in style_name:
            try:
                heading_level = int(style_name.split()[-1])
                is_heading = True
            except (ValueError, IndexError):
                is_heading = "heading" in style_name
                heading_level = 1

        # Also detect ALL CAPS lines as potential section headers
        if not is_heading and text and len(text) > 3 and text == text.upper() and not text.startswith("["):
            is_heading = True
            heading_level = 1

        if is_heading and text:
            # Save previous section
            if current_section is not None:
                word_count = sum(len(w.split()) for w in word_buffer)
                current_section.estimated_word_count = word_count
                current_section.estimated_paragraph_count = para_count
                current_section.source_excerpt_sample = " ".join(word_buffer)[:200]
                sections.append(current_section)

            # Start new section
            section_id = _normalize_section_id(text)
            current_section = TemplateSectionSpec(
                section_id=section_id,
                title=text,
                normalized_title=_normalize_title(text),
                heading_level=heading_level,
                order_index=len(sections),
            )
            word_buffer = []
            para_count = 0
        elif text and current_section is not None:
            word_buffer.append(text)
            para_count += 1

    # Save last section
    if current_section is not None:
        word_count = sum(len(w.split()) for w in word_buffer)
        current_section.estimated_word_count = word_count
        current_section.estimated_paragraph_count = para_count
        current_section.source_excerpt_sample = " ".join(word_buffer)[:200]
        sections.append(current_section)

    # Enrich sections with content kind
    for s in sections:
        s.content_kind = _infer_content_kind(s)
        s.citation_density_hint = _infer_citation_density(s)
        s.requires_evidence = s.citation_density_hint in ("medium", "high")

    return sections


def _normalize_section_id(title: str) -> str:
    """Create a normalized section ID from a heading title."""
    # Remove numbering prefixes
    cleaned = re.sub(r"^\d+[\.\)]\s*", "", title)
    cleaned = re.sub(r"^[A-Z]\d*[\.\)]\s*", "", cleaned)
    cleaned = re.sub(r"^Stage\s+\d+\s*[:\-\u2014]?\s*", "stage_", cleaned, flags=re.IGNORECASE)
    # Normalize
    normalized = re.sub(r"[^a-z0-9]+", "_", cleaned.lower()).strip("_")
    return normalized[:60] or "untitled"


def _normalize_title(title: str) -> str:
    """Create a condition-agnostic normalized title."""
    normalized = title
    for patterns in _CONDITION_PATTERNS.values():
        for pattern in patterns:
            normalized = re.sub(pattern, "{CONDITION}", normalized, flags=re.IGNORECASE)
    return normalized.strip()


def _infer_content_kind(section: TemplateSectionSpec) -> str:
    """Infer the content kind of a section."""
    title_lower = section.title.lower()
    if any(kw in title_lower for kw in ["checklist", "scoring", "form"]):
        return "checklist"
    if any(kw in title_lower for kw in ["reference", "citation", "bibliography"]):
        return "reference_list"
    if section.table_expected or section.table_count > 0:
        return "table_heavy"
    if section.estimated_word_count < 50:
        return "brief"
    return "prose"


def _infer_citation_density(section: TemplateSectionSpec) -> str:
    """Infer expected citation density from section characteristics."""
    title_lower = section.title.lower()
    if any(kw in title_lower for kw in ["pathophysiology", "evidence", "protocol", "mechanism"]):
        return "high"
    if any(kw in title_lower for kw in ["overview", "anatomy", "brain region", "network"]):
        return "medium"
    if any(kw in title_lower for kw in ["safety", "contraindication", "monitoring"]):
        return "medium"
    return "low"


# ── Table Extraction ──────────────────────────────────────────────────────

def _extract_tables(doc: Document, sections: list[TemplateSectionSpec]) -> list[TableSpec]:
    """Extract table schemas from the document."""
    tables = []
    for idx, table in enumerate(doc.tables):
        headers = []
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]

        # Try to determine parent section
        parent_section = ""
        if sections:
            # Approximate: assign to section by order
            ratio = idx / max(len(doc.tables), 1)
            section_idx = min(int(ratio * len(sections)), len(sections) - 1)
            parent_section = sections[section_idx].section_id
            sections[section_idx].table_expected = True
            sections[section_idx].table_count += 1

        purpose = _infer_table_purpose(headers)

        tables.append(TableSpec(
            table_index=idx,
            parent_section_id=parent_section,
            headers=headers,
            col_count=len(headers),
            row_count=len(table.rows),
            has_header_row=True,
            estimated_purpose=purpose,
        ))

    return tables


def _infer_table_purpose(headers: list[str]) -> str:
    """Infer table purpose from headers."""
    header_text = " ".join(h.lower() for h in headers)
    if any(kw in header_text for kw in ["score", "rating", "0-10", "severity"]):
        return "scoring"
    if any(kw in header_text for kw in ["protocol", "target", "parameter", "montage"]):
        return "protocol_params"
    if any(kw in header_text for kw in ["modality", "device", "evidence"]):
        return "modality_summary"
    if any(kw in header_text for kw in ["inclusion", "exclusion", "criteria"]):
        return "criteria_list"
    if any(kw in header_text for kw in ["field", "value", "patient", "date"]):
        return "demographics"
    if any(kw in header_text for kw in ["yes", "no", "check"]):
        return "checklist"
    return "general"


# ── Figure Detection ──────────────────────────────────────────────────────

def _extract_figures(doc: Document, sections: list[TemplateSectionSpec]) -> list[FigureSpec]:
    """Detect figure/image placements in the document."""
    figures = []
    fig_idx = 0

    for para in doc.paragraphs:
        # Check for inline images
        has_image = False
        for run in para.runs:
            if run._element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                has_image = True
            if run._element.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
                has_image = True

        # Also check for figure caption patterns
        text = para.text.strip().lower()
        is_figure_ref = text.startswith("figure") or "[figure" in text or "fig." in text

        if has_image or is_figure_ref:
            # Determine parent section
            parent = ""
            if sections:
                # Approximate by position
                for s in reversed(sections):
                    if s.order_index <= fig_idx:
                        parent = s.section_id
                        s.visual_expected = True
                        break

            # Infer visual type from context
            visual_type = _infer_visual_type(para.text)

            figures.append(FigureSpec(
                figure_index=fig_idx,
                parent_section_id=parent,
                caption_text=para.text.strip() if is_figure_ref else "",
                estimated_type=visual_type,
                approximate_position="inline",
            ))
            fig_idx += 1

    return figures


def _infer_visual_type(context: str) -> str:
    """Infer visual type from surrounding text."""
    context_lower = context.lower()
    for vtype, keywords in _VISUAL_KEYWORDS.items():
        if any(kw in context_lower for kw in keywords):
            return vtype
    return "general"


# ── Formatting Extraction ─────────────────────────────────────────────────

def _extract_formatting(doc: Document) -> FormattingProfile:
    """Extract formatting profile from document styles and content."""
    fonts = []
    sizes = []
    heading_sizes: dict[int, list[float]] = {}

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        for run in para.runs:
            if run.font.name:
                fonts.append(run.font.name)
            if run.font.size:
                size_pt = run.font.size.pt
                sizes.append(size_pt)

                # Track heading sizes
                if "heading" in style_name:
                    try:
                        level = int(style_name.split()[-1])
                        heading_sizes.setdefault(level, []).append(size_pt)
                    except (ValueError, IndexError):
                        pass

    # Determine primary font
    primary_font = "Calibri"
    if fonts:
        from collections import Counter
        primary_font = Counter(fonts).most_common(1)[0][0]

    # Determine body size
    body_size = 11.0
    if sizes:
        body_size = statistics.median(sizes)

    # Heading sizes
    avg_heading_sizes = {}
    for level, s_list in heading_sizes.items():
        avg_heading_sizes[level] = round(statistics.mean(s_list), 1)

    return FormattingProfile(
        primary_font=primary_font,
        heading_font=primary_font,
        body_font_size_pt=body_size,
        heading_sizes_pt=avg_heading_sizes or {1: 16.0, 2: 14.0, 3: 12.0},
    )


# ── Tone Profiling ────────────────────────────────────────────────────────

def _build_tone_profile(doc: Document, sections: list[TemplateSectionSpec]) -> ToneProfile:
    """Build a tone profile from document prose analysis."""
    all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    sample_paras = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 50][:10]

    # Detect voice
    voice = "third_person"
    if re.search(r"\b(we|our|us)\b", all_text, re.IGNORECASE):
        voice = "first_person_plural"
    if re.search(r"\b(ensure|document|assess|evaluate|monitor)\b", all_text, re.IGNORECASE):
        voice = "imperative"

    # Detect formality
    formality = "formal"
    if re.search(r"\b(don't|can't|won't|it's)\b", all_text):
        formality = "semi_formal"

    # Detect citation style
    citation_style = "inline_pmid"
    if re.search(r"PMID[:\s]*\d+", all_text):
        citation_style = "inline_pmid"
    elif re.search(r"\[\d+\]", all_text):
        citation_style = "numbered"
    elif re.search(r"\(\w+\s+et\s+al\.,?\s*\d{4}\)", all_text):
        citation_style = "author_year"

    # Section depth
    avg_words = 0
    if sections:
        word_counts = [s.estimated_word_count for s in sections if s.estimated_word_count > 0]
        if word_counts:
            avg_words = statistics.mean(word_counts)

    if avg_words > 500:
        depth = "comprehensive"
    elif avg_words > 200:
        depth = "detailed"
    elif avg_words > 80:
        depth = "moderate"
    else:
        depth = "brief"

    return ToneProfile(
        audience="clinician",
        formality=formality,
        voice=voice,
        section_depth=depth,
        citation_style=citation_style,
        uses_abbreviations=bool(re.search(r"\b[A-Z]{2,6}\b", all_text)),
        uses_clinical_warnings=bool(re.search(r"WARNING|CAUTION|MANDATORY|OFF.LABEL", all_text, re.IGNORECASE)),
        uses_governance_boxes=bool(re.search(r"GOVERNANCE|governance rule", all_text, re.IGNORECASE)),
        sample_excerpts=sample_paras[:5],
    )


# ── Placeholder Detection ─────────────────────────────────────────────────

def _detect_placeholders(doc: Document) -> dict[str, list[str]]:
    """Detect placeholder patterns in the document."""
    placeholders: dict[str, list[str]] = {}
    patterns = [
        r"\[(?:PLACEHOLDER|TODO|REVIEW|INSERT|TBD)[^\]]*\]",
        r"\{(?:CONDITION|PATIENT|DATE|NAME)[^\}]*\}",
        r"___+",  # Blank fill lines
    ]
    combined = "|".join(patterns)

    current_section = "document"
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if "heading" in style_name and para.text.strip():
            current_section = _normalize_section_id(para.text.strip())

        matches = re.findall(combined, para.text)
        if matches:
            placeholders.setdefault(current_section, []).extend(matches)

    return placeholders


# ── Detection Helpers ─────────────────────────────────────────────────────

def _detect_condition(doc: Document) -> Optional[str]:
    """Detect which condition the template is about."""
    text = " ".join(p.text for p in doc.paragraphs[:30])  # Check first 30 paragraphs
    for slug, patterns in _CONDITION_PATTERNS.items():
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return slug
    return None


def _detect_tier(doc: Document) -> Optional[str]:
    """Detect tier from document content."""
    text = " ".join(p.text for p in doc.paragraphs[:20]).lower()
    if "partners" in text or "fnon" in text:
        return "partners"
    if "fellow" in text:
        return "fellow"
    return None


def _infer_doc_type(sections: list[TemplateSectionSpec]) -> tuple[Optional[str], str]:
    """Infer document type from section structure."""
    section_text = " ".join(s.title.lower() for s in sections)

    best_type = None
    best_score = 0
    template_type = "document"

    for doc_type, keywords in _DOC_TYPE_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in section_text)
        if score > best_score:
            best_score = score
            best_type = doc_type

    # Map to template_type
    type_map = {
        "handbook": "handbook",
        "evidence_based_protocol": "protocol",
        "all_in_one_protocol": "protocol",
        "clinical_exam": "assessment",
        "phenotype_classification": "assessment",
        "responder_tracking": "assessment",
        "psych_intake": "assessment",
        "network_assessment": "assessment",
    }
    template_type = type_map.get(best_type, "document")

    return best_type, template_type


def _compute_fingerprint(path: Path) -> str:
    """Compute a content fingerprint for the template."""
    content = path.read_bytes()
    return hashlib.sha256(content).hexdigest()[:16]
