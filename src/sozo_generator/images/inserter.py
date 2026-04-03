"""
Document Image Inserter.

Inserts curated images into DOCX documents at appropriate section locations,
with proper captions, attributions, and formatting.

Usage:
    from sozo_generator.images.inserter import DocumentImageInserter
    from sozo_generator.images.curator import ImageCurator

    curator = ImageCurator()
    manifest = curator.curate_for_document("parkinsons", "evidence_based_protocol", "partners")

    inserter = DocumentImageInserter()
    inserter.insert_images(doc, manifest)
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .curator import CuratedImage, DocumentImageManifest

logger = logging.getLogger(__name__)

# ── Formatting constants ────────────────────────────────────────────────────
try:
    from ..docx.styles import COLOR_GRAY, FONT_BODY
except ImportError:
    from docx.shared import RGBColor
    COLOR_GRAY = RGBColor(0x88, 0x88, 0x88)
    FONT_BODY = "Calibri"


class DocumentImageInserter:
    """Inserts curated images into DOCX documents.

    Handles:
    - Image sizing and positioning
    - Caption formatting
    - Source attribution (required for all non-proprietary images)
    - Fallback to placeholder if image file is missing
    """

    def __init__(
        self,
        default_width_inches: float = 5.0,
        include_attribution: bool = True,
        figure_counter_start: int = 1,
    ):
        self.default_width = default_width_inches
        self.include_attribution = include_attribution
        self._figure_counter = figure_counter_start

    def insert_images(
        self,
        doc: Document,
        manifest: DocumentImageManifest,
        section_filter: list[str] | None = None,
    ) -> int:
        """Insert all curated images into a document.

        Images are appended to the end of the document in section order.
        For more precise placement, use insert_image_at_section().

        Args:
            doc: python-docx Document object
            manifest: ImageManifest with curated images
            section_filter: Only insert images for these sections (None = all)

        Returns:
            Number of images successfully inserted
        """
        inserted = 0
        sections = section_filter or list(
            set(img.section_id for img in manifest.images)
        )

        for section_id in sections:
            images = manifest.images_for_section(section_id)
            for img in images:
                success = self._insert_single(doc, img)
                if success:
                    inserted += 1

        logger.info(f"Inserted {inserted} images into document")
        return inserted

    def insert_image_for_section(
        self,
        doc: Document,
        manifest: DocumentImageManifest,
        section_id: str,
    ) -> int:
        """Insert images for a specific section."""
        images = manifest.images_for_section(section_id)
        inserted = 0
        for img in images:
            if self._insert_single(doc, img):
                inserted += 1
        return inserted

    def _insert_single(self, doc: Document, img: CuratedImage) -> bool:
        """Insert a single curated image into the document."""
        path = Path(img.local_path)
        if not path.exists():
            logger.warning(f"Image file not found: {path}")
            self._insert_placeholder(doc, img)
            return False

        try:
            # Image paragraph
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            width = Inches(img.width_inches or self.default_width)
            run.add_picture(str(path), width=width)
            img_para.paragraph_format.space_after = Pt(4)

            # Caption
            fig_num = self._figure_counter
            self._figure_counter += 1
            caption_text = f"Figure {fig_num}: {img.caption}"

            cap_para = doc.add_paragraph()
            cap_run = cap_para.add_run(caption_text)
            cap_run.font.name = FONT_BODY
            cap_run.font.size = Pt(9)
            cap_run.font.italic = True
            cap_run.font.color.rgb = COLOR_GRAY
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Attribution (if external image)
            if self.include_attribution and img.source != "precurated":
                attr_text = self._build_attribution(img)
                if attr_text:
                    attr_para = doc.add_paragraph()
                    attr_run = attr_para.add_run(attr_text)
                    attr_run.font.name = FONT_BODY
                    attr_run.font.size = Pt(7)
                    attr_run.font.italic = True
                    attr_run.font.color.rgb = COLOR_GRAY
                    attr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    attr_para.paragraph_format.space_after = Pt(10)

            return True

        except Exception as e:
            logger.error(f"Failed to insert image {path}: {e}")
            self._insert_placeholder(doc, img)
            return False

    def _build_attribution(self, img: CuratedImage) -> str:
        """Build attribution text for an image."""
        parts = []
        if img.attribution:
            parts.append(f"Source: {img.attribution}")
        if img.license and img.license != "unknown":
            parts.append(f"License: {img.license}")
        if img.original_url:
            parts.append(f"URL: {img.original_url}")
        return " | ".join(parts) if parts else ""

    def _insert_placeholder(self, doc: Document, img: CuratedImage):
        """Insert a placeholder for a missing image."""
        text = f"[FIGURE: {img.caption} — image not available]"
        para = doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            run.font.color.rgb = COLOR_GRAY
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)
