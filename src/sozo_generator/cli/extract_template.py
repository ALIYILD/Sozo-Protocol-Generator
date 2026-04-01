"""Extract and analyze a template directory."""
from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Optional

import typer

extract_app = typer.Typer(
    name="extract-template",
    help="Extract and analyze template directories.",
    add_completion=False,
)


def _analyze_docx(docx_path: Path) -> dict:
    """Extract paragraph count, heading structure, and table count from a .docx file."""
    try:
        from docx import Document  # type: ignore
        doc = Document(str(docx_path))

        paragraphs = doc.paragraphs
        para_count = len(paragraphs)

        headings = []
        for para in paragraphs:
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                level = style_name.replace("Heading", "").strip()
                headings.append({"level": level, "text": para.text.strip()})

        table_count = len(doc.tables)

        return {
            "path": str(docx_path),
            "paragraph_count": para_count,
            "heading_count": len(headings),
            "headings": headings,
            "table_count": table_count,
        }
    except Exception as e:
        return {
            "path": str(docx_path),
            "error": str(e),
            "paragraph_count": 0,
            "heading_count": 0,
            "headings": [],
            "table_count": 0,
        }


@extract_app.command("extract")
def extract_template(
    input_path: Path = typer.Option(..., "--input", help="Path to template folder"),
    output_path: Optional[Path] = typer.Option(None, "--output", help="Where to write analysis JSON"),
):
    """Walk a template directory, analyze all .docx files, and produce a structured report."""
    if not input_path.exists():
        typer.echo(typer.style(f"Input path does not exist: {input_path}", fg=typer.colors.RED))
        raise typer.Exit(code=1)

    if not input_path.is_dir():
        typer.echo(typer.style(f"Input path is not a directory: {input_path}", fg=typer.colors.RED))
        raise typer.Exit(code=1)

    typer.echo(
        typer.style(f"Scanning template directory: {input_path}", fg=typer.colors.CYAN)
    )

    # Walk directory and collect .docx files
    docx_files: list[Path] = []
    for root, _dirs, files in os.walk(input_path):
        for fname in files:
            if fname.lower().endswith(".docx"):
                docx_files.append(Path(root) / fname)

    if not docx_files:
        typer.echo(typer.style("No .docx files found in the input directory.", fg=typer.colors.YELLOW))
        raise typer.Exit(code=0)

    # Group by subfolder relative to input_path
    # Detect Fellow / Partners grouping by folder name
    fellow_docs: list[dict] = []
    partners_docs: list[dict] = []
    other_docs: list[dict] = []

    for docx_path in sorted(docx_files):
        relative = docx_path.relative_to(input_path)
        parts = relative.parts

        typer.echo(f"  Analyzing: {relative}")
        analysis = _analyze_docx(docx_path)
        analysis["relative_path"] = str(relative)

        # Categorize by subfolder keywords
        parts_lower = [p.lower() for p in parts]
        if any("fellow" in p for p in parts_lower):
            fellow_docs.append(analysis)
        elif any("partner" in p for p in parts_lower):
            partners_docs.append(analysis)
        else:
            other_docs.append(analysis)

    total_count = len(docx_files)
    fellow_count = len(fellow_docs)
    partners_count = len(partners_docs)
    other_count = len(other_docs)

    report = {
        "input_directory": str(input_path),
        "total_documents": total_count,
        "fellow_count": fellow_count,
        "partners_count": partners_count,
        "other_count": other_count,
        "fellow_documents": fellow_docs,
        "partners_documents": partners_docs,
        "other_documents": other_docs,
    }

    if output_path:
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(report, f, indent=2, ensure_ascii=False)
            typer.echo("")
            typer.echo(
                typer.style(f"Analysis written to: {output_path}", fg=typer.colors.GREEN)
            )
        except Exception as e:
            typer.echo(typer.style(f"Failed to write output: {e}", fg=typer.colors.RED))
            raise typer.Exit(code=1)
    else:
        typer.echo("")
        typer.echo(json.dumps(report, indent=2, ensure_ascii=False))

    typer.echo("")
    typer.echo(
        typer.style(
            f"Found {total_count} documents, {fellow_count} Fellow / {partners_count} Partners"
            + (f" / {other_count} Other" if other_count else ""),
            fg=typer.colors.GREEN,
        )
    )
