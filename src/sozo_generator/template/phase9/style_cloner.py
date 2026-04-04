"""
style_cloner.py — Injects DocumentBlueprint styles into a python-docx Document.

Takes a DocumentBlueprint and writes all its StyleSpec definitions into the
styles table of a python-docx Document object, so any document created via
Document() inherits the uploaded template's exact visual identity.

Phase 9, pipeline_version: phase9_v1
"""

from __future__ import annotations

import logging
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from .blueprint_schema import DocumentBlueprint, PageLayout, StyleSpec

logger = logging.getLogger(__name__)


class StyleCloner:
    """
    Injects styles from a DocumentBlueprint into a python-docx Document.

    Usage:
        doc = Document()
        cloner = StyleCloner(blueprint)
        doc = cloner.apply(doc)
        # doc now has all blueprint styles applied
    """

    _ALIGNMENT_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def __init__(self, blueprint: DocumentBlueprint) -> None:
        self.blueprint = blueprint

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def apply(self, doc: Document) -> Document:
        """
        Apply all blueprint styles to the document.
        Returns the modified document (same object, mutated in place).

        Applies the page layout first, then iterates through all StyleSpec
        entries in the blueprint.  Any style that fails to apply is logged
        as a WARNING and skipped — the rest of the styles are still applied.
        """
        self.apply_page_layout(doc)

        total = len(self.blueprint.styles)
        success = 0
        for spec in self.blueprint.styles:
            try:
                self._apply_style(doc, spec)
                success += 1
            except Exception as exc:  # noqa: BLE001
                logger.warning(
                    "StyleCloner: failed to apply style '%s': %s",
                    spec.name,
                    exc,
                )

        logger.info(
            "StyleCloner: applied %d/%d styles from blueprint %s",
            success,
            total,
            self.blueprint.source_filename,
        )
        return doc

    def apply_page_layout(self, doc: Document) -> None:
        """Apply page size, margins, and orientation from blueprint.

        Applies A4/Letter dimensions (portrait or landscape), then sets the
        four margins from the blueprint's PageLayout.  Operates on the first
        section of the document.
        """
        layout: PageLayout = self.blueprint.page_layout
        section = doc.sections[0]

        # Page dimensions
        if layout.size == "A4":
            if layout.orientation == "portrait":
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
            else:
                section.page_width = Cm(29.7)
                section.page_height = Cm(21.0)
        elif layout.size == "Letter":
            if layout.orientation == "portrait":
                section.page_width = Cm(21.59)
                section.page_height = Cm(27.94)
            else:
                section.page_width = Cm(27.94)
                section.page_height = Cm(21.59)
        elif layout.size == "A3":
            if layout.orientation == "portrait":
                section.page_width = Cm(29.7)
                section.page_height = Cm(42.0)
            else:
                section.page_width = Cm(42.0)
                section.page_height = Cm(29.7)
        elif layout.size == "Legal":
            if layout.orientation == "portrait":
                section.page_width = Cm(21.59)
                section.page_height = Cm(35.56)
            else:
                section.page_width = Cm(35.56)
                section.page_height = Cm(21.59)

        # Margins
        section.top_margin = Cm(layout.margin_top_cm)
        section.bottom_margin = Cm(layout.margin_bottom_cm)
        section.left_margin = Cm(layout.margin_left_cm)
        section.right_margin = Cm(layout.margin_right_cm)

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _apply_style(self, doc: Document, spec: StyleSpec) -> None:
        """Apply one StyleSpec to the document's style definitions.

        Creates the named style if it does not already exist, then writes
        all font and paragraph-format properties from the StyleSpec.
        """
        style_obj = self._get_or_create_style(doc, spec.name, spec.style_type)
        self._apply_font_props(style_obj.font, spec)
        if spec.style_type == "paragraph":
            self._apply_para_format(style_obj.paragraph_format, spec)

    def _get_or_create_style(self, doc: Document, name: str, style_type: str):
        """Return the named style from the document, creating it if absent.

        Args:
            doc: The python-docx Document to search.
            name: The Word style name.
            style_type: One of 'paragraph', 'character', 'table', 'numbering'.

        Returns:
            The python-docx style object (existing or newly created).
        """
        try:
            return doc.styles[name]
        except KeyError:
            # Map style_type string to WD_STYLE_TYPE enum
            _type_map = {
                "paragraph": WD_STYLE_TYPE.PARAGRAPH,
                "character": WD_STYLE_TYPE.CHARACTER,
                "table": WD_STYLE_TYPE.TABLE,
                "numbering": WD_STYLE_TYPE.LIST,
            }
            wst = _type_map.get(style_type, WD_STYLE_TYPE.PARAGRAPH)
            return doc.styles.add_style(name, wst)

    def _apply_font_props(self, font, spec: StyleSpec) -> None:
        """Apply font name, size, bold, italic, and color to a python-docx Font.

        Each property is applied independently; a failure on one property does
        not prevent the others from being written.
        """
        try:
            if spec.font_name:
                font.name = spec.font_name
        except Exception as exc:  # noqa: BLE001
            logger.debug("font.name error for '%s': %s", spec.name, exc)

        try:
            if spec.font_size_pt:
                font.size = Pt(spec.font_size_pt)
        except Exception as exc:  # noqa: BLE001
            logger.debug("font.size error for '%s': %s", spec.name, exc)

        try:
            if spec.bold is not None:
                font.bold = spec.bold
        except Exception as exc:  # noqa: BLE001
            logger.debug("font.bold error for '%s': %s", spec.name, exc)

        try:
            if spec.italic is not None:
                font.italic = spec.italic
        except Exception as exc:  # noqa: BLE001
            logger.debug("font.italic error for '%s': %s", spec.name, exc)

        try:
            if spec.color_hex:
                font.color.rgb = RGBColor.from_string(spec.color_hex)
        except Exception as exc:  # noqa: BLE001
            logger.debug("font.color error for '%s': %s", spec.name, exc)

    def _apply_para_format(self, para_format, spec: StyleSpec) -> None:
        """Apply spacing, alignment, and indent to a python-docx ParagraphFormat.

        Each property is applied independently; a failure on one does not
        prevent the others from being written.
        """
        try:
            if spec.space_before_pt is not None:
                para_format.space_before = Pt(spec.space_before_pt)
        except Exception as exc:  # noqa: BLE001
            logger.debug("space_before error for '%s': %s", spec.name, exc)

        try:
            if spec.space_after_pt is not None:
                para_format.space_after = Pt(spec.space_after_pt)
        except Exception as exc:  # noqa: BLE001
            logger.debug("space_after error for '%s': %s", spec.name, exc)

        try:
            if spec.alignment:
                para_format.alignment = self._ALIGNMENT_MAP.get(spec.alignment)
        except Exception as exc:  # noqa: BLE001
            logger.debug("alignment error for '%s': %s", spec.name, exc)

        try:
            if spec.left_indent_cm is not None:
                para_format.left_indent = Cm(spec.left_indent_cm)
        except Exception as exc:  # noqa: BLE001
            logger.debug("left_indent error for '%s': %s", spec.name, exc)

        try:
            para_format.keep_with_next = spec.keep_with_next
        except Exception as exc:  # noqa: BLE001
            logger.debug("keep_with_next error for '%s': %s", spec.name, exc)

        try:
            para_format.page_break_before = spec.page_break_before
        except Exception as exc:  # noqa: BLE001
            logger.debug("page_break_before error for '%s': %s", spec.name, exc)


# ---------------------------------------------------------------------------
# Module-level convenience function
# ---------------------------------------------------------------------------


def clone_styles_from_blueprint(
    blueprint: DocumentBlueprint,
    doc: Optional[Document] = None,
) -> Document:
    """Apply blueprint styles to a new (or existing) Document.

    Creates a fresh python-docx Document if none is supplied, applies all
    styles from the blueprint, and returns the resulting Document.

    Args:
        blueprint: The :class:`DocumentBlueprint` whose styles to inject.
        doc: An existing Document to inject into.  If ``None``, a new
            Document is created.

    Returns:
        The Document with all blueprint styles applied.
    """
    if doc is None:
        doc = Document()
    return StyleCloner(blueprint).apply(doc)
