"""
blueprint_extractor.py — Reads an uploaded DOCX and extracts a DocumentBlueprint.

Captures page layout, all named paragraph/table styles, color palette, header/footer
templates, and section structure from a DOCX file without storing any content — only
design and structure.

Phase 9, pipeline_version: phase9_v1
"""

from __future__ import annotations

import logging
import re
import uuid
from pathlib import Path
from typing import Optional

from .blueprint_schema import (
    ColorPaletteEntry,
    DocumentBlueprint,
    HeaderFooterBlueprint,
    LayoutRegion,
    PageLayout,
    SectionBlueprint,
    StyleSpec,
    TableSchemaBlueprint,
    empty_blueprint,
)

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Emu, Pt, RGBColor
    import lxml.etree as etree
except ImportError:
    raise ImportError("python-docx is required. Install with: pip install python-docx")

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Module-level helpers
# ---------------------------------------------------------------------------

_EMU_PER_CM = 360_000
_EMU_PER_PT = 12_700

_ALIGN_MAP: dict[object, str] = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


def _emu_color_to_hex(rgb_color: object) -> Optional[str]:
    """Convert a python-docx RGBColor object to a 6-char uppercase hex string.

    Args:
        rgb_color: A ``docx.shared.RGBColor`` instance, or None.

    Returns:
        6-character uppercase hex string without '#', or None if the value is
        absent, auto-color, or cannot be converted.
    """
    if rgb_color is None:
        return None
    try:
        # RGBColor.__str__ returns a 6-char hex string like '1F3864'
        hex_str = str(rgb_color).upper()
        if len(hex_str) == 6:
            int(hex_str, 16)  # validate it's actually hex
            return hex_str
    except Exception:  # noqa: BLE001
        pass
    return None


def _xml_fill_color(cell: object) -> Optional[str]:
    """Extract the ``w:shd @w:fill`` background color from a table cell's XML.

    Args:
        cell: A ``docx.table._Cell`` instance.

    Returns:
        6-character uppercase hex string without '#', or None if absent or
        the fill is 'auto' / transparent.
    """
    try:
        tc = cell._tc  # type: ignore[attr-defined]
        tc_pr = tc.find(qn("w:tcPr"))
        if tc_pr is None:
            return None
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            return None
        # The fill attribute lives in the w: namespace
        fill = shd.get(qn("w:fill"))
        if fill and fill.upper() not in ("AUTO", "FFFFFF", "000000FF", ""):
            fill = fill.lstrip("#").upper()
            if len(fill) == 6:
                int(fill, 16)  # validate hex
                return fill
    except Exception:  # noqa: BLE001
        pass
    return None


# ---------------------------------------------------------------------------
# BlueprintExtractor
# ---------------------------------------------------------------------------


class BlueprintExtractor:
    """Extracts a DocumentBlueprint from an uploaded DOCX file.

    Reads all named paragraph and table styles, page layout, header/footer
    templates, and section hierarchy — without touching document content.

    Usage::

        extractor = BlueprintExtractor("my_template.docx")
        blueprint = extractor.extract()
        blueprint_json = blueprint.model_dump_json(indent=2)
    """

    # Headings in Word are "Heading 1" through "Heading 9"
    HEADING_STYLE_PATTERN = re.compile(r"^heading\s+(\d)$", re.IGNORECASE)

    # Styles to always include even if rarely used
    CORE_STYLES: frozenset[str] = frozenset(
        {"Normal", "Default Paragraph Font", "Table Normal"}
    )

    # Style name fragments that indicate a custom SOZO style (lowercase)
    SOZO_STYLE_MARKERS: frozenset[str] = frozenset(
        {"sozo", "protocol", "clinical", "evidence", "device"}
    )

    def __init__(self, docx_path: Path | str) -> None:
        """Initialise the extractor and open the DOCX file.

        Args:
            docx_path: Path to the ``.docx`` file to analyse.

        Raises:
            FileNotFoundError: If the file does not exist at the given path.
        """
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"Template not found: {self.docx_path}")
        self.doc = Document(str(self.docx_path))

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def extract(self) -> DocumentBlueprint:
        """Extract the full DocumentBlueprint from the loaded DOCX.

        Never raises — logs warnings for any extraction failures and continues.

        Returns:
            A :class:`DocumentBlueprint` (may be partial if the document is
            unusual or malformed).
        """
        warnings: list[str] = []

        # 1. Page layout
        try:
            page_layout = self._extract_page_layout()
        except Exception as exc:  # noqa: BLE001
            logger.warning("Page layout extraction failed: %s", exc)
            warnings.append(f"page_layout: {exc}")
            page_layout = PageLayout()

        # 2. Styles
        try:
            styles = self._extract_styles()
        except Exception as exc:  # noqa: BLE001
            logger.warning("Style extraction failed: %s", exc)
            warnings.append(f"styles: {exc}")
            styles = []

        # 3. Color palette
        try:
            color_palette = self._extract_color_palette(styles)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Color palette extraction failed: %s", exc)
            warnings.append(f"color_palette: {exc}")
            color_palette = []

        # 4. Table schemas
        try:
            table_schemas = self._extract_table_schemas()
        except Exception as exc:  # noqa: BLE001
            logger.warning("Table schema extraction failed: %s", exc)
            warnings.append(f"table_schemas: {exc}")
            table_schemas = []

        # 5. Header / footer
        try:
            header_bp, footer_bp = self._extract_header_footer()
        except Exception as exc:  # noqa: BLE001
            logger.warning("Header/footer extraction failed: %s", exc)
            warnings.append(f"header_footer: {exc}")
            header_bp, footer_bp = None, None

        # 6. Sections
        try:
            sections = self._detect_sections()
        except Exception as exc:  # noqa: BLE001
            logger.warning("Section detection failed: %s", exc)
            warnings.append(f"sections: {exc}")
            sections = []

        # 7. Document type
        try:
            document_type = self._detect_document_type(sections)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Document type detection failed: %s", exc)
            warnings.append(f"document_type: {exc}")
            document_type = None

        # 8. Detect flags
        has_cover = any(s.section_type == "cover" for s in sections)
        has_toc = any(s.section_type == "toc" for s in sections)
        has_header = header_bp is not None
        has_footer = footer_bp is not None

        # 9. Assemble blueprint
        blueprint = DocumentBlueprint(
            source_filename=self.docx_path.name,
            page_layout=page_layout,
            styles=styles,
            color_palette=color_palette,
            table_schemas=table_schemas,
            sections=sections,
            header=header_bp,
            footer=footer_bp,
            has_cover_page=has_cover,
            has_toc=has_toc,
            has_header=has_header,
            has_footer=has_footer,
            total_styles_extracted=len(styles),
            total_sections_detected=len(sections),
            document_type_guess=document_type,
            extraction_warnings=warnings,
        )

        logger.info(
            "Blueprint extracted: %d styles, %d sections from %s",
            len(styles),
            len(sections),
            self.docx_path.name,
        )
        return blueprint

    # ------------------------------------------------------------------
    # Private extraction helpers
    # ------------------------------------------------------------------

    def _extract_page_layout(self) -> PageLayout:
        """Extract physical page dimensions, margins, and column layout.

        Returns:
            A :class:`PageLayout` populated from the first section of the
            document.  Falls back to a default ``PageLayout()`` on any error.
        """
        try:
            section = self.doc.sections[0]

            width_emu: int = section.page_width or 0
            height_emu: int = section.page_height or 0

            # Determine paper size from width
            if 6_900_000 < width_emu < 8_000_000:
                size: str = "A4"
            elif 8_000_000 < width_emu < 9_000_000:
                size = "Letter"
            elif width_emu > 10_000_000:
                size = "A3"
            else:
                logger.warning(
                    "Unrecognised page width %d EMU for '%s'; defaulting to A4.",
                    width_emu,
                    self.docx_path.name,
                )
                size = "A4"

            orientation = "portrait" if height_emu >= width_emu else "landscape"

            def _cm(emu: Optional[int]) -> float:
                return round((emu or 0) / _EMU_PER_CM, 4)

            # Column count
            columns = 1
            try:
                cols_elems = section._sectPr.findall(qn("w:cols"))  # type: ignore[attr-defined]
                if cols_elems:
                    col_attr = cols_elems[0].get(qn("w:num"))
                    if col_attr and col_attr.isdigit():
                        columns = int(col_attr)
            except Exception:  # noqa: BLE001
                pass

            return PageLayout(
                size=size,  # type: ignore[arg-type]
                orientation=orientation,  # type: ignore[arg-type]
                margin_top_cm=_cm(section.top_margin),
                margin_bottom_cm=_cm(section.bottom_margin),
                margin_left_cm=_cm(section.left_margin),
                margin_right_cm=_cm(section.right_margin),
                columns=columns,
                header_distance_cm=_cm(section.header_distance),
                footer_distance_cm=_cm(section.footer_distance),
            )
        except Exception:  # noqa: BLE001
            raise

    def _extract_styles(self) -> list[StyleSpec]:
        """Extract all paragraph and character styles from the document.

        SOZO-branded and core styles are always included; built-in styles with
        no customisation are skipped unless they belong to CORE_STYLES.

        Returns:
            List of :class:`StyleSpec` objects, SOZO/custom styles first.
        """
        specs: list[StyleSpec] = []
        sozo_specs: list[StyleSpec] = []
        builtin_specs: list[StyleSpec] = []

        for style in self.doc.styles:
            try:
                if style.type not in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
                    continue

                name: str = style.name
                name_lower = name.lower()

                # Decide whether to include this style
                is_core = name in self.CORE_STYLES
                is_sozo = any(marker in name_lower for marker in self.SOZO_STYLE_MARKERS)
                is_heading_style = bool(self.HEADING_STYLE_PATTERN.match(name))

                # Check if it has any font customisation
                has_custom_font = (
                    style.font.name is not None
                    or style.font.size is not None
                    or style.font.bold is not None
                    or style.font.italic is not None
                )

                if not (is_core or is_sozo or is_heading_style or has_custom_font):
                    continue

                # Font properties
                font_name: Optional[str] = style.font.name
                font_size_pt: Optional[float] = None
                try:
                    if style.font.size:
                        font_size_pt = style.font.size.pt
                except Exception:  # noqa: BLE001
                    pass

                bold: Optional[bool] = style.font.bold
                italic: Optional[bool] = style.font.italic
                underline: Optional[bool] = style.font.underline

                color_hex: Optional[str] = None
                try:
                    color_type = style.font.color.type
                    if color_type is not None:
                        color_hex = _emu_color_to_hex(style.font.color.rgb)
                except Exception:  # noqa: BLE001
                    pass

                # Paragraph format (paragraph styles only)
                space_before_pt: Optional[float] = None
                space_after_pt: Optional[float] = None
                alignment: Optional[str] = None
                left_indent_cm: Optional[float] = None
                keep_with_next: bool = False
                page_break_before: bool = False

                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    try:
                        pf = style.paragraph_format
                        if pf.space_before:
                            space_before_pt = pf.space_before.pt
                        if pf.space_after:
                            space_after_pt = pf.space_after.pt
                        if pf.alignment is not None:
                            alignment = _ALIGN_MAP.get(pf.alignment)
                        if pf.left_indent:
                            left_indent_cm = pf.left_indent.cm
                        keep_with_next = bool(pf.keep_with_next)
                        page_break_before = bool(pf.page_break_before)
                    except Exception:  # noqa: BLE001
                        pass

                # Heading detection
                is_heading = False
                heading_level: Optional[int] = None
                m = self.HEADING_STYLE_PATTERN.match(name)
                if m:
                    is_heading = True
                    heading_level = int(m.group(1))

                # Base style
                base_style: Optional[str] = None
                try:
                    if style.base_style:
                        base_style = style.base_style.name
                except Exception:  # noqa: BLE001
                    pass

                style_type_str = (
                    "paragraph"
                    if style.type == WD_STYLE_TYPE.PARAGRAPH
                    else "character"
                )

                spec = StyleSpec(
                    name=name,
                    style_type=style_type_str,  # type: ignore[arg-type]
                    font_name=font_name,
                    font_size_pt=font_size_pt,
                    bold=bold,
                    italic=italic,
                    underline=underline,
                    color_hex=color_hex,
                    space_before_pt=space_before_pt,
                    space_after_pt=space_after_pt,
                    alignment=alignment,  # type: ignore[arg-type]
                    left_indent_cm=left_indent_cm,
                    keep_with_next=keep_with_next,
                    page_break_before=page_break_before,
                    is_heading=is_heading,
                    heading_level=heading_level,
                    base_style=base_style,
                )

                if is_sozo:
                    sozo_specs.append(spec)
                else:
                    builtin_specs.append(spec)

            except Exception as exc:  # noqa: BLE001
                logger.warning("Skipping style '%s' due to error: %s", getattr(style, "name", "?"), exc)

        return sozo_specs + builtin_specs

    def _extract_color_palette(self, styles: list[StyleSpec]) -> list[ColorPaletteEntry]:
        """Build a semantic color palette from extracted styles and table XML.

        Maps colors to roles:
        - Heading 1 → 'primary'
        - Heading 2 → 'secondary'
        - Heading 3+ → 'accent'
        - Normal/body → 'body'
        - Table header rows → 'table_header'
        - Table alternating rows → 'table_stripe_light'

        Args:
            styles: List of extracted :class:`StyleSpec` objects.

        Returns:
            Up to 8 :class:`ColorPaletteEntry` objects deduplicated by hex value.
        """
        palette: list[ColorPaletteEntry] = []
        seen_hex: set[str] = set()

        def _add(role: str, hex_val: Optional[str]) -> None:
            if not hex_val:
                return
            hex_upper = hex_val.upper()
            if hex_upper in seen_hex:
                return
            if len(hex_upper) != 6:
                return
            try:
                int(hex_upper, 16)
            except ValueError:
                return
            seen_hex.add(hex_upper)
            palette.append(ColorPaletteEntry(role=role, hex=hex_upper))

        # Map style colors to palette roles
        style_map: dict[str, StyleSpec] = {s.name.lower(): s for s in styles}

        # Heading roles
        heading_role_map: dict[int, str] = {
            1: "primary",
            2: "secondary",
        }

        for style in styles:
            if not style.is_heading or not style.color_hex:
                continue
            level = style.heading_level or 0
            role = heading_role_map.get(level, f"heading_{level}" if level >= 3 else "accent")
            if level == 3:
                role = "accent"
            _add(role, style.color_hex)

        # Body text color
        for name in ("normal", "body text", "default paragraph font"):
            s = style_map.get(name)
            if s and s.color_hex:
                _add("body", s.color_hex)
                break

        # Table XML colors
        try:
            for table in self.doc.tables:
                rows = table.rows
                if not rows:
                    continue

                # First row → table_header
                try:
                    first_row_cells = rows[0].cells
                    if first_row_cells:
                        fill = _xml_fill_color(first_row_cells[0])
                        _add("table_header", fill)
                except Exception:  # noqa: BLE001
                    pass

                # Second row → table_stripe_light
                try:
                    if len(rows) > 1:
                        second_row_cells = rows[1].cells
                        if second_row_cells:
                            fill = _xml_fill_color(second_row_cells[0])
                            _add("table_stripe_light", fill)
                except Exception:  # noqa: BLE001
                    pass

                if len(palette) >= 8:
                    break
        except Exception as exc:  # noqa: BLE001
            logger.warning("Table color extraction failed: %s", exc)

        return palette[:8]

    def _extract_table_schemas(self) -> list[TableSchemaBlueprint]:
        """Extract structural and visual schemas from the first 5 tables.

        Returns:
            List of unique :class:`TableSchemaBlueprint` objects (deduplicated
            by style name).
        """
        schemas: list[TableSchemaBlueprint] = []
        seen_styles: set[Optional[str]] = set()

        for table in self.doc.tables[:5]:
            try:
                style_name: Optional[str] = None
                try:
                    style_name = table.style.name if table.style else None
                except Exception:  # noqa: BLE001
                    pass

                if style_name in seen_styles:
                    continue
                seen_styles.add(style_name)

                # Column count and proportional widths
                col_count = 0
                col_widths_pct: list[float] = []
                try:
                    cols = table.columns
                    col_count = len(cols)
                    widths = []
                    for col in cols:
                        try:
                            w = col.width or 0
                            widths.append(w)
                        except Exception:  # noqa: BLE001
                            widths.append(0)
                    total_width = sum(widths) or 1
                    col_widths_pct = [round(w / total_width * 100, 2) for w in widths]
                except Exception:  # noqa: BLE001
                    pass

                # Header row attributes
                header_bg: Optional[str] = None
                header_font_color: Optional[str] = None
                row_alt_bg: Optional[str] = None

                rows = table.rows
                if rows:
                    try:
                        first_cells = rows[0].cells
                        if first_cells:
                            header_bg = _xml_fill_color(first_cells[0])
                            # Font color from first run in first cell
                            try:
                                paras = first_cells[0].paragraphs
                                for para in paras:
                                    for run in para.runs:
                                        color_elem = run._r.find(qn("w:rPr"))  # type: ignore[attr-defined]
                                        if color_elem is not None:
                                            c = color_elem.find(qn("w:color"))
                                            if c is not None:
                                                val = c.get(qn("w:val"))
                                                if val and val.upper() not in ("AUTO", ""):
                                                    header_font_color = val.upper()[:6]
                                        if header_font_color:
                                            break
                                    if header_font_color:
                                        break
                            except Exception:  # noqa: BLE001
                                pass
                    except Exception:  # noqa: BLE001
                        pass

                    try:
                        if len(rows) > 1:
                            second_cells = rows[1].cells
                            if second_cells:
                                row_alt_bg = _xml_fill_color(second_cells[0])
                    except Exception:  # noqa: BLE001
                        pass

                schemas.append(
                    TableSchemaBlueprint(
                        style_name=style_name,
                        header_bg_hex=header_bg,
                        header_font_color=header_font_color,
                        row_alt_bg_hex=row_alt_bg,
                        col_count=col_count,
                        typical_col_widths_pct=col_widths_pct,
                    )
                )

            except Exception as exc:  # noqa: BLE001
                logger.warning("Table schema extraction error: %s", exc)

        return schemas

    def _extract_header_footer(
        self,
    ) -> tuple[Optional[HeaderFooterBlueprint], Optional[HeaderFooterBlueprint]]:
        """Extract header and footer layout templates from the first section.

        Parses text from up to 3 paragraphs (mapping to left/center/right
        zones), checks for inline logo images, and records font information.

        Returns:
            A tuple ``(header_blueprint, footer_blueprint)``, each of which may
            be ``None`` if the section has no header or footer.
        """
        try:
            section = self.doc.sections[0]
        except Exception:  # noqa: BLE001
            return None, None

        def _parse_hf(hf_object: object) -> Optional[HeaderFooterBlueprint]:
            try:
                # is_linked_to_previous means it inherits from prior section — skip
                if getattr(hf_object, "is_linked_to_previous", True):
                    return None
                paragraphs = getattr(hf_object, "paragraphs", [])
                if not paragraphs:
                    return None

                # Check for logo images across all paragraphs and runs
                has_logo = False
                try:
                    has_logo = any(
                        run.inline_shapes
                        for para in paragraphs
                        for run in getattr(para, "runs", [])
                    )
                except Exception:  # noqa: BLE001
                    pass

                # Up to 3 paragraphs → left, center, right
                texts: list[str] = []
                for para in paragraphs[:3]:
                    try:
                        texts.append(para.text.strip())
                    except Exception:  # noqa: BLE001
                        texts.append("")

                # Also check if there's an embedded 3-column table for hf layout
                hf_tables = getattr(hf_object, "tables", [])
                if hf_tables:
                    try:
                        first_table = hf_tables[0]
                        if first_table.rows:
                            row_cells = first_table.rows[0].cells
                            if len(row_cells) >= 3:
                                texts = [
                                    row_cells[0].text.strip(),
                                    row_cells[1].text.strip(),
                                    row_cells[2].text.strip(),
                                ]
                    except Exception:  # noqa: BLE001
                        pass

                left_text = texts[0] if len(texts) > 0 else None
                center_text = texts[1] if len(texts) > 1 else None
                right_text = texts[2] if len(texts) > 2 else None

                # Font from first run of first paragraph
                font_name: Optional[str] = None
                font_size_pt: Optional[float] = None
                color_hex: Optional[str] = None

                try:
                    first_para = paragraphs[0]
                    if first_para.runs:
                        run = first_para.runs[0]
                        font_name = run.font.name
                        if run.font.size:
                            font_size_pt = run.font.size.pt
                        try:
                            if run.font.color.type is not None:
                                color_hex = _emu_color_to_hex(run.font.color.rgb)
                        except Exception:  # noqa: BLE001
                            pass
                except Exception:  # noqa: BLE001
                    pass

                logo_position: Optional[str] = None
                if has_logo:
                    # Guess position from which column/paragraph contains the image
                    logo_position = "left"

                return HeaderFooterBlueprint(
                    left_text=left_text or None,
                    center_text=center_text or None,
                    right_text=right_text or None,
                    has_logo=has_logo,
                    logo_position=logo_position,  # type: ignore[arg-type]
                    font_name=font_name,
                    font_size_pt=font_size_pt,
                    color_hex=color_hex,
                )
            except Exception as exc:  # noqa: BLE001
                logger.warning("Header/footer parse error: %s", exc)
                return None

        header_bp = _parse_hf(section.header)
        footer_bp = _parse_hf(section.footer)
        return header_bp, footer_bp

    def _detect_sections(self) -> list[SectionBlueprint]:
        """Walk all paragraphs and detect logical sections from heading styles.

        Also detects cover pages (Title style or large font) and tables of
        contents (TOC styles or matching text).

        Returns:
            Ordered list of :class:`SectionBlueprint` objects.
        """
        sections: list[SectionBlueprint] = []
        order = 0
        paragraphs = self.doc.paragraphs

        has_cover = False
        has_toc = False

        # Check first paragraph for cover detection
        if paragraphs:
            try:
                first_para = paragraphs[0]
                first_style_name = (first_para.style.name or "").lower()
                first_font_size = None
                try:
                    if first_para.runs:
                        fs = first_para.runs[0].font.size
                        if fs:
                            first_font_size = fs.pt
                except Exception:  # noqa: BLE001
                    pass

                if "title" in first_style_name or (
                    first_font_size is not None and first_font_size > 20
                ):
                    has_cover = True
                    sections.append(
                        SectionBlueprint(
                            section_id=str(uuid.uuid4())[:8],
                            section_type="cover",
                            heading_text=first_para.text.strip() or "Cover",
                            heading_style=first_para.style.name,
                            heading_level=1,
                            order=order,
                            layout_regions=[
                                LayoutRegion(
                                    region_id=f"cover_text_{order}",
                                    region_type="text_block",
                                    content_key="cover_content",
                                    order=0,
                                )
                            ],
                            page_break_before=False,
                        )
                    )
                    order += 1
            except Exception as exc:  # noqa: BLE001
                logger.warning("Cover page detection error: %s", exc)

        # TOC detection across all paragraphs
        for i, para in enumerate(paragraphs):
            try:
                style_name_lower = (para.style.name or "").lower()
                text_lower = para.text.strip().lower()

                if "toc" in style_name_lower or "table of contents" in text_lower or "contents" == text_lower:
                    if not has_toc:
                        has_toc = True
                        sections.append(
                            SectionBlueprint(
                                section_id=str(uuid.uuid4())[:8],
                                section_type="toc",
                                heading_text="Table of Contents",
                                heading_style=para.style.name,
                                heading_level=1,
                                order=order,
                                layout_regions=[
                                    LayoutRegion(
                                        region_id=f"toc_block_{order}",
                                        region_type="text_block",
                                        content_key="toc_content",
                                        order=0,
                                    )
                                ],
                            )
                        )
                        order += 1
            except Exception:  # noqa: BLE001
                pass

        # Walk all paragraphs for heading-based sections
        for i, para in enumerate(paragraphs):
            try:
                style_name = para.style.name or ""
                style_name_lower = style_name.lower()

                # Match heading pattern
                m = self.HEADING_STYLE_PATTERN.match(style_name)
                if not m:
                    continue

                heading_level = int(m.group(1))
                heading_text = para.text.strip()
                if not heading_text:
                    continue

                section_type = _guess_section_type(heading_text)

                # Build layout regions
                regions: list[LayoutRegion] = [
                    LayoutRegion(
                        region_id=f"text_{order}",
                        region_type="text_block",
                        content_key=f"section_{order}_content",
                        order=0,
                    )
                ]

                # Check if next paragraph is a table
                try:
                    next_para = paragraphs[i + 1] if i + 1 < len(paragraphs) else None
                    if next_para is not None:
                        # Check if the very next element is a table by inspecting siblings
                        para_elem = para._element  # type: ignore[attr-defined]
                        next_sibling = para_elem.getnext()
                        if next_sibling is not None:
                            tag = etree.QName(next_sibling.tag).localname if next_sibling.tag else ""
                            if tag == "tbl":
                                regions.append(
                                    LayoutRegion(
                                        region_id=f"table_{order}",
                                        region_type="table",
                                        content_key=f"section_{order}_table",
                                        order=1,
                                    )
                                )
                except Exception:  # noqa: BLE001
                    pass

                sections.append(
                    SectionBlueprint(
                        section_id=str(uuid.uuid4())[:8],
                        section_type=section_type,  # type: ignore[arg-type]
                        heading_text=heading_text,
                        heading_style=style_name,
                        heading_level=heading_level,
                        order=order,
                        layout_regions=regions,
                        page_break_before=(heading_level == 1),
                    )
                )
                order += 1

            except Exception as exc:  # noqa: BLE001
                logger.warning("Section detection error at paragraph %d: %s", i, exc)

        return sections

    def _detect_document_type(
        self, sections: list[SectionBlueprint]
    ) -> Optional[str]:
        """Infer the document type from the distribution of section types.

        Args:
            sections: Detected :class:`SectionBlueprint` list.

        Returns:
            One of ``'protocol'``, ``'assessment'``, ``'handbook'``,
            ``'unknown'``, or ``None`` if no sections are present.
        """
        if not sections:
            return None

        type_counts: dict[str, int] = {}
        for sec in sections:
            type_counts[sec.section_type] = type_counts.get(sec.section_type, 0) + 1

        protocol_score = (
            type_counts.get("protocol_parameters", 0)
            + type_counts.get("session_plan", 0)
            + type_counts.get("evidence_summary", 0)
        )
        assessment_score = (
            type_counts.get("assessment_tools", 0)
            + type_counts.get("scoring_guide", 0)
        )
        handbook_score = (
            type_counts.get("handbook_guidance", 0)
            + type_counts.get("overview", 0)
        )

        max_score = max(protocol_score, assessment_score, handbook_score)
        if max_score == 0:
            return "unknown"

        if protocol_score == max_score:
            return "protocol"
        if assessment_score == max_score:
            return "assessment"
        if handbook_score == max_score:
            return "handbook"
        return "unknown"


# ---------------------------------------------------------------------------
# Section type guesser (module-level for reuse)
# ---------------------------------------------------------------------------

_SECTION_TYPE_KEYWORDS: list[tuple[frozenset[str], str]] = [
    (frozenset({"overview", "introduction", "background"}), "overview"),
    (frozenset({"protocol", "parameters", "stimulation"}), "protocol_parameters"),
    (frozenset({"session", "schedule", "programme", "program"}), "session_plan"),
    (frozenset({"evidence", "literature", "research"}), "evidence_summary"),
    (frozenset({"device", "equipment", "specifications", "specification"}), "device_specs"),
    (frozenset({"contraindication", "safety", "exclusion", "precaution"}), "contraindications"),
    (frozenset({"assessment", "scale", "measure", "scoring", "score"}), "assessment_tools"),
    (frozenset({"reference", "bibliography", "references"}), "reference_list"),
    (frozenset({"handbook", "guide", "guidance"}), "handbook_guidance"),
    (frozenset({"appendix", "appendices"}), "appendix"),
]


def _guess_section_type(heading_text: str) -> str:
    """Guess the semantic section type from a heading string.

    Performs a simple word-level keyword match against known section types.

    Args:
        heading_text: The heading text to classify.

    Returns:
        One of the :class:`SectionBlueprint` ``section_type`` literals, or
        ``'custom'`` if no keyword matches.
    """
    words = set(re.sub(r"[^\w\s]", "", heading_text).lower().split())
    for keywords, section_type in _SECTION_TYPE_KEYWORDS:
        if words & keywords:
            return section_type
    return "custom"


# ---------------------------------------------------------------------------
# Convenience function
# ---------------------------------------------------------------------------


def extract_blueprint(docx_path: Path | str) -> DocumentBlueprint:
    """Extract a DocumentBlueprint from a DOCX file.

    Convenience wrapper around :class:`BlueprintExtractor`.

    Args:
        docx_path: Path to the ``.docx`` file to analyse.

    Returns:
        :class:`DocumentBlueprint`: the extracted visual DNA of the document.

    Raises:
        FileNotFoundError: if the file does not exist.
        ImportError: if python-docx is not installed.
    """
    return BlueprintExtractor(docx_path).extract()
