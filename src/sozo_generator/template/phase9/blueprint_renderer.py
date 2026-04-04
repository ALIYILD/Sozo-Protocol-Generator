"""
blueprint_renderer.py — Assembles a styled DOCX from a DocumentBlueprint + DocumentSpec.

The BlueprintRenderer bridges the template DNA (DocumentBlueprint) with the
generated content (DocumentSpec) to produce a visually consistent, fully
assembled DOCX document.

Phase 9, pipeline_version: phase9_v1
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor

from .blueprint_schema import DocumentBlueprint, SectionBlueprint, LayoutRegion
from .style_cloner import StyleCloner
from ...schemas.documents import DocumentSpec, SectionContent

logger = logging.getLogger(__name__)


class BlueprintRenderer:
    """
    Renders a DocumentSpec into a styled DOCX Document using a DocumentBlueprint
    as the visual template.

    The renderer:
    1. Creates a new empty Document
    2. Applies all styles from the blueprint (via StyleCloner)
    3. Renders sections in blueprint order, matching section_type to DocumentSpec sections
    4. For unmatched sections: uses the blueprint heading style + generic body text style
    5. Returns the fully assembled Document ready for save/export

    Usage:
        blueprint = extract_blueprint("template.docx")
        renderer = BlueprintRenderer(blueprint)
        doc = renderer.render(document_spec)
        doc.save("output.docx")
    """

    def __init__(self, blueprint: DocumentBlueprint) -> None:
        self.blueprint = blueprint

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def render(
        self,
        spec: DocumentSpec,
        output_path: Optional[Path] = None,
    ) -> Document:
        """Render a DocumentSpec into a styled DOCX using the blueprint.

        Steps:
        1. Creates a fresh Document.
        2. Injects all blueprint styles via StyleCloner.
        3. Renders cover page if blueprint.has_cover_page.
        4. Inserts a TOC placeholder if blueprint.has_toc.
        5. Renders each SectionBlueprint in order, matching against DocumentSpec sections.
        6. Optionally saves to output_path.

        Args:
            spec: The :class:`DocumentSpec` carrying all content to render.
            output_path: If provided, the rendered document is saved here.

        Returns:
            The assembled python-docx :class:`Document`.
        """
        doc = Document()
        StyleCloner(self.blueprint).apply(doc)

        # Cover page
        if self.blueprint.has_cover_page:
            self._render_cover(doc, spec)

        # TOC placeholder
        if self.blueprint.has_toc:
            try:
                doc.add_paragraph("Table of Contents", style="TOC Heading")
            except Exception:  # noqa: BLE001
                doc.add_paragraph("Table of Contents")

        # Render each section in blueprint order
        _skip_types = {"cover", "toc"}
        for section_bp in self.blueprint.ordered_sections():
            if section_bp.section_type in _skip_types:
                continue
            content = self._match_spec_section(section_bp, spec)
            self._render_section(doc, section_bp, content)

        if output_path is not None:
            doc.save(str(output_path))

        logger.info(
            "Rendered %s — %s → %d sections",
            spec.condition_name,
            spec.document_type.value,
            len(self.blueprint.sections),
        )
        return doc

    # ------------------------------------------------------------------
    # Rendering helpers
    # ------------------------------------------------------------------

    def _render_cover(self, doc: Document, spec: DocumentSpec) -> None:
        """Render the cover page block (title, subtitle, date, confidentiality mark).

        Adds four paragraphs using blueprint heading/normal styles:
        title, subtitle (if present), date label, and confidentiality mark.
        """
        title_style = self._resolve_style(
            self.blueprint.heading_styles()[0].name if self.blueprint.heading_styles() else None,
            fallback="Heading 1",
        )
        self._add_paragraph(doc, spec.title, style_name=title_style)

        if spec.subtitle:
            self._add_paragraph(doc, spec.subtitle, style_name=self._resolve_style(None, "Normal"))

        if spec.date_label:
            self._add_paragraph(doc, spec.date_label, style_name=self._resolve_style(None, "Normal"))

        if spec.confidentiality_mark:
            self._add_paragraph(
                doc,
                spec.confidentiality_mark,
                style_name=self._resolve_style(None, "Normal"),
            )

        doc.add_page_break()

    def _render_section(
        self,
        doc: Document,
        section_bp: SectionBlueprint,
        content: Optional[SectionContent],
    ) -> None:
        """Render one section: heading + content blocks.

        Adds the section heading, then either a placeholder paragraph (when
        no matching content is found) or the full content body.  Evidence
        callout is added for 'evidence_summary' sections.  A page break is
        inserted before the heading when section_bp.page_break_before is set.

        Args:
            doc: The Document to write into.
            section_bp: The blueprint definition for this section.
            content: Matching :class:`SectionContent`, or None if unmatched.
        """
        if section_bp.page_break_before:
            doc.add_page_break()

        self._add_heading(
            doc,
            section_bp.heading_text,
            section_bp.heading_level,
            section_bp.heading_style,
        )

        if content is None:
            self._add_paragraph(
                doc,
                f"[Section: {section_bp.heading_text} — content pending]",
                style_name="Normal",
            )
            return

        # Body text
        body_style = self._resolve_style(None, "Normal")
        self._render_text_content(doc, content, body_style)

        # Tables
        if content.tables:
            # Find a table layout region to pass; use the first if available
            table_region: Optional[LayoutRegion] = None
            for region in section_bp.layout_regions:
                if region.region_type == "table":
                    table_region = region
                    break
            if table_region is None and section_bp.layout_regions:
                table_region = section_bp.layout_regions[0]
            self._render_table_content(doc, content.tables, region=table_region)

        # Evidence callout
        if section_bp.section_type == "evidence_summary":
            self._render_evidence_callout(doc, content)

    def _render_text_content(
        self,
        doc: Document,
        content: SectionContent,
        style_name: str,
    ) -> None:
        """Add body text paragraphs from section content.

        Splits the content string on double newlines (paragraph breaks).
        Paragraphs beginning with '• ' or '- ' are rendered with the
        'List Bullet' style when that style is available.

        Args:
            doc: The Document to write into.
            content: The :class:`SectionContent` whose ``content`` field to render.
            style_name: The default paragraph style to apply.
        """
        if not content.content:
            return

        has_list_bullet = self.blueprint.has_style("List Bullet")
        paragraphs = content.content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith("• ") or para_text.startswith("- "):
                effective_style = "List Bullet" if has_list_bullet else style_name
            else:
                effective_style = style_name
            self._add_paragraph(doc, para_text, style_name=effective_style)

    def _render_table_content(
        self,
        doc: Document,
        tables: list[dict],
        region: Optional[LayoutRegion],
    ) -> None:
        """Add tables from a list of table dicts.

        Each dict must have the keys ``"headers"`` (list of str) and
        ``"rows"`` (list of lists of str).  The blueprint's first table
        schema style is applied when available.

        Args:
            doc: The Document to write into.
            tables: List of table dicts with ``headers`` and ``rows``.
            region: The :class:`LayoutRegion` hint (may be None).
        """
        # Resolve table style
        table_style = "Table Grid"
        if self.blueprint.table_schemas:
            candidate = self.blueprint.table_schemas[0].style_name
            if candidate:
                table_style = candidate

        for table_dict in tables:
            headers: list = table_dict.get("headers", [])
            rows: list[list] = table_dict.get("rows", [])
            if not headers:
                continue

            col_count = len(headers)
            row_count = 1 + len(rows)

            try:
                table = doc.add_table(rows=row_count, cols=col_count)
                try:
                    table.style = table_style
                except Exception:  # noqa: BLE001
                    pass

                # Header row
                hdr_row = table.rows[0]
                for idx, header_text in enumerate(headers):
                    if idx < len(hdr_row.cells):
                        hdr_row.cells[idx].text = str(header_text)

                # Data rows
                for row_idx, row_data in enumerate(rows):
                    docx_row = table.rows[row_idx + 1]
                    for col_idx, cell_value in enumerate(row_data):
                        if col_idx < len(docx_row.cells):
                            docx_row.cells[col_idx].text = str(cell_value)

            except Exception as exc:  # noqa: BLE001
                logger.warning("Table render error: %s", exc)

    def _render_evidence_callout(
        self,
        doc: Document,
        content: SectionContent,
    ) -> None:
        """Add a styled callout box for evidence citations.

        Uses the 'SOZO Callout' style when present in the blueprint,
        otherwise falls back to 'Normal'.  Displays PMID count and
        confidence label.

        Args:
            doc: The Document to write into.
            content: The :class:`SectionContent` carrying evidence metadata.
        """
        callout_style = self._resolve_style("SOZO Callout", fallback="Normal")
        pmid_count = len(content.evidence_pmids) if content.evidence_pmids else 0
        confidence = content.confidence_label or "unrated"
        callout_text = (
            f"Evidence: {pmid_count} citations | Confidence: {confidence}"
        )
        self._add_paragraph(doc, callout_text, style_name=callout_style)

    # ------------------------------------------------------------------
    # Low-level paragraph/heading helpers
    # ------------------------------------------------------------------

    def _add_heading(
        self,
        doc: Document,
        text: str,
        level: int,
        style_name: Optional[str],
    ) -> None:
        """Add a heading paragraph with blueprint style.

        Prefers the named style from the blueprint when it exists; otherwise
        falls back to python-docx's built-in heading at the specified level.

        Args:
            doc: The Document to write into.
            text: The heading text.
            level: Heading level (1–6).
            style_name: Blueprint style name to try first.
        """
        if style_name and self.blueprint.has_style(style_name):
            try:
                p = doc.add_paragraph(text)
                p.style = doc.styles[style_name]
                return
            except Exception:  # noqa: BLE001
                pass
        # Fallback to built-in heading
        try:
            doc.add_heading(text, level=level)
        except Exception:  # noqa: BLE001
            doc.add_paragraph(text)

    def _add_paragraph(
        self,
        doc: Document,
        text: str,
        style_name: Optional[str] = "Normal",
    ) -> None:
        """Add a body text paragraph.

        Applies the named style if it exists in the document styles; silently
        falls back to python-docx default on failure.

        Args:
            doc: The Document to write into.
            text: The paragraph text.
            style_name: The Word style name to apply.
        """
        try:
            if style_name:
                doc.add_paragraph(text, style=style_name)
            else:
                doc.add_paragraph(text)
        except Exception:  # noqa: BLE001
            try:
                doc.add_paragraph(text)
            except Exception:  # noqa: BLE001
                logger.warning("Failed to add paragraph: %.80s", text)

    # ------------------------------------------------------------------
    # Matching and resolution helpers
    # ------------------------------------------------------------------

    def _match_spec_section(
        self,
        section_bp: SectionBlueprint,
        spec: DocumentSpec,
    ) -> Optional[SectionContent]:
        """Find the DocumentSpec section that best matches this blueprint section.

        Matching priority:
        1. Exact section_id match.
        2. Case-insensitive substring match between blueprint heading_text
           and section title.

        Args:
            section_bp: The blueprint section to match against.
            spec: The :class:`DocumentSpec` holding all content sections.

        Returns:
            The matching :class:`SectionContent`, or ``None`` if not found.
        """
        # 1. Exact section_id match
        for s in spec.sections:
            if s.section_id == section_bp.section_id:
                return s

        # 2. Case-insensitive substring match
        bp_heading_lower = section_bp.heading_text.lower()
        for s in spec.sections:
            if bp_heading_lower in s.title.lower() or s.title.lower() in bp_heading_lower:
                return s

        return None

    def _resolve_style(
        self,
        preferred: Optional[str],
        fallback: str = "Normal",
    ) -> str:
        """Return preferred style if it exists in the blueprint, else fallback.

        Args:
            preferred: The preferred style name to check.
            fallback: The style to use when preferred is absent or None.

        Returns:
            A style name string guaranteed to be non-empty.
        """
        if preferred and self.blueprint.has_style(preferred):
            return preferred
        return fallback


# ---------------------------------------------------------------------------
# Module-level convenience function
# ---------------------------------------------------------------------------


def render_document(
    blueprint: DocumentBlueprint,
    spec: DocumentSpec,
    output_path: Optional[Path] = None,
) -> Document:
    """Render a DocumentSpec to a styled DOCX using the given blueprint.

    Convenience wrapper around :class:`BlueprintRenderer`.

    Args:
        blueprint: The :class:`DocumentBlueprint` providing visual identity.
        spec: The :class:`DocumentSpec` carrying all content to render.
        output_path: If provided, the document is saved to this path.

    Returns:
        The assembled python-docx :class:`Document`.
    """
    return BlueprintRenderer(blueprint).render(spec, output_path=output_path)
