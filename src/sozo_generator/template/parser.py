"""
Template parser — extracts structure from Gold Standard .docx template files.
Identifies section headers, table locations, placeholder fields, and review flags.
"""
import logging
import re
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Placeholder detection patterns
PLACEHOLDER_PATTERNS = [
    re.compile(r"\[PLACEHOLDER[^\]]*\]", re.IGNORECASE),
    re.compile(r"\[TODO[^\]]*\]", re.IGNORECASE),
    re.compile(r"\[REVIEW[^\]]*\]", re.IGNORECASE),
    re.compile(r"\[INSERT[^\]]*\]", re.IGNORECASE),
    re.compile(r"\[TBD[^\]]*\]", re.IGNORECASE),
    re.compile(r"\[CONDITION NAME\]", re.IGNORECASE),
    re.compile(r"\[PATIENT NAME\]", re.IGNORECASE),
    re.compile(r"\[DATE\]", re.IGNORECASE),
]

# Section header detection: heading-style text in ALL CAPS or Title Case followed by colon or newline
SECTION_HEADER_PATTERN = re.compile(
    r"^(?:SECTION\s+\d+[:\s]|STAGE\s+\d+[:\s]|STEP\s+\d+[:\s]|"
    r"[A-Z][A-Z\s/&\-]{4,}(?::|$))",
    re.MULTILINE,
)


class ParsedSection:
    """Represents a parsed section from a template document."""

    def __init__(
        self,
        section_id: str,
        title: str,
        raw_content: str = "",
        placeholder_count: int = 0,
        has_table: bool = False,
        page_number: Optional[int] = None,
        heading_level: int = 1,
    ):
        self.section_id = section_id
        self.title = title
        self.raw_content = raw_content
        self.placeholder_count = placeholder_count
        self.has_table = has_table
        self.page_number = page_number
        self.heading_level = heading_level
        self.placeholders: list[str] = []

    def __repr__(self) -> str:
        return f"ParsedSection(id={self.section_id!r}, title={self.title!r}, placeholders={self.placeholder_count})"


class TemplateParser:
    """
    Parses Gold Standard .docx template files to extract structure,
    sections, and placeholder locations.
    """

    def __init__(self, template_path: Path | str):
        self.template_path = Path(template_path)
        self._sections: list[ParsedSection] = []
        self._raw_text: str = ""
        self._parsed: bool = False

    def parse(self) -> list[ParsedSection]:
        """
        Parse the template file and extract sections.
        Returns list of ParsedSection objects.
        """
        if not self.template_path.exists():
            logger.warning(f"Template not found: {self.template_path}")
            return []

        if self.template_path.suffix.lower() == ".docx":
            return self._parse_docx()
        elif self.template_path.suffix.lower() == ".txt":
            return self._parse_text()
        else:
            logger.warning(f"Unsupported template format: {self.template_path.suffix}")
            return []

    def _parse_docx(self) -> list[ParsedSection]:
        """Parse a .docx template file using python-docx."""
        try:
            import docx
        except ImportError:
            logger.error("python-docx not installed — cannot parse .docx templates")
            return []

        try:
            doc = docx.Document(self.template_path)
            sections = []
            current_section: Optional[ParsedSection] = None
            current_content_lines: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect heading style
                if para.style.name.startswith("Heading"):
                    # Save previous section
                    if current_section is not None:
                        current_section.raw_content = "\n".join(current_content_lines)
                        current_section.placeholder_count = self._count_placeholders(current_section.raw_content)
                        current_section.placeholders = self._find_placeholders(current_section.raw_content)
                        sections.append(current_section)

                    heading_level = int(para.style.name.replace("Heading ", "") or "1")
                    section_id = self._title_to_id(text)
                    current_section = ParsedSection(
                        section_id=section_id,
                        title=text,
                        heading_level=heading_level,
                    )
                    current_content_lines = []
                else:
                    current_content_lines.append(text)

            # Flush last section
            if current_section is not None:
                current_section.raw_content = "\n".join(current_content_lines)
                current_section.placeholder_count = self._count_placeholders(current_section.raw_content)
                current_section.placeholders = self._find_placeholders(current_section.raw_content)
                sections.append(current_section)

            # Check tables
            for table in doc.tables:
                table_text = " ".join(cell.text for row in table.rows for cell in row.cells)
                for section in sections:
                    if any(word in table_text for word in section.title.split()[:3] if len(word) > 3):
                        section.has_table = True
                        break

            self._sections = sections
            self._parsed = True
            logger.info(f"Parsed {len(sections)} sections from {self.template_path.name}")
            return sections

        except Exception as e:
            logger.error(f"Error parsing template {self.template_path}: {e}")
            return []

    def _parse_text(self) -> list[ParsedSection]:
        """Parse a plain text template file."""
        with open(self.template_path, "r", encoding="utf-8") as f:
            self._raw_text = f.read()

        sections = []
        lines = self._raw_text.split("\n")
        current_section: Optional[ParsedSection] = None
        current_content_lines: list[str] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current_content_lines:
                    current_content_lines.append("")
                continue

            # Detect section headers (lines ending with colon or ALL CAPS)
            if self._is_section_header(stripped):
                if current_section is not None:
                    current_section.raw_content = "\n".join(current_content_lines).strip()
                    current_section.placeholder_count = self._count_placeholders(current_section.raw_content)
                    sections.append(current_section)

                title = stripped.rstrip(":")
                section_id = self._title_to_id(title)
                current_section = ParsedSection(section_id=section_id, title=title)
                current_content_lines = []
            else:
                current_content_lines.append(stripped)

        if current_section is not None:
            current_section.raw_content = "\n".join(current_content_lines).strip()
            current_section.placeholder_count = self._count_placeholders(current_section.raw_content)
            sections.append(current_section)

        self._sections = sections
        self._parsed = True
        return sections

    def _is_section_header(self, text: str) -> bool:
        """Heuristic: section header detection."""
        if text.endswith(":") and len(text) < 80:
            return True
        if text.isupper() and len(text.split()) >= 2:
            return True
        return False

    def _title_to_id(self, title: str) -> str:
        """Convert title text to a section_id slug."""
        import re
        slug = re.sub(r"[^\w\s]", "", title.lower())
        slug = re.sub(r"\s+", "_", slug.strip())
        return slug[:60]

    def _count_placeholders(self, text: str) -> int:
        """Count placeholder occurrences in text."""
        count = 0
        for pattern in PLACEHOLDER_PATTERNS:
            count += len(pattern.findall(text))
        return count

    def _find_placeholders(self, text: str) -> list[str]:
        """Find all placeholder strings in text."""
        found = []
        for pattern in PLACEHOLDER_PATTERNS:
            found.extend(pattern.findall(text))
        return found

    def get_section(self, section_id: str) -> Optional[ParsedSection]:
        """Get a specific section by ID."""
        for section in self._sections:
            if section.section_id == section_id:
                return section
        return None

    def get_all_placeholders(self) -> dict[str, list[str]]:
        """Return all placeholders indexed by section_id."""
        return {
            s.section_id: s.placeholders
            for s in self._sections
            if s.placeholders
        }

    def get_section_ids(self) -> list[str]:
        """Return all section IDs in document order."""
        return [s.section_id for s in self._sections]

    @property
    def sections(self) -> list[ParsedSection]:
        return self._sections

    @property
    def total_placeholders(self) -> int:
        return sum(s.placeholder_count for s in self._sections)
