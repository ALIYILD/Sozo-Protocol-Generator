"""
Document ingester — reads existing SOZO DOCX files and extracts a structured
fingerprint of their content, formatting, and layout patterns.

This is NOT template parsing. This ingests ALREADY GENERATED documents
to learn what the system has produced, so future generation stays consistent.
"""
from __future__ import annotations

import hashlib
import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class TableFingerprint:
    """Structural fingerprint of one table."""
    index: int = 0
    headers: list[str] = field(default_factory=list)
    row_count: int = 0
    col_count: int = 0
    caption: str = ""
    parent_section: str = ""


@dataclass
class SectionFingerprint:
    """Structural fingerprint of one section in a document."""
    title: str = ""
    section_id: str = ""
    level: int = 1  # heading level
    word_count: int = 0
    has_table: bool = False
    has_bullets: bool = False
    has_numbered_list: bool = False
    is_boilerplate: bool = False  # same across all conditions
    content_hash: str = ""  # hash of text content
    paragraph_count: int = 0


@dataclass
class DocumentFingerprint:
    """Complete structural fingerprint of one SOZO document."""
    file_path: str = ""
    file_name: str = ""
    condition_slug: str = ""
    doc_type: str = ""
    tier: str = ""
    # Title block
    title_block: list[str] = field(default_factory=list)
    title: str = ""
    # Structure
    sections: list[SectionFingerprint] = field(default_factory=list)
    tables: list[TableFingerprint] = field(default_factory=list)
    total_paragraphs: int = 0
    total_tables: int = 0
    total_words: int = 0
    # Formatting
    fonts_used: list[str] = field(default_factory=list)
    styles_used: list[str] = field(default_factory=list)
    font_sizes: list[float] = field(default_factory=list)
    # Content
    boilerplate_sections: list[str] = field(default_factory=list)
    variable_sections: list[str] = field(default_factory=list)
    references_count: int = 0
    content_hash: str = ""

    def to_dict(self) -> dict:
        import dataclasses
        return dataclasses.asdict(self)


# ── Boilerplate patterns — text that is identical across conditions ──────

_BOILERPLATE_MARKERS = [
    "sozo brain center",
    "evidence-based neuromodulation",
    "governance rule",
    "© 2026 sozo brain center",
    "confidential",
    "for use by sozo fellow",
    "for use by sozo partner",
    "for authorized sozo",
    "this document is for authorized",
    "no fellow or clinical assistant may independently",
    "doctor authorization",
]

_TITLE_BLOCK_MARKERS = [
    "sozo brain center",
    "evidence-based neuromodulation",
    "fellow tier",
    "partners tier",
    "confidential",
    "version",
]


def ingest_document(file_path: Path) -> DocumentFingerprint:
    """Ingest a single SOZO DOCX and return its structural fingerprint."""
    import docx

    file_path = Path(file_path)
    doc = docx.Document(str(file_path))

    fp = DocumentFingerprint(
        file_path=str(file_path),
        file_name=file_path.name,
    )

    # Infer condition, doc_type, tier from path
    fp.condition_slug, fp.doc_type, fp.tier = _infer_metadata_from_path(file_path)

    fp.total_paragraphs = len(doc.paragraphs)
    fp.total_tables = len(doc.tables)

    # ── Extract title block (first ~8 non-empty paragraphs) ──
    title_paras = []
    for p in doc.paragraphs[:12]:
        t = p.text.strip()
        if t:
            title_paras.append(t)
        if len(title_paras) >= 8:
            break
    fp.title_block = title_paras

    # Find the actual document title (longest line in title block, or the one with —)
    for t in title_paras:
        if "—" in t or "–" in t:
            fp.title = t
            break
    if not fp.title and len(title_paras) >= 4:
        fp.title = title_paras[3]  # usually line 4

    # ── Extract sections by detecting bold headings ──
    current_section = None
    current_words = []
    current_para_count = 0
    all_words = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        all_words += len(text.split())

        # Detect section heading (bold text, not too long)
        is_heading = (
            any(r.bold for r in p.runs if r.bold)
            and len(text) < 100
            and text not in fp.title_block[:4]
        )

        if is_heading:
            # Save previous section
            if current_section:
                content = " ".join(current_words)
                current_section.word_count = len(current_words)
                current_section.paragraph_count = current_para_count
                current_section.content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
                current_section.has_bullets = any(
                    w.startswith(("•", "- ", "* ")) for w in current_words
                )
                current_section.is_boilerplate = _is_boilerplate(content)
                fp.sections.append(current_section)

            current_section = SectionFingerprint(
                title=text,
                section_id=_slugify(text),
            )
            current_words = []
            current_para_count = 0
        else:
            current_words.append(text)
            current_para_count += 1

    # Flush last section
    if current_section:
        content = " ".join(current_words)
        current_section.word_count = len(current_words)
        current_section.paragraph_count = current_para_count
        current_section.content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
        current_section.is_boilerplate = _is_boilerplate(content)
        fp.sections.append(current_section)

    fp.total_words = all_words

    # Classify sections
    for s in fp.sections:
        if s.is_boilerplate:
            fp.boilerplate_sections.append(s.section_id)
        else:
            fp.variable_sections.append(s.section_id)

    # ── Extract table fingerprints ──
    for i, table in enumerate(doc.tables):
        tf = TableFingerprint(index=i)
        if table.rows:
            tf.headers = [cell.text.strip()[:40] for cell in table.rows[0].cells]
            tf.row_count = len(table.rows) - 1
            tf.col_count = len(table.rows[0].cells)
        fp.tables.append(tf)

    # ── Extract formatting ──
    styles = set()
    fonts = set()
    sizes = set()
    for p in doc.paragraphs:
        if p.style:
            styles.add(p.style.name)
        for run in p.runs:
            if run.font.name:
                fonts.add(run.font.name)
            if run.font.size:
                sizes.add(run.font.size.pt)

    fp.styles_used = sorted(styles)
    fp.fonts_used = sorted(fonts)
    fp.font_sizes = sorted(sizes)

    # Content hash of full document
    full_text = " ".join(p.text for p in doc.paragraphs)
    fp.content_hash = hashlib.sha256(full_text.encode()).hexdigest()[:16]

    logger.info("Ingested %s: %d sections, %d tables, %d words",
                file_path.name, len(fp.sections), fp.total_tables, fp.total_words)

    return fp


def ingest_directory(directory: Path) -> list[DocumentFingerprint]:
    """Ingest all DOCX files in a directory tree."""
    directory = Path(directory)
    results = []
    for docx_file in sorted(directory.rglob("*.docx")):
        try:
            fp = ingest_document(docx_file)
            results.append(fp)
        except Exception as e:
            logger.warning("Failed to ingest %s: %s", docx_file, e)
    logger.info("Ingested %d documents from %s", len(results), directory)
    return results


def save_fingerprints(fingerprints: list[DocumentFingerprint], output_path: Path) -> None:
    """Save fingerprints to JSON."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    data = [fp.to_dict() for fp in fingerprints]
    output_path.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
    logger.info("Saved %d fingerprints to %s", len(data), output_path)


def load_fingerprints(input_path: Path) -> list[DocumentFingerprint]:
    """Load fingerprints from JSON."""
    data = json.loads(Path(input_path).read_text(encoding="utf-8"))
    results = []
    for d in data:
        fp = DocumentFingerprint(**{
            k: v for k, v in d.items()
            if k in DocumentFingerprint.__dataclass_fields__
        })
        results.append(fp)
    return results


# ── Helpers ──────────────────────────────────────────────────────────────

def _slugify(text: str) -> str:
    slug = re.sub(r"[^\w\s]", "", text.lower())
    return re.sub(r"\s+", "_", slug.strip())[:60]


def _is_boilerplate(text: str) -> bool:
    text_lower = text.lower()
    return any(marker in text_lower for marker in _BOILERPLATE_MARKERS)


def _infer_metadata_from_path(path: Path) -> tuple[str, str, str]:
    """Infer condition_slug, doc_type, tier from file path."""
    parts = path.parts
    name = path.stem.lower()

    # Tier from folder or filename
    tier = "unknown"
    for p in parts:
        pl = p.lower()
        if pl == "fellow":
            tier = "fellow"
        elif pl == "partners":
            tier = "partners"
    if tier == "unknown":
        if "fellow" in name:
            tier = "fellow"
        elif "partners" in name:
            tier = "partners"

    # Condition from parent folder
    condition = "unknown"
    for p in parts:
        pl = p.lower().replace("_", "").replace(" ", "")
        # Check known condition slugs
        for slug in ["parkinsons", "depression", "anxiety", "adhd", "alzheimers",
                      "strokerehab", "tbi", "chronicpain", "ptsd", "ocd",
                      "ms", "asd", "longcovid", "tinnitus", "insomnia"]:
            if slug in pl:
                condition = slug.replace("strokerehab", "stroke_rehab").replace(
                    "chronicpain", "chronic_pain").replace("longcovid", "long_covid")
                break
        if condition != "unknown":
            break

    # Doc type from filename
    doc_type = "unknown"
    doc_type_patterns = {
        "evidence_based_protocol": ["evidence_based", "evidence based", "ebp"],
        "all_in_one_protocol": ["all_in_one", "all in one", "aio"],
        "handbook": ["handbook", "clinical_handbook"],
        "clinical_exam": ["clinical_examination", "exam_checklist", "clinical_exam"],
        "phenotype_classification": ["phenotype"],
        "responder_tracking": ["responder"],
        "psych_intake": ["psychological_intake", "psych_intake", "prs_baseline"],
        "network_assessment": ["6network", "bedside_assessment", "network_assessment"],
    }
    for dt, patterns in doc_type_patterns.items():
        if any(p in name for p in patterns):
            doc_type = dt
            break

    return condition, doc_type, tier
