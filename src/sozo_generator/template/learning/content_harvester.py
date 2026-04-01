"""
Content harvester — extracts rich section content from existing SOZO documents
and stores it as reusable content blocks keyed by (doc_type, section_id).

This allows new document generation to use the actual clinical prose from
existing documents rather than generating skeletal content.

SAFETY: Harvested content comes from EXISTING clinician-reviewed documents.
The system adapts condition-specific details but preserves the prose quality.
"""
from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import docx

logger = logging.getLogger(__name__)


@dataclass
class HarvestedSection:
    """A rich section extracted from an existing document."""
    section_id: str = ""
    title: str = ""
    content: str = ""  # full prose text
    word_count: int = 0
    has_table: bool = False
    table_data: list[dict] = field(default_factory=list)  # [{headers, rows, caption}]
    source_file: str = ""
    source_condition: str = ""
    source_doc_type: str = ""
    source_tier: str = ""
    # Which parts are condition-specific (need replacement)
    condition_specific_phrases: list[str] = field(default_factory=list)


@dataclass
class HarvestedDocument:
    """All harvested content from one document."""
    file_path: str = ""
    condition_slug: str = ""
    doc_type: str = ""
    tier: str = ""
    sections: list[HarvestedSection] = field(default_factory=list)
    total_words: int = 0


@dataclass
class ContentLibrary:
    """Library of harvested content indexed by doc_type + section pattern."""
    # Key: (doc_type, normalized_section_id) → list of HarvestedSection from different conditions
    sections: dict[str, list[HarvestedSection]] = field(default_factory=dict)
    doc_types: list[str] = field(default_factory=list)
    total_sections: int = 0
    total_documents: int = 0


# Condition names to strip/replace when adapting content
_CONDITION_NAMES = {
    "parkinsons": ["Parkinson's Disease", "Parkinson's", "Parkinsons", "PD"],
    "depression": ["Major Depressive Disorder", "Depression", "MDD"],
    "anxiety": ["Generalized Anxiety Disorder", "Anxiety Disorders", "GAD"],
    "adhd": ["Attention Deficit Hyperactivity Disorder", "ADHD", "ADD"],
    "alzheimers": ["Alzheimer's Disease", "Alzheimer's", "Alzheimers", "MCI",
                   "Mild Cognitive Impairment"],
    "stroke_rehab": ["Post-Stroke Rehabilitation", "Stroke Rehabilitation", "Stroke"],
    "tbi": ["Traumatic Brain Injury", "TBI", "Post-Concussion"],
    "chronic_pain": ["Chronic Pain", "Fibromyalgia", "Chronic Pain / Fibromyalgia"],
    "ptsd": ["Post-Traumatic Stress Disorder", "PTSD"],
    "ocd": ["Obsessive-Compulsive Disorder", "OCD"],
    "ms": ["Multiple Sclerosis", "MS"],
    "asd": ["Autism Spectrum Disorder", "ASD", "Autism"],
    "long_covid": ["Long COVID", "Post-COVID", "Post-Acute Sequelae"],
    "tinnitus": ["Chronic Tinnitus", "Tinnitus"],
    "insomnia": ["Chronic Insomnia", "Insomnia", "Sleep Disorder"],
}


class ContentHarvester:
    """Extracts rich content from existing SOZO documents."""

    def harvest_document(self, file_path: Path) -> HarvestedDocument:
        """Extract all sections with their full content from a DOCX file."""
        file_path = Path(file_path)
        doc = docx.Document(str(file_path))

        condition_slug, doc_type, tier = self._infer_metadata(file_path)

        result = HarvestedDocument(
            file_path=str(file_path),
            condition_slug=condition_slug,
            doc_type=doc_type,
            tier=tier,
        )

        # Parse sections by detecting bold headings
        current_title = ""
        current_content_lines: list[str] = []
        current_section_id = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = (
                any(r.bold for r in para.runs if r.bold)
                and len(text) < 120
                and len(text.split()) < 20
            )

            if is_heading:
                # Save previous section
                if current_title and current_content_lines:
                    content = "\n".join(current_content_lines)
                    section = HarvestedSection(
                        section_id=current_section_id,
                        title=current_title,
                        content=content,
                        word_count=len(content.split()),
                        source_file=file_path.name,
                        source_condition=condition_slug,
                        source_doc_type=doc_type,
                        source_tier=tier,
                    )
                    # Find condition-specific phrases
                    if condition_slug in _CONDITION_NAMES:
                        for name in _CONDITION_NAMES[condition_slug]:
                            if name.lower() in content.lower():
                                section.condition_specific_phrases.append(name)
                    result.sections.append(section)

                current_title = text
                current_section_id = self._slugify(text)
                current_content_lines = []
            else:
                current_content_lines.append(text)

        # Flush last section
        if current_title and current_content_lines:
            content = "\n".join(current_content_lines)
            section = HarvestedSection(
                section_id=current_section_id,
                title=current_title,
                content=content,
                word_count=len(content.split()),
                source_file=file_path.name,
                source_condition=condition_slug,
                source_doc_type=doc_type,
                source_tier=tier,
            )
            result.sections.append(section)

        # Harvest tables
        for i, table in enumerate(doc.tables):
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])
            table_data = {"headers": headers, "rows": rows, "index": i}

            # Associate table with nearest previous section
            if result.sections:
                result.sections[-1].has_table = True
                result.sections[-1].table_data.append(table_data)

        result.total_words = sum(s.word_count for s in result.sections)
        return result

    def harvest_directory(self, directory: Path) -> list[HarvestedDocument]:
        """Harvest all DOCX files in a directory."""
        results = []
        for docx_file in sorted(Path(directory).rglob("*.docx")):
            try:
                harvested = self.harvest_document(docx_file)
                if harvested.sections:
                    results.append(harvested)
            except Exception as e:
                logger.warning("Failed to harvest %s: %s", docx_file, e)
        logger.info("Harvested %d documents from %s", len(results), directory)
        return results

    def build_content_library(self, documents: list[HarvestedDocument]) -> ContentLibrary:
        """Build a reusable content library from harvested documents."""
        library = ContentLibrary()
        library.total_documents = len(documents)

        doc_types = set()
        for doc in documents:
            doc_types.add(doc.doc_type)
            for section in doc.sections:
                # Normalize section ID for cross-condition matching
                norm_id = self._normalize_section_id(section.section_id)
                key = f"{doc.doc_type}::{norm_id}"
                if key not in library.sections:
                    library.sections[key] = []
                library.sections[key].append(section)
                library.total_sections += 1

        library.doc_types = sorted(doc_types)
        return library

    def save_library(self, library: ContentLibrary, output_path: Path) -> None:
        """Save content library to JSON."""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        data = {
            "doc_types": library.doc_types,
            "total_sections": library.total_sections,
            "total_documents": library.total_documents,
            "sections": {},
        }
        for key, sections in library.sections.items():
            data["sections"][key] = [
                {
                    "section_id": s.section_id,
                    "title": s.title,
                    "content": s.content,
                    "word_count": s.word_count,
                    "has_table": s.has_table,
                    "table_data": s.table_data,
                    "source_file": s.source_file,
                    "source_condition": s.source_condition,
                    "source_doc_type": s.source_doc_type,
                    "source_tier": s.source_tier,
                    "condition_specific_phrases": s.condition_specific_phrases,
                }
                for s in sections
            ]
        output_path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
        logger.info("Saved content library (%d sections) to %s", library.total_sections, output_path)

    def load_library(self, input_path: Path) -> ContentLibrary:
        """Load content library from JSON."""
        data = json.loads(Path(input_path).read_text(encoding="utf-8"))
        library = ContentLibrary(
            doc_types=data.get("doc_types", []),
            total_sections=data.get("total_sections", 0),
            total_documents=data.get("total_documents", 0),
        )
        for key, sections_data in data.get("sections", {}).items():
            library.sections[key] = [
                HarvestedSection(**{k: v for k, v in s.items()
                                    if k in HarvestedSection.__dataclass_fields__})
                for s in sections_data
            ]
        return library

    def adapt_content(
        self,
        section: HarvestedSection,
        target_condition_slug: str,
        target_condition_name: str,
    ) -> str:
        """Adapt harvested content from one condition to another.

        Replaces condition-specific names/phrases while preserving
        the clinical prose structure.
        """
        content = section.content
        source_slug = section.source_condition

        # Replace source condition names with target condition name
        if source_slug in _CONDITION_NAMES:
            for source_name in _CONDITION_NAMES[source_slug]:
                # Case-insensitive replacement preserving the target name's casing
                content = re.sub(
                    re.escape(source_name),
                    target_condition_name,
                    content,
                    flags=re.IGNORECASE,
                )

        return content

    def get_best_section(
        self,
        library: ContentLibrary,
        doc_type: str,
        section_id: str,
        prefer_condition: str = "",
    ) -> Optional[HarvestedSection]:
        """Get the richest section content for a given doc_type and section_id.

        Prefers content from the specified condition if available,
        otherwise picks the one with the most words (richest content).
        """
        norm_id = self._normalize_section_id(section_id)
        key = f"{doc_type}::{norm_id}"

        candidates = library.sections.get(key, [])
        if not candidates:
            # Try partial match
            for lib_key, lib_sections in library.sections.items():
                if lib_key.startswith(f"{doc_type}::") and norm_id in lib_key:
                    candidates = lib_sections
                    break

        if not candidates:
            return None

        # Prefer the specified condition
        if prefer_condition:
            for c in candidates:
                if c.source_condition == prefer_condition:
                    return c

        # Otherwise pick the richest
        return max(candidates, key=lambda s: s.word_count)

    def _infer_metadata(self, path: Path) -> tuple[str, str, str]:
        """Infer condition, doc_type, tier from file path."""
        from .document_ingester import _infer_metadata_from_path
        return _infer_metadata_from_path(path)

    def _slugify(self, text: str) -> str:
        slug = re.sub(r"[^\w\s]", "", text.lower())
        return re.sub(r"\s+", "_", slug.strip())[:80]

    def _normalize_section_id(self, section_id: str) -> str:
        """Normalize for cross-condition matching."""
        from .consistency_scorer import _normalize_section_id
        return _normalize_section_id(section_id)
