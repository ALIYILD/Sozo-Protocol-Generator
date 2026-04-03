"""
Upload Extraction Pipeline — converts uploaded PDF/DOCX into structured Sozo protocol schema.

Pipeline:
  1. INGEST: Read file, extract raw text
  2. CHUNK: Split into ~2000-char chunks preserving paragraph boundaries
  3. EXTRACT: LLM extracts structured fields per chunk
  4. ASSEMBLE: Merge extractions into unified protocol schema
  5. PRESENT: Return confidence-annotated draft for clinician correction
"""
from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Result of extracting a protocol from an uploaded document."""
    upload_id: str = ""
    filename: str = ""
    file_type: str = ""
    raw_text: str = ""
    chunks: list[str] = field(default_factory=list)
    extracted_sections: list[dict] = field(default_factory=list)
    assembled_schema: dict = field(default_factory=dict)
    overall_confidence: str = "medium"
    warnings: list[str] = field(default_factory=list)
    error: str = ""

    @property
    def success(self) -> bool:
        return bool(self.assembled_schema) and not self.error


class ProtocolExtractor:
    """Extracts structured protocol data from uploaded documents."""

    def extract(self, file_path: str | Path) -> ExtractionResult:
        """Full extraction pipeline: ingest → chunk → extract → assemble."""
        path = Path(file_path)
        result = ExtractionResult(filename=path.name, file_type=path.suffix.lstrip("."))

        # 1. Ingest
        try:
            if path.suffix.lower() == ".docx":
                result.raw_text = self._extract_docx_text(path)
            elif path.suffix.lower() == ".pdf":
                result.raw_text = self._extract_pdf_text(path)
            elif path.suffix.lower() in (".txt", ".md"):
                result.raw_text = path.read_text(encoding="utf-8")
            else:
                result.error = f"Unsupported file type: {path.suffix}"
                return result
        except Exception as e:
            result.error = f"File read error: {e}"
            return result

        if not result.raw_text.strip():
            result.error = "No text could be extracted from file"
            return result

        # 2. Chunk
        result.chunks = self._chunk_text(result.raw_text)
        logger.info(f"Chunked {path.name} into {len(result.chunks)} chunks")

        # 3. Extract structure from chunks (rule-based MVP, LLM in V2)
        result.extracted_sections = self._extract_sections(result.chunks, result.raw_text)

        # 4. Assemble into protocol schema
        result.assembled_schema = self._assemble_schema(result.extracted_sections, result.raw_text)
        result.overall_confidence = self._compute_confidence(result.assembled_schema)

        return result

    def _extract_docx_text(self, path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        doc = Document(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading-like formatting
                style = (para.style.name or "").lower()
                if "heading" in style:
                    paragraphs.append(f"\n## {text}\n")
                else:
                    paragraphs.append(text)
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    def _extract_pdf_text(self, path: Path) -> str:
        """Extract text from PDF. Try PyMuPDF first, fallback to basic."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            return text
        except ImportError:
            pass

        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except ImportError:
            pass

        return f"[PDF extraction requires PyMuPDF or pdfplumber: {path.name}]"

    def _chunk_text(self, text: str, max_chars: int = 2000) -> list[str]:
        """Split text into chunks preserving paragraph boundaries."""
        paragraphs = text.split("\n")
        chunks = []
        current = []
        current_len = 0

        for para in paragraphs:
            if current_len + len(para) > max_chars and current:
                chunks.append("\n".join(current))
                current = []
                current_len = 0
            current.append(para)
            current_len += len(para)

        if current:
            chunks.append("\n".join(current))

        return chunks

    def _extract_sections(self, chunks: list[str], full_text: str) -> list[dict]:
        """Extract protocol sections from text (rule-based MVP)."""
        sections = []

        # Detect condition
        condition = self._detect_condition(full_text)
        if condition:
            sections.append({
                "section_slug": "condition",
                "extracted_content": condition["name"],
                "extracted_data": condition,
                "confidence": condition.get("confidence", "medium"),
            })

        # Detect modality and parameters
        for modality_data in self._detect_modalities(full_text):
            sections.append({
                "section_slug": f"protocol_{modality_data['modality']}",
                "extracted_content": modality_data.get("raw_text", ""),
                "extracted_data": modality_data,
                "confidence": modality_data.get("confidence", "medium"),
            })

        # Detect safety/contraindications
        safety = self._detect_safety(full_text)
        if safety:
            sections.append({
                "section_slug": "safety",
                "extracted_content": "; ".join(safety),
                "extracted_data": {"contraindications": safety},
                "confidence": "medium",
            })

        # Detect references/PMIDs
        pmids = self._detect_pmids(full_text)
        if pmids:
            sections.append({
                "section_slug": "references",
                "extracted_content": f"{len(pmids)} references found",
                "extracted_data": {"pmids": pmids},
                "confidence": "high",
            })

        return sections

    def _detect_condition(self, text: str) -> Optional[dict]:
        """Detect which condition the protocol is about."""
        conditions = {
            "parkinsons": (r"parkinson", "Parkinson's Disease", "G20"),
            "depression": (r"depress|MDD|major depressive", "Major Depressive Disorder", "F32"),
            "adhd": (r"ADHD|attention.deficit", "ADHD", "F90"),
            "anxiety": (r"anxiety|GAD|generali[sz]ed anxiety", "Generalized Anxiety Disorder", "F41"),
            "alzheimers": (r"alzheimer|MCI|mild cognitive", "Alzheimer's Disease", "G30"),
            "stroke_rehab": (r"stroke|post.stroke|cerebrovascular", "Post-Stroke Rehabilitation", "I63"),
            "tbi": (r"traumatic brain|TBI|concussion", "Traumatic Brain Injury", "S06"),
            "ptsd": (r"PTSD|post.traumatic", "PTSD", "F43.1"),
            "ocd": (r"OCD|obsessive.compulsive", "OCD", "F42"),
            "migraine": (r"migraine", "Migraine", "G43"),
        }
        text_lower = text[:5000].lower()
        for slug, (pattern, name, icd10) in conditions.items():
            if re.search(pattern, text_lower, re.IGNORECASE):
                return {"slug": slug, "name": name, "icd10": icd10, "confidence": "high"}
        return None

    def _detect_modalities(self, text: str) -> list[dict]:
        """Detect stimulation modalities and parameters."""
        modalities = []
        text_lower = text.lower()

        patterns = {
            "tdcs": {
                "pattern": r"tdcs|transcranial direct current|direct current stimulation",
                "intensity": r"(\d+(?:\.\d+)?)\s*m[Aa]",
                "duration": r"(\d+)\s*min",
            },
            "tps": {
                "pattern": r"tps|transcranial pulse|neurolith|focused ultrasound",
                "energy": r"(\d+(?:\.\d+)?)\s*mJ",
                "pulses": r"(\d+(?:,\d+)?)\s*pulses",
            },
            "tavns": {
                "pattern": r"tavns|vagus nerve|auricular",
                "intensity": r"(\d+(?:\.\d+)?)\s*m[Aa]",
            },
            "ces": {
                "pattern": r"ces|cranial electrotherapy|alpha.stim",
                "intensity": r"(\d+)\s*[µu][Aa]",
            },
        }

        for modality, config in patterns.items():
            if re.search(config["pattern"], text_lower):
                data = {"modality": modality, "confidence": "medium"}
                # Extract parameters
                for param_name, param_pattern in config.items():
                    if param_name == "pattern":
                        continue
                    match = re.search(param_pattern, text)
                    if match:
                        data[param_name] = match.group(1)
                        data["confidence"] = "high"
                modalities.append(data)

        return modalities

    def _detect_safety(self, text: str) -> list[str]:
        """Detect contraindications and safety mentions."""
        safety_terms = [
            r"contraindication", r"seizure", r"epilepsy", r"pacemaker",
            r"metal implant", r"pregnancy", r"adverse", r"off.label",
        ]
        found = []
        for term in safety_terms:
            if re.search(term, text, re.IGNORECASE):
                found.append(term.replace(r"\b", "").replace(".", "-"))
        return found

    def _detect_pmids(self, text: str) -> list[str]:
        """Extract PMID references from text."""
        pmids = re.findall(r"PMID[:\s]*(\d{7,9})", text)
        return list(set(pmids))

    def _assemble_schema(self, sections: list[dict], full_text: str) -> dict:
        """Assemble extracted sections into a unified protocol schema."""
        schema: dict[str, Any] = {
            "source": "uploaded",
            "sections": [],
            "condition": {},
            "modalities": [],
            "safety": [],
            "references": [],
        }

        for sec in sections:
            slug = sec["section_slug"]
            data = sec.get("extracted_data", {})

            if slug == "condition":
                schema["condition"] = data
            elif slug.startswith("protocol_"):
                schema["modalities"].append(data)
            elif slug == "safety":
                schema["safety"] = data.get("contraindications", [])
            elif slug == "references":
                schema["references"] = [{"pmid": p} for p in data.get("pmids", [])]

            schema["sections"].append({
                "section_slug": slug,
                "content": sec.get("extracted_content", ""),
                "confidence": sec.get("confidence", "low"),
            })

        # Add raw text length for reference
        schema["raw_text_length"] = len(full_text)
        schema["chunks_processed"] = len(sections)

        return schema

    def _compute_confidence(self, schema: dict) -> str:
        """Compute overall extraction confidence."""
        if not schema.get("condition"):
            return "low"
        has_modality = len(schema.get("modalities", [])) > 0
        has_safety = len(schema.get("safety", [])) > 0
        has_refs = len(schema.get("references", [])) > 0

        score = sum([bool(schema.get("condition")), has_modality, has_safety, has_refs])
        if score >= 3:
            return "high"
        if score >= 2:
            return "medium"
        return "low"
