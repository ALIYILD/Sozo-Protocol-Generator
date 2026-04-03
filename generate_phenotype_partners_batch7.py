# DEPRECATED: This script is superseded by the canonical generation pipeline.
# Use instead: GenerationService.generate(condition="...", tier="...", doc_type="...")
# Or CLI: PYTHONPATH=src python -m sozo_generator.cli.main build condition --condition <slug> --tier <tier> --doc-type <type>
# See docs/MIGRATION_PLAN.md for details.

from pathlib import Path
from docx import Document
from docx.shared import RGBColor

_PROJECT_ROOT = Path(__file__).resolve().parent

C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)


def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None
    size = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except Exception:
        color = None
    for r in para.runs:
        r.text = ""
    if fr:
        fr.text = replaced
        fr.font.bold = bold
        if size:
            fr.font.size = size
        if color:
            fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type:
            run.font.color.rgb = C_BLACK


def _replace_table(table, rows):
    for ri in range(len(table.rows)):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells:
                _cell_write(cell, "")
            continue
        is_hdr = ri == 0
        for ci, cell in enumerate(row.cells):
            _cell_write(
                cell,
                str(rows[ri][ci]) if ci < len(rows[ri]) else "",
                bold=is_hdr,
                white=is_hdr,
            )


def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs):
            _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs):
                        _para_replace(para, old, new)


def build(c):
    TEMPLATE = _PROJECT_ROOT / "templates" / "gold_standard" / "Phenotype_Classification.docx"
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs
    tables = doc.tables
    for old, new in [
        ("Parkinson's Disease (PD)", c["full"]),
        ("Parkinson's Disease", c["full"]),
        ("Parkinson's", c["short"]),
        (" (PD)", f" ({c['abbr']})"),
        (" PD ", f" {c['abbr']} "),
        (" PD.", f" {c['abbr']}."),
        (" PD,", f" {c['abbr']},"),
        ("for PD", f"for {c['abbr']}"),
        ("H&Y Stage", c["severity_field"]),
        ("Disease Duration", c["duration_field"]),
        ("PD Phenotype", f"{c['abbr']} Phenotype"),
        ("Hoehn & Yahr", "disease severity rating"),
        ("levodopa", c.get("med_name", "primary medication")),
    ]:
        _global_replace(doc, old, new)

    _para_set(paras[1], f"{c['short']} Phenotype Classification & Protocol Mapping")
    _para_set(paras[7], f"STEP 1 \u2014 {c['short']} Phenotype Determination")
    _para_set(paras[11], f"1.2 {c['short']} Phenotype Classification")
    _para_set(paras[14], f"ASSIGNED PHENOTYPE: {c['phenotype_codes']}")
    _para_set(paras[19], c["g1"])
    _para_set(paras[21], c["g2"])
    _para_set(paras[23], c["g3"])
    _para_set(paras[27], c["tps_ol"])
    _para_set(paras[40], c["sel_opt"])

    _replace_table(
        tables[0],
        [
            ["Field", "Details"],
            ["Patient Name", ""],
            ["File Number", ""],
            [c["severity_field"], ""],
            ["Date", ""],
            ["Treating Doctor", ""],
            [c["duration_field"], "___ years / ___ months"],
        ],
    )

    for ti, key in [
        (1, "dc"),
        (2, "ph"),
        (3, "nw"),
        (5, "t5"),
        (6, "t6"),
        (7, "t7"),
        (9, "t9"),
        (12, "tavns"),
        (13, "ces"),
        (14, "seq"),
    ]:
        _replace_table(tables[ti], c[key])

    _cell_write(tables[15].rows[1].cells[0], c["resp"])
    _cell_write(
        tables[15].rows[1].cells[1],
        "\u25a1 Completed | Score: ___ / baseline: ___",
    )
    _cell_write(tables[18].rows[1].cells[0], f"{c['abbr']} Phenotype")

    out = (
        Path("outputs/documents")
        / c["slug"]
        / "partners"
        / f"Phenotype_Classification_Partners_{c['slug']}.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  [OK] {out}")
    return out


TINNITUS = {
    "slug": "tinnitus",
    "short": "Tinnitus",
    "abbr": "TIN",
    "full": "Chronic Tinnitus (TIN)",
    "severity_field": "THI / TFI Severity",
    "duration_field": "Tinnitus Duration",
    "med_name": "tinnitus medication",
    "phenotype_codes": "\u25a1 TO  \u25a1 NO  \u25a1 PU  \u25a1 RE  \u25a1 SO  \u25a1 HY  \u25a1 DE",
    "g1": "3.1 Tonal / Noise-Type (TO / NO)",
    "g2": "3.2 Reactive / Somatic / Hyperacusis (RE / SO / HY)",
    "g3": "3.3 Pulsatile / Depression-Comorbid (PU / DE)",
    "tps_ol": "TPS is OFF-LABEL for Tinnitus. Requires Doctor authorisation and off-label consent.",
    "sel_opt": "SELECTED OPTION: \u25a1 A (Tonal/Central)  \u25a1 B (Reactive/Somatic)  \u25a1 C (Hyperacusis)  \u25a1 D (Mixed)",
    "resp": "Reassess SOZO PRS + THI / TFI",
    "dc": [
        ["Data Point", "Completed", "Source"],
        ["THI / TFI severity score", "\u25a1", "Rating Scales"],
        ["Tinnitus character: tonal/noise/pulsatile, laterality", "\u25a1", "Clinical Interview"],
        ["Audiological assessment (PTA, tinnitus matching)", "\u25a1", "Audiologist"],
        ["Hyperacusis severity (HQ, LDL testing)", "\u25a1", "Audiological"],
        ["Somatic modulation testing (jaw, cervical)", "\u25a1", "Clinical Exam"],
        ["Comorbid depression and anxiety (PHQ-9, GAD-7)", "\u25a1", "Rating Scales"],
        ["Sleep disruption secondary to tinnitus (ISI)", "\u25a1", "ISI / Interview"],
        ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"],
    ],
    "ph": [
        ["Dominant Features", "Assign Phenotype", "Code"],
        ["Tonal, pure-tone quality tinnitus, stable pitch-matched", "Tonal (TO)", "TO"],
        ["Broadband noise, hissing, rushing quality dominant", "Noise-Type (NO)", "NO"],
        ["Pulsatile, rhythmic, heartbeat-synchronous tinnitus", "Pulsatile (PU)", "PU"],
        ["Tinnitus loudness/character fluctuates with sound exposure", "Reactive (RE)", "RE"],
        ["Tinnitus modulated by jaw, neck, or somatic manoeuvres", "Somatic (SO)", "SO"],
        ["Severe sound sensitivity, hyperacusis dominant feature", "Hyperacusis (HY)", "HY"],
        ["Tinnitus-related depression, emotional distress dominant", "Depression-Comorbid (DE)", "DE"],
    ],
    "nw": [
        ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
        ["TO", "Auditory (L-A1 cathodal)", "DMN", "CEN (L-DLPFC)"],
        ["NO", "Auditory (bilateral A1)", "SN", "CEN"],
        ["PU", "Auditory + vascular", "SN", "CEN"],
        ["RE", "Auditory (A1)", "SN (Insula)", "CEN (DLPFC)"],
        ["SO", "Auditory + SMN (cervical)", "SN", "CEN"],
        ["HY", "Auditory (A1 hyper)", "SN (Insula)", "CEN (L-DLPFC)"],
        ["DE", "CEN (L-DLPFC)", "Limbic", "Auditory"],
    ],
    "t5": [
        ["Goal", "Anode", "Cathode"],
        ["L-A1 cathodal suppression (TO/NO)", "Left DLPFC (F3)", "Left temporal (T3/A1) cathodal"],
        ["L-DLPFC top-down auditory suppression", "Left DLPFC (F3)", "Right supraorbital"],
        ["Bilateral A1 inhibition (NO/HY)", "Left DLPFC", "Bilateral temporal cathodal"],
    ],
    "t6": [
        ["Goal", "Anode", "Cathode"],
        ["R-DLPFC cathodal distress reduction (DE)", "Left DLPFC", "Right DLPFC cathodal"],
        ["Somatosensory cortex modulation (SO)", "Left S1 (C3)", "Right supraorbital"],
    ],
    "t7": [
        ["Goal", "Anode", "Cathode"],
        ["Insula/SN hyperacusis modulation (HY)", "Left DLPFC", "Right temporal cathodal"],
        ["L-DLPFC antidepressant (DE)", "Left DLPFC (F3)", "Right supraorbital"],
    ],
    "t9": [
        ["Clinical Target", "TPS Protocol", "Phenotype"],
        ["L-A1 / auditory cortex inhibition", "FT1", "TO, NO, RE"],
        ["L-DLPFC top-down control", "FT3", "All phenotypes"],
        ["Bilateral auditory network", "FT4", "NO, HY"],
        ["SN/Insula distress circuit", "FT5", "HY, DE"],
        ["Somatosensory cortex (SO)", "FT7", "SO"],
        ["Limbic-distress circuit", "FT8", "DE"],
        ["Thalamic auditory relay", "FT9", "TO, PU"],
        ["A1 + DLPFC combined", "Multiple", "RE / HY"],
        ["Multi-network tinnitus", "Multiple", "MX / severe"],
    ],
    "tavns": [
        ["Domain", "Details"],
        ["Indications", "Tinnitus distress, depression comorbidity, sleep disruption, autonomic dysregulation"],
        ["Protocol", "20\u201330 min with paired auditory tones if possible, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
        ["Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "ces": [
        ["Domain", "Details"],
        ["Indications", "Anxiety-tinnitus comorbidity, insomnia, hyperacusis, depression-comorbid phenotype"],
        ["Protocol", "20\u201360 min daily, Alpha-Stim AID, especially DE, HY, RE phenotypes"],
        ["Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "seq": [
        ["Option", "Sequence", "Phenotype"],
        ["OPTION A \u2014 Tonal/Central", "Wk 1\u20134: L-A1 cathodal + L-DLPFC anodal tDCS (5\u00d7/wk) \u2192 Wk 2\u20136: TPS A1/DLPFC", "TO / NO"],
        ["OPTION B \u2014 Reactive/Somatic", "Wk 1\u20134: S1/DLPFC tDCS \u2192 Wk 3\u20136: TPS somatosensory/A1", "RE / SO"],
        ["OPTION C \u2014 Hyperacusis", "Wk 1\u20132: CES stabilization \u2192 Wk 2\u20136: tDCS bilateral A1 cathodal + TPS A1/SN", "HY"],
        ["OPTION D \u2014 Depression/Mixed", "Sequential: TPS (highest burden) \u2192 Maintenance L-DLPFC tDCS + taVNS", "DE / MX"],
    ],
}

INSOMNIA = {
    "slug": "insomnia",
    "short": "Insomnia",
    "abbr": "INS",
    "full": "Chronic Insomnia Disorder (INS)",
    "severity_field": "ISI / PSQI Severity",
    "duration_field": "Insomnia Duration",
    "med_name": "hypnotic/sleep medication",
    "phenotype_codes": "\u25a1 SO  \u25a1 MA  \u25a1 EM  \u25a1 HY  \u25a1 PS  \u25a1 PA  \u25a1 CI",
    "g1": "3.1 Sleep-Onset / Maintenance (SO / MA)",
    "g2": "3.2 Hyperarousal / Psychophysiological (HY / PS)",
    "g3": "3.3 Early-Morning / Paradoxical / Circadian (EM / PA / CI)",
    "tps_ol": "TPS is OFF-LABEL for Insomnia. Requires Doctor authorisation and off-label consent.",
    "sel_opt": "SELECTED OPTION: \u25a1 A (Sleep-Onset)  \u25a1 B (Hyperarousal)  \u25a1 C (Psychiatric)  \u25a1 D (Circadian/Mixed)",
    "resp": "Reassess SOZO PRS + ISI / PSQI / sleep diary",
    "dc": [
        ["Data Point", "Completed", "Source"],
        ["ISI / PSQI severity score", "\u25a1", "Rating Scales"],
        ["Sleep diary: SOL, WASO, TST, SE (2 weeks)", "\u25a1", "Sleep Diary"],
        ["Hyperarousal assessment (PSAS, FIRST)", "\u25a1", "Arousal Scales"],
        ["Actigraphy / sleep study findings", "\u25a1", "Actigraphy / PSG"],
        ["Psychiatric comorbidity (PHQ-9, GAD-7)", "\u25a1", "Rating Scales"],
        ["Circadian preference (MEQ, dim-light melatonin)", "\u25a1", "Chronotype Assessment"],
        ["Paradoxical insomnia / sleep misperception", "\u25a1", "Clinical Interview"],
        ["6-Network Bedside Assessment scores", "\u25a1", "Assessment Form"],
    ],
    "ph": [
        ["Dominant Features", "Assign Phenotype", "Code"],
        ["Difficulty initiating sleep, prolonged SOL dominant", "Sleep-Onset (SO)", "SO"],
        ["Frequent night wakings, difficulty returning to sleep (WASO)", "Maintenance (MA)", "MA"],
        ["Early morning awakening, unable to return to sleep", "Early-Morning (EM)", "EM"],
        ["Physiological and cognitive hyperarousal dominant", "Hyperarousal (HY)", "HY"],
        ["Insomnia secondary to psychiatric disorder (depression/anxiety)", "Psychiatric (PS)", "PS"],
        ["Sleep state misperception, paradoxical insomnia", "Paradoxical (PA)", "PA"],
        ["Circadian rhythm disruption, phase delay/advance dominant", "Circadian (CI)", "CI"],
    ],
    "nw": [
        ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
        ["SO", "SN (hyperarousal)", "CEN (R-DLPFC cathodal)", "Thalamic"],
        ["MA", "SN", "DMN", "CEN (R-DLPFC cathodal)"],
        ["EM", "SN + Limbic", "DMN", "CEN"],
        ["HY", "SN (Insula/dACC)", "CEN (R-DLPFC)", "SMN"],
        ["PS", "CEN (L-DLPFC)", "Limbic", "SN"],
        ["PA", "DMN (perception)", "SN", "CEN"],
        ["CI", "Hypothalamic-SCN", "SN", "DMN"],
    ],
    "t5": [
        ["Goal", "Anode", "Cathode"],
        ["R-DLPFC cathodal arousal reduction (SO/HY)", "Left DLPFC anodal", "Right DLPFC cathodal"],
        ["SN/dACC downregulation (HY/MA)", "Left DLPFC", "Right DLPFC cathodal"],
        ["L-DLPFC psychiatric comorbidity (PS)", "Left DLPFC (F3)", "Right supraorbital"],
    ],
    "t6": [
        ["Goal", "Anode", "Cathode"],
        ["Frontal delta/slow-wave induction (SO/MA)", "Left DLPFC", "Right DLPFC cathodal"],
        ["Limbic modulation early-morning (EM)", "Left DLPFC", "Right supraorbital"],
    ],
    "t7": [
        ["Goal", "Anode", "Cathode"],
        ["Bilateral frontal relaxation (HY/PA)", "Right DLPFC cathodal", "Left DLPFC anodal"],
        ["L-DLPFC antidepressant for PS", "Left DLPFC (F3)", "Right supraorbital"],
    ],
    "t9": [
        ["Clinical Target", "TPS Protocol", "Phenotype"],
        ["R-DLPFC cathodal arousal suppression", "FT1", "SO, HY, MA"],
        ["SN/Insula hyperarousal", "FT3", "HY"],
        ["Limbic-prefrontal (EM/PS)", "FT4", "EM, PS"],
        ["Thalamic arousal gating", "FT5", "SO, MA"],
        ["L-DLPFC psychiatric circuit", "FT7", "PS"],
        ["DMN sleep onset (PA)", "FT8", "PA"],
        ["Hypothalamic-frontal circadian", "FT9", "CI"],
        ["Bilateral frontal arousal", "Multiple", "HY / severe"],
        ["Multi-network insomnia", "Multiple", "MX / chronic"],
    ],
    "tavns": [
        ["Domain", "Details"],
        ["Indications", "Hyperarousal, autonomic activation, sleep-onset difficulty, all insomnia phenotypes"],
        ["Protocol", "30\u201345 min pre-sleep, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
        ["Add taVNS?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "ces": [
        ["Domain", "Details"],
        ["Indications", "Insomnia first-line adjunct for all phenotypes \u2014 especially HY, SO, PS"],
        ["Protocol", "20\u201360 min pre-sleep daily, Alpha-Stim AID, titrate to sedation response"],
        ["Add CES?", "\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "seq": [
        ["Option", "Sequence", "Phenotype"],
        ["OPTION A \u2014 Sleep-Onset", "Wk 1\u20134: R-DLPFC cathodal tDCS (pre-sleep, 5\u00d7/wk) + CES \u2192 Wk 2\u20136: TPS SN/thalamic", "SO / MA"],
        ["OPTION B \u2014 Hyperarousal", "Wk 1\u20132: CES/taVNS stabilization \u2192 Wk 2\u20136: Bilateral DLPFC tDCS + TPS HY circuit", "HY"],
        ["OPTION C \u2014 Psychiatric", "Wk 1\u20134: L-DLPFC tDCS + Daily CES \u2192 Wk 3\u20136: TPS L-DLPFC/Limbic", "PS / EM"],
        ["OPTION D \u2014 Circadian/Mixed", "Sequential: TPS (highest burden) \u2192 Maintenance R-DLPFC cathodal + taVNS", "CI / MX"],
    ],
}

if __name__ == "__main__":
    print("Generating Partners Phenotype Classification documents — Batch 7")
    print("=" * 60)
    results = []
    for condition in [TINNITUS, INSOMNIA]:
        print(f"\nBuilding: {condition['short']} ({condition['abbr']})")
        out = build(condition)
        results.append(out)

    print("\n" + "=" * 60)
    print("Output file sizes:")
    for out in results:
        p = Path(out)
        if p.exists():
            size_kb = p.stat().st_size / 1024
            print(f"  {p} — {size_kb:.1f} KB")
        else:
            print(f"  MISSING: {p}")
    print("\nDone.")
