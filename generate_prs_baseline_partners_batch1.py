# DEPRECATED: This script is superseded by the canonical generation pipeline.
# Use instead: GenerationService.generate(condition="...", tier="...", doc_type="...")
# Or CLI: PYTHONPATH=src python -m sozo_generator.cli.main build condition --condition <slug> --tier <tier> --doc-type <type>
# See docs/MIGRATION_PLAN.md for details.

"""
Generates Partners Tier Psychological_Intake_PRS_Baseline_Partners DOCX files
for 7 neuromodulation conditions (Batch 1).
"""

from pathlib import Path
from docx import Document
from docx.shared import RGBColor

_PROJECT_ROOT = Path(__file__).resolve().parent

C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)


def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None
    size = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except Exception:
        color = None
    for r in para.runs:
        r.text = ""
    if fr:
        fr.text = replaced
        fr.font.bold = bold
        if size:
            fr.font.size = size
        if color:
            fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type:
            run.font.color.rgb = C_BLACK


def _replace_table(table, rows):
    for ri in range(len(table.rows)):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells:
                _cell_write(cell, "")
            continue
        is_hdr = (ri == 0)
        for ci, cell in enumerate(row.cells):
            _cell_write(
                cell,
                str(rows[ri][ci]) if ci < len(rows[ri]) else "",
                bold=is_hdr,
                white=is_hdr,
            )


def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs):
            _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs):
                        _para_replace(para, old, new)


def build_prs(c):
    TEMPLATE = _PROJECT_ROOT / "templates" / "gold_standard" / "Psychological_Intake_PRS.docx"
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs
    tables = doc.tables

    # Global text replacements (order matters: longest/most specific first)
    for old, new in [
        ("Parkinson's Disease (PD)", c["full"]),
        ("Parkinson's Disease", c["full"]),
        ("Parkinson's", c["short"]),
        (" (PD)", f" ({c['abbr']})"),
        (" PD ", f" {c['abbr']} "),
        (" PD.", f" {c['abbr']}."),
        (" PD,", f" {c['abbr']},"),
        ("for PD", f"for {c['abbr']}"),
        ("PD Phenotype", f"{c['abbr']} Phenotype"),
        ("Hoehn & Yahr", "disease severity rating"),
        ("levodopa", c.get("med_name", "primary medication")),
    ]:
        _global_replace(doc, old, new)

    # Paragraph replacements
    _para_set(paras[1], f"{c['short']} Psychological Intake & PRS Baseline")
    _para_set(paras[21], "B1. Primary Symptoms Domain")

    # Table[3]: condition-specific primary symptoms
    _replace_table(tables[3], c["primary_symptoms"])

    # Table[5] row[2] cell[1]: phenotype checkboxes
    _cell_write(tables[5].rows[2].cells[1], c["phenotype_checkboxes"])

    slug = c["slug"]
    out = (
        Path("outputs/documents")
        / slug
        / "partners"
        / f"Psychological_Intake_PRS_Baseline_Partners_{slug}.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  [OK] {out}")
    return out


CONDITIONS = [
    {
        "slug": "depression",
        "short": "Depression",
        "abbr": "MDD",
        "full": "Major Depressive Disorder (MDD)",
        "med_name": "antidepressant",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Low mood / persistent sadness", "", "\u25a1", ""],
            ["Anhedonia / loss of pleasure", "", "\u25a1", ""],
            ["Fatigue / energy loss", "", "\u25a1", ""],
            ["Sleep disturbance (insomnia or hypersomnia)", "", "\u25a1", ""],
            ["Psychomotor changes (retardation / agitation)", "", "\u25a1", ""],
            ["Concentration / cognitive difficulties", "", "\u25a1", ""],
            ["Appetite / weight changes", "", "\u25a1", ""],
            ["Hopelessness / worthlessness / guilt", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Melancholic (ME)  \u25a1 Atypical (AT)  \u25a1 Anxious (AX)  \u25a1 TR  \u25a1 PM  \u25a1 CG  \u25a1 AN",
    },
    {
        "slug": "anxiety",
        "short": "Anxiety",
        "abbr": "GAD",
        "full": "Generalized Anxiety Disorder (GAD)",
        "med_name": "anxiolytic/SSRI",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Worry severity / uncontrollability", "", "\u25a1", ""],
            ["Muscle tension / physical tension", "", "\u25a1", ""],
            ["Restlessness / feeling on edge", "", "\u25a1", ""],
            ["Concentration difficulty", "", "\u25a1", ""],
            ["Irritability", "", "\u25a1", ""],
            ["Sleep disturbance", "", "\u25a1", ""],
            ["Somatic symptoms (headache, GI, palpitations)", "", "\u25a1", ""],
            ["Avoidance behaviour", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Cognitive-Worry (CW)  \u25a1 Somatic (ST)  \u25a1 Mixed (MX)  \u25a1 PA  \u25a1 SO  \u25a1 IN  \u25a1 AU",
    },
    {
        "slug": "adhd",
        "short": "ADHD",
        "abbr": "ADHD",
        "full": "Attention Deficit Hyperactivity Disorder (ADHD)",
        "med_name": "stimulant medication",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Inattention / distractibility", "", "\u25a1", ""],
            ["Hyperactivity / motor restlessness", "", "\u25a1", ""],
            ["Impulsivity / acting without thinking", "", "\u25a1", ""],
            ["Working memory difficulties", "", "\u25a1", ""],
            ["Executive function (planning / organisation)", "", "\u25a1", ""],
            ["Emotional dysregulation / frustration", "", "\u25a1", ""],
            ["Sleep difficulties", "", "\u25a1", ""],
            ["Academic / occupational impairment", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Inattentive (IA)  \u25a1 Hyperactive (HI)  \u25a1 Combined (CO)  \u25a1 EX  \u25a1 EM  \u25a1 SC  \u25a1 WM",
    },
    {
        "slug": "alzheimers",
        "short": "Alzheimer's",
        "abbr": "AD",
        "full": "Alzheimer's Disease / MCI (AD)",
        "med_name": "cholinesterase inhibitor",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Memory loss (recent events, repetition)", "", "\u25a1", ""],
            ["Word-finding / naming difficulty", "", "\u25a1", ""],
            ["Executive function decline (planning)", "", "\u25a1", ""],
            ["Orientation difficulties (time / place)", "", "\u25a1", ""],
            ["Visuospatial problems", "", "\u25a1", ""],
            ["Personality / behaviour changes", "", "\u25a1", ""],
            ["Activities of daily living decline", "", "\u25a1", ""],
            ["Navigation / driving difficulties", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Amnestic (AM)  \u25a1 Multi-Domain (MD)  \u25a1 Early-Onset (EO)  \u25a1 LO  \u25a1 PC  \u25a1 LG  \u25a1 AP",
    },
    {
        "slug": "stroke",
        "short": "Stroke",
        "abbr": "CVA",
        "full": "Post-Stroke Rehabilitation (CVA)",
        "med_name": "anticoagulant/antiplatelet",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Motor weakness (affected limb/side)", "", "\u25a1", ""],
            ["Gait / walking difficulty", "", "\u25a1", ""],
            ["Speech / language difficulty", "", "\u25a1", ""],
            ["Balance / coordination impairment", "", "\u25a1", ""],
            ["Cognitive changes post-stroke", "", "\u25a1", ""],
            ["Post-stroke fatigue", "", "\u25a1", ""],
            ["Emotional changes (depression / lability)", "", "\u25a1", ""],
            ["Swallowing difficulty (dysphagia)", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Motor-UL (MU)  \u25a1 Motor-LL (ML)  \u25a1 Aphasia (AP)  \u25a1 NE  \u25a1 CG  \u25a1 SP  \u25a1 DE",
    },
    {
        "slug": "tbi",
        "short": "TBI",
        "abbr": "TBI",
        "full": "Traumatic Brain Injury (TBI)",
        "med_name": "neuroprotective agent",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Headache severity", "", "\u25a1", ""],
            ["Cognitive difficulties (attention / memory)", "", "\u25a1", ""],
            ["Fatigue (physical and cognitive)", "", "\u25a1", ""],
            ["Emotional dysregulation / irritability", "", "\u25a1", ""],
            ["Dizziness / vestibular symptoms", "", "\u25a1", ""],
            ["Sleep disruption", "", "\u25a1", ""],
            ["Behavioural changes", "", "\u25a1", ""],
            ["Noise / light sensitivity", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Cognitive (CG)  \u25a1 Executive (EX)  \u25a1 Headache (HA)  \u25a1 BE  \u25a1 VE  \u25a1 FA  \u25a1 ME",
    },
    {
        "slug": "chronic_pain",
        "short": "Chronic Pain",
        "abbr": "CP",
        "full": "Chronic Pain / Fibromyalgia (CP)",
        "med_name": "analgesic/pain medication",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Average pain intensity (NRS)", "", "\u25a1", ""],
            ["Worst pain (last 7 days)", "", "\u25a1", ""],
            ["Pain interference with daily activities", "", "\u25a1", ""],
            ["Sleep disruption from pain", "", "\u25a1", ""],
            ["Fatigue related to pain", "", "\u25a1", ""],
            ["Mood effects of pain (depression/anxiety)", "", "\u25a1", ""],
            ["Mobility / physical function limitation", "", "\u25a1", ""],
            ["Medication effectiveness / side effects", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Nociceptive (NO)  \u25a1 Neuropathic (NE)  \u25a1 Nociplastic (NP)  \u25a1 FM  \u25a1 CR  \u25a1 MS  \u25a1 HA",
    },
]


if __name__ == "__main__":
    print("=" * 60)
    print("SOZO Partners — Psychological_Intake_PRS_Baseline_Partners")
    print("Generating 7 condition DOCX files...")
    print("=" * 60)

    generated = []
    errors = []

    for cond in CONDITIONS:
        try:
            out = build_prs(cond)
            generated.append(out)
        except Exception as e:
            print(f"  [ERROR] {cond['slug']}: {e}")
            errors.append((cond["slug"], str(e)))

    print()
    print("=" * 60)
    print("VERIFICATION — File sizes:")
    print("=" * 60)
    for out in generated:
        p = Path(out)
        if p.exists():
            size_kb = p.stat().st_size / 1024
            print(f"  {p.name:<60} {size_kb:>8.1f} KB")
        else:
            print(f"  MISSING: {out}")

    print()
    print(f"Generated : {len(generated)}/7")
    if errors:
        print(f"Errors    : {len(errors)}")
        for slug, msg in errors:
            print(f"  - {slug}: {msg}")
    else:
        print("Errors    : 0")
    print("Done.")
