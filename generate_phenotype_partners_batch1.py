# DEPRECATED: This script is superseded by the canonical generation pipeline.
# Use instead: GenerationService.generate(condition="...", tier="...", doc_type="...")
# Or CLI: PYTHONPATH=src python -m sozo_generator.cli.main build condition --condition <slug> --tier <tier> --doc-type <type>
# See docs/MIGRATION_PLAN.md for details.

"""
Generate Partners Tier Phenotype Classification & Protocol Mapping DOCX files
Batch 1: Depression (MDD) and Anxiety (GAD)
"""

from pathlib import Path
from docx import Document
from docx.shared import RGBColor

_PROJECT_ROOT = Path(__file__).resolve().parent

# ─── Colour constants ────────────────────────────────────────────────────────
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)

TEMPLATE = _PROJECT_ROOT / "templates" / "gold_standard" / "Phenotype_Classification.docx"


# ─── Helper functions ────────────────────────────────────────────────────────

def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None
    size = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except Exception:
        color = None
    for r in para.runs:
        r.text = ""
    if fr:
        fr.text = replaced
        fr.font.bold = bold
        if size:
            fr.font.size = size
        if color:
            fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        try:
            if run.font.color.type:
                run.font.color.rgb = C_BLACK
        except Exception:
            pass


def _replace_table(table, rows):
    n_tmpl = len(table.rows)
    for ri in range(n_tmpl):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells:
                _cell_write(cell, "")
            continue
        is_hdr = (ri == 0)
        for ci, cell in enumerate(row.cells):
            text = str(rows[ri][ci]) if ci < len(rows[ri]) else ""
            _cell_write(cell, text, bold=is_hdr, white=is_hdr)


def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs):
            _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs):
                        _para_replace(para, old, new)


# ─── Condition data ──────────────────────────────────────────────────────────

CONDITIONS = [
    # ── DEPRESSION (MDD) ──────────────────────────────────────────────────
    {
        "slug": "depression",
        "short": "Depression",
        "abbr": "MDD",
        "full_name": "Major Depressive Disorder",
        "severity_field": "PHQ-9 / MADRS Severity",
        "duration_field": "Episode Duration",
        "phenotype_codes": "□ ME  □ AT  □ AX  □ TR  □ PM  □ CG  □ AN",
        "group1_label": "3.1 CEN-Primary Phenotypes (ME / CG / AN)",
        "group2_label": "3.2 Limbic-Affective Phenotypes (AT / AX)",
        "group3_label": "3.3 Psychomotor / Treatment-Resistant (PM / TR)",
        "opt_a": "CEN-Primary",
        "opt_b": "Limbic",
        "opt_c": "TR",
        "opt_d": "Mixed",
        "tps_offlabel": "TPS is OFF-LABEL for MDD. Requires Doctor authorisation and off-label consent.",
        "table1": [
            ["Data Point", "Completed", "Source"],
            ["PHQ-9 / MADRS / HDRS severity score", "□", "Rating Scales"],
            ["Melancholic vs atypical feature screening", "□", "Clinical Interview"],
            ["Cognitive complaints and concentration", "□", "MoCA / Interview"],
            ["Anxiety comorbidity (GAD-7, HAM-A)", "□", "Rating Scales"],
            ["Anhedonia assessment (SHAPS, BAS)", "□", "SHAPS / Interview"],
            ["Treatment history (antidepressants, ECT, TMS)", "□", "Medical History"],
            ["Sleep quality and pattern (ISI)", "□", "ISI / Interview"],
            ["6-Network Bedside Assessment scores", "□", "Assessment Form"],
        ],
        "table2": [
            ["Dominant Features", "Assign Phenotype", "Code"],
            ["Pervasive low mood, anhedonia, diurnal variation, early morning waking", "Melancholic (ME)", "ME"],
            ["Mood reactivity, hypersomnia, hyperphagia, rejection sensitivity", "Atypical (AT)", "AT"],
            ["Prominent anxiety, worry, somatic tension, cognitive interference", "Anxious (AX)", "AX"],
            ["≥2 adequate antidepressant failures, chronic/refractory course", "Treatment-Resistant (TR)", "TR"],
            ["Psychomotor retardation or agitation as dominant feature", "Psychomotor (PM)", "PM"],
            ["Concentration, memory, executive dysfunction foregrounded", "Cognitive (CG)", "CG"],
            ["Profound loss of pleasure, emotional numbing, motivational deficit", "Anhedonic (AN)", "AN"],
        ],
        "table3": [
            ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
            ["ME", "DMN (vmPFC/sgACC)", "CEN (L-DLPFC)", "Limbic (OFC)"],
            ["AT", "Limbic (OFC/BLA)", "SN (Insula)", "CEN (R-DLPFC cathodal)"],
            ["AX", "SN (dACC/Insula)", "Limbic (BLA)", "CEN (R-DLPFC cathodal)"],
            ["TR", "CEN bilateral", "SN + Limbic", "DMN (vmPFC)"],
            ["PM", "SMN (M1/SMA)", "CEN (L-DLPFC)", "Limbic"],
            ["CG", "CEN (L-DLPFC)", "DAN (parietal)", "DMN"],
            ["AN", "Limbic (NAcc-OFC)", "CEN (L-DLPFC)", "SN"],
        ],
        "table5": [
            ["Goal", "Anode", "Cathode"],
            ["Enhance left DLPFC excitability (ME/CG/AN)", "Left DLPFC (F3)", "Right supraorbital"],
            ["Bilateral DLPFC activation (TR/PM)", "Left DLPFC", "Right DLPFC (bi-cephalic)"],
            ["Reduce right hemisphere over-activation", "Right DLPFC (cathodal)", "Left supraorbital"],
        ],
        "table6": [
            ["Goal", "Anode", "Cathode"],
            ["Modulate OFC/vmPFC-limbic circuit (AT)", "vmPFC (Fpz)", "Right mastoid"],
            ["Reduce SN hyperactivity (AX)", "Left DLPFC", "Right DLPFC cathodal"],
        ],
        "table7": [
            ["Goal", "Anode", "Cathode"],
            ["Augment antidepressant effect (TR)", "Left DLPFC", "Right supraorbital"],
            ["Motor-affective combined (PM)", "Left M1/SMA", "Left DLPFC (bi-frontal)"],
        ],
        "table9": [
            ["Clinical Target", "TPS Protocol", "Phenotype"],
            ["L-DLPFC excitability enhancement", "FT1", "ME, CG, AN"],
            ["Bilateral CEN network normalization", "FT3", "TR, PM"],
            ["Limbic-OFC circuit modulation", "FT4", "AT, AX"],
            ["Salience network downregulation", "FT5", "AX, TR"],
            ["vmPFC / subgenual ACC targeting", "FT7", "ME, AT"],
            ["Ventral striatum (anhedonia circuit)", "FT8", "AN"],
            ["Dorsal ACC / cognitive control", "FT9", "CG"],
            ["Bilateral prefrontal augmentation", "Multiple", "TR (TMS-naive)"],
            ["Multi-network distributed depression", "Multiple", "MX / severe"],
        ],
        "table12": [
            ["Domain", "Details"],
            ["Indications", "Anhedonia, autonomic dysregulation, sleep disruption, TR depression"],
            ["Protocol", "20–30 min daily, 25 Hz, 250 µs, left tragus, 0.5 mA"],
            ["Add taVNS?", "□ Yes  □ No | Rationale: __________"],
        ],
        "table13": [
            ["Domain", "Details"],
            ["Indications", "Anxiety-depression comorbidity, insomnia, TR phenotype"],
            ["Protocol", "20–60 min daily, Alpha-Stim AID, adjunct to all phenotypes"],
            ["Add CES?", "□ Yes  □ No | Rationale: __________"],
        ],
        "table14": [
            ["Option", "Sequence", "Phenotype"],
            ["OPTION A — CEN-Primary", "Wk 1–4: L-DLPFC tDCS (5×/wk) → Wk 2–6: TPS L-DLPFC/vmPFC", "ME / CG / AN"],
            ["OPTION B — Limbic-Affective", "Wk 1–4: R-DLPFC cathodal tDCS + Daily taVNS → Wk 3–6: TPS Limbic/SN", "AT / AX"],
            ["OPTION C — Treatment-Resistant", "Wk 1–2: CES/taVNS stabilization → Wk 2–6: Bilateral tDCS + TPS bilateral", "TR"],
            ["OPTION D — Mixed/Psychomotor", "Sequential: TPS (highest burden network) → Maintenance L-DLPFC tDCS", "PM / MX"],
        ],
        "table15_row1": ["Reassess SOZO PRS + PHQ-9/MADRS", "□ Completed | Score: ___ / baseline: ___"],
        "table18_phenotype": "MDD Phenotype",
    },

    # ── ANXIETY (GAD) ─────────────────────────────────────────────────────
    {
        "slug": "anxiety",
        "short": "Anxiety",
        "abbr": "GAD",
        "full_name": "Generalized Anxiety Disorder",
        "severity_field": "GAD-7 / HAM-A Severity",
        "duration_field": "Disorder Duration",
        "phenotype_codes": "□ CW  □ ST  □ MX  □ PA  □ SO  □ IN  □ AU",
        "group1_label": "3.1 Cognitive-Worry Phenotypes (CW / SO / MX)",
        "group2_label": "3.2 Somatic-Autonomic Phenotypes (ST / AU)",
        "group3_label": "3.3 Panic / Insomnia Dominant (PA / IN)",
        "opt_a": "Cognitive-Worry",
        "opt_b": "Somatic",
        "opt_c": "Panic/Insomnia",
        "opt_d": "Mixed",
        "tps_offlabel": "TPS is OFF-LABEL for GAD. Requires Doctor authorisation and off-label consent.",
        "table1": [
            ["Data Point", "Completed", "Source"],
            ["GAD-7 / HAM-A severity score", "□", "Rating Scales"],
            ["Cognitive worry vs somatic tension profiling", "□", "Clinical Interview"],
            ["Panic attack history and frequency", "□", "Clinical Interview"],
            ["Social anxiety screening (LSAS)", "□", "LSAS / Interview"],
            ["Insomnia assessment (ISI, actigraphy)", "□", "ISI / Sleep Diary"],
            ["Autonomic symptom inventory (palpitations, sweating)", "□", "SOZO PRS"],
            ["Comorbid depression screening (PHQ-9)", "□", "PHQ-9"],
            ["6-Network Bedside Assessment scores", "□", "Assessment Form"],
        ],
        "table2": [
            ["Dominant Features", "Assign Phenotype", "Code"],
            ["Excessive uncontrollable worry, cognitive rumination, concentration impairment", "Cognitive-Worry (CW)", "CW"],
            ["Muscle tension, somatic complaints, physical restlessness dominant", "Somatic-Tension (ST)", "ST"],
            ["Mixed cognitive and somatic features equally prominent", "Mixed (MX)", "MX"],
            ["Discrete panic attacks with autonomic surge", "Panic-Subtype (PA)", "PA"],
            ["Social evaluation fear, avoidance, performance anxiety dominant", "Social (SO)", "SO"],
            ["Sleep-onset difficulty, hyperarousal, nocturnal rumination", "Insomnia-Subtype (IN)", "IN"],
            ["Prominent autonomic instability: HR variability, GI, sweating", "Autonomic (AU)", "AU"],
        ],
        "table3": [
            ["Phenotype", "Primary Network", "Secondary Network", "Tertiary"],
            ["CW", "CEN (L-DLPFC / ACC)", "DMN (vmPFC)", "SN (dACC)"],
            ["ST", "SN (Insula/dACC)", "SMN", "Limbic (BLA)"],
            ["MX", "SN (dACC)", "CEN (R-DLPFC cathodal)", "Limbic"],
            ["PA", "SN (Insula/BLA)", "Limbic (amygdala)", "CEN"],
            ["SO", "Limbic (BLA-vmPFC)", "CEN (R-DLPFC)", "SN"],
            ["IN", "SN (arousal)", "CEN (R-DLPFC cathodal)", "DMN"],
            ["AU", "SN (Insula)", "SMN (autonomic cortex)", "Limbic"],
        ],
        "table5": [
            ["Goal", "Anode", "Cathode"],
            ["Reduce right DLPFC hyperactivity (CW/MX)", "Left DLPFC", "Right DLPFC (cathodal)"],
            ["Enhance inhibitory prefrontal control (SO)", "Left DLPFC", "Right supraorbital"],
            ["Bilateral prefrontal balancing", "Left DLPFC", "Right DLPFC (bi-cephalic)"],
        ],
        "table6": [
            ["Goal", "Anode", "Cathode"],
            ["Downregulate SN/insula hyperactivity (ST/AU)", "Right DLPFC cathodal", "Left DLPFC anodal"],
            ["vmPFC modulation for fear extinction (PA)", "vmPFC (Fpz)", "Right mastoid"],
        ],
        "table7": [
            ["Goal", "Anode", "Cathode"],
            ["Reduce nocturnal hyperarousal (IN)", "Right DLPFC cathodal", "Left frontal"],
            ["Autonomic regulation (AU/ST)", "Left DLPFC", "Extracephalic (shoulder)"],
        ],
        "table9": [
            ["Clinical Target", "TPS Protocol", "Phenotype"],
            ["R-DLPFC inhibition / L-DLPFC excitation", "FT1", "CW, MX"],
            ["SN / dACC downregulation", "FT3", "ST, AU"],
            ["Amygdala-vmPFC fear circuit", "FT4", "PA, SO"],
            ["Insula hyperactivity suppression", "FT5", "ST, AU, PA"],
            ["vmPFC-hippocampal extinction circuit", "FT7", "PA, SO"],
            ["CEN bilateral normalization", "FT8", "TR-GAD, MX"],
            ["Thalamic arousal modulation", "FT9", "IN"],
            ["Right frontal cathodal targeting", "Multiple", "CW, IN"],
            ["Multi-network distributed anxiety", "Multiple", "MX / severe"],
        ],
        "table12": [
            ["Domain", "Details"],
            ["Indications", "Autonomic dysregulation, panic, sleep onset difficulty, GAD with CES non-response"],
            ["Protocol", "20–30 min daily, 25 Hz, 250 µs, left tragus, 0.5 mA"],
            ["Add taVNS?", "□ Yes  □ No | Rationale: __________"],
        ],
        "table13": [
            ["Domain", "Details"],
            ["Indications", "All GAD phenotypes — first-line adjunct for anxiety and insomnia"],
            ["Protocol", "20–60 min daily, Alpha-Stim AID, especially IN and AU phenotypes"],
            ["Add CES?", "□ Yes  □ No | Rationale: __________"],
        ],
        "table14": [
            ["Option", "Sequence", "Phenotype"],
            ["OPTION A — Cognitive-Worry", "Wk 1–4: R-DLPFC cathodal tDCS (5×/wk) → Wk 2–6: TPS CEN/SN", "CW / MX"],
            ["OPTION B — Somatic-Autonomic", "Wk 1–4: L-DLPFC anodal tDCS + Daily CES → Wk 3–6: TPS SN/Insula", "ST / AU"],
            ["OPTION C — Panic/Insomnia", "Wk 1–2: CES/taVNS stabilization → Wk 2–6: tDCS vmPFC + TPS BLA", "PA / IN"],
            ["OPTION D — Mixed/Social", "Sequential: TPS (highest burden) → Maintenance bilateral tDCS", "SO / MX"],
        ],
        "table15_row1": ["Reassess SOZO PRS + GAD-7/HAM-A", "□ Completed | Score: ___ / baseline: ___"],
        "table18_phenotype": "GAD Phenotype",
    },
]


# ─── Builder ─────────────────────────────────────────────────────────────────

def build_phenotype_classification(c):
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs
    tables = doc.tables

    short = c["short"]
    abbr  = c["abbr"]
    full  = c["full_name"]

    # ── Global replacements (PD / Parkinson's variants) ──────────────────
    _global_replace(doc, "Parkinson's Disease", full)
    _global_replace(doc, "Parkinson Disease", full)
    _global_replace(doc, "Parkinson's", full)
    _global_replace(doc, "Parkinsons", full)
    # Keep "PD" replacement last and scoped (abbr) — avoids partial hits in URLs etc.
    _global_replace(doc, " PD ", f" {abbr} ")
    _global_replace(doc, "(PD)", f"({abbr})")
    _global_replace(doc, "for PD", f"for {abbr}")

    # ── Paragraph replacements ────────────────────────────────────────────
    # [1]  '{SHORT} Phenotype Classification & Protocol Mapping'
    _para_set(paras[1], f"{short} Phenotype Classification & Protocol Mapping")

    # [7]  STEP 1 heading
    _para_set(paras[7], f"STEP 1 — {short} Phenotype Determination")

    # [8]  Instruction line (keep as-is but replace {SHORT} if present)
    _para_replace(paras[8], "{SHORT}", short)

    # [11] Phenotype Classification sub-heading
    _para_set(paras[11], f"1.2 {short} Phenotype Classification")

    # [14] Assigned phenotype checkboxes
    _para_set(paras[14], f"ASSIGNED PHENOTYPE: {c['phenotype_codes']}")

    # [19] tDCS group labels
    _para_set(paras[19], c["group1_label"])
    _para_set(paras[21], c["group2_label"])
    _para_set(paras[23], c["group3_label"])

    # [27] TPS off-label notice
    _para_set(paras[27], c["tps_offlabel"])

    # [40] Sequencing option selection line
    _para_set(
        paras[40],
        f"SELECTED OPTION: □ A ({c['opt_a']})  □ B ({c['opt_b']})  "
        f"□ C ({c['opt_c']})  □ D ({c['opt_d']})"
    )

    # ── Table[0]: Patient info — replace severity & duration rows ─────────
    t0 = tables[0]
    # row[3][0] = "H&Y Stage" → severity field
    _cell_write(t0.rows[3].cells[0], c["severity_field"])
    # row[6][0] = "Disease Duration" → duration field
    _cell_write(t0.rows[6].cells[0], c["duration_field"])

    # ── Table[1]: Data Collection Checklist ───────────────────────────────
    _replace_table(tables[1], c["table1"])

    # ── Table[2]: Phenotype Classification ───────────────────────────────
    _replace_table(tables[2], c["table2"])

    # ── Table[3]: Network Prioritization ─────────────────────────────────
    _replace_table(tables[3], c["table3"])

    # ── Table[5,6,7]: tDCS montage groups ────────────────────────────────
    _replace_table(tables[5], c["table5"])
    _replace_table(tables[6], c["table6"])
    _replace_table(tables[7], c["table7"])

    # ── Table[9]: TPS Protocol Assignment ────────────────────────────────
    _replace_table(tables[9], c["table9"])

    # ── Table[12]: taVNS ─────────────────────────────────────────────────
    _replace_table(tables[12], c["table12"])

    # ── Table[13]: CES ───────────────────────────────────────────────────
    _replace_table(tables[13], c["table13"])

    # ── Table[14]: Sequencing Options ────────────────────────────────────
    _replace_table(tables[14], c["table14"])

    # ── Table[15]: Response Evaluation — row[1] only ─────────────────────
    t15 = tables[15]
    row1_cells = t15.rows[1].cells
    for ci, val in enumerate(c["table15_row1"]):
        if ci < len(row1_cells):
            _cell_write(row1_cells[ci], val)

    # ── Table[18]: Protocol Summary — row[1][0] phenotype label ──────────
    t18 = tables[18]
    _cell_write(t18.rows[1].cells[0], c["table18_phenotype"])

    # ── Save ──────────────────────────────────────────────────────────────
    out_dir = _PROJECT_ROOT / "outputs" / "documents" / c['slug'] / "partners"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"Phenotype_Classification_Partners_{c['slug']}.docx"
    doc.save(str(out_path))
    print(f"  Saved: {out_path}  ({out_path.stat().st_size:,} bytes)")
    return out_path


# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Partners Tier - Phenotype Classification & Protocol Mapping")
    print("=" * 62)
    for cond in CONDITIONS:
        print(f"\nBuilding: {cond['full_name']} ({cond['abbr']}) ...")
        build_phenotype_classification(cond)
    print("\nDone — 2 files generated.")
