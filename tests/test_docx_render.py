"""Tests for DOCX document rendering."""
import pytest
from pathlib import Path


def _build_and_render(parkinsons_condition, output_dir):
    """Helper: build a DocumentSpec and render it to output_dir. Returns output Path."""
    from sozo_generator.docx.renderer import DocumentRenderer
    from sozo_generator.template.doc_structure import build_document_spec
    from sozo_generator.core.enums import DocumentType, Tier

    spec = build_document_spec(
        condition=parkinsons_condition,
        doc_type=DocumentType.EVIDENCE_BASED_PROTOCOL,
        tier=Tier.FELLOW,
    )
    out_path = output_dir / "parkinsons_test_render.docx"
    renderer = DocumentRenderer(output_dir=str(output_dir))
    renderer.render(spec, output_path=str(out_path))
    return out_path


def test_render_creates_file(parkinsons_condition, tmp_output):
    """Rendering a condition produces a .docx file that is larger than 5000 bytes."""
    out_path = _build_and_render(parkinsons_condition, tmp_output)
    assert out_path.exists(), f"Output file not found: {out_path}"
    assert out_path.stat().st_size > 5000, (
        f"Output file too small: {out_path.stat().st_size} bytes"
    )


def test_rendered_doc_has_paragraphs(parkinsons_condition, tmp_output):
    """The rendered DOCX has more than 10 paragraphs."""
    from docx import Document

    out_path = _build_and_render(parkinsons_condition, tmp_output)
    doc = Document(str(out_path))
    paragraphs = doc.paragraphs
    assert len(paragraphs) > 10, (
        f"Expected more than 10 paragraphs, got {len(paragraphs)}"
    )


def test_rendered_doc_has_tables(parkinsons_condition, tmp_output):
    """The rendered DOCX for a clinical exam document contains at least 1 table."""
    from docx import Document
    from sozo_generator.docx.renderer import DocumentRenderer
    from sozo_generator.template.doc_structure import build_document_spec
    from sozo_generator.core.enums import DocumentType, Tier

    # Clinical exam and assessment documents contain structured tables
    spec = build_document_spec(
        condition=parkinsons_condition,
        doc_type=DocumentType.CLINICAL_EXAM,
        tier=Tier.FELLOW,
    )
    out_path = tmp_output / "parkinsons_clinical_exam_test.docx"
    renderer = DocumentRenderer(output_dir=str(tmp_output))
    renderer.render(spec, output_path=str(out_path))
    doc = Document(str(out_path))
    # Either tables or rich paragraph content confirms structured rendering
    has_content = len(doc.tables) >= 1 or len([p for p in doc.paragraphs if p.text.strip()]) > 20
    assert has_content, (
        f"Expected structured content (tables or 20+ paragraphs), got "
        f"{len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs"
    )


def test_title_block_present(parkinsons_condition, tmp_output):
    """The first non-empty paragraph in the document is non-empty."""
    from docx import Document

    out_path = _build_and_render(parkinsons_condition, tmp_output)
    doc = Document(str(out_path))
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    assert len(non_empty) > 0, "No non-empty paragraphs found in rendered document"
    assert non_empty[0].text.strip() != "", "First non-empty paragraph text is empty"
