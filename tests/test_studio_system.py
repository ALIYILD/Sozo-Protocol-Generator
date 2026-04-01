"""Tests for the document studio system: visual generation, condition matching,
template matching, extended planner, and end-to-end job flows."""
from __future__ import annotations
import pytest
from pathlib import Path


class TestMontageGeneration:
    def test_generate_montage_diagram(self, tmp_path):
        from sozo_generator.visuals.montage_diagrams import MontageDiagramGenerator
        gen = MontageDiagramGenerator()
        path = gen.generate_montage_diagram("F3", "F4", "Depression", "DLPFC Protocol", output_dir=tmp_path)
        assert path.exists()
        assert path.stat().st_size > 10000  # real PNG

    def test_generate_target_diagram(self, tmp_path):
        from sozo_generator.visuals.montage_diagrams import MontageDiagramGenerator
        gen = MontageDiagramGenerator()
        path = gen.generate_target_diagram(["DLPFC", "M1"], "Parkinsons", output_dir=tmp_path)
        assert path.exists()

    def test_generate_protocol_card(self, tmp_path):
        from sozo_generator.visuals.montage_diagrams import MontageDiagramGenerator
        gen = MontageDiagramGenerator()
        path = gen.generate_protocol_card("Test Protocol", "tDCS", "M1",
            {"anode": "C3", "cathode": "C4", "intensity": "2mA"}, output_dir=tmp_path)
        assert path.exists()

    def test_generate_all_for_condition(self, tmp_path, parkinsons_condition):
        from sozo_generator.visuals.montage_diagrams import MontageDiagramGenerator
        gen = MontageDiagramGenerator()
        paths = gen.generate_all_for_condition(parkinsons_condition, tmp_path)
        assert len(paths) >= 3  # at least targets + some protocol cards

    def test_electrode_positions_complete(self):
        from sozo_generator.visuals.montage_diagrams import ELECTRODES_10_20
        assert len(ELECTRODES_10_20) >= 20
        # Standard positions exist
        for pos in ["Fp1", "Fp2", "F3", "F4", "C3", "C4", "P3", "P4", "O1", "O2", "Fz", "Cz", "Pz"]:
            assert pos in ELECTRODES_10_20


class TestConditionMatchAgent:
    def test_explicit_conditions(self):
        from sozo_generator.agents.condition_match_agent import ConditionMatchAgent
        agent = ConditionMatchAgent()
        result = agent.execute({"conditions": ["parkinsons", "depression"]}, "/tmp")
        assert result.success
        assert "parkinsons" in result.output_data["validated_conditions"]
        assert "depression" in result.output_data["validated_conditions"]

    def test_prompt_inference(self):
        from sozo_generator.agents.condition_match_agent import ConditionMatchAgent
        agent = ConditionMatchAgent()
        result = agent.execute({"prompt": "Generate handbook for ADHD"}, "/tmp")
        assert result.success
        assert "adhd" in result.output_data["validated_conditions"]

    def test_unknown_condition_flagged(self):
        from sozo_generator.agents.condition_match_agent import ConditionMatchAgent
        agent = ConditionMatchAgent()
        result = agent.execute({"conditions": ["narcolepsy"]}, "/tmp")
        assert "narcolepsy" in result.output_data["draft_conditions"]

    def test_all_conditions_from_prompt(self):
        from sozo_generator.agents.condition_match_agent import ConditionMatchAgent
        agent = ConditionMatchAgent()
        result = agent.execute({"prompt": "Generate all 15 conditions"}, "/tmp")
        assert len(result.output_data["validated_conditions"]) == 15


class TestTemplateMatcher:
    def test_match_returns_result(self):
        from sozo_generator.template.learning.template_matcher import TemplateMatcher
        # Create a minimal test docx
        import docx
        import tempfile
        doc = docx.Document()
        doc.add_heading("Evidence-Based Protocol", 1)
        doc.add_heading("Pathophysiology", 1)
        doc.add_heading("Safety and Monitoring", 1)
        doc.add_heading("References", 1)
        path = Path(tempfile.mkdtemp()) / "test.docx"
        doc.save(str(path))

        matcher = TemplateMatcher()
        result = matcher.match(path)
        assert result.matched_doc_type == "evidence_based_protocol"
        assert result.confidence > 0.3

    def test_low_confidence_warning(self):
        from sozo_generator.template.learning.template_matcher import TemplateMatcher
        import docx, tempfile
        doc = docx.Document()
        doc.add_heading("Random Title", 1)
        path = Path(tempfile.mkdtemp()) / "random.docx"
        doc.save(str(path))

        matcher = TemplateMatcher()
        result = matcher.match(path)
        # Should have low confidence or warning
        assert result.confidence < 0.5 or result.warnings


class TestVisualAgent:
    def test_visual_agent_runs(self, tmp_path):
        from sozo_generator.agents.visual_agent import VisualProtocolAgent
        agent = VisualProtocolAgent()
        result = agent.execute({"conditions": ["parkinsons"]}, str(tmp_path))
        assert result.success
        assert result.output_data["total_visuals"] >= 1

    def test_visual_agent_registered(self):
        from sozo_generator.agents.registry import get_agent
        # Import to trigger registration
        from sozo_generator.agents import visual_agent  # noqa: F401
        agent = get_agent("visual_agent")
        assert agent is not None


class TestExtendedPlanner:
    def test_plan_includes_condition_match(self):
        from sozo_generator.jobs.models import Job
        from sozo_generator.jobs.planner import JobPlanner
        job = Job(job_id="test", source_prompt="Generate for parkinsons")
        plan = JobPlanner().plan(job)
        names = [t.agent_name for t in plan.tasks]
        assert "condition_match_agent" in names

    def test_plan_includes_visual_agent(self):
        from sozo_generator.jobs.models import Job
        from sozo_generator.jobs.planner import JobPlanner
        job = Job(job_id="test", target_conditions=["parkinsons"])
        plan = JobPlanner().plan(job)
        names = [t.agent_name for t in plan.tasks]
        assert "visual_agent" in names

    def test_visual_after_draft_before_qa(self):
        from sozo_generator.jobs.models import Job
        from sozo_generator.jobs.planner import JobPlanner
        job = Job(job_id="test", target_conditions=["parkinsons"])
        plan = JobPlanner().plan(job)
        names = [t.agent_name for t in plan.tasks]
        if "visual_agent" in names and "draft_agent" in names and "qa_agent" in names:
            assert names.index("visual_agent") > names.index("draft_agent")
            assert names.index("qa_agent") > names.index("visual_agent")


class TestEndToEndJob:
    @pytest.mark.slow
    def test_full_job_parkinsons_handbook(self, tmp_path):
        """Full job: create -> plan -> execute -> check artifacts."""
        from sozo_generator.jobs.manager import JobManager
        from sozo_generator.jobs.planner import JobPlanner
        from sozo_generator.agents.executor import AgentExecutor
        from sozo_generator.jobs.models import JobStatus
        from sozo_generator.agents import (  # noqa: F401
            template_agent, content_agent, evidence_agent,
            draft_agent, qa_agent, review_pack_agent,
            visual_agent, condition_match_agent,
        )

        mgr = JobManager(tmp_path / "jobs", tmp_path / "ws")
        job = mgr.create_job(
            source_prompt="Generate handbook for parkinsons",
            target_conditions=["parkinsons"],
            target_doc_types=["handbook"],
            target_tiers=["fellow"],
        )
        plan = JobPlanner().plan(job)
        job.plan = plan
        mgr._save_job(job)

        executor = AgentExecutor(mgr)
        result = executor.execute_job(job)

        assert result.status in (JobStatus.AWAITING_REVIEW, JobStatus.COMPLETED)
        assert result.total_documents >= 1
        # Check for DOCX artifacts
        docx_arts = [a for a in result.artifacts if a.get("type") == "document"]
        assert len(docx_arts) >= 1
        # Check for visual artifacts
        vis_arts = [a for a in result.artifacts if a.get("type") == "visual"]
        assert len(vis_arts) >= 1


class TestRegressions:
    def test_existing_chat_engine_works(self, tmp_path):
        from sozo_generator.ai.chat_engine import ChatEngine
        engine = ChatEngine(output_dir=str(tmp_path))
        r = engine.process_message("list conditions")
        assert r.success

    def test_existing_qa_works(self):
        from sozo_generator.qa.engine import QAEngine
        from sozo_generator.conditions.registry import get_registry
        engine = QAEngine()
        pk = get_registry().get("parkinsons")
        report = engine.run_condition_qa(pk)
        report.compute_counts()
        assert report.block_count == 0

    def test_app_syntax(self):
        import ast
        with open("app.py") as f:
            ast.parse(f.read())
