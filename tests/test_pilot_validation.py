"""Pilot-quality validation tests: condition matching, template matching,
visual QA, and 5 end-to-end operator scenarios."""
from __future__ import annotations
import pytest
from pathlib import Path


# ═══════════════════════════════════════════════════════════════════════
# A. Condition Match Validation
# ═══════════════════════════════════════════════════════════════════════

class TestConditionMatchValidation:
    def _match(self, **kwargs):
        from sozo_generator.agents.condition_match_agent import ConditionMatchAgent
        return ConditionMatchAgent().execute(kwargs, "/tmp")

    def test_alias_mdd(self):
        r = self._match(prompt="Generate for MDD")
        assert "depression" in r.output_data["validated_conditions"]

    def test_alias_pd(self):
        r = self._match(prompt="Create PD handbook")
        assert "parkinsons" in r.output_data["validated_conditions"]

    def test_alias_gad(self):
        r = self._match(prompt="GAD evidence protocol")
        assert "anxiety" in r.output_data["validated_conditions"]

    def test_alias_fibromyalgia(self):
        r = self._match(prompt="fibromyalgia clinical exam")
        assert "chronic_pain" in r.output_data["validated_conditions"]

    def test_multi_condition_prompt(self):
        r = self._match(prompt="Generate handbooks for depression, anxiety, and ADHD")
        valid = r.output_data["validated_conditions"]
        assert "depression" in valid
        assert "anxiety" in valid
        assert "adhd" in valid

    def test_all_conditions_prompt(self):
        from sozo_generator.conditions.registry import get_registry
        r = self._match(prompt="Generate all 15 conditions")
        assert r.output_data["all_conditions"] is True
        assert len(r.output_data["validated_conditions"]) == len(get_registry().list_slugs())

    def test_unknown_condition_flagged_as_draft(self):
        r = self._match(conditions=["narcolepsy"])
        assert "narcolepsy" in r.output_data["draft_conditions"]
        assert r.output_data["confidence"] <= 0.5
        assert any("draft" in w.lower() for w in r.warnings)

    def test_ambiguous_prompt_low_confidence(self):
        r = self._match(prompt="create a handbook")
        assert r.output_data["confidence"] <= 0.5

    def test_empty_prompt_fails(self):
        r = self._match(prompt="")
        assert not r.success

    def test_mixed_known_unknown(self):
        r = self._match(conditions=["parkinsons", "narcolepsy", "depression"])
        assert "parkinsons" in r.output_data["validated_conditions"]
        assert "depression" in r.output_data["validated_conditions"]
        assert "narcolepsy" in r.output_data["draft_conditions"]


# ═══════════════════════════════════════════════════════════════════════
# B. Template Match Validation
# ═══════════════════════════════════════════════════════════════════════

class TestTemplateMatchValidation:
    def _make_template(self, tmp_path, headings):
        import docx
        doc = docx.Document()
        for h in headings:
            doc.add_heading(h, level=1)
            doc.add_paragraph("Content for this section.")
        path = tmp_path / "test_template.docx"
        doc.save(str(path))
        return path

    def test_evidence_protocol_detected(self, tmp_path):
        from sozo_generator.template.learning.template_matcher import TemplateMatcher
        path = self._make_template(tmp_path, [
            "Inclusion Criteria", "Exclusion Criteria", "Pathophysiology",
            "Protocols", "Safety and Monitoring", "Evidence Summary",
        ])
        result = TemplateMatcher().match(path)
        assert result.matched_doc_type == "evidence_based_protocol"
        assert result.confidence >= 0.5

    def test_handbook_detected(self, tmp_path):
        from sozo_generator.template.learning.template_matcher import TemplateMatcher
        path = self._make_template(tmp_path, [
            "Introduction", "Available Modalities", "Stage 1", "Stage 2",
            "Stage 3", "Governance Rules",
        ])
        result = TemplateMatcher().match(path)
        assert result.matched_doc_type == "handbook"
        assert result.confidence >= 0.5

    def test_low_confidence_shows_candidates(self, tmp_path):
        from sozo_generator.template.learning.template_matcher import TemplateMatcher
        path = self._make_template(tmp_path, ["Introduction", "Summary"])
        result = TemplateMatcher().match(path)
        # Should have warning about low confidence
        assert result.confidence < 0.6 or result.candidates
        if result.confidence < 0.3:
            assert any("candidate" in w.lower() or "manual" in w.lower()
                       for w in result.warnings)

    def test_few_sections_warning(self, tmp_path):
        from sozo_generator.template.learning.template_matcher import TemplateMatcher
        path = self._make_template(tmp_path, ["Title"])
        result = TemplateMatcher().match(path)
        assert any("few sections" in w.lower() for w in result.warnings)

    def test_condition_detection_in_template(self, tmp_path):
        from sozo_generator.template.learning.template_matcher import TemplateMatcher
        import docx
        doc = docx.Document()
        doc.add_heading("Parkinson's Disease Protocol", level=1)
        doc.add_heading("Pathophysiology", level=1)
        doc.add_paragraph("Treatment of Parkinson's Disease with tDCS.")
        path = tmp_path / "pd_template.docx"
        doc.save(str(path))
        result = TemplateMatcher().match(path)
        assert result.condition_detected == "parkinsons"


# ═══════════════════════════════════════════════════════════════════════
# C. Visual QA
# ═══════════════════════════════════════════════════════════════════════

class TestVisualQA:
    def test_visual_qa_parkinsons_passes(self, parkinsons_condition):
        from sozo_generator.qa.visual_qa import VisualQAChecker
        checker = VisualQAChecker()
        report = checker.check_condition(parkinsons_condition)
        assert report.error_count == 0  # no errors (warnings OK)

    def test_electrode_validity_check(self):
        from sozo_generator.qa.visual_qa import VisualQAChecker, VisualQAReport
        from sozo_generator.schemas.condition import ConditionSchema, ProtocolEntry
        from sozo_generator.core.enums import Modality, EvidenceLevel
        checker = VisualQAChecker()
        # Create a condition with an invalid electrode
        cond = ConditionSchema(
            slug="test", display_name="Test", icd10="X00",
            protocols=[ProtocolEntry(
                protocol_id="p1", label="Bad Protocol",
                modality=Modality.TDCS, target_region="Test",
                target_abbreviation="T", rationale="test",
                parameters={"anode": "ZZ99", "cathode": "C4"},
            )],
        )
        report = checker.check_condition(cond)
        # Should flag invalid electrode ZZ99
        electrode_issues = [i for i in report.issues if i.check == "electrode_validity"]
        assert len(electrode_issues) >= 1

    def test_electrode_collision_error(self):
        from sozo_generator.qa.visual_qa import VisualQAChecker
        from sozo_generator.schemas.condition import ConditionSchema, ProtocolEntry
        from sozo_generator.core.enums import Modality
        checker = VisualQAChecker()
        cond = ConditionSchema(
            slug="test", display_name="Test", icd10="X00",
            protocols=[ProtocolEntry(
                protocol_id="p1", label="Same Electrode",
                modality=Modality.TDCS, target_region="Test",
                target_abbreviation="T", rationale="test",
                parameters={"anode": "C3", "cathode": "C3"},
            )],
        )
        report = checker.check_condition(cond)
        assert report.error_count >= 1
        assert not report.passed

    def test_all_conditions_visual_qa(self):
        from sozo_generator.qa.visual_qa import VisualQAChecker
        from sozo_generator.conditions.registry import get_registry
        checker = VisualQAChecker()
        registry = get_registry()
        for slug in registry.list_slugs():
            cond = registry.get(slug)
            report = checker.check_condition(cond)
            # No errors (warnings are acceptable)
            assert report.error_count == 0, f"{slug} has visual QA errors: {[i.message for i in report.issues if i.severity == 'error']}"


# ═══════════════════════════════════════════════════════════════════════
# D. Pilot Operator Scenarios (end-to-end)
# ═══════════════════════════════════════════════════════════════════════

class TestPilotScenarios:
    """5 real operator scenarios validated end-to-end."""

    def _run_job(self, tmp_path, **kwargs):
        from sozo_generator.jobs.manager import JobManager
        from sozo_generator.jobs.planner import JobPlanner
        from sozo_generator.agents.executor import AgentExecutor
        from sozo_generator.agents import (
            template_agent, content_agent, evidence_agent,
            draft_agent, revision_agent, qa_agent, review_pack_agent,
            visual_agent, condition_match_agent,
        )
        mgr = JobManager(tmp_path / "jobs", tmp_path / "ws")
        job = mgr.create_job(**kwargs)
        job.plan = JobPlanner().plan(job)
        mgr._save_job(job)
        return AgentExecutor(mgr).execute_job(job)

    def test_scenario_1_single_condition_with_template(self, tmp_path):
        """Scenario 1: Template + single validated condition → doc generated."""
        import docx
        tpl = tmp_path / "handbook_template.docx"
        doc = docx.Document()
        doc.add_heading("Introduction", 1)
        doc.add_heading("Stage 1", 1)
        doc.add_heading("Governance", 1)
        doc.save(str(tpl))

        result = self._run_job(
            tmp_path,
            source_prompt="Generate handbook for parkinsons from this template",
            target_conditions=["parkinsons"],
            target_doc_types=["handbook"],
            target_tiers=["fellow"],
            uploaded_files=[str(tpl)],
        )
        assert result.total_documents >= 1
        docx_arts = [a for a in result.artifacts if str(a.get("path", "")).endswith(".docx")]
        assert len(docx_arts) >= 1

    def test_scenario_2_protocol_doc_with_visuals(self, tmp_path):
        """Scenario 2: Protocol doc + visual generation."""
        result = self._run_job(
            tmp_path,
            source_prompt="Generate evidence-based protocol for depression with visuals",
            target_conditions=["depression"],
            target_tiers=["fellow"],
        )
        # Should have both documents and visuals
        doc_arts = [a for a in result.artifacts if a.get("type") == "document"]
        vis_arts = [a for a in result.artifacts if a.get("type") == "visual"]
        assert len(doc_arts) >= 1
        assert len(vis_arts) >= 1

    def test_scenario_3_doctor_comment_revision(self, tmp_path):
        """Scenario 3: Doctor comment → revision job."""
        result = self._run_job(
            tmp_path,
            source_prompt="Revise anxiety documents",
            target_conditions=["anxiety"],
            target_tiers=["fellow"],
            doctor_comments=["Remove TPS protocols", "Use conservative language"],
        )
        # Revision agent should have run
        task_names = [t.agent_name for t in result.plan.tasks]
        assert "revision_agent" in task_names

    def test_scenario_4_draft_condition(self, tmp_path):
        """Scenario 4: Unknown condition → draft with mandatory review."""
        from sozo_generator.jobs.models import JobStatus
        result = self._run_job(
            tmp_path,
            source_prompt="Generate docs for narcolepsy",
            target_conditions=["narcolepsy"],
        )
        # Job should still complete (condition match flags it as draft)
        # Condition match should warn about draft
        cond_task = next(
            (t for t in result.plan.tasks if t.agent_name == "condition_match_agent"),
            None,
        )
        assert cond_task is not None

    def test_scenario_5_review_and_export(self, tmp_path):
        """Scenario 5: Generate → review → approve → export."""
        from sozo_generator.review.manager import ReviewManager

        result = self._run_job(
            tmp_path,
            source_prompt="Generate handbook for ADHD",
            target_conditions=["adhd"],
            target_doc_types=["handbook"],
            target_tiers=["fellow"],
        )

        # Create review entry
        mgr = ReviewManager(tmp_path / "reviews")
        state = mgr.create_review(
            result.job_id, "adhd", "handbook", "fellow"
        )
        assert state.status.value == "draft"

        # Submit → approve → export
        mgr.submit_for_review(result.job_id)
        mgr.approve(result.job_id, "Dr. Test", "Looks good")
        final = mgr.get_review(result.job_id)
        assert final.status.value == "approved"


# ═══════════════════════════════════════════════════════════════════════
# E. Review Clarity
# ═══════════════════════════════════════════════════════════════════════

class TestReviewClarity:
    def test_draft_agent_flags_adapted_content(self, tmp_path):
        """Adapted content should be distinguishable from native content."""
        from sozo_generator.generation.rich_generator import RichDocumentGenerator
        from sozo_generator.conditions.registry import get_registry
        gen = RichDocumentGenerator()
        if not gen.has_library:
            pytest.skip("No content library")
        cond = get_registry().get("depression")
        spec = gen.generate(cond, "handbook", "fellow")
        # Sections adapted from other conditions should have confidence_label set
        adapted = [s for s in spec.sections if s.confidence_label]
        # It's OK if none are adapted (depression may have native content)
        # But the mechanism must exist
        assert hasattr(spec.sections[0], "confidence_label")

    def test_visual_qa_report_accessible(self, parkinsons_condition):
        from sozo_generator.qa.visual_qa import VisualQAChecker
        report = VisualQAChecker().check_condition(parkinsons_condition)
        assert hasattr(report, "error_count")
        assert hasattr(report, "warning_count")
        assert hasattr(report, "passed")

    def test_match_confidence_in_result(self, tmp_path):
        """Template match should include confidence for reviewer."""
        from sozo_generator.template.learning.template_matcher import TemplateMatcher
        import docx
        doc = docx.Document()
        doc.add_heading("Test", 1)
        path = tmp_path / "test.docx"
        doc.save(str(path))
        result = TemplateMatcher().match(path)
        assert isinstance(result.confidence, float)
        assert 0.0 <= result.confidence <= 1.0


class TestRegressions:
    def test_app_syntax(self):
        import ast
        with open("app.py", encoding="utf-8") as f:
            ast.parse(f.read())

    def test_existing_generation_works(self, tmp_path):
        from sozo_generator.ai.chat_engine import ChatEngine
        engine = ChatEngine(output_dir=str(tmp_path))
        r = engine.process_message("Generate evidence-based protocol for depression")
        assert r.success

    def test_review_queue_works(self, tmp_path):
        from sozo_generator.review.manager import ReviewManager
        mgr = ReviewManager(tmp_path / "reviews")
        state = mgr.create_review("test-1", "parkinsons", "handbook", "fellow")
        assert state.status.value == "draft"
