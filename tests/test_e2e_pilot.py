"""
End-to-end integration tests + golden dataset for pilot readiness.

Tests the full pipeline: generate → validate safety → check evidence → export → verify output.
"""
import json
import pytest
from pathlib import Path


class TestSafetyValidator:
    def test_validate_parkinsons(self):
        from sozo_generator.knowledge.safety import validate_condition_protocols
        report = validate_condition_protocols("parkinsons")
        assert report is not None
        assert report.passed
        assert report.total_checks >= 5

    def test_validate_all_conditions(self):
        from sozo_generator.knowledge.safety import validate_condition_protocols
        from sozo_generator.knowledge.base import KnowledgeBase
        kb = KnowledgeBase()
        kb.load_all()
        for slug in kb.list_conditions():
            report = validate_condition_protocols(slug)
            assert report is not None
            assert report.passed, f"Safety BLOCKED for {slug}"

    def test_off_label_tps_warning(self):
        from sozo_generator.knowledge.safety import validate_condition_protocols
        report = validate_condition_protocols("parkinsons")
        off_label = [c for c in report.checks if c.check_type == "off_label"]
        assert len(off_label) >= 1
        assert "OFF-LABEL" in off_label[0].message

    def test_dangerous_parameter_blocked(self):
        from sozo_generator.knowledge.safety import SafetyValidator
        validator = SafetyValidator()
        report = validator.validate_protocol(
            condition_slug="test",
            protocols=[{"modality": "tdcs", "label": "bad", "parameters": {"intensity_ma": 50.0}}],
        )
        blocked = [c for c in report.checks if c.severity == "block"]
        assert len(blocked) >= 1


class TestExportHardening:
    @pytest.fixture(scope="class")
    def pd_docx_path(self):
        from sozo_generator.generation.service import GenerationService
        svc = GenerationService(with_visuals=False, with_qa=False)
        result = svc.generate_canonical("parkinsons", "evidence_based_protocol", "fellow")
        assert result.success
        return result.output_path

    def test_docx_has_disclaimer(self, pd_docx_path):
        from docx import Document
        doc = Document(pd_docx_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Clinical Decision Support Disclaimer" in text

    def test_docx_has_provenance(self, pd_docx_path):
        from docx import Document
        doc = Document(pd_docx_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Condition: parkinsons" in text

    def test_docx_has_references(self, pd_docx_path):
        from docx import Document
        doc = Document(pd_docx_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "PMID" in text

    def test_docx_substantial(self, pd_docx_path):
        from docx import Document
        doc = Document(pd_docx_path)
        assert len(doc.paragraphs) >= 50


class TestGoldenDataset:
    @pytest.mark.parametrize("condition", ["parkinsons", "depression", "adhd", "asd", "migraine"])
    def test_full_pipeline(self, condition):
        from sozo_generator.generation.service import GenerationService
        from sozo_generator.knowledge.safety import validate_condition_protocols

        svc = GenerationService(with_visuals=False, with_qa=False)
        result = svc.generate_canonical(condition, "evidence_based_protocol", "fellow")
        assert result.success, f"{condition} generation failed: {result.error}"

        safety = validate_condition_protocols(condition)
        assert safety.passed, f"{condition} safety blocked"

        prov_path = Path(result.output_path).with_suffix(".provenance.json")
        assert prov_path.exists()
        prov = json.loads(prov_path.read_text())
        assert prov["readiness"] in ("ready", "review_required")
        assert Path(result.output_path).stat().st_size > 20000

    @pytest.mark.parametrize("condition", ["parkinsons", "depression", "adhd", "asd", "migraine"])
    def test_evidence_traceability(self, condition):
        from sozo_generator.knowledge.base import KnowledgeBase
        from sozo_generator.knowledge.assembler import CanonicalDocumentAssembler

        kb = KnowledgeBase()
        kb.load_all()
        assembler = CanonicalDocumentAssembler(kb)
        _, prov = assembler.assemble(condition, "evidence_based_protocol", "fellow")

        for sec in prov.sections:
            if sec.evidence_criticality in ("critical", "recommended"):
                assert sec.actual_pmid_count > 0, (
                    f"{condition}/{sec.section_slug}: {sec.evidence_criticality} but 0 PMIDs"
                )


class TestRegressionSafety:
    def test_all_16_generate(self):
        from sozo_generator.generation.service import GenerationService
        from sozo_generator.knowledge.base import KnowledgeBase

        kb = KnowledgeBase()
        kb.load_all()
        svc = GenerationService(with_visuals=False, with_qa=False)
        for slug in kb.list_conditions():
            result = svc.generate_canonical(slug, "evidence_based_protocol", "fellow")
            assert result.success, f"REGRESSION: {slug} failed"

    def test_all_16_pass_safety(self):
        from sozo_generator.knowledge.safety import validate_condition_protocols
        from sozo_generator.knowledge.base import KnowledgeBase

        kb = KnowledgeBase()
        kb.load_all()
        for slug in kb.list_conditions():
            report = validate_condition_protocols(slug)
            assert report.passed, f"REGRESSION: {slug} safety blocked"

    def test_knowledge_valid(self):
        from sozo_generator.knowledge.validate import validate_all
        report = validate_all()
        assert report.passed
