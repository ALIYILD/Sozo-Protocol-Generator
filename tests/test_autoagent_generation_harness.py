"""Tests for AutoAgent-Clinical GenerationService harness adapter."""
from __future__ import annotations

import os
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from autoagent_clinical.harnesses.generation_service_harness import GenerationServiceHarness
from autoagent_clinical.registry import get_harness, list_harnesses, register_harness
from autoagent_clinical.runner import load_all_benchmark_files, run_suite
from autoagent_clinical.schemas import BenchmarkCase, BenchmarkExpectations, BenchmarkSuite


def test_generation_service_harness_registered() -> None:
    assert "generation_service" in list_harnesses()
    assert get_harness("generation_service").harness_id == "generation_service"


def test_docx_normalize_emits_atx_headings(tmp_path: Path) -> None:
    from autoagent_clinical.normalize import docx_to_benchmark_markdown

    p = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Supporting text with PMID: 12345678.")
    doc.save(p)
    md = docx_to_benchmark_markdown(p)
    assert "# Overview" in md
    assert "PMID: 12345678" in md


def test_generation_service_mock_success(tmp_path: Path) -> None:
    from sozo_generator.generation.service import GenerationResult

    docx = tmp_path / "out.docx"
    d = Document()
    d.add_heading("Stimulation parameters", level=2)
    d.add_paragraph("Example 1 mA tDCS for internal QA. PMID: 99999999")
    d.save(docx)

    def factory():
        svc = MagicMock()

        def _gen(**_kwargs):
            return [
                GenerationResult(
                    condition_slug="parkinsons",
                    tier="fellow",
                    doc_type="evidence_based_protocol",
                    success=True,
                    output_path=str(docx),
                    build_id="mock-build",
                )
            ]

        svc.generate.side_effect = _gen
        return svc

    h = GenerationServiceHarness(service_factory=factory)
    case = BenchmarkCase(
        id="mock_ok",
        category="generation_adapter",
        title="mock",
        inputs={
            "condition": "parkinsons",
            "tier": "fellow",
            "doc_type": "evidence_based_protocol",
            "relaxed_citation_check": True,
        },
        fixture_markdown="# ignore",
        expectations=BenchmarkExpectations(),
    )
    res = h.run(case)
    assert res.case_id == "mock_ok"
    assert "# Generation failed" not in res.output_text
    assert "## Stimulation parameters" in res.output_text
    assert any("parkinsons" in n for n in res.metadata.notes)


def test_generation_service_mock_failure_surfaces_error_doc() -> None:
    from sozo_generator.generation.service import GenerationResult

    def factory():
        svc = MagicMock()
        svc.generate.return_value = [
            GenerationResult(
                condition_slug="__x__",
                tier="fellow",
                doc_type="evidence_based_protocol",
                success=False,
                error="Knowledge base not available",
            )
        ]
        return svc

    h = GenerationServiceHarness(service_factory=factory)
    case = BenchmarkCase(
        id="mock_fail",
        category="generation_adapter",
        title="mock",
        inputs={"condition": "x", "tier": "fellow", "doc_type": "evidence_based_protocol"},
        fixture_markdown="stub",
        expectations=BenchmarkExpectations(),
    )
    res = h.run(case)
    assert "# Generation failed" in res.output_text
    assert "Knowledge base not available" in res.output_text


def test_generation_service_missing_condition_field() -> None:
    h = GenerationServiceHarness(service_factory=lambda: MagicMock())
    case = BenchmarkCase(
        id="no_cond",
        category="generation_adapter",
        title="x",
        inputs={"tier": "fellow"},
        fixture_markdown="stub",
        expectations=BenchmarkExpectations(),
    )
    res = h.run(case)
    assert "# Generation failed" in res.output_text
    assert "condition" in res.output_text.lower()


def test_adapter_yaml_excluded_by_default() -> None:
    a = load_all_benchmark_files(include_adapter_benchmarks=False)
    b = load_all_benchmark_files(include_adapter_benchmarks=True)
    assert len(b.cases) == len(a.cases) + 2
    assert all(c.id != "gen_pd_fellow_ebp_v1" for c in a.cases)


def test_regression_gate_yaml_excluded_by_default() -> None:
    base = load_all_benchmark_files()
    with_gates = load_all_benchmark_files(include_regression_gates=True)
    assert len(with_gates.cases) == len(base.cases) + 2
    assert all(c.id != "gate_pd_ebp_relaxed_invariants" for c in base.cases)


def test_opt_in_suite_sizes() -> None:
    base = load_all_benchmark_files()
    adapter = load_all_benchmark_files(include_adapter_benchmarks=True)
    gates = load_all_benchmark_files(include_regression_gates=True)
    both = load_all_benchmark_files(
        include_adapter_benchmarks=True,
        include_regression_gates=True,
    )
    assert len(adapter.cases) == len(base.cases) + 2
    assert len(gates.cases) == len(base.cases) + 2
    assert len(both.cases) == len(base.cases) + 4


@pytest.mark.integration
@pytest.mark.live_generation
def test_live_generation_service_harness_parkinsons(tmp_path: Path) -> None:
    """In-process GenerationService (requires knowledge base + data on disk)."""
    if os.environ.get("SOZO_AUTOAGENT_LIVE") != "1":
        pytest.skip("Set SOZO_AUTOAGENT_LIVE=1 for live GenerationService tests")

    def factory():
        from sozo_generator.generation.service import GenerationService

        return GenerationService(
            output_dir=str(tmp_path),
            with_visuals=False,
            with_qa=False,
        )

    h = GenerationServiceHarness(service_factory=factory)
    case = BenchmarkCase(
        id="live_pd",
        category="generation_adapter",
        title="live",
        inputs={
            "condition": "parkinsons",
            "tier": "fellow",
            "doc_type": "evidence_based_protocol",
            "relaxed_citation_check": True,
        },
        fixture_markdown="# stub",
        expectations=BenchmarkExpectations(),
    )
    res = h.run(case)
    if "# Generation failed" in res.output_text:
        pytest.skip(f"generation unavailable in this environment: {res.metadata.notes[:3]}")
    assert len(res.output_text) > 200
    assert "#" in res.output_text


def test_runner_with_injected_generation_harness(tmp_path: Path) -> None:
    from sozo_generator.generation.service import GenerationResult

    docx = tmp_path / "bench.docx"
    d = Document()
    d.add_heading("Overview", level=1)
    d.add_paragraph("Content PMID: 88888888")
    d.save(docx)

    def factory():
        svc = MagicMock()

        def _gen(**_kwargs):
            return [
                GenerationResult(
                    condition_slug="parkinsons",
                    tier="fellow",
                    doc_type="evidence_based_protocol",
                    success=True,
                    output_path=str(docx),
                    build_id="inj",
                )
            ]

        svc.generate.side_effect = _gen
        return svc

    register_harness("generation_service_mocked", lambda: GenerationServiceHarness(service_factory=factory))
    try:
        suite = BenchmarkSuite(
            cases=[
                BenchmarkCase(
                    id="inj1",
                    category="generation_adapter",
                    title="injected",
                    inputs={
                        "condition": "parkinsons",
                        "tier": "fellow",
                        "doc_type": "evidence_based_protocol",
                        "relaxed_citation_check": True,
                    },
                    fixture_markdown="# x",
                    expectations=BenchmarkExpectations(),
                )
            ]
        )
        report = run_suite(suite, "generation_service_mocked")
        r0 = report.case_results[0]
        assert r0.case_id == "inj1"
        assert len(r0.validator_reports) >= 1
        h = GenerationServiceHarness(service_factory=factory)
        hr = h.run(suite.cases[0])
        assert "# Overview" in hr.output_text
    finally:
        from autoagent_clinical import registry as reg_module

        reg_module._REGISTRY.pop("generation_service_mocked", None)


def test_mock_failure_propagates_through_runner() -> None:
    from sozo_generator.generation.service import GenerationResult

    def factory():
        svc = MagicMock()
        svc.generate.return_value = [
            GenerationResult(
                condition_slug="z",
                tier="fellow",
                doc_type="evidence_based_protocol",
                success=False,
                error="UNIT_TEST_FAILURE_SIGNAL",
            )
        ]
        return svc

    register_harness("generation_service_failmock", lambda: GenerationServiceHarness(service_factory=factory))
    try:
        suite = BenchmarkSuite(
            cases=[
                BenchmarkCase(
                    id="fail1",
                    category="generation_adapter",
                    title="f",
                    inputs={"condition": "z", "tier": "fellow", "doc_type": "evidence_based_protocol"},
                    fixture_markdown="stub",
                    expectations=BenchmarkExpectations(),
                )
            ]
        )
        report = run_suite(suite, "generation_service_failmock")
        assert len(report.case_results) == 1
        h = GenerationServiceHarness(service_factory=factory)
        hr = h.run(suite.cases[0])
        assert "# Generation failed" in hr.output_text
        assert "UNIT_TEST_FAILURE_SIGNAL" in hr.output_text
    finally:
        from autoagent_clinical import registry as reg_module

        reg_module._REGISTRY.pop("generation_service_failmock", None)
