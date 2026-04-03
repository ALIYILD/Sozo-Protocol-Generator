"""Tests for the image search, curation, and insertion system."""
import pytest
from pathlib import Path


class TestImageSearchEngine:
    def test_import(self):
        from sozo_generator.images.search import (
            ImageSearchEngine,
            ImageResult,
            build_queries,
        )

    def test_build_queries_parkinsons_networks(self):
        from sozo_generator.images.search import build_queries
        queries = build_queries("parkinsons", "networks")
        assert len(queries) >= 2
        assert any("Parkinson" in q for q in queries)
        assert any("network" in q.lower() for q in queries)

    def test_build_queries_depression_pathophysiology(self):
        from sozo_generator.images.search import build_queries
        queries = build_queries("depression", "pathophysiology")
        assert len(queries) >= 1
        assert any("Depressive" in q for q in queries)

    def test_build_queries_with_extra_terms(self):
        from sozo_generator.images.search import build_queries
        queries = build_queries("parkinsons", "overview", extra_terms=["dopamine pathway"])
        assert any("dopamine pathway" in q for q in queries)

    def test_build_queries_fallback_for_unknown_section(self):
        from sozo_generator.images.search import build_queries
        queries = build_queries("parkinsons", "nonexistent_section_xyz")
        assert len(queries) >= 1  # Should use fallback templates

    def test_image_result_properties(self):
        from sozo_generator.images.search import ImageResult
        r = ImageResult(
            url="https://example.com/brain.png",
            title="Brain Map",
            source="pmc",
            license="CC-BY",
        )
        assert r.is_open_access is True
        assert r.safe_filename.startswith("img_")
        assert r.safe_filename.endswith(".png")

    def test_image_result_non_open_access(self):
        from sozo_generator.images.search import ImageResult
        r = ImageResult(url="https://example.com/img.jpg", license="unknown")
        assert r.is_open_access is False

    def test_search_engine_init(self, tmp_path):
        from sozo_generator.images.search import ImageSearchEngine
        engine = ImageSearchEngine(cache_dir=tmp_path / "cache")
        assert engine.cache_dir.exists()

    def test_section_query_templates_coverage(self):
        from sozo_generator.images.search import SECTION_QUERY_TEMPLATES
        # Verify key sections have templates
        assert "networks" in SECTION_QUERY_TEMPLATES
        assert "pathophysiology" in SECTION_QUERY_TEMPLATES
        assert "protocols" in SECTION_QUERY_TEMPLATES
        assert "eeg_qeeg" in SECTION_QUERY_TEMPLATES
        assert "mri_fmri" in SECTION_QUERY_TEMPLATES
        assert "connectivity" in SECTION_QUERY_TEMPLATES

    def test_all_15_conditions_have_display_names(self):
        from sozo_generator.images.search import CONDITION_DISPLAY_NAMES
        expected = [
            "parkinsons", "depression", "anxiety", "adhd", "alzheimers",
            "stroke_rehab", "tbi", "chronic_pain", "ptsd", "ocd",
            "ms", "asd", "long_covid", "tinnitus", "insomnia",
        ]
        for slug in expected:
            assert slug in CONDITION_DISPLAY_NAMES, f"Missing display name for {slug}"


class TestImageCurator:
    def test_import(self):
        from sozo_generator.images.curator import (
            ImageCurator,
            DocumentImageManifest,
            CuratedImage,
            SECTION_IMAGE_NEEDS,
        )

    def test_section_needs_defined(self):
        from sozo_generator.images.curator import SECTION_IMAGE_NEEDS
        assert "overview" in SECTION_IMAGE_NEEDS
        assert "networks" in SECTION_IMAGE_NEEDS
        assert "pathophysiology" in SECTION_IMAGE_NEEDS
        assert "protocols" in SECTION_IMAGE_NEEDS

    def test_curator_init(self, tmp_path):
        from sozo_generator.images.curator import ImageCurator
        curator = ImageCurator(
            image_dir=tmp_path / "images",
            cache_dir=tmp_path / "cache",
            enable_web_search=False,
        )
        assert curator.image_dir == tmp_path / "images"

    def test_curate_empty_no_crash(self, tmp_path):
        from sozo_generator.images.curator import ImageCurator
        curator = ImageCurator(
            image_dir=tmp_path / "images",
            cache_dir=tmp_path / "cache",
            enable_web_search=False,
        )
        manifest = curator.curate_for_document(
            "parkinsons", "evidence_based_protocol", "partners"
        )
        assert manifest.condition_slug == "parkinsons"
        # No precurated/cached images, search disabled → empty
        assert len(manifest.images) == 0

    def test_precurated_images_found(self, tmp_path):
        from sozo_generator.images.curator import ImageCurator
        # Create a precurated image
        img_dir = tmp_path / "images" / "parkinsons" / "networks"
        img_dir.mkdir(parents=True)
        fake_img = img_dir / "brain_network.png"
        fake_img.write_bytes(b"\x89PNG" + b"\x00" * 1000)  # Minimal PNG header
        meta = img_dir / "brain_network.json"
        meta.write_text('{"caption": "PD Network Dysfunction", "license": "CC-BY"}')

        curator = ImageCurator(
            image_dir=tmp_path / "images",
            cache_dir=tmp_path / "cache",
            enable_web_search=False,
        )
        manifest = curator.curate_for_document(
            "parkinsons", "evidence_based_protocol", "partners",
            sections=["networks"],
        )
        assert len(manifest.images) == 1
        assert manifest.images[0].source == "precurated"
        assert manifest.images[0].caption == "PD Network Dysfunction"

    def test_manifest_to_dict(self, tmp_path):
        from sozo_generator.images.curator import DocumentImageManifest, CuratedImage
        manifest = DocumentImageManifest(
            condition_slug="parkinsons",
            document_type="handbook",
            tier="partners",
            images=[
                CuratedImage(
                    local_path=Path("/tmp/test.png"),
                    caption="Test",
                    attribution="Test Source",
                    source="searched",
                    section_id="networks",
                    condition_slug="parkinsons",
                )
            ],
        )
        d = manifest.to_dict()
        assert d["condition_slug"] == "parkinsons"
        assert len(d["images"]) == 1
        assert d["images"][0]["section_id"] == "networks"


def _has_docx():
    try:
        import docx
        return True
    except ImportError:
        return False


class TestDocumentImageInserter:
    def test_import(self):
        from sozo_generator.images.inserter import DocumentImageInserter

    @pytest.mark.skipif(
        not _has_docx(), reason="python-docx not installed"
    )
    def test_insert_placeholder_for_missing(self, tmp_path):
        from docx import Document
        from sozo_generator.images.inserter import DocumentImageInserter
        from sozo_generator.images.curator import CuratedImage, DocumentImageManifest

        doc = Document()
        doc.add_paragraph("Test document")

        manifest = DocumentImageManifest(
            condition_slug="parkinsons",
            document_type="handbook",
            tier="fellow",
            images=[
                CuratedImage(
                    local_path=Path("/nonexistent/image.png"),
                    caption="Missing Image Test",
                    attribution="Test",
                    source="searched",
                    section_id="networks",
                    condition_slug="parkinsons",
                )
            ],
        )

        inserter = DocumentImageInserter()
        n = inserter.insert_images(doc, manifest)
        assert n == 0  # Image file doesn't exist → placeholder inserted

        # Verify placeholder text was added
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "FIGURE" in text
