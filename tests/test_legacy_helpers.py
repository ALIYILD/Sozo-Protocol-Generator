"""Tests for the shared DOCX legacy helper functions."""
import pytest


class TestLegacyHelpersImport:
    """Verify the shared helper module can be imported and has expected API."""

    def test_import_all_functions(self):
        from sozo_generator.docx.legacy_helpers import (
            para_replace,
            para_set,
            cell_write,
            replace_table,
            global_replace,
            apply_global_replacements,
            C_WHITE,
            C_BLACK,
        )
        assert callable(para_replace)
        assert callable(para_set)
        assert callable(cell_write)
        assert callable(replace_table)
        assert callable(global_replace)
        assert callable(apply_global_replacements)

    def test_color_constants(self):
        from sozo_generator.docx.legacy_helpers import C_WHITE, C_BLACK
        # RGBColor is a tuple-like of (R, G, B)
        assert C_WHITE == (0xFF, 0xFF, 0xFF)
        assert C_BLACK == (0x00, 0x00, 0x00)


try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestParaReplace:
    """Test para_replace on real Document objects."""

    def _make_para(self, text: str):
        """Create a temporary document with one paragraph."""
        doc = Document()
        para = doc.add_paragraph(text)
        return para

    def test_basic_replace(self):
        from sozo_generator.docx.legacy_helpers import para_replace
        para = self._make_para("Hello World")
        para_replace(para, "World", "Earth")
        assert "".join(r.text for r in para.runs) == "Hello Earth"

    def test_no_match_noop(self):
        from sozo_generator.docx.legacy_helpers import para_replace
        para = self._make_para("Hello World")
        para_replace(para, "Missing", "Replaced")
        assert "".join(r.text for r in para.runs) == "Hello World"

    def test_empty_paragraph(self):
        from sozo_generator.docx.legacy_helpers import para_replace
        doc = Document()
        para = doc.add_paragraph("")
        para_replace(para, "", "new text")
        assert "".join(r.text for r in para.runs) == "new text"


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestParaSet:
    def test_overwrite(self):
        from sozo_generator.docx.legacy_helpers import para_set
        doc = Document()
        para = doc.add_paragraph("Old text here")
        para_set(para, "Brand new text")
        assert "".join(r.text for r in para.runs) == "Brand new text"


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestCellWrite:
    def test_write_to_cell(self):
        from sozo_generator.docx.legacy_helpers import cell_write
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell_write(cell, "Test Content", bold=True)
        text = "".join(r.text for r in cell.paragraphs[0].runs)
        assert text == "Test Content"
        assert cell.paragraphs[0].runs[0].font.bold is True


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestReplaceTable:
    def test_replace_rows(self):
        from sozo_generator.docx.legacy_helpers import replace_table
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        rows = [
            ["Header A", "Header B"],
            ["Row 1A", "Row 1B"],
            ["Row 2A", "Row 2B"],
        ]
        replace_table(table, rows)
        for ri, row_data in enumerate(rows):
            for ci, expected in enumerate(row_data):
                cell_text = "".join(
                    r.text for r in table.rows[ri].cells[ci].paragraphs[0].runs
                )
                assert cell_text == expected


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
class TestGlobalReplace:
    def test_replaces_in_paragraphs(self):
        from sozo_generator.docx.legacy_helpers import global_replace
        doc = Document()
        doc.add_paragraph("Parkinson's Disease affects many people.")
        doc.add_paragraph("Another paragraph about Parkinson's Disease.")
        global_replace(doc, "Parkinson's Disease", "Depression")
        for para in doc.paragraphs:
            text = "".join(r.text for r in para.runs)
            assert "Parkinson's Disease" not in text
            assert "Depression" in text
