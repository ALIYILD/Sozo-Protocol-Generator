"""Tests for Phase 9 Template DNA System.

Covers blueprint_schema, blueprint_extractor, style_cloner, and
blueprint_renderer — the full Phase 9 pipeline from DOCX upload to
styled document generation.
"""
from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.document import Document as DocxDocument

from sozo_generator.template.phase9.blueprint_schema import (
    ColorPaletteEntry,
    DocumentBlueprint,
    LayoutRegion,
    PageLayout,
    SectionBlueprint,
    StyleSpec,
    TableSchemaBlueprint,
    empty_blueprint,
)
from sozo_generator.template.phase9.blueprint_extractor import (
    BlueprintExtractor,
    extract_blueprint,
)
from sozo_generator.template.phase9.style_cloner import (
    StyleCloner,
    clone_styles_from_blueprint,
)
from sozo_generator.template.phase9.blueprint_renderer import (
    BlueprintRenderer,
    render_document,
)
from sozo_generator.schemas.documents import DocumentSpec, SectionContent
from sozo_generator.core.enums import DocumentType, Tier


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def minimal_blueprint() -> DocumentBlueprint:
    """A minimal valid DocumentBlueprint for testing."""
    return empty_blueprint("test_template.docx")


@pytest.fixture
def rich_blueprint() -> DocumentBlueprint:
    """Blueprint with sections, styles, and color palette."""
    bp = empty_blueprint("rich_template.docx")
    bp.styles.append(
        StyleSpec(
            name="SOZO Heading 1",
            font_name="Calibri",
            font_size_pt=16.0,
            bold=True,
            color_hex="1F3864",
            is_heading=True,
            heading_level=1,
            space_before_pt=12.0,
            space_after_pt=6.0,
        )
    )
    bp.color_palette.append(ColorPaletteEntry(role="table_header", hex="2E4057"))
    bp.sections.append(
        SectionBlueprint(
            section_id="overview",
            section_type="overview",
            heading_text="Clinical Overview",
            heading_style="SOZO Heading 1",
            heading_level=1,
            order=1,
        )
    )
    bp.sections.append(
        SectionBlueprint(
            section_id="protocol",
            section_type="protocol_parameters",
            heading_text="Protocol Parameters",
            heading_style="SOZO Heading 1",
            heading_level=1,
            order=2,
        )
    )
    bp.has_header = True
    bp.document_type_guess = "protocol"
    return bp


@pytest.fixture
def simple_docx_path(tmp_path) -> Path:
    """Create a minimal .docx file for extraction testing."""
    doc = Document()
    doc.add_heading("Clinical Overview", level=1)
    doc.add_paragraph("This is the clinical overview section.")
    doc.add_heading("Protocol Parameters", level=2)
    doc.add_paragraph("TMS parameters go here.")
    t = doc.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "Parameter"
    t.rows[0].cells[1].text = "Value"
    t.rows[0].cells[2].text = "Evidence"
    path = tmp_path / "test_template.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def minimal_docspec() -> DocumentSpec:
    """A DocumentSpec with two sections for renderer tests."""
    return DocumentSpec(
        document_type=DocumentType.EVIDENCE_BASED_PROTOCOL,
        tier=Tier.FELLOW,
        condition_slug="depression",
        condition_name="Major Depressive Disorder",
        title="tDCS Protocol for MDD",
        sections=[
            SectionContent(
                section_id="overview",
                title="Clinical Overview",
                content="Major Depressive Disorder is characterised by...",
                evidence_pmids=["28805473"],
                confidence_label="high",
            ),
            SectionContent(
                section_id="protocol",
                title="Protocol Parameters",
                content="Anodal tDCS at 2mA applied to left DLPFC for 20 minutes.",
                tables=[
                    {
                        "headers": ["Parameter", "Value"],
                        "rows": [["Intensity", "2 mA"], ["Duration", "20 min"]],
                    }
                ],
            ),
        ],
    )


# ---------------------------------------------------------------------------
# TestDocumentBlueprintSchema
# ---------------------------------------------------------------------------


class TestDocumentBlueprintSchema:
    """Unit tests for DocumentBlueprint and related schema models."""

    def test_empty_blueprint_is_valid(self):
        """empty_blueprint() should return a DocumentBlueprint with a non-empty blueprint_id."""
        bp = empty_blueprint()
        assert isinstance(bp, DocumentBlueprint)
        assert bp.blueprint_id  # UUID must be non-empty
        assert bp.source_filename == "unknown.docx"

    def test_empty_blueprint_has_default_styles(self):
        """empty_blueprint() should include 'Normal' and 'Heading 1' styles."""
        bp = empty_blueprint()
        style_names = bp.style_names()
        assert "Normal" in style_names
        assert "Heading 1" in style_names

    def test_get_style_returns_correct_spec(self):
        """get_style('Normal') should return the Normal StyleSpec with font_name='Calibri'."""
        bp = empty_blueprint()
        style = bp.get_style("Normal")
        assert style is not None
        assert style.font_name == "Calibri"

    def test_get_style_returns_none_for_missing(self):
        """get_style() should return None for a style that does not exist."""
        bp = empty_blueprint()
        result = bp.get_style("NonExistentStyle999")
        assert result is None

    def test_has_style_true_and_false(self):
        """has_style() should return True for existing styles and False for absent ones."""
        bp = empty_blueprint()
        assert bp.has_style("Normal") is True
        assert bp.has_style("Heading 1") is True
        assert bp.has_style("GhostStyle404") is False

    def test_color_palette_entry_strips_hash(self):
        """ColorPaletteEntry should store hex without the leading '#' and uppercase."""
        entry = ColorPaletteEntry(role="primary", hex="#1F3864")
        assert entry.hex == "1F3864"

    def test_color_palette_entry_rejects_invalid_hex(self):
        """ColorPaletteEntry should raise ValidationError for non-hex input."""
        from pydantic import ValidationError

        with pytest.raises(ValidationError):
            ColorPaletteEntry(role="primary", hex="ZZZ")

    def test_primary_color_returns_first_primary(self):
        """primary_color() should return the hex of the first 'primary' palette entry."""
        bp = empty_blueprint()
        # empty_blueprint already seeds a primary entry at '1F3864'
        assert bp.primary_color() == "1F3864"

    def test_ordered_sections(self, rich_blueprint):
        """ordered_sections() should return sections sorted ascending by order."""
        # Add an out-of-order section
        rich_blueprint.sections.append(
            SectionBlueprint(
                section_id="early",
                section_type="overview",
                heading_text="Early Section",
                heading_level=1,
                order=0,
            )
        )
        ordered = rich_blueprint.ordered_sections()
        orders = [s.order for s in ordered]
        assert orders == sorted(orders)

    def test_to_jinja2_context_keys(self, minimal_blueprint):
        """to_jinja2_context() should include all required bp_ prefixed keys."""
        ctx = minimal_blueprint.to_jinja2_context()
        required_keys = {
            "bp_primary_color",
            "bp_heading_font",
            "bp_body_font",
            "bp_page_size",
            "bp_section_count",
        }
        assert required_keys.issubset(set(ctx.keys()))

    def test_blueprint_serialises_to_json(self, minimal_blueprint):
        """model_dump_json() should produce valid JSON for a DocumentBlueprint."""
        json_str = minimal_blueprint.model_dump_json()
        parsed = json.loads(json_str)
        assert "blueprint_id" in parsed
        assert "source_filename" in parsed

    def test_blueprint_round_trips_json(self, minimal_blueprint):
        """A blueprint should survive a dump-and-reload round-trip without data loss."""
        json_str = minimal_blueprint.model_dump_json()
        reloaded = DocumentBlueprint.model_validate_json(json_str)
        assert reloaded.blueprint_id == minimal_blueprint.blueprint_id
        assert reloaded.source_filename == minimal_blueprint.source_filename
        assert len(reloaded.styles) == len(minimal_blueprint.styles)


# ---------------------------------------------------------------------------
# TestBlueprintExtractor
# ---------------------------------------------------------------------------


class TestBlueprintExtractor:
    """Integration tests for BlueprintExtractor against a real DOCX fixture."""

    def test_extract_from_real_docx(self, simple_docx_path):
        """extract() on a real DOCX should return a DocumentBlueprint with sections > 0."""
        bp = BlueprintExtractor(simple_docx_path).extract()
        assert isinstance(bp, DocumentBlueprint)
        assert len(bp.sections) > 0

    def test_extract_page_layout(self, simple_docx_path):
        """Extracted blueprint should contain a valid PageLayout."""
        bp = BlueprintExtractor(simple_docx_path).extract()
        assert isinstance(bp.page_layout, PageLayout)

    def test_extract_styles_not_empty(self, simple_docx_path):
        """Extracted blueprint should have at least one style."""
        bp = BlueprintExtractor(simple_docx_path).extract()
        assert len(bp.styles) > 0

    def test_extract_detects_heading_sections(self, simple_docx_path):
        """Extractor should detect at least one section from heading paragraphs."""
        bp = BlueprintExtractor(simple_docx_path).extract()
        heading_sections = [s for s in bp.sections if s.section_type != "toc"]
        assert len(heading_sections) >= 1

    def test_extract_detects_table_schema(self, simple_docx_path):
        """Extractor should detect the table in the DOCX and produce table_schemas."""
        bp = BlueprintExtractor(simple_docx_path).extract()
        assert len(bp.table_schemas) > 0
        assert isinstance(bp.table_schemas[0], TableSchemaBlueprint)

    def test_extract_handles_missing_file(self, tmp_path):
        """BlueprintExtractor should raise FileNotFoundError for a missing path."""
        missing = tmp_path / "no_such_file.docx"
        with pytest.raises(FileNotFoundError):
            BlueprintExtractor(missing).extract()

    def test_extract_function_alias(self, simple_docx_path):
        """extract_blueprint() convenience function should return a DocumentBlueprint."""
        bp = extract_blueprint(simple_docx_path)
        assert isinstance(bp, DocumentBlueprint)

    def test_blueprint_pipeline_version(self, simple_docx_path):
        """Extracted blueprint pipeline_version should be 'phase9_v1'."""
        bp = extract_blueprint(simple_docx_path)
        assert bp.pipeline_version == "phase9_v1"


# ---------------------------------------------------------------------------
# TestStyleCloner
# ---------------------------------------------------------------------------


class TestStyleCloner:
    """Unit tests for StyleCloner and clone_styles_from_blueprint."""

    def test_clone_applies_page_layout(self, minimal_blueprint):
        """StyleCloner.apply() should set the page margins on doc.sections[0]."""
        doc = Document()
        StyleCloner(minimal_blueprint).apply(doc)
        section = doc.sections[0]
        expected_top = minimal_blueprint.page_layout.margin_top_cm
        # Margin is stored as EMU; convert back to cm for comparison (tolerance 0.1 cm)
        from docx.shared import Cm
        assert abs(section.top_margin - Cm(expected_top)) < 10_000  # ~0.001 cm tolerance

    def test_clone_creates_custom_style(self, rich_blueprint):
        """After apply(), the document should contain all styles listed in the blueprint."""
        doc = Document()
        StyleCloner(rich_blueprint).apply(doc)
        style_names_in_doc = [s.name for s in doc.styles]
        assert "SOZO Heading 1" in style_names_in_doc

    def test_clone_returns_document_instance(self, minimal_blueprint):
        """StyleCloner.apply() must return a python-docx Document object."""
        doc = Document()
        result = StyleCloner(minimal_blueprint).apply(doc)
        assert isinstance(result, DocxDocument)

    def test_clone_styles_function_alias(self, minimal_blueprint):
        """clone_styles_from_blueprint() convenience function should return a Document."""
        result = clone_styles_from_blueprint(minimal_blueprint)
        assert isinstance(result, DocxDocument)

    def test_clone_does_not_crash_on_empty_blueprint(self):
        """Applying an empty_blueprint() to a new Document should not raise."""
        bp = empty_blueprint()
        doc = Document()
        # Should complete without exception
        result = StyleCloner(bp).apply(doc)
        assert isinstance(result, DocxDocument)

    def test_clone_applies_font_name(self, rich_blueprint):
        """The 'SOZO Heading 1' style in the cloned document should have font name 'Calibri'."""
        doc = Document()
        StyleCloner(rich_blueprint).apply(doc)
        sozo_style = doc.styles["SOZO Heading 1"]
        assert sozo_style.font.name == "Calibri"


# ---------------------------------------------------------------------------
# TestBlueprintRenderer
# ---------------------------------------------------------------------------


class TestBlueprintRenderer:
    """Unit tests for BlueprintRenderer and render_document."""

    def test_render_returns_document(self, rich_blueprint, minimal_docspec):
        """render_document() should return a python-docx Document."""
        result = render_document(rich_blueprint, minimal_docspec)
        assert isinstance(result, DocxDocument)

    def test_render_saves_to_path(self, rich_blueprint, minimal_docspec, tmp_path):
        """When output_path is supplied, the file should exist after rendering."""
        out = tmp_path / "output.docx"
        render_document(rich_blueprint, minimal_docspec, output_path=out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_render_includes_sections(self, rich_blueprint, minimal_docspec):
        """The rendered document should contain at least one paragraph."""
        doc = render_document(rich_blueprint, minimal_docspec)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert len(doc.paragraphs) > 0
        assert len(all_text) > 0

    def test_render_section_headings_present(self, rich_blueprint, minimal_docspec):
        """The rendered document should contain the heading text 'Clinical Overview'."""
        doc = render_document(rich_blueprint, minimal_docspec)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Clinical Overview" in all_text

    def test_render_handles_missing_content(self, rich_blueprint):
        """Rendering a spec with no sections should not crash and should produce a Document."""
        empty_spec = DocumentSpec(
            document_type=DocumentType.EVIDENCE_BASED_PROTOCOL,
            tier=Tier.FELLOW,
            condition_slug="test",
            condition_name="Test Condition",
            title="Empty Test",
            sections=[],
        )
        doc = render_document(rich_blueprint, empty_spec)
        assert isinstance(doc, DocxDocument)
        # Placeholder paragraphs should be present for the two blueprint sections
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "content pending" in all_text

    def test_render_with_minimal_blueprint(self, minimal_blueprint, minimal_docspec):
        """A minimal blueprint (no custom styles, no sections) should still render without error."""
        doc = render_document(minimal_blueprint, minimal_docspec)
        assert isinstance(doc, DocxDocument)
        # The spec sections should still produce some output even with no blueprint sections
        # (blueprint has no sections, so no section rendering — but the doc itself is valid)
        assert doc is not None
