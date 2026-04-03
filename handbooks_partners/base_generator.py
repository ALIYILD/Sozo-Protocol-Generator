"""
SOZO Clinical Handbook Generator — Partners Tier
Clones PD Partners template and applies condition-specific replacements.
Usage: from handbooks_partners.base_generator import build_handbook
       build_handbook(cdata)

Template profile: 478 paragraphs, 16 tables, ~3,621 words
Key replaced tables: [1],[4],[7],[11],[12],[14]
Key replaced para ranges: [10],[37],[103],[131-148],[152],[287],[314],[367],[433-455]
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

_PROJECT_ROOT = Path(__file__).resolve().parent.parent

# Ensure src/ is importable for shared helpers
sys.path.insert(0, str(_PROJECT_ROOT / "src"))
from sozo_generator.docx.legacy_helpers import (
    para_replace,
    para_set,
    cell_write,
    replace_table,
    apply_global_replacements,
)

TEMPLATE = _PROJECT_ROOT / "templates" / "gold_standard" / "Clinical_Handbook.docx"
OUTPUT_ROOT = Path("outputs/documents")


# ─── global text replacements ────────────────────────────────────────────────

def _apply_global_replacements(doc, c):
    """Replace PD-specific phrases with condition-specific equivalents throughout."""
    reps = [
        # Name variants (longest first to avoid partial replacement)
        ("PARKINSON'S DISEASE",         c["cover_title"]),
        ("Parkinson's Disease (PD)",     c["name"]),
        ("Parkinson's disease (PD)",     c["name"]),
        ("Parkinson's Disease",          c["name"]),
        ("Parkinson's disease",          c["name"]),
        ("Parkinson's",                  c["short"]),
        # Abbreviation
        (" (PD)",                        f" ({c['abbr']})"),
        ("(PD) ",                        f"({c['abbr']}) "),
        (" PD ",                         f" {c['abbr']} "),
        (" PD.",                         f" {c['abbr']}."),
        (" PD,",                         f" {c['abbr']},"),
        (" PD)",                         f" {c['abbr']})"),
        (" PD:",                         f" {c['abbr']}:"),
        ("PD diagnosis",                 f"{c['abbr']} diagnosis"),
        (" PD is",                       f" {c['abbr']} is"),
        ("PD or",                        f"{c['abbr']} or"),
        ("for PD ",                      f"for {c['abbr']} "),
        ("for PD.",                      f"for {c['abbr']}."),
        ("for PD,",                      f"for {c['abbr']},"),
        # PD-specific clinical references in boilerplate
        ("Hoehn & Yahr stage",           "disease severity rating"),
        ("Hoehn and Yahr stage",         "disease severity rating"),
        ("Document levodopa timing specifically", "Document primary medication timing specifically"),
        ("levodopa",                     c.get("med_name", "primary medication")),
        # Cover subtitle (condition already handled above via name replacements)
        ("for Parkinson's Disease (PD) using", f"for {c['name']} using"),
    ]
    apply_global_replacements(doc, reps)


# ─── paragraph-level targeted replacements ───────────────────────────────────

def _apply_para_replacements(doc, c):
    """Replace full paragraph text at known indices."""
    paras = doc.paragraphs

    # [10] PARTNERS TIER callout (purple)
    para_set(paras[10], c["partners_tier_body"])

    # [37] TPS off-label warning
    para_set(paras[37], c["tps_offlabel_body"])

    # [103] MANDATORY consent warning
    para_set(paras[103], c["mandatory_consent_body"])

    # [131] Primary symptoms domain header
    para_set(paras[131], c["prs_primary_label"])

    # [132]-[139] Primary PRS items (8)
    for i, item in enumerate(c["prs_primary_items"][:8]):
        para_set(paras[132 + i], item)

    # [140] Secondary symptoms domain header
    para_set(paras[140], c["prs_secondary_label"])

    # [141]-[148] Secondary PRS items (8)
    for i, item in enumerate(c["prs_secondary_items"][:8]):
        para_set(paras[141 + i], item)

    # [152] Preliminary phenotype assignment line
    para_set(paras[152], c["phenotype_prelim"])

    # [287] Pre-session medication check
    para_set(paras[287], c["pre_session_med_check"])

    # [314] Session medication documentation item
    para_set(paras[314], c["session_med_doc"])

    # [367] False classification warning
    para_set(paras[367], c["false_class_body"])

    # [433]-[438] Inclusion criteria items (6)
    for i, item in enumerate(c["inclusion_items"][:6]):
        para_set(paras[433 + i], item)

    # [440]-[446] Absolute exclusion items (7)
    for i, item in enumerate(c["exclusion_items"][:7]):
        para_set(paras[440 + i], item)

    # [448]-[455] Conditional criteria items (8)
    for i, item in enumerate(c["conditional_items"][:8]):
        para_set(paras[448 + i], item)


# ─── table-level replacements ────────────────────────────────────────────────

def _apply_table_replacements(doc, c):
    tables = doc.tables

    # Table[1]: Available Modalities (5r x 4c)
    replace_table(tables[1], c["modality_table"])

    # Table[4]: Off-Label Disclosure Summary (5r x 4c)
    replace_table(tables[4], c["offlabel_table"])

    # Table[7]: Phenotype Identification (9r x 2c)
    replace_table(tables[7], c["phenotype_table"])

    # Table[11]: In-Clinic Task Pairing (6r x 3c)
    replace_table(tables[11], c["task_pairing_table"])

    # Table[12]: Response Domains (5r x 2c)
    replace_table(tables[12], c["response_domains_table"])

    # Table[14]: Appendix B Phenotype-to-Montage (10r x 2c)
    replace_table(tables[14], c["montage_table"])


# ─── main entry point ────────────────────────────────────────────────────────

def build_handbook(c: dict):
    """Build Partners handbook for one condition. c = condition data dict."""
    doc = Document(str(TEMPLATE))

    _apply_global_replacements(doc, c)
    _apply_para_replacements(doc, c)
    _apply_table_replacements(doc, c)

    slug = c["slug"]
    out_path = OUTPUT_ROOT / slug / "partners" / f"Clinical_Handbook_Partners_{slug}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  [OK] {out_path}")
    return out_path
