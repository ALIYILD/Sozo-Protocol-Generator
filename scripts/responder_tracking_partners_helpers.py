"""Helper functions for Responder Tracking Partners generator."""
from docx.shared import RGBColor

C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK  = RGBColor(0x00, 0x00, 0x00)
C_NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
C_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
C_PURPLE = RGBColor(0x7B, 0x2D, 0x8E)
C_RED    = RGBColor(0xCC, 0x00, 0x00)


def _para_replace(para, old: str, new: str):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold  = fr.font.bold if fr else None
    size  = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except Exception:
        color = None
    for r in para.runs:
        r.text = ""
    if fr:
        fr.text = replaced
        fr.font.bold = bold
        if size:  fr.font.size = size
        if color: fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text: str):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text: str, bold: bool = False, white: bool = False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    run  = para.runs[0] if para.runs else para.add_run()
    run.text = text
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type:
            run.font.color.rgb = C_BLACK


def _replace_table(table, rows: list):
    n_template_rows = len(table.rows)
    n_new_rows = len(rows)
    for ri in range(n_template_rows):
        row = table.rows[ri]
        if ri >= n_new_rows:
            for cell in row.cells:
                _cell_write(cell, "")
            continue
        is_hdr = (ri == 0)
        new_row = rows[ri]
        for ci, cell in enumerate(row.cells):
            text = str(new_row[ci]) if ci < len(new_row) else ""
            _cell_write(cell, text, bold=is_hdr, white=is_hdr)
