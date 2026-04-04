#!/usr/bin/env python3
"""
Generate Responder Tracking & Classification (Partners Tier) for all 14+ conditions.
Built from scratch with python-docx (no template dependency).
Usage: python scripts/generate_responder_tracking_partners.py [slug]
       omit slug to generate all conditions.
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── add project root to path ─────────────────────────────────────────────────
sys.path.insert(0, str(Path(__file__).parent))
from responder_conditions_data import CONDITIONS
from responder_conditions_data_2 import CONDITIONS_2
CONDITIONS.update(CONDITIONS_2)

_PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_ROOT = _PROJECT_ROOT / "outputs" / "documents"

# ── colour palette ────────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
C_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
C_PURPLE = RGBColor(0x7B, 0x2D, 0x8E)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK  = RGBColor(0x00, 0x00, 0x00)
C_LGRAY  = RGBColor(0xEB, 0xF0, 0xF7)

HEX_NAVY   = "1B3A5C"
HEX_BLUE   = "2E75B6"
HEX_PURPLE = "7B2D8E"
HEX_LGRAY  = "EBF0F7"
HEX_WHITE  = "FFFFFF"


# ── xml helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_col_width(table, col_idx: int, width_inches: float):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


# ── paragraph/table factories ─────────────────────────────────────────────────

def _add_title_banner(doc, text: str):
    """Full-width navy banner paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = C_WHITE
    # shade the paragraph background via a 1x1 table trick — simpler: use a table
    # (paragraph shading via XML)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), HEX_NAVY)
    pPr.append(shd)
    return p


def _add_section_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(11)
        run.font.color.rgb = C_NAVY
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = C_BLUE
    return p


def _add_body_text(doc, text: str, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    if italic:
        run.italic = True
    return p


def _make_table(doc, rows_data: list, navy_header: bool = True):
    """Build a table from rows_data (list of lists). First row = header."""
    if not rows_data:
        return None
    n_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_vals in enumerate(rows_data):
        row = table.rows[ri]
        is_hdr = (ri == 0 and navy_header)
        shade = not is_hdr and (ri % 2 == 1)
        for ci in range(n_cols):
            cell = row.cells[ci]
            val = str(row_vals[ci]) if ci < len(row_vals) else ""
            if is_hdr:
                _set_cell_bg(cell, HEX_NAVY)
            elif shade:
                _set_cell_bg(cell, HEX_LGRAY)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.bold = is_hdr
            run.font.size = Pt(9)
            if is_hdr:
                run.font.color.rgb = C_WHITE
    return table


# ── build one condition ───────────────────────────────────────────────────────

def build(c: dict):
    slug = c["slug"]
    name = c["name"]
    short = c["short"]
    print(f"  Generating: {short} ({slug})...")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # ── Header banner ─────────────────────────────────────────────────────────
    _add_title_banner(doc, "SOZO BRAIN CENTER \u2014 PARTNERS TIER")
    _add_title_banner(doc, c["subtitle"])
    _add_title_banner(doc, c["sec2_subtitle"])

    p = doc.add_paragraph()
    run = p.add_run("FNON + Evidence-Based Assessment \u2014 Includes Brain Network Analysis")
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = C_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Patient Information Table ─────────────────────────────────────────────
    _add_section_heading(doc, "Patient Information", level=2)
    patient_rows = [
        ["Field", "Details"],
        ["Patient Name / ID", ""],
        ["Date of Birth", ""],
        [c["phenotype_row"][0], c["phenotype_row"][1]],
        ["Symptom Severity / Stage", ""],
        ["Condition Name", name],
        ["Medication Class", c.get("med_class", "")],
        ["Referring Clinician", ""],
        ["Assessment Date", ""],
    ]
    _make_table(doc, patient_rows)
    doc.add_paragraph()

    # ── Section 1: Response Definition ───────────────────────────────────────
    _add_section_heading(doc, "1. SOZO Operational Response Definition")
    _add_body_text(doc, (
        "Define response at end of block (typically 10\u201320 sessions). "
        "Patient must meet criteria in at least one primary domain AND show no deterioration in others."
    ))
    _add_body_text(doc, c["false_nonresponder_note"], italic=True)
    doc.add_paragraph()

    # ── Section 2: Likely Responder Profiles ─────────────────────────────────
    _add_section_heading(doc, "2. Likely Responder Profiles")
    _add_body_text(doc, (
        "FNON-based responder profiling maps treatment targets to network hypotheses "
        "for optimised protocol selection."
    ))
    _make_table(doc, c["responder_profiles"])
    doc.add_paragraph()

    # ── Section 3: Predictors of Better Response ─────────────────────────────
    _add_section_heading(doc, "3. Predictors of Better Response")
    _add_section_heading(doc, "3A. Clinical Predictors", level=2)
    _make_table(doc, c["clinical_predictors"])
    doc.add_paragraph()

    _add_section_heading(doc, "3B. Protocol Predictors (Modifiable)", level=2)
    _make_table(doc, c["protocol_predictors"])
    doc.add_paragraph()

    _add_section_heading(doc, "3C. Measurement Predictors", level=2)
    _add_body_text(doc, c["measurement_confound_warning"], italic=True)
    doc.add_paragraph()

    # ── Section 4: Response Classification ───────────────────────────────────
    _add_section_heading(doc, "4. SOZO Response Classification")
    _add_body_text(doc, "Use at Session 10 (or 8\u201310) and again at end of block:")
    _make_table(doc, c["response_classification"])
    _add_body_text(doc, (
        "PATIENT CLASSIFICATION: \u25a1 Responder  \u25a1 Partial Responder  \u25a1 Non-Responder"
    ))
    doc.add_paragraph()

    # ── Section 5: Non-Responder Profiles ────────────────────────────────────
    _add_section_heading(doc, "5. Likely Non-Responder Profiles")
    _make_table(doc, c["non_responder_profiles"])
    doc.add_paragraph()

    # ── Section 6: FNON Pathway ───────────────────────────────────────────────
    _add_section_heading(doc, "6. SOZO Non-Responder Pathway (FNON Level 5)")
    _add_body_text(doc, (
        "If no meaningful improvement after 8\u201310 sessions, follow FNON Level 5 adjustment pathway: "
        "Reassess phenotype, review medication state, adjust montage, consider multi-modal pathway."
    ))
    _add_body_text(doc, (
        "DO NOT escalate intensity beyond safety limits. "
        "Always follow SOZO FNON Protocol safety parameters."
    ))
    doc.add_paragraph()

    # ── Section 7: Medication-tDCS Documentation ──────────────────────────────
    _add_section_heading(doc, c["sec7_title"])
    _add_body_text(doc, c["sec7_intro"])

    _add_section_heading(doc, c["sec7a_title"], level=2)
    _make_table(doc, c["med_state_comparison"])
    doc.add_paragraph()

    _add_section_heading(doc, "7B. SOZO Scheduling Rules", level=2)
    _add_body_text(doc, c["scheduling_rule"])
    doc.add_paragraph()

    _add_section_heading(doc, c["sec7c_title"], level=2)
    _make_table(doc, c["phenotype_pairing"])
    doc.add_paragraph()

    _add_section_heading(doc, c["sec7d_title"], level=2)
    _make_table(doc, c["doc_checklist"])
    doc.add_paragraph()

    # ── Section 8: Signature / Data Collection ────────────────────────────────
    _add_section_heading(doc, "8. SOZO Responder Signature Data Collection")
    _make_table(doc, c["med_documentation"])
    doc.add_paragraph()

    # ── Regulator-Safe Language Guide ─────────────────────────────────────────
    _add_section_heading(doc, "Regulator-Safe Language Guide")
    _make_table(doc, c["language_guide"])
    doc.add_paragraph()

    # ── Response Domains (Section 1 table) ────────────────────────────────────
    _add_section_heading(doc, "Response Domain Assessment", level=2)
    _make_table(doc, c["response_domains"])
    doc.add_paragraph()

    # ── Clinical Notes ────────────────────────────────────────────────────────
    _add_section_heading(doc, "Clinical Notes & Rationale:", level=2)
    _add_body_text(doc, "_" * 60)
    _add_body_text(doc, "_" * 60)

    # ── Sign-off table ────────────────────────────────────────────────────────
    signoff_rows = [
        ["Role", "Name", "Signature", "Date"],
        ["Treating Clinician", "", "", ""],
    ]
    _make_table(doc, signoff_rows)

    # ── Save ──────────────────────────────────────────────────────────────────
    out = OUTPUT_ROOT / slug / "partners" / f"Responder_Tracking_Partners_{slug}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  [OK] {out}")
    return out


# ── entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else None
    if target:
        if target not in CONDITIONS:
            print(f"Unknown slug: {target}. Available: {list(CONDITIONS.keys())}")
            sys.exit(1)
        build(CONDITIONS[target])
    else:
        print(f"Generating Responder Tracking (Partners) for {len(CONDITIONS)} conditions...")
        ok, fail = [], []
        for slug, data in CONDITIONS.items():
            try:
                build(data)
                ok.append(slug)
            except Exception as e:
                import traceback
                print(f"  [FAIL] {slug}: {e}")
                traceback.print_exc()
                fail.append(slug)
        print(f"\nDone: {len(ok)} OK, {len(fail)} failed.")
        if fail:
            print(f"Failed: {fail}")
