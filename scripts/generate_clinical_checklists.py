"""
Generate Clinical Examination Checklist (Fellow Edition) for all 14 conditions.
Exact replica of PD template structure with condition-specific adaptations.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x3A, 0x5C)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)
CB = "\u2610"  # checkbox


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def add_header_row(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, "1B3A5C")
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_data_row(table, row_idx, values, shade=False):
    row = table.rows[row_idx]
    for i, v in enumerate(values):
        cell = row.cells[i]
        if shade:
            set_cell_bg(cell, "EBF0F7")
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(v))
        run.font.size = Pt(9)


def make_table(doc, headers, data_rows):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, headers)
    for i, row in enumerate(data_rows):
        add_data_row(table, i + 1, row, shade=(i % 2 == 0))
    return table


def add_heading1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_normal(doc, text, color=None, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def build_document(cond):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Apply heading styles from template
    h1_style = doc.styles["Heading 1"]
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.color.rgb = NAVY

    h2_style = doc.styles["Heading 2"]
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.font.color.rgb = BLUE

    # ── Cover Header ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SOZO BRAIN CENTER \u2014 CYPRUS")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FELLOW CLINICAL ASSESSMENT")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Clinical Examination Checklist")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Stage 4 \u2014 Patient Journey | Evidence-Based Neurological Examination")
    r.font.color.rgb = GRAY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cond["display_name"].upper())
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(
        f"This checklist guides the Fellow through a structured neurological examination for "
        f"{cond['display_name']}. The focus is on standardised scoring, phenotype identification, "
        f"and selection of appropriate validated assessment instruments."
    )
    r.font.color.rgb = GRAY

    doc.add_paragraph()

    # ── Patient Information ──
    add_heading1(doc, "Patient Information")
    make_table(doc, ["Field", "Value", "Field", "Value"], [
        ["Patient Name", "", "Date", ""],
        ["DOB", "", "Clinician", ""],
        ["File Number", "", "Assessment Type", f"{CB} Baseline  {CB} Follow-Up"],
    ])

    doc.add_paragraph()
    make_table(doc, ["Field", "Value", "Field", "Value"], [
        [cond["status_field_1"], "", cond["status_field_2"], ""],
        ["Symptom Duration", "___ years / months", "Dominant Side", f"{CB} Left  {CB} Right  {CB} N/A"],
        ["Referring Clinician", "", "Session #", ""],
        ["Medication / Treatment", "", "Setting", f"{CB} Clinic  {CB} Remote"],
    ])

    doc.add_paragraph()

    # ── Section A ──
    add_heading1(doc, f"Section A: {cond['section_a_title']}")
    add_normal(doc, "Score using the following scale unless otherwise specified:")
    p = doc.add_paragraph()
    r = p.add_run("0 = Normal  1 = Slight  2 = Mild  3 = Moderate  4 = Severe")
    r.bold = True

    for sub in cond["section_a_subsections"]:
        add_heading2(doc, sub["title"])
        make_table(doc, sub["headers"], sub["rows"])
        doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(f"{cond['section_a_title']} Total: _____ / _____")
    r.bold = True

    doc.add_paragraph()

    # ── Section B ──
    add_heading1(doc, "Section B: Standard Neurological Examination")
    make_table(doc, ["Examination Domain", "Finding", "Details / Notes"],
               cond.get("section_b_rows", [
        ["Pyramidal tract signs", f"{CB} Normal  {CB} Abnormal", ""],
        ["Extrapyramidal signs", f"{CB} Normal  {CB} Abnormal", ""],
        ["Frontal release signs", f"{CB} Normal  {CB} Abnormal", ""],
        ["Cerebellar examination", f"{CB} Normal  {CB} Abnormal", ""],
        ["Cranial nerves / visual-oculomotor", f"{CB} Normal  {CB} Abnormal", ""],
        ["Vestibular-oculomotor", f"{CB} Normal  {CB} Abnormal", ""],
        ["Sensory examination", f"{CB} Normal  {CB} Abnormal", ""],
        ["Deep tendon reflexes", f"{CB} Normal  {CB} Abnormal", ""],
        ["Plantar responses", f"{CB} Flexor (normal)  {CB} Extensor", ""],
        ["Romberg test", f"{CB} Negative  {CB} Positive", ""],
    ]))

    doc.add_paragraph()

    # ── Section C ──
    add_heading1(doc, "Section C: Cognitive Screening")
    add_normal(doc, "Use standardised tools. Record scores below:")
    make_table(doc, ["Assessment", "Score", "Cut-off", "Interpretation"], cond["section_c_rows"])

    doc.add_paragraph()

    # ── Section D ──
    add_heading1(doc, "Section D: Mood & Behavioural Screening")
    make_table(doc, ["Assessment", "Score", "Clinical Threshold", "Action"], cond["section_d_rows"])

    doc.add_paragraph()

    # ── Section E ──
    add_heading1(doc, "Section E: Phenotype Identification & Clinical Summary")
    make_table(doc, ["Domain", "Assessment"], cond["section_e_rows"])

    doc.add_paragraph()
    add_heading2(doc, "Clinical Notes")
    for _ in range(5):
        doc.add_paragraph("_" * 90)

    doc.add_paragraph()
    make_table(doc, ["Assessed By", "Signature", "Date"], [["", "", ""]])

    return doc


# ══════════════════════════════════════════════════════════════════════════════
# CONDITION DATA — 14 CONDITIONS
# ══════════════════════════════════════════════════════════════════════════════

CONDITIONS = [

    # 1 ─ DEPRESSION ──────────────────────────────────────────────────────────
    {
        "slug": "depression",
        "display_name": "Major Depressive Disorder (MDD)",
        "status_field_1": "Current Episode Duration",
        "status_field_2": "Episode Number",
        "section_a_title": "Mood & Affective Examination",
        "section_a_subsections": [
            {
                "title": "A1. Mood & Affect Assessment",
                "headers": ["Item", "Observation / Rating", "Score (0–4)", "Notes"],
                "rows": [
                    ["Subjective mood (patient report)", "0=Normal, 4=Severely depressed", "", ""],
                    ["Observed affect", "0=Normal range, 4=Flat/blunted", "", ""],
                    ["Mood reactivity", "0=Normal, 4=Completely non-reactive", "", ""],
                    ["Diurnal variation", "0=None, 4=Severe AM worsening", "", ""],
                    ["Emotional lability", "0=None, 4=Uncontrollable", "", ""],
                ],
            },
            {
                "title": "A2. Psychomotor Assessment",
                "headers": ["Item", "Side / Context", "Score (0–4)", "Notes"],
                "rows": [
                    ["Psychomotor retardation", "—", "", "Speech latency, slow movement"],
                    ["Psychomotor agitation", "—", "", "Restlessness, handwringing"],
                    ["Speech rate & volume", "—", "", "0=Normal, 4=Mute/whisper"],
                    ["Reaction time (tapping test)", "L / R", "", ""],
                    ["Gait observation", "—", "", "0=Normal, 4=Shuffling/stooped"],
                ],
            },
            {
                "title": "A3. Neurovegetative Symptoms",
                "headers": ["Symptom", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Sleep disturbance", "0=None, 4=Total insomnia", "", f"{CB} Insomnia  {CB} Hypersomnia"],
                    ["Appetite change", "0=None, 4=Complete loss/hyperphagia", "", "Weight: ___ kg  Change: ___"],
                    ["Fatigue / energy level", "0=None, 4=Cannot function", "", ""],
                    ["Libido reduction", "0=None, 4=Complete absence", "", ""],
                    ["Somatic complaints", "0=None, 4=Disabling", "", ""],
                ],
            },
            {
                "title": "A4. Anhedonia & Motivation",
                "headers": ["Item", "Context", "Score (0–4)", "Notes"],
                "rows": [
                    ["Anhedonia — pleasurable activities", "0=Normal, 4=Complete", "", ""],
                    ["Social withdrawal", "0=None, 4=Complete isolation", "", ""],
                    ["Motivation for daily tasks", "0=Normal, 4=None", "", ""],
                    ["Engagement in hobbies", "0=Normal, 4=Completely ceased", "", ""],
                ],
            },
            {
                "title": "A5. Suicidality & Safety Screen",
                "headers": ["Item", "Response", "Score (0–4)", "Action Required"],
                "rows": [
                    ["Passive suicidal ideation", f"{CB} None  {CB} Present", "", ""],
                    ["Active suicidal ideation", f"{CB} None  {CB} Present", "", f"{CB} Immediate referral if present"],
                    ["Suicide plan", f"{CB} None  {CB} Present", "", "Document plan details"],
                    ["Intent / means", f"{CB} None  {CB} Present", "", "Safety plan required"],
                    ["Recent self-harm", f"{CB} None  {CB} Present", "", ""],
                ],
            },
            {
                "title": "A6. Functional Impairment",
                "headers": ["Domain", "Level of Impairment", "Score (0–4)", "Notes"],
                "rows": [
                    ["Occupational / academic functioning", "0=None, 4=Unable to work", "", ""],
                    ["Social functioning", "0=None, 4=Complete isolation", "", ""],
                    ["Self-care & ADLs", "0=None, 4=Dependent on others", "", ""],
                    ["Concentration / cognition", "0=None, 4=Cannot sustain attention", "", ""],
                    ["Global functional level (GAF estimate)", "0–100 scale", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["PHQ-9 (Patient Health Questionnaire)", "___ / 27", "≥10 moderate", f"{CB} Minimal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["HDRS-17 (Hamilton Depression Rating Scale)", "___ / 52", "≥18 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["BDI-II (Beck Depression Inventory)", "___ / 63", "≥29 severe", ""],
            ["MoCA (Cognitive Screening)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
            ["Trail Making A", "___ sec", "", ""],
            ["Trail Making B", "___ sec", "", ""],
        ],
        "section_d_rows": [
            ["GAD-7 (Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["C-SSRS (Suicidality)", "___", "Any score → action", f"{CB} No ideation  {CB} Ideation  {CB} Behaviour"],
            ["Sheehan Disability Scale", "___ / 30", "≥5 significant", f"{CB} Low  {CB} Moderate  {CB} High impairment"],
            ["WHOQOL-BREF (Quality of Life)", "___", "", ""],
            ["Sleep quality (subjective)", "", "", f"{CB} Good  {CB} Disturbed  {CB} Severe insomnia"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Melancholic  {CB} Atypical  {CB} Anxious  {CB} Psychomotor-retarded"],
            ["Additional Specifier", f"{CB} Treatment-resistant  {CB} Cognitive  {CB} First episode  {CB} Recurrent"],
            ["Episode Severity", f"{CB} Mild  {CB} Moderate  {CB} Severe  {CB} With psychotic features"],
            ["Chronicity", "___ years since first episode"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} PHQ-9  {CB} HDRS-17  {CB} BDI-II  {CB} MoCA  {CB} GAD-7"],
        ],
    },

    # 2 ─ ANXIETY ─────────────────────────────────────────────────────────────
    {
        "slug": "anxiety",
        "display_name": "Generalized Anxiety Disorder (GAD)",
        "status_field_1": "Anxiety Duration",
        "status_field_2": "Disorder Chronicity",
        "section_a_title": "Anxiety & Somatic Examination",
        "section_a_subsections": [
            {
                "title": "A1. Anxiety Severity Assessment",
                "headers": ["Item", "Observation / Rating", "Score (0–4)", "Notes"],
                "rows": [
                    ["Subjective anxiety level", "0=None, 4=Incapacitating", "", ""],
                    ["Observed anxiety (clinician-rated)", "0=Calm, 4=Panic-level", "", ""],
                    ["Anticipatory anxiety", "0=None, 4=Constant", "", ""],
                    ["Generalised worry — frequency", "0=Rare, 4=Uncontrollable constant", "", ""],
                    ["Duration of worry episodes", "0=<5 min, 4=>2 hours/day", "", ""],
                ],
            },
            {
                "title": "A2. Somatic Tension Examination",
                "headers": ["Item", "Side / Region", "Score (0–4)", "Notes"],
                "rows": [
                    ["Muscle tension (neck/shoulders)", "—", "", "Palpation"],
                    ["Muscle tension (jaw/facial)", "—", "", "Bruxism noted?"],
                    ["Tremulousness / fine tremor", "L / R hands", "", ""],
                    ["Restlessness / inability to sit still", "—", "", ""],
                    ["Physical tension — global", "—", "", "0=None, 4=Rigid throughout"],
                ],
            },
            {
                "title": "A3. Autonomic Symptoms",
                "headers": ["Symptom", "Frequency", "Score (0–4)", "Notes"],
                "rows": [
                    ["Palpitations / tachycardia", "0=Never, 4=Constant", "", "HR at rest: ___"],
                    ["Sweating / diaphoresis", "0=None, 4=Drenching", "", ""],
                    ["Dyspnoea / chest tightness", "0=None, 4=Disabling", "", ""],
                    ["Gastrointestinal symptoms", "0=None, 4=Daily nausea/pain", "", ""],
                    ["Dizziness / lightheadedness", "0=None, 4=Near-syncope", "", ""],
                ],
            },
            {
                "title": "A4. Avoidance Behaviours",
                "headers": ["Behaviour", "Context", "Score (0–4)", "Notes"],
                "rows": [
                    ["Avoidance of triggers", "0=None, 4=Housebound", "", ""],
                    ["Social avoidance", "0=None, 4=Complete isolation", "", ""],
                    ["Reassurance-seeking behaviour", "0=None, 4=Constant", "", ""],
                    ["Procrastination / task-avoidance", "0=None, 4=All tasks avoided", "", ""],
                ],
            },
            {
                "title": "A5. Worry Frequency & Content",
                "headers": ["Item", "Rating", "Score (0–4)", "Notes"],
                "rows": [
                    ["Number of worry domains", "Health / Finance / Work / Relationships", "", ""],
                    ["Difficulty controlling worry", "0=Controllable, 4=Impossible", "", ""],
                    ["Intrusive thoughts", "0=None, 4=Constant", "", ""],
                    ["Catastrophising tendency", "0=None, 4=Pervasive", "", ""],
                ],
            },
            {
                "title": "A6. Sleep Disruption",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Sleep onset difficulty (worry)", "0=None, 4=>2 hours latency", "", ""],
                    ["Nocturnal awakening (anxiety)", "0=None, 4=Multiple/night", "", ""],
                    ["Early morning waking", "0=None, 4=>2 hours early", "", ""],
                    ["Daytime fatigue from poor sleep", "0=None, 4=Cannot function", "", ""],
                    ["Sleep quality — subjective", "0=Excellent, 4=Terrible", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Montreal Cognitive Assessment)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["Trail Making A", "___ sec", "", ""],
            ["Trail Making B", "___ sec", "", "Attention / set-shifting"],
            ["Digit Span (forward)", "___", "≥5 normal", ""],
            ["Digit Span (backward)", "___", "≥3 normal", ""],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
        ],
        "section_d_rows": [
            ["GAD-7 (Generalized Anxiety Disorder Scale)", "___ / 21", "≥10 moderate", f"{CB} Minimal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["HAMA (Hamilton Anxiety Rating Scale)", "___ / 56", "≥18 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["BAI (Beck Anxiety Inventory)", "___ / 63", "≥26 severe", ""],
            ["PHQ-9 (Comorbid Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["PSQI (Sleep Quality)", "___ / 21", ">5 poor sleep", f"{CB} Good  {CB} Poor sleeper"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Cognitive-worry  {CB} Somatic-tension  {CB} Mixed  {CB} Panic-prone"],
            ["Additional Specifier", f"{CB} Social overlap  {CB} Insomnia-linked  {CB} Health anxiety  {CB} OCD features"],
            ["Severity", f"{CB} Mild  {CB} Moderate  {CB} Severe  {CB} Panic disorder comorbidity"],
            ["Chronicity", "___ years since onset"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} GAD-7  {CB} HAMA  {CB} BAI  {CB} PHQ-9  {CB} PSQI"],
        ],
    },

    # 3 ─ ADHD ─────────────────────────────────────────────────────────────────
    {
        "slug": "adhd",
        "display_name": "Attention Deficit Hyperactivity Disorder (ADHD)",
        "status_field_1": "Age at Diagnosis",
        "status_field_2": "ADHD Presentation",
        "section_a_title": "Attention & Executive Function Examination",
        "section_a_subsections": [
            {
                "title": "A1. Attention & Concentration Testing",
                "headers": ["Test", "Result / Observation", "Score (0–4)", "Notes"],
                "rows": [
                    ["Digit Span Forward", "___ / 9", "", "Sustained attention"],
                    ["Digit Span Backward", "___ / 8", "", "Working memory"],
                    ["Serial 7s subtraction", "___ errors / ___ time", "", "0=Flawless, 4=Cannot attempt"],
                    ["Continuous performance (tap on A)", "", "", "Commission & omission errors"],
                    ["Distractibility during interview", "0=None, 4=Cannot maintain", "", ""],
                ],
            },
            {
                "title": "A2. Hyperactivity Observation",
                "headers": ["Item", "Context", "Score (0–4)", "Notes"],
                "rows": [
                    ["Motor restlessness during interview", "0=None, 4=Constant movement", "", ""],
                    ["Fidgeting (hands/feet/legs)", "0=None, 4=Cannot stop", "", ""],
                    ["Leaving seat inappropriately", "0=None, 4=Repeatedly", "", ""],
                    ["Talkativeness / verbosity", "0=Normal, 4=Cannot interrupt", "", ""],
                    ["On-the-go / driven behaviour", "0=None, 4=Always", "", ""],
                ],
            },
            {
                "title": "A3. Impulsivity Assessment",
                "headers": ["Item", "Rating", "Score (0–4)", "Notes"],
                "rows": [
                    ["Interrupting / blurting out answers", "0=None, 4=Constant", "", ""],
                    ["Difficulty waiting turn", "0=None, 4=Cannot wait", "", ""],
                    ["Impulsive decision-making", "0=None, 4=Daily harmful decisions", "", ""],
                    ["Emotional impulsivity", "0=None, 4=Explosive outbursts", "", ""],
                    ["Risk-taking behaviour", "0=None, 4=Dangerous level", "", ""],
                ],
            },
            {
                "title": "A4. Executive Function Screen",
                "headers": ["Function", "Test / Observation", "Score (0–4)", "Notes"],
                "rows": [
                    ["Planning & organisation", "TMT-B / clinical obs", "", ""],
                    ["Working memory", "Digit span backward", "", ""],
                    ["Cognitive flexibility", "TMT B-A difference", "", ""],
                    ["Inhibitory control", "Stroop interference (est)", "", ""],
                    ["Task initiation & follow-through", "0=Normal, 4=Cannot initiate", "", ""],
                ],
            },
            {
                "title": "A5. Emotional Dysregulation",
                "headers": ["Item", "Frequency", "Score (0–4)", "Notes"],
                "rows": [
                    ["Emotional lability", "0=Stable, 4=Multiple daily", "", ""],
                    ["Low frustration tolerance", "0=Normal, 4=Extreme", "", ""],
                    ["Rejection sensitive dysphoria", "0=None, 4=Incapacitating", "", ""],
                    ["Mood swings duration", "0=Hours, 4=Days", "", ""],
                ],
            },
            {
                "title": "A6. Time Management & Daily Function",
                "headers": ["Domain", "Impairment Level", "Score (0–4)", "Notes"],
                "rows": [
                    ["Time blindness / poor time sense", "0=None, 4=Cannot manage time", "", ""],
                    ["Chronic lateness / missing deadlines", "0=Rare, 4=Daily", "", ""],
                    ["Task completion rate", "0=Good, 4=Rarely completes", "", ""],
                    ["Organisational skills", "0=Good, 4=Chaotic/dysfunctional", "", ""],
                    ["Impact on occupational function", "0=None, 4=Job loss/academic failure", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["Digit Span (Forward + Backward)", "___ / 18", "", "Working memory proxy"],
            ["Trail Making A", "___ sec", "<29 sec normal", "Processing speed"],
            ["Trail Making B", "___ sec", "<75 sec normal", "Cognitive flexibility"],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
            ["Clock Drawing Test", "___ / 10", "", "Executive / visuospatial"],
            ["BRIEF-A (self-report summary T-score)", "___", "T>65 elevated", f"{CB} Normal  {CB} Elevated"],
        ],
        "section_d_rows": [
            ["ASRS v1.1 (Adult ADHD Self-Report)", "___ / 72", "≥14 Part A = likely", f"{CB} Below threshold  {CB} Likely ADHD"],
            ["Conners Adult ADHD Rating Scale", "___", "T>65 clinical", f"{CB} Normal  {CB} Inattentive  {CB} Hyperactive  {CB} Combined"],
            ["PHQ-9 (Comorbid Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Comorbid Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
            ["WURS (Wender Utah — childhood retrospective)", "___", "≥46 = likely ADHD", ""],
        ],
        "section_e_rows": [
            ["Identified Presentation", f"{CB} Inattentive (ADHD-I)  {CB} Hyperactive-Impulsive (ADHD-HI)  {CB} Combined (ADHD-C)"],
            ["Additional Specifier", f"{CB} Executive dysfunction  {CB} Emotional dysregulation  {CB} SCT  {CB} Sluggish cognitive tempo"],
            ["Severity", f"{CB} Mild  {CB} Moderate  {CB} Severe"],
            ["Age at Symptom Onset", "___ years (childhood onset required)"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} ASRS  {CB} Conners  {CB} BRIEF-A  {CB} TMT  {CB} PHQ-9"],
        ],
    },

    # 4 ─ ALZHEIMER'S / MCI ───────────────────────────────────────────────────
    {
        "slug": "alzheimers",
        "display_name": "Alzheimer's Disease / Mild Cognitive Impairment (MCI)",
        "status_field_1": "CDR Global Score",
        "status_field_2": "MMSE / MoCA Score",
        "section_a_title": "Cognitive & Functional Examination",
        "section_a_subsections": [
            {
                "title": "A1. Memory Assessment",
                "headers": ["Test", "Result", "Score (0–4)", "Notes"],
                "rows": [
                    ["3-word recall (immediate)", "___ / 3", "", "Words: ___"],
                    ["3-word recall (5 min delay)", "___ / 3", "", "0=All correct, 4=None recalled"],
                    ["Story recall (logical memory)", "___", "", "Immediate vs delayed"],
                    ["Autobiographical memory", "0=Intact, 4=Severely impaired", "", "Recent vs remote"],
                    ["Prospective memory", "0=Normal, 4=Cannot plan/remember", "", ""],
                ],
            },
            {
                "title": "A2. Orientation",
                "headers": ["Domain", "Response", "Score (0–4)", "Notes"],
                "rows": [
                    ["Orientation to time (date/day/month/year)", "___ / 5", "", ""],
                    ["Orientation to place (city/country/building)", "___ / 3", "", ""],
                    ["Orientation to person (own name/age/DOB)", "___ / 3", "", ""],
                    ["Temporal orientation (morning/afternoon/evening)", "", "", ""],
                ],
            },
            {
                "title": "A3. Visuospatial Function",
                "headers": ["Test", "Result", "Score (0–4)", "Notes"],
                "rows": [
                    ["Clock Drawing Test (CDT)", "___ / 10", "", "Command + copy"],
                    ["Intersecting pentagons", f"{CB} Pass  {CB} Fail", "", ""],
                    ["Rey Figure (copy)", "___", "", "Optional if available"],
                    ["Visual object recognition", "0=Normal, 4=Agnosia", "", ""],
                    ["Spatial navigation", "0=Normal, 4=Gets lost in familiar", "", ""],
                ],
            },
            {
                "title": "A4. Language & Naming",
                "headers": ["Test", "Result", "Score (0–4)", "Notes"],
                "rows": [
                    ["Boston Naming Test (short: 15 items)", "___ / 15", "", ""],
                    ["Verbal fluency — animals (1 min)", "___", "≥15 normal", ""],
                    ["Verbal fluency — letter F (1 min)", "___", "≥12 normal", ""],
                    ["Repetition (complex sentence)", f"{CB} Pass  {CB} Fail", "", ""],
                    ["Comprehension (3-step command)", "___ / 3", "", ""],
                ],
            },
            {
                "title": "A5. Executive Function",
                "headers": ["Test", "Result", "Score (0–4)", "Notes"],
                "rows": [
                    ["Trail Making A", "___ sec", "<29 sec", ""],
                    ["Trail Making B", "___ sec", "<75 sec", ""],
                    ["Frontal Assessment Battery (FAB)", "___ / 18", "≥15 normal", ""],
                    ["Luria alternating sequences", f"{CB} Pass  {CB} Fail", "", ""],
                    ["Go/No-Go", "___ / 10", "", "Inhibitory control"],
                ],
            },
            {
                "title": "A6. Activities of Daily Living",
                "headers": ["ADL Domain", "Dependence Level", "Score (0–4)", "Notes"],
                "rows": [
                    ["Basic ADLs (dressing, bathing, feeding)", "0=Independent, 4=Fully dependent", "", ""],
                    ["Instrumental ADLs (finances, medications)", "0=Independent, 4=Fully dependent", "", ""],
                    ["Driving / transportation", f"{CB} Drives  {CB} Ceased  {CB} Never drove", "", ""],
                    ["Home management", "0=Independent, 4=Unable", "", ""],
                    ["Use of technology/telephone", "0=Normal, 4=Cannot use", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Montreal Cognitive Assessment)", "___ / 30", "18–25 MCI; <18 dementia", f"{CB} Normal  {CB} MCI  {CB} Dementia"],
            ["MMSE (Mini-Mental State Examination)", "___ / 30", "24–27 MCI; <24 dementia", f"{CB} Normal  {CB} MCI  {CB} Dementia"],
            ["ADAS-Cog (Alzheimer Disease Assessment Scale-Cog)", "___ / 70", ">18 impaired", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["CDR (Clinical Dementia Rating)", "___ (0/0.5/1/2/3)", "0.5=MCI; 1=Mild AD", f"{CB} 0  {CB} 0.5  {CB} 1  {CB} 2  {CB} 3"],
            ["Clock Drawing Test", "___ / 10", "", ""],
            ["FAB (Frontal Assessment Battery)", "___ / 18", "≥15 normal", ""],
        ],
        "section_d_rows": [
            ["GDS-15 (Geriatric Depression Scale)", "___ / 15", "≥6 depression", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["NPI (Neuropsychiatric Inventory — carer)", "___ / 144", "", "Domains: ___"],
            ["Apathy Evaluation Scale", "___", "", f"{CB} Normal  {CB} Apathetic"],
            ["CSDD (Cornell Scale for Depression in Dementia)", "___", "≥8 significant", f"{CB} Normal  {CB} Present"],
            ["Sleep disturbance (subjective / carer report)", "", "", f"{CB} Good  {CB} Disturbed  {CB} Sundowning"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Amnestic MCI  {CB} Multi-domain MCI  {CB} Early-onset AD  {CB} Late-onset AD"],
            ["Additional Specifier", f"{CB} PCA (Posterior Cortical Atrophy)  {CB} Logopenic  {CB} Frontal variant"],
            ["Dementia Severity", f"{CB} MCI  {CB} Mild AD  {CB} Moderate AD  {CB} Severe AD"],
            ["Disease Duration", "___ years since symptom onset"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} MoCA  {CB} MMSE  {CB} ADAS-Cog  {CB} CDR  {CB} NPI"],
        ],
    },

    # 5 ─ STROKE REHAB ─────────────────────────────────────────────────────────
    {
        "slug": "stroke_rehab",
        "display_name": "Post-Stroke Rehabilitation",
        "status_field_1": "Stroke Date / Days Post-Stroke",
        "status_field_2": "mRS Score",
        "section_a_title": "Neurological & Motor Rehabilitation Examination",
        "section_a_subsections": [
            {
                "title": "A1. Upper Limb Motor Strength (MRC Scale)",
                "headers": ["Muscle Group", "Side", "Score (0–5 MRC)", "Notes"],
                "rows": [
                    ["Shoulder abduction / flexion", "L / R", "", ""],
                    ["Elbow flexion / extension", "L / R", "", ""],
                    ["Wrist extension / flexion", "L / R", "", ""],
                    ["Finger extension / grip", "L / R", "", "Grip dynamometer: ___"],
                    ["Pinch strength", "L / R", "", ""],
                ],
            },
            {
                "title": "A2. Lower Limb Motor Strength (MRC Scale)",
                "headers": ["Muscle Group", "Side", "Score (0–5 MRC)", "Notes"],
                "rows": [
                    ["Hip flexion / extension", "L / R", "", ""],
                    ["Knee extension / flexion", "L / R", "", ""],
                    ["Ankle dorsiflexion / plantarflexion", "L / R", "", ""],
                    ["Toe extension", "L / R", "", ""],
                    ["Overall lower limb function", "L / R", "", ""],
                ],
            },
            {
                "title": "A3. Spasticity Assessment (Modified Ashworth Scale)",
                "headers": ["Muscle Group", "Side", "MAS Score (0–4)", "Notes"],
                "rows": [
                    ["Elbow flexors", "L / R", "", "0=No increase, 4=Rigid"],
                    ["Wrist flexors", "L / R", "", ""],
                    ["Finger flexors", "L / R", "", ""],
                    ["Knee extensors", "L / R", "", ""],
                    ["Ankle plantar flexors", "L / R", "", "Clonus: present / absent"],
                ],
            },
            {
                "title": "A4. Sensory & Perceptual Deficits",
                "headers": ["Modality", "Side", "Score (0–4)", "Notes"],
                "rows": [
                    ["Light touch (cotton wool)", "L / R face/arm/leg", "", ""],
                    ["Proprioception (joint position)", "L / R", "", "Finger/toe"],
                    ["Two-point discrimination", "L / R", "", "Fingertip: ___mm"],
                    ["Hemispatial neglect screen", "—", "", f"{CB} Absent  {CB} Present (line bisection)"],
                    ["Hemianopia screen", "—", "", f"{CB} Absent  {CB} Present  {CB} Refer ophthalmology"],
                ],
            },
            {
                "title": "A5. Speech & Language",
                "headers": ["Domain", "Finding", "Score (0–4)", "Notes"],
                "rows": [
                    ["Fluency of speech", "0=Normal, 4=Non-fluent/mute", "", ""],
                    ["Comprehension", "0=Normal, 4=Severely impaired", "", ""],
                    ["Repetition", "0=Normal, 4=Cannot repeat", "", ""],
                    ["Naming (Boston Naming — short)", "___ / 15", "", ""],
                    ["Dysarthria severity", "0=None, 4=Unintelligible", "", f"{CB} SLT referral if ≥2"],
                ],
            },
            {
                "title": "A6. Functional Mobility & Swallowing",
                "headers": ["Test", "Result", "Normative / Score", "Notes"],
                "rows": [
                    ["10-Metre Walk Test (10MWT)", "___ seconds", "<10s community ambulant", ""],
                    ["Timed Up and Go (TUG)", "___ seconds", "<12s normal", ""],
                    ["Barthel Index", "___ / 100", "100=Independent", ""],
                    ["Swallowing screen (3oz water test)", f"{CB} Pass  {CB} Fail", "", f"{CB} Refer SLT if fail"],
                    ["Berg Balance Scale", "___ / 56", "≥45 low fall risk", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Montreal Cognitive Assessment)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Post-stroke cognitive impairment"],
            ["MMSE", "___ / 30", "≥24 normal", ""],
            ["Trail Making A", "___ sec", "", "Attention / processing"],
            ["Trail Making B", "___ sec", "", "Executive function"],
            ["Clock Drawing Test", "___ / 10", "", ""],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
        ],
        "section_d_rows": [
            ["PHQ-9 (Post-Stroke Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["HADS (Hospital Anxiety & Depression Scale)", "___ / 42", "≥8 per subscale", ""],
            ["Apathy Evaluation Scale", "___", "", f"{CB} Normal  {CB} Apathetic"],
            ["NIHSS (NIH Stroke Scale)", "___ / 42", "0=No deficit", "Motor/sensory/speech/neglect"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Motor UL  {CB} Motor LL  {CB} Aphasia  {CB} Neglect  {CB} Cognitive  {CB} Spasticity"],
            ["Stroke Type", f"{CB} Ischaemic  {CB} Haemorrhagic  {CB} Lacunar  {CB} Cortical  {CB} Subcortical"],
            ["Hemi-side Affected", f"{CB} Left hemiplegia  {CB} Right hemiplegia  {CB} Bilateral  {CB} Brainstem"],
            ["Rehab Stage", f"{CB} Acute (<3 months)  {CB} Subacute (3–6 months)  {CB} Chronic (>6 months)"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} FMA  {CB} NIHSS  {CB} mRS  {CB} Barthel  {CB} MoCA"],
        ],
    },

    # 6 ─ TBI ──────────────────────────────────────────────────────────────────
    {
        "slug": "tbi",
        "display_name": "Traumatic Brain Injury (TBI)",
        "status_field_1": "Injury Date / Days Post-Injury",
        "status_field_2": "Initial GCS Score",
        "section_a_title": "TBI-Specific Neurological & Cognitive Examination",
        "section_a_subsections": [
            {
                "title": "A1. Consciousness & Arousal",
                "headers": ["Item", "Finding", "Score (0–4)", "Notes"],
                "rows": [
                    ["Current GCS (E/V/M)", "___ / 15", "", "E:___ V:___ M:___"],
                    ["Orientation (time/place/person)", "___ / 3", "", ""],
                    ["Post-traumatic amnesia status", "", "", f"{CB} Resolved  {CB} Ongoing — ___ days"],
                    ["Level of arousal", "0=Alert, 4=Comatose", "", ""],
                    ["Awareness of deficits (anosognosia)", "0=None, 4=Complete unawareness", "", ""],
                ],
            },
            {
                "title": "A2. Post-Traumatic Amnesia (PTA)",
                "headers": ["Domain", "Result", "Score (0–4)", "Notes"],
                "rows": [
                    ["Anterograde amnesia severity", "0=None, 4=Dense", "", "Galveston Orientation: ___"],
                    ["Retrograde amnesia extent", "0=None, 4=>1 year", "", ""],
                    ["Day-to-day memory continuity", "0=Normal, 4=No continuity", "", ""],
                    ["GOAT score (if applicable)", "___ / 100", "≥76 = out of PTA", ""],
                ],
            },
            {
                "title": "A3. Post-Traumatic Headache",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Headache frequency (days/week)", "0=0, 4=Daily/constant", "", ""],
                    ["Headache severity (NRS)", "___ / 10", "", "Peak: ___ / 10"],
                    ["Headache character", "", "", f"{CB} Migraine-like  {CB} Tension-type  {CB} Cervicogenic"],
                    ["Photophobia / phonophobia", "0=None, 4=Disabling", "", ""],
                    ["Impact on function", "0=None, 4=Housebound", "", "HIT-6: ___ / 78"],
                ],
            },
            {
                "title": "A4. Balance & Vestibular Assessment",
                "headers": ["Test", "Result", "Score (0–4)", "Notes"],
                "rows": [
                    ["Romberg (eyes open / closed)", f"{CB} Normal  {CB} Abnormal", "", ""],
                    ["Tandem stance (10 sec)", f"{CB} Pass  {CB} Fail", "", ""],
                    ["BESS (Balance Error Scoring System)", "___ errors", "0–5 normal", ""],
                    ["Dizziness Handicap Inventory (DHI)", "___ / 100", ">28 moderate", ""],
                    ["Gait observation — tandem walk", "0=Normal, 4=Cannot perform", "", ""],
                ],
            },
            {
                "title": "A5. Executive Function & Processing Speed",
                "headers": ["Test", "Result", "Score (0–4)", "Notes"],
                "rows": [
                    ["Trail Making A", "___ sec", "<29 sec normal", ""],
                    ["Trail Making B", "___ sec", "<75 sec normal", ""],
                    ["Verbal Fluency — letter F", "___", "≥12 normal", ""],
                    ["Digit Span (backward)", "___", "≥3 normal", ""],
                    ["Symbol Digit Modalities Test (SDMT)", "___", ">44 normal", "Processing speed"],
                ],
            },
            {
                "title": "A6. Behavioural & Emotional Changes",
                "headers": ["Symptom", "Frequency / Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Irritability / aggression", "0=None, 4=Daily dangerous", "", ""],
                    ["Emotional lability", "0=None, 4=Constant", "", ""],
                    ["Impulsivity", "0=None, 4=Cannot control", "", ""],
                    ["Disinhibition", "0=None, 4=Severe social impairment", "", ""],
                    ["Apathy / motivation loss", "0=None, 4=Complete", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["RBANS (Repeatable Battery for Neuropsychological Status)", "___ (index score)", ">80 normal range", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe impairment"],
            ["MoCA", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["SDMT (Symbol Digit Modalities Test)", "___", ">44 normal", "Processing speed"],
            ["Trail Making A", "___ sec", "", ""],
            ["Trail Making B", "___ sec", "", ""],
            ["Digit Span (F + B)", "___ / 18", "", ""],
        ],
        "section_d_rows": [
            ["PHQ-9 (Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
            ["PCL-5 (PTSD Checklist)", "___ / 80", "≥33 probable PTSD", f"{CB} Below threshold  {CB} Probable PTSD"],
            ["NSI (Neurobehavioural Symptom Inventory)", "___ / 44", "", "22 post-concussive symptoms"],
            ["RPQ (Rivermead Post-Concussion Symptoms)", "___ / 64", "≥16 significant", ""],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Mild/Concussion  {CB} Moderate  {CB} Severe  {CB} DOC (Disorder of Consciousness)"],
            ["Dominant Syndrome", f"{CB} Executive  {CB} Post-traumatic headache  {CB} Behavioural  {CB} Vestibular/balance"],
            ["Injury Mechanism", f"{CB} Blunt trauma  {CB} Blast/explosion  {CB} MVA  {CB} Sports  {CB} Fall"],
            ["Time Post-Injury", f"{CB} Acute (<2 weeks)  {CB} Subacute (2w–3m)  {CB} Chronic (>3 months)"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} GOSE  {CB} RPQ  {CB} NSI  {CB} RBANS  {CB} PHQ-9"],
        ],
    },

    # 7 ─ CHRONIC PAIN ─────────────────────────────────────────────────────────
    {
        "slug": "chronic_pain",
        "display_name": "Chronic Pain / Fibromyalgia",
        "status_field_1": "Pain Duration (years/months)",
        "status_field_2": "Primary Diagnosis",
        "section_a_title": "Chronic Pain & Sensory Examination",
        "section_a_subsections": [
            {
                "title": "A1. Pain Location & Intensity Mapping",
                "headers": ["Region", "NRS (0–10)", "Score (0–4)", "Notes"],
                "rows": [
                    ["Head / face / jaw", "___ / 10", "", ""],
                    ["Neck / cervical", "___ / 10", "", ""],
                    ["Upper back / thoracic", "___ / 10", "", ""],
                    ["Lower back / lumbar", "___ / 10", "", ""],
                    ["Upper extremities", "___ / 10", "", "L:___ R:___"],
                    ["Lower extremities", "___ / 10", "", "L:___ R:___"],
                    ["Widespread (≥4 body regions)", f"{CB} Yes  {CB} No", "", "WPI score: ___"],
                ],
            },
            {
                "title": "A2. Tender Point Examination (ACR Criteria)",
                "headers": ["Tender Point Site", "Side", "Tender (0–4)", "Notes"],
                "rows": [
                    ["Occiput (suboccipital insertions)", "L / R", "", "4 kg pressure"],
                    ["Low cervical (C5–C7 anterior)", "L / R", "", ""],
                    ["Trapezius (midpoint upper border)", "L / R", "", ""],
                    ["Supraspinatus (above scapular spine)", "L / R", "", ""],
                    ["Second rib (costochondral junction)", "L / R", "", ""],
                    ["Lateral epicondyle (2 cm distal)", "L / R", "", ""],
                    ["Gluteal (upper outer quadrant)", "L / R", "", ""],
                    ["Greater trochanter", "L / R", "", ""],
                    ["Knee (medial fat pad)", "L / R", "", ""],
                    ["Total tender points", "___ / 18", "", "≥11 = fibromyalgia (ACR 1990)"],
                ],
            },
            {
                "title": "A3. Sensory Testing (Quantitative)",
                "headers": ["Modality", "Side / Region", "Score (0–4)", "Notes"],
                "rows": [
                    ["Allodynia (light touch)", "L / R", "", f"{CB} Absent  {CB} Present — region: ___"],
                    ["Hyperalgesia (pinprick)", "L / R", "", f"{CB} Absent  {CB} Present"],
                    ["Temporal summation (wind-up)", "—", "", "0=None, 4=Severe"],
                    ["Cold allodynia", "L / R", "", ""],
                    ["Conditioned Pain Modulation (CPM)", "—", "", "Adequate  /  Impaired"],
                ],
            },
            {
                "title": "A4. Functional Capacity",
                "headers": ["Domain", "Impairment Level", "Score (0–4)", "Notes"],
                "rows": [
                    ["Mobility — walking distance", "0=Normal, 4=Bed-bound", "", "___ metres"],
                    ["Upper limb function", "0=Normal, 4=Cannot use", "", ""],
                    ["Work / occupational capacity", "0=Full-time, 4=Cannot work", "", ""],
                    ["ADLs (dressing, bathing)", "0=Independent, 4=Dependent", "", ""],
                    ["Exercise tolerance", "0=Normal, 4=Any activity intolerable", "", ""],
                ],
            },
            {
                "title": "A5. Sleep Quality",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Sleep onset difficulty (pain)", "0=None, 4=>2h latency", "", ""],
                    ["Nocturnal pain awakenings", "0=None, 4=>5/night", "", ""],
                    ["Non-restorative sleep", "0=Refreshed, 4=Always exhausted", "", ""],
                    ["PSQI global score", "___ / 21", ">5 poor sleep", ""],
                ],
            },
            {
                "title": "A6. Fatigue Severity",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Fatigue severity (NRS)", "___ / 10", "", ""],
                    ["Post-exertional malaise", "0=None, 4=Days of crash", "", ""],
                    ["Mental fatigue / brain fog", "0=None, 4=Cannot think", "", ""],
                    ["FSS (Fatigue Severity Scale)", "___ / 63", "≥36 significant", ""],
                    ["Impact on daily activities", "0=None, 4=Housebound", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Cognitive — pain & medication effects)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["Trail Making A", "___ sec", "", "Processing speed"],
            ["Trail Making B", "___ sec", "", "Executive / set-shifting"],
            ["Digit Span (forward + backward)", "___ / 18", "", "Working memory"],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
        ],
        "section_d_rows": [
            ["PHQ-9 (Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
            ["PCS (Pain Catastrophising Scale)", "___ / 52", "≥30 high", f"{CB} Low  {CB} Moderate  {CB} High catastrophising"],
            ["FIQ-R (Revised Fibromyalgia Impact Questionnaire)", "___ / 100", ">59 severe", ""],
            ["BPI (Brief Pain Inventory)", "Severity: ___ / 10  Interference: ___ / 10", "", ""],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Nociceptive  {CB} Neuropathic  {CB} Nociplastic  {CB} Fibromyalgia  {CB} CRPS"],
            ["Additional Specifier", f"{CB} Musculoskeletal  {CB} Visceral  {CB} Central sensitisation  {CB} Mixed"],
            ["Severity", f"{CB} Mild (NRS 1–3)  {CB} Moderate (4–6)  {CB} Severe (7–10)"],
            ["Pain Duration", "___ months / years"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} VAS  {CB} BPI  {CB} FIQ-R  {CB} PCS  {CB} PHQ-9"],
        ],
    },

    # 8 ─ PTSD ─────────────────────────────────────────────────────────────────
    {
        "slug": "ptsd",
        "display_name": "Post-Traumatic Stress Disorder (PTSD)",
        "status_field_1": "Trauma Date / Years Since Event",
        "status_field_2": "Trauma Type",
        "section_a_title": "PTSD Symptom & Trauma Response Examination",
        "section_a_subsections": [
            {
                "title": "A1. Trauma Exposure Screen",
                "headers": ["Item", "Response", "Score (0–4)", "Notes"],
                "rows": [
                    ["Nature of index trauma", "", "", f"{CB} Combat  {CB} Assault  {CB} Accident  {CB} Disaster  {CB} Other"],
                    ["Number of traumatic events", "___", "", "Polytrauma: Y / N"],
                    ["Time since index trauma", "___ years / months", "", ""],
                    ["Subjective perceived threat to life", "0=Low, 4=Extreme", "", ""],
                    ["Degree of social support at time", "0=Strong, 4=None", "", ""],
                ],
            },
            {
                "title": "A2. Re-experiencing Symptoms",
                "headers": ["Symptom", "Frequency", "Score (0–4)", "Notes"],
                "rows": [
                    ["Intrusive memories / flashbacks", "0=Never, 4=Daily/hourly", "", ""],
                    ["Trauma-related nightmares", "0=Never, 4=Every night", "", ""],
                    ["Psychological distress to cues", "0=None, 4=Incapacitating", "", ""],
                    ["Physiological reactivity to cues", "0=None, 4=Severe (panic)", "", "HR response observed: ___"],
                    ["Dissociative flashback episodes", "0=Never, 4=Multiple daily", "", ""],
                ],
            },
            {
                "title": "A3. Avoidance Behaviours",
                "headers": ["Behaviour", "Degree", "Score (0–4)", "Notes"],
                "rows": [
                    ["Avoidance of trauma-related thoughts", "0=None, 4=Total suppression", "", ""],
                    ["Avoidance of trauma-related places/people", "0=None, 4=Housebound", "", ""],
                    ["Emotional numbing", "0=None, 4=Complete numbing", "", ""],
                    ["Restricted range of affect", "0=Normal, 4=Flat/detached", "", ""],
                    ["Diminished interest in activities", "0=None, 4=Complete anhedonia", "", ""],
                ],
            },
            {
                "title": "A4. Hyperarousal Signs",
                "headers": ["Sign", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Hypervigilance", "0=None, 4=Constant scanning", "", "Observable during session: Y/N"],
                    ["Exaggerated startle response", "0=None, 4=Extreme", "", ""],
                    ["Irritability / anger outbursts", "0=None, 4=Daily violent", "", ""],
                    ["Concentration difficulties", "0=None, 4=Cannot focus at all", "", ""],
                    ["Sleep disturbance (trauma-related)", "0=None, 4=Cannot sleep", "", ""],
                ],
            },
            {
                "title": "A5. Dissociation Screen",
                "headers": ["Item", "Frequency", "Score (0–4)", "Notes"],
                "rows": [
                    ["Depersonalisation (detachment from self)", "0=Never, 4=Constant", "", ""],
                    ["Derealisation (unreality of environment)", "0=Never, 4=Constant", "", ""],
                    ["Dissociative amnesia (trauma gaps)", "0=None, 4=Extensive gaps", "", ""],
                    ["Identity confusion / fragmentation", "0=None, 4=Severe", "", f"{CB} Refer for C-PTSD assessment"],
                ],
            },
            {
                "title": "A6. Functional Impact",
                "headers": ["Domain", "Impairment", "Score (0–4)", "Notes"],
                "rows": [
                    ["Occupational / academic functioning", "0=None, 4=Cannot work", "", ""],
                    ["Social / relationship functioning", "0=None, 4=Complete isolation", "", ""],
                    ["Intimate / family relationships", "0=None, 4=Broken", "", ""],
                    ["Physical health impact", "0=None, 4=Severe somatisation", "", ""],
                    ["Quality of life (subjective)", "0=Good, 4=Terrible", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Cognitive — dissociation / attention impact)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["Trail Making A", "___ sec", "", "Attention / processing"],
            ["Trail Making B", "___ sec", "", "Executive function"],
            ["Digit Span (forward + backward)", "___ / 18", "", "Working memory"],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
        ],
        "section_d_rows": [
            ["PCL-5 (PTSD Checklist for DSM-5)", "___ / 80", "≥33 probable PTSD", f"{CB} Below threshold  {CB} Probable PTSD"],
            ["CAPS-5 (Clinician-Administered PTSD Scale)", "___ / 80", "≥23 probable PTSD", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["IES-R (Impact of Event Scale — Revised)", "___ / 88", "≥33 PTSD possible", ""],
            ["PHQ-9 (Comorbid Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["DES-II (Dissociative Experiences Scale)", "___", ">30 elevated", f"{CB} Normal  {CB} Elevated  {CB} Pathological dissociation"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Re-experiencing dominant  {CB} Avoidance dominant  {CB} Hyperarousal dominant"],
            ["Additional Specifier", f"{CB} Dissociative subtype  {CB} Complex PTSD (C-PTSD)  {CB} Comorbid TBI  {CB} Military/combat"],
            ["Chronicity", f"{CB} Acute (<3 months)  {CB} Chronic (≥3 months)  {CB} Delayed onset (>6 months)"],
            ["Trauma Type", f"{CB} Single incident  {CB} Repeated  {CB} Childhood  {CB} Interpersonal violence"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} CAPS-5  {CB} PCL-5  {CB} IES-R  {CB} PHQ-9  {CB} DES-II"],
        ],
    },

    # 9 ─ OCD ──────────────────────────────────────────────────────────────────
    {
        "slug": "ocd",
        "display_name": "Obsessive-Compulsive Disorder (OCD)",
        "status_field_1": "OCD Duration (years)",
        "status_field_2": "Treatment History",
        "section_a_title": "OCD Symptom Severity Examination",
        "section_a_subsections": [
            {
                "title": "A1. Obsession Severity & Frequency",
                "headers": ["Item", "Rating", "Score (0–4)", "Notes"],
                "rows": [
                    ["Time occupied by obsessions (hrs/day)", "0=None, 4=>8 hours", "", ""],
                    ["Frequency of intrusive thoughts", "0=Never, 4=Constant", "", ""],
                    ["Intensity / distress caused", "0=None, 4=Incapacitating", "", ""],
                    ["Ego-dystonicity of obsessions", "0=Not ego-dystonic, 4=Fully", "", ""],
                    ["Resistance to obsessions", "0=Always resists, 4=Never resists", "", ""],
                ],
            },
            {
                "title": "A2. Compulsion Severity & Frequency",
                "headers": ["Item", "Rating", "Score (0–4)", "Notes"],
                "rows": [
                    ["Time spent on compulsions (hrs/day)", "0=None, 4=>8 hours", "", ""],
                    ["Frequency of compulsive behaviours", "0=Never, 4=Constant", "", ""],
                    ["Distress if compulsion resisted", "0=None, 4=Extreme panic", "", ""],
                    ["Sense of compulsion as excessive", "0=Fully recognised, 4=Not recognised", "", ""],
                    ["Reduction of anxiety post-compulsion", "0=No relief, 4=Complete relief", "", ""],
                ],
            },
            {
                "title": "A3. Resistance Capacity & Control",
                "headers": ["Item", "Rating", "Score (0–4)", "Notes"],
                "rows": [
                    ["Ability to resist obsessions (ERP capacity)", "0=Can resist, 4=Cannot at all", "", ""],
                    ["Ability to delay compulsions", "0=Hours, 4=Cannot delay at all", "", ""],
                    ["Control over compulsive urges", "0=Full control, 4=No control", "", ""],
                    ["Self-monitoring ability", "0=Excellent, 4=Cannot self-monitor", "", ""],
                ],
            },
            {
                "title": "A4. Insight Assessment",
                "headers": ["Level of Insight", "Finding", "Score (0–4)", "Notes"],
                "rows": [
                    ["Recognition obsessions/compulsions are excessive", "0=Full, 4=None (overvalued ideas)", "", ""],
                    ["Attribution to OCD vs real threat", "0=OCD, 4=Believes threat is real", "", ""],
                    ["Belief change potential (metacognition)", "0=Open, 4=Fixed/delusional", "", ""],
                    ["Overvalued ideation", "0=None, 4=Delusional intensity", "", f"{CB} OBQ-44 indicated"],
                ],
            },
            {
                "title": "A5. Avoidance & Reassurance-Seeking",
                "headers": ["Behaviour", "Frequency", "Score (0–4)", "Notes"],
                "rows": [
                    ["Situational avoidance (triggers)", "0=None, 4=Housebound", "", ""],
                    ["Reassurance-seeking (others)", "0=None, 4=Constant", "", ""],
                    ["Mental neutralisation rituals", "0=None, 4=Constant", "", ""],
                    ["Accommodation by family/carers", "0=None, 4=Family lives around OCD", "", ""],
                ],
            },
            {
                "title": "A6. Functional Impairment",
                "headers": ["Domain", "Impairment Level", "Score (0–4)", "Notes"],
                "rows": [
                    ["Occupational / academic functioning", "0=None, 4=Cannot work/attend", "", ""],
                    ["Social functioning", "0=None, 4=Complete isolation", "", ""],
                    ["Home management", "0=None, 4=Cannot manage home", "", ""],
                    ["Time lost to OCD per day (hrs)", "0=<1, 4=>8 hours", "", ""],
                    ["Quality of life (subjective)", "0=Good, 4=Terrible", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Cognitive — frontal/executive)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["Trail Making A", "___ sec", "", ""],
            ["Trail Making B", "___ sec", "", "Executive function / flexibility"],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
            ["WCST (Wisconsin Card Sorting — perseveration estimate)", "___", "", "Perseverative errors"],
        ],
        "section_d_rows": [
            ["Y-BOCS (Yale-Brown Obsessive Compulsive Scale)", "___ / 40", "≥16 moderate OCD", f"{CB} Subclinical  {CB} Mild  {CB} Moderate  {CB} Severe  {CB} Extreme"],
            ["OCI-R (OCD Inventory — Revised)", "___ / 72", "≥18 OCD likely", ""],
            ["PHQ-9 (Comorbid Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["CGI-S (Clinical Global Impression — Severity)", "___ / 7", "≥4 moderately ill", f"{CB} Normal  {CB} Borderline  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Comorbid Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Contamination  {CB} Symmetry/ordering  {CB} Harm/checking  {CB} Hoarding"],
            ["Additional Specifier", f"{CB} Pure obsessional  {CB} Treatment-resistant  {CB} Tic-related  {CB} Sensory phenomena"],
            ["Severity", f"{CB} Mild (Y-BOCS 8–15)  {CB} Moderate (16–23)  {CB} Severe (24–31)  {CB} Extreme (32–40)"],
            ["Duration", "___ years"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} Y-BOCS  {CB} OCI-R  {CB} CGI-S  {CB} PHQ-9"],
        ],
    },

    # 10 ─ MS ──────────────────────────────────────────────────────────────────
    {
        "slug": "ms",
        "display_name": "Multiple Sclerosis (MS)",
        "status_field_1": "EDSS Score",
        "status_field_2": "MS Type (RRMS/SPMS/PPMS)",
        "section_a_title": "MS Neurological & Functional Examination",
        "section_a_subsections": [
            {
                "title": "A1. Motor Function Assessment",
                "headers": ["Item", "Side", "Score (0–4)", "Notes"],
                "rows": [
                    ["Upper limb weakness (MRC)", "L / R", "", "0=Full power, 4=No movement"],
                    ["Lower limb weakness (MRC)", "L / R", "", ""],
                    ["Fatigue impact on motor function", "—", "", "0=None, 4=Cannot perform tasks"],
                    ["Fine motor dexterity (9HPT)", "L / R ___ sec", "", "Normative: <50 sec"],
                    ["Grip strength (dynamometer)", "L / R ___ kg", "", ""],
                ],
            },
            {
                "title": "A2. Spasticity (Modified Ashworth Scale)",
                "headers": ["Muscle Group", "Side", "MAS Score (0–4)", "Notes"],
                "rows": [
                    ["Hip adductors", "L / R", "", ""],
                    ["Knee extensors", "L / R", "", ""],
                    ["Ankle plantar flexors", "L / R", "", "Clonus: Y / N"],
                    ["Upper limb flexors", "L / R", "", ""],
                    ["Spasms (flexor / extensor)", "—", "", f"{CB} Absent  {CB} Mild  {CB} Moderate  {CB} Severe"],
                ],
            },
            {
                "title": "A3. Coordination & Ataxia",
                "headers": ["Test", "Side", "Score (0–4)", "Notes"],
                "rows": [
                    ["Finger-to-nose", "L / R", "", "0=Normal, 4=Unable"],
                    ["Heel-to-shin", "L / R", "", ""],
                    ["Rapid alternating movements", "L / R", "", "Dysdiadochokinesis"],
                    ["Tandem gait", "—", "", f"{CB} Normal  {CB} Ataxic  {CB} Cannot perform"],
                    ["SARA score (Gait subscale)", "___ / 8", "", "Scale for Assessment of Ataxia"],
                ],
            },
            {
                "title": "A4. Sensory Examination",
                "headers": ["Modality", "Side / Region", "Score (0–4)", "Notes"],
                "rows": [
                    ["Light touch (Semmes-Weinstein)", "L / R face/arm/leg", "", ""],
                    ["Vibration sense (128Hz tuning fork)", "L / R", "", "Ankle: present / absent"],
                    ["Joint position sense", "L / R", "", ""],
                    ["Pain / temperature discrimination", "L / R", "", ""],
                    ["Lhermitte's sign", "—", "", f"{CB} Absent  {CB} Present → cervical cord"],
                ],
            },
            {
                "title": "A5. Visual Acuity & Fields",
                "headers": ["Test", "Side", "Score (0–4)", "Notes"],
                "rows": [
                    ["Visual acuity (Snellen)", "L / R", "", "Corrected: ___"],
                    ["Colour vision (Ishihara plates)", "L / R", "", "___ / 15 plates"],
                    ["Visual fields (confrontation)", "L / R", "", f"{CB} Full  {CB} Hemianopia  {CB} Scotoma"],
                    ["Optic disc appearance", "L / R", "", f"{CB} Normal  {CB} Pallor — optic neuritis history"],
                    ["ONTT / RNFL (OCT if available)", "L / R", "", "___ μm"],
                ],
            },
            {
                "title": "A6. Bladder, Bowel & Fatigue",
                "headers": ["Domain", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Urinary urgency / frequency", "0=None, 4=Incontinence", "", ""],
                    ["Urinary hesitancy / retention", "0=None, 4=Catheter-dependent", "", ""],
                    ["Bowel dysfunction (constipation / urgency)", "0=None, 4=Incontinence", "", ""],
                    ["MS fatigue (MFIS)", "___ / 84", ">38 significant", ""],
                    ["Heat sensitivity (Uhthoff's phenomenon)", "0=None, 4=Disabling", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["SDMT (Symbol Digit Modalities Test)", "___ (raw score)", ">44 normal", f"{CB} Normal  {CB} Impaired — cognitive processing speed"],
            ["MSNQ (MS Neuropsychological Questionnaire)", "___ / 90", ">27 impaired", ""],
            ["Trail Making A", "___ sec", "", ""],
            ["Trail Making B", "___ sec", "", "Executive / set-shifting"],
            ["Verbal Fluency (COWAT — animals)", "___", "≥15 normal", ""],
            ["MoCA", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
        ],
        "section_d_rows": [
            ["PHQ-9 (Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
            ["FSS (Fatigue Severity Scale)", "___ / 63", "≥36 significant", f"{CB} Normal  {CB} Significant fatigue"],
            ["MFIS (Modified Fatigue Impact Scale)", "___ / 84", ">38 significant", "Cognitive / Physical / Psychosocial subscales"],
            ["MSQOL-54 (Quality of Life)", "___ / 100", "", "Mental / Physical health subscales"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} RRMS  {CB} SPMS  {CB} PPMS  {CB} Radiologically isolated syndrome (RIS)"],
            ["Dominant Syndrome", f"{CB} Fatigue-predominant  {CB} Cognitive  {CB} Motor/spastic  {CB} Cerebellar  {CB} Visual"],
            ["EDSS Score", f"___ (0–10)  {CB} Ambulatory (0–5.5)  {CB} Walking aid (6.0–6.5)  {CB} Wheelchair (7.0+)"],
            ["Disease Duration", "___ years since diagnosis"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} EDSS  {CB} T25FW  {CB} 9HPT  {CB} SDMT  {CB} FSS"],
        ],
    },

    # 11 ─ ASD ─────────────────────────────────────────────────────────────────
    {
        "slug": "asd",
        "display_name": "Autism Spectrum Disorder (ASD)",
        "status_field_1": "Age at ASD Diagnosis",
        "status_field_2": "Support Level (DSM-5 Level 1/2/3)",
        "section_a_title": "ASD-Specific Clinical Assessment",
        "section_a_subsections": [
            {
                "title": "A1. Social Communication Assessment",
                "headers": ["Item", "Observation / Rating", "Score (0–4)", "Notes"],
                "rows": [
                    ["Eye contact (frequency / quality)", "0=Normal, 4=Absent", "", ""],
                    ["Joint attention (pointing, sharing)", "0=Normal, 4=Absent", "", ""],
                    ["Reciprocal conversation", "0=Normal, 4=Cannot engage", "", ""],
                    ["Facial affect recognition", "0=Normal, 4=Cannot recognise", "", ""],
                    ["Use of gesture / non-verbal communication", "0=Normal, 4=Absent", "", ""],
                ],
            },
            {
                "title": "A2. Restricted & Repetitive Behaviours",
                "headers": ["Behaviour", "Frequency", "Score (0–4)", "Notes"],
                "rows": [
                    ["Repetitive motor movements (stimming)", "0=None, 4=Constant / disruptive", "", ""],
                    ["Insistence on sameness / routines", "0=None, 4=Extreme / disruptive", "", ""],
                    ["Highly restricted interests (preoccupation)", "0=None, 4=Consumes all time", "", ""],
                    ["Ritualistic / compulsive behaviour", "0=None, 4=Cannot deviate", "", ""],
                    ["Unusual sensory preoccupation", "0=None, 4=Severe", "", ""],
                ],
            },
            {
                "title": "A3. Sensory Processing",
                "headers": ["Modality", "Pattern", "Score (0–4)", "Notes"],
                "rows": [
                    ["Auditory sensitivity", f"{CB} Hyper  {CB} Hypo  {CB} Normal", "", "Triggers: ___"],
                    ["Tactile sensitivity", f"{CB} Hyper  {CB} Hypo  {CB} Normal", "", ""],
                    ["Visual sensitivity", f"{CB} Hyper  {CB} Hypo  {CB} Normal", "", "Fluorescent lights / patterns"],
                    ["Olfactory / gustatory sensitivity", f"{CB} Hyper  {CB} Hypo  {CB} Normal", "", "Food selectivity: ___"],
                    ["Vestibular / proprioceptive seeking/avoiding", f"{CB} Seeking  {CB} Avoiding  {CB} Normal", "", ""],
                ],
            },
            {
                "title": "A4. Language & Pragmatics",
                "headers": ["Domain", "Level", "Score (0–4)", "Notes"],
                "rows": [
                    ["Expressive language level", "0=Age-appropriate, 4=Non-verbal", "", ""],
                    ["Receptive language level", "0=Age-appropriate, 4=Minimal comprehension", "", ""],
                    ["Pragmatic language (conversation)", "0=Normal, 4=Severely impaired", "", ""],
                    ["Prosody / tone variation", "0=Normal, 4=Monotone / echolalia", "", ""],
                    ["Literal interpretation tendency", "0=Normal, 4=Exclusively literal", "", ""],
                ],
            },
            {
                "title": "A5. Executive Function",
                "headers": ["Function", "Assessment", "Score (0–4)", "Notes"],
                "rows": [
                    ["Cognitive flexibility", "TMT-B / set-shifting tasks", "", ""],
                    ["Planning & problem solving", "Tower task (clinical obs)", "", ""],
                    ["Working memory", "Digit span backward", "", ""],
                    ["Inhibitory control", "Go/No-Go (clinical obs)", "", ""],
                    ["Initiation / task completion", "0=Normal, 4=Cannot initiate", "", ""],
                ],
            },
            {
                "title": "A6. Adaptive Behaviour",
                "headers": ["Domain", "Level", "Score (0–4)", "Notes"],
                "rows": [
                    ["Communication skills (daily)", "0=Age-appropriate, 4=Severely limited", "", ""],
                    ["Daily living skills (self-care)", "0=Independent, 4=Fully dependent", "", ""],
                    ["Socialisation (peer relationships)", "0=Normal, 4=None", "", ""],
                    ["Academic / vocational participation", "0=Full, 4=Cannot participate", "", ""],
                    ["Adaptive behaviour composite (ABAS/Vineland)", "___", "", "Standard score: ___"],
                ],
            },
        ],
        "section_c_rows": [
            ["ADOS-2 (Autism Diagnostic Obs Schedule) — if available", "Module: ___  Score: ___", "", "Gold standard observational"],
            ["MoCA (Cognitive)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["Trail Making A", "___ sec", "", ""],
            ["Trail Making B", "___ sec", "", "Cognitive flexibility"],
            ["Vineland Adaptive Behaviour Scales", "Composite: ___", "", ""],
            ["BRIEF-2 / BRIEF-A (Executive)", "Global T-score: ___", "T>65 elevated", ""],
        ],
        "section_d_rows": [
            ["SRS-2 (Social Responsiveness Scale)", "T-score: ___", "T>65 elevated", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe ASD symptoms"],
            ["PHQ-9 (Comorbid Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["SCARED / GAD-7 (Comorbid Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
            ["ASRS (Comorbid ADHD)", "___ / 72", "≥14 Part A", f"{CB} Normal  {CB} Likely ADHD comorbidity"],
            ["RBS-R (Repetitive Behaviour Scale)", "___ / 129", "", "6 subscales"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Social communication dominant  {CB} Restricted/repetitive dominant  {CB} Language delay"],
            ["Additional Specifier", f"{CB} Sensory processing  {CB} Executive dysfunction  {CB} ADHD comorbidity  {CB} Pathological demand avoidance"],
            ["DSM-5 Support Level", f"{CB} Level 1 (requiring support)  {CB} Level 2  {CB} Level 3 (requiring substantial support)"],
            ["Intellectual Ability", f"{CB} Average / above  {CB} Borderline  {CB} Intellectual disability"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} SRS-2  {CB} ADOS-2  {CB} RBS-R  {CB} BRIEF  {CB} Vineland"],
        ],
    },

    # 12 ─ LONG COVID ──────────────────────────────────────────────────────────
    {
        "slug": "long_covid",
        "display_name": "Long COVID (Post-Acute Sequelae of SARS-CoV-2)",
        "status_field_1": "Acute COVID Date / Months Post-Infection",
        "status_field_2": "Acute COVID Severity",
        "section_a_title": "Long COVID Multi-System Examination",
        "section_a_subsections": [
            {
                "title": "A1. Cognitive Assessment — Brain Fog",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Subjective cognitive complaint", "0=None, 4=Cannot think", "", ""],
                    ["Digit Span (forward)", "___ / 9", "", "Sustained attention"],
                    ["Digit Span (backward)", "___ / 8", "", "Working memory"],
                    ["Serial 7s (concentration)", "___ errors", "", ""],
                    ["Word-finding difficulty", "0=None, 4=Cannot name common objects", "", ""],
                ],
            },
            {
                "title": "A2. Fatigue Severity",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Fatigue level at rest (NRS)", "___ / 10", "", ""],
                    ["Post-exertional malaise (PEM)", "0=None, 4=Crash lasting days", "", ""],
                    ["PEM trigger threshold (activity level)", "0=Heavy exercise, 4=Minimal activity", "", ""],
                    ["Recovery time after activity", "0=<1 hour, 4=>72 hours", "", ""],
                    ["Impact on daily activities", "0=None, 4=Bed-bound", "", ""],
                ],
            },
            {
                "title": "A3. Autonomic Function (Dysautonomia)",
                "headers": ["Test / Symptom", "Finding", "Score (0–4)", "Notes"],
                "rows": [
                    ["Orthostatic symptoms (POTS screen)", f"{CB} None  {CB} Present", "", "HR supine: ___ bpm"],
                    ["HR change on standing (10 min)", "___ bpm change", "≥30 bpm = POTS criteria", ""],
                    ["BP lying vs standing", "___ / ___ → ___ / ___", "", "Orthostatic hypotension?"],
                    ["Palpitations frequency", "0=Never, 4=Constant", "", ""],
                    ["Temperature dysregulation", "0=None, 4=Severe", "", ""],
                ],
            },
            {
                "title": "A4. Exercise Tolerance",
                "headers": ["Test", "Result", "Score (0–4)", "Notes"],
                "rows": [
                    ["6-Minute Walk Test (6MWT)", "___ metres", ">400m normal", ""],
                    ["1-minute sit-to-stand", "___ reps", "", ""],
                    ["Borg RPE at peak exertion", "___ / 20", "", "Perceived exertion"],
                    ["Oxygen saturation at rest / exertion", "___% / ___%", ">95% normal", ""],
                    ["Heart rate recovery (1 min post)", "___ bpm drop", ">12 bpm normal", ""],
                ],
            },
            {
                "title": "A5. Headache & Pain",
                "headers": ["Symptom", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Headache frequency (days/week)", "0=None, 4=Daily", "", ""],
                    ["Headache severity (NRS)", "___ / 10", "", ""],
                    ["Chest pain / pressure", "0=None, 4=Severe", "", f"{CB} Cardiac excluded"],
                    ["Myalgia / arthralgia", "0=None, 4=Disabling", "", "Location: ___"],
                    ["Neuropathic symptoms (burning/pins & needles)", "0=None, 4=Severe", "", ""],
                ],
            },
            {
                "title": "A6. Mood & Sleep Screen",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Depressive symptoms (subjective)", "0=None, 4=Severe depression", "", ""],
                    ["Anxiety symptoms", "0=None, 4=Incapacitating", "", ""],
                    ["Sleep quality (subjective)", "0=Good, 4=Terrible", "", ""],
                    ["Sleep duration (hours/night)", "___ hours", "7–9 hours normal", ""],
                    ["Dysautonomia effect on sleep", "0=None, 4=Cannot sleep", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Brain fog / cognitive assessment)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["CFQ (Cognitive Failures Questionnaire)", "___ / 100", ">43 significant", "Self-reported cognitive failures"],
            ["Trail Making A", "___ sec", "", "Processing speed"],
            ["Trail Making B", "___ sec", "", "Executive / flexibility"],
            ["Digit Span (forward + backward)", "___ / 18", "", ""],
            ["PCFS (Post-COVID Functional Status Scale)", "___ (0–4)", "0=No limitation, 4=Severe", f"{CB} Grade 0  {CB} 1  {CB} 2  {CB} 3  {CB} 4"],
        ],
        "section_d_rows": [
            ["MFI-20 (Multidimensional Fatigue Inventory)", "___ / 100", ">60 severe fatigue", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["PHQ-9 (Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
            ["PCFS (Post-COVID Functional Status)", "___ / 4", "≥2 significant impact", f"{CB} Grade 0  {CB} 1  {CB} 2  {CB} 3  {CB} 4"],
            ["PSQI (Sleep Quality)", "___ / 21", ">5 poor sleep", f"{CB} Good  {CB} Poor sleeper"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Cognitive fog  {CB} Fatigue-predominant  {CB} Dysautonomia  {CB} Multi-system  {CB} Neuropsychiatric"],
            ["Additional Specifier", f"{CB} Headache/pain  {CB} Cardiorespiratory  {CB} Anosmia/ageusia persistent  {CB} Post-POTS"],
            ["Severity", f"{CB} Mild (PCFS 1)  {CB} Moderate (PCFS 2)  {CB} Severe (PCFS 3)  {CB} Very severe (PCFS 4)"],
            ["Duration Post-COVID", "___ months since acute infection"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} MoCA  {CB} MFI-20  {CB} PCFS  {CB} CFQ  {CB} PHQ-9"],
        ],
    },

    # 13 ─ TINNITUS ────────────────────────────────────────────────────────────
    {
        "slug": "tinnitus",
        "display_name": "Tinnitus",
        "status_field_1": "Tinnitus Duration (months/years)",
        "status_field_2": "Onset Type",
        "section_a_title": "Tinnitus Characterisation & Auditory Examination",
        "section_a_subsections": [
            {
                "title": "A1. Tinnitus Characterisation",
                "headers": ["Item", "Finding", "Score (0–4)", "Notes"],
                "rows": [
                    ["Perceived laterality", f"{CB} Left  {CB} Right  {CB} Bilateral  {CB} Central (in head)", "", ""],
                    ["Tinnitus loudness match (audiometer)", "___ dBSL or ___ dBHL", "", ""],
                    ["Tinnitus pitch match (audiometer)", "___ Hz", "", ""],
                    ["Masking level (minimum masking level)", "___ dBSL", "", ""],
                    ["Residual inhibition duration", "___ seconds after masking", "", f"{CB} Complete  {CB} Partial  {CB} None"],
                ],
            },
            {
                "title": "A2. Hearing Assessment",
                "headers": ["Domain", "Side", "Score (0–4)", "Notes"],
                "rows": [
                    ["Audiogram — pure tone average (PTA)", "L: ___ dB  R: ___ dB", "", "0.5–4 kHz average"],
                    ["High-frequency hearing loss (4–8 kHz)", "L / R", "", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
                    ["Speech discrimination score", "L: ___%  R: ___%", "", ""],
                    ["Otoscopy", "L / R", "", f"{CB} Normal  {CB} Wax  {CB} Perforation  {CB} Other"],
                    ["Tympanometry", "L / R", "", "Type A/B/C: ___"],
                ],
            },
            {
                "title": "A3. Hyperacusis Screen",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Loudness discomfort levels (LDL)", "L: ___ dBHL  R: ___ dBHL", "", "<90 dBHL = hyperacusis"],
                    ["Hyperacusis Questionnaire (HQ)", "___ / 42", "≥28 significant", f"{CB} Absent  {CB} Mild  {CB} Moderate  {CB} Severe"],
                    ["Sound tolerance in daily environments", "0=Normal, 4=Cannot leave home", "", ""],
                    ["Misophonia features", "0=None, 4=Severe", "", f"{CB} Specific triggers: ___"],
                ],
            },
            {
                "title": "A4. Somatic Modulation",
                "headers": ["Manoeuvre", "Effect", "Score (0–4)", "Notes"],
                "rows": [
                    ["Jaw movements (open/lateral)", f"{CB} No change  {CB} Increases  {CB} Decreases", "", "Temporomandibular joint exam"],
                    ["Neck movements (flexion/rotation)", f"{CB} No change  {CB} Increases  {CB} Decreases", "", "Cervicogenic tinnitus?"],
                    ["Scalp / periauricular pressure", f"{CB} No change  {CB} Modulates", "", ""],
                    ["Eye movements (lateral gaze)", f"{CB} No change  {CB} Modulates", "", ""],
                ],
            },
            {
                "title": "A5. Sleep Impact",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Sleep onset difficulty due to tinnitus", "0=None, 4=>2 hour latency", "", ""],
                    ["Nocturnal awakening from tinnitus", "0=None, 4=Multiple/night", "", ""],
                    ["Tinnitus louder in quiet / at night", "0=No, 4=Severely worse", "", ""],
                    ["PSQI score", "___ / 21", ">5 poor sleep", ""],
                ],
            },
            {
                "title": "A6. Concentration & Quality of Life Impact",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Concentration impairment from tinnitus", "0=None, 4=Cannot concentrate", "", ""],
                    ["Reading / work performance", "0=Normal, 4=Cannot work", "", ""],
                    ["Social impact / communication", "0=None, 4=Avoids all social", "", ""],
                    ["Emotional distress from tinnitus", "0=None, 4=Suicidal ideation", "", f"{CB} Refer urgently if SI present"],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Cognitive — attention / concentration)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["Trail Making A", "___ sec", "", "Attention / processing speed"],
            ["Digit Span (forward + backward)", "___ / 18", "", ""],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
        ],
        "section_d_rows": [
            ["THI (Tinnitus Handicap Inventory)", "___ / 100", "≥58 severe", f"{CB} Slight  {CB} Mild  {CB} Moderate  {CB} Severe  {CB} Catastrophic"],
            ["TFI (Tinnitus Functional Index)", "___ / 100", "≥25 significant", "8 subscale domains"],
            ["PHQ-9 (Comorbid Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Comorbid Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
            ["VAS Tinnitus Loudness (subjective)", "___ / 10", "", "Separate from audiometric matching"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Tonal  {CB} Noise-like (broadband)  {CB} Pulsatile  {CB} Reactive tinnitus"],
            ["Additional Specifier", f"{CB} Somatic/somatosensory  {CB} With hyperacusis  {CB} With misophonia  {CB} Audiological origin"],
            ["Severity", f"{CB} Slight (THI 0–16)  {CB} Mild (18–36)  {CB} Moderate (38–56)  {CB} Severe (58–76)  {CB} Catastrophic (78–100)"],
            ["Hearing Loss Comorbidity", f"{CB} No hearing loss  {CB} Mild  {CB} Moderate  {CB} Severe / profound"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Not available"],
            ["EEG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} THI  {CB} TFI  {CB} VAS loudness  {CB} PHQ-9  {CB} PSQI"],
        ],
    },

    # 14 ─ INSOMNIA ────────────────────────────────────────────────────────────
    {
        "slug": "insomnia",
        "display_name": "Insomnia Disorder",
        "status_field_1": "Insomnia Duration (months/years)",
        "status_field_2": "Insomnia Type",
        "section_a_title": "Sleep Architecture & Insomnia Examination",
        "section_a_subsections": [
            {
                "title": "A1. Sleep Onset Latency",
                "headers": ["Item", "Measurement", "Score (0–4)", "Notes"],
                "rows": [
                    ["Subjective sleep onset latency (SOL)", "___ minutes", ">30 min = abnormal", ""],
                    ["Sleep diary average SOL (last 2 weeks)", "___ minutes", "", ""],
                    ["Pre-sleep arousal (cognitive)", "0=None, 4=Racing thoughts prevent sleep", "", ""],
                    ["Pre-sleep arousal (somatic)", "0=Relaxed, 4=Tense/restless/cannot lie still", "", ""],
                    ["Frequency of delayed onset (nights/week)", "0=Never, 4=Every night", "", ""],
                ],
            },
            {
                "title": "A2. Sleep Maintenance",
                "headers": ["Item", "Measurement", "Score (0–4)", "Notes"],
                "rows": [
                    ["Number of nocturnal awakenings / night", "___ awakenings", "0=None, 4=>5 awakenings", ""],
                    ["Total time awake after sleep onset (WASO)", "___ minutes", ">30 min = abnormal", ""],
                    ["Ease of returning to sleep", "0=Immediately, 4=Cannot return", "", ""],
                    ["Duration of awake periods", "___ minutes average", "", ""],
                    ["Frequency (nights/week)", "0=Never, 4=Every night", "", ""],
                ],
            },
            {
                "title": "A3. Early Morning Awakening",
                "headers": ["Item", "Measurement", "Score (0–4)", "Notes"],
                "rows": [
                    ["Habitual wake time", "___ AM", "", ""],
                    ["Desired wake time", "___ AM", "", ""],
                    ["Early awakening (>30 min before desired)", "0=Never, 4=Every day", "", ""],
                    ["Inability to return to sleep after early waking", "0=Always returns, 4=Never returns", "", ""],
                    ["Total sleep time (TST)", "___ hours / night", "7–9 hours normal", ""],
                ],
            },
            {
                "title": "A4. Daytime Sleepiness & Impairment",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Epworth Sleepiness Scale (ESS)", "___ / 24", "≥11 excessive sleepiness", ""],
                    ["Daytime fatigue (not sleepiness)", "0=None, 4=Incapacitating", "", ""],
                    ["Concentration / memory impairment", "0=None, 4=Cannot function", "", ""],
                    ["Mood impact (irritability)", "0=None, 4=Severe", "", ""],
                    ["Occupational / social impairment", "0=None, 4=Cannot work", "", ""],
                ],
            },
            {
                "title": "A5. Sleep Hygiene Assessment",
                "headers": ["Factor", "Status", "Score (0–4)", "Notes"],
                "rows": [
                    ["Irregular sleep schedule", "0=Consistent, 4=Highly irregular", "", "Bedtime range: ___"],
                    ["Screen use before bed", "0=None, 4=>2 hours in bed", "", ""],
                    ["Caffeine / alcohol use (evening)", "0=None, 4=Daily heavy use", "", ""],
                    ["Napping frequency", "0=Never, 4=Daily long naps", "", "___ naps/day, ___ min"],
                    ["Bedroom environment (light, noise, temp)", "0=Optimal, 4=Highly disruptive", "", ""],
                ],
            },
            {
                "title": "A6. Hyperarousal & Conditioned Arousal",
                "headers": ["Item", "Severity", "Score (0–4)", "Notes"],
                "rows": [
                    ["Conditioned arousal to bed/bedroom", "0=None, 4=Dread entering bedroom", "", ""],
                    ["Sleep-related worry / sleep anxiety", "0=None, 4=Panic at bedtime", "", ""],
                    ["24-hour physiological arousal", "0=Normal, 4=Constantly keyed-up", "", ""],
                    ["Dysfunctional beliefs about sleep", "0=None, 4=Catastrophising", "", f"{CB} DBAS-16 indicated"],
                    ["Nocturnal cognitive hyperarousal", "0=None, 4=Cannot switch off thoughts", "", ""],
                ],
            },
        ],
        "section_c_rows": [
            ["MoCA (Cognitive — sleep deprivation effects)", "___ / 30", "≥26 normal", f"{CB} Normal  {CB} Impaired"],
            ["Trail Making A", "___ sec", "", "Processing speed / attention"],
            ["Trail Making B", "___ sec", "", "Executive function"],
            ["Digit Span (forward + backward)", "___ / 18", "", "Working memory"],
            ["Verbal Fluency (animals / 1 min)", "___", "≥15 normal", ""],
        ],
        "section_d_rows": [
            ["ISI (Insomnia Severity Index)", "___ / 28", "≥15 moderate insomnia", f"{CB} No / subthreshold  {CB} Moderate  {CB} Severe insomnia"],
            ["PSQI (Pittsburgh Sleep Quality Index)", "___ / 21", ">5 poor sleep", "7 component scores"],
            ["ESS (Epworth Sleepiness Scale)", "___ / 24", "≥11 excessive", f"{CB} Normal  {CB} Borderline  {CB} Excessive daytime sleepiness"],
            ["PHQ-9 (Comorbid Depression)", "___ / 27", "≥10 moderate", f"{CB} Normal  {CB} Mild  {CB} Moderate  {CB} Severe"],
            ["GAD-7 (Comorbid Anxiety)", "___ / 21", "≥10 moderate", f"{CB} Normal  {CB} Present"],
        ],
        "section_e_rows": [
            ["Identified Phenotype", f"{CB} Sleep-onset  {CB} Sleep-maintenance  {CB} Early-morning awakening  {CB} Combined"],
            ["Additional Specifier", f"{CB} Hyperarousal-predominant  {CB} Comorbid psychiatric  {CB} Paradoxical insomnia  {CB} Conditioned arousal"],
            ["Severity", f"{CB} Subthreshold (ISI <8)  {CB} Mild (8–14)  {CB} Moderate (15–21)  {CB} Severe (22–28)"],
            ["Duration", f"{CB} Acute (<3 months)  {CB} Chronic (≥3 months)"],
            ["EEG Performed", f"{CB} Yes  {CB} No  {CB} Polysomnography instead"],
            ["EEG / PSG Findings Summary", ""],
            ["Preliminary NIBS Strategy", ""],
            ["Recommended Validated Scales", f"{CB} ISI  {CB} PSQI  {CB} ESS  {CB} PHQ-9  {CB} GAD-7"],
        ],
    },
]


def main():
    base_output = os.path.join(
        os.path.dirname(__file__), "..", "outputs", "documents"
    )

    for cond in CONDITIONS:
        slug = cond["slug"]
        out_dir = os.path.join(base_output, slug, "fellow")
        os.makedirs(out_dir, exist_ok=True)

        out_path = os.path.join(out_dir, "Clinical_Examination_Checklist_Fellow.docx")
        doc = build_document(cond)
        doc.save(out_path)
        print(f"[OK] {slug}: {out_path}")


if __name__ == "__main__":
    main()
