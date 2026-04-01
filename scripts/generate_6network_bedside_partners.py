#!/usr/bin/env python3
"""
Generate 6-Network Bedside Assessment (Partners Tier) for all 14 conditions.
Run: python scripts/generate_6network_bedside_partners.py
"""
import os, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ────────────────────────────────────────────────────────────────────
NAVY   = "1B3A5C";  NAVY_RGB   = RGBColor(0x1B,0x3A,0x5C)
BLUE   = "2E75B6";  BLUE_RGB   = RGBColor(0x2E,0x75,0xB6)
PURPLE = "7B2D8E";  PURPLE_RGB = RGBColor(0x7B,0x2D,0x8E)
RED    = "CC0000";  RED_RGB    = RGBColor(0xCC,0x00,0x00)
WHITE_RGB = RGBColor(0xFF,0xFF,0xFF)
LGRAY  = "F2F2F2";  WHITE = "FFFFFF"
FONT   = "Calibri"

OUTPUT_BASE = Path(__file__).parent.parent / "outputs" / "documents"

# ── XML helpers ───────────────────────────────────────────────────────────────
def _shade(cell, hex6):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), hex6.lstrip("#")); pr.append(shd)

def _borders(table):
    tbl=table._tbl; tp=tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tb=OxmlElement("w:tblBorders")
    for e in ("top","left","bottom","right","insideH","insideV"):
        b=OxmlElement(f"w:{e}"); b.set(qn("w:val"),"single")
        b.set(qn("w:sz"),"4"); b.set(qn("w:space"),"0")
        b.set(qn("w:color"),"1B3A5C"); tb.append(b)
    tp.append(tb)

def _para_shade(para, hex6):
    pPr=para._p.get_or_add_pPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex6.lstrip("#"))
    pPr.append(shd)

def _run(para, text, bold=False, color=None, size=11, italic=False):
    r=para.add_run(text); r.font.name=FONT; r.font.size=Pt(size)
    r.font.bold=bold; r.font.italic=italic
    if color: r.font.color.rgb=color
    return r

def _heading(doc, text, level=1):
    colors={1:NAVY_RGB,2:BLUE_RGB,3:PURPLE_RGB}
    sizes={1:14,2:12,3:11}
    p=doc.add_paragraph()
    _run(p,text,bold=True,color=colors[level],size=sizes[level])
    p.paragraph_format.space_before=Pt(10 if level==1 else 6)
    p.paragraph_format.space_after=Pt(4)
    return p

def _purple_callout(doc, text):
    p=doc.add_paragraph()
    _para_shade(p, PURPLE)
    _run(p,"🟣 FNON NOTE: ",bold=True,color=WHITE_RGB,size=10)
    _run(p,text,color=WHITE_RGB,size=10)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(6)

def _red_callout(doc, text):
    p=doc.add_paragraph()
    _para_shade(p, RED)
    _run(p,"⚠ CRITICAL: ",bold=True,color=WHITE_RGB,size=10)
    _run(p,text,color=WHITE_RGB,size=10)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(6)

def _add_6col_table(doc, tests):
    """tests = list of (name, how_to, normal, hypo, hyper) — Score col blank."""
    headers=["Bedside Test","How to Perform","Normal","HYPO (1)","HYPER (2)","Score"]
    widths=[1.1, 1.9, 1.1, 1.3, 1.3, 0.6]
    t=doc.add_table(rows=len(tests)+1, cols=6)
    t.alignment=WD_TABLE_ALIGNMENT.LEFT
    t.style="Table Grid"
    _borders(t)
    # header row
    for j,(h,w) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[j]; _shade(c,NAVY)
        p=c.paragraphs[0]; p.clear()
        _run(p,h,bold=True,color=WHITE_RGB,size=9)
        c.width=Inches(w)
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    # data rows
    for i,(name,how_to,normal,hypo,hyper) in enumerate(tests):
        row=t.rows[i+1]; bg=LGRAY if i%2==0 else WHITE
        for j,(txt,w) in enumerate(zip([name,how_to,normal,hypo,hyper,""],widths)):
            c=row.cells[j]; _shade(c,bg)
            p=c.paragraphs[0]; p.clear()
            _run(p,txt,size=9)
            c.width=Inches(w)
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    doc.add_paragraph()

def _network_block(doc, net_name, max_score, tests, relevance_text, short):
    _heading(doc, f"{net_name} — Max Score: {max_score}", level=2)
    _add_6col_table(doc, tests)
    _purple_callout(doc, f"{short} Relevance: {relevance_text}")
    # total line
    p=doc.add_paragraph()
    _run(p,f"{net_name} Total: _____ / {max_score}",bold=True,color=NAVY_RGB,size=11)
    p.paragraph_format.space_after=Pt(8)

def _patient_info_table(doc):
    t=doc.add_table(rows=11, cols=2); t.style="Table Grid"; _borders(t)
    fields=[
        ("Patient Name:",""),("Date of Birth:",""),("Date of Assessment:",""),
        ("Clinician:",""),("Session Number:",""),("Medication State:","ON  /  OFF  /  Wash-out"),
        ("Primary Diagnosis:",""),("Secondary Diagnosis:",""),("Current Medications:",""),
        ("Referral Source:",""),("Notes:",""),
    ]
    for i,(label,val) in enumerate(fields):
        r=t.rows[i]
        _shade(r.cells[0], LGRAY)
        p0=r.cells[0].paragraphs[0]; p0.clear()
        _run(p0,label,bold=True,color=NAVY_RGB,size=10)
        p1=r.cells[1].paragraphs[0]; p1.clear()
        _run(p1,val,size=10)
    doc.add_paragraph()

def _summary_table(doc, short):
    headers=["Network","Max Score","Raw Score","% of Max","Dysfunction Pattern","Priority"]
    rows=[
        ["Default Mode Network (DMN)","14","","","",""],
        ["Central Executive Network (CEN/FPN)","14","","","",""],
        ["Salience Network (SN)","14","","","",""],
        ["Sensorimotor Network (SMN)","16","","","",""],
        ["Limbic / Emotional Network","14","","","",""],
        ["Attention Networks (DAN/VAN)","16","","","",""],
        ["GRAND TOTAL","88","","","",""],
    ]
    t=doc.add_table(rows=len(rows)+1, cols=6)
    t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style="Table Grid"; _borders(t)
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; _shade(c,NAVY)
        p=c.paragraphs[0]; p.clear()
        _run(p,h,bold=True,color=WHITE_RGB,size=9)
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    for i,row in enumerate(rows):
        bg="E8D5F0" if i==6 else (LGRAY if i%2==0 else WHITE)
        bold=(i==6)
        for j,txt in enumerate(row):
            c=t.rows[i+1].cells[j]; _shade(c,bg)
            p=c.paragraphs[0]; p.clear()
            _run(p,txt,bold=bold,size=9)
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    doc.add_paragraph()

# ── Main document builder ─────────────────────────────────────────────────────
def build_document(cond):
    doc=Document()
    # Page setup: Letter, 1-inch margins
    sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(sec,attr,Inches(1))

    short=cond["short"]; full=cond["full"]

    # ── HEADER ──
    p=doc.add_paragraph()
    _run(p,"SOZO BRAIN CENTER — PARTNERS TIER",bold=True,color=NAVY_RGB,size=16)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    p=doc.add_paragraph()
    _run(p,f"{short}-Specific 6-Network Bedside Assessment",bold=True,color=BLUE_RGB,size=14)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    p=doc.add_paragraph()
    _para_shade(p, PURPLE)
    _run(p,"🟣 FNON + Evidence-Based Assessment — Includes Brain Network Analysis",
         bold=True,color=WHITE_RGB,size=11)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _patient_info_table(doc)

    # Scoring key
    p=doc.add_paragraph()
    _run(p,"Scoring Key: ",bold=True,color=NAVY_RGB,size=11)
    _run(p,"HYPO (1)  |  ",color=BLUE_RGB,size=11)
    _run(p,"NORMAL (0)  |  ",color=NAVY_RGB,size=11)
    _run(p,"HYPER (2)",color=RED_RGB,size=11)
    _run(p,"          Total: 44 Tests (Max Score: 88)",bold=True,color=NAVY_RGB,size=11)

    p=doc.add_paragraph()
    _para_shade(p, PURPLE)
    _run(p,f"🟣 {short}-Adapted Version: ",bold=True,color=WHITE_RGB,size=10)
    _run(p,cond["adapted_note"],color=WHITE_RGB,size=10)

    doc.add_paragraph()

    # ── 6 NETWORK BLOCKS ──
    networks=[
        ("Default Mode Network (DMN)",        14, cond["dmn_tests"],    cond["dmn_rel"]),
        ("Central Executive Network (CEN/FPN)",14, cond["cen_tests"],    cond["cen_rel"]),
        ("Salience Network (SN)",              14, cond["sn_tests"],     cond["sn_rel"]),
        ("Sensorimotor Network (SMN)",         16, cond["smn_tests"],    cond["smn_rel"]),
        ("Limbic / Emotional Network",         14, cond["limbic_tests"], cond["limbic_rel"]),
        ("Attention Networks (DAN/VAN)",       16, cond["attn_tests"],   cond["attn_rel"]),
    ]
    for net_name, max_score, tests, relevance in networks:
        _network_block(doc, net_name, max_score, tests, relevance, short)

    # ── NETWORK SUMMARY PROFILE ──
    _heading(doc,"Network Summary Profile",level=1)
    _purple_callout(doc,"Map network scores to phenotype for FNON-based montage and protocol selection")
    _summary_table(doc, short)

    p=doc.add_paragraph()
    _run(p,"Dominant Network Dysfunction: ",bold=True,color=NAVY_RGB)
    _run(p,"_______________________________   Score: _____ / _____")
    doc.add_paragraph()
    p=doc.add_paragraph()
    _run(p,"Secondary Network: ",bold=True,color=NAVY_RGB)
    _run(p,"_______________________________   Tertiary: _______________________________")

    # ── PHENOTYPE ASSIGNMENT ──
    doc.add_paragraph()
    _heading(doc,"Phenotype Assignment",level=1)
    p=doc.add_paragraph()
    _run(p,cond["phenotype_text"],size=11)
    doc.add_paragraph()
    p=doc.add_paragraph()
    _run(p,"Assigned Phenotype: ",bold=True,color=NAVY_RGB)
    _run(p,"________________________________________")
    doc.add_paragraph()
    p=doc.add_paragraph()
    _run(p,"FNON Strategy: ",bold=True,color=PURPLE_RGB)
    _run(p,cond["fnon_strategy"],size=11)

    # ── MEDICATION & CLINICAL NOTES ──
    doc.add_paragraph()
    _heading(doc,"Medication & Clinical Notes",level=1)
    _red_callout(doc,"CRITICAL: Record medication state at each assessment. Results may differ significantly ON vs OFF medication.")
    p=doc.add_paragraph()
    _run(p,"Medication State at Assessment: ",bold=True,color=NAVY_RGB)
    _run(p,"□ ON Medication    □ OFF Medication    □ Medication Wash-out    □ Unmedicated")
    doc.add_paragraph()
    for label in ["Medication Name/Dose/Timing:","Time Since Last Dose:","Clinical Observations:","Network Interpretation Notes:","Next Assessment Date:"]:
        p=doc.add_paragraph()
        _run(p,f"{label} ",bold=True,color=NAVY_RGB,size=10)
        _run(p,"_"*50,size=10)

    # Save
    slug=cond["slug"]
    out_dir=OUTPUT_BASE / slug / "partners"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path=out_dir / f"6Network_Bedside_Assessment_{slug}.docx"
    doc.save(str(out_path))
    print(f"  + Saved: {out_path}")
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# CONDITION DATA
# ═══════════════════════════════════════════════════════════════════════════════

# ── 1. DEPRESSION ─────────────────────────────────────────────────────────────
DEPRESSION = {
  "full":"Major Depressive Disorder (MDD)","short":"Depression","slug":"depression",
  "adapted_note":"Includes rumination and anhedonia markers, psychomotor retardation tests, and Depression phenotype markers",
  "dmn_rel":"Depression hallmark is DMN HYPERACTIVITY — pathological self-referential rumination, negative autobiographical memory bias, and failure to deactivate during goal-directed tasks. mPFC-PCC hyperconnectivity drives persistent negative self-focus.",
  "cen_rel":"CEN HYPOACTIVITY in Depression impairs executive control, working memory, and cognitive flexibility. Patients struggle to override negative automatic thoughts due to DLPFC hypofrontality and reduced top-down cognitive regulation.",
  "sn_rel":"SN shows HYPER-sensitivity to negative emotional content and HYPO-sensitivity to positive/rewarding stimuli. ACC dysfunction reduces error monitoring and motivation. Insula hyperactivation drives negative interoceptive awareness.",
  "smn_rel":"Psychomotor retardation reflects SMN HYPOACTIVITY — slowed motor initiation, reduced movement speed, and decreased motor variability. Basal ganglia hypofrontality contributes to bradykinesia-like motor slowing in severe MDD.",
  "limbic_rel":"Amygdala HYPERACTIVITY and hippocampal volume reduction characterize limbic dysfunction in Depression. Negative emotional memory bias, impaired fear extinction, and anhedonia (blunted reward circuit) drive core affective symptoms.",
  "attn_rel":"Depression causes HYPO-activity of dorsal attention network with reduced top-down attentional control and increased susceptibility to mind-wandering. Concentration difficulties reflect CEN-DAN dyscoordination with DMN intrusion.",
  "dmn_tests":[
    ("Rumination Probe","Ask: 'When feeling sad, do thoughts repeat involuntarily for >5 min?' Rate frequency 0–10 and duration","Minimal intrusive thoughts; able to redirect within 2 min","Blunted self-reflection; empty thought content; emotional numbness","Persistent negative intrusive thoughts >5 min; unable to redirect; ruminative loops"),
    ("Autobiographical Recall","Ask for 3 specific personal memories (childhood, recent, positive event). Rate detail, accuracy, and emotional tone","Vivid, specific memories with mixed emotional content; appropriate affect","Sparse memory detail; no associated emotion; apathetic recall","Predominantly negative memory retrieval; catastrophic reinterpretation of neutral events"),
    ("Self-Reference Memory","Read 10 adjectives with 'Does this describe you?' frame. Free recall after 2 min distractor. Count self-referenced words.","≥6/10 recalled; balanced positive and negative self-descriptors","<4/10 recalled; emotional flatness; difficulty with self-referential processing","All negative self-descriptors endorsed and recalled; positive words rejected or not recalled"),
    ("Mind Wandering Probe","Give 2-min reading passage; interrupt 4 times randomly: 'Were you thinking about the text?' Note task-unrelated thought rate","≤1/4 off-task responses","0/4 off-task — blunted ideation; emotionally flat or dissociated","3–4/4 off-task; reports persistent sad/worry thoughts during reading"),
    ("Prospective Memory","Tell patient: 'Remind me to ask about your sleep in 5 minutes.' Continue assessment. Note spontaneous recall.","Independently recalls within 5±1 min","Forgets entirely; no self-monitoring; poor future orientation","Perseveratively asks before 5 min; intrusive future-oriented worry"),
    ("Default Narrative","Ask: 'Tell me something about yourself from the past week.' Rate: coherence, self-reference, negative bias, time-orientation","Balanced narrative; positive/neutral events described; normal temporal orientation","Minimal narrative content; flat affect; difficulty generating self-relevant material","Exclusively negative content; rumination on failures/losses; catastrophic framing"),
    ("Semantic Fluency — Self","Name personal possessions or meaningful places in 1 min. Count items; note emotional valence of examples","≥8 items; mixed emotional valence; no perseveration","<4 items; flat responses; poor self-schema access","Predominantly negative/loss-associated items; perseveration on depressive themes"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Read digit sequences; ask reverse recall. Start 2-digits, increase by 1 each trial. Score max length achieved.","Backward span ≥5 digits","Backward span ≤3; poor working memory; rapid forgetting","Span intact but slow; perseverative responses; minor executive intrusions"),
    ("Trail Making B","Draw alternating number-letter connections (1-A-2-B-3-C...) on prepared sheet. Record time and errors.","Completion <90 sec, ≤2 errors","Extremely slow >180 sec; multiple errors; gives up mid-task","Impulsive sequence errors; rushes without accuracy; mild planning deficits"),
    ("Verbal Fluency FAS","Name words beginning with F in 60 sec. Repeat for A and S. Count unique words; note repetitions.","≥12 words per letter; <3 repetitions","<6 words; flat verbal output; poor initiating","Normal count but prolonged pauses; intrusions of depressive content"),
    ("Stroop Interference","Show color-word cards. Name ink color only, ignoring printed word. Record time and errors for 20 items.","<45 sec; ≤3 errors","Extremely slow; disengaged; requires prompting to continue","Moderate interference errors; emotional Stroop effect on negative words"),
    ("Serial 7 Subtraction","Count backward from 100 by 7s. Record 5 subtractions. Note time, errors, and patient effort/initiation.","5/5 correct; completed within 30 sec","≤2/5 correct; poor initiation; requires encouragement; effortful","4–5/5 correct but slow; perseverative self-criticism on errors"),
    ("Go/No-Go Task","Read 20 letters at 1/sec; patient taps for every letter except X. Count commission errors (tapped X) and omission errors.","≤2 commission; ≤2 omission errors","High omission errors (>5); poor inhibitory initiation; psychomotor slowing","Moderate commission errors; impulsive responding despite depressed affect"),
    ("Category Switching","Name alternating fruits then vegetables (apple, broccoli, pear...) for 60 sec. Count switches and intrusion errors.","≥10 correct alternations; ≤2 intrusions","<5 alternations; difficulty switching; flat verbal output","Normal switching but slowed initiation; mild set-maintenance errors"),
  ],
  "sn_tests":[
    ("Heartbeat Counting","Count own heartbeats mentally for 25 sec without touching body. Compare to actual count. Calculate interoceptive accuracy score.","Interoceptive accuracy >0.75 (within ±5 beats)","Low accuracy <0.50; poor body awareness; alexithymia pattern","Overestimation of heartbeats; hyperawareness of somatic sensations; health anxiety"),
    ("Error Monitoring","Read sentence with 2 embedded factual errors aloud. Note whether patient spontaneously self-corrects or identifies errors.","Spontaneously catches ≥1 error; self-correction","Misses both errors; poor self-monitoring; apathetic processing","Extreme distress on noticing own errors; self-critical rumination; catastrophizing"),
    ("Emotional Salience Rating","Show 6 emotional (3 negative, 3 positive) and 6 neutral images. Rate personal salience/importance 0–10.","Balanced salience across valences; negatives ≤7, positives ≥4","Flat salience ratings across all images; anhedonic; emotional blunting","Negative images rated ≥9; positive images rated ≤2; extreme negativity bias"),
    ("Threat Appraisal","Present 3 ambiguous social scenarios in writing. Rate perceived threat 0–10.","Threat ratings 2–5 for ambiguous scenarios; balanced interpretation","Threat ratings ≤2 across all scenarios including objectively negative ones; emotional blunting","All scenarios rated ≥7; catastrophic interpretation of neutral social events"),
    ("Sensory Switching","Alternate between visual tracking (follow moving pen) and auditory counting (count tones). Measure switch latency.","Switch delay <2 sec; smooth transitions","Slow switching >5 sec; poor disengagement; psychomotor inertia","Impulsive switching; insufficient task maintenance; distractibility"),
    ("Pain Sensitivity","Apply firm thumb pressure to fingernail bed for 3 sec. Rate intensity 0–10. Compare bilateral.","Rating 3–6; bilateral symmetry; appropriate facial response","Low pain ratings ≤2; emotional blunting; reduced somatic awareness","Rating ≥8; hyperalgesia; somatic amplification consistent with depression-pain comorbidity"),
    ("Orienting Response","Deliver unexpected auditory tone (clap) mid-task. Observe behavioral interruption, reorientation time, and recovery.","Brief interruption; reorientation within 3 sec","Minimal orienting; blunted startle; disengaged vigilance","Prolonged behavioral disruption; exaggerated distress; slow recovery"),
  ],
  "smn_tests":[
    ("Finger Tapping Rate","Index finger tap on table as fast as possible for 10 sec. Count taps. Test dominant then non-dominant hand.","≥50 taps/10 sec dominant; ≥45 non-dominant","<35 taps; psychomotor retardation; slowed movement initiation","≥60 taps but with agitation; restless tapping; psychomotor agitation subtype"),
    ("Grip Assessment","Ask patient to squeeze examiner's two fingers as firmly as possible. Compare left vs right. Rate bilateral symmetry.","Firm symmetric grip; no hesitation","Weak bilateral grip; slow force generation; reduced motor drive","Normal grip but perseverative squeezing; agitated, can't release"),
    ("Fine Motor Sequencing","Thumb-to-finger opposition sequence (1-2-3-4-4-3-2-1) at maximum speed. Count complete sequences in 10 sec.","≥4 complete sequences in 10 sec","≤2 sequences; sequential errors; reduced motor speed and fluency","3–4 sequences but with agitated interruptions; motor restlessness"),
    ("Tandem Gait","Walk heel-to-toe on a straight line for 10 steps. Count steps off line. Observe posture and gait pattern.","≤1 step off line; upright posture; normal pace","Slow pace; stooped posture; multiple step failures; psychomotor slowing","Restless gait; unable to slow to tandem; agitated movement"),
    ("Romberg Test","Stand feet together, arms crossed, eyes closed for 30 sec. Observe postural sway, balance, and stance.","Minimal sway; stable stance for 30 sec","Excessive sway; requires wall support; poor proprioceptive engagement","Fidgeting/shifting weight during standing; unable to remain still; agitation"),
    ("Rapid Alternating Movements","Rapidly pronate and supinate forearm alternating for 10 sec on lap. Observe rhythm, speed, and bilateral coordination.","Regular rhythm; ≥14 cycles in 10 sec","Slow; irregular rhythm; reduced motor amplitude; effortful movements","Irregular rhythm; agitated rapid but uncoordinated movements"),
    ("Writing Sample","Write name and address in cursive; draw 10 connected circles. Assess size consistency, pressure, and writing speed.","Legible; consistent letter size; moderate pressure; smooth circles","Micrographia; reduced writing speed; poor pressure; incomplete circles","Normal size but tremor/scrawl from agitation; irregular letter spacing"),
    ("Reaction Time","Lift dominant hand from table as fast as possible when examiner drops pen. Measure 3 trials and average.","Average reaction time <250 ms","Reaction time >450 ms; psychomotor slowing; delayed motor initiation","Anticipatory responses (lifts before drop); motor restlessness; impulsive"),
  ],
  "limbic_tests":[
    ("Emotional Face Recognition","Show 6 photos representing: joy, sadness, fear, anger, disgust, surprise. Ask to identify each. Score /6.","≥5/6 correct identification","<4/6; blunted affect recognition; emotional numbing pattern","Hypersensitivity to sad/fearful faces; misidentification of neutral as negative"),
    ("Affect Labeling","Describe 3 emotional scenarios (loss of a loved one, receiving praise, unexpected noise). Ask: 'What emotion would you feel?'","Names appropriate distinct emotions for each scenario","Flat responses; 'I don't know' or 'nothing'; alexithymia; emotional blunting","Excessive negative emotion labeling; catastrophic responses to all scenarios"),
    ("Emotional Memory Bias","Read 10 emotional (5 negative, 5 positive) and 10 neutral words. Free recall after 5-min distractor task.","Equal or slight positive recall advantage; ≥8/20 total","Flat recall; poor emotional enhancement of memory; anhedonic","Selective negative word recall; positive words suppressed; depressive memory bias"),
    ("Fear Response Assessment","Present 3 mildly threatening scenarios verbally. Rate anxiety 0–10. Observe physiological signs (breathing, posture).","Anxiety ratings 2–5; proportionate responses; quick recovery","Flat anxiety response ≤2 even to objectively threatening scenarios; blunted fear","Anxiety ≥8 to all scenarios; catastrophic fear response; breathing changes"),
    ("Reward Sensitivity — Anhedonia Screen","Rate interest in 5 previously enjoyed activities (hobbies, food, social interaction) on 0–10 scale.","Average interest ≥6/10; motivational engagement maintained","Average ≤3/10; pervasive anhedonia; loss of pleasure (core Depression marker)","Average ≥8 but with anxious over-engagement; manic-mixed features"),
    ("Social Comfort Rating","Rate comfort in 5 social scenarios: crowded room, intimate conversation, meeting strangers, being observed, being alone.","Comfort ≥5/10 for ≥4 scenarios; social engagement maintained","All scenarios rated ≤3; social withdrawal; isolation; depressive avoidance","Anxious social discomfort ≥8 for all scenarios; social anxiety comorbidity"),
    ("Affective Flexibility","Recall a sad personal memory (rate mood). Then recall a happy memory (rate mood). Measure shift magnitude and recovery time.","Mood shift ≥3 points and return to baseline within 3 min","Flat mood across both memories; no emotional responsivity; anhedonic rigidity","Sad memory causes prolonged low mood >10 min; unable to access positive affect"),
  ],
  "attn_tests":[
    ("Sustained Attention — SART","Count backward from 50 aloud continuously. Rate performance consistency over 60 sec. Note fadeouts.","Counts consistently; ≤2 errors or stops in 60 sec","Frequent pauses; stops mid-count; poor sustained engagement; effortful","Monotonous perfect count with agitated pace; pressured speech subtype"),
    ("Selective Attention","In a list of 20 mixed letters, circle only the letter 'A'. Record time and commission/omission errors.","Completes in <45 sec; ≤2 errors","Slow completion >90 sec; multiple omissions; poor attentional engagement","Fast but imprecise; multiple false positives; attentional capture"),
    ("Divided Attention","While counting backward by 2s from 30, patient names objects examiner points to. Assess dual-task performance.","Maintains both tasks with ≤2 errors each","Abandons one task; poor cognitive resource sharing; effortful mono-tasking","Switches rapidly between tasks; agitated but inaccurate dual-task performance"),
    ("Visual Search","Find a target symbol (star) among 20 distractors (circles) on paper. Record time and accuracy.","Finds target in <15 sec; systematic search pattern",">45 sec; unsystematic; gives up; poor attentional sweep","Fast but random; impulsive pointing; misidentifies distractors"),
    ("Alerting Response","Examiner says 'ready' before a stimulus (tap on table). Measure preparatory behavioral response and readiness.","Postural readiness; focused gaze; appropriate anticipatory state","Blunted alerting; no preparatory response; disengaged; psychomotor inertia","Exaggerated startle; over-alerting; anxious anticipatory response"),
    ("Spatial Orienting","Valid/invalid arrow cue (drawn on paper) before peripheral target. Note accuracy and reaction time difference.","Orienting benefit ≥100 ms for valid cues; accurate","Absent orienting benefit; poor spatial attention direction","Hyper-orienting; inability to disengage from cued location; perseverative attention"),
    ("Response Inhibition Cancel","Cancel all circles on a sheet with mixed shapes (circles, squares, triangles). Speed and accuracy.","Completes in <60 sec; ≤2 omissions or false cancellations","Slow; >120 sec; multiple omissions; poor inhibitory drive","Fast but with multiple false cancellations; impulsive cancelling of non-targets"),
    ("Attentional Fatigue Probe","Letter cancellation for 5 min. Compare hit rate first 2.5 min vs second 2.5 min. Calculate fatigue index.","Performance drop <10% between first and second half",">30% decline; rapid attentional fatigue; poor sustained engagement","No decline but accelerating speed with increasing errors; agitated performance"),
  ],
  "phenotype_text":"Depression phenotypes based on network profile: (1) RUMINATIVE subtype — DMN dominant HYPER with CEN HYPO; (2) ANHEDONIC subtype — Limbic-Reward HYPO dominant; (3) PSYCHOMOTOR subtype — SMN HYPO dominant with CEN HYPO; (4) ANXIOUS-DEPRESSIVE — SN+Limbic HYPER with DMN HYPER. Map highest deficit network to primary protocol target.",
  "fnon_strategy":"Primary: Inhibit DMN hyperactivity (cathodal/low-frequency over mPFC or right PFC). Facilitate CEN (anodal/high-frequency left DLPFC). Secondary: Normalize limbic reactivity via left DLPFC upregulation. Tertiary: SMN activation for psychomotor subtypes.",
}

# ── 2. ANXIETY ────────────────────────────────────────────────────────────────
ANXIETY = {
  "full":"Generalized Anxiety Disorder (GAD)","short":"Anxiety","slug":"anxiety",
  "adapted_note":"Includes worry severity markers, interoceptive hypervigilance tests, and Anxiety phenotype markers",
  "dmn_rel":"GAD is characterized by DMN HYPERACTIVITY with pathological worry — future-oriented negative self-referential processing. Unlike Depression's past-focus, Anxiety drives prospective DMN over-engagement: repetitive 'what if' scenarios and catastrophic future simulations.",
  "cen_rel":"CEN HYPOACTIVITY in GAD reflects executive resources hijacked by worry processes. DLPFC fails to adequately suppress amygdala-driven threat appraisals. Working memory is congested with anxious content, reducing cognitive flexibility and problem-solving capacity.",
  "sn_rel":"SN is the core dysfunction in GAD — HYPERACTIVITY of the salience detection system misidentifies neutral stimuli as threatening. ACC and anterior insula HYPER-engagement creates persistent false alarm signaling, interoceptive hypersensitivity, and autonomic arousal.",
  "smn_rel":"SMN shows HYPER-arousal in GAD — muscle tension, physiological hyperactivation, tremor, and restlessness reflect elevated sympathetic tone. Somatic symptoms (tension headache, GI symptoms) arise from chronic SMN/autonomic HYPER-engagement.",
  "limbic_rel":"Amygdala HYPERACTIVITY drives the fear-based core of GAD. Hyperactive threat detection with impaired prefrontal fear regulation results in chronically elevated emotional reactivity, fear generalization, and avoidance behavior.",
  "attn_rel":"Attention Networks show HYPER-vigilance (DAN/VAN HYPER) toward threat-related stimuli alongside reduced top-down attentional control. Patients cannot disengage from threat cues (attentional lock), while divided and sustained attention are impaired by worry.",
  "dmn_tests":[
    ("Worry Frequency Probe","Ask: 'How often do worrying thoughts occur in a typical day?' Rate 0–10. Ask: 'Can you stop them when you want?'","Occasional controllable worry; 0–3/10; able to redirect","Minimal worry (HYPO: emotional blunting or dissociation pattern)","Worry rated ≥8/10; uncontrollable; occurs >6 hours/day; GAD hallmark"),
    ("Autobiographical Future Simulation","Ask patient to describe their life 1 year from now. Rate positive vs negative content and level of catastrophizing.","Balanced future projection; realistic optimism; some uncertainty tolerated","Flat or absent future thinking; poor prospective simulation; depressive overlap","Exclusively catastrophic futures; 'what if' elaboration; multiple disaster scenarios"),
    ("Self-Reference Memory","Read 10 adjectives with 'Does this describe you?' frame. Recall after 2-min distractor. Note anxious vs calm descriptors recalled.","≥6/10 recalled; balanced positive-negative self-descriptors","<4/10 recalled; poor self-schema engagement; emotional avoidance","Negative/threat self-descriptors preferentially recalled; positive self-views suppressed"),
    ("Mind Wandering Probe","2-min reading passage with 4 random probes. Note task-unrelated thought content — specifically worry themes.","≤1/4 off-task; neutral or task-relevant mind-wandering","0/4 off-task — suppression or emotional blunting","3–4/4 off-task; worry content; future-oriented catastrophic themes"),
    ("Prospective Memory","Tell patient to remind you of a specific word in 5 min. Note: does patient check the time repeatedly? Rate anxiety about forgetting.","Recalls within 5±1 min; no time-checking anxiety","Forgets entirely; poor future self-monitoring","Repeatedly checks time or asks if it's been 5 min yet; hypervigilant to future task"),
    ("Default Narrative","Ask: 'Tell me something about your plans for next week.' Rate coherence and anxiety content.","Balanced description; some flexibility; manageable uncertainty","Flat response; poor planning; disinterested in future","Excessive catastrophic scenarios; unable to describe neutral plans; catastrophizing"),
    ("Semantic Fluency — Future","Name things that could go wrong vs go right in a hypothetical trip. Count positive vs negative items generated.","Balanced generation; ≥4 positive and ≥4 negative items","Flat generation; unable to generate either valence; emotional blunting","≥8 negative items; ≤2 positive; strong threat generation bias"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Read digit sequences; ask reverse recall. Start at 2 digits and increase. Score maximum span achieved.","Backward span ≥5 digits","Span ≤3; working memory overwhelmed by anxiety; cognitive interference","Span intact but performance fluctuates; anxiety disrupts under time pressure"),
    ("Trail Making B","Alternating number-letter connections on paper. Record time and errors.","<90 sec; ≤2 errors","Extremely slow; gives up; decision anxiety paralysis","Fast but with multiple errors; impulsive choices; anxiety-driven rushing"),
    ("Verbal Fluency FAS","Name words beginning with F in 60 sec. Note if patient generates anxiety/threat-related words preferentially.","≥12 words; <3 repetitions; mixed content","<7 words; poor retrieval; cognitive interference from worry","≥12 words but many are threat/worry-related (fears, failures, fight...)"),
    ("Stroop Interference","Emotional Stroop variant: include anxiety-related words (danger, threat). Compare interference on neutral vs anxiety words.","Uniform response time across word types; ≤3 errors","No Stroop effect; flat processing; emotional blunting","Selectively slowed on anxiety-related words; emotional Stroop effect prominent"),
    ("Serial 7 Subtraction","Backward 100-7 subtraction for 5 steps. Note accuracy and whether errors trigger distress or self-criticism.","5/5 correct; calm error acceptance if any","≤2/5; poor concentration; cognitive load from worry","4–5/5 correct but extreme distress on any error; perfectionism; self-criticism"),
    ("Go/No-Go Task","Read 20 letters at 1/sec; tap for all except X. Count commission and omission errors.","≤2 commission; ≤2 omission errors","High omission errors; poor sustained initiation","High commission errors; impulsive responding; over-vigilant false alarms"),
    ("Category Switching","Alternating fruits and vegetables for 60 sec. Note switching speed and intrusion errors.","≥10 correct alternations; ≤2 intrusions","<5 alternations; rigid category maintenance; inflexibility","Normal count but rushed, anxious pacing; perseverative corrections"),
  ],
  "sn_tests":[
    ("Heartbeat Counting — Interoception","Count own heartbeats for 25 sec without touching body. Compare to measured pulse.","Accuracy >0.75; appropriate body awareness","Accuracy <0.50; poor interoception; body disconnection","Accuracy >0.90 with extreme focus; hyperawareness; health anxiety; palpitation focus"),
    ("Threat Appraisal — GAD Screen","Present 5 ambiguous scenarios (e.g., boss calls you in, friend doesn't reply). Rate threat 0–10.","Ambiguous scenarios rated 2–4; tolerates uncertainty","All scenarios rated ≤2; denial or minimization pattern","All scenarios rated ≥8; catastrophic interpretations; intolerance of uncertainty"),
    ("Error Monitoring","Read passage with 2 embedded errors. Note spontaneous error detection and emotional response to errors found.","Catches ≥1 error calmly; mild satisfaction","Misses both errors; poor ACC self-monitoring","Catches errors but excessive distress; replays passage to find more; perfectionism"),
    ("Sensory Switching","Alternate visual tracking (pen) and auditory counting (tones) tasks. Measure switching cost.","Switch delay <2 sec; smooth transitions","Slow switching >5 sec; poor disengagement; inertia","Rapid switching but with anticipatory anxiety about next task; hypervigilant"),
    ("Pain/Somatic Sensitivity","Apply light pressure to fingernail bed for 3 sec. Rate intensity 0–10. Ask about somatic tension (head, shoulders, stomach).","Pain rating 3–6; minimal somatic complaints","Pain rating ≤2; poor somatic awareness; body disconnection","Pain rating ≥8; somatic amplification; reports tension in ≥3 body areas"),
    ("Surprise/Orienting Response","Deliver unexpected auditory tone mid-task. Observe startle magnitude, reorientation time, and physiological signs.","Brief interruption; reorientation within 3 sec; calm recovery","Blunted startle; poor orienting; disengaged vigilance","Exaggerated startle; prolonged recovery >10 sec; autonomic activation signs"),
    ("Emotional Salience Rating","Show 6 emotional (3 threat-related, 3 positive) and 6 neutral images. Rate salience 0–10.","Balanced ratings; threat ≤6, positive ≥4","Flat ratings; emotional blunting across all image types","Threat images rated ≥9; positive/neutral rated ≤3; strong threat salience bias"),
  ],
  "smn_tests":[
    ("Finger Tapping Rate","Index finger tapping for 10 sec each hand. Count taps. Note tremor, restlessness, or irregular rhythm.","≥50 taps/10 sec; regular rhythm; no tremor","<35 taps; slow; reduced motor drive","Normal count but with visible tremor; restless hand movements; motor tension"),
    ("Grip Assessment","Firm squeeze of examiner's fingers bilaterally. Note strength and whether patient can release smoothly.","Firm symmetric grip; smooth release","Weak bilateral grip; disengaged muscle tone","Intense grip; difficulty releasing; white-knuckle tension; gripping response"),
    ("Muscle Tension Survey","Palpate/observe tension in trapezius, jaw (TMJ), and forearms. Rate visible tension 0–10 per region.","Minimal tension ≤3 in all regions; relaxed posture","Flat; no tension; hypotonic; dissociated body awareness","Tension ≥7 in ≥2 regions; teeth clenching; rigid posture; somatic hyperactivation"),
    ("Tandem Gait","Heel-to-toe walk for 10 steps on a line. Count steps off line and observe postural tension.","≤1 step off line; relaxed gait","Multiple step failures; poor proprioceptive engagement; slow pace","Steps performed rigidly; tense upright posture; hypervigilant gait"),
    ("Romberg Test","Feet together, eyes closed for 30 sec. Observe sway and note visible tension/fidgeting.","Minimal sway; stable 30 sec; relaxed posture","Excessive sway; poor proprioception; motor disengagement","Stable but with visible tension; fidgeting; unable to relax into stance"),
    ("Rapid Alternating Movements","Pronate/supinate forearm for 10 sec. Observe rhythm, speed, and tremor.","Regular rhythm; ≥14 cycles; no tremor","Slow; reduced amplitude; poor motor engagement","Normal speed but tremulous; irregular rhythm; tension-induced dysrhythmia"),
    ("Writing Sample","Write name and address; draw 10 circles. Assess pressure, size, and tremor.","Legible; consistent; moderate pressure; smooth circles","Reduced pressure; micrographia; flat motor engagement","Heavy pressure; irregular size; visible tremor; rushed scrawl from anxiety"),
    ("Reaction Time","Hand lift from table when pen drops. 3 trials averaged.","Average <250 ms; consistent across trials",">400 ms; slow motor initiation; poor readiness","<150 ms average; anticipatory lifts; motor hypervigilance; false starts"),
  ],
  "limbic_tests":[
    ("Emotional Face Recognition","Identify 6 emotions (joy, sadness, fear, anger, disgust, surprise) from facial photos. Score /6.","≥5/6 correct","<4/6; poor emotional recognition; blunted affect","Misidentifies neutral as fearful/angry; hypersensitivity to threatening expressions"),
    ("Fear Generalization","Present 3 mildly threatening scenarios then 3 neutral scenarios. Rate anxiety 0–10 for each.","Anxiety 5–7 for threatening, ≤3 for neutral; appropriate differentiation","Flat anxiety across all scenarios; blunted fear response","High anxiety (≥7) for neutral scenarios; fear generalization; threat overgeneralization"),
    ("Emotional Memory — Threat Bias","Read 10 words (5 threat-related: danger, attack; 5 positive/neutral). Free recall after 5-min distractor.","Balanced recall or slight positive advantage","Flat recall; poor emotional memory enhancement","Threat words preferentially recalled; positive words suppressed; anxiety memory bias"),
    ("Affect Labeling","Describe 3 emotional scenarios. Ask: 'What emotion would you feel?' Note: does patient label anxiety in non-threatening scenarios?","Accurate, differentiated emotional labels","'Nothing' or unable to identify emotions; alexithymia","Labels anxiety/fear even in positive scenarios; over-activation of fear response"),
    ("Reward Sensitivity","Rate interest in 5 pleasant activities 0–10. Note whether anxiety about activities (not anhedonia) drives avoidance.","Average interest ≥6/10; approach motivation maintained","Average ≤3; anhedonia (screen for Depression overlap)","Average ≥7 but with anxious over-planning; avoidance of spontaneous pleasure"),
    ("Social Comfort Rating","Rate comfort in 5 social scenarios. Note whether anxiety (vs depression) is the reported barrier to social engagement.","Comfort ≥5/10 for ≥4 scenarios","Comfort ≤3 due to depression/withdrawal (not anxiety)","Comfort ≤3 for ≥4 scenarios; social anxiety; fear of negative evaluation"),
    ("Affective Flexibility","Sad then happy memory recall. Rate mood before/after. Note: does anxious anticipation interfere with emotional switching?","Mood shift ≥3 points; returns to baseline within 3 min","Flat mood; no emotional responsivity; anhedonic rigidity","Unable to access positive memory; anxiety about 'letting guard down'; hypervigilant affect"),
  ],
  "attn_tests":[
    ("Sustained Attention","Count backward from 50 aloud for 60 sec. Note consistency and whether anxiety disrupts sustained performance.","≤2 errors; consistent pace","Frequent pauses; poor sustained engagement","Anxious rushing; accelerating pace; anxious self-monitoring of performance"),
    ("Attentional Bias — Threat","In a 20-item letter list, circle all A's. Include 5 threat-related words in context. Note if threat words cause attentional capture.","Uniform attention; ≤2 errors; no content-specific slowing","Slow overall completion; poor attentional drive","Selective slowing near threat words; attentional capture; anxiety-driven distractibility"),
    ("Divided Attention","Count backward by 2s while naming pointed objects. Dual-task performance assessment.","Maintains both tasks; ≤2 errors each","Abandons one task; poor resource sharing","Anxious performance; monitors own performance; self-interrupts to check accuracy"),
    ("Visual Search","Find star target among 20 distractors. Time and accuracy.","<15 sec; systematic search",">45 sec; unsystematic; gives up","Fast but hypervigilant; finds target quickly but continues checking for more threats"),
    ("Alerting Response","Examiner says 'ready' before stimulus. Assess preparatory state and arousal level.","Appropriate alerting; focused; controlled readiness","Blunted alerting; disengaged; poor preparatory arousal","Hyper-alerting; over-aroused; exaggerated anticipatory response; vigilant scanning"),
    ("Spatial Orienting","Arrow cue (valid/invalid) before peripheral target. Measure orienting benefit and disengagement speed.","Orienting benefit ≥100 ms; smooth disengagement","Absent orienting benefit; poor spatial attention","Attentional lock; difficulty disengaging from cued location; hypervigilant orienting"),
    ("Response Inhibition Cancel","Cancel all circles on mixed-shape sheet. Speed and accuracy.","<60 sec; ≤2 errors","Slow; >120 sec; omissions; poor inhibitory drive","Fast with false alarms; impulsive cancelling; attentional hyperactivation"),
    ("Attentional Fatigue","Letter cancellation 5 min; compare first vs last 2.5 min hit rate.","Performance drop <10%",">30% decline; rapid fatigue; poor sustained drive","No decline but accelerating errors; hypervigilant pace maintained at cost of accuracy"),
  ],
  "phenotype_text":"Anxiety phenotypes: (1) WORRY-DOMINANT — DMN HYPER + CEN HYPO; (2) SOMATIC — SN+SMN HYPER dominant; (3) FEAR-AVOIDANCE — Limbic HYPER + attention bias; (4) HYPERVIGILANCE — Attention HYPER + SN HYPER. Map highest-scoring network to primary protocol target.",
  "fnon_strategy":"Primary: Inhibit SN hyperactivity (cathodal/low-frequency over ACC/insula, or right DLPFC). Facilitate CEN (anodal left DLPFC) to enhance prefrontal regulation of amygdala. Secondary: Normalize Limbic reactivity via DLPFC upregulation.",
}

# ── 3. ADHD ───────────────────────────────────────────────────────────────────
ADHD = {
  "full":"ADHD","short":"ADHD","slug":"adhd",
  "adapted_note":"Includes inattention/hyperactivity subtype differentiation, executive function markers, and ADHD phenotype indicators",
  "dmn_rel":"ADHD involves paradoxical DMN HYPERACTIVITY during task engagement — failure to suppress default mode when attention is required. The mPFC-PCC circuit fails to deactivate, causing mind-wandering, task-irrelevant thoughts, and attentional lapses.",
  "cen_rel":"CEN is the CORE HYPOACTIVE network in ADHD — DLPFC, posterior parietal, and ACC hypofunction underlie working memory deficits, poor response inhibition, and executive dysfunction. Mesocortical dopamine/noradrenaline deficiency drives CEN underperformance.",
  "sn_rel":"SN shows mixed dysfunction: HYPO error-monitoring (poor ACC engagement, lack of self-correction) in inattentive subtype; HYPER sensory reactivity and impulsive salience detection in hyperactive/combined subtype.",
  "smn_rel":"ADHD-Combined shows SMN HYPER with hyperactivity, fidgeting, and motor restlessness. ADHD-Inattentive may show mild SMN HYPO (sluggish cognitive tempo). Cerebellar-motor circuit involvement produces coordination and sequencing difficulties.",
  "limbic_rel":"Limbic network shows HYPO reward sensitivity (striatal dopamine deficiency causing reduced motivation) alongside HYPER emotional reactivity (ADHD emotional dysregulation). Rejection dysphoria reflects limbic hyperresponsivity.",
  "attn_rel":"Attention Networks show complex dual dysfunction: HYPO sustained/selective attention and HYPER distractibility. DAN HYPO reflects poor top-down attentional control; VAN HYPER reflects excessive bottom-up attentional capture by irrelevant stimuli.",
  "dmn_tests":[
    ("Task-Unrelated Thought Probe","Give 2-min reading passage; probe 4 times: 'Were you thinking about the text?' Count off-task responses.","≤1/4 off-task responses","0/4 off-task (over-compliance, anxiety, or rigid task focus)","3–4/4 off-task; mind-wandering; ADHD hallmark DMN failure to suppress"),
    ("Autobiographical Memory","Ask 3 personal memories. Note coherence, narrative structure, and chronological organization.","Specific, organized memories with temporal sequence","Sparse memories; poor narrative structure; effortful recall","Disorganized, tangential narrative; jumps between unrelated memories; loses thread"),
    ("Prospective Memory","Tell patient to remind you of a specific word in 5 min. Continue assessment. Note if remembered.","Independent recall within 5±1 min","Completely forgets; poor prospective monitoring","Repeatedly interrupts assessment to ask if it's time yet; impulsive time monitoring"),
    ("Self-Reference Memory","Read 10 adjectives with 'Does this describe you?' frame. Recall after 2-min distractor.","≥6/10 recalled; balanced content","<4/10; poor self-schema engagement","Normal recall but dominated by ADHD-related self-descriptors (restless, forgetful, distracted)"),
    ("Default Narrative","Ask: 'Tell me about something you did this week.' Rate: narrative coherence, topic maintenance, tangentiality.","Coherent narrative; ≤2 topic shifts; resolves story","Short vague response; flat affect; poor self-monitoring","Highly tangential; ≥4 topic shifts; never returns to original story; ADHD narrative pattern"),
    ("Mind-Wandering During Waiting","Ask patient to sit quietly for 60 sec without speaking. Observe: fidgeting, vocalizing, looking around.","Sits relatively still; may look around but self-regulates","Sits motionless but dissociated; flat; inattentive subtype pattern","Fidgets, vocalizes, or interrupts within 30 sec; hyperactive pattern"),
    ("Semantic Fluency","Name as many sports as possible in 1 min. Note: clustering, switching, and total count.","≥12 items; appropriate clustering and switching","<7 items; poor fluency; sluggish cognitive tempo","≥15 items but with poor organization; rapid random output; impulsive generation"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Start at 2 digits; increase until failure. Score maximum span.","Backward span ≥5 digits","Span ≤3; working memory impairment; ADHD core deficit","Span ≤3 with frustration; impulsive early responses; starts before sequence ends"),
    ("Trail Making B","Alternating number-letter connections. Record time and errors. Note impulsive wrong connections.","<90 sec; ≤2 errors","Extremely slow; sluggish cognitive tempo subtype","Multiple impulsive sequence errors; rushes without planning; restarts frequently"),
    ("Verbal Fluency FAS","F-A-S word generation for 60 sec each. Count unique words, repetitions, and rule violations.","≥12 words/letter; <3 repetitions; no rule violations","<7 words; slow initiation; poor verbal output","≥12 words but with repetitions; rule violations (proper nouns); impulsive output"),
    ("Stroop Color-Word","Name ink color ignoring printed word. Record interference errors and time.","<45 sec; ≤3 errors","Slow but accurate; inhibitory dysfunction pattern","Multiple impulsive errors; names word instead of color; poor response inhibition"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps. Note accuracy, self-correction, and frustration tolerance.","5/5 correct; self-corrects errors calmly","≤2/5; slow; poor working memory maintenance","3–4/5 correct but impulsive; skips steps; becomes frustrated or disengaged"),
    ("Go/No-Go Task","Tap for all letters except X. 20 letters at 1/sec. Count commission and omission errors.","≤2 commission; ≤2 omission errors","High omission errors >5; poor engagement; inattentive subtype","High commission errors >5; impulsive tapping; fails to inhibit X response; hyperactive"),
    ("Category Switching","Alternating fruits and vegetables for 60 sec. Count correct switches and intrusion errors.","≥10 correct alternations; ≤2 intrusions","<5 alternations; rigid; poor cognitive flexibility","Normal count but perseverative intrusions; set-shifting difficulty; executive error pattern"),
  ],
  "sn_tests":[
    ("Heartbeat Counting — Interoception","Count heartbeats for 25 sec without touching body. Calculate interoceptive accuracy.","Accuracy >0.75","Accuracy <0.50; poor body awareness; inattentive subtype pattern","Variable accuracy; hyperawareness at times then complete unawareness; inconsistent salience"),
    ("Error Monitoring","Read sentence with 2 errors. Note spontaneous detection and response.","Catches ≥1 error; calm correction","Misses both errors; poor ACC monitoring; inattentive subtype","Catches errors impulsively; interrupts reading; may over-correct non-errors"),
    ("Emotional Salience Rating","Rate 6 emotional and 6 neutral images for personal salience 0–10.","Balanced ratings; appropriate emotional differentiation","Flat ratings; poor emotional engagement; sluggish cognitive tempo","Extreme ratings; over-assigns salience to novelty/interesting images; impulsive rating"),
    ("Sensory Switching","Alternate visual tracking and auditory counting. Measure switch latency and accuracy.","Switch delay <2 sec; smooth transitions","Slow switching; inattentive disengagement from first task","Rapid but inaccurate switching; abandons first task before instructed; impulsive"),
    ("Threat Appraisal","Rate threat in 5 ambiguous social scenarios 0–10.","Ratings 2–5; balanced interpretation","Flat ratings ≤2; poor social awareness; inattentive pattern","Extreme variability; oscillates 0–10 within same scenario set; impulsive appraisals"),
    ("Surprise/Orienting Response","Unexpected auditory tone mid-task. Observe startle, reorientation, and recovery.","Brief interruption; reorientation within 3 sec","Minimal startle; poor orienting; inattentive subtype","Exaggerated startle; prolonged behavioral disruption; attentional capture by novel stimulus"),
    ("Pain Sensitivity","Light pressure to fingernail bed for 3 sec. Rate intensity 0–10.","Rating 3–6; bilateral symmetry","Rating ≤2; poor somatic awareness; inattentive type","Rating variable; may ignore during distraction then over-report when focused; inconsistent salience"),
  ],
  "smn_tests":[
    ("Finger Tapping Rate","Index finger tapping 10 sec each hand. Count taps and note rhythm consistency.","≥50 taps/10 sec; regular rhythm","<40 taps; sluggish; poor motor initiation","≥55 taps but with irregular rhythm; variable rate; motor impersistence"),
    ("Fidgeting Observation","During the full 10-min assessment, count self-directed movements (leg bouncing, hand touching face, seat shifting, object handling).","≤3 self-directed movements per 10 min","0 self-directed movements; rigid stillness; over-inhibited","≥8 self-directed movements; motor restlessness; hyperactive subtype marker"),
    ("Fine Motor Sequencing","Thumb-to-finger opposition sequence at maximum speed for 10 sec.","≥4 complete sequences","≤2 sequences; sluggish motor execution","3–4 sequences but with error and restart; loses sequence; impulsive motor pattern"),
    ("Tandem Gait","Heel-to-toe walk 10 steps. Count steps off line and note pace consistency.","≤1 step off line; consistent pace","Multiple steps off; slow; poor motor attention","≥3 steps off; variable pace; distracted gait; ADHD motor coordination pattern"),
    ("Romberg Test","Feet together eyes closed 30 sec. Observe sway and movement.","Minimal sway; stable","Excessive sway; poor proprioception; disconnected","Fidgets; shifts weight repeatedly; cannot maintain still stance for 30 sec"),
    ("Rapid Alternating Movements","Pronate/supinate 10 sec. Rate rhythm, speed, amplitude.","Regular rhythm; ≥14 cycles","Slow; reduced amplitude; poor motor drive","Fast but irregular; motor impersistence; amplitude variation; hyperkinetic pattern"),
    ("Writing Sample","Write name and address; draw 10 circles. Note size consistency and organization on page.","Legible; consistent size; organized on page","Reduced pressure; incomplete; motor disengagement","Variable size; circles wander across page; erratic pressure; ADHD motor variability"),
    ("Reaction Time","Hand lift 3 trials. Average response time. Note false starts.","Average <250 ms; ≤1 false start",">400 ms; slow reaction; poor motor readiness","<200 ms average with ≥2 false starts; impulsive; cannot wait for stimulus"),
  ],
  "limbic_tests":[
    ("Reward Sensitivity — ADHD","Rate motivation for 5 activities: immediate reward vs delayed reward scenarios. Note discounting of delayed rewards.","Consistent motivation regardless of delay; approach motivation maintained","Flat motivation for all activities; anhedonic pattern (screen for Depression)","Strong immediate preference; significant delay discounting; reward-seeking impulsivity"),
    ("Emotional Face Recognition","Identify 6 emotions from facial photos. Score /6.","≥5/6 correct","<4/6; poor affect recognition; social emotional HYPO","Normal accuracy but impulsive responses; guesses before examining fully"),
    ("Affect Labeling","3 emotional scenarios. Name own emotional response.","Accurate, differentiated labels","Flat labels; 'I don't know'; poor affect awareness","Impulsive emotional labeling; over-identifies with extreme emotions; emotional dysregulation"),
    ("Rejection Dysphoria Screen","Ask: 'How do you feel when criticized or when you feel you've failed someone?' Rate emotional intensity 0–10.","Intensity 3–6; proportionate; resolves within 1 hour","Flat response ≤2; emotional blunting; avoidant","Intensity ≥9; 'world-ending' quality; hours-long emotional response; RSD marker"),
    ("Fear Response Assessment","3 mildly threatening scenarios. Rate anxiety 0–10.","Anxiety 2–5; proportionate","Flat ≤2; poor fear response; avoidant pattern","Extreme variability; some scenarios ≥9; ADHD emotional dysregulation"),
    ("Social Comfort Rating","Rate comfort in 5 social scenarios 0–10.","≥5/10 for ≥4 scenarios","≤3 for all; withdrawn; avoidant","Variable; some scenarios very high, others very low; impulsive social engagement"),
    ("Affective Flexibility","Sad then happy memory recall. Rate mood shift magnitude and recovery.","Mood shift ≥3 points; returns to baseline <3 min","Flat shift; no emotional responsivity","Rapid large shifts; emotional lability; ADHD emotional dysregulation pattern"),
  ],
  "attn_tests":[
    ("Sustained Attention — Vigilance","Count backward from 50 aloud for 60 sec. Note fadeouts, errors, and consistency of pace.","≤2 errors; consistent pace","Monotonous consistent count; over-controlled","Frequent pauses; pace failures; mind-wandering; ADHD sustained attention core deficit"),
    ("Selective Attention","Circle all A's in 20-item letter list. Record time and errors.","<45 sec; ≤2 errors","Slow; poor attentional engagement","Fast but multiple omissions AND false alarms; impulsive/inattentive combined"),
    ("Divided Attention","Count backward by 2s while naming pointed objects. Dual-task performance.","≤2 errors per task","Abandons one task; poor engagement","Completely abandons one task; impulsive task switching; ADHD divided attention failure"),
    ("Visual Search","Find star target among 20 distractors. Time and accuracy.","<15 sec; systematic search pattern",">45 sec; gives up; poor search initiation","Fast but random; misses target due to impulsive incomplete search"),
    ("Alerting Response","Examiner says 'ready' before stimulus. Assess preparatory state.","Appropriate alerting; focused readiness","Blunted alerting; not ready; disengaged","Over-alerting; impulsive response before stimulus; cannot wait; false starts"),
    ("Spatial Orienting","Arrow cue before peripheral target. Note orienting benefit and disengagement.","Orienting benefit ≥100 ms; smooth disengagement","Absent benefit; poor spatial attention direction","Normal orienting but difficulty disengaging from invalid cues; attentional capture"),
    ("Response Inhibition Cancel","Cancel all circles on mixed-shape sheet. Speed and accuracy.","<60 sec; ≤2 errors","Slow; >120 sec; multiple omissions","Fast with multiple false cancellations; impulsive cancelling of non-targets; ADHD marker"),
    ("Attentional Fatigue","Letter cancellation 5 min; compare first vs last 2.5 min.","Performance drop <10%","Consistent low performance throughout; flat engagement","Sharp >40% decline; rapid attentional fatigue; vigilance decrement; ADHD fatigue pattern"),
  ],
  "phenotype_text":"ADHD phenotypes: (1) INATTENTIVE — CEN HYPO + DMN HYPER + Attention HYPO; (2) HYPERACTIVE-IMPULSIVE — SMN HYPER + SN HYPER + Limbic HYPER; (3) COMBINED — all networks affected; (4) SLUGGISH COGNITIVE TEMPO — SMN HYPO + Attention HYPO + DMN HYPO. Subtype guides stimulation laterality and frequency selection.",
  "fnon_strategy":"Primary: Facilitate CEN via anodal/high-frequency left DLPFC stimulation to enhance top-down executive control. Secondary: Inhibit DMN (right PFC or mPFC cathodal). For hyperactive subtype: add SN regulation. Stimulate during cognitive tasks for optimal efficacy.",
}

# ── 4. ALZHEIMER'S / MCI ──────────────────────────────────────────────────────
ALZHEIMERS = {
  "full":"Alzheimer's Disease / MCI","short":"Alzheimer's","slug":"alzheimers",
  "adapted_note":"Includes progressive memory loss markers, default mode failure indicators, and Alzheimer's/MCI phenotype staging",
  "dmn_rel":"Alzheimer's disease begins with DMN FAILURE — amyloid deposition preferentially in DMN hubs (mPFC, PCC, angular gyrus) disrupts default mode connectivity. Early AD shows DMN HYPO with autobiographical memory loss, spatial disorientation, and impaired self-referential processing.",
  "cen_rel":"CEN becomes progressively HYPOACTIVE as AD advances. Executive dysfunction, poor working memory, and impaired semantic retrieval reflect frontal-parietal network failure. MCI may show preserved CEN with isolated DMN dysfunction — critical for early staging.",
  "sn_rel":"SN shows complex dysfunction in AD: HYPO error monitoring (poor self-correction of errors) alongside HYPO salience discrimination (unable to identify what's important in environment). Late-stage HYPER may emerge as behavioral symptoms (agitation, wandering).",
  "smn_rel":"SMN remains relatively spared in early AD but shows progressive HYPO as disease advances: reduced gait speed, balance impairment, and extrapyramidal signs in moderate-severe stages. Motor preservation in early disease helps differentiate AD from DLB.",
  "limbic_rel":"Medial temporal lobe degeneration causes profound LIMBIC HYPO: episodic memory failure, emotional memory loss, and reduced affective response. In behavioral variant, paradoxical HYPER emotional reactions (apathy, agitation) may alternate.",
  "attn_rel":"Attention networks show progressive HYPO: visuospatial attention, spatial orientation, and sustained attention all degrade. Parietal degeneration disrupts DAN function; executive-attention dysregulation impairs VAN response inhibition.",
  "dmn_tests":[
    ("Episodic Memory — MMSE Recall","Tell 3 unrelated words (apple, table, penny). Distract for 5 min with other tests. Ask for free recall.","Recalls 3/3 without cues; consistent across session","Unable to recall even with semantic cues; severe DMN failure (AD marker)","Recalls 3/3 but perseveratively repeats words before asked; confabulatory"),
    ("Clock Drawing Test","Draw a clock showing 10:10. Score: circle (1), all numbers (1), correct placement (1), both hands correct (1). Max 4.","Score 4/4; appropriate size; symmetric layout","Score 0–2; missing numbers/hands; spatial disorientation (DMN hub failure)","Score 3–4 but slow; anxious; excessively checked (MCI compensatory behavior)"),
    ("Autobiographical Memory","Ask 3 personal memories (childhood, young adult, recent). Rate specificity, chronology, and accuracy if verifiable.","Specific, episodically rich memories; correct temporal sequence","Remote memories intact but recent memories lost (temporal gradient); AD pattern","Confabulates recent events with distant memories; temporal disorientation"),
    ("Orientation — Time and Place","Ask: current date, year, season, place (building, city, state, country). Score /8.","8/8 correct; no hesitation","<6/8; disoriented to time then place; AD progression marker","Answers correctly but repetitively asks for confirmation; MCI anxiety"),
    ("Verbal Semantic Memory","Category fluency: name as many animals as possible in 1 min. Count and note intrusions.","≥15 animals; <3 repetitions; no intrusions","<8 animals; semantic memory degradation; AD fluency pattern","12–14 animals but with 3–5 perseverative repetitions; semantic network degradation"),
    ("Spatial Self-Reference","Draw an arrow pointing to the exit. Draw a map of the room. Rate spatial self-orientation.","Correct arrow and rough room map; spatial reference maintained","Unable to indicate exit or map room; severe spatial disorientation","Correct arrow but confuses own position in room; egocentric spatial failure"),
    ("Story Recall","Read a 3-sentence story aloud. Ask to repeat it immediately then after 5 min delay.","Immediate: ≥80% details; Delayed: ≥60% details","Immediate <50%, delayed 0%; encoding failure; AD episodic memory marker","Immediate 80% but contaminated with intrusions; confabulatory elaboration"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Start 2 digits; increase.","Backward span ≥5","Span ≤3; working memory failure; moderate AD","Span ≥5 but perseverative errors; intrusion from prior digits; executive intrusion"),
    ("Trail Making B","Alternating number-letter connections. Time and errors.","<90 sec; ≤2 errors","Unable to sequence; loses alternating rule; moderate-severe AD","Multiple self-corrected errors; slow; MCI executive vulnerability"),
    ("Verbal Fluency FAS","Letter fluency for F, A, S. Count unique words per letter.","≥9 words/letter","<5 words/letter; poor letter fluency; frontal CEN failure","7–9 words but with rule violations and intrusions; CEN degradation"),
    ("Semantic Category Sorting","Present 6 items (2 fruits, 2 vehicles, 2 tools). Ask patient to sort into groups. Note category errors.","Correct 3-group sorting; no cross-contamination","Cannot identify categories; semantic system failure; moderate-severe AD","Correct sorting but slow; requires multiple prompts; MCI executive degradation"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct; <30 sec","≤2/5 correct; loses count; serial arithmetic failure; AD CEN deficit","3–4/5 with self-corrections; slow; MCI executive vulnerability"),
    ("Go/No-Go Task","Tap for all letters except X. 20 letters at 1/sec.","≤2 commission; ≤2 omission errors","High omissions; poor sustained initiation; executive-attention failure","Understands rule but makes multiple commission errors; response inhibition failure"),
    ("Category Switching","Alternate fruits and vegetables for 60 sec.","≥10 alternations; ≤2 intrusions","<5 alternations; category perseveration; cognitive rigidity","5–8 alternations with intrusion errors; set-shifting vulnerability; CEN degradation"),
  ],
  "sn_tests":[
    ("Error Monitoring","Read sentence with 2 errors. Note whether patient self-corrects.","Spontaneously catches ≥1 error","Misses both; unconcerned; anosognosia pattern; AD-specific SN failure","Catches errors but excessively distressed; MCI hypervigilance to mistakes"),
    ("Anosognosia Screen","Ask: 'Do you notice any memory changes?' Then ask family/caregiver same question. Note discrepancy.","Acknowledges memory changes appropriately; insight preserved","Denies memory problems despite clear deficits; anosognosia; AD SN failure","Over-reports deficits beyond what's observed; MCI anxiety amplification"),
    ("Threat Appraisal","3 ambiguous social scenarios. Rate threat 0–10.","Appropriate ratings 2–5","Inappropriately low ratings; poor threat detection; social unawareness","Inappropriately high ratings; catastrophic interpretation; behavioral AD"),
    ("Sensory Switching","Alternate visual tracking and auditory counting. Measure switch accuracy.","Switch delay <2 sec; smooth transitions","Slow switching; loses task set; executive perseveration","Cannot disengage from first task; perseveration; AD frontal release"),
    ("Orienting Response","Unexpected auditory tone mid-task. Observe behavioral interruption.","Brief interruption; reorientation within 3 sec","No orienting; flat response; disengagement; advanced AD","Exaggerated startle; prolonged recovery; behavioral BPSD pattern"),
    ("Pain Sensitivity","Light fingernail pressure 3 sec. Rate 0–10.","Rating 3–6; bilateral symmetry","Low ratings; poor somatic awareness; advanced AD","Normal rating but forgets pain stimulus immediately; episodic binding failure"),
    ("Emotional Salience Rating","Rate 6 emotional and 6 neutral images for personal salience 0–10.","Balanced ratings; appropriate differentiation","Flat ratings; emotional blunting; advanced AD","Exaggerated emotional responses; BPSD pattern; behavioral AD variant"),
  ],
  "smn_tests":[
    ("Finger Tapping Rate","Index finger tapping 10 sec each hand. Count and note rhythm.","≥45 taps/10 sec; regular rhythm","<35 taps; motor slowing; moderate-severe AD","Normal rate but irregular rhythm; motor impersistence"),
    ("Grip Assessment","Bilateral squeeze of examiner's fingers.","Firm symmetric grip; smooth release","Weak bilateral grip; reduced motor drive","Normal grip but primitive grasp reflex present; frontal release sign"),
    ("Fine Motor Sequencing","Thumb-to-finger opposition 10 sec.","≥4 complete sequences","≤2 sequences; forgets sequence mid-task; working memory-motor failure","3–4 sequences but perseverative; repeats first finger pair repeatedly"),
    ("Tandem Gait","Heel-to-toe walk 10 steps.","≤1 step off; consistent pace","Multiple failures; wide-based; gait apraxia; moderate-severe AD","Slow but intact; mild gait hesitancy; MCI motor vulnerability"),
    ("Romberg Test","Feet together eyes closed 30 sec.","Minimal sway; stable 30 sec","Excessive sway; falls; requires support; motor system involvement","Mild sway; compensatory widening of stance; balance vulnerability"),
    ("Rapid Alternating Movements","Pronate/supinate 10 sec. Note rhythm.","Regular rhythm; ≥12 cycles","Dysdiadochokinesia; slow; irregular; extrapyramidal involvement","Slightly irregular; mildly slowed; early motor system involvement"),
    ("Writing Sample","Write name and address; draw circles.","Legible; organized; consistent","Micrographia; cannot write name; severe AD","Spelling errors; writing tremor; mild motor degradation; MCI marker"),
    ("Gait Speed","Walk 3 meters at comfortable pace. Time it. (Slow gait predicts AD progression.)","<5 sec for 3 m (>0.6 m/sec)",">8 sec; slow gait (<0.6 m/sec); motor involvement","5–7 sec; borderline gait speed; MCI-motor overlap"),
  ],
  "limbic_tests":[
    ("Emotional Face Recognition","Identify 6 emotions from facial photos. Score /6.","≥5/6 correct","<4/6; emotional agnosia; advancing AD","4–5/6 with hesitation; relies on prominent features; MCI affect recognition"),
    ("Affect Labeling","3 emotional scenarios. Name own response.","Accurate, differentiated labels","Flat labels; 'nothing'; poor affect awareness; emotional blunting","Appropriate labels but disconnected from personal relevance; semantic-emotional dissociation"),
    ("Emotional Memory","10 emotional and 10 neutral words. Free recall after 5-min distractor.","Emotional advantage: ≥8 emotional recalled vs ≥6 neutral","Flat recall; lost emotional memory enhancement; limbic-hippocampal failure","Emotional words recalled but patient cannot recall context of learning; binding failure"),
    ("Fear Response","3 threatening scenarios. Rate anxiety 0–10.","Anxiety 4–7; proportionate","Flat ≤2; blunted fear; emotional blunting; moderate AD","Extreme ≥8; behavioral AD; agitation; catastrophic emotional response"),
    ("Reward Sensitivity — Apathy","Rate interest in 5 previously enjoyed activities 0–10.","Average ≥6/10; motivation preserved","Average ≤3/10; apathy; loss of initiative; AD behavioral symptom","Average ≥7 but patient forgets having been asked; episodic memory failure"),
    ("Caregiver Emotional Burden Screen","Ask patient: 'How do you feel about needing help with daily tasks?' Note: awareness, emotional response, and coping.","Acknowledges need for help; appropriate adjustment affect","Denies need for help; anosognosia; no emotional response","Extreme distress; catastrophic emotional response to acknowledgment; MCI anxiety"),
    ("Affective Flexibility","Sad then happy memory recall. Rate mood before/after.","Mood shift ≥3 points; recovery within 3 min","Flat across both; emotional blunting; advanced AD","Stuck in sad mood; cannot access positive state; limbic rigidity"),
  ],
  "attn_tests":[
    ("Sustained Attention","Count backward from 50 for 60 sec. Note fadeouts.","≤2 errors; consistent pace","Frequent pauses; fails to sustain; moderate AD","Monotonous slow count; perseverative; effortful but sustained"),
    ("Selective Attention","Circle all A's in 20-item letter list.","<60 sec; ≤3 errors","Slow; misses target; omissions; attentional failure","Normal time but multiple false alarms; poor response inhibition"),
    ("Visual Cancellation","Cross out all circles on mixed-shape sheet. Systematic search?","Systematic row-by-row; <60 sec; ≤3 errors","Random or incomplete; gives up; severe attentional failure","Systematic but extreme slowing; compensatory strategy use; MCI"),
    ("Spatial Attention — Line Bisection","Bisect a 20cm horizontal line. Measure deviation from center.","Within ±1 cm of center; bilateral","Deviation >2 cm right (left hemisphere): contralateral attention failure","Deviation >2 cm left (right hemisphere): ipsilateral attention bias; AD parietal"),
    ("Alerting Response","Say 'ready' before stimulus. Assess preparatory readiness.","Appropriate alerting; focused","Absent alerting; no preparatory response; disengaged","Normal alerting but immediately forgets set; sustained alerting failure"),
    ("Orienting — Spatial Cue","Arrow cue before target location. Orienting benefit and disengagement.","Orienting benefit ≥100 ms","Absent orienting benefit; poor spatial attention; parietal failure","Cannot disengage from invalid cue; perseverative orienting; frontal-parietal failure"),
    ("Response Inhibition Cancel","Cancel circles on mixed-shape sheet.","<60 sec; ≤2 errors","Slow; multiple omissions; poor initiation","Multiple false cancellations; cannot inhibit non-target responses"),
    ("Attentional Fatigue","Letter cancellation 5 min; compare halves.","Drop <15%","Large drop >40%; rapid fatigue; AD attentional depletion","Consistent slowing throughout; no fatigue but no speed either; flat trajectory"),
  ],
  "phenotype_text":"Alzheimer's phenotypes: (1) AMNESTIC — DMN dominant HYPO (episodic memory); (2) EXECUTIVE — CEN HYPO + frontal variant; (3) POSTERIOR CORTICAL — Attention/SMN HYPO (visuospatial); (4) BEHAVIORAL — Limbic HYPER (agitation, apathy). MCI staging: isolated DMN HYPO = early; multi-network = moderate-advanced.",
  "fnon_strategy":"Primary: Facilitate DMN/hippocampal memory networks via temporal-parietal anodal stimulation. Facilitate CEN (left DLPFC) for executive function. tDCS + cognitive training. Consider high-frequency rTMS over left DLPFC for MCI stage.",
}


# ── 5. STROKE ─────────────────────────────────────────────────────────────────
STROKE = {
  "full":"Post-Stroke Rehabilitation","short":"Stroke","slug":"stroke_rehab",
  "adapted_note":"Includes lesion-dependent network testing, lateralized motor/language assessment, and post-stroke phenotype mapping",
  "dmn_rel":"Post-stroke DMN disruption depends on lesion location. Cortical strokes cause focal DMN HYPO with self-awareness deficits, neglect, or anosognosia. Subcortical white matter strokes cause diffuse DMN HYPO via disconnection, manifesting as flat affect and reduced self-referential processing.",
  "cen_rel":"CEN is frequently affected post-stroke. Left hemisphere strokes cause CEN HYPO with aphasia and executive dysfunction; right hemisphere causes CEN HYPO with visuospatial and attention deficits. Frontal and parietal lesions directly disrupt CEN hubs.",
  "sn_rel":"Post-stroke SN dysfunction reflects insular and ACC involvement. Insular strokes cause interoceptive failure and autonomic dysregulation. Bilateral SN HYPO impairs error monitoring and reduces rehabilitation motivation. HYPER emerges as agitation in behavioral post-stroke syndrome.",
  "smn_rel":"SMN is typically HYPO contralateral to lesion: spasticity (motor HYPER upper motor neuron), weakness (motor HYPO), and coordination deficits reflect corticospinal and cerebellar pathway disruption. Motor rehabilitation directly targets SMN neuroplasticity.",
  "limbic_rel":"Post-stroke limbic dysfunction includes depression (30-40% post-stroke), emotional lability, and reduced motivation. Left anterior lesions increase depression risk. Right hemisphere strokes impair emotional recognition and regulation. Limbic HYPO drives apathy.",
  "attn_rel":"Hemispatial neglect (right parietal lesion = left neglect) represents DAN HYPO. Post-stroke attention deficits span all domains. Early attention normalization is a strong predictor of rehabilitation outcomes.",
  "dmn_tests":[
    ("Anosognosia Screen","Ask: 'Do you notice any weakness or difficulty since your stroke?' Compare to objective exam findings.","Full awareness; appropriate insight and adjustment","Denies deficits despite clear objective findings; anosognosia; right hemisphere stroke marker","Over-reports deficits beyond objective findings; catastrophic awareness; depression"),
    ("Clock Drawing Test","Draw clock showing 10:10. Score /4. Observe hemispatial neglect pattern.","4/4; symmetric layout; balanced number placement","<2/4; all numbers on one side; hemispatial neglect; right parietal stroke","3/4; mild spatial error; mild planning deficit; anterior stroke"),
    ("Autobiographical Memory","Ask 3 personal memories. Note: intact remote vs impaired recent; peri-stroke amnesia is expected.","Intact pre-stroke memories; peri-stroke gap is normal; integrating post-stroke","Remote and recent memory loss; global amnesia; bilateral hippocampal involvement","Confabulates peri-stroke events; fills memory gaps with invented content"),
    ("Self-Awareness Rating","Ask to rate pre-stroke vs current abilities in 3 domains (walking, reading, memory).","Accurate relative self-assessment; appropriate grieving of loss","No recognition of change; anosognosia; right hemisphere stroke","Catastrophically overestimates loss; severe post-stroke depression"),
    ("Orientation — Time and Place","Date, year, season, building, city, state. Score /6.","6/6 correct","<5/6; post-stroke confusion; acute phase disorientation","Correct but slow; post-stroke processing speed reduction"),
    ("Default Narrative","Ask patient to describe their stroke experience and goals for recovery.","Coherent narrative; realistic goals; appropriate emotional content","Sparse; flat affect; post-stroke apathy; poor self-reflection","Highly catastrophic; no future goals; post-stroke depression"),
    ("Verbal Fluency","Animals in 1 min. Note word-finding pauses or circumlocutions.","≥12 animals; fluent retrieval; no circumlocutions","<8; word retrieval failure; left hemisphere stroke; anomia","12-14 with circumlocutions; anomia; left posterior stroke"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Maximum span achieved.","Span ≥5","Span ≤3; working memory failure; frontal or parietal stroke","Normal span but very slow; processing speed reduction; white matter"),
    ("Trail Making B","Alternating number-letter connections. Time and errors.","<90 sec; ≤2 errors","Unable to alternate; loses rule; frontal lobe stroke","Slow >150 sec; self-corrected errors; executive slowing post-stroke"),
    ("Verbal Fluency FAS","Letter fluency 60 sec per letter. Count unique words.","≥9 words/letter","<5; left frontal stroke; Broca area involvement","7-9 with word-finding pauses; left perisylvian stroke"),
    ("Stroop Interference","Color-word cards. Name ink color.","<45 sec; ≤3 errors","Cannot perform; language comprehension failure; left hemisphere stroke","Slow; multiple self-corrections; post-stroke executive slowing"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct; <30 sec","≤2/5; post-stroke arithmetic failure; left parietal involvement","3-4/5 with self-corrections; executive slowing post-stroke"),
    ("Go/No-Go Task","Tap for all except X. 20 letters at 1/sec.","≤2 commission; ≤2 omission","High omissions; motor aphasia or frontal HYPO","High commission errors; frontal disinhibition post-stroke"),
    ("Category Switching","Alternating fruits and vegetables 60 sec.","≥10 alternations; ≤2 intrusions","<5; perseveration; frontal lobe stroke","5-8 alternations; mildly impaired; post-stroke CEN slowing"),
  ],
  "sn_tests":[
    ("Error Monitoring","Read sentence with 2 errors. Note spontaneous self-correction.","Catches ≥1 error; self-corrects","Misses both; anosognosia; no self-monitoring; right stroke","Catches errors; becomes significantly distressed; depression amplifier"),
    ("Pain-Sensory Awareness — Bilateral","Light fingernail pressure bilaterally. Compare sides. Note hemisensory loss.","Symmetric ratings 3-6 bilaterally","Absent or markedly reduced one side; sensory stroke; thalamic or cortical","Allodynia; post-stroke central pain; hyperalgesia on affected side"),
    ("Interoception — Heartbeat Counting","Count heartbeats 25 sec without touch.","Accuracy >0.75","Accuracy <0.50; insular stroke; poor autonomic awareness","High accuracy but hypervigilant; post-stroke health anxiety"),
    ("Threat Appraisal","3 ambiguous social scenarios. Rate threat 0-10.","Appropriate 2-5 ratings","Flat ratings; poor threat detection; post-stroke behavioral HYPO","All rated ≥8; catastrophic appraisal; post-stroke anxiety"),
    ("Sensory Switching","Alternate visual tracking and auditory counting. Switch latency.","<2 sec switch; smooth transitions","Slow >5 sec; post-stroke switching failure; executive-attentional","Impulsive switching; frontal disinhibition post-stroke"),
    ("Surprise Orienting Response","Unexpected auditory tone mid-task.","Brief interruption; reorientation <3 sec","Absent orienting; flat affect post-stroke; or advanced impairment","Exaggerated startle; autonomic arousal; post-stroke hypervigilance"),
    ("Emotional Salience Rating","Rate 6 emotional and 6 neutral images 0-10.","Balanced; appropriate differentiation","Flat; emotional blunting; post-stroke apathy","Exaggerated; emotional lability; post-stroke emotional dysregulation"),
  ],
  "smn_tests":[
    ("Grip Strength — Bilateral Comparison","Bilateral grip: squeeze examiner's fingers. Compare paretic vs intact hand.","Symmetric; firm bilateral grip","Marked weakness or paresis on affected side; corticospinal stroke marker","Spastic over-grip; cannot release; upper motor neuron spasticity"),
    ("Finger Tapping — Bilateral","Index tapping 10 sec each hand. Compare paretic vs intact side.","≥45 taps/10 sec; ≤10% asymmetry","<25 taps on paretic side; severe motor deficit; corticospinal involvement","Normal rate but irregular rhythm; mirror movements; motor cortex reorganization"),
    ("Fine Motor Sequencing — Bilateral","Thumb-to-finger opposition each hand separately 10 sec.","≥4 sequences each hand","≤1 sequence paretic side; finger individualization failure; M1 stroke","2-3 sequences paretic; recovering corticospinal tract"),
    ("Tandem Gait","Heel-to-toe walk 10 steps. Note lateralized failures.","≤1 step off; symmetric gait","Cannot perform; hemiparetic gait; requires support","Compensatory wide-base; slow; safety assessment needed"),
    ("Romberg Test","Feet together eyes closed 30 sec. Note one-sided sway.","Stable; symmetric sway","Falls or requires support; balance system involvement","Asymmetric sway toward paretic side; proprioceptive loss"),
    ("Rapid Alternating Movements — Bilateral","Pronate/supinate each forearm. Compare sides.","Regular ≥12 cycles each side","Dysdiadochokinesia on affected side; cerebellar or motor involvement","Slow irregular bilaterally; cerebellum or basal ganglia involvement"),
    ("Writing with Affected Hand","Draw circles with paretic hand. Note size, pressure, motor quality.","Legible attempt; improving with trials","Cannot hold pen; severe motor deficit; acute stroke","Abnormal pen grip; compensatory pressure; early motor recovery attempt"),
    ("Gait Speed — 3-Meter Walk","Walk 3 meters at comfortable pace. Time it.","<5 sec (>0.6 m/sec); community walking","≥10 sec (<0.3 m/sec); severely limited; assistance required","5-8 sec; limited community mobility; rehabilitation target"),
  ],
  "limbic_tests":[
    ("Post-Stroke Depression Screen","Ask: 'In the past 2 weeks, have you felt down or hopeless, or lost interest in things?' Rate 0-6.","0-2: Appropriate mood given situation","0: Flat; no emotional response; apathy; left anterior stroke","4-6: Moderate-severe depression; limbic HYPER; left anterior stroke risk"),
    ("Emotional Face Recognition","Identify 6 emotions from photos. Score /6.","≥5/6 correct","<4/6; emotional agnosia; right hemisphere stroke","4-5/6 with hesitation; bilateral or right posterior involvement"),
    ("Emotional Lability Screen","Ask: 'Do you cry or laugh more easily than before your stroke?' Rate 0-10.","0-2: No lability; emotional control maintained","0: Complete emotional flat affect; apathy; left anterior stroke","≥7: Pseudobulbar affect; involuntary laughing/crying; bilateral corticobulbar tract"),
    ("Affect Labeling","3 emotional scenarios. Name emotional response.","Appropriate proportionate labels","'Nothing'; apathy; flat affect; limbic HYPO","Extreme emotional labels; post-stroke emotional dysregulation"),
    ("Motivation for Rehabilitation","Rate motivation to attend therapy and complete exercises 0-10.","≥7; engaged in rehabilitation goals","≤3; apathy; post-stroke limbic failure; poor prognosis","High motivation but catastrophic fear of failure; performance anxiety"),
    ("Social Withdrawal Screen","Rate comfort in 5 social scenarios since stroke 0-10.","≥5/10 for ≥4 scenarios; social engagement maintained","≤3 all scenarios; severe withdrawal; post-stroke depression","≤3 in public only; social anxiety; self-consciousness about visible deficits"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no affect change; post-stroke apathy","Stuck in sad mood; cannot shift; post-stroke depression"),
  ],
  "attn_tests":[
    ("Sustained Attention","Count backward from 50 for 60 sec. Note fadeouts.","≤2 errors; consistent","Frequent pauses; severe post-stroke attention deficit","Very slow; processing speed reduction; white matter involvement"),
    ("Hemispatial Neglect Screen — Line Bisection","Mark midpoint of a 20 cm horizontal line. Measure deviation from true center.","Within ±1 cm; no systematic bias","Deviation >2 cm right (left neglect); right parietal stroke; DAN HYPO","Deviation ±1.5 cm; mild bias; recovering attention system"),
    ("Letter Cancellation — Neglect Screen","Cross out all target letters on a sheet. Count omissions per visual field.","≤2 omissions per side; symmetric","≥8 omissions one side; hemispatial neglect; right parietal stroke","3-5 omissions one side; mild attentional bias; early recovery"),
    ("Divided Attention","Count backward by 2s while naming pointed objects.","≤2 errors each task","Abandons one task; post-stroke attention narrowing","Both performed but at reduced speed; dual-task vulnerability"),
    ("Visual Search","Find star target among 20 distractors.","<20 sec; systematic search","Searches one visual field only; neglect; right parietal stroke","Slow >40 sec; systematic but effortful; attentional recovery"),
    ("Alerting Response","Say 'ready' before stimulus. Assess readiness.","Appropriate alerting; focused","Blunted alerting; poor readiness; post-stroke alerting HYPO","Over-alerting; vigilant; post-stroke anxiety pattern"),
    ("Spatial Orienting — Cue Validity","Arrow cue before target. Orienting benefit and disengagement speed.","Orienting benefit ≥100 ms; smooth disengagement","Absent orienting; parietal failure; neglect side","Deviation toward lesion; neglect on disengagement; parietal stroke"),
    ("Attentional Fatigue","Letter cancellation 5 min; compare first vs last 2.5 min.","Drop <15%","Drop >40%; rapid post-stroke attentional fatigue","Consistent slowing; flat trajectory; processing speed limitation"),
  ],
  "phenotype_text":"Post-Stroke phenotypes: (1) MOTOR-DOMINANT — SMN HYPO contralateral; (2) COGNITIVE — CEN HYPO + Attention HYPO; (3) AFFECTIVE — Limbic HYPER (post-stroke depression); (4) NEGLECT — Attention HYPO right parietal; (5) APHASIC — CEN HYPO left perisylvian. Map to lesion location for FNON target selection.",
  "fnon_strategy":"Primary: Facilitate affected SMN via anodal stimulation over ipsilesional motor cortex. Inhibit contralesional motor cortex to reduce maladaptive compensation (for severe chronic cases). Facilitate left DLPFC for post-stroke depression. Address CEN for cognitive subtypes.",
}

# ── 6. TBI ────────────────────────────────────────────────────────────────────
TBI = {
  "full":"Traumatic Brain Injury (TBI)","short":"TBI","slug":"tbi",
  "adapted_note":"Includes post-concussion symptom markers, white matter disconnection tests, and TBI severity phenotype indicators",
  "dmn_rel":"TBI causes DMN HYPO through white matter disconnection — axonal shearing disrupts PCC-mPFC connectivity underlying self-referential processing. Chronic TBI shows DMN failure with poor autobiographical integration, emotional numbing, and identity disruption.",
  "cen_rel":"CEN HYPOACTIVITY is the hallmark of TBI cognitive impairment: frontal axonal damage, working memory failure, and processing speed reduction. Executive dysfunction reflects CEN white matter disconnection more than focal lesion in most mTBI cases.",
  "sn_rel":"TBI causes SN HYPERACTIVITY: central sensitization, heightened pain and sensory sensitivity (light, sound). Insular involvement contributes to autonomic dysregulation and post-traumatic headache. SN HYPER drives post-concussion syndrome.",
  "smn_rel":"SMN dysfunction in TBI depends on severity: motor slowing, balance deficits, and coordination impairment are common. Vestibular involvement causes gait and balance HYPO. mTBI may show subtle SMN HYPO not detectable on standard exams.",
  "limbic_rel":"Limbic HYPERACTIVITY drives post-traumatic emotional dysregulation, irritability, and anxiety. Amygdala hyperreactivity from direct injury or CEN deafferentation causes emotional volatility. Depression and PTSD comorbidity are common limbic-TBI presentations.",
  "attn_rel":"Attention Networks show HYPO across all domains post-TBI: processing speed, sustained attention, and divided attention. White matter integrity loss in fronto-parietal networks causes attentional failure. Attentional fatigue is a core TBI complaint.",
  "dmn_tests":[
    ("Post-Traumatic Narrative","Ask patient to describe the TBI event and recovery. Note coherence, gaps, temporal sequence.","Coherent narrative; appropriate emotional processing; peri-injury gap is expected","Sparse or absent narrative; poor self-awareness; anosognosia; frontal TBI","Highly distressing; intrusive re-experiencing; PTSD-TBI overlap"),
    ("Autobiographical Memory","3 personal memories (pre-injury, recent, positive). Rate specificity and accuracy.","Pre-injury memories intact; peri-injury gap normal; post-injury integrating","Retrograde amnesia extends before injury; severe DAI pattern","Confabulates injury events; fills memory gaps with plausible but inaccurate content"),
    ("Clock Drawing Test","Draw clock showing 10:10. Score /4.","4/4; organized layout","<2/4; executive-spatial failure; frontal or parietal TBI","3/4; mild planning error; executive vulnerability"),
    ("Self-Reference Memory","10 adjectives with self-reference frame. Recall after 2-min distractor.","≥6/10 recalled; balanced content","<4/10; poor self-schema; TBI identity disruption","Normal recall but dominated by injury-related self-concepts"),
    ("Mind Wandering Probe","2-min reading with 4 probes. Note off-task content.","≤1/4 off-task","0/4; over-focused; compensatory rigidity","3-4/4; poor sustained attention; TBI DMN-CEN dyscoordination"),
    ("Prospective Memory","5-min reminder task. Note if recalled independently.","Independent recall within 5±1 min","Completely forgets; TBI prospective memory failure","Asks repeatedly before time; compensatory over-monitoring; anxiety"),
    ("Orientation","Date, year, season, place /6.","6/6 correct","<5/6; post-TBI disorientation; acute confusion","6/6 but slow; processing speed reduction"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Maximum span achieved.","Span ≥5","Span ≤3; TBI working memory failure","Span 4; borderline; processing speed interferes"),
    ("Trail Making B","Alternating connections. Time and errors.","<90 sec; ≤2 errors","Slow >150 sec or multiple errors; frontal white matter TBI","100-150 sec; mild-moderate slowing; post-concussion"),
    ("Processing Speed — Symbol Copying","Copy 3 symbols under 3 number pairs; complete 10 pairs in 60 sec. Count correct.","≥8/10 in 60 sec","<4/10; severe processing speed impairment; TBI core CEN deficit","5-7/10; mild-moderate speed reduction; post-concussion syndrome"),
    ("Verbal Fluency FAS","Letter fluency 60 sec per letter.","≥10 words/letter","<6; TBI frontal-executive failure","6-9 words; mild reduction; processing speed contribution"),
    ("Stroop Interference","Name ink color. Time and errors.","<45 sec; ≤3 errors","Slow >90 sec; severe interference; frontal TBI","45-75 sec; moderate; post-concussion CEN slowing"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct","≤2/5; working memory failure; TBI","3-4/5; mild errors; processing speed and attention interaction"),
    ("Go/No-Go Task","Tap for all except X. 20 letters.","≤2 commission; ≤2 omission","High commission errors; frontal disinhibition; TBI","2-4 commission errors; borderline inhibitory control; mTBI"),
  ],
  "sn_tests":[
    ("Post-Concussion Sensory Screen","Rate current headache, photophobia, and phonophobia 0-10 each.","0-2 on all; no sensitization","0 on all; complete sensory blunting; severe TBI or dissociation","≥6 on ≥2 measures; post-traumatic sensitization; SN HYPER"),
    ("Light Sensitivity Test","Observe in normally lit room; ask about discomfort with ceiling lights.","No discomfort with normal lighting","No sensitivity; sensory blunting; severe TBI","Significant discomfort; light avoidance; post-concussion SN hyperactivation"),
    ("Interoception — Heartbeat Counting","Count heartbeats 25 sec without touch.","Accuracy >0.75","Accuracy <0.50; autonomic dysregulation; brainstem/insular TBI","High accuracy; hypervigilant to somatic signals; post-TBI health anxiety"),
    ("Error Monitoring","Sentence with 2 errors. Spontaneous correction.","Catches ≥1; calm correction","Misses both; anosognosia; frontal TBI","Extreme distress on errors; perfectionism; post-TBI emotional dysregulation"),
    ("Sensory Switching","Alternate visual and auditory tasks. Note if switching triggers nausea/pain.","<2 sec; smooth; no symptoms","Slow; poor switching; frontal disconnection","Fast but triggers vestibular symptoms; sensory overload; post-concussion TBI"),
    ("Pain Sensitivity — Bilateral","Fingernail pressure 3 sec. Rate and compare bilaterally.","3-6; bilateral symmetry","Low or absent; sensory loss; focal TBI","High bilateral ≥8; central sensitization; diffuse TBI"),
    ("Orienting Response — Exaggerated Startle","Unexpected tone mid-task. Measure startle magnitude and recovery.","Brief interruption; reorientation <3 sec","Absent orienting; severe TBI blunting","Exaggerated startle; prolonged recovery >15 sec; PTSD-TBI overlap"),
  ],
  "smn_tests":[
    ("Finger Tapping Rate","Index tapping 10 sec each hand. Count and note rhythm.","≥45 taps; regular","<35 taps; motor slowing; DAI or focal motor TBI","35-45 taps; mild reduction; processing speed contribution"),
    ("Grip Assessment","Bilateral grip. Compare sides.","Firm symmetric grip","Asymmetric weakness; focal motor pathway TBI","Normal but slow force generation; psychomotor slowing"),
    ("Fine Motor Sequencing","Thumb-to-finger 10 sec.","≥4 sequences","≤2; motor-sequencing failure; cerebellar or frontal TBI","2-3 sequences; mild-moderate fine motor impairment; mTBI"),
    ("Balance — Modified Romberg","Tandem stance eyes closed 20 sec. Count sway or steps off.","Stable 20 sec; ≤1 step adjustment","Cannot maintain; falls; vestibular or cerebellar TBI","2-3 adjustments; mild balance impairment; post-concussion vestibular"),
    ("Gait Pattern","Walk 5 m at normal pace. Observe stride length, arm swing, stability.","Symmetric; normal stride; bilateral arm swing","Hemiparetic gait; step asymmetry; focal motor TBI","Mildly reduced stride; cautious gait; post-concussion vestibular"),
    ("Rapid Alternating Movements","Pronate/supinate 10 sec. Note rhythm.","Regular; ≥14 cycles","Dysrhythmic; cerebellar TBI","Slightly irregular; mild rhythm disruption; cerebellar vulnerability"),
    ("Writing Sample","Write name/address; draw circles.","Legible; consistent; smooth","Cannot hold pen; severe motor TBI","Slight irregularity; fine motor vulnerability post-TBI"),
    ("Reaction Time — Simple","Hand lift when pen drops. 3 trials averaged.","<250 ms","≥500 ms; severe processing speed impairment; DAI marker","250-400 ms; mild-moderate slowing; post-concussion processing speed"),
  ],
  "limbic_tests":[
    ("Irritability Screen","Ask: 'Since your injury, do you get angry more easily?' Rate 0-10.","0-3; appropriate emotion regulation","0; emotional blunting; limbic HYPO; severe TBI","≥8; post-traumatic irritability; limbic HYPER; amygdala disinhibition"),
    ("Emotional Face Recognition","6 emotions from photos. Score /6.","≥5/6 correct","<4/6; affect agnosia; right hemisphere or amygdala TBI","4-5/6; mild; frontal-limbic disconnection"),
    ("Emotional Memory","10 emotional and 10 neutral words. Free recall after 5-min distractor.","Emotional advantage; ≥8 emotional recalled","Flat recall; emotional memory loss; limbic-hippocampal TBI","Negative emotional words preferentially recalled; PTSD-TBI overlap"),
    ("Affect Labeling","3 emotional scenarios. Name response.","Accurate differentiated labels","'Nothing'; alexithymia; limbic blunting","Intense labels; explosive emotional responses; dysregulation"),
    ("Reward Sensitivity","Rate interest in 5 activities 0-10.","Average ≥6","Average ≤3; apathy; frontal-limbic TBI","Average ≥7 but explosive short engagement; impulse dysregulation"),
    ("Fear Response","3 threatening scenarios. Rate anxiety 0-10.","Anxiety 4-6; proportionate","Flat ≤2; emotional blunting; frontal-limbic TBI","≥8; catastrophic; PTSD-TBI overlap; amygdala hyperreactivity"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no shift; emotional blunting; limbic HYPO","Stuck in negative mood; poor positive affect access; post-TBI depression"),
  ],
  "attn_tests":[
    ("Sustained Attention — Fatigue Probe","Count backward from 50 for 60 sec. Note rapid fatigue onset.","≤2 errors; consistent pace","Stops within 30 sec; severe attentional failure","Completes but dramatically decelerates; attentional fatigue; post-concussion hallmark"),
    ("Selective Attention","Circle all A's in 20-item letter list. Time and errors.","<60 sec; ≤3 errors","Slow; multiple omissions; severe post-TBI attention failure","Completes; 45-75 sec; processing speed bottleneck"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one task; TBI attention resource depletion","Performs both with 3-5 errors each; dual-task vulnerability; TBI hallmark"),
    ("Visual Search","Find star among 20 distractors.","<20 sec; systematic","Fails to find target; disorganized; severe TBI","20-40 sec; systematic but slow; processing speed and attention interaction"),
    ("Alerting Response","'Ready' before stimulus. Preparatory state.","Appropriate alerting; focused","Blunted; not ready; post-TBI alerting failure","Hyper-alerting; hypervigilant; PTSD-TBI overlap"),
    ("Spatial Orienting","Arrow cue before target. Orienting benefit and disengagement.","Orienting benefit ≥100 ms","Absent orienting; spatial attention failure; parietal TBI","Normal orienting; cannot disengage; frontal perseveration TBI"),
    ("Response Inhibition Cancel","Cancel circles on mixed-shape sheet.","<60 sec; ≤2 errors","Slow; >120 sec; omissions","Fast with false alarms; impulsive; frontal inhibitory failure TBI"),
    ("Attentional Fatigue Index","Letter cancellation 5 min; compare first vs last 2.5 min.","Drop <15%","Drop >50%; rapid attentional exhaustion; TBI hallmark","20-35% drop; mild-moderate fatigue; post-concussion syndrome marker"),
  ],
  "phenotype_text":"TBI phenotypes: (1) COGNITIVE-FATIGUE — CEN HYPO + Attention HYPO (most common mTBI); (2) AFFECTIVE — Limbic HYPER (emotional TBI); (3) SENSORY-PAIN — SN HYPER (post-concussion sensitization); (4) MOTOR — SMN HYPO (moderate-severe TBI); (5) PTSD-TBI — Limbic HYPER + SN HYPER. Severity stage guides FNON selection.",
  "fnon_strategy":"Primary: Facilitate CEN (left DLPFC anodal) for executive dysfunction and processing speed. Address SN hyperactivation (cathodal right prefrontal). For PTSD-TBI overlap: target limbic regulation. Avoid stimulation during acute phase (<3 months without clearance).",
}

# ── 7. CHRONIC PAIN / FIBROMYALGIA ────────────────────────────────────────────
CHRONIC_PAIN = {
  "full":"Chronic Pain / Fibromyalgia","short":"Chronic Pain","slug":"chronic_pain",
  "adapted_note":"Includes central sensitization markers, pain catastrophizing tests, and Chronic Pain phenotype indicators",
  "dmn_rel":"Chronic Pain drives DMN HYPERACTIVITY through pain rumination and catastrophizing. The pain-default mode loop maintains suffering at rest. Pain-related DMN regions (mPFC, PCC) show increased connectivity, sustaining pain chronification beyond peripheral nociception.",
  "cen_rel":"CEN HYPOACTIVITY in Chronic Pain reflects attentional resource depletion by pain. DLPFC fails to apply top-down analgesia, reducing endogenous pain inhibition. Working memory and executive function are impaired by constant pain-related cognitive load.",
  "sn_rel":"Central sensitization is characterized by SN HYPERACTIVITY: the ACC and anterior insula become hyper-responsive to all sensory inputs, creating widespread allodynia, hyperalgesia, and emotional hypersensitivity. This is the neurobiological hallmark of fibromyalgia.",
  "smn_rel":"Chronic Pain causes SMN HYPOACTIVITY through pain-inhibited movement, guarding, and disuse. The fear-avoidance model drives progressive SMN under-engagement. Movement-induced pain amplification in fibromyalgia reflects central sensitization of SMN output.",
  "limbic_rel":"Limbic HYPERACTIVITY drives fear-avoidance, pain catastrophizing, and emotional amplification of pain signals. Amygdala-ACC hyperconnectivity creates a pain-emotion feedback loop. Depression and anxiety are common comorbidities driven by limbic chronification.",
  "attn_rel":"Attention Networks show HYPO for non-pain content (pain captures attentional resources) alongside HYPER-vigilance for body sensations and pain cues. DAN HYPO reflects reduced top-down attentional control; VAN HYPER reflects excessive pain-signal capture.",
  "dmn_tests":[
    ("Pain Rumination Probe","Ask: 'When pain is worst, do thoughts about it repeat involuntarily?' Rate duration and controllability 0-10.","Pain thoughts controllable; 0-4/10; brief episodes","Minimal rumination despite high pain; dissociation from pain","Uncontrollable pain thoughts ≥8/10 for >2 hours; chronic pain DMN hallmark"),
    ("Pain Autobiography","Ask patient to narrate pain history: onset, evolution, current impact.","Coherent; some positive reframing; functional goals mentioned","Sparse; denies pain impact; stoic suppression","Entirely suffering-focused; hopeless; no positive elements; catastrophizing"),
    ("Pain Identity Integration","Ask: 'Would you say chronic pain is part of who you are?' Rate 0-10.","0-4; pain acknowledged but not identity-defining; adaptive coping","0; denies pain significance; avoidant coping","≥8; pain as complete identity; loss of non-pain self-concept"),
    ("Mind Wandering — Pain Content","2-min reading with 4 probes. Note: when off-task, is content pain-related?","≤1/4 off-task; neutral mind-wandering","0/4; hypervigilant task focus; suppression coping","3-4/4 off-task; pain-specific thought content; pain-DMN capture"),
    ("Pain Catastrophizing Screen","Ask 3 questions: 'Will pain never get better? Do you feel helpless? Do you magnify pain?' Rate 0-10 each.","All ≤3; realistic pain appraisal; self-efficacy maintained","0-1 across all; denial; minimization; avoidant","≥7 on ≥2 items; high catastrophizing; pain DMN hyperactivation"),
    ("Prospective Memory","5-min reminder task during assessment.","Independent recall within 5±1 min","Forgets; cognitive load from pain impairs prospective memory","Repeatedly reminds before time; pain-driven hypermonitoring"),
    ("Default Narrative — Pain Impact","Ask: 'What activities have you given up because of pain?' Rate acceptance vs avoidance framing.","Acknowledges losses but maintains goals; ACT-compatible","Denies giving up activities; minimization; avoidant","Extensive list; nothing maintained; complete pain-driven life restriction"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Maximum span.","Span ≥5","Span ≤3; cognitive dysfunction from chronic pain","Span 4; borderline; pain-related interference"),
    ("Trail Making B","Alternating connections. Time and errors.","<90 sec; ≤2 errors","Slow >150 sec; pain-related executive dysfunction","90-150 sec; moderate; pain-cognition interference"),
    ("Verbal Fluency FAS","60 sec per letter. Count unique words.","≥10 words/letter","<6; poor fluency; pain-related executive HYPO","6-9; mild reduction; attentional drain from pain"),
    ("Stroop — Pain Emotional Words","Include pain-related words (hurt, ache, sharp). Compare interference.","Uniform response time; ≤3 errors","No Stroop effect; dissociation from pain","Selectively slowed on pain words; pain-emotional Stroop effect"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct","≤2/5; cognitive load from pain","3-4/5; mild errors; pain-attention interaction"),
    ("Pain Self-Efficacy — Decision Task","Offer 2 task options. Note if pain appraisal impairs decision-making.","Decisive; considers pain but engages","Cannot decide; helplessness; pain-driven executive failure","Rapid pain-avoidant decisions; pain drives all choices"),
    ("Category Switching","Alternating fruits and vegetables 60 sec.","≥10 alternations; ≤2 intrusions","<5; poor flexibility; pain-cognitive interference","Normal count; slow pacing; effortful due to pain"),
  ],
  "sn_tests":[
    ("Widespread Pain Index — Bedside","Ask patient to indicate pain locations on body. Count regions affected.","0-3 pain regions; localized","0 regions; denial; avoidant coping; analgesia","≥7 widespread regions; fibromyalgia pattern; central sensitization; SN HYPER"),
    ("Allodynia Screen","Apply light cotton touch to forearm. Ask if painful.","Light touch not painful; no allodynia; normal threshold","No pain to any stimulus; sensory blunting; hypoalgesia; HYPO","Light touch reported as painful; allodynia; central sensitization; SN HYPER"),
    ("Interoception — Pain Hyperawareness","Count heartbeats 25 sec. Ask: 'Are you aware of pain during concentration?'","Accuracy 0.75; pain remains in background during concentration","Low accuracy; poor body awareness; dissociation","High accuracy AND heightened pain awareness during focus; hyperinteroception"),
    ("Error Monitoring","Sentence with 2 errors. Spontaneous correction.","Catches ≥1; calm","Misses both; cognitive load from pain","Catches errors; uses for self-criticism; pain-mood interaction"),
    ("Pressure Algometry — Tender Point","Apply thumb pressure to trapezius and calf bilaterally. Rate pain 0-10 each.","≤4 at all sites; no widespread tenderness","0 at all sites; hypoalgesia; analgesia; HYPO","≥6 at ≥3 sites; widespread tenderness; fibromyalgia SN HYPER"),
    ("Sensory Switching","Alternate visual and auditory tasks. Note if switching triggers pain increase.","<2 sec; smooth; no pain increase","Poor switching; cognitive inertia; pain-cognitive rigidity","Switching triggers pain or headache; sensory overload; fibromyalgia SN HYPER"),
    ("Orienting Response — Pain Startle","Unexpected auditory tone. Note if tone triggers pain response.","Brief interruption; reorientation <3 sec; no pain increase","Absent orienting; disengaged; hyporesponsive","Exaggerated startle; tone triggers pain; central sensitization of startle-pain circuit"),
  ],
  "smn_tests":[
    ("Finger Tapping — Pain During Movement","Index tapping 10 sec each hand. Note pain with movement.","≥45 taps; regular; no movement-induced pain","<35 taps; pain-inhibited movement; kinesiophobia","Normal rate; patient reports pain increase during tapping; movement hyperalgesia"),
    ("Grip Assessment — Pain with Gripping","Bilateral grip. Note pain with gripping.","Firm symmetric; no pain with gripping","Weak; pain avoidance; fear-avoidance model","Normal grip but pain ≥7 during gripping; movement-induced pain amplification"),
    ("Fine Motor Sequencing","Thumb-to-finger 10 sec. Note pain interference.","≥4 sequences; no pain increase","≤2; pain-inhibited fine motor; disuse pattern","3-4 sequences; pain awareness heightened during fine motor"),
    ("Functional Movement — Chair Rise","Rise from chair without arm support; turn; sit back. Time the sequence.","<12 sec; pain-free or minimal","Cannot perform; severe pain-inhibited movement; kinesiophobia","Performs but >20 sec; reports pain ≥7; fear-avoidance"),
    ("Romberg Test","Feet together eyes closed 30 sec.","Stable; minimal sway","Excessive sway; fibromyalgia proprioceptive HYPO","Stable but excessive tensing; guarding; protective motor pattern"),
    ("Widespread Tenderness Survey","Rate pressure at 4 tender points (bilateral trapezius, bilateral knee) 0-10.","0-2 at each site; no diffuse tenderness","0 pain at all sites; hypoalgesia; analgesia; HYPO","≥6 at ≥3 sites; fibromyalgia diagnosis support; SN HYPER"),
    ("Activity Pacing Assessment","Ask: 'Do you stop before pain starts (pacing) or push until pain stops you?'","Paces; activity matched to capacity; adaptive","Denies any pain limitation; avoidant minimization","Boom-bust pattern; over-does then crashes; maladaptive pacing; pain chronification"),
    ("Reaction Time","Hand lift when pen drops. 3 trials average.","<300 ms; pain does not delay initiation","≥500 ms; pain-inhibited motor readiness","<250 ms but pain increases after fast movement; reactive hyperalgesia"),
  ],
  "limbic_tests":[
    ("Pain Catastrophizing — Fear Screen","Ask: 'Are you afraid of what your pain means for your future?' Rate 0-10.","Fear ≤4; realistic; acceptance-based coping","0; denial; avoidant; catastrophizing not accessible","≥8; catastrophic pain fear; kinesiophobia; limbic HYPER core feature"),
    ("Emotional Face Recognition","6 emotions from photos. Score /6.","≥5/6 correct","<4/6; emotional blunting; pain-social disconnection","4-5/6; mild; pain-related emotional preoccupation"),
    ("Affect Labeling","3 emotional scenarios. Name response.","Accurate differentiated labels","Flat; 'nothing'; pain has numbed emotional response","All emotions flavored by pain-frustration; labels anxiety across all scenarios"),
    ("Emotional Memory — Pain Bias","10 words (5 pain/negative, 5 positive). Recall after 5-min distractor.","Balanced recall; slight emotional advantage","Flat recall; pain cognitive load impairs emotional processing","Pain/negative words preferentially recalled; positive suppressed; pain-mood bias"),
    ("Reward Sensitivity — Pain Barrier","Rate interest in 5 enjoyable activities 0-10. Note: avoidance from pain fear vs anhedonia.","Average ≥6; maintains pleasurable activities","Average ≤3; pain-driven anhedonia; avoidance from fear","High interest but pain prevents participation; preserved motivation with pain barrier"),
    ("Social Impact of Pain","Rate comfort in 5 social scenarios 0-10. Note pain as barrier.","≥5/10 for ≥4; social engagement maintained","≤3; social withdrawal; depression; limbic HYPO","≤3; avoidance from pain in social contexts; shame; isolation"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no shift; pain-induced emotional blunting","Sad memory activates pain; pain prevents positive affect; limbic-pain loop"),
  ],
  "attn_tests":[
    ("Pain-Attention Interference","Count backward from 50. At pain moment: how does performance change? Rate interference 0-10.","Maintains counting with mild pain; ≤2 errors","Doesn't stop despite severe pain; dissociation possible","Completely stops at pain flare; pain captures all attentional resources"),
    ("Selective Attention","Circle all A's in 20-item list.","<60 sec; ≤3 errors","Slow; multiple omissions; cognitive interference from pain","Normal but interspersed with pain rubs; pain breaks attentional continuity"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one; pain absorbs cognitive resources","Completes but reports pain increase with dual-task; cognitive-pain interaction"),
    ("Visual Search","Find star among 20 distractors.","<20 sec; systematic","Fails; gives up; pain-driven disengagement","20-40 sec; systematic but interrupted by pain attention"),
    ("Body Hypervigilance Screen","Observe attention to body during assessment. Count body-directed behaviors (touching painful area, wincing).","≤2 body-directed behaviors per 5 min; balanced attention","0-1; avoidant body attention; suppression coping","≥5 body-directed behaviors; body hypervigilance; VAN HYPER for pain cues"),
    ("Orienting Response — Pain Context","Unexpected tone. Note if it triggers body awareness increase.","Brief interruption; no pain increase; reorientation <3 sec","Absent orienting; sensory blunting; HYPO","Tone triggers body scan; pain awareness increase; SN-VAN hyperactivation"),
    ("Response Inhibition Cancel","Cancel circles on mixed-shape sheet.","<60 sec; ≤2 errors","Slow; pain-driven psychomotor slowing","Normal time but interrupted by pain breaks; continuity disruption"),
    ("Attentional Fatigue","Letter cancellation 5 min; compare halves.","Drop <15%","Drop >40%; pain-driven rapid attentional fatigue","20-35% drop; moderate; pain-attention resource depletion"),
  ],
  "phenotype_text":"Chronic Pain phenotypes: (1) CATASTROPHIZING — DMN HYPER + Limbic HYPER; (2) CENTRAL SENSITIZATION — SN HYPER dominant (fibromyalgia); (3) FEAR-AVOIDANCE — Limbic HYPER + SMN HYPO; (4) COGNITIVE-PAIN — CEN HYPO + Attention HYPO. Map to primary network for FNON protocol selection.",
  "fnon_strategy":"Primary: Inhibit SN hyperactivation (cathodal/low-frequency over ACC or contralateral S1/M1). Facilitate CEN (left DLPFC) for top-down pain modulation. Anodal M1 tDCS reduces central sensitization in fibromyalgia. Combine with graded activity and pain education.",
}

# ── 8. PTSD ───────────────────────────────────────────────────────────────────
PTSD = {
  "full":"PTSD","short":"PTSD","slug":"ptsd",
  "adapted_note":"Includes trauma-network dysregulation markers, hyperarousal/avoidance phenotyping, and PTSD network profile indicators",
  "dmn_rel":"PTSD creates complex DMN dysfunction: HYPO during acute re-experiencing (dissociation from self-referential processing) and HYPER during ruminative re-processing (intrusive thoughts, trauma-related mind-wandering). The DMN becomes 'trauma-colonized' losing access to non-trauma self-referential content.",
  "cen_rel":"CEN HYPOACTIVITY reflects prefrontal down-regulation by the hyperactive amygdala. vmPFC/DLPFC fails to extinguish fear memories and regulate amygdala responses. Poor working memory under stress and impaired extinction are CEN HYPO signatures of PTSD.",
  "sn_rel":"SN HYPERACTIVITY is the hallmark of PTSD: the amygdala-ACC-insula circuit is in a chronic threat-detection state, misidentifying safety as danger. Interoceptive hyperactivation (racing heart, physical tension) and emotional hyperreactivity reflect SN chronic over-engagement.",
  "smn_rel":"PTSD SMN shows HYPERACTIVATION in the sympathetic-motor axis: heightened startle, postural tension, freeze/flight motor readiness, and somatic tension. Chronic SMN hyperactivation contributes to physical symptoms (pain, fatigue, GI disturbance).",
  "limbic_rel":"Amygdala HYPERACTIVITY is the core limbic feature of PTSD: fear generalization, impaired extinction, and emotional hyperreactivity. Hippocampal HYPO (context-fear discrimination failure) allows fear responses to generalize beyond the trauma context.",
  "attn_rel":"Attention shows HYPER-vigilance (VAN HYPER) for threat cues alongside HYPO for non-threat sustained and divided attention. The DAN fails to direct attention away from trauma-related stimuli. Attentional narrowing in trauma-relevant contexts is a core PTSD feature.",
  "dmn_tests":[
    ("Trauma Narrative Coherence","Ask patient to briefly describe their understanding of the trauma. Rate coherence, fragmentation, integration.","Coherent, integrated account; temporal distance; some distress acceptable","Flat, emotionally absent account; numbing; avoidant PTSD phenotype","Highly fragmented; present-tense re-experiencing; unable to narrate as past event"),
    ("Autobiographical Memory — Trauma Context","Ask 3 personal memories: pre-trauma, neutral recent, positive. Rate retrieval vs trauma contamination.","Pre-trauma memories intact; non-trauma memories accessible","Unable to access non-trauma memories; trauma colonizes autobiographical recall","Intrusions of trauma content into unrelated memory retrieval; overgeneral memory"),
    ("Self-Reference — Post-Trauma Identity","Ask: 'How has your sense of who you are changed since the trauma?' Rate shift 0-10.","0-3; acknowledges change; integrated; growth possible","0; no self-change sense; denial; avoidant; poor self-awareness","≥8; profound identity destruction; 'I died'; depersonalization; complex PTSD"),
    ("Mind Wandering — Trauma Content","2-min reading with 4 probes. Note: is off-task content trauma-related?","≤1/4 off-task; non-trauma content","0/4; hypervigilant task focus; suppression; avoidant phenotype","3-4/4 off-task; trauma-specific intrusions; re-experiencing"),
    ("Safety Anchor Assessment","Ask: 'Where are you right now? What day is it? Are you safe?' Rate grounding 0-10.","Fully oriented; sense of safety; grounded in present","Partially disoriented during distress; dissociative state","Normal orientation but cannot feel safe; pervasive threat sense; hyperarousal PTSD"),
    ("Prospective Memory","5-min reminder task.","Independent recall within 5±1 min","Forgets; PTSD avoidance of future-orientation","Trauma intrusion interrupts task; re-experiencing mid-assessment"),
    ("Foreshortened Future Screen","Ask: 'What do you hope for in the next year?' Rate ability to imagine positive future.","Articulates specific positive goals; realistic","'I don't know'; avoidant; inability to imagine future","'Nothing good will happen'; foreshortened future; PTSD criterion E"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Note: does stress trigger intrusion?","Span ≥5; stable under low-stress assessment","Span ≤3; PTSD hyperarousal impairs working memory","Normal span but drops when stressed; context-dependent WM failure"),
    ("Trail Making B","Alternating connections. Time and errors.","<90 sec; ≤2 errors","Slow >150 sec; PTSD executive impairment","Normal but errors increase with duration; sustained CEN failure under pressure"),
    ("Verbal Fluency FAS","Letter fluency 60 sec per letter.","≥10 words/letter","<6; poor verbal output; CEN HYPO; avoidant processing","7-9 words; moderate; PTSD cognitive avoidance"),
    ("Stroop — Trauma Words","Include trauma-relevant words. Compare interference vs neutral.","Uniform time; ≤3 errors across all word types","No Stroop effect; emotional numbing; dissociative PTSD","Selectively slowed on trauma words; trauma-specific Stroop effect"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct","≤2/5; PTSD CEN HYPO","3-4/5; hyperarousal-cognition interference"),
    ("Go/No-Go Task","Tap for all except X. 20 letters.","≤2 commission; ≤2 omission","High omissions; avoidant PTSD; poor initiation","High commission errors; hypervigilant impulsive; hyperarousal PTSD"),
    ("Category Switching","Alternating fruits/vegetables 60 sec.","≥10 alternations; ≤2 intrusions","<5; executive rigidity; avoidant PTSD","Normal but trauma associations contaminate word choices; semantic intrusion"),
  ],
  "sn_tests":[
    ("Hyperarousal Screen","Ask: 'Are you constantly on alert for danger even when safe?' Rate 0-10.","0-3; appropriate alertness; safety sense maintained","0; emotional numbing; avoidant/dissociative PTSD","≥8; pervasive threat-readiness; cannot experience safety; SN HYPER hallmark"),
    ("Startle Response — Quantified","Drop pen unexpectedly mid-assessment. Rate startle magnitude 0-10, recovery time, and secondary emotional response.","Mild startle; recovery <5 sec; returns to calm","Absent startle; blunted response; dissociative PTSD; numbing","Extreme startle ≥8; prolonged recovery >15 sec; secondary anger/fear; PTSD criterion D"),
    ("Interoception — Heartbeat Counting","Count heartbeats 25 sec without touch.","Accuracy >0.70","Low accuracy <0.50; body disconnection; dissociative PTSD","High accuracy; hypervigilance to heartbeat; palpitation focus; somatic hyperawareness"),
    ("Threat Appraisal — Safety Testing","Rate threat in 3 ambiguous + 1 objectively safe scenario.","Ambiguous ≤5; safe ≤2; safety perception intact","Flat low ratings; numbing; emotional blunting; dissociation","Safe scenario rated ≥6; pervasive threat misidentification; PTSD SN hallmark"),
    ("Sensory Sensitivity Screen","Rate sensitivity to sounds, smells, touch 0-10 each.","All ≤4; no hyperreactivity","All ≤2; sensory blunting; numbing PTSD phenotype","≥7 in ≥2 modalities; multi-modal sensory hypersensitivity; SN HYPER PTSD"),
    ("Error Monitoring","Sentence with 2 errors. Spontaneous correction.","Catches ≥1; calm correction","Misses both; avoidant processing; numbing","Catches errors; secondary shame; self-critical; perfectionism PTSD"),
    ("Pain Sensitivity","Fingernail pressure 3 sec. Rate 0-10.","3-6; bilateral symmetry","Low ≤2; pain numbing; dissociation; freeze response","≥8; pain hypersensitivity; somatic PTSD; SN-pain HYPER"),
  ],
  "smn_tests":[
    ("Startle-Motor Response Measure","Drop pen unexpectedly. Measure: body movement extent, time to stillness, vocalizations.","Small controlled flinch; still within 2 sec; no vocalization","No motor response; freeze PTSD; motor numbing; dissociation","Large whole-body startle; ≥5 sec to stillness; vocalizes; PTSD hyperarousal motor"),
    ("Postural Tension Assessment","Observe posture throughout assessment. Rate visible tension in shoulders, jaw, hands 0-10.","Relaxed posture; ≤3 tension in observable areas","Flat; collapsed posture; frozen stance; hypoarousal PTSD","Rigid tense posture; clenched hands; jaw tension; defensive body posture"),
    ("Finger Tapping Rate","Tapping 10 sec each hand.","≥45 taps; regular","<35; psychomotor slowing; freeze state; dissociative PTSD","Normal but highly irregular; tremor; sympathetic arousal interfering with motor"),
    ("Grip Assessment — Proximity Comfort","Bilateral grip. Note anxiety during close physical proximity to examiner.","Firm symmetric; comfortable with examiner proximity","Weak; avoids examiner contact; hypervigilance to proximity","Grips tightly; difficulty releasing; freeze-grip; hyperactivated protective motor"),
    ("Tandem Gait","Heel-to-toe 10 steps. Note environmental scanning during walking.","≤1 step off; forward-focused gaze","Multiple failures; dissociated; poor proprioceptive engagement","Normal walking but scans environment; hypervigilant orienting during movement"),
    ("Romberg Test","Feet together eyes closed 30 sec.","Minimal sway; stable","Excessive sway; proprioceptive HYPO; freeze","Hypertense stance; minimal sway but extreme muscle activation; exhausting to maintain"),
    ("Rapid Alternating Movements","Pronate/supinate 10 sec.","Regular; ≥14 cycles","Slow; reduced amplitude; freeze state motor HYPO","Normal rate but visible tremor; sympathetic activation interfering with smooth output"),
    ("Reaction Time","Hand lift when pen drops. 3 trials.","<250 ms; consistent","≥500 ms; freeze motor delay; hypoarousal PTSD","<150 ms; false starts; hypervigilant anticipatory motor readiness; PTSD"),
  ],
  "limbic_tests":[
    ("Emotional Face Recognition — Threat Bias","Identify 6 emotions from photos. Note: neutral faces misidentified as threatening?","≥5/6 correct; no systematic bias","<4/6; emotional recognition failure; numbing PTSD","Misidentifies neutral as angry/fearful; threat-bias; amygdala HYPER"),
    ("Fear Generalization Assessment","Rate fear in 2 objectively threatening and 2 objectively safe scenarios.","Threatening ≥6; safe ≤3; discriminates threat from safety","Flat ≤3; fear numbing; avoidant/dissociative PTSD","Safe scenarios rated ≥6; fear generalization; hippocampal context failure; PTSD"),
    ("Affect Labeling — Trauma Adjacent","3 scenarios including 1 trauma-adjacent. Name emotion.","Accurate differentiated labels; tolerates trauma-adjacent scenario","'Nothing'; emotional numbing; avoidant PTSD","Extreme reaction on trauma-adjacent scenario; intense limbic reactivity; re-experiencing"),
    ("Emotional Memory — Trauma Bias","10 words (5 trauma-related, 5 positive). Recall after 5-min distractor.","Balanced recall; slight positive advantage","Flat recall; emotional numbing; dissociative PTSD","Trauma words exclusively recalled; positive words inaccessible; PTSD memory bias"),
    ("Reward Sensitivity — Anhedonia Screen","Rate interest in 5 activities 0-10.","Average ≥5; some preserved pleasure","Average ≤2; severe anhedonia; limbic HYPO; complex PTSD","Average ≥6 but accompanied by guilt for enjoying; survivor guilt"),
    ("Anger/Irritability Screen","Ask: 'Are you more irritable or have sudden anger since the trauma?' Rate 0-10.","0-3; normal emotion regulation; no marked change","0; emotional blunting; numbing; avoidant PTSD","≥8; explosive anger; limbic HYPER; PTSD criterion D2"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no shift; emotional blunting; dissociation","Sad memory triggers trauma re-experiencing; cannot access happy memory"),
  ],
  "attn_tests":[
    ("Sustained Attention — Trauma Interference","Count backward from 50 for 60 sec. Note intrusion interruptions.","≤2 errors; continuous; no intrusions","Monotonous; hypervigilant suppression of intrusions","Stops 1-3 times due to intrusive thoughts; PTSD attentional intrusion"),
    ("Threat Vigilance — Attentional Bias","Letter cancellation; include 3 trauma-relevant words embedded. Note time near those words.","Uniform time; ≤3 omissions","Slow overall; poor engagement; avoidant","Selectively slows near trauma words; attentional capture; PTSD VAN HYPER"),
    ("Selective Attention","Circle all A's in 20-item letter list.","<60 sec; ≤3 errors","Slow; multiple omissions; avoidant disengagement","Normal speed but with intrusion breaks; attention disrupted by re-experiencing"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one; avoidant; poor resource allocation","Completes but hypervigilance to examiner consumes attentional resources"),
    ("Visual Search","Find star among 20 distractors.","<20 sec; systematic","Fails; gives up; poor attentional engagement","Fast but hypervigilantly scans beyond found target; cannot disengage search"),
    ("Alerting — Safety vs Threat","'Ready' before stimulus. Assess alerting quality.","Appropriate alerting; calibrated readiness","Absent alerting; no preparatory response; dissociation","Chronically over-alerted; cannot down-regulate vigilance; PTSD hyperarousal"),
    ("Orienting — Threat-Laden Environment","Count orienting responses to new stimuli (movement, door sound) during 10-min assessment.","≤2 orienting per 10 min; returns to task quickly","0-1 orienting; hypervigilant suppression; over-controlled","≥6 orienting per 10 min; hypervigilant scanning; PTSD criterion D"),
    ("Attentional Fatigue","Letter cancellation 5 min; compare halves.","Drop <15%","Drop >40%; rapid attentional depletion; chronic PTSD","20-30% drop; moderate; hyperarousal sustaining at cost of accuracy"),
  ],
  "phenotype_text":"PTSD phenotypes: (1) HYPERAROUSAL — SN HYPER + SMN HYPER + Attention HYPER (hypervigilance); (2) AVOIDANT/NUMBING — Limbic HYPO + DMN HYPO (dissociative); (3) RE-EXPERIENCING — DMN HYPER + Limbic HYPER (intrusive); (4) COMPLEX PTSD — all networks dysregulated. Map dominant cluster for FNON protocol selection.",
  "fnon_strategy":"Primary: Inhibit amygdala-driven SN hyperactivity (right prefrontal cathodal; mPFC anodal to facilitate fear extinction). Facilitate CEN (left DLPFC) to restore prefrontal regulation. For dissociative subtype: facilitate Limbic engagement before inhibiting arousal. Combine with trauma-focused therapy.",
}

# ── 9. OCD ────────────────────────────────────────────────────────────────────
OCD = {
  "full":"OCD","short":"OCD","slug":"ocd",
  "adapted_note":"Includes obsessive-compulsive cycle markers, error hypermonitoring tests, and OCD network loop phenotype indicators",
  "dmn_rel":"OCD involves DMN HYPERACTIVITY dominated by egodystonic intrusive obsessional thoughts — unwanted self-discordant ideation that invades default mode processing. The DMN becomes colonized by obsessional content, creating intrusive thought loops experienced as alien to self-concept.",
  "cen_rel":"CEN dysfunction in OCD is complex: HYPER error-monitoring loops (OFC-ACC hyperactivation) alongside HYPO inhibitory control (inability to suppress compulsive behavior). The CEN cannot terminate obsessional sequences — cognitive inflexibility and poor set-shifting drive the compulsive loop.",
  "sn_rel":"SN HYPERACTIVITY is the engine of OCD: ACC generates excessive error signals ('not just right' experiences) and anterior insula drives visceral discomfort motivating compulsions. OCD's SN creates false alarm error detection without true pathology.",
  "smn_rel":"OCD SMN shows HYPER compulsive motor output: repetitive movements (checking, washing, ordering) reflect cortico-striato-thalamo-cortical loop hyperactivation. Motor inhibition is HYPO — the patient cannot stop compulsive motor sequences once initiated.",
  "limbic_rel":"Limbic HYPERACTIVITY in OCD drives the anxiety-compulsion cycle: anxiety from obsessions motivates compulsions for temporary relief. Amygdala hyperreactivity to contamination, harm, and symmetry violations creates emotional urgency sustaining OCD.",
  "attn_rel":"Attention in OCD shows HYPER for obsessional cues (hypervigilance for triggers) and HYPO for flexible attention (unable to disengage from obsessional focus). VAN HYPER captures and locks attention onto obsession-relevant stimuli; DAN HYPO reduces top-down control.",
  "dmn_tests":[
    ("Intrusive Thought Probe","Ask: 'Do you have unwanted thoughts that feel foreign to you?' Rate frequency and egodystonicity 0-10.","Occasional intrusive thoughts ≤3; egodystonic ≤4; normal","Minimal intrusive thoughts; poor insight; avoidant; or pure-obsessional OCD","≥8 frequency; ≥8 egodystonicity; classic OCD DMN intrusion pattern"),
    ("Autobiographical Memory — OCD Contamination","3 personal memories. Note: does OCD content (contamination, harm) intrude into memory retrieval?","Specific memories; minimal OCD theme contamination","Sparse; avoidant; poor self-referential processing; suppression","OCD themes contaminate non-related memories; memory as obsessional trigger"),
    ("Ego-Dystonicity Self-Assessment","Ask: 'Do your OCD thoughts feel like you or not like you?' Rate self-discordance 0-10.","0-3; insight into ego-dystonic nature; healthy ego distance","0; thoughts feel ego-syntonic; OCD-personality overlap; poor insight variant","8-10; extreme ego-dystonicity; thoughts feel like alien intrusions"),
    ("Mind Wandering — OCD Theme","2-min reading with 4 probes. Note: is off-task thinking OCD-themed?","≤1/4 off-task; non-OCD content","0/4; hypervigilant suppression; thought suppression paradox","3-4/4 off-task; OCD-specific obsessional content; DMN-OCD capture"),
    ("Uncertainty Tolerance Test","Ask: 'Can you leave a task uncertain or incomplete without checking?' Rate tolerance 0-10.","≥6 tolerance; task completion without checking","≥8 tolerance through avoidance of triggering tasks; not adaptive","≤2; extreme uncertainty intolerance; OCD core cognitive signature"),
    ("Prospective Memory — Checking Behavior","5-min reminder task. Note: does patient check repeatedly that they haven't forgotten?","Independent recall within 5±1 min; no checking","Forgets entirely; poor prospective monitoring","Asks repeatedly before 5 min; checking behavior; OCD prospective uncertainty"),
    ("Symmetry/Order Observation","Observe if patient adjusts objects during assessment. Rate compulsive ordering 0-10.","0-2; no ordering; tolerates asymmetry","0; complete indifference to order; contamination-type OCD","≥7; adjusts objects; distress at asymmetry; symmetry/ordering OCD subtype"),
  ],
  "cen_tests":[
    ("Digit Span Backward — OCD Doubt","Backward digit recall. Note: does OCD doubt cause re-checking of sequence?","Span ≥5; no re-checking","Span ≤3; cognitive rigidity; executive dysfunction","Normal span but asks to repeat or correct; doubt-driven checking"),
    ("Trail Making B — Compulsive Retracing","Alternating connections. Note: does patient retrace completed connections?","<90 sec; ≤2 errors; no retracing","Slow >150 sec; poor cognitive flexibility; OCD CEN rigidity","Slow due to retracing and checking; >150 sec from compulsive retracing"),
    ("Verbal Fluency FAS — Rule Checking","Letter fluency 60 sec per letter. Note: does OCD doubt interrupt generation?","≥10 words; no self-doubt interruptions","<6; cognitive rigidity; poor fluency","Normal count but interrupted by 'does that count?'; rule-checking during fluency"),
    ("Stroop — OCD Content Words","Include contamination or harm words. Compare interference.","<45 sec; ≤3 errors; no content-specific slowing","No Stroop effect; avoidant OCD","Selectively slowed on OCD-relevant words; OCD emotional Stroop"),
    ("Serial 7 Subtraction — Rechecking","Backward 100-7 for 5 steps. Note: does patient recheck answers?","5/5 correct; no rechecking","≤2/5; arithmetic failure; executive HYPO","4-5/5 but asks to 'check'; compulsive rechecking; OCD CEN-checking loop"),
    ("Go/No-Go Task — Error Distress","Tap for all except X. Note: do errors cause extreme distress?","≤2 commission; ≤2 omission; accepts errors calmly","High omissions; avoidant; poor inhibitory initiation","1-2 commission errors cause extreme distress; error hypersensitivity; OCD hallmark"),
    ("Cognitive Flexibility — Rule Switch","Sort 6 objects by size then color. Switch rule without announcement. Note perseveration.","Detects rule change within 2 trials; <3 perseverative errors","Cannot detect change; extreme perseveration; OCD cognitive rigidity","Detects change but repeats old rule; needs explicit confirmation before switching"),
  ],
  "sn_tests":[
    ("Error Hypermonitoring Screen","Ask: 'Do you feel you've made a mistake even when you know you haven't?' Rate 0-10.","0-3; appropriate error monitoring; accepts non-errors","0-1; no error monitoring; poor self-evaluation; insight deficit","≥8; constant error sense; ACC hyperactivation; 'not just right' OCD hallmark"),
    ("'Not Just Right' Experience","Present 3 tasks (stack books, align pens, close box). Rate discomfort if not perfectly aligned 0-10.","0-3; tolerates imperfection; ≤3 sensory discomfort","0; complete indifference; contamination OCD; not symmetry-type","≥7 on ≥2 tasks; extreme 'not just right'; SN HYPER OCD signature"),
    ("Interoception — Contamination Focus","Count heartbeats 25 sec. Then ask: 'Are you aware of contamination sensations on your hands?'","Accurate heartbeat count; no contamination hyperawareness","Low accuracy; poor interoception; avoidant suppression","High accuracy AND hyperaware of somatic contamination cues; SN contamination circuit"),
    ("Doubt Generation Test","Ask patient to lock/close something then walk away. Rate return urge 0-10.","No return urge; completion accepted; uncertainty tolerated","No checking urge; avoidant; hasn't engaged with trigger","Strong return urge ≥8; cannot tolerate completion uncertainty; OCD SN hallmark"),
    ("Contamination Sensitivity","Ask patient to touch a common object (pen). Rate disgust/contamination discomfort 0-10.","0-3; no contamination sensitivity; tolerates neutral touch","0; complete indifference; other OCD type; suppression","≥7; contamination disgust; hand-washing urge; contamination OCD subtype"),
    ("Sensory Switching — OCD Doubt on Transition","Alternate visual and auditory tasks. Note: does switching trigger doubt about previous task?","<2 sec switch; smooth; no backward doubt","Slow switching; rigid; poor cognitive flexibility","Switches but immediately doubts previous task; OCD doubt on transition"),
    ("Harm OCD — Threat Overestimation","Rate probability that touching a doorknob will cause illness 0-100%.","<5%; realistic probability assessment","0%; denial; minimization; not harm-type OCD","≥40%; extreme threat overestimation; harm OCD; SN-ACC hyperactivation"),
  ],
  "smn_tests":[
    ("Compulsive Motor Survey","Ask: 'Do you perform repeated physical actions to reduce anxiety (checking, washing, ordering, tapping)?' Rate categories 0-10.","0-2 ritual categories; brief, controllable","0; pure obsessional OCD; mental rituals only; no motor compulsions","≥5 ritual categories; hours of compulsive motor behavior/day; SMN HYPER"),
    ("Motor Inhibition — Compulsion Stop Test","Start writing 5 letters; ask patient to stop mid-sequence. Rate urge to complete 0-10.","Stops immediately; urge ≤3; accepts incompletion","Stops; no urge; avoidant or pure obsessional type","Cannot stop; 'must finish'; urge ≥8; compulsive motor completion; corticostriatal loop"),
    ("Finger Tapping — Symmetry Check","Tapping 10 sec each hand. Note: does asymmetry in counts trigger checking?","≥45 taps each; accepts count difference; no checking","<40; motor slowing; suppression state","Normal tapping but asks for count comparison; adjusts to match; symmetry OCD motor"),
    ("Repetitive Movement Observation","Count skin picking, nail checking, object adjustment, hair touching during 10-min assessment.","0-2 self-directed behaviors per 10 min","0; frozen; avoidant; hypervigilant self-control","≥5 repetitive self-directed behaviors per 10 min; motor compulsive pattern; OCD SMN"),
    ("Tandem Gait — Symmetry OCD","Heel-to-toe 10 steps. Note: does step asymmetry trigger compulsive re-stepping?","≤1 step off; completes without repeating","Multiple failures; poor coordination","Steps off then re-does; symmetry OCD triggers motor compulsion during gait"),
    ("Romberg Test","Feet together eyes closed 30 sec.","Stable; minimal sway","Excessive sway; poor proprioception; motor disengagement","Stable but extreme vigilance about body position; hyperaware of sway; interoceptive OCD"),
    ("Rapid Alternating Movements — Rhythm Checking","Pronate/supinate 10 sec. Note: does patient check rhythm quality?","Regular; ≥14 cycles; accepts rhythm variations","Slow; rigid; poor motor engagement","Regular but asks if rhythm was correct; performance checking; OCD motor doubt"),
    ("Writing Sample — Perfectionism","Write name and address. Note: erasing, restarting, extreme dissatisfaction?","Legible; complete; accepts imperfection","Sparse; avoidant of writing task","Rewrites ≥2 times; extreme perfectionism; compulsive writing; OCD motor"),
  ],
  "limbic_tests":[
    ("Anxiety Without Compulsion","Rate anxiety when NOT performing compulsion after obsession triggers 0-10.","0-4; moderate anxiety; ERP-amenable; proportionate","0; no anxiety; pure obsessional or avoidant; no compulsions","≥8; intolerable anxiety; driving all compulsive behavior; limbic HYPER OCD core"),
    ("Emotional Face Recognition","6 emotions from photos. Note threat/disgust biases.","≥5/6 correct","<4/6; emotional recognition failure; OCD emotional blunting","Misidentifies neutral as disgusted/fearful; emotion misattribution; OCD threat bias"),
    ("Disgust Sensitivity Screen","Rate disgust to 5 scenarios (touching trash, raw meat, spider, stranger's pen, public toilet) 0-10 each.","Average ≤4; proportionate","Average ≤2; disgust blunting; not contamination OCD type","Average ≥7; extreme disgust sensitivity; contamination OCD limbic HYPER"),
    ("Affect Labeling — OCD Context","3 scenarios including 1 OCD-adjacent. Name emotional response.","Accurate, proportionate labels","Flat; 'nothing'; OCD avoidant suppression","Extreme disgust/anxiety on OCD scenario; disproportionate emotional magnitude"),
    ("Fear Response — Harm OCD","3 harm-adjacent scenarios. Rate anxiety 0-10.","Anxiety ≤5; proportionate to actual harm risk","Flat; denial; poor harm appraisal","≥8; catastrophic; hyperresponsive to all harm-adjacent scenarios; harm OCD limbic"),
    ("Reward Sensitivity","Rate interest in 5 activities 0-10.","Average ≥6; maintains enjoyment","Average ≤3; OCD has consumed hedonic capacity; limbic HYPO from exhaustion","High interest expressed but OCD rituals prevent engagement; barrier vs anhedonia"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no affect shift; OCD-emotional numbing","Sad memory triggers OCD rumination; positive memory contaminated by OCD doubt"),
  ],
  "attn_tests":[
    ("Sustained Attention — OCD Checking","Count backward from 50 for 60 sec. Note: does patient check progress or restart after errors?","≤2 errors; continues without restarting","Poor sustained engagement; avoidant; stops","Makes 1 error; restarts from 100; compulsive restart; OCD attention-checking loop"),
    ("Selective Attention — OCD Trigger","Circle all A's in 20-item list. Include OCD-relevant word. Note attentional capture.","<60 sec; ≤3 errors; uniform across items","Slow; multiple omissions; poor attentional engagement","Slows near OCD-relevant word; attentional capture; VAN HYPER for OCD cues"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one; rigid; cannot divide attention","Normal performance but demands certainty in each task; cannot tolerate dual-task uncertainty"),
    ("Visual Search — Over-Checking","Find star among 20 distractors. Note: does patient continue searching after finding target?","<20 sec; systematic; stops when found","Fails; gives up; poor attentional drive","Finds target but continues checking to 'make sure'; OCD attentional over-checking"),
    ("Alerting — Compulsion Readiness","Observe: is patient in heightened readiness for OCD triggers throughout assessment?","Appropriate relaxed readiness; not hypervigilant","Absent alerting; disengaged; avoidant OCD","Chronically hyper-alerted; scanning for triggers; OCD hypervigilance throughout"),
    ("Attentional Lock — OCD Theme","Show OCD-relevant image (dirty object). Time to voluntarily look away.","Looks away within 5 sec on request; smooth disengagement","Immediate avoidant gaze aversion; cannot look; avoidant OCD","Cannot look away for >15 sec; attentional lock; OCD SN-VAN capture"),
    ("Response Inhibition Cancel — Double Check","Cancel circles on mixed-shape sheet. Note: does patient double-cancel items?","<60 sec; ≤2 errors; no double-cancellation","Slow; omissions; poor inhibitory drive","Normal time but double-cancels ≥3 items; compulsive checking during task"),
    ("Attentional Fatigue","Letter cancellation 5 min; compare halves.","Drop <15%","Drop >40%; rapid attentional fatigue; OCD cognitive exhaustion","Performance maintained but checking behavior increases over time; compensatory"),
  ],
  "phenotype_text":"OCD phenotypes: (1) CONTAMINATION — SN HYPER (disgust-contamination) + Limbic HYPER; (2) HARM/CHECKING — SN HYPER (error-monitoring) + Limbic HYPER; (3) SYMMETRY/ORDERING — SMN HYPER + SN 'not just right'; (4) PURE-O — DMN HYPER + CEN HYPO (mental rituals). Map dominant obsessional theme to primary FNON target.",
  "fnon_strategy":"Primary: Inhibit SN hyperactivation (cathodal over ACC/OFC or supplementary motor area). Facilitate right DLPFC for top-down inhibitory control of compulsive sequences. For symmetry-type: target SMN inhibition (M1 cathodal). Combine with ERP therapy.",
}


# ── 10. MS ────────────────────────────────────────────────────────────────────
MS = {
  "full":"Multiple Sclerosis (MS)","short":"MS","slug":"ms",
  "adapted_note":"Includes MS fatigue markers, relapse-remission network testing, and MS phenotype indicators",
  "dmn_rel":"MS causes DMN HYPOACTIVITY through white matter demyelination disrupting PCC-mPFC connectivity. Cognitive fatigue in MS reflects impaired DMN efficiency. Reduced default mode engagement manifests as apathy, reduced self-monitoring, and impaired autobiographical memory consolidation.",
  "cen_rel":"CEN HYPOACTIVITY in MS reflects demyelination of frontal-parietal white matter: slowed processing speed, reduced working memory capacity, and executive dysfunction. Cognitive MS (cognitively impaired MS) shows CEN failure disproportionate to motor severity.",
  "sn_rel":"SN shows complex MS dysfunction: HYPO salience discrimination from cortical lesions alongside HYPER in some pain and fatigue phenotypes. Central sensitization in progressive MS drives SN HYPER contributing to MS-related pain syndromes.",
  "smn_rel":"SMN HYPOACTIVITY is a cardinal feature of MS: demyelination of corticospinal tracts, cerebellar pathways, and proprioceptive fibers cause weakness, ataxia, spasticity, and coordination failure. SMN dysfunction directly maps to disability stage.",
  "limbic_rel":"MS causes limbic dysfunction including depression (50% lifetime prevalence), emotional lability (pseudobulbar affect), and cognitive fatigue affecting affective processing. Periventricular lesions near limbic circuits contribute to emotional dysregulation.",
  "attn_rel":"Attention Networks show HYPO from MS-related cognitive fatigue and white matter lesions: processing speed reduction, sustained attention failure, and divided attention impairment. Fatigue-driven attention failure is the most common MS cognitive complaint.",
  "dmn_tests":[
    ("Cognitive Fatigue Screen — MS","Ask: 'Do you feel mentally exhausted by tasks that shouldn't be tiring?' Rate cognitive fatigue 0-10.","Fatigue ≤3; not disproportionate to effort","Fatigue ≤1; denial; stoic suppression; or very early MS","≥7; disproportionate cognitive fatigue; MS-specific DMN HYPO marker"),
    ("Autobiographical Memory","3 personal memories. Note slowed retrieval speed and coherence.","Specific memories; appropriate retrieval speed","Sparse; very slow; poor retrieval; moderate-advanced MS cognitive","Memories intact but retrieval extremely effortful; MS cognitive fatigue"),
    ("Self-Reference Memory","10 adjectives with self-reference frame. Recall after 2-min distractor.","≥6/10 recalled; balanced content","<4/10; poor self-schema; MS cognitive HYPO","Normal recall but very slow; processing speed contribution"),
    ("Mind Wandering Probe","2-min reading with 4 probes. Note off-task content.","≤1/4 off-task; neutral content","0/4; over-focused; rigid; compensatory hypervigilance","3-4/4 off-task; cognitive fatigue causes task disengagement; MS DMN drift"),
    ("Prospective Memory","5-min reminder task.","Independent recall within 5±1 min","Completely forgets; MS prospective memory failure","Recalls but very slowly; MS processing speed impairs prospective memory"),
    ("Default Narrative","Ask: 'Tell me about your life with MS.' Rate coherence, fatigue impact, adjustment.","Coherent; adaptive coping; realistic acknowledgment of MS impact","Sparse; flat; MS apathy; poor self-reflection","Overwhelmed; catastrophic framing; focuses entirely on disability"),
    ("Semantic Fluency","Animals in 1 min. Note retrieval speed and total.","≥12; fluent retrieval; normal pace","<7; poor fluency; MS semantic network dysfunction","8-11; reduced but present; MS cognitive slowing"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Maximum span.","Span ≥5","Span ≤3; MS working memory failure","Span 4; borderline; MS cognitive impairment"),
    ("Trail Making B","Alternating connections. Time and errors.","<90 sec; ≤2 errors","Slow >180 sec; multiple errors; MS CEN white matter failure","90-150 sec; MS processing speed reduction; executive vulnerability"),
    ("Symbol Digit Modalities","Copy symbols under numbers for 90 sec. Count correct pairs.","≥45 pairs in 90 sec","<25 pairs; MS processing speed core deficit; SDMT criterion","25-44 pairs; mild-moderate MS cognitive impairment; borderline"),
    ("Verbal Fluency FAS","Letter fluency 60 sec per letter.","≥10 words/letter","<6; MS frontal-executive failure","6-9 words; MS cognitive slowing"),
    ("Stroop Interference","Name ink color. Time and errors.","<45 sec; ≤3 errors","Slow >90 sec; severe MS executive failure","45-75 sec; MS processing speed + executive interaction"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct","≤2/5; MS working memory and executive failure","3-4/5; mild errors; MS cognitive involvement"),
    ("Go/No-Go Task","Tap for all except X. 20 letters.","≤2 commission; ≤2 omission","High omissions; poor initiation; MS motor-cognitive interaction","2-4 errors; MS cognitive slowing affects response inhibition"),
  ],
  "sn_tests":[
    ("MS Fatigue — SN Interoception","Ask: 'Do you feel a wave of whole-body fatigue distinct from muscle weakness?' Rate 0-10.","Fatigue ≤3; not disproportionate","0; no MS fatigue; physical fatigue absent or minimal","≥7; disproportionate whole-body fatigue; MS-specific SN-interoceptive fatigue signal"),
    ("Pain Sensitivity — MS Central Pain","Fingernail pressure bilaterally. Ask about spontaneous pain. Note dysesthesia.","3-6; bilateral symmetry; no spontaneous pain","Low bilateral; poor somatic awareness; advanced MS sensory loss","≥7; hyperalgesia; dysesthetic pain; MS central pain syndrome"),
    ("Interoception — Heartbeat Counting","Count heartbeats 25 sec without touch.","Accuracy >0.70","Accuracy <0.50; MS autonomic dysregulation; sensory pathway failure","High accuracy; hyperaware of MS-related somatic changes"),
    ("Error Monitoring","Sentence with 2 errors. Spontaneous correction.","Catches ≥1; calm","Misses both; cognitive fatigue impairs MS error monitoring","Catches errors; significant distress; perfectionism; MS-depression interaction"),
    ("Sensory Switching","Alternate visual tracking and auditory counting. Switch latency.","<2 sec; smooth transitions","Slow >5 sec; MS white matter switching failure","2-4 sec switch; mild slowing; MS processing speed"),
    ("Heat Sensitivity Screen","Ask: 'Do hot temperatures worsen your symptoms (Uhthoff phenomenon)?'","No temperature-dependent symptom worsening","No awareness; early MS or remission","Yes; heat causes functional decline; MS SN pathway sensitization"),
    ("Orienting Response","Unexpected auditory tone mid-task.","Brief interruption; reorientation <3 sec","Absent orienting; MS sensory pathway failure or fatigue","Exaggerated; prolonged recovery; MS-fatigue sensory hypersensitivity"),
  ],
  "smn_tests":[
    ("Finger Tapping Rate — Bilateral","Index tapping 10 sec each hand. Note asymmetry and fatigue over trials.","≥45 taps; ≤10% asymmetry; no fatigue across 3 trials","<35 taps; marked slowing; MS corticospinal tract involvement","Normal first trial; significant slowing by trial 3; MS motor fatigue"),
    ("Grip Strength — Bilateral","Bilateral grip. Compare strength and note fatigue.","Firm symmetric; no fatigability on 3 repetitions","Weak bilateral; MS motor pathway failure","Normal first grip; marked reduction on third; MS motor fatigability"),
    ("Fine Motor Sequencing","Thumb-to-finger 10 sec. Note fatigue and coordination.","≥4 sequences; consistent rhythm","≤2; MS fine motor failure; cerebellar or corticospinal","3 sequences; irregular rhythm; mild MS motor involvement"),
    ("Tandem Gait","Heel-to-toe 10 steps. Note ataxia and truncal sway.","≤1 step off; symmetric; no ataxia","Cannot perform; cerebellar MS; ataxic gait","2-3 step failures; mild ataxia; cerebellar MS involvement"),
    ("Romberg Test","Feet together eyes closed 30 sec. Note Romberg sign.","Stable; minimal sway","Falls; positive Romberg; dorsal column MS lesion; proprioceptive failure","Excessive sway; borderline Romberg; MS sensory pathway involvement"),
    ("Rapid Alternating Movements","Pronate/supinate 10 sec. Note dysdiadochokinesia.","Regular; ≥14 cycles; no dysdiadochokinesia","Marked dysdiadochokinesia; cerebellar MS tract involvement","Mildly irregular; early cerebellar MS involvement"),
    ("MS-Specific Coordination — Finger-Nose","Point to examiner's finger then own nose, 5 repetitions. Note intention tremor.","Accurate; no intention tremor; smooth movement","Cannot perform; severe cerebellar MS; dysmetria","Mild intention tremor on approach; early cerebellar MS"),
    ("Gait Speed — 25-Foot Walk","Walk 25 feet (7.6 m) at maximum safe speed. Time it. (MS T25FW standard.)","<5 sec (>1.5 m/sec); normal MS function","≥20 sec; severely impaired mobility; MS EDSS 5+","5-10 sec; moderately impaired; EDSS 3-4 range"),
  ],
  "limbic_tests":[
    ("MS Depression Screen","Ask: 'How has your mood been since your MS diagnosis?' Rate depression impact 0-10.","0-3; appropriate adjustment; positive coping","0; denial; emotional suppression; stoic","≥7; significant depression; limbic HYPER; MS depression comorbidity (50% lifetime)"),
    ("Emotional Face Recognition","6 emotions from photos. Score /6.","≥5/6 correct","<4/6; MS emotional recognition failure; periventricular limbic involvement","4-5/6; mild; fatigue-related reduced emotional processing"),
    ("Emotional Lability Screen","Ask: 'Do you cry or laugh more easily since MS?' Rate 0-10.","0-2; emotional control maintained","0; emotional blunting; MS apathy","≥7; pseudobulbar affect; bilateral MS corticobulbar tract involvement"),
    ("Affect Labeling","3 emotional scenarios. Name response.","Accurate differentiated labels","Flat; 'nothing'; MS apathy; limbic HYPO","Appropriate labels but fatigued; emotional processing effortful with MS"),
    ("Reward Sensitivity — MS Apathy","Rate interest in 5 activities 0-10.","Average ≥5; motivation maintained","Average ≤2; MS apathy; limbic HYPO; poor rehabilitation engagement","Average ≥6 but fatigue prevents participation; MS fatigue barrier vs anhedonia"),
    ("Fear Response — MS Uncertainty","3 scenarios including 1 MS-related (relapse). Rate anxiety 0-10.","Appropriate 3-6; manages MS uncertainty adaptively","Flat ≤2; denial; MS uncertainty avoidance","≥8 on MS-related scenario; MS anxiety; catastrophic fear of progression"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no shift; MS apathy; emotional blunting","Stuck in sad mood; MS depression contributing to limbic rigidity"),
  ],
  "attn_tests":[
    ("Sustained Attention — MS Fatigue Probe","Count backward from 50 for 60 sec. Note fatigue-driven performance decline.","≤2 errors; consistent pace","Stops within 30 sec; MS attentional failure from fatigue","Consistent but dramatically slows over 60 sec; MS attentional fatigue"),
    ("Selective Attention","Circle all A's in 20-item letter list.","<60 sec; ≤3 errors","Slow; multiple omissions; MS attention failure","Completes; 45-75 sec; MS processing speed reduction"),
    ("Paced Auditory Serial Addition Test — Simplified","Add pairs of digits read at 3-sec pace (1+2=? then 2+3=? etc.) for 5 pairs. Score correct.","4-5/5 correct; <3 sec per addition","≤1/5 correct; MS PASAT failure; auditory working memory + processing speed","2-3/5; mild MS cognitive impairment; MS clinical benchmark"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one; MS attentional resource depletion","Performs both; 3-5 errors each; dual-task vulnerability; MS hallmark"),
    ("Visual Search","Find star among 20 distractors.","<20 sec; systematic","Fails; disorganized; MS severe attention failure","20-40 sec; systematic but slow; MS processing speed"),
    ("Alerting Response","'Ready' before stimulus. Preparatory readiness.","Appropriate alerting; focused","Blunted; not ready; MS fatigue-driven alerting HYPO","Normal alerting but rapidly fatigues; MS sustained alerting failure"),
    ("Orienting","Arrow cue before target. Orienting benefit.","Orienting benefit ≥100 ms","Absent orienting; MS parietal white matter failure","Present but diminished; MS attention slowing"),
    ("Attentional Fatigue Index","Letter cancellation 5 min; compare halves.","Drop <15%","Drop >40%; rapid MS attentional fatigue; core MS complaint","20-35% drop; moderate; MS fatigue-attention interaction"),
  ],
  "phenotype_text":"MS phenotypes: (1) MOTOR-DOMINANT — SMN HYPO (pyramidal/cerebellar); (2) COGNITIVE — CEN HYPO + Attention HYPO (cognitively impaired MS); (3) FATIGUE — SN-DMN HYPO (MS fatigue syndrome); (4) AFFECTIVE — Limbic HYPER (depression) or HYPO (apathy). Correlate with MRI lesion burden for FNON targeting.",
  "fnon_strategy":"Primary: Facilitate CEN (left DLPFC) for cognitive rehabilitation. Facilitate SMN (anodal M1) for motor deficits. Address MS fatigue via SN modulation. tDCS during cognitive training enhances neuroplasticity in remitting MS. Avoid stimulation during acute relapse.",
}

# ── 11. ASD ───────────────────────────────────────────────────────────────────
ASD = {
  "full":"Autism Spectrum Disorder (ASD)","short":"ASD","slug":"asd",
  "adapted_note":"Includes social cognition markers, sensory processing atypicality tests, and ASD network phenotype indicators",
  "dmn_rel":"ASD is characterized by DMN HYPOACTIVITY for social/self-referential processing — reduced mentalizing, theory of mind, and self-reflective thought. The mPFC-TPJ circuit (social default mode) is underconnected. However, ASD may show DMN HYPER for nonsocial restricted interests.",
  "cen_rel":"CEN shows mixed ASD dysfunction: HYPO for flexible task-switching and social executive function; variable or HYPER for restricted, systematized interests. Working memory for nonsocial content may be relatively preserved while social-executive function is HYPO.",
  "sn_rel":"SN shows ATYPICAL/HYPO social salience detection — ASD individuals fail to automatically detect social stimuli as salient. Sensory processing differences reflect HYPER in some modalities (sensory hypersensitivity) and HYPO in others (sensory hyposensitivity), often in the same individual.",
  "smn_rel":"ASD SMN shows atypicalities: coordination difficulties, gait differences, and reduced motor imitation. HYPER repetitive motor behaviors (stereotypies, self-stimulation) reflect motor circuit over-activation. HYPO motor imitation reflects SMN social-motor integration failure.",
  "limbic_rel":"ASD Limbic shows HYPO emotional recognition (alexithymia) and HYPO emotional face processing alongside HYPER emotional dysregulation (meltdowns, rigid emotional responses). The limbic system fails to integrate emotional context with social cognition.",
  "attn_rel":"Attention in ASD: HYPER for restricted interests (hyperfocus/detail-focused processing — weak central coherence) and HYPO for joint attention and social attention orienting. VAN HYPO for social cues; VAN HYPER for non-social special interest stimuli.",
  "dmn_tests":[
    ("Theory of Mind — False Belief","Tell a simple false belief story: 'Sally puts her ball in a basket. Anne moves it. Where will Sally look?' Note correct mentalizing.","Correct: 'In the basket'; mentalizing intact; DMN social circuit preserved","Gives random answer; global cognitive impairment; or very young","Correct answer but excessive processing time or requires repeated explanation; ASD mentalizing delay"),
    ("Self-Reference Memory","10 adjectives with self-reference frame. Recall after 2-min distractor.","≥6/10 recalled; self-reference enhances memory","<4/10; reduced self-reference effect; ASD self-schema HYPO; alexithymia","Normal recall but only non-social self-descriptors (systematic, rule-following) recalled"),
    ("Autobiographical Memory — ASD","Ask 3 personal memories. Note: are memories semantic/fact-based or episodically rich with emotional content?","Episodically rich; balanced social and nonsocial content","Sparse; very brief; poor emotional autobiographical content; ASD pattern","Highly detailed factual memories but lacking social-emotional context; ASD memory style"),
    ("Mind Wandering — Special Interest Content","2-min reading with 4 probes. Note: is off-task thinking related to restricted interests?","≤1/4 off-task; varied content","0/4; rigidly task-focused; ASD perseverative engagement","3-4/4 off-task; content specifically restricted interest topics; ASD DMN special interest"),
    ("Perspective Taking","Ask: 'What might the examiner be thinking right now?' Rate ToM response quality.","Generates plausible examiner perspective; mentalization intact","'I don't know'; refuses social perspective; ASD mentalizing HYPO","Generates literal/factual guess only; no social inference; ASD perspective-taking deficit"),
    ("Prospective Memory","5-min reminder task.","Independent recall within 5±1 min","Completely forgets; poor prospective monitoring","Recalls exactly at 5 min or recites rule repeatedly; ASD rule-following recall"),
    ("Default Narrative — Social Content","Ask: 'Tell me about a social situation from this week.' Rate social content and emotional integration.","Coherent social narrative; emotional perspective-taking present","Minimal social content; prefers nonsocial topics; avoids social narrative","Factual social narrative without emotional integration; ASD social-DMN pattern"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Maximum span.","Span ≥5; may be higher in ASD (strong rote memory)","Span ≤3; intellectual disability overlap; global cognitive HYPO","Span ≥6; ASD enhanced rote working memory; preserved to superior"),
    ("Trail Making B","Alternating connections. Time and errors.","<90 sec; ≤2 errors","Slow >150 sec; ASD set-shifting difficulty; executive HYPO","Fast with good accuracy; ASD rule-following strength; executive pattern strength"),
    ("Verbal Fluency FAS","Letter fluency 60 sec per letter. Note: does ASD systematize output?","≥10 words; fluid switching","<6; poor verbal fluency; language delay overlap","≥12 words but from restricted semantic field; ASD systematic fluency pattern"),
    ("Stroop Interference","Name ink color. Note: does ASD rule-following paradoxically reduce Stroop interference?","<45 sec; ≤3 errors","Slow; impulsive errors; poor rule maintenance","Fast; minimal interference; ASD rigid rule-following reduces Stroop effect"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct","≤2/5; arithmetic difficulty; ASD-ID overlap","5/5 rapidly; ASD systematic arithmetic strength; enhanced CEN for rule-based tasks"),
    ("Go/No-Go Task","Tap for all except X. 20 letters.","≤2 commission; ≤2 omission","High omissions; poor task engagement; ASD social-task disengagement","Very low errors; rigid rule-following; ASD rule-adherence"),
    ("Category Switching","Alternating fruits and vegetables for 60 sec.","≥10 alternations; ≤2 intrusions","<5 alternations; cognitive inflexibility; ASD category rigidity","Normal count but insists on predictable alternating rule; ASD systematizing"),
  ],
  "sn_tests":[
    ("Social Salience Detection","Show 5 images: 3 with social scenes, 2 nonsocial. Ask what patient noticed first in each.","Social elements noticed first in social images; appropriate orienting","No social orienting; focuses exclusively on objects or background; ASD SN social HYPO","Social elements noticed but without emotional significance attributed; ASD social-SN HYPO"),
    ("Sensory Hypersensitivity Screen","Rate sensory discomfort in 5 modalities (sound, light, touch, smell, taste) 0-10 each.","Average ≤4; no significant hypersensitivity","Average ≤2; sensory hyposensitivity; ASD sensory HYPO subtype","≥7 in ≥2 modalities; sensory hypersensitivity; ASD SN HYPER sensory subtype"),
    ("Interoception — Heartbeat Counting","Count heartbeats 25 sec without touch.","Accuracy >0.70","Low accuracy <0.50; ASD interoceptive awareness HYPO; alexithymia","High accuracy; strong interoceptive awareness; ASD interoceptive HYPER subtype"),
    ("Error Monitoring — Social Context","Note reaction to making an error during assessment. Rate distress vs. neutral response.","Mild appropriate response; tolerates error; moves on","No response; does not register social significance of error; ASD SN HYPO","Extreme distress; rigid perfectionism; cannot move on; ASD OCD overlap"),
    ("Sensory Processing — Tactile","Apply light fingertip touch to back of hand with cotton. Rate: aversiveness 0-10.","0-3; not aversive; tactile neutral","0; no tactile response; hyposensitivity; ASD tactile HYPO","≥7; tactile aversiveness; tactile defensiveness; ASD sensory SN HYPER"),
    ("Sensory Switching","Alternate visual and auditory tasks. Note sensory channel preference and switching.","<2 sec; smooth transitions","Cannot switch from preferred sensory modality; ASD sensory perseveration","Switches correctly but with visible distress at modality change; ASD sensory inflexibility"),
    ("Emotional Salience Rating","Rate 6 emotional (3 social, 3 nonsocial interesting) and 6 neutral images 0-10.","Balanced; social emotional images rated appropriately salient","Flat; no differential salience; ASD emotional-social blunting","Nonsocial/interesting images rated far higher than social emotional; ASD salience inversion"),
  ],
  "smn_tests":[
    ("Motor Imitation","Examiner demonstrates 3 novel hand gestures. Ask patient to imitate immediately.","Accurate imitation of ≥2/3 gestures; automatic mirroring","Refuses or unable to imitate; global motor impairment","Accurate but delayed; requires explicit instruction; ASD imitation delay (not loss)"),
    ("Stereotypy Observation","Observe throughout assessment for repetitive motor behaviors (hand-flapping, rocking, finger-rubbing, pacing).","0-1 stereotyped movements per 10 min; socially modulated","0; no stereotypy; ASD without prominent motor stereotypy; or masking","≥3 stereotyped movement types; increased with emotional arousal; ASD SMN HYPER"),
    ("Finger Tapping Rate","Index tapping 10 sec each hand.","≥45 taps; regular rhythm","<35 taps; motor slowing; DCD-ASD overlap","Normal rate but unusual rhythm pattern; ASD motor variability"),
    ("Fine Motor Sequencing","Thumb-to-finger 10 sec.","≥4 sequences","≤2; motor sequencing failure; DCD-ASD overlap","Normal but excessively precise; ASD systematic motor execution"),
    ("Tandem Gait","Heel-to-toe 10 steps.","≤1 step off; consistent","Multiple failures; ASD-DCD motor coordination deficit","Normal with excessively rigid posture; ASD motor rigidity"),
    ("Grip Assessment","Bilateral grip.","Firm symmetric","Weak bilateral; motor coordination deficit","Normal grip but unusual grip style; excessively stiff; ASD motor variability"),
    ("Rapid Alternating Movements","Pronate/supinate 10 sec.","Regular; ≥14 cycles","Slow; irregular; ASD-DCD motor coordination","Regular but with unusual motor quality; ASD motor singularity"),
    ("Gait Observation — ASD","Observe natural gait for 5 steps. Note: toe-walking, arm-swing asymmetry, unusual body posture.","Symmetric; heel-toe; bilateral arm swing; normal posture","Normal gait but reduced spontaneity; low motor engagement","Toe-walking; reduced arm swing; unusual posture; ASD-specific gait markers"),
  ],
  "limbic_tests":[
    ("Emotional Face Recognition — ASD","Identify 6 emotions from facial photos. Score /6. Note fixation on mouth vs eyes.","≥5/6 correct; uses eye region for recognition","<4/6; ASD facial affect recognition failure; limbic-social HYPO","4/6; uses mouth or non-diagnostic features; ASD non-typical recognition strategy"),
    ("Alexithymia Screen","Ask: 'When you feel something, can you describe what that feeling is?' Rate 0-10.","≥6; can label emotions with some description; emotional awareness maintained","0-3; poor emotional awareness; alexithymia; ASD limbic HYPO","Extreme difficulty labeling own emotions despite preserved emotion recognition"),
    ("Affect Labeling","3 emotional scenarios. Name response.","Identifies appropriate basic emotions","Flat; 'nothing'; limited emotional vocabulary; ASD alexithymia","Correct but over-literal; misses emotional nuance; ASD concrete limbic processing"),
    ("Social Empathy Assessment","Ask: 'How do you think that person felt in this story?' Rate perspective-taking quality.","Generates plausible emotional perspective; empathic reasoning","Cannot generate emotional perspective; ASD cognitive empathy HYPO","Generates factual description only; 'they would be cold' not 'frightened'; ASD empathy"),
    ("Reward Sensitivity — Special Interest vs Social","Rate interest in 5 activities: 3 related to possible special interest, 2 social activities.","Balanced interest across types; ≥5/10 for social activities","Low interest in all; global anhedonia; ASD-depression overlap","10/10 for special interest; 0-2 for social; ASD reward circuit asymmetry"),
    ("Meltdown/Emotional Regulation Screen","Ask: 'When overwhelmed, what do you do?' Rate emotional regulation strategies 0-10.","Describes ≥2 coping strategies; emotion regulation maintained","Describes avoidance as primary strategy; limited regulation repertoire","Reports meltdowns; overwhelm; inability to regulate; ASD limbic HYPER in overload"),
    ("Affective Flexibility","Sad then happy memory. Rate mood shift.","Shift ≥3 points; recovery <3 min","Flat; no shift; ASD emotional blunting; alexithymia pattern","Mood shift but extreme difficulty initiating; ASD emotional inertia"),
  ],
  "attn_tests":[
    ("Joint Attention Observation","Point to object across room. Note: does patient follow gaze/point? Rate joint attention 0-10.","Follows point; establishes joint attention; social orienting maintained","Does not follow; reduced joint attention; ASD VAN social HYPO","Follows but does not share attention socially; looks at object only; ASD joint attention"),
    ("Hyperfocus Screen","Ask: 'Do you ever become so absorbed in an activity you lose track of time?' Rate hyperfocus frequency 0-10.","0-4; normal absorbed attention; controllable","0; no hyperfocus; inattentive ASD subtype; or flat engagement","≥8; hyperfocus; cannot disengage; ASD attention HYPER for special interests"),
    ("Selective Attention","Circle all A's in 20-item letter list.","<60 sec; ≤3 errors","Slow; multiple omissions; ASD attention engagement deficit","Fast; very accurate; ASD detail-focused attention strength"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one; ASD attention-splitting difficulty","Performs both but excessively slow; ASD monotropic attention pattern"),
    ("Visual Search — ASD Detail Advantage","Find target symbol among 20 distractors.","<20 sec; systematic","Fails; does not engage with search task","<10 sec; faster than neurotypical; ASD superior visual search (embedded figures strength)"),
    ("Alerting Response","'Ready' before stimulus.","Appropriate alerting; focused","Absent alerting; poor preparatory readiness; ASD social-task disengagement","Over-alerting; rigid anticipatory state; ASD anxiety-attention overlap"),
    ("Social Attention — Eye Contact Duration","Observe eye contact during 5-min assessment. Rate ASD-specific social attention pattern.","Natural variable eye contact; appropriate social attention","Absent eye contact; severe social attention HYPO; classic ASD marker","Prolonged staring; atypical social attention pattern; ASD social attention anomaly"),
    ("Attentional Flexibility — Task Switch","Switch from letter cancellation to symbol copying mid-task without warning. Measure switch cost.","Adapts within 5 sec; switch cost <30 sec","Cannot switch; rigid; requires extensive instruction; ASD inflexibility HYPO","Switches but with extreme distress at transition; ASD cognitive inflexibility"),
  ],
  "phenotype_text":"ASD phenotypes: (1) SOCIAL-COGNITIVE — DMN HYPO + Limbic HYPO (social cognition core); (2) SENSORY-REGULATORY — SN HYPER (sensory hypersensitivity); (3) RESTRICTED/REPETITIVE — SMN HYPER + Attention HYPER (special interest); (4) EXECUTIVE — CEN HYPO (cognitive flexibility). Map to dominant presentation for FNON protocol selection.",
  "fnon_strategy":"Primary: Facilitate DMN social circuits (anodal TPJ/mPFC) to enhance mentalizing and social cognition. For sensory-regulatory subtype: inhibit SN hyperactivation. For executive subtype: facilitate left DLPFC. Combine with behavioral social skills intervention and sensory integration.",
}

# ── 12. LONG COVID ────────────────────────────────────────────────────────────
LONG_COVID = {
  "full":"Long COVID","short":"Long COVID","slug":"long_covid",
  "adapted_note":"Includes brain fog markers, post-exertional malaise indicators, and Long COVID network profile phenotype markers",
  "dmn_rel":"Long COVID causes DMN HYPOACTIVITY through neuroinflammation-mediated default mode disruption: brain fog reflects impaired default mode efficiency, poor spontaneous cognition, and reduced self-referential processing. The 'mental blankness' of brain fog maps to DMN HYPO.",
  "cen_rel":"CEN HYPOACTIVITY is the core of Long COVID brain fog: working memory failure, slowed processing speed, and executive dysfunction reflect frontal-network neuroinflammation. Patients describe 'thinking through fog' — CEN HYPO from microglial activation and cerebrovascular dysfunction.",
  "sn_rel":"Long COVID SN shows HYPER/dysregulated function: dysautonomia (POTS, orthostatic intolerance), interoceptive hypersensitivity, and autonomic volatility reflect SN-autonomic dysfunction. Some patients show sensory hypersensitivity (light, sound, smell) from central sensitization.",
  "smn_rel":"Long COVID SMN shows HYPOACTIVITY through post-exertional malaise: movement worsens symptoms, creating a cycle of deconditioning and SMN under-engagement. Muscle fatigue, weakness, and coordination difficulty reflect peripheral and central SMN dysfunction.",
  "limbic_rel":"Long COVID creates Limbic HYPERACTIVITY through neuroinflammatory anxiety and depression comorbidity (40-60% prevalence). Amygdala reactivity increases from neuroinflammation. Limbic HYPO (anhedonia, flat affect) emerges in prolonged severe cases.",
  "attn_rel":"Attention Networks show profound HYPO in Long COVID brain fog: sustained attention failure, processing speed reduction, and divided attention impairment. DAN HYPO reflects prefrontal neuroinflammation; attentional fatigue is the most common and disabling Long COVID cognitive complaint.",
  "dmn_tests":[
    ("Brain Fog — Default Mode Assessment","Ask: 'Does your mind feel blank or empty when you're not doing anything?' Rate mental blankness 0-10.","0-3; spontaneous thought present; mental activity maintained","0; no brain fog; mild or recovered Long COVID","≥7; pervasive mental blankness; spontaneous thought absent; Long COVID DMN HYPO hallmark"),
    ("Autobiographical Memory — Brain Fog","Ask 3 personal memories. Note: are memories accessible or trapped 'behind fog'?","Accessible memories; normal retrieval effort","Sparse; very slow; 'can't find them'; brain fog memory retrieval failure","Memories present but take extreme effort; 'fog between me and memories'; Long COVID pattern"),
    ("Self-Reference Memory","10 adjectives with self-reference frame. Recall after 2-min distractor.","≥6/10 recalled; self-reference enhances memory","<4/10; brain fog impairs self-schema encoding","Normal recall on good day; fluctuates with symptom severity; Long COVID variability"),
    ("Mind Wandering Probe","2-min reading with 4 probes. Note off-task content.","≤1/4 off-task; varied mind-wandering","0/4; mental blankness even off-task; no spontaneous cognition","3-4/4 off-task; cannot maintain reading; brain fog attention failure"),
    ("Symptom Fluctuation History","Ask: 'Do your cognitive symptoms vary day to day or within a day?'","Stable baseline; consistent performance","No fluctuation; possibly early/mild Long COVID","Significant daily fluctuation; worse post-exertion (PEM); Long COVID DMN variability"),
    ("Prospective Memory","5-min reminder task.","Independent recall within 5±1 min","Completely forgets; brain fog prospective memory failure","Recalls but extremely effortful; 'had to work very hard to remember'; Long COVID"),
    ("Default Narrative — Long COVID Impact","Ask: 'How has Long COVID changed your daily life?' Rate coherence and emotional content.","Coherent; realistic adjustment; some positive coping","Sparse; minimizes impact; suppression","Overwhelmed narrative; profound disability; grief over lost function; Long COVID"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Maximum span. Note effort vs performance.","Span ≥5","Span ≤3; brain fog working memory failure; Long COVID CEN HYPO","Span 4; effortful; 'harder than it should be'; Long COVID CEN"),
    ("Trail Making B","Alternating connections. Time and errors.","<90 sec; ≤2 errors","Slow >150 sec; multiple errors; Long COVID CEN failure","90-150 sec; moderate slowing; brain fog CEN vulnerability"),
    ("Processing Speed — Symbol Digit","Copy symbols under numbers for 60 sec. Count correct pairs.","≥40 pairs in 60 sec","<20 pairs; severe processing speed failure; Long COVID brain fog core","20-39 pairs; mild-moderate; brain fog processing speed"),
    ("Verbal Fluency FAS","Letter fluency 60 sec per letter.","≥10 words/letter","<6; brain fog verbal fluency failure","6-9 words; reduced; Long COVID brain fog"),
    ("Stroop Interference","Name ink color. Time and errors.","<45 sec; ≤3 errors","Slow >90 sec; brain fog executive failure","45-75 sec; moderate slowing; Long COVID processing speed + executive"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct","≤2/5; brain fog working memory failure","3-4/5; mild errors; brain fog CEN involvement"),
    ("Go/No-Go Task","Tap for all except X. 20 letters.","≤2 commission; ≤2 omission","High omissions; brain fog poor initiation","2-4 errors; mild; brain fog response inhibition"),
  ],
  "sn_tests":[
    ("Dysautonomia Screen — POTS Proxy","Ask: 'Do you feel dizzy or worse when standing?' Rate orthostatic symptoms 0-10.","0-2; no orthostatic symptoms; autonomic stability","0; no dysautonomia; mild or recovered Long COVID","≥6; orthostatic intolerance; Long COVID autonomic SN dysfunction"),
    ("Sensory Sensitivity Screen","Rate discomfort with light, sound, smell 0-10 each.","All ≤4; no hypersensitivity","All ≤2; sensory blunting; severe or atypical Long COVID","≥7 in ≥2 modalities; sensory hypersensitivity; Long COVID SN HYPER"),
    ("Interoception — Heartbeat Counting","Count heartbeats 25 sec without touch.","Accuracy >0.70","Low accuracy <0.50; poor body awareness despite symptoms; dissociation","High accuracy but hyperaware of cardiac symptoms; Long COVID health anxiety"),
    ("Post-Exertional Malaise Screen","Ask: 'Do symptoms worsen after physical or mental effort?' Rate PEM severity 0-10.","0-2; normal fatigue after exertion; no PEM","0; no PEM; mild or early Long COVID","≥7; significant PEM; hallmark Long COVID SN-autonomic feature"),
    ("Error Monitoring","Sentence with 2 errors. Spontaneous correction.","Catches ≥1; calm correction","Misses both; brain fog impairs ACC error monitoring","Catches errors; significant distress; 'can't believe I missed that'; Long COVID mood"),
    ("Sensory Switching","Alternate visual and auditory tasks. Note symptom increase with switching.","<2 sec; smooth; no symptom increase","Slow; brain fog switching failure","Fast but triggers headache or cognitive fatigue; Long COVID SN overload"),
    ("Orienting Response","Unexpected auditory tone mid-task.","Brief interruption; reorientation <3 sec","Absent orienting; mental fog; blunted response","Exaggerated; triggers cognitive fatigue; sensory overload"),
  ],
  "smn_tests":[
    ("Post-Exertional Motor Assessment","Finger tapping 10 sec each hand. Then wait 2 min and repeat. Compare pre-post exertion.","≥45 taps; no significant reduction on repeat","<35 taps at baseline; severe Long COVID motor impairment","Normal baseline; ≥20% decline on repeat; post-exertional motor fatigue; Long COVID SMN"),
    ("Grip Assessment — Fatigue","Bilateral grip. Perform 3 repetitions. Note fatigue across repetitions.","Firm symmetric; no fatigue across 3 reps","Weak bilateral; Long COVID motor weakness","Normal first grip; significant reduction by third; Long COVID motor fatigability"),
    ("Fine Motor Sequencing","Thumb-to-finger 10 sec.","≥4 sequences; consistent","≤2; brain fog fine motor failure","3 sequences; effortful; post-exertional susceptibility"),
    ("Tandem Gait","Heel-to-toe 10 steps.","≤1 step off; consistent pace","Multiple failures; Long COVID motor ataxia or extreme fatigue","Normal but very slow; highly cautious; Long COVID deconditioning pattern"),
    ("Romberg Test","Feet together eyes closed 30 sec.","Stable; minimal sway","Excessive sway; Long COVID autonomic/proprioceptive involvement","Stable but patient reports dizziness; dysautonomia contribution to balance"),
    ("Rapid Alternating Movements","Pronate/supinate 10 sec. Note fatigue.","Regular; ≥14 cycles","Slow; Long COVID motor fatigue","Normal first 5 sec; slows last 5 sec; Long COVID motor fatigability"),
    ("Activity Tolerance Assessment","Ask: 'How much walking before symptom worsening?' Rate activity tolerance 0-10.","≥7; able to perform normal daily activities","≤2; severe activity intolerance; Long COVID disability","3-5; moderate activity intolerance; Long COVID deconditioning + PEM"),
    ("Reaction Time","Hand lift when pen drops. 3 trials.","<250 ms","≥500 ms; brain fog psychomotor slowing","250-400 ms; mild slowing; Long COVID processing speed and motor"),
  ],
  "limbic_tests":[
    ("Long COVID Anxiety Screen","Ask: 'Do you feel more anxious or worried about your health since Long COVID?' Rate 0-10.","0-3; appropriate health concern; adaptive coping","0; denial; suppression; minimizing Long COVID symptoms","≥7; significant health anxiety; Long COVID neuroinflammatory limbic HYPER"),
    ("Emotional Face Recognition","6 emotions from photos. Score /6.","≥5/6 correct","<4/6; neuroinflammatory emotional recognition impairment","4-5/6; mild; brain fog reduces emotional processing efficiency"),
    ("Affect Labeling","3 emotional scenarios. Name response.","Accurate differentiated labels","Flat; 'nothing'; Long COVID emotional blunting","Appropriate labels but effortful; 'brain fog makes it hard to find the word'"),
    ("Anhedonia Screen — Brain Fog","Rate interest in 5 activities 0-10. Note: is low interest from anhedonia or fatigue barrier?","Average ≥5; motivation present; activity limited by symptoms not interest","Average ≤2; Long COVID anhedonia; neuroinflammatory limbic HYPO","High interest but symptoms prevent engagement; preserved motivation with functional barrier"),
    ("Emotional Memory","10 emotional and 10 neutral words. Recall after 5-min distractor.","Emotional advantage; ≥7 emotional recalled","Flat recall; brain fog impairs emotional memory encoding","Normal emotional advantage but overall recall reduced; Long COVID memory-emotion"),
    ("Fear Response — Health Anxiety","3 health-related and 2 neutral scenarios. Rate anxiety 0-10.","Appropriate 2-5 for health scenarios; not catastrophic","Flat ≤2; health concern denied","≥8 on health scenarios; catastrophic health anxiety; Long COVID limbic HYPER"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no shift; Long COVID emotional blunting","Stuck in negative mood; brain fog prevents positive affect access; Long COVID depression"),
  ],
  "attn_tests":[
    ("Brain Fog Attention Screen","Ask: 'How often does your mind go blank mid-task?' Rate 0-10.","0-3; normal attentional continuity","0; no brain fog; mild or recovered","≥7; frequent mid-task blank outs; Long COVID brain fog attention hallmark"),
    ("Selective Attention","Circle all A's in 20-item letter list. Note effort and errors.","<60 sec; ≤3 errors","Slow; multiple omissions; brain fog selective attention failure","Completes but 60-90 sec; effortful; 'harder than it should be'; Long COVID"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one; brain fog attentional resource failure","Performs both but 4-6 errors each; Long COVID dual-task vulnerability"),
    ("Visual Search","Find star among 20 distractors.","<20 sec; systematic","Fails; brain fog disorganized search","20-45 sec; systematic but very slow; brain fog processing speed"),
    ("Alerting Response","'Ready' before stimulus.","Appropriate alerting; focused readiness","Blunted; not ready; brain fog alerting HYPO","Normal alerting but immediately fatigued; Long COVID sustained alerting failure"),
    ("Post-Cognitive-Exertion Symptom Check","After 10 min of assessment: rate cognitive fatigue, headache, and dizziness 0-10 each.","All ≤3; no symptom increase with cognitive effort","No increase; mild or recovered Long COVID","≥5 on ≥2 measures; cognitive PEM; Long COVID hallmark response to mental exertion"),
    ("Response Inhibition Cancel","Cancel circles on mixed-shape sheet.","<60 sec; ≤2 errors","Slow; brain fog psychomotor slowing","Normal start; significant slowing by end; fatigue-driven inhibition failure"),
    ("Attentional Fatigue Index","Letter cancellation 5 min; compare halves.","Drop <15%","Drop >50%; extremely rapid brain fog attentional fatigue; Long COVID hallmark","20-40% drop; moderate; Long COVID attentional fatigue core complaint"),
  ],
  "phenotype_text":"Long COVID phenotypes: (1) BRAIN FOG — CEN HYPO + Attention HYPO + DMN HYPO (cognitive predominant); (2) AUTONOMIC — SN HYPER/dysregulated (dysautonomia/POTS); (3) FATIGUE-MOTOR — SMN HYPO + post-exertional malaise; (4) AFFECTIVE — Limbic HYPER (anxiety/depression). Map to dominant symptom cluster for FNON selection.",
  "fnon_strategy":"Primary: Facilitate CEN (left DLPFC anodal) for brain fog and executive dysfunction. Address SN dysregulation (cathodal right prefrontal or insula). Avoid over-stimulation triggering PEM — start with low-intensity, short sessions. Pace cognitive rehabilitation to avoid post-exertional worsening.",
}

# ── 13. TINNITUS ──────────────────────────────────────────────────────────────
TINNITUS = {
  "full":"Tinnitus","short":"Tinnitus","slug":"tinnitus",
  "adapted_note":"Includes tinnitus-related distress markers, auditory salience network tests, and Tinnitus network phenotype indicators",
  "dmn_rel":"Tinnitus chronifies through DMN HYPERACTIVITY: the phantom sound becomes incorporated into self-referential thought, leading to tinnitus-related rumination, identity disruption ('my life is ruined by this sound'), and persistent default mode engagement with tinnitus as the signal.",
  "cen_rel":"CEN shows HYPOACTIVITY in tinnitus from attentional resource consumption by the tinnitus signal. Executive function, working memory, and cognitive flexibility are reduced when tinnitus is severe. Patients report difficulty concentrating ('the ringing takes up all my thinking').",
  "sn_rel":"SN HYPERACTIVITY is the neural mechanism of tinnitus chronification: the auditory cortex generates an artificial salience signal (the phantom sound) that is detected and amplified by the ACC-insula as highly salient and personally relevant. This creates an irreversible salience loop.",
  "smn_rel":"Tinnitus SMN shows mild HYPO changes: reduced head and neck motor control, jaw muscle tension, and somatic tinnitus components (temporomandibular, cervicogenic tinnitus). Somatic tinnitus responds to SMN/cervical manipulation. General motor function is typically preserved.",
  "limbic_rel":"Limbic HYPERACTIVITY drives tinnitus distress: the emotional significance attributed to the phantom sound determines disability. Amygdala hyperconnectivity with auditory cortex creates distress-sound coupling. Depression (40%) and anxiety (55%) are the most common tinnitus comorbidities.",
  "attn_rel":"Attention Networks show HYPO for non-tinnitus content (tinnitus captures attentional resources) and HYPER-vigilance for the tinnitus signal itself. Sustained attention and working memory are disrupted by tinnitus intrusion. Attentional training (attention-focused sound therapy) is a key intervention.",
  "dmn_tests":[
    ("Tinnitus Rumination Probe","Ask: 'How much time do you spend thinking about your tinnitus per day?' Rate 0-10.","<30 min/day; tinnitus doesn't dominate thought; 0-3/10","0; no rumination; tinnitus not intrusive; mild tinnitus phenotype","≥3 hours/day; constant tinnitus rumination; DMN HYPER tinnitus colonization; severe distress"),
    ("Tinnitus Identity Integration","Ask: 'Would you say tinnitus has become part of who you are?' Rate 0-10.","0-3; tinnitus acknowledged but not identity-defining; adaptive","0; denies tinnitus impact; suppression or mild tinnitus","≥8; tinnitus as core identity; catastrophizing; 'I am my tinnitus'; DMN colonization"),
    ("Autobiographical Memory","3 personal memories. Note: are memories contaminated by tinnitus-related negative emotion?","Specific memories; appropriate emotional content; tinnitus not central","Sparse; poor memory retrieval; attention from tinnitus impairs encoding","Memories contaminated with tinnitus distress; 'can't remember when it was quiet'"),
    ("Mind Wandering Probe","2-min reading with 4 probes. Note: is off-task content tinnitus-related?","≤1/4 off-task; neutral mind-wandering","0/4; hypervigilant task focus; suppression of tinnitus awareness","3-4/4 off-task; tinnitus-specific content; DMN-tinnitus capture"),
    ("Pre-Tinnitus vs Current Quality of Life","Ask patient to rate life satisfaction before vs after tinnitus onset 0-10.","Both ≥6; or appropriate minimal decline; good coping","No difference; denial; minimization","Pre ≥7; current ≤3; significant QoL decline; tinnitus distress; DMN negative impact"),
    ("Prospective Memory","5-min reminder task.","Independent recall within 5±1 min","Completely forgets; tinnitus-attentional load impairs prospective memory","Recalls but effortfully; tinnitus attention competition"),
    ("Default Narrative — Tinnitus Impact","Ask: 'How has tinnitus changed your daily life?' Rate narrative.","Acknowledges impact; maintains positive elements; adaptive coping","Denies any impact; suppression; avoidant","Overwhelmed; 'ruined my life'; no positive elements; catastrophizing; DMN HYPER"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Maximum span.","Span ≥5","Span ≤3; tinnitus-related cognitive interference; CEN HYPO","Span 4; borderline; tinnitus attention competition with working memory"),
    ("Trail Making B","Alternating connections. Time and errors.","<90 sec; ≤2 errors","Slow >150 sec; tinnitus CEN interference","90-150 sec; moderate; tinnitus attentional drain on executive"),
    ("Verbal Fluency FAS","Letter fluency 60 sec per letter.","≥10 words/letter","<6; poor fluency; tinnitus CEN HYPO","7-9 words; tinnitus attentional competition with verbal output"),
    ("Stroop Interference","Name ink color. Time and errors.","<45 sec; ≤3 errors","Slow >90 sec; severe tinnitus executive interference","45-75 sec; moderate; tinnitus processing speed competition"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct","≤2/5; tinnitus cognitive load failure","3-4/5; tinnitus-attention interaction with working memory"),
    ("Go/No-Go Task","Tap for all except X. 20 letters.","≤2 commission; ≤2 omission","High omissions; tinnitus-related inattention","2-4 errors; tinnitus attentional competition"),
    ("Category Switching","Alternating fruits and vegetables for 60 sec.","≥10 alternations; ≤2 intrusions","<5; cognitive rigidity; tinnitus CEN HYPO","Normal count; mildly slowed; tinnitus cognitive drain"),
  ],
  "sn_tests":[
    ("Tinnitus Loudness Rating","Ask patient to rate tinnitus loudness right now 0-10.","0-3; mild tinnitus; SN engagement minimal","0; no tinnitus at assessment moment; assessment captures quiet period","≥7; loud tinnitus at assessment; SN HYPER active; most distressing presentation"),
    ("Tinnitus Distress Rating","Ask: 'How much does tinnitus bother you right now?' Rate 0-10.","0-3; tinnitus present but not distressing; SN not over-engaged","0; no distress; habituated; successful SN down-regulation","≥7; highly distressing; SN HYPER; emotional-salience coupling active"),
    ("Auditory Salience Hyperactivation","In quiet room: ask if tinnitus is louder. In noisy environment: ask if tinnitus is masked.","Tinnitus reduced in noise; normal auditory masking; SN modulation intact","Absent tinnitus in noise; complete masking; mild tinnitus","Tinnitus same in noise; poor masking; SN HYPER; central tinnitus generation"),
    ("Interoception — Sound Hyperawareness","Count heartbeats 25 sec. Note: is patient also hyperaware of tinnitus during body focus?","Accurate heartbeat count; tinnitus in background","Low accuracy; poor interoception; tinnitus-body disconnection","Both heartbeat and tinnitus simultaneously hyperattended; SN dual hyperactivation"),
    ("Somatic Tinnitus Screen","Ask patient to perform jaw clenching, neck rotation, and jaw protrusion. Note any tinnitus change.","No tinnitus modulation with jaw/neck movements; pure audiogenic","Tinnitus absent with somatic maneuver; somatic origin","Tinnitus changes (louder or quieter) with jaw/neck; somatic tinnitus; SMN-SN interaction"),
    ("Error Monitoring","Sentence with 2 errors. Spontaneous correction.","Catches ≥1; calm correction","Misses both; tinnitus-related inattention; poor monitoring","Catches; distressed; perfectionism; tinnitus-depression interaction"),
    ("Sensory Switching","Alternate visual and auditory tasks. Note: does auditory task increase tinnitus awareness?","<2 sec; smooth; auditory task does not worsen tinnitus","Poor switching; tinnitus-related cognitive rigidity","Auditory task markedly increases tinnitus loudness; auditory salience hyperactivation"),
  ],
  "smn_tests":[
    ("Jaw/TMJ Assessment — Somatic Tinnitus","Palpate masseter muscle bilaterally. Note: tenderness, clenching, or TMJ dysfunction.","No tenderness; no TMJ symptoms; normal masticatory function","No TMJ findings; non-somatic tinnitus; audiogenic origin","Bilateral tenderness; TMJ dysfunction; bruxism; somatic tinnitus SMN component"),
    ("Cervical Muscle Tension","Palpate trapezius and cervical paraspinals. Rate tension 0-10 per side.","Tension ≤3; symmetric; no cervicogenic involvement","0 tension; no cervicogenic component","Tension ≥7; cervicogenic tinnitus; somatic SMN-tinnitus link"),
    ("Finger Tapping Rate","Index tapping 10 sec each hand.","≥45 taps; regular","<35 taps; unrelated motor slowing; assess for other pathology","Normal tapping; tinnitus does not affect distal motor; expected finding"),
    ("Grip Assessment","Bilateral grip.","Firm symmetric grip","Weak grip; unrelated to tinnitus; assess other pathology","Normal grip; tinnitus does not affect hand motor; expected"),
    ("Fine Motor Sequencing","Thumb-to-finger 10 sec.","≥4 sequences","≤2; unrelated to tinnitus; assess other pathology","Normal; tinnitus does not affect fine motor"),
    ("Head/Neck Motor Control","Rotate head left then right to end range. Rate range and smoothness.","Full symmetric rotation; no tinnitus change","Reduced range; cervical pathology; assess further","Reduced range with tinnitus modulation; cervicogenic tinnitus SMN component"),
    ("Rapid Alternating Movements","Pronate/supinate 10 sec.","Regular; ≥14 cycles","Slow; unrelated pathology","Normal; tinnitus typically does not affect limb motor function"),
    ("Reaction Time","Hand lift when pen drops. 3 trials.","<250 ms","≥450 ms; tinnitus-related cognitive slowing; brain fog overlap","250-400 ms; mild slowing; tinnitus attentional competition with motor readiness"),
  ],
  "limbic_tests":[
    ("Tinnitus Distress — Emotional Impact","Ask: 'What emotion does your tinnitus bring up?' Rate emotional intensity 0-10.","0-4; mild annoyance; no significant emotional distress; adaptive coping","0; no emotional response to tinnitus; complete habituation or denial","≥8; extreme distress; despair, anger, or hopelessness about tinnitus; limbic HYPER"),
    ("Emotional Face Recognition","6 emotions from photos. Score /6.","≥5/6 correct","<4/6; tinnitus-related emotional blunting from depression","4-5/6; mild; tinnitus distress reduces emotional processing capacity"),
    ("Affect Labeling","3 scenarios. Name emotional response.","Accurate differentiated labels","Flat; 'nothing'; tinnitus-related emotional blunting","Appropriate labels but tinnitus distress colors all responses with frustration/despair"),
    ("Tinnitus-Related Depression Screen","Ask: 'Has tinnitus caused you to feel hopeless or lose interest in life?' Rate 0-10.","0-3; tinnitus does not cause hopelessness","0; no depression; tinnitus not distressing; habituated","≥7; tinnitus-driven depression; limbic HYPER; comorbid major depression"),
    ("Reward Sensitivity — Tinnitus Barrier","Rate interest in 5 activities 0-10. Note: is avoidance from tinnitus distress or anhedonia?","Average ≥6; maintains enjoyment despite tinnitus","Average ≤3; tinnitus anhedonia; limbic HYPO from severe distress","High interest but tinnitus prevents quiet activities; preserved motivation with sound barrier"),
    ("Social Withdrawal — Tinnitus","Rate comfort in 5 social scenarios. Note if tinnitus drives withdrawal.","≥5/10 for ≥4 scenarios; social engagement maintained","≤3 all scenarios; unrelated depression or severe tinnitus","≤3 in social/noisy scenarios; tinnitus-driven social avoidance; limbic-social withdrawal"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no shift; tinnitus-related emotional blunting; depression","Stuck in tinnitus-distress mood; cannot access positive affect; limbic HYPER tinnitus"),
  ],
  "attn_tests":[
    ("Tinnitus Attentional Capture","During assessment: count how many times patient spontaneously mentions or rubs ear (tinnitus-directed behavior).","0-1 tinnitus-directed behaviors in 10 min; background awareness","0; tinnitus fully habituated; attention not captured","≥3 tinnitus-directed behaviors; attentional capture; VAN HYPER for tinnitus signal"),
    ("Selective Attention","Circle all A's in 20-item letter list.","<60 sec; ≤3 errors","Slow; multiple omissions; tinnitus-related attention failure","Normal but effortful; tinnitus competes for attentional resources"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one; tinnitus attentional drain","Completes but 3-5 errors each; tinnitus competing for attentional resources"),
    ("Visual Search","Find star among 20 distractors.","<20 sec; systematic","Fails; tinnitus-related attention failure","20-40 sec; systematic but slow; tinnitus attentional interference"),
    ("Quiet Environment Attention Test","Perform letter cancellation in quiet room. Note: does quiet environment increase tinnitus awareness and decrease performance?","Normal; quiet room does not worsen performance","No change; tinnitus habituated; quiet not triggering","Worse in quiet; quiet increases tinnitus salience; SN paradoxical quiet sensitivity"),
    ("Alerting Response","'Ready' before stimulus.","Appropriate alerting; focused","Blunted; tinnitus-related disengagement; depression","Over-alerting; tinnitus hypervigilance; anxiety-tinnitus overlap"),
    ("Response Inhibition Cancel","Cancel circles on mixed-shape sheet.","<60 sec; ≤2 errors","Slow; tinnitus-related psychomotor slowing","Normal time but interrupted by ear-touching; tinnitus-directed behavior"),
    ("Attentional Fatigue","Letter cancellation 5 min; compare halves.","Drop <15%","Drop >40%; tinnitus-rapid attentional fatigue","20-35% drop; moderate; tinnitus attentional drain over time"),
  ],
  "phenotype_text":"Tinnitus phenotypes: (1) DISTRESS-DOMINANT — Limbic HYPER + DMN HYPER (catastrophizing); (2) ATTENTIONAL — Attention HYPO + SN HYPER (salience); (3) SOMATIC — SMN involvement (cervicogenic/TMJ); (4) COGNITIVE — CEN HYPO (brain fog overlap). Map to primary network dysfunction for FNON protocol selection.",
  "fnon_strategy":"Primary: Inhibit auditory cortex SN hyperactivation (cathodal/low-frequency over left temporal cortex or A1). Facilitate DLPFC for top-down auditory attention control. For distress-dominant: address limbic regulation via DLPFC facilitation. Combine with tinnitus retraining therapy and sound therapy.",
}

# ── 14. INSOMNIA ──────────────────────────────────────────────────────────────
INSOMNIA = {
  "full":"Insomnia","short":"Insomnia","slug":"insomnia",
  "adapted_note":"Includes hyperarousal markers, sleep-related cognitive tests, and Insomnia phenotype indicators",
  "dmn_rel":"Insomnia is driven by DMN HYPERACTIVITY at sleep onset: pre-sleep rumination, mind-racing, and failure to suppress default mode activity represents the core 'hyperarousal' model of insomnia. The brain cannot transition from active wakefulness (DMN suppression needed) to sleep-permissive states.",
  "cen_rel":"Insomnia causes CEN HYPOACTIVITY through chronic sleep deprivation: working memory, executive function, and cognitive flexibility are all impaired by insufficient sleep. Performance variability (sometimes fine, sometimes impaired) reflects day-to-day sleep quality fluctuation.",
  "sn_rel":"SN HYPERACTIVITY is the physiological correlate of insomnia hyperarousal: elevated cortical arousal, heightened interoceptive awareness (scanning for signs of not-sleeping), and hypervigilance to sleep-related cues. The bedroom becomes a threat-conditioned environment.",
  "smn_rel":"Insomnia SMN shows HYPOACTIVITY from fatigue and deconditioning but paradoxically HYPER arousal (inability to physically relax). The body cannot achieve the motor quiescence required for sleep onset. Daytime fatigue coexists with physiological hyperarousal.",
  "limbic_rel":"Limbic HYPERACTIVITY drives the anxiety-insomnia cycle: sleep anxiety, anticipatory arousal at bedtime, and catastrophic beliefs about consequences of not sleeping. Amygdala hyperreactivity increases with sleep loss, creating a reinforcing cycle.",
  "attn_rel":"Attention shows HYPO for sustained performance (daytime fatigue-driven) alongside HYPER-vigilance for sleep-related cues (clock-watching, body scanning). Attentional narrowing toward sleep-related stimuli at night reflects VAN HYPER in the pre-sleep period.",
  "dmn_tests":[
    ("Pre-Sleep Rumination Screen","Ask: 'When trying to sleep, do thoughts race or repeat in your mind?' Rate 0-10.","0-3; mind quiets at bedtime; DMN deactivates normally","0; no pre-sleep rumination; early insomnia or denial","≥7; racing thoughts at bedtime; DMN HYPER hallmark; cognitive hyperarousal insomnia"),
    ("Sleep-Related Autobiographical Memory","Ask 3 personal memories including 1 about a good night's sleep vs worst sleep. Note emotional content.","Specific memories; balanced; good sleep memories accessible","Sparse; poor recall; sleep deprivation impairs memory encoding","Good sleep memories inaccessible; 'can't remember sleeping well'; insomnia identity"),
    ("Self-Reference — Sleep Identity","Ask: 'Has insomnia become part of who you are?' Rate 0-10.","0-3; insomnia acknowledged but not identity-defining","0; denies sleep problems; avoidant; minimization","≥8; insomnia as core identity; catastrophizing; 'I am an insomniac'"),
    ("Mind Wandering Probe","2-min reading with 4 probes. Note: is off-task content sleep-related?","≤1/4 off-task; neutral mind-wandering","0/4; hypervigilant task focus; suppression or very mild insomnia","3-4/4 off-task; sleep-related worries during daytime; insomnia DMN generalization"),
    ("Sleep Catastrophizing Screen","Ask: 'What do you think will happen if you don't sleep well tonight?' Rate catastrophic thinking 0-10.","0-3; realistic; accepts occasional poor sleep; no catastrophizing","0; no concern; does not monitor sleep; early insomnia","≥7; catastrophic predictions; 'I'll fail tomorrow'; 'I'll get sick'; DMN HYPER insomnia"),
    ("Prospective Memory","5-min reminder task.","Independent recall within 5±1 min","Forgets; sleep deprivation impairs prospective memory","Recalls but effortfully; insomnia sleep-deprivation cognitive impact"),
    ("Default Narrative — Sleep History","Ask: 'Tell me about your sleep patterns over the past month.' Rate coherence and catastrophic framing.","Coherent; variability acknowledged; balanced coping","Minimizes sleep problem; avoidant; poor self-monitoring","Catastrophic framing; every night described as disaster; insomnia DMN narrative"),
  ],
  "cen_tests":[
    ("Digit Span Backward","Backward digit recall. Note performance variability.","Span ≥5","Span ≤3; sleep deprivation working memory failure","Span 4; borderline; sleep-deprivation cognitive interference"),
    ("Trail Making B","Alternating connections. Time and errors.","<90 sec; ≤2 errors","Slow >150 sec; sleep deprivation executive function failure","90-150 sec; moderate; insomnia sleep deprivation CEN slowing"),
    ("Verbal Fluency FAS","Letter fluency 60 sec per letter.","≥10 words/letter","<6; sleep deprivation verbal fluency failure","7-9 words; moderate; sleep deprivation cognitive drain"),
    ("Stroop Interference","Name ink color. Include sleep-related words (tired, rest, awake). Compare interference.","<45 sec; ≤3 errors; no content-specific slowing","Slow >90 sec; sleep deprivation executive failure","Selectively slowed on sleep-related words; hypervigilance to sleep content; insomnia CEN"),
    ("Serial 7 Subtraction","Backward 100-7 for 5 steps.","5/5 correct","≤2/5; sleep deprivation working memory failure","3-4/5; sleep deprivation CEN vulnerability"),
    ("Performance Variability Assessment","Administer digit span twice (beginning and end of assessment). Compare performance.","Consistent span across sessions; stable CEN function","Low across both; global sleep deprivation failure","Variable; better early then worse late; sleep deprivation attentional fatigue"),
    ("Go/No-Go Task","Tap for all except X. 20 letters.","≤2 commission; ≤2 omission","High omissions; sleep deprivation poor initiation","High commission errors; microsleep-related impulsive response; insomnia"),
  ],
  "sn_tests":[
    ("Physiological Hyperarousal Screen","Ask: 'Do you feel physically alert or wired even when tired?' Rate 0-10.","0-3; appropriate fatigue; no physiological over-activation","0; no arousal; pure sleep-scheduling insomnia or mild","≥7; physiological hyperarousal despite fatigue; SN HYPER insomnia hallmark"),
    ("Sleep Environment Threat Conditioning","Ask: 'Do you feel more awake/anxious when in your bedroom?' Rate 0-10.","0-3; bedroom associated with relaxation and sleep","0; no bedroom anxiety; behavioral insomnia early stage","≥7; bedroom triggers arousal; conditioned arousal; psychophysiological insomnia"),
    ("Interoception — Sleep Hyperawareness","Count heartbeats 25 sec. Note: do patients hypermonitor body for sleep readiness?","Accuracy >0.70; body focus does not trigger arousal","Low accuracy; poor body awareness; not hypervigilant","High accuracy AND reports monitoring body for sleep readiness; SN HYPER insomnia"),
    ("Clock-Watching Screen","Ask: 'How often do you check the time when trying to sleep?' Rate 0-10.","0-2; occasional time checking; not excessive","0; no time checking; non-hyperarousal pattern","≥7; repeated time checking; SN HYPER for sleep cues; insomnia perpetuating behavior"),
    ("Sensory Sensitivity at Night","Ask: 'Are you bothered by sounds, light, or temperature when trying to sleep?' Rate 0-10.","0-3; normal sleep environment sensitivity","0; no sensitivity; possible hyposomnia pattern","≥7; extreme sensory sensitivity at night; SN HYPER; insomnia sensory hyperarousal"),
    ("Error Monitoring","Sentence with 2 errors. Spontaneous correction.","Catches ≥1; calm correction","Misses both; sleep deprivation monitoring failure","Catches errors; perfectionism; 'can't let it go'; insomnia-anxiety perfectionism"),
    ("Orienting Response","Unexpected auditory tone mid-task.","Brief interruption; reorientation <3 sec","Absent orienting; sleep-deprived blunting","Exaggerated startle; hyperarousal; insomnia SN HYPER baseline arousal"),
  ],
  "smn_tests":[
    ("Physical Relaxation Assessment","Ask: 'Can you relax your body on demand?' Rate ability to achieve muscle relaxation 0-10.","≥6; able to relax on demand; progressive relaxation accessible","0; not monitored; unaware of body tension","≤3; unable to physically relax; somatic hyperarousal; insomnia SMN HYPER"),
    ("Muscle Tension Survey","Palpate/observe trapezius, jaw, and forearms. Rate visible tension 0-10 per region.","Tension ≤3 all regions; relaxed posture","0; hypotonic; poor somatic awareness","Tension ≥7 in ≥2 regions; somatic hyperarousal; insomnia SMN HYPER marker"),
    ("Finger Tapping Rate","Index tapping 10 sec each hand.","≥45 taps; regular","<35 taps; sleep deprivation psychomotor slowing","Normal rate but with fatigue markers; slow blink rate; sleepiness signs"),
    ("Grip Assessment — Fatigue vs Tension","Bilateral grip. Note: is grip weak (fatigue) or tense (hyperarousal)?","Firm; relaxed release; no fatigue","Weak; sleep deprivation motor fatigue; HYPO","Normal or above-normal grip; tense; cannot fully release; somatic hyperarousal"),
    ("Fine Motor Sequencing","Thumb-to-finger 10 sec.","≥4 sequences","≤2; sleep deprivation motor failure","Normal but with tremor; physiological hyperarousal in motor output"),
    ("Tandem Gait","Heel-to-toe 10 steps.","≤1 step off; consistent","Multiple failures; sleep deprivation balance and coordination","Normal; mild balance instability; sleep deprivation proprioception effect"),
    ("Romberg Test","Feet together eyes closed 30 sec.","Stable; minimal sway","Excessive sway; sleep deprivation balance HYPO","Stable but excessive muscle tension; cannot achieve restful stance; somatic HYPER"),
    ("Reaction Time","Hand lift when pen drops. 3 trials.","<250 ms; consistent across trials","≥450 ms; sleep deprivation psychomotor slowing","Highly variable (150-500 ms); microsleep-related lapses; insomnia reaction time variability"),
  ],
  "limbic_tests":[
    ("Sleep Anxiety Screen","Ask: 'Do you feel anxious as bedtime approaches?' Rate anticipatory sleep anxiety 0-10.","0-3; bedtime associated with relaxation","0; no sleep anxiety; behavioral or circadian insomnia type","≥7; severe sleep anxiety; Limbic HYPER; psychophysiological insomnia core feature"),
    ("Emotional Face Recognition","6 emotions from photos. Score /6.","≥5/6 correct","<4/6; sleep deprivation emotional recognition failure","4-5/6; mild; sleep deprivation reduces emotional face processing speed"),
    ("Affect Labeling","3 emotional scenarios. Name response.","Accurate differentiated labels","Flat; sleep deprivation blunts emotional processing","Appropriate labels but irritability/frustration colors all; sleep deprivation limbic HYPER"),
    ("Emotional Memory — Sleep Deprivation","10 emotional and 10 neutral words. Recall after 5-min distractor.","Emotional advantage; ≥7 recalled","Flat recall; sleep deprivation impairs emotional memory encoding","Normal emotional advantage but overall recall reduced; sleep deprivation effect"),
    ("Reward Sensitivity — Daytime Fatigue","Rate interest in 5 activities 0-10. Note: is low interest from fatigue or true anhedonia?","Average ≥5; motivation present; fatigue limits but doesn't eliminate","Average ≤2; severe daytime fatigue-driven anhedonia from insomnia","High interest; fatigue prevents engagement; preserved motivation with fatigue barrier"),
    ("Catastrophic Beliefs About Sleep","Ask: 'What is the worst thing about not sleeping?' Rate catastrophizing 0-10.","0-3; realistic; accepts occasional poor sleep consequences","0; no beliefs; not a thought-maintaining insomnia pattern","≥7; catastrophic; 'not sleeping will destroy my health/career/life'; limbic HYPER"),
    ("Affective Flexibility","Sad then happy memory. Mood shift and recovery.","Shift ≥3; recovery <3 min","Flat; no shift; sleep deprivation emotional blunting","Shift achieved but slow; sleep-deprived emotional regulation is sluggish; limbic efficiency HYPO"),
  ],
  "attn_tests":[
    ("Daytime Sleepiness Probe","Ask: 'How likely are you to fall asleep in these situations: watching TV, sitting quietly, in a car?'","Not likely in active situations; normal daytime alertness maintained","No sleepiness; possible circadian or behavioral insomnia type","Very likely in all situations; severe daytime sleepiness; sleep deprivation Epworth proxy"),
    ("Selective Attention — Sleep Deprived","Circle all A's in 20-item letter list. Note performance variability.","<60 sec; ≤3 errors","Slow; multiple omissions; sleep deprivation attention failure","Variable; normal then fails; microsleep-related attentional lapses"),
    ("Sustained Attention — Vigilance Probe","Count backward from 50 for 60 sec. Note: any lapses, pauses, or restarts.","≤2 errors; consistent pace","Stops; sleep deprivation sustained attention failure","2-4 lapses; microsleeps; insomnia sustained attention hallmark"),
    ("Divided Attention","Backward by 2s while naming objects.","≤2 errors each task","Abandons one; sleep deprivation attention resource failure","Performs both with 3-5 errors; sleep deprivation dual-task vulnerability"),
    ("Visual Search","Find star among 20 distractors.","<20 sec; systematic","Fails; sleep deprivation disorganized attention","20-40 sec; systematic but with 1-2 pauses; microsleep interruptions"),
    ("Clock-Face Attention — Sleep Hypervigilance","Show analog clock. Ask what time it is. Note: does patient stare at clock beyond task?","Reads time immediately; moves on","Cannot read clock; sleep deprivation visual attention failure","Reads correctly but stares; clock hypervigilance; SN-VAN HYPER for sleep cues"),
    ("Response Inhibition Cancel","Cancel circles on mixed-shape sheet.","<60 sec; ≤2 errors","Slow; sleep deprivation psychomotor slowing","Variable speed; fast then pauses; microsleep-related; insomnia"),
    ("Attentional Fatigue Index","Letter cancellation 5 min; compare halves.","Drop <15%","Drop >40%; sleep deprivation rapid attentional fatigue","20-35% drop; moderate; insomnia sleep deprivation core attentional feature"),
  ],
  "phenotype_text":"Insomnia phenotypes: (1) HYPERAROUSAL — SN HYPER + Limbic HYPER + DMN HYPER (psychophysiological insomnia); (2) COGNITIVE — CEN HYPO + Attention HYPO (sleep deprivation sequelae); (3) SOMATIC — SMN HYPER (physiological tension); (4) BEHAVIORAL — conditioned arousal (SN HYPER to bed cues). Map dominant phenotype for FNON protocol selection.",
  "fnon_strategy":"Primary: Inhibit cortical hyperarousal (cathodal/slow-wave transcranial stimulation over frontal cortex at sleep onset). Facilitate slow-wave sleep networks via cathodal frontal or delta-band entrainment. Address limbic hyperarousal (right DLPFC cathodal). Combine with CBT-I for hyperarousal phenotype.",
}



# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
ALL_CONDITIONS = [
    DEPRESSION, ANXIETY, ADHD, ALZHEIMERS,
    STROKE, TBI, CHRONIC_PAIN, PTSD, OCD,
    MS, ASD, LONG_COVID, TINNITUS, INSOMNIA,
]

if __name__ == "__main__":
    print("SOZO 6-Network Bedside Assessment - Partners Tier")
    print("=" * 55)
    errors = []
    for cond in ALL_CONDITIONS:
        slug = cond["slug"]
        short = cond["short"]
        print(f"\n[{short}] generating...")
        # Validate test counts
        for net, key, expected in [
            ("DMN","dmn_tests",7),("CEN","cen_tests",7),("SN","sn_tests",7),
            ("SMN","smn_tests",8),("Limbic","limbic_tests",7),("Attention","attn_tests",8),
        ]:
            actual = len(cond[key])
            if actual != expected:
                msg = f"  ⚠ {short} {net}: expected {expected} tests, got {actual}"
                print(str(msg).encode("ascii","replace").decode()); errors.append(msg)
        try:
            out = build_document(cond)
        except Exception as e:
            msg = f"  ✗ {short}: {e}"
            print(str(msg).encode("ascii","replace").decode()); errors.append(msg)

    print("\n" + "=" * 55)
    if errors:
        print(f"Completed with {len(errors)} warnings:")
        for e in errors: print(f"  {e}")
    else:
        print(f"All {len(ALL_CONDITIONS)} documents generated successfully.")
    print(f"\nOutput: {OUTPUT_BASE}")
