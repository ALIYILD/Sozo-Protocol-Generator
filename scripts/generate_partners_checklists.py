"""
Generate Clinical Examination Checklist (PARTNERS TIER) for all 14 conditions.
Exact replica of PD Partners template with FNON 6-network bedside assessment.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
PURPLE = RGBColor(0x7B, 0x2D, 0x8E)
BLACK  = RGBColor(0x00, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
CB = "\u2610"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    ex = tcPr.find(qn("w:shd"))
    if ex is not None: tcPr.remove(ex)
    tcPr.append(shd)

def add_header_row(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, "1B3A5C")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(9)

def add_data_row(table, row_idx, values, shade=False):
    row = table.rows[row_idx]
    for i, v in enumerate(values):
        cell = row.cells[i]
        if shade: set_cell_bg(cell, "EBF0F7")
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(v))
        run.font.size = Pt(9)

def make_table(doc, headers, data_rows):
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, headers)
    for i, row in enumerate(data_rows):
        add_data_row(table, i + 1, row, shade=(i % 2 == 0))
    return table

def add_h1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_h2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_bold_para(doc, bold_text, normal_text="", color=BLACK):
    p = doc.add_paragraph()
    r = p.add_run(bold_text); r.bold = True; r.font.color.rgb = color; r.font.size = Pt(10)
    if normal_text:
        r2 = p.add_run(normal_text); r2.font.size = Pt(10)
    return p

def build_document(cond):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(sec, attr, Inches(1))

    h1s = doc.styles["Heading 1"]
    h1s.font.size = Pt(14); h1s.font.bold = True; h1s.font.color.rgb = NAVY
    h2s = doc.styles["Heading 2"]
    h2s.font.size = Pt(12); h2s.font.bold = True; h2s.font.color.rgb = BLUE

    # ── Cover ──
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SOZO BRAIN CENTER \u2014 PARTNERS TIER")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Clinical Examination Checklist")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = BLUE

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Stage 4 \u2014 Patient Journey | FNON Network-Based Assessment")
    r.font.size = Pt(11); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\U0001f7e3 FNON + Evidence-Based Assessment \u2014 Includes Brain Network Analysis")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = PURPLE

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cond["display_name"].upper())
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

    doc.add_paragraph()

    # Patient info table
    make_table(doc, ["Field", "Details"], [
        ["Patient Name", ""],
        ["DOB", ""],
        ["Examining Doctor", ""],
        ["Date", ""],
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Scoring: "); r.bold = True; r.font.size = Pt(10)
    r2 = p.add_run("HYPO (1)  |  NORMAL (0)  |  HYPER (2)")
    r2.font.size = Pt(10)

    doc.add_paragraph()

    # ── Section A ──
    add_h1(doc, "Section A: 6-Network Bedside Assessment")
    add_bold_para(doc,
        "Network assessment maps patient symptoms to functional brain networks for FNON-based montage selection.",
        color=PURPLE)

    NET_HUBS = {
        "dmn":  ("Self-referential processing, autobiographical memory", "mPFC, PCC/Precuneus, Angular gyrus, Hippocampus", "Pz, P3/P4, Cz"),
        "cen":  ("Working memory, planning, cognitive control", "DLPFC, PPC/IPS", "F3/F4, P3/P4, Fz"),
        "sn":   ("Threat detection, interoception, DMN\u2194CEN switching", "Anterior Insula, dACC", "Fz, FCz, F3/F4"),
        "smn":  ("Motor output, somatosensation, proprioception", "M1, S1, SMA, Premotor", "C3/C4, Cz"),
        "limbic":("Affect generation/regulation, reward, stress", "OFC/mPFC, ACC, Amygdala", "Fp1/Fp2, AF3/AF4, Fz"),
        "attn": ("Sustained/selective attention, orienting", "DAN: IPS/FEF; VAN: rTPJ/rIFG", "P3/P4, PO7/PO8, F7/F8"),
    }
    NET_LABELS = [
        ("dmn",   "1. Default Mode Network (DMN)",          14, 7),
        ("cen",   "2. Central Executive Network (CEN/FPN)", 14, 7),
        ("sn",    "3. Salience Network (SN)",               14, 7),
        ("smn",   "4. Sensorimotor Network (SMN)",          16, 8),
        ("limbic","5. Limbic / Emotional Network",          14, 7),
        ("attn",  "6. Attention Networks (DAN / VAN)",      16, 8),
    ]
    NET_ABBREV = {
        "dmn":"DMN","cen":"CEN","sn":"SN","smn":"SMN","limbic":"Limbic","attn":"Attention"
    }

    for key, label, max_score, n_rows in NET_LABELS:
        add_h2(doc, f"{label} \u2014 Max Score: {max_score}")
        core, hubs, eeg = NET_HUBS[key]
        add_bold_para(doc, f"Core: {core} | Hubs: {hubs} | EEG sites: {eeg}")
        rel = cond["network_relevance"][key]
        add_bold_para(doc, f"{cond['short_name']} Relevance: ", rel, color=PURPLE)
        rows = cond["network_tests"][key]
        make_table(doc, ["Test", "Scoring Guide", "Score (0/1/2)"],
                   [[r[0], r[1], CB] for r in rows])
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(f"{NET_ABBREV[key]} Total: _____ / {max_score}")
        r.bold = True
        doc.add_paragraph()

    # ── Section B ──
    add_h1(doc, "Section B: Network Imbalance Profile")
    add_bold_para(doc,
        "Map network scores to identify dominant dysfunction pattern for FNON montage selection.",
        color=PURPLE)
    nibs = cond["nibs_targets"]
    make_table(doc, ["Network","Score","Max","Severity","State",f"{cond['short_name']}-Specific NIBS Target"], [
        ["1. DMN",             "", "14", f"{CB} Normal {CB} Mild {CB} Mod {CB} Severe", f"{CB} Hypo {CB} Hyper", nibs["dmn"]],
        ["2. CEN / FPN",       "", "14", f"{CB} Normal {CB} Mild {CB} Mod {CB} Severe", f"{CB} Hypo {CB} Hyper", nibs["cen"]],
        ["3. Salience (SN)",   "", "14", f"{CB} Normal {CB} Mild {CB} Mod {CB} Severe", f"{CB} Hypo {CB} Hyper", nibs["sn"]],
        ["4. SMN",             "", "16", f"{CB} Normal {CB} Mild {CB} Mod {CB} Severe", f"{CB} Hypo {CB} Hyper", nibs["smn"]],
        ["5. Limbic",          "", "14", f"{CB} Normal {CB} Mild {CB} Mod {CB} Severe", f"{CB} Hypo {CB} Hyper", nibs["limbic"]],
        ["6. Attention",       "", "16", f"{CB} Normal {CB} Mild {CB} Mod {CB} Severe", f"{CB} Hypo {CB} Hyper", nibs["attn"]],
        ["TOTAL",              "", "88", "0-17 Normal | 18-35 Mild | 36-53 Mod | 54+ Severe", "", ""],
    ])

    doc.add_paragraph()

    # ── Section C ──
    add_h1(doc, "Section C: Neurological Examination Summary")
    make_table(doc, ["Examination Domain", "Finding", "Details / Notes"],
               cond.get("neuro_rows", [
        ["Pyramidal tract signs",          f"{CB} Normal {CB} Abnormal", ""],
        ["Extrapyramidal signs",           f"{CB} Normal {CB} Abnormal", ""],
        ["Frontal release signs",          f"{CB} Normal {CB} Abnormal", ""],
        ["Cerebellar examination",         f"{CB} Normal {CB} Abnormal", ""],
        ["Cranial nerves / visual-oculomotor", f"{CB} Normal {CB} Abnormal", ""],
        ["Vestibular-oculomotor",          f"{CB} Normal {CB} Abnormal", ""],
        ["Sensory examination",            f"{CB} Normal {CB} Abnormal", ""],
        ["Deep tendon reflexes",           f"{CB} Normal {CB} Abnormal", ""],
    ]))

    doc.add_paragraph()

    # ── Section D ──
    add_h1(doc, "Section D: Phenotype Identification & FNON Strategy")
    add_bold_para(doc,
        "Map examination findings to FNON phenotype and network hypothesis for montage selection.",
        color=PURPLE)
    make_table(doc, ["Domain", "Assessment"], [
        ["Identified Phenotype",           cond["phenotype_options"]],
        ["Network Dysfunction Hypothesis", ""],
        ["Primary Dysfunctional Network",  ""],
        ["Secondary Dysfunctional Network",""],
        ["EEG Performed",                  f"{CB} Yes {CB} No {CB} Not available"],
        ["EEG Findings Summary",           ""],
        ["Preliminary FNON NIBS Strategy", ""],
        ["Proposed tDCS Protocol ID",      ""],
        ["Proposed TPS Protocol ID",       ""],
    ])

    doc.add_paragraph()
    make_table(doc, ["Role","Name","Signature","Date"],
               [["Examining Doctor","","",""]])
    return doc

from _partners_conditions import CONDITIONS

def main():
    base = os.path.join(os.path.dirname(__file__), "..", "outputs", "documents")
    for cond in CONDITIONS:
        slug = cond["slug"]
        out_dir = os.path.join(base, slug, "partners")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, f"Clinical_Exam_Checklist_Partners_{slug}.docx")
        doc = build_document(cond)
        doc.save(out_path)
        print(f"[OK] {slug}: {out_path}")

if __name__ == "__main__":
    main()
