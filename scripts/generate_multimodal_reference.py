"""
Generate SOZO Multi-Modality Protocol Reference DOCX for each of the 15 conditions.
Saves to: outputs/documents/{slug}/multimodal/Multimodal_Protocol_Reference_{slug}.docx
Run from project root: python scripts/generate_multimodal_reference.py
"""
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
PURPLE = RGBColor(0x7B, 0x2D, 0x8E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
DGRAY  = RGBColor(0x44, 0x44, 0x44)

BG = {
    "navy":     "1B3A5C",
    "blue":     "2E75B6",
    "purple":   "7B2D8E",
    "white":    "FFFFFF",
    "alt_row":  "F0F4F8",
    "lgray":    "F5F5F5",
    "lila":     "F0E6F6",
}

# ── Condition definitions ──────────────────────────────────────────────────────
CONDITIONS = [
    ("parkinsons",   "Parkinson's Disease"),
    ("depression",   "Depression (MDD)"),
    ("anxiety",      "Anxiety / GAD"),
    ("adhd",         "ADHD"),
    ("alzheimers",   "Alzheimer's Disease"),
    ("stroke_rehab", "Stroke Rehabilitation"),
    ("tbi",          "TBI / Concussion"),
    ("chronic_pain", "Chronic Pain"),
    ("ptsd",         "PTSD"),
    ("ocd",          "OCD"),
    ("ms",           "Multiple Sclerosis"),
    ("asd",          "Autism Spectrum Disorder"),
    ("long_covid",   "Long COVID"),
    ("tinnitus",     "Tinnitus"),
    ("insomnia",     "Insomnia"),
]

# Substrings to match (case-insensitive) for each slug
MATCH_TERMS = {
    "parkinsons":   ["parkinson"],
    "depression":   ["depression", "mdd"],
    "anxiety":      ["anxiety", "gad"],
    "adhd":         ["adhd", "attention"],
    "alzheimers":   ["alzheimer", "dementia", "mci"],
    "stroke_rehab": ["stroke"],
    "tbi":          ["tbi", "concussion"],
    "chronic_pain": ["pain"],
    "ptsd":         ["ptsd", "post-traumatic"],
    "ocd":          ["ocd", "obsessive"],
    "ms":           ["multiple sclerosis", " ms "],
    "asd":          ["autism", "asd"],
    "long_covid":   ["long covid", "covid"],
    "tinnitus":     ["tinnitus"],
    "insomnia":     ["insomnia", "sleep"],
}

# Modality display labels (JSON key → friendly label)
MODALITY_LABELS = {
    "TPS":        "TPS (Theta Burst)",
    "TMS_rTMS":   "TMS / rTMS",
    "tDCS":       "tDCS",
    "taVNS_tVNS": "taVNS / tVNS",
    "CES":        "CES",
    "tACS":       "tACS",
    "PBM":        "PBM (Photobiomodulation)",
    "PEMF":       "PEMF",
    "LIFU_tFUS":  "LIFU / tFUS",
    "tRNS":       "tRNS",
    "DBS":        "DBS",
}

# Matrix key aliases (conditions_matrix uses slightly different names)
MATRIX_ALIASES = {
    "TMS_rTMS":   "TMS",
    "taVNS_tVNS": "taVNS",
    "LIFU_tFUS":  "LIFU",
}

# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), sz)
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _set_table_borders(table, color="CCCCCC", sz="4"):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color, sz)


def _no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


def _set_margins(doc, top=2.0, bottom=2.0, left=2.0, right=2.0):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    _no_space_para(p)


def _add_run(para, text, bold=False, size=9, color=None, italic=False):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color
    if italic:
        r.font.italic = True
    return r


# ── Band heading (full-width coloured row) ─────────────────────────────────────

def _band(doc, text, bg="1B3A5C", fg=WHITE, size=13, bold=True):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("  " + text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = fg
    r.font.name = "Calibri"
    return tbl


# ── Callout box ────────────────────────────────────────────────────────────────

def _callout_box(doc, label, text, bg="7B2D8E", fg_label=WHITE, fg_text=WHITE):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg)
    _set_cell_borders(cell, bg, "6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    if label:
        rl = p.add_run(label + "  ")
        rl.bold = True
        rl.font.size = Pt(9)
        rl.font.color.rgb = fg_label
        rl.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = fg_text
    r.font.name = "Calibri"
    return tbl


# ── Data table ─────────────────────────────────────────────────────────────────

def _data_table(doc, headers, rows, header_bg="1B3A5C", alt_bg="F0F4F8",
                col_widths_cm=None, font_size=8.5):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")

    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(str(val) if val is not None else "")
            r.font.size = Pt(font_size)
            r.font.color.rgb = NAVY
            r.font.name = "Calibri"

    if col_widths_cm:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[i])

    return tbl


# ── Small note paragraph ───────────────────────────────────────────────────────

def _note_para(doc, label, text, size=8.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    rl = p.add_run(label + " ")
    rl.bold = True
    rl.italic = True
    rl.font.size = Pt(size)
    rl.font.color.rgb = DGRAY
    rl.font.name = "Calibri"
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(size)
    r.font.color.rgb = DGRAY
    r.font.name = "Calibri"
    return p


# ── Fix encoding artefacts from excel dump ─────────────────────────────────────

def _fix(text: str) -> str:
    if not isinstance(text, str):
        return str(text) if text is not None else ""
    # Common mojibake from Windows-1252 read as latin-1
    replacements = [
        ("\u00e2\u20ac\u201d", "\u2014"),   # â€" → —
        ("\u00e2\u20ac\u2122", "\u2019"),   # â€™ → '
        ("\u00e2\u20ac\u0153", "\u201c"),   # â€œ → "
        ("\u00e2\u20ac\u009d", "\u201d"),   # â€ → "
        ("\u00c3\u00a9", "\u00e9"),         # Ã© → é
        ("\u00c3\u00b3", "\u00f3"),         # Ã³ → ó
        ("\u00c3\u00bc", "\u00fc"),         # Ã¼ → ü
        ("\u00c3\u0097", "\u00d7"),         # Ã— → ×
        ("\u00c3\u00a0", "\u00e0"),         # Ã  → à
        ("\u00c2\u00b1", "\u00b1"),         # Â± → ±
        ("\u00c2\u00b5", "\u00b5"),         # Âµ → µ
    ]
    for bad, good in replacements:
        text = text.replace(bad, good)
    return text


# ── Condition matching ─────────────────────────────────────────────────────────

def _matches_condition(proto_condition: str, slug: str) -> bool:
    c = proto_condition.lower()
    for term in MATCH_TERMS[slug]:
        if term.lower() in c:
            return True
    return False


# ── Main generator ─────────────────────────────────────────────────────────────

def _matrix_paper_count(cm_entry: dict, mod_key: str) -> int:
    """Extract integer paper count from conditions_matrix entry for a modality."""
    matrix_key = MATRIX_ALIASES.get(mod_key, mod_key)
    val = cm_entry.get(matrix_key)
    if val is None:
        return 0
    try:
        return int(str(val).replace("+", "").replace(",", "").strip())
    except (ValueError, TypeError):
        return 0


def _find_cm_entry(conditions_matrix: dict, display_name: str, slug: str):
    """Find the conditions_matrix entry for this condition (fuzzy)."""
    # Try exact display_name first
    if display_name in conditions_matrix:
        return conditions_matrix[display_name]
    # Try slug-based terms
    terms = MATCH_TERMS[slug]
    for key, val in conditions_matrix.items():
        k_lower = key.lower()
        for term in terms:
            if term.lower() in k_lower:
                return val
    return {}


def generate_condition_doc(slug: str, display_name: str, data: dict, out_dir: Path):
    protocols_by_mod = data["protocols"]
    conditions_matrix = data.get("conditions_matrix", {})
    best_modalities = data.get("best_modalities", {})

    # Find best modality string
    best_mod_str = ""
    for key, val in best_modalities.items():
        if any(t.lower() in key.lower() for t in MATCH_TERMS[slug]):
            best_mod_str = val
            break
    if not best_mod_str:
        best_mod_str = "See evidence summary below"

    # Find conditions_matrix entry
    cm_entry = _find_cm_entry(conditions_matrix, display_name, slug)

    # Gather matching protocols per modality
    matched: dict[str, list] = {}
    for mod_key, protos in protocols_by_mod.items():
        hits = [p for p in protos if _matches_condition(p.get("condition", ""), slug)]
        if hits:
            matched[mod_key] = hits

    # Build evidence summary rows: Modality | Papers | Evidence Level
    evidence_rows = []
    for mod_key in MODALITY_LABELS:
        papers = _matrix_paper_count(cm_entry, mod_key)
        if papers == 0:
            # Also check matched protocols for any evidence_level
            protos_here = matched.get(mod_key, [])
            if not protos_here:
                continue
            papers_str = "—"
        else:
            papers_str = str(papers)

        protos_here = matched.get(mod_key, [])
        ev_levels = list(dict.fromkeys(
            p.get("evidence_level", "") for p in protos_here
            if p.get("evidence_level")
        ))
        ev_str = "; ".join(ev_levels) if ev_levels else "—"
        evidence_rows.append((mod_key, MODALITY_LABELS[mod_key], papers_str, ev_str, papers if isinstance(papers, int) else 0))

    # Sort by paper count desc
    evidence_rows.sort(key=lambda x: x[4], reverse=True)

    # ── Build document ─────────────────────────────────────────────────────────
    doc = Document()
    _set_margins(doc, top=1.8, bottom=1.8, left=1.8, right=1.8)

    # 1. Title band
    _band(doc,
          "SOZO Brain Center — Multi-Modality Neuromodulation Protocol Reference",
          bg="1B3A5C", fg=WHITE, size=14)
    _band(doc, display_name, bg="2E75B6", fg=WHITE, size=12)
    _spacer(doc, 8)

    # 2. Best modalities callout
    _callout_box(doc,
                 "Recommended First-Line Modalities:",
                 best_mod_str,
                 bg="7B2D8E", fg_label=WHITE, fg_text=WHITE)
    _spacer(doc, 10)

    # 3. Evidence Summary table
    _band(doc, "Evidence Summary by Modality", bg="1B3A5C", fg=WHITE, size=11)
    _spacer(doc, 4)

    if evidence_rows:
        ev_table_rows = [(r[1], r[2], r[3]) for r in evidence_rows]
        _data_table(doc,
                    ["Modality", "Papers (Literature Count)", "Evidence Level"],
                    ev_table_rows,
                    header_bg="1B3A5C",
                    alt_bg="F0F4F8",
                    col_widths_cm=[5.5, 4.5, 8.0],
                    font_size=8.5)
    else:
        p = doc.add_paragraph()
        _add_run(p, "No evidence data available for this condition.", size=9, color=DGRAY)

    _spacer(doc, 10)

    # 4. Protocol sections per modality
    _band(doc, "Detailed Protocol Specifications", bg="1B3A5C", fg=WHITE, size=11)
    _spacer(doc, 6)

    proto_headers = [
        "Brain Target", "EEG Position", "Protocol Summary",
        "Intensity", "Frequency", "Sessions",
        "Devices", "Regulatory Status",
    ]
    # Column widths in cm (total ~17.4 cm usable for A4 with 1.8cm margins each side)
    proto_col_widths = [2.5, 2.0, 4.5, 2.0, 1.8, 2.2, 2.5, 2.5]

    protocol_count = 0
    for mod_key in MODALITY_LABELS:
        protos = matched.get(mod_key, [])
        if not protos:
            continue

        mod_label = MODALITY_LABELS[mod_key]
        _band(doc, mod_label, bg="2E75B6", fg=WHITE, size=10)
        _spacer(doc, 3)

        # Build table rows
        table_rows = []
        side_effects_list = []
        key_refs_list = []

        for p in protos:
            row = [
                _fix(p.get("brain_target", "")),
                _fix(p.get("eeg_position", "")),
                _fix(p.get("protocol_summary", "")),
                _fix(p.get("intensity", "")),
                _fix(p.get("frequency", "")),
                _fix(p.get("total_sessions", "")),
                _fix(p.get("devices", "")),
                _fix(p.get("regulatory_status", "")),
            ]
            table_rows.append(row)
            se = _fix(p.get("side_effects", ""))
            if se:
                side_effects_list.append(se)
            kr = _fix(p.get("key_references", ""))
            if kr:
                key_refs_list.append(kr)
            protocol_count += 1

        _data_table(doc, proto_headers, table_rows,
                    header_bg="2E75B6",
                    alt_bg="F0F4F8",
                    col_widths_cm=proto_col_widths,
                    font_size=8.0)

        if side_effects_list:
            _note_para(doc, "Side Effects:", " | ".join(dict.fromkeys(side_effects_list)))

        if key_refs_list:
            _note_para(doc, "Key References:", " | ".join(dict.fromkeys(key_refs_list)))

        _spacer(doc, 8)

    if protocol_count == 0:
        p = doc.add_paragraph()
        _add_run(p, "No specific protocol entries found for this condition in the current dataset.",
                 size=9, color=DGRAY, italic=True)
        _spacer(doc, 6)

    # Save
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"Multimodal_Protocol_Reference_{slug}.docx"
    doc.save(str(out_path))
    return out_path, protocol_count


def main():
    json_path = ROOT / "scripts" / "excel_knowledge_dump.json"
    if not json_path.exists():
        print(f"ERROR: {json_path} not found", file=sys.stderr)
        sys.exit(1)

    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    print("Generating Multi-Modality Protocol Reference documents...")
    print(f"{'Condition':<30} {'Protocols':>10}  Path")
    print("-" * 80)

    created = []
    for slug, display_name in CONDITIONS:
        out_dir = ROOT / "outputs" / "documents" / slug / "multimodal"
        out_path, proto_count = generate_condition_doc(slug, display_name, data, out_dir)
        rel = out_path.relative_to(ROOT)
        print(f"{display_name:<30} {proto_count:>10}  {rel}")
        created.append(out_path)

    print()
    print(f"Done. {len(created)} DOCX files generated.")
    return created


if __name__ == "__main__":
    main()
