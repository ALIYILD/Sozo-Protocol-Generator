# -*- coding: utf-8 -*-
"""
Generate SOZO Multimodal Neuromodulation Comparison DOCX for all conditions.
Outputs: outputs/documents/{slug}/shared/Multimodal_Comparison_{SLUG}.docx
Run: python scripts/generate_multimodal_comparison.py
"""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from excel_protocol_data import MODALITY_PROTOCOLS, CONDITIONS_MATRIX, SLUG_TO_CONDITION

NAVY  = RGBColor(0x0D, 0x21, 0x37)
TEAL  = RGBColor(0x1A, 0x7A, 0x8A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

BG = {
    "navy": "0D2137", "teal": "1A7A8A", "white": "FFFFFF",
    "lteal": "E8F4F6", "lgray": "F5F5F5", "green": "D5F5E3",
    "lgreen": "EAFAF1", "mid": "D0E8EC", "lred": "FDECEA", "yellow": "FEF9E7",
}

ALL_MODALITIES = ["TPS", "TMS", "tDCS", "taVNS", "CES", "tACS", "PBM", "PEMF", "LIFU", "tRNS", "DBS"]

SLUG_KEYWORDS = {
    "alzheimers":   ["alzheimer"],
    "parkinsons":   ["parkinson"],
    "depression":   ["depression", "mdd"],
    "anxiety":      ["anxiety", "gad"],
    "adhd":         ["adhd"],
    "stroke_rehab": ["stroke"],
    "tbi":          ["tbi", "concussion"],
    "chronic_pain": ["chronic pain", "pain"],
    "ptsd":         ["ptsd"],
    "ocd":          ["ocd"],
    "ms":           ["multiple sclerosis"],
    "asd":          ["autism", "asd"],
    "long_covid":   ["long covid"],
    "tinnitus":     ["tinnitus"],
    "insomnia":     ["insomnia"],
}


def get_protocols_for_slug(slug):
    """Return {modality_key: (cond_name, protocol_dict)} for all modalities with data."""
    kws = SLUG_KEYWORDS.get(slug, [])
    if not kws:
        return {}
    result = {}
    for mk, conds in MODALITY_PROTOCOLS.items():
        for ck, p in conds.items():
            if any(kw in ck.lower() for kw in kws):
                if mk not in result:
                    result[mk] = (ck, p)
                break
    return result


# ── XML helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, h):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), h.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color="1A7A8A", sz="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), sz)
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _col_widths(table, wids):
    for row in table.rows:
        for cell, w in zip(row.cells, wids):
            cell.width = Cm(w)


def _set_margins(doc, top=1.9, bottom=1.9, left=1.9, right=1.9):
    s = doc.sections[0]
    s.top_margin = Cm(top); s.bottom_margin = Cm(bottom)
    s.left_margin = Cm(left); s.right_margin = Cm(right)


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)


def _heading_band(doc, text, bg="navy", fg=None, size=11):
    if fg is None:
        fg = WHITE
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    _set_cell_bg(cell, BG[bg])
    cell.width = Cm(17.8)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = fg


def _notice_box(doc, text, bg="lteal"):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    _set_cell_bg(cell, BG[bg])
    _set_cell_borders(cell, "1A7A8A", "6")
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY


def _data_table(doc, headers, rows, col_widths_cm, header_bg="teal", alt_bg="lgray"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        _set_cell_bg(cell, BG[header_bg])
        _set_cell_borders(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(h)
        run.bold = True; run.font.size = Pt(8); run.font.color.rgb = WHITE
    for i, dr in enumerate(rows):
        tr = table.rows[i + 1]
        bg = "white" if i % 2 == 0 else alt_bg
        for j, val in enumerate(dr):
            cell = tr.cells[j]
            _set_cell_bg(cell, BG.get(bg, "white"))
            _set_cell_borders(cell, "CCCCCC", "4")
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(str(val) if val is not None else "")
            run.font.size = Pt(8)
            run.font.color.rgb = NAVY
    _col_widths(table, col_widths_cm)
    return table


def _evidence_table(doc, matrix_data):
    best_str = matrix_data.get("best", "").lower()
    rows = []
    for mod in ALL_MODALITIES:
        cnt = matrix_data.get(mod)
        if cnt is None:
            continue
        is_best = mod.lower() in best_str
        rows.append([mod, str(cnt), "YES" if is_best else ""])
    if not rows:
        return
    _heading_band(doc, "Modality Evidence Overview (Published Paper Counts)", "teal", WHITE, 9)
    _spacer(doc, 3)
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(["Modality", "Published Papers", "Recommended?"]):
        cell = table.rows[0].cells[j]
        _set_cell_bg(cell, BG["teal"])
        _set_cell_borders(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE
    for i, (mod, cnt, best) in enumerate(rows):
        tr = table.rows[i + 1]
        bgk = "lgreen" if best == "YES" else ("white" if i % 2 == 0 else "lgray")
        for j, val in enumerate([mod, cnt, best]):
            cell = tr.cells[j]
            _set_cell_bg(cell, BG[bgk])
            _set_cell_borders(cell, "CCCCCC", "4")
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9); run.font.color.rgb = NAVY
            if j == 0:
                run.bold = True
            if best == "YES" and j == 2:
                run.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x8B, 0x4C)
    _col_widths(table, [5.5, 5.5, 4.5])


def build_comparison_document(slug):
    protocols = get_protocols_for_slug(slug)
    if len(protocols) < 2:
        return None

    matrix_cond = SLUG_TO_CONDITION.get(slug)
    matrix_data = CONDITIONS_MATRIX.get(matrix_cond, {}) if matrix_cond else {}
    display_name = matrix_cond or slug.replace("_", " ").title()

    doc = Document()
    _set_margins(doc, 1.9, 1.9, 1.9, 1.9)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)

    # Title
    _heading_band(doc, "SOZO BRAIN CENTER", "navy", WHITE, 14)
    _heading_band(doc, "Multimodal Neuromodulation Comparison", "teal", WHITE, 12)
    _heading_band(doc, display_name.upper(), "navy", WHITE, 11)
    _spacer(doc, 8)
    _notice_box(doc,
        f"All available modalities for {display_name}. "
        "Source: SOZO Master Neuromodulation Protocol Database, April 2026. "
        "For use by qualified SOZO practitioners under Doctor supervision only.",
        "lteal")
    _spacer(doc, 10)

    # Section 1: Evidence Overview
    _heading_band(doc, "SECTION 1 - EVIDENCE OVERVIEW BY MODALITY", "navy", WHITE, 10)
    _spacer(doc, 4)
    if matrix_data:
        _evidence_table(doc, matrix_data)
        if matrix_data.get("best"):
            _spacer(doc, 4)
            _notice_box(doc, f"Recommended modality/modalities: {matrix_data['best']}", "green")
    else:
        _notice_box(doc, "Evidence matrix data not available for this condition.", "yellow")
    _spacer(doc, 10)
    _page_break(doc)

    # Section 2: Protocol Parameters
    _heading_band(doc, "SECTION 2 - PROTOCOL PARAMETERS COMPARISON", "navy", WHITE, 10)
    _spacer(doc, 4)
    proto_rows = []
    for mk in ALL_MODALITIES:
        if mk not in protocols:
            continue
        cn, p = protocols[mk]
        dev = p.get("devices", "")
        dev = dev[:40] + ("..." if len(dev) > 40 else "")
        proto_rows.append([mk, p.get("brain_target", ""), p.get("eeg_position", ""),
            p.get("intensity", ""), p.get("frequency", ""), p.get("total_sessions", ""),
            dev, p.get("regulatory_status", "")])
    if proto_rows:
        _data_table(doc,
            headers=["Modality", "Brain Target", "EEG Position", "Intensity", "Frequency",
                     "Sessions", "Device(s)", "Regulatory"],
            rows=proto_rows,
            col_widths_cm=[2.0, 3.2, 2.2, 2.0, 1.8, 1.8, 3.0, 2.5])
    _spacer(doc, 8)

    _heading_band(doc, "Evidence Levels by Modality", "teal", WHITE, 9)
    _spacer(doc, 3)
    ev_rows = []
    for mk in ALL_MODALITIES:
        if mk not in protocols:
            continue
        _, p = protocols[mk]
        ev_rows.append([mk, p.get("evidence_level", ""), p.get("literature_count", "")])
    if ev_rows:
        _data_table(doc,
            headers=["Modality", "Evidence Level", "Literature Count"],
            rows=ev_rows,
            col_widths_cm=[2.5, 11.5, 2.5])
    _spacer(doc, 10)
    _page_break(doc)

    # Section 3: Key References
    _heading_band(doc, "SECTION 3 - KEY REFERENCES BY MODALITY", "navy", WHITE, 10)
    _spacer(doc, 4)
    for mk in ALL_MODALITIES:
        if mk not in protocols:
            continue
        _, p = protocols[mk]
        refs = p.get("key_references", "")
        if refs:
            _heading_band(doc, mk, "teal", WHITE, 9)
            _spacer(doc, 2)
            _notice_box(doc, refs, "lteal")
            _spacer(doc, 4)

    # Section 4: Side Effects
    _heading_band(doc, "SECTION 4 - SIDE EFFECTS AND SAFETY BY MODALITY", "navy", WHITE, 10)
    _spacer(doc, 4)
    se_rows = []
    for mk in ALL_MODALITIES:
        if mk not in protocols:
            continue
        _, p = protocols[mk]
        se_rows.append([mk, p.get("side_effects", "")])
    if se_rows:
        _data_table(doc,
            headers=["Modality", "Side Effects & Safety Profile"],
            rows=se_rows,
            col_widths_cm=[2.5, 15.0])

    out_dir = ROOT / "outputs" / "documents" / slug / "shared"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"Multimodal_Comparison_{slug.upper()}.docx"
    doc.save(str(out_path))
    return out_path


def main():
    slugs = list(SLUG_TO_CONDITION.keys())
    if len(sys.argv) > 1:
        slugs = sys.argv[1:]
    total = len(slugs)
    generated = 0
    for i, slug in enumerate(slugs, 1):
        print(f"[{i}/{total}] {slug} ...", end=" ", flush=True)
        try:
            path = build_comparison_document(slug)
            if path is None:
                print("SKIP (< 2 modalities)")
            else:
                kb = path.stat().st_size // 1024
                print(f"OK -> {path.relative_to(ROOT)} ({kb} KB)")
                generated += 1
        except Exception as e:
            import traceback
            print(f"FAILED - {e}")
            traceback.print_exc()
    print(f"\nDone. Generated: {generated}")


if __name__ == "__main__":
    main()
