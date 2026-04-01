"""
Generate SOZO Partners-Tier ALL-IN-ONE Protocol Compendium for all 14 conditions.
Saves to: outputs/documents/{slug}/partners/All_In_One_Partner_{slug}.docx
Run from project root: python scripts/generate_partners_all_in_one.py
"""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from partners_all_in_one_data import PARTNERS_CONDITIONS

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
PURPLE = RGBColor(0x7B, 0x2D, 0x8E)
RED    = RGBColor(0xCC, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
DGRAY  = RGBColor(0x44, 0x44, 0x44)

BG = {
    "navy":      "1B3A5C",
    "blue":      "2E75B6",
    "purple":    "7B2D8E",
    "white":     "FFFFFF",
    "lila":      "F0E6F6",
    "fnon_bg":   "FAF5FF",
    "task_bg":   "EBF5FB",
    "device_bg": "D6E4F0",
    "s_green":   "E8F5E9",
    "o_blue":    "E3F2FD",
    "z_purple":  "F3E5F5",
    "o_orange":  "FFF3E0",
    "alt_row":   "F0F4F8",
    "red_warn":  "FFEBEE",
    "lgray":     "F5F5F5",
    "mid_blue":  "D6E4F0",
}

# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color="1B3A5C", sz="6"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), sz)
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _set_table_borders(table, color="CCCCCC", sz="4"):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color, sz)


def _no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


def _col_widths(table, widths_cm: list):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _set_margins(doc, top=2.54, bottom=2.54, left=2.54, right=2.54):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def _page_break(doc):
    doc.add_page_break()


def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    _no_space_para(p)


def _hr(doc, color="1B3A5C"):
    p = doc.add_paragraph()
    _no_space_para(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Text helpers ──────────────────────────────────────────────────────────────

def _add_run(para, text, bold=False, size=9, color=None, italic=False):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color
    if italic:
        r.font.italic = True
    return r


def _para(doc, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    _add_run(p, text, bold=bold, size=size, color=color or NAVY)
    return p


# ── Band headings ─────────────────────────────────────────────────────────────

def _band(doc, text, bg_key="navy", fg=WHITE, size=13, left_pad=True):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, BG[bg_key])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(("  " if left_pad else "") + text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = fg
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    return tbl


def _colored_para(doc, text, color=NAVY, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


# ── Notice boxes ──────────────────────────────────────────────────────────────

def _notice_box(doc, text, bg_key="lila", label="", label_color=None):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, BG[bg_key])
    _set_cell_borders(cell, "BBBBBB", "4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    if label:
        lc = label_color or PURPLE
        rl = p.add_run(label + "  ")
        rl.bold = True
        rl.font.size = Pt(9)
        rl.font.color.rgb = lc
        rl.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    return tbl


def _warning_box(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, BG["red_warn"])
    _set_cell_borders(cell, "CC0000", "6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    rl = p.add_run("⚠ WARNING  ")
    rl.bold = True
    rl.font.size = Pt(9)
    rl.font.color.rgb = RED
    rl.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.color.rgb = RED
    r.font.name = "Calibri"
    return tbl


# ── Generic data table ────────────────────────────────────────────────────────

def _data_table(doc, headers, rows, header_bg="navy", alt_bg="alt_row",
                col_widths=None, font_size=8.5):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")

    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, BG[header_bg])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "white"
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            _set_cell_bg(cell, BG[bg])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(str(val))
            r.font.size = Pt(font_size)
            r.font.color.rgb = NAVY
            r.font.name = "Calibri"

    if col_widths:
        _col_widths(tbl, col_widths)

    return tbl


# ── Protocol card (1×2) ───────────────────────────────────────────────────────

def _protocol_card(doc, left_text, right_text, left_bg="navy", right_bg="white",
                   left_w=5.0, right_w=13.0):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")

    cl = tbl.cell(0, 0)
    cr = tbl.cell(0, 1)
    _set_cell_bg(cl, BG[left_bg])
    _set_cell_bg(cr, BG[right_bg])
    cl.width = Cm(left_w)
    cr.width = Cm(right_w)
    cl.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cr.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Left cell content
    pl = cl.paragraphs[0]
    pl.paragraph_format.space_before = Pt(4)
    pl.paragraph_format.space_after  = Pt(4)
    for line in left_text.split("\n"):
        if line == left_text.split("\n")[0]:
            p = pl
        else:
            p = cl.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
        bold_line = line.startswith("**") and line.endswith("**")
        txt = line.strip("* ")
        r = p.add_run(txt)
        r.bold = bold_line or (line == left_text.split("\n")[0])
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE if left_bg in ("navy", "blue", "purple") else NAVY
        r.font.name = "Calibri"

    # Right cell content
    pr = cr.paragraphs[0]
    pr.paragraph_format.space_before = Pt(4)
    pr.paragraph_format.space_after  = Pt(4)
    for i, line in enumerate(right_text.split("\n")):
        if i == 0:
            p = pr
        else:
            p = cr.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
        is_label = line.endswith(":") or (line and line[0].isupper() and len(line.split()) <= 3 and ":" not in line and i == 0)
        r = p.add_run(line)
        r.bold = is_label
        r.font.size = Pt(8.5)
        r.font.color.rgb = NAVY
        r.font.name = "Calibri"

    return tbl


# ── S-O-Z-O bar ───────────────────────────────────────────────────────────────

def _sozo_bar(doc, sozo):
    labels = [("S", "S — Stabilize", "s_green"), ("O", "O — Optimize", "o_blue"),
              ("Z", "Z — Zone Target", "z_purple"), ("O2", "O — Outcome", "o_orange")]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")
    for i, (key, label, bg_key) in enumerate(labels):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, BG[bg_key])
        cell.width = Cm(4.5)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        rl = p.add_run(label + "\n")
        rl.bold = True
        rl.font.size = Pt(8.5)
        rl.font.color.rgb = NAVY
        rl.font.name = "Calibri"
        r = p.add_run(sozo.get(key, ""))
        r.font.size = Pt(8)
        r.font.color.rgb = NAVY
        r.font.name = "Calibri"
    return tbl


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION BUILDERS
# ═══════════════════════════════════════════════════════════════════════════════

def _build_cover(doc, cond):
    _spacer(doc, 24)
    _colored_para(doc, "SOZO BRAIN CENTER", NAVY, bold=True, size=20,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    _colored_para(doc, "CYPRUS", BLUE, bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    _spacer(doc, 8)
    _colored_para(doc, f"{cond['short'].upper()} ALL-IN-ONE PROTOCOL COMPENDIUM",
                  NAVY, bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    _colored_para(doc, "Functional Network-Oriented Neuromodulation",
                  PURPLE, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _spacer(doc, 8)
    _colored_para(doc, cond["display_name"], NAVY, bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    _colored_para(doc, "Partners Tier  |  Version 1.0  |  March 2026",
                  PURPLE, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _spacer(doc, 12)
    _hr(doc, "7B2D8E")
    _spacer(doc, 8)

    # Protocol coverage summary
    _data_table(doc,
        headers=["Section", "Protocols Included", "Tier"],
        rows=[
            ["Section 1 — Classic TPS (T1-T5)", "5 protocols · NEUROLITH®", "All Tiers"],
            ["Section 2 — FNON TPS (FT1-FT9)", "9 network-oriented TPS protocols", "Partners Only"],
            ["Section 3 — Newronika HDCkit tDCS (C1-C8)", "8 in-clinic tDCS protocols", "All Tiers"],
            ["Section 4 — FNON tDCS (F1-F6)", "6 network-oriented tDCS protocols", "Partners Only"],
            ["Section 5 — PlatoScience tDCS (C1-PS–C8-PS)", "8 home-based tDCS protocols", "All Tiers"],
            ["Section 6 — FNON PlatoScience (FP1-FP6)", "6 FNON home-based protocols", "Partners Only"],
            ["Section 7 — Multimodal S-O-Z-O (P1-P9)", "9 phenotype-specific full protocols", "Partners Only"],
            ["Section 8 — Safety & Monitoring", "Side effects, grading, contraindications", "All Tiers"],
            ["Section 9 — Follow-Up & Decision-Making", "Reassessment schedule & algorithm", "All Tiers"],
        ],
        header_bg="navy", alt_bg="alt_row",
        col_widths=[9.0, 6.5, 2.5],
        font_size=8.5,
    )
    _spacer(doc, 8)

    # FNON network overview
    _band(doc, "FNON Framework — Six Target Networks", bg_key="purple", size=11)
    _spacer(doc, 4)
    _data_table(doc,
        headers=["Network", "Abbreviation", "Key Nodes", "Clinical Role"],
        rows=[
            ["Default Mode Network", "DMN", "mPFC, PCC, angular gyrus", "Self-referential processing, rumination, autobiographical memory"],
            ["Central Executive Network", "CEN", "DLPFC, PPC, lateral PFC", "Working memory, cognitive control, goal-directed behavior"],
            ["Salience Network", "SN", "ACC, anterior insula, amygdala", "Threat detection, network switching, emotional tagging"],
            ["Sensorimotor Network", "SMN", "M1, S1, SMA, cerebellum", "Motor output, proprioception, sensorimotor integration"],
            ["Limbic Network", "LIMBIC", "Hippocampus, amygdala, vmPFC, OFC", "Affect regulation, reward processing, memory consolidation"],
            ["Attention Networks", "DAN/VAN", "FEF, IPS, TPJ, MFG", "Sustained attention, vigilance, attentional orienting"],
        ],
        header_bg="purple", alt_bg="lila",
        col_widths=[4.5, 2.5, 4.5, 6.5],
        font_size=8.5,
    )
    _page_break(doc)


def _build_toc(doc, cond):
    _band(doc, "TABLE OF CONTENTS", bg_key="navy", size=12)
    _spacer(doc, 6)
    toc_items = [
        "Cover Page & Protocol Coverage Summary",
        "FNON Framework Overview",
        "SECTION 1:  Classic TPS Protocols (T1-T5)",
        "SECTION 2:  FNON TPS Protocols — Network-Oriented (FT1-FT9)  [PARTNERS ONLY]",
        "SECTION 3:  Newronika HDCkit tDCS Protocols (C1-C8)",
        "SECTION 4:  FNON tDCS Protocols — Network-Oriented (F1-F6)  [PARTNERS ONLY]",
        "SECTION 5:  PlatoScience tDCS Protocols (C1-PS–C8-PS)",
        "SECTION 6:  FNON PlatoScience Protocols (FP1-FP6)  [PARTNERS ONLY]",
        "SECTION 7:  Multimodal S-O-Z-O Protocols (P1-P9)  [PARTNERS ONLY]",
        "SECTION 8:  Safety, Side Effects & Monitoring",
        "SECTION 9:  Follow-Up & Decision-Making",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        bullet = "  ▸  " if "SECTION" in item else "      "
        r = p.add_run(bullet + item)
        r.font.size = Pt(9.5)
        r.font.color.rgb = PURPLE if "PARTNERS ONLY" in item else NAVY
        r.font.bold = "SECTION" in item
        r.font.name = "Calibri"
    _page_break(doc)


def _build_tps_section(doc, cond):
    _band(doc, "SECTION 1 — CLASSIC TPS PROTOCOLS (T1-T5)", bg_key="navy", size=13)
    _spacer(doc, 4)
    _warning_box(doc, cond.get("tps_warning",
        f"TPS use in {cond['display_name']} is INVESTIGATIONAL / OFF-LABEL in most jurisdictions. "
        "NEUROLITH® is CE-marked for musculoskeletal indications. Use in neurological/psychiatric "
        "conditions is supported by emerging clinical evidence only. Informed consent mandatory."))
    _spacer(doc, 4)
    _notice_box(doc,
        "Device: NEUROLITH® (Storz Medical AG). Focused ultrasound pulse propagation. "
        "Parameters: 0.20–0.25 mJ/mm² energy density · 800–1,200 pulses/session standard · "
        "3–5 Hz repetition rate · Scalp coupling gel required. Doctor configures all parameters.",
        bg_key="device_bg", label="NOTE")
    _spacer(doc, 6)

    for proto in cond.get("tps", []):
        _band(doc,
              f"{proto['code']}  ·  {proto['symptom']}  [{proto.get('evidence','Emerging')}]",
              bg_key="blue", size=10)
        _spacer(doc, 3)
        left = f"{proto['code']}\nScalp Positions\n\nTargets:\n" + "\n".join(proto.get("targets", []))
        right = (f"Symptom Target: {proto['symptom']}\n"
                 f"ROIs: {', '.join(proto.get('targets', []))}\n"
                 f"Parameters: {proto.get('params','')}\n"
                 f"Evidence Level: {proto.get('evidence','Emerging')}\n"
                 f"Rationale: {proto.get('rationale','')}")
        _protocol_card(doc, left, right, left_bg="navy", right_bg="white")
        _spacer(doc, 8)

    _page_break(doc)


def _build_fnon_tps_section(doc, cond):
    _band(doc, "SECTION 2 — FNON TPS PROTOCOLS — NETWORK-ORIENTED (FT1-FT9)", bg_key="purple", size=13)
    _spacer(doc, 4)
    _warning_box(doc, cond.get("tps_warning",
        f"FNON TPS protocols for {cond['display_name']} are INVESTIGATIONAL. "
        "Network-oriented targeting is based on functional connectivity research and should be used "
        "only by SOZO-trained clinicians with explicit Doctor supervision and patient consent."))
    _spacer(doc, 4)
    _notice_box(doc, "✦ PARTNERS ONLY  —  These network-oriented TPS protocols integrate FNON phenotyping "
                "with targeted neuromodulation. Not available in Fellow tier.",
                bg_key="lila", label="✦ PARTNERS ONLY", label_color=PURPLE)
    _spacer(doc, 6)

    for proto in cond.get("fnon_tps", []):
        _band(doc,
              f"{proto['code']}  ·  {proto['name']}  ·  Network: {proto.get('network','')}",
              bg_key="purple", size=10)
        _spacer(doc, 3)
        left = (f"{proto['code']}\n{proto['name']}\n\n"
                f"Network:\n{proto.get('network','')}\n\n"
                f"Phenotype:\n{proto.get('phenotype','')}")
        right = (f"Network Target: {proto.get('network','')}\n"
                 f"Phenotype: {proto.get('phenotype','')}\n"
                 f"Scalp Positions: {', '.join(proto.get('targets', []))}\n"
                 f"Parameters: {proto.get('params','')}\n"
                 f"S-O-Z-O Integration: {proto.get('sozo_integration','')}\n"
                 f"Rationale: {proto.get('rationale','')}")
        _protocol_card(doc, left, right, left_bg="purple", right_bg="fnon_bg")
        _spacer(doc, 8)

    _page_break(doc)


def _build_tdcs_section(doc, cond):
    _band(doc, "SECTION 3 — NEWRONIKA HDCkit tDCS PROTOCOLS (C1-C8)", bg_key="blue", size=13)
    _spacer(doc, 4)
    _data_table(doc,
        headers=["Device", "Specification"],
        rows=[
            ["Newronika HDCkit", "CE Class IIa (CE 0476) · 2-channel anodal/cathodal tDCS"],
            ["Current", "1–2 mA (standard 2 mA unless noted)"],
            ["Electrode", "25 cm² (5×5 cm) saline-soaked sponge pads"],
            ["Ramp", "30 s ramp-up, 30 s ramp-down (automatic)"],
            ["Session", "20–30 min per protocol"],
            ["Programming", "Doctor programs via HDCprog software — patient cannot modify"],
        ],
        header_bg="blue", alt_bg="device_bg",
        col_widths=[5.0, 13.0], font_size=8.5,
    )
    _spacer(doc, 6)

    for proto in cond.get("tdcs", []):
        _band(doc, f"{proto['code']}  ·  {proto['symptom']}", bg_key="blue", size=10)
        _spacer(doc, 3)
        anodes   = ", ".join(proto.get("anodes", []))
        cathodes = ", ".join(proto.get("cathodes", []))
        left = (f"{proto['code']}\n10-20 Positions\n\n"
                f"Anode (+):\n{anodes}\n\n"
                f"Cathode (−):\n{cathodes}")
        right = (f"Symptom Target: {proto['symptom']}\n"
                 f"Anode (+): {anodes}\n"
                 f"Cathode (−): {cathodes}\n"
                 f"Parameters: {proto.get('params','')}\n"
                 f"Rationale: {proto.get('rationale','')}")
        _protocol_card(doc, left, right, left_bg="navy", right_bg="white")
        _spacer(doc, 8)

    # Safety table
    _band(doc, "tDCS Safety & Contraindications", bg_key="navy", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Category", "tDCS Consideration", "Action"],
        rows=[
            ["Metal cranial implants", "ABSOLUTE contraindication", "Do not treat"],
            ["Pacemaker / ICD", "ABSOLUTE contraindication", "Consult cardiologist first"],
            ["Active epilepsy", "ABSOLUTE contraindication", "Medical clearance required"],
            ["Pregnancy", "RELATIVE — frontal montages", "Doctor risk-benefit decision"],
            ["Open scalp wounds", "ABSOLUTE at affected site", "Avoid electrode placement over wound"],
            ["Photosensitivity (frontal)", "RELATIVE", "Monitor for phosphenes; reposition"],
        ],
        header_bg="navy", alt_bg="alt_row",
        col_widths=[5.5, 7.0, 5.5], font_size=8.5,
    )
    _page_break(doc)


def _build_fnon_tdcs_section(doc, cond):
    _band(doc, "SECTION 4 — FNON tDCS PROTOCOLS — NETWORK-ORIENTED (F1-F6)", bg_key="purple", size=13)
    _spacer(doc, 4)
    _notice_box(doc, "✦ PARTNERS ONLY  —  FNON tDCS protocols apply Functional Network-Oriented "
                "Neuromodulation principles to tDCS montage selection. Each protocol targets a "
                "specific dysfunctional network rather than a single symptom.",
                bg_key="lila", label="✦ PARTNERS ONLY", label_color=PURPLE)
    _spacer(doc, 6)

    for proto in cond.get("fnon_tdcs", []):
        _band(doc,
              f"{proto['code']}  ·  {proto['name']}  ·  {proto.get('network','')}",
              bg_key="purple", size=10)
        _spacer(doc, 3)
        anodes   = ", ".join(proto.get("anodes", []))
        cathodes = ", ".join(proto.get("cathodes", []))
        left = (f"{proto['code']}\n{proto['name']}\n\n"
                f"Network:\n{proto.get('network','')}\n\n"
                f"Anode (+):\n{anodes}\n\n"
                f"Cathode (−):\n{cathodes}")
        right = (f"Network Target: {proto.get('network','')}\n"
                 f"Clinical Focus: {proto['name']}\n"
                 f"Anode (+): {anodes}\n"
                 f"Cathode (−): {cathodes}\n"
                 f"Parameters: {proto.get('params','')}\n"
                 f"Rationale: {proto.get('rationale','')}")
        _protocol_card(doc, left, right, left_bg="purple", right_bg="fnon_bg")
        _spacer(doc, 8)

    # FNON Classification table
    _band(doc, "FNON Classification System — Network Dysfunction & Protocol Selection", bg_key="purple", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Network", "Dysfunction Role", "Clinical Presentation", "FNON Protocol"],
        rows=cond.get("fnon_classification", [
            ["CEN", "Executive/WM deficits", "Cognitive impairment, poor planning", "F1"],
            ["DMN", "Hyperconnectivity", "Rumination, self-focused ideation", "F2"],
            ["SN", "Dysregulation", "Emotional reactivity, threat overdetection", "F3"],
            ["SMN", "Disruption", "Motor/sensory symptoms", "F4"],
            ["LIMBIC", "Dysregulation", "Mood, reward, memory dysfunction", "F5"],
            ["ATTENTION", "Hypofunction", "Inattention, vigilance deficits", "F6"],
        ]),
        header_bg="purple", alt_bg="lila",
        col_widths=[3.0, 4.0, 5.5, 5.5], font_size=8.5,
    )
    _spacer(doc, 6)

    # FNON Montage Selection
    _band(doc, "FNON Montage Selection Guide", bg_key="purple", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Clinical Profile", "Recommended Montage", "Network Target"],
        rows=cond.get("fnon_montage_selection", [
            ["Cognitive dominant", "F3 anode, Pz cathode", "CEN upregulation"],
            ["Mood dominant", "F3 anode, F4 cathode (asymmetric)", "LIMBIC + CEN balance"],
            ["Anxiety dominant", "vmPFC (Fz) anode, Oz cathode", "SN downregulation"],
            ["Motor dominant", "C3/C4 anode, Fp1/Fp2 cathode", "SMN facilitation"],
            ["Attention dominant", "F4 anode, Fp1 cathode", "ATTENTION upregulation"],
            ["Mixed/complex", "Bilateral DLPFC (F3+F4)", "CEN + SN dual-target"],
        ]),
        header_bg="purple", alt_bg="lila",
        col_widths=[5.0, 6.0, 7.0], font_size=8.5,
    )
    _page_break(doc)


def _build_plato_section(doc, cond):
    _band(doc, "SECTION 5 — PLATOSCIENCE tDCS PROTOCOLS (C1-PS–C8-PS)", bg_key="blue", size=13)
    _spacer(doc, 4)
    _data_table(doc,
        headers=["Device", "Specification"],
        rows=[
            ["PlatoWork Headset", "Single-channel tDCS · 1.2–1.6 mA clinical range"],
            ["Placement", "Snap-in variants A–G (see placement guide below)"],
            ["Duration", "20 min fixed per session · 30 s ramp-up/ramp-down"],
            ["Control", "PlatoApp (iOS/Android) — Doctor prescribes program"],
            ["Constraint", "Single-channel ONLY — bilateral montages not available"],
        ],
        header_bg="blue", alt_bg="device_bg",
        col_widths=[5.0, 13.0], font_size=8.5,
    )
    _spacer(doc, 5)

    # PlatoApp programs
    _band(doc, "PlatoApp Built-in Programs", bg_key="blue", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Program", "tDCS Effect", "Indication"],
        rows=cond.get("plato_programs", [
            ["Focus", "Anodal L-DLPFC stimulation", "Attention, executive function"],
            ["Memory", "Anodal temporal stimulation", "Memory encoding/retrieval"],
            ["Mood+", "L→R asymmetric prefrontal", "Low mood, motivation"],
            ["Calm", "Cathodal prefrontal inhibition", "Anxiety, hyperarousal"],
            ["Energy", "Bilateral frontal upregulation", "Fatigue, mental energy"],
            ["Sleep", "Bilateral frontal downregulation", "Insomnia, sleep onset"],
            ["Pain Relief", "Anodal M1 stimulation", "Chronic pain, sensitivity"],
            ["Performance", "DLPFC bilateral activation", "Cognitive performance"],
            ["Recovery", "Cathodal motor cortex", "Post-exertion recovery"],
        ]),
        header_bg="blue", alt_bg="alt_row",
        col_widths=[3.5, 6.5, 8.0], font_size=8.5,
    )
    _spacer(doc, 5)

    # PlatoWork variants
    _band(doc, "PlatoWork Placement Variants", bg_key="blue", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Variant", "EEG Position", "Electrode Coverage", "Target Region", "Notes"],
        rows=[
            ["A", "Fp1-F3", "Left prefrontal", "L-DLPFC", "Standard mood/cognition"],
            ["B", "Fp2-F4", "Right prefrontal", "R-DLPFC", "Anxiety/inhibition protocols"],
            ["C", "F3-F4", "Bilateral prefrontal", "Bi-DLPFC", "Bilateral executive"],
            ["D", "F3-P3", "Left frontoparietal", "DLPFC + PPC", "Memory + executive"],
            ["E", "C3-Cz", "Left centroparietal", "L-M1 + SMA", "Motor/gait protocols"],
            ["F", "T3-Cz", "Left temporoparietal", "Temporal + SMA", "Language/auditory"],
            ["G", "Fz-Cz", "Midline frontal", "mPFC + SMA", "Midline/medial protocols"],
        ],
        header_bg="blue", alt_bg="alt_row",
        col_widths=[2.0, 2.5, 3.5, 3.5, 6.5], font_size=8.5,
    )
    _spacer(doc, 6)

    for proto in cond.get("plato", []):
        _band(doc, f"{proto['code']}  ·  {proto['symptom']}", bg_key="blue", size=10)
        _spacer(doc, 3)
        anodes   = ", ".join(proto.get("anodes", []))
        cathodes = ", ".join(proto.get("cathodes", []))
        left = (f"{proto['code']}\nPlatoWork\n{proto.get('variant','Variant A')}\n\n"
                f"Program:\n{proto.get('program','Focus')}\n\n"
                f"Anode (+):\n{anodes}")
        right = (f"Symptom Target: {proto['symptom']}\n"
                 f"PlatoWork Variant: {proto.get('variant','')}\n"
                 f"PlatoApp Program: {proto.get('program','')}\n"
                 f"Anode (+): {anodes}\n"
                 f"Parameters: {proto.get('params','')}\n"
                 f"Rationale: {proto.get('rationale','')}")
        _protocol_card(doc, left, right, left_bg="blue", right_bg="white")
        _spacer(doc, 8)

    # Quick reference
    _band(doc, "Quick Reference — PlatoScience Protocol Summary", bg_key="navy", size=10)
    _spacer(doc, 3)
    plato_protos = cond.get("plato", [])
    qr_rows = [(p["code"], p["symptom"],
                ", ".join(p.get("anodes", [])),
                p.get("program",""), p.get("params","")) for p in plato_protos]
    _data_table(doc,
        headers=["Protocol", "Symptom Target", "Anode (+)", "PlatoApp", "Parameters"],
        rows=qr_rows,
        header_bg="navy", alt_bg="alt_row",
        col_widths=[2.0, 4.5, 3.0, 3.0, 5.5], font_size=8.0,
    )
    _page_break(doc)


def _build_fnon_plato_section(doc, cond):
    _band(doc, "SECTION 6 — FNON PLATOSCIENCE PROTOCOLS (FP1-FP6)", bg_key="purple", size=13)
    _spacer(doc, 4)
    _notice_box(doc, "✦ PARTNERS ONLY  —  FNON PlatoScience protocols map the F1-F6 network-oriented "
                "montages onto PlatoWork headset positions. Same FNON framework, optimised for "
                "home-based administration under remote SOZO Doctor oversight.",
                bg_key="lila", label="✦ PARTNERS ONLY", label_color=PURPLE)
    _spacer(doc, 6)

    for proto in cond.get("fnon_plato", []):
        _band(doc,
              f"{proto['code']}  ·  {proto['name']}  ·  Derived: {proto.get('derived','')}",
              bg_key="purple", size=10)
        _spacer(doc, 3)
        left = (f"{proto['code']}\n{proto['name']}\n\n"
                f"Derived from:\n{proto.get('derived','')}\n\n"
                f"Network:\n{proto.get('network','')}\n\n"
                f"Variant: {proto.get('variant','')}")
        right = (f"Network: {proto.get('network','')}\n"
                 f"PlatoWork Variant: {proto.get('variant','')}\n"
                 f"PlatoApp Program: {proto.get('program','')}\n"
                 f"Parameters: {proto.get('params','')}\n"
                 f"Derived from: {proto.get('derived','')} (HDCkit montage)\n"
                 f"Rationale: {proto.get('rationale','')}")
        _protocol_card(doc, left, right, left_bg="purple", right_bg="fnon_bg")
        _spacer(doc, 8)

    # Variant mapping
    _band(doc, "FNON → PlatoWork Variant Mapping Table", bg_key="purple", size=10)
    _spacer(doc, 3)
    fp_protos = cond.get("fnon_plato", [])
    vm_rows = [(p["code"], p.get("network",""), p.get("derived",""),
                p.get("variant",""), p.get("program",""), p.get("params",""))
               for p in fp_protos]
    _data_table(doc,
        headers=["FP Code", "Network", "F-Code", "PlatoWork Variant", "PlatoApp", "Parameters"],
        rows=vm_rows,
        header_bg="purple", alt_bg="lila",
        col_widths=[1.8, 3.2, 1.8, 3.2, 2.5, 5.5], font_size=8.0,
    )
    _page_break(doc)


def _build_multimodal_section(doc, cond):
    _band(doc, "SECTION 7 — MULTIMODAL S-O-Z-O PROTOCOLS (P1-P9)", bg_key="purple", size=13)
    _spacer(doc, 4)
    _notice_box(doc,
        "✦ PARTNERS ONLY  —  These multimodal phenotype protocols integrate all SOZO modalities "
        "(tDCS, TPS, taVNS, CES) within the S-O-Z-O sequencing framework. Each protocol targets "
        "a specific clinical phenotype through its underlying network dysfunction signature.",
        bg_key="lila", label="✦ PARTNERS ONLY", label_color=PURPLE)
    _spacer(doc, 6)

    # S-O-Z-O key
    _band(doc, "S-O-Z-O Sequence Key", bg_key="navy", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Phase", "Modality", "Duration", "Purpose"],
        rows=[
            ["S — Stabilize", "taVNS or CES (Alpha-Stim AID)", "10–15 min", "Downregulate autonomic arousal, reduce cortical noise, prepare neural substrate"],
            ["O — Optimize", "tDCS (HDCkit or PlatoScience)", "20–30 min", "Prime target network; anodal excitation or cathodal inhibition of key nodes"],
            ["Z — Zone Target", "TPS (NEUROLITH®)", "20–30 min · 800–1,200 pulses", "Deep focal stimulation of network hubs; temporal summation with tDCS"],
            ["O — Outcome", "taVNS/CES + Rehab Tasks", "30–45 min", "Consolidate plasticity window; condition-specific cognitive/motor/behavioral tasks"],
        ],
        header_bg="navy", alt_bg="alt_row",
        col_widths=[3.0, 4.5, 3.5, 7.0], font_size=8.5,
    )
    _spacer(doc, 8)

    for pheno in cond.get("phenotypes", []):
        # Phenotype header
        _band(doc,
              f"{pheno['code']}  —  {pheno['name']}    {pheno.get('ft_code','')}",
              bg_key="purple", size=12)
        _spacer(doc, 4)

        # S-O-Z-O sequence header
        p = doc.add_paragraph()
        r = p.add_run("  S-O-Z-O Treatment Sequence")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
        r.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

        # S-O-Z-O bar
        _sozo_bar(doc, pheno.get("sozo_bar", {}))
        _spacer(doc, 4)

        # Protocol card
        _protocol_card(doc,
                       pheno.get("card_left", ""),
                       pheno.get("card_right", ""),
                       left_bg="purple", right_bg="white")
        _spacer(doc, 4)

        # Brain Targets table
        _band(doc, f"Brain Targets — {pheno['name']}", bg_key="navy", size=10)
        _spacer(doc, 3)
        brain_rows = pheno.get("brain_targets", [])
        _data_table(doc,
            headers=["Brain Region / Network Node", "Role & Rationale"],
            rows=brain_rows,
            header_bg="navy", alt_bg="alt_row",
            col_widths=[5.0, 13.0], font_size=8.5,
        )
        _spacer(doc, 4)

        # Multimodal Combinations table
        _band(doc, f"Multimodal Combinations — {pheno['code']}", bg_key="navy", size=10)
        _spacer(doc, 3)
        _data_table(doc,
            headers=["Combination", "Mechanistic Rationale", "Timing", "Clinical Indication"],
            rows=pheno.get("combinations", []),
            header_bg="navy", alt_bg="alt_row",
            col_widths=[4.0, 5.5, 3.5, 5.0], font_size=8.5,
        )
        _spacer(doc, 4)

        # Task Pairing box
        _notice_box(doc,
                    pheno.get("task_pairing", "Concurrent tasks to be prescribed by clinician."),
                    bg_key="task_bg",
                    label="🎯 Concurrent Task Pairing")
        _spacer(doc, 4)

        # FNON Hypothesis box
        _notice_box(doc,
                    pheno.get("fnon_hypothesis", "FNON network hypothesis to be documented."),
                    bg_key="fnon_bg",
                    label="✦ FNON Network Hypothesis", label_color=PURPLE)
        _spacer(doc, 10)
        _hr(doc, "7B2D8E")
        _spacer(doc, 6)

    _page_break(doc)


def _build_safety_section(doc, cond):
    _band(doc, "SECTION 8 — SAFETY, SIDE EFFECTS & MONITORING", bg_key="navy", size=13)
    _spacer(doc, 6)

    _band(doc, "Common Side Effects by Modality", bg_key="navy", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Side Effect", "Modality", "Frequency", "Management"],
        rows=[
            ["Tingling / itching at electrode", "tDCS", "Common (40–70%)", "Reassure; verify electrode contact; saline re-moisten"],
            ["Headache", "tDCS · TPS · taVNS", "Uncommon (5–15%)", "Reduce intensity; hydration; paracetamol PRN; rest"],
            ["Skin redness / erythema", "tDCS", "Common at electrode site", "Inspect post-session; petroleum jelly barrier; rotate sites"],
            ["Transient phosphenes", "tDCS (frontal)", "Uncommon", "Reposition electrodes; avoid OFC montage; reduce current"],
            ["Ear discomfort / irritation", "taVNS", "Common (20–40%)", "Adjust clip; reduce current; 5-min break; change ear"],
            ["Drowsiness / fatigue post-session", "CES · taVNS · TPS", "Common", "Allow 15-min rest post-session; do not drive after session"],
        ],
        header_bg="navy", alt_bg="alt_row",
        col_widths=[4.5, 3.0, 3.5, 7.0], font_size=8.5,
    )
    _spacer(doc, 6)

    _band(doc, "Adverse Event Grading & Response", bg_key="navy", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Grade", "Description", "Required Action"],
        rows=[
            ["Grade 1 — Mild", "Self-resolving within 30 min; no functional impairment", "Continue session; document in clinical notes; monitor at next visit"],
            ["Grade 2 — Moderate", "Persists >30 min or recurs at multiple sessions", "Pause or terminate session; reassess protocol; adjust parameters"],
            ["Grade 3 — Severe", "Requires medical attention; significant distress/impairment", "STOP all stimulation; escalate to Doctor; file incident report"],
            ["Grade 4 — Serious (Seizure / LOC)", "Seizure, loss of consciousness, or cardiovascular event", "STOP immediately; call emergency services; notify Doctor at once"],
        ],
        header_bg="navy", alt_bg="alt_row",
        col_widths=[4.0, 6.0, 8.0], font_size=8.5,
    )
    _spacer(doc, 6)

    _band(doc, "Contraindications — tDCS vs TPS", bg_key="navy", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Category", "tDCS Contraindication", "TPS Contraindication"],
        rows=[
            ["Metal implants (cranial)", "ABSOLUTE — all montages", "ABSOLUTE — all targets"],
            ["Pacemaker / ICD / neuromodulator implant", "ABSOLUTE — consult cardiologist", "ABSOLUTE — all targets"],
            ["Active uncontrolled epilepsy", "ABSOLUTE", "ABSOLUTE"],
            ["Pregnancy (frontal)", "RELATIVE — Doctor decision", "RELATIVE — Doctor decision"],
            ["Open scalp wounds / active skin infection", "ABSOLUTE at affected site", "ABSOLUTE at affected site"],
            ["Recent neurosurgery (<6 months)", "RELATIVE", "ABSOLUTE"],
        ],
        header_bg="navy", alt_bg="alt_row",
        col_widths=[5.0, 6.5, 6.5], font_size=8.5,
    )
    _spacer(doc, 6)

    _notice_box(doc,
        "HOME-USE SAFETY  —  PlatoScience protocols prescribed for home use require: "
        "(1) Completed supervised in-clinic baseline (minimum 3 sessions). "
        "(2) Written device orientation and safety checklist completion. "
        "(3) SOZO Doctor remote monitoring active. "
        "(4) Follow-up call within 48 hours of first home session. "
        "(5) Emergency escalation pathway communicated to patient and carer.",
        bg_key="o_orange", label="🏠 Home-Use Safety Protocol", label_color=NAVY)
    _spacer(doc, 6)

    _band(doc, "In-Clinic Task Pairing — Stimulation During Rehabilitation", bg_key="navy", size=10)
    _spacer(doc, 3)
    task_rows = cond.get("task_pairing_rows", [
        ["C1 / C2 (DLPFC)", "Primary symptom target", "Condition-specific cognitive tasks", "Stimulation + task coupling during plasticity window"],
        ["C3 (Cognition)", "Executive / working memory", "N-back, PASAT, verbal fluency", "DLPFC excitability + concurrent cognitive load"],
        ["C4 (Mood/Affect)", "Affective regulation", "Positive affect induction, imagery", "Prefrontal asymmetry + emotional learning"],
        ["C5/C6", "Fatigue / energy / sleep", "Paced breathing, mindfulness", "Autonomic downregulation concurrent with stimulation"],
        ["C7/C8 (Advanced)", "Symptom-specific", "Condition-specific tasks", "Targeted engagement at peak stimulation window"],
        ["FT/F protocols", "Network-oriented targets", "FNON-matched rehabilitation tasks", "Network-level plasticity via task-conditioned stimulation"],
    ])
    _data_table(doc,
        headers=["Montage", "Target Domain", "Concurrent Tasks", "Rationale"],
        rows=task_rows,
        header_bg="navy", alt_bg="alt_row",
        col_widths=[3.0, 3.5, 5.5, 6.0], font_size=8.5,
    )
    _page_break(doc)


def _build_followup_section(doc, cond):
    _band(doc, "SECTION 9 — FOLLOW-UP & DECISION-MAKING", bg_key="navy", size=13)
    _spacer(doc, 6)

    _colored_para(doc,
        "SOZO's iterative reassessment model ensures protocols are refined continuously based on "
        "clinical response, biomarker data (EEG, cognitive testing, validated rating scales), and "
        "patient-reported outcomes. The following schedule applies to the Partners tier.",
        color=NAVY, size=9)

    _band(doc, "Recommended Follow-Up Schedule", bg_key="navy", size=10)
    _spacer(doc, 3)
    followup_rows = cond.get("followup_rows", [
        ("Phase 1 — Induction (Sessions 1–10)", "Daily or 5×/week", "Baseline validated rating scale + SOZO-PRS", "Establish tolerability; confirm phenotype mapping"),
        ("Phase 2 — Consolidation (Sessions 11–20)", "3–5×/week", "Mid-treatment rating scale; cognitive battery", "Adjust protocol code if <20% improvement; consider phenotype switch"),
        ("Phase 3 — Maintenance (Month 2–3)", "2–3×/week", "Full outcome battery + patient-reported outcomes", "Taper intensity; introduce home PlatoScience protocols"),
        ("Phase 4 — Long-Term (Month 3+)", "Weekly booster or PRN", "Quarterly full reassessment + network mapping", "Maintenance protocol; re-induce if relapse; add FNON refinement"),
    ])
    _data_table(doc,
        headers=["Phase", "Frequency", "Assessment", "Protocol Adjustment"],
        rows=followup_rows,
        header_bg="navy", alt_bg="alt_row",
        col_widths=[4.5, 3.0, 4.5, 6.0], font_size=8.5,
    )
    _spacer(doc, 6)

    _band(doc, "Decision Algorithm — Protocol Non-Response", bg_key="navy", size=10)
    _spacer(doc, 3)
    _data_table(doc,
        headers=["Criterion", "Definition", "Action"],
        rows=[
            ["Non-response", "<20% improvement on primary scale at session 10", "Switch protocol code within same section OR switch phenotype label"],
            ["Partial response", "20–49% improvement at session 15", "Augment with FNON protocol; add second modality"],
            ["Good response", "≥50% improvement at session 20", "Consolidate; begin taper to maintenance frequency"],
            ["Remission", "Below clinical threshold on primary scale", "Transition to monthly booster; document for audit"],
            ["Relapse after remission", "Return to >50% of baseline severity", "Re-induce with original protocol; consider FNON reassessment"],
        ],
        header_bg="navy", alt_bg="alt_row",
        col_widths=[4.0, 5.0, 9.0], font_size=8.5,
    )
    _spacer(doc, 6)

    _notice_box(doc,
        "All clinical decisions remain the sole responsibility of the treating SOZO Doctor. "
        "This compendium is a clinical reference tool — it does not replace clinical judgment. "
        "Protocols must be adapted to the individual patient's presentation, comorbidities, "
        "medications, and response trajectory. SOZO Brain Center, Cyprus.",
        bg_key="lila", label="⚕ Clinical Governance", label_color=PURPLE)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_document(cond_data: dict, slug: str) -> Document:
    doc = Document()
    _set_margins(doc)

    # Remove default empty paragraph
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    _build_cover(doc, cond_data)
    _build_toc(doc, cond_data)
    _build_tps_section(doc, cond_data)
    _build_fnon_tps_section(doc, cond_data)
    _build_tdcs_section(doc, cond_data)
    _build_fnon_tdcs_section(doc, cond_data)
    _build_plato_section(doc, cond_data)
    _build_fnon_plato_section(doc, cond_data)
    _build_multimodal_section(doc, cond_data)
    _build_safety_section(doc, cond_data)
    _build_followup_section(doc, cond_data)

    return doc


def main():
    root = Path(__file__).resolve().parent.parent
    generated = []
    failed = []

    for slug, cond_data in PARTNERS_CONDITIONS.items():
        try:
            out_dir = root / "outputs" / "documents" / slug / "partners"
            out_dir.mkdir(parents=True, exist_ok=True)
            out_path = out_dir / f"All_In_One_Partner_{slug}.docx"

            print(f"  Generating {slug} ...", end=" ", flush=True)
            doc = build_document(cond_data, slug)
            doc.save(str(out_path))
            size_kb = out_path.stat().st_size // 1024
            print(f"OK  {out_path.name}  ({size_kb} KB)")
            generated.append(slug)
        except Exception as e:
            print(f"FAIL  {slug}: {e}")
            import traceback; traceback.print_exc()
            failed.append((slug, str(e)))

    print(f"\n{'='*60}")
    print(f"Generated: {len(generated)}/{len(PARTNERS_CONDITIONS)}")
    if failed:
        print("Failed:")
        for slug, err in failed:
            print(f"  {slug}: {err}")
    else:
        print("All documents generated successfully.")


if __name__ == "__main__":
    main()
