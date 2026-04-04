"""
Generate SOZO Neuromodulation Modality Selection Guide (master DOCX).
Output: outputs/documents/shared/SOZO_Modality_Selection_Guide.docx
Run from project root: python scripts/generate_modality_selection_guide.py
"""
import sys
import json
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
PURPLE = RGBColor(0x7B, 0x2D, 0x8E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
DGRAY  = RGBColor(0x44, 0x44, 0x44)
LGRAY  = RGBColor(0xF0, 0xF4, 0xF8)

BG = {
    "navy":     "1B3A5C",
    "blue":     "2E75B6",
    "purple":   "7B2D8E",
    "white":    "FFFFFF",
    "lila":     "F0E6F6",
    "alt_row":  "F0F4F8",
    "lgray":    "F5F5F5",
    "mid_blue": "D6E4F0",
    "cover_bg": "1B3A5C",
}

# ── Modality columns & display names ──────────────────────────────────────────
MATRIX_COLS = ["TPS", "TMS", "tDCS", "taVNS", "CES", "tACS", "PBM", "PEMF", "LIFU", "tRNS", "DBS"]

# Map from matrix column name → protocol dict key in protocols{}
PROTO_KEY_MAP = {
    "TPS":   "TPS",
    "TMS":   "TMS_rTMS",
    "tDCS":  "tDCS",
    "taVNS": "taVNS_tVNS",
    "CES":   "CES",
    "tACS":  "tACS",
    "PBM":   "PBM",
    "PEMF":  "PEMF",
    "LIFU":  "LIFU_tFUS",
    "tRNS":  "tRNS",
    "DBS":   "DBS",
}

MODALITY_DISPLAY = {
    "TPS":      "TPS — Transcranial Pulse Stimulation",
    "TMS_rTMS": "TMS / rTMS — Transcranial Magnetic Stimulation",
    "tDCS":     "tDCS — Transcranial Direct Current Stimulation",
    "taVNS_tVNS": "taVNS / tVNS — Transcutaneous Auricular Vagus Nerve Stimulation",
    "CES":      "CES — Cranial Electrotherapy Stimulation",
    "tACS":     "tACS — Transcranial Alternating Current Stimulation",
    "PBM":      "PBM — Photobiomodulation",
    "PEMF":     "PEMF — Pulsed Electromagnetic Field Therapy",
    "LIFU_tFUS": "LIFU / tFUS — Low-Intensity Focused Ultrasound",
    "tRNS":     "tRNS — Transcranial Random Noise Stimulation",
    "DBS":      "DBS — Deep Brain Stimulation",
}

MODALITY_SHORT = {
    "TPS":      "TPS",
    "TMS_rTMS": "TMS/rTMS",
    "tDCS":     "tDCS",
    "taVNS_tVNS": "taVNS",
    "CES":      "CES",
    "tACS":     "tACS",
    "PBM":      "PBM",
    "PEMF":     "PEMF",
    "LIFU_tFUS": "LIFU/tFUS",
    "tRNS":     "tRNS",
    "DBS":      "DBS",
}


# ══════════════════════════════════════════════════════════════════════════════
# LOW-LEVEL XML HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), sz)
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _set_table_borders(table, color="CCCCCC", sz="4"):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color, sz)


def _no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


def _col_widths(table, widths_cm: list):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _set_margins(doc, top=1.8, bottom=1.8, left=1.8, right=1.8):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def _page_break(doc):
    doc.add_page_break()


def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    _no_space_para(p)


def _hr(doc, color="1B3A5C"):
    p = doc.add_paragraph()
    _no_space_para(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# TEXT + BLOCK HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _add_run(para, text, bold=False, size=9, color=None, italic=False):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color
    if italic:
        r.font.italic = True
    return r


def _colored_para(doc, text, color=NAVY, bold=False, size=11,
                  align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def _band(doc, text, bg_key="navy", fg=WHITE, size=12, left_pad=True):
    """Full-width coloured band heading."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, BG[bg_key])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(("  " if left_pad else "") + text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = fg
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    return tbl


def _notice_box(doc, text, bg_key="lila", label="", label_color=None):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, BG[bg_key])
    _set_cell_borders(cell, "BBBBBB", "4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    if label:
        lc = label_color or PURPLE
        rl = p.add_run(label + "  ")
        rl.bold = True
        rl.font.size = Pt(9)
        rl.font.color.rgb = lc
        rl.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    return tbl


def _data_table(doc, headers, rows, header_bg="navy", alt_bg="alt_row",
                col_widths_cm=None, font_size=8.5, header_font_size=None):
    if header_font_size is None:
        header_font_size = font_size
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")

    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, BG[header_bg])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(header_font_size)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        bg = BG[alt_bg] if ri % 2 == 0 else BG["white"]
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(str(val) if val is not None else "")
            r.font.size = Pt(font_size)
            r.font.color.rgb = NAVY
            r.font.name = "Calibri"

    if col_widths_cm:
        _col_widths(tbl, col_widths_cm)

    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# DATA HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _parse_count(val):
    """Return integer paper count or 0 if None/non-numeric."""
    if val is None:
        return 0
    s = str(val).strip().rstrip("+").replace(",", "")
    try:
        return int(s)
    except ValueError:
        return 0


def _display_count(val):
    """Human-readable count string; None -> dash."""
    if val is None:
        return "\u2014"
    s = str(val).strip()
    if s == "0":
        return "\u2014"
    return s


def _find_protocol(protocols, proto_key, condition_name):
    """Find first protocol for a modality that matches a condition (case-insensitive partial match)."""
    if proto_key not in protocols:
        return None
    cond_lower = condition_name.lower()
    # Try substring match on condition field
    for p in protocols[proto_key]:
        pcond = p.get("condition", "").lower()
        # Use first few words of condition name for matching
        cond_words = cond_lower.split()[:3]
        if any(w in pcond for w in cond_words if len(w) > 3):
            return p
    # Fallback: first protocol in modality
    return None


def _truncate(text, max_len=200):
    if not text:
        return ""
    text = str(text)
    if len(text) <= max_len:
        return text
    return text[:max_len].rstrip() + "..."


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: COVER
# ══════════════════════════════════════════════════════════════════════════════

def _build_cover(doc):
    # Large navy background band as cover
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, BG["navy"])
    _set_cell_borders(cell, "1B3A5C", "0")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run("SOZO Brain Center")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = WHITE
    r.font.name = "Calibri"

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    r2 = p2.add_run("Neuromodulation Modality Selection Guide")
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0xB8, 0xD4, 0xF0)
    r2.font.name = "Calibri"

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(16)
    p3.paragraph_format.space_after  = Pt(4)
    r3 = p3.add_run("Evidence-Based Clinical Decision Support \u2014 April 2026")
    r3.bold = False
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    r3.font.name = "Calibri"

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(8)
    p4.paragraph_format.space_after  = Pt(40)
    r4 = p4.add_run("22 Conditions  \u2022  11 Modalities  \u2022  Master Evidence Matrix + Protocol Protocols")
    r4.bold = False
    r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)
    r4.font.name = "Calibri"

    _page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: HOW TO USE THIS GUIDE
# ══════════════════════════════════════════════════════════════════════════════

def _build_how_to_use(doc):
    _band(doc, "How to Use This Guide", bg_key="blue", size=13)
    _spacer(doc, 6)

    intro = (
        "This guide is designed to help SOZO Brain Center clinicians rapidly identify the most "
        "evidence-supported neuromodulation modality for each condition encountered in clinical practice. "
        "It integrates data from the SOZO evidence matrix (sourced from PubMed, Google Scholar, and "
        "regulatory databases) across 22 neurological and psychiatric conditions and 11 non-invasive "
        "and minimally-invasive brain stimulation modalities."
    )
    _colored_para(doc, intro, NAVY, size=11, space_before=4, space_after=6)

    how_text = (
        "Paper counts in the Evidence Matrix are approximate totals from indexed literature searches "
        "and reflect relative research volume, not quality. Higher counts generally indicate a more "
        "mature evidence base, but should be interpreted alongside evidence level (RCT, meta-analysis, "
        "case series, etc.) and regulatory status. The Best Modality column identifies the first-line "
        "recommendation based on a synthesis of evidence strength, regulatory approval, safety profile, "
        "and clinical availability."
    )
    _colored_para(doc, how_text, NAVY, size=11, space_before=2, space_after=6)

    _notice_box(doc,
        "Deep-Dive sections provide condition-specific modality rankings and protocol summaries. "
        "Modality Reference Cards provide the complete list of conditions each modality treats, "
        "with brain targets, protocols, and key references for clinical planning.",
        bg_key="mid_blue",
        label="Guide Structure:",
        label_color=NAVY)
    _spacer(doc, 8)

    # Quick legend table
    _colored_para(doc, "Evidence Level Key", BLUE, bold=True, size=11, space_before=6, space_after=4)
    _data_table(doc,
        headers=["Level", "Description"],
        rows=[
            ["RCT, Meta-analysis", "Highest — multiple randomised controlled trials and/or published meta-analyses"],
            ["RCT, Systematic Review", "High — at least one RCT plus systematic review of evidence"],
            ["RCT", "Moderate-High — at least one published randomised controlled trial"],
            ["Pilot RCT, Case Series", "Moderate — small-scale RCTs or case series; promising but limited"],
            ["Case Report, Observational", "Emerging — preliminary evidence only; interpret with caution"],
            ["Expert Opinion", "Theoretical/consensus — no controlled trials; clinical extrapolation"],
        ],
        header_bg="blue",
        alt_bg="alt_row",
        col_widths_cm=[5.0, 13.5],
        font_size=9,
    )
    _spacer(doc, 6)
    _page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: MASTER EVIDENCE MATRIX TABLE
# ══════════════════════════════════════════════════════════════════════════════

def _build_evidence_matrix(doc, conditions_matrix):
    _band(doc, "Master Evidence Matrix \u2014 All Conditions \u00d7 All Modalities", bg_key="navy", size=13)
    _spacer(doc, 4)
    _colored_para(doc,
        "Paper counts are approximate PubMed/Scholar totals. Bold = highest count per row. "
        "'\u2014' = negligible or no indexed literature. Best Modality = first-line recommendation.",
        DGRAY, size=9, space_before=2, space_after=4)

    headers = ["Condition"] + MATRIX_COLS + ["Best Modality"]
    n_cols = len(headers)
    conditions = list(conditions_matrix.keys())
    n_rows = len(conditions)

    tbl = doc.add_table(rows=1 + n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")

    # Header row
    hrow = tbl.rows[0]
    for ci, h in enumerate(headers):
        cell = hrow.cells[ci]
        _set_cell_bg(cell, BG["navy"])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    # Data rows
    for ri, cond in enumerate(conditions):
        entry = conditions_matrix[cond]
        counts = [_parse_count(entry.get(col)) for col in MATRIX_COLS]
        max_count = max(counts) if any(c > 0 for c in counts) else 0
        bg = BG["alt_row"] if ri % 2 == 0 else BG["white"]

        drow = tbl.rows[ri + 1]

        # Condition name cell
        cell0 = drow.cells[0]
        _set_cell_bg(cell0, bg)
        p0 = cell0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after  = Pt(2)
        r0 = p0.add_run(cond)
        r0.bold = True
        r0.font.size = Pt(7.5)
        r0.font.color.rgb = NAVY
        r0.font.name = "Calibri"

        # Modality count cells
        for ci, col in enumerate(MATRIX_COLS):
            cell = drow.cells[ci + 1]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            raw = entry.get(col)
            cnt = _parse_count(raw)
            display = _display_count(raw)
            is_max = (cnt == max_count and cnt > 0)
            r = p.add_run(display)
            r.bold = is_max
            r.font.size = Pt(7.5)
            r.font.color.rgb = NAVY
            r.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Best modality cell (purple text)
        cell_bm = drow.cells[n_cols - 1]
        _set_cell_bg(cell_bm, bg)
        p_bm = cell_bm.paragraphs[0]
        p_bm.paragraph_format.space_before = Pt(2)
        p_bm.paragraph_format.space_after  = Pt(2)
        best = entry.get("Best Modality/Modalities", "")
        r_bm = p_bm.add_run(best or "\u2014")
        r_bm.bold = True
        r_bm.font.size = Pt(7.5)
        r_bm.font.color.rgb = PURPLE
        r_bm.font.name = "Calibri"

    # Column widths — total ~18.5 cm usable at 1.8 cm margins on A4 (21 - 3.6 = 17.4 + a bit)
    # Condition col wide, modality cols narrow, best-modality wider
    cw = [4.0] + [1.1] * 11 + [3.8]
    _col_widths(tbl, cw)

    _spacer(doc, 8)
    _page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: PER-CONDITION DEEP DIVE
# ══════════════════════════════════════════════════════════════════════════════

def _build_condition_deep_dives(doc, conditions_matrix, protocols):
    _band(doc, "Per-Condition Deep Dive", bg_key="navy", size=14)
    _spacer(doc, 6)
    _colored_para(doc,
        "Each subsection ranks modalities by evidence volume, highlights the recommended first-line "
        "modality, and provides protocol summaries for the top three evidence-supported approaches.",
        DGRAY, size=10, space_before=2, space_after=6)
    _page_break(doc)

    for cond in conditions_matrix:
        entry = conditions_matrix[cond]
        best = entry.get("Best Modality/Modalities", "")

        # Section heading
        _band(doc, f"Condition: {cond}", bg_key="blue", size=12)
        _spacer(doc, 4)

        # Best modality callout
        _notice_box(doc,
            best or "See matrix for details",
            bg_key="lila",
            label="Recommended First-Line Modality:",
            label_color=PURPLE)
        _spacer(doc, 4)

        # Build ranked modality list (>0 papers only)
        ranked = []
        for col in MATRIX_COLS:
            raw = entry.get(col)
            cnt = _parse_count(raw)
            if cnt > 0:
                # Find evidence level from protocols
                proto_key = PROTO_KEY_MAP[col]
                ev_level = ""
                p = _find_protocol(protocols, proto_key, cond)
                if p:
                    ev_level = p.get("evidence_level", "")
                ranked.append((col, raw, cnt, ev_level))
        ranked.sort(key=lambda x: x[2], reverse=True)

        if ranked:
            _colored_para(doc, "Evidence by Modality (descending paper count)", BLUE, bold=True,
                          size=10, space_before=2, space_after=3)
            mini_rows = []
            for col, raw, cnt, ev_level in ranked:
                mini_rows.append([col, _display_count(raw), ev_level or "\u2014"])
            _data_table(doc,
                headers=["Modality", "Papers", "Evidence Level"],
                rows=mini_rows,
                header_bg="blue",
                alt_bg="alt_row",
                col_widths_cm=[3.0, 2.5, 12.5],
                font_size=9,
            )
            _spacer(doc, 5)

        # Recommended protocol summaries — top 3 modalities
        top3 = ranked[:3]
        if top3:
            _colored_para(doc, "Recommended Protocol Summaries (Top 3 Modalities)", BLUE, bold=True,
                          size=10, space_before=2, space_after=3)
            proto_rows = []
            for col, raw, cnt, _ in top3:
                proto_key = PROTO_KEY_MAP[col]
                p = _find_protocol(protocols, proto_key, cond)
                if p:
                    proto_rows.append([
                        col,
                        p.get("brain_target", "\u2014"),
                        _truncate(p.get("protocol_summary", "\u2014"), 180),
                        p.get("devices", "\u2014"),
                        p.get("regulatory_status", "\u2014"),
                    ])
            if proto_rows:
                _data_table(doc,
                    headers=["Modality", "Brain Target", "Protocol Summary", "Devices", "Regulatory"],
                    rows=proto_rows,
                    header_bg="navy",
                    alt_bg="alt_row",
                    col_widths_cm=[2.0, 3.0, 6.5, 3.5, 3.0],
                    font_size=8.5,
                )
        _spacer(doc, 8)
        _page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: MODALITY REFERENCE CARDS
# ══════════════════════════════════════════════════════════════════════════════

def _build_modality_reference_cards(doc, protocols):
    _band(doc, "Modality Reference Cards", bg_key="navy", size=14)
    _spacer(doc, 6)
    _colored_para(doc,
        "One card per modality. Each card lists all conditions the modality has published protocols "
        "for, with brain targets, protocol summaries, session counts, evidence levels, and key "
        "references. Use these cards for rapid modality-specific clinical planning.",
        DGRAY, size=10, space_before=2, space_after=6)
    _page_break(doc)

    MODALITY_KEYS = list(MODALITY_DISPLAY.keys())

    for proto_key in MODALITY_KEYS:
        display_name = MODALITY_DISPLAY[proto_key]
        short_name = MODALITY_SHORT[proto_key]
        proto_list = protocols.get(proto_key, [])

        _band(doc, display_name, bg_key="blue", size=12)
        _spacer(doc, 4)

        if not proto_list:
            _colored_para(doc, "No protocol data available for this modality.", DGRAY, size=10)
            _spacer(doc, 6)
            _page_break(doc)
            continue

        _colored_para(doc,
            f"{short_name} treats {len(proto_list)} condition(s) with published protocols. "
            f"All protocols sourced from indexed literature and regulatory databases.",
            NAVY, size=10, space_before=2, space_after=4)

        # Full protocols table
        rows = []
        for p in proto_list:
            rows.append([
                p.get("condition", "\u2014"),
                p.get("brain_target", "\u2014"),
                _truncate(p.get("protocol_summary", "\u2014"), 160),
                p.get("total_sessions", "\u2014"),
                p.get("evidence_level", "\u2014"),
                _truncate(p.get("key_references", "\u2014"), 120),
            ])

        _data_table(doc,
            headers=["Condition", "Brain Target", "Protocol Summary",
                     "Sessions", "Evidence Level", "Key References"],
            rows=rows,
            header_bg="navy",
            alt_bg="alt_row",
            col_widths_cm=[3.2, 2.8, 5.5, 2.0, 2.5, 4.0],
            font_size=8,
        )
        _spacer(doc, 8)
        _page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    # ── Load data ────────────────────────────────────────────────────────────
    data_path = ROOT / "scripts" / "excel_knowledge_dump.json"
    if not data_path.exists():
        print(f"ERROR: Data file not found: {data_path}")
        sys.exit(1)

    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)

    conditions_matrix = data["conditions_matrix"]   # condition -> {modality: count, ...}
    protocols         = data["protocols"]            # modality_key -> [proto_dict, ...]

    conditions    = list(conditions_matrix.keys())   # 22 conditions
    modality_keys = list(protocols.keys())           # 11 modality keys

    print(f"Loaded {len(conditions)} conditions, {len(modality_keys)} modality keys")

    # ── Prepare output path ──────────────────────────────────────────────────
    out_dir = ROOT / "outputs" / "documents" / "shared"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "SOZO_Modality_Selection_Guide.docx"

    # ── Build document ───────────────────────────────────────────────────────
    doc = Document()
    _set_margins(doc, top=1.8, bottom=1.8, left=1.8, right=1.8)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    section_count = 0

    # 1. Cover
    _build_cover(doc)
    section_count += 1
    print(f"  [1/5] Cover page done")

    # 2. How to Use
    _build_how_to_use(doc)
    section_count += 1
    print(f"  [2/5] How-to-use section done")

    # 3. Evidence Matrix
    _build_evidence_matrix(doc, conditions_matrix)
    section_count += 1
    print(f"  [3/5] Evidence matrix table done ({len(conditions)} conditions x {len(MATRIX_COLS)} modalities)")

    # 4. Per-condition deep dives
    _build_condition_deep_dives(doc, conditions_matrix, protocols)
    section_count += 1
    print(f"  [4/5] Per-condition deep dives done ({len(conditions)} conditions)")

    # 5. Modality reference cards
    _build_modality_reference_cards(doc, protocols)
    section_count += 1
    print(f"  [5/5] Modality reference cards done ({len(modality_keys)} modalities)")

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(str(out_path))

    total_proto_count = sum(len(v) for v in protocols.values())
    print()
    print("=" * 65)
    print("SOZO Modality Selection Guide — Generation Complete")
    print(f"  Output : {out_path}")
    print(f"  Sections: {section_count} (Cover, How-To-Use, Evidence Matrix,")
    print(f"             Per-Condition Deep Dives, Modality Reference Cards)")
    print(f"  Conditions: {len(conditions)}")
    print(f"  Modalities: {len(modality_keys)}")
    print(f"  Protocols indexed: {total_proto_count}")
    print("=" * 65)


if __name__ == "__main__":
    main()
