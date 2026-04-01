"""
Generate FELLOW ALL-IN-ONE protocol documents for all 15 SOZO conditions.
Replicates the PD template structure: TPS, HDCkit tDCS, PlatoScience, Multimodal F1-F9, Safety.
Outputs: outputs/documents/{slug}/fellow/FELLOW_ALL_IN_ONE_{SLUG}.docx
"""

import io
import sys
import datetime
from pathlib import Path

# Allow running from scripts/ or repo root
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

import matplotlib
matplotlib.use("Agg")

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from eeg_map import generate_eeg_map
from protocol_data import (
    CONDITIONS, TPS_WARNING_TEMPLATE, TPS_NOTE, HDCKIT_NOTE,
    PLATOSCIENCE_DEVICE, HOME_NOTE, SOZO_SEQUENCE_HEADER,
    EVIDENCE_LEVELS, SAFETY_SIDE_EFFECTS, ADVERSE_GRADING,
    CONTRAINDICATIONS_BASE,
)

# ── Colour palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x21, 0x37)
TEAL   = RGBColor(0x1A, 0x7A, 0x8A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RED_W  = RGBColor(0xC0, 0x39, 0x2B)
LGRAY  = RGBColor(0xF5, 0xF5, 0xF5)
LTEAL  = RGBColor(0xE8, 0xF4, 0xF6)
LRED   = RGBColor(0xFD, 0xEC, 0xEA)
DGRAY  = RGBColor(0x4A, 0x60, 0x70)

# hex versions for XML cell shading
_HEX = {
    "navy":  "0D2137",
    "teal":  "1A7A8A",
    "white": "FFFFFF",
    "lteal": "E8F4F6",
    "lred":  "FDECEA",
    "lgray": "F5F5F5",
    "mid":   "D0E8EC",
}


# ── Low-level XML helpers ───────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color="1A7A8A", sz="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), sz)
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _set_table_borders(table, color="1A7A8A", sz="8"):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color, sz)


def _no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


def _col_widths(table, widths_cm: list):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def _set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(Cm(height_cm).emu / 914.4)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


# ── Document-level helpers ─────────────────────────────────────────────────

def _set_margins(doc, top=1.9, bottom=1.9, left=1.9, right=1.9):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def _page_break(doc):
    doc.add_page_break()


def _hr(doc, color="1A7A8A"):
    """Thin teal horizontal rule."""
    p = doc.add_paragraph()
    _no_space_para(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    _no_space_para(p)


# ── Styled paragraph helpers ───────────────────────────────────────────────

def _add_heading_band(doc, text: str, bg: str = "navy", fg: RGBColor = WHITE,
                       font_size: int = 14):
    """Full-width navy/teal band heading."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, _HEX[bg])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = fg
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    cell.width = Cm(17.8)
    return tbl


def _add_notice_box(doc, text: str, bg: str = "lteal", label: str = ""):
    """Coloured info/warning box."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, _HEX[bg])
    _set_cell_borders(cell, "AAAAAA", "4")
    p = cell.paragraphs[0]
    if label:
        run_lbl = p.add_run(label + "  ")
        run_lbl.bold = True
        run_lbl.font.size = Pt(8.5)
        run_lbl.font.color.rgb = RED_W if bg == "lred" else NAVY
        run_lbl.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def _add_protocol_header(doc, code: str, symptom: str, evidence: str = ""):
    """Teal protocol code + symptom sub-header row."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c0 = tbl.cell(0, 0)
    c1 = tbl.cell(0, 1)
    _set_cell_bg(c0, _HEX["teal"])
    _set_cell_bg(c1, _HEX["lteal"])
    c0.width = Cm(2.2)
    c1.width = Cm(15.6)

    p0 = c0.paragraphs[0]
    r0 = p0.add_run(code)
    r0.bold = True
    r0.font.size = Pt(13)
    r0.font.color.rgb = WHITE
    r0.font.name = "Calibri"
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after  = Pt(4)

    p1 = c1.paragraphs[0]
    r1 = p1.add_run(symptom)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r1.font.name = "Calibri"
    if evidence:
        r_ev = p1.add_run(f"   [{evidence}]")
        r_ev.font.size = Pt(8.5)
        r_ev.font.color.rgb = DGRAY
        r_ev.font.name = "Calibri"
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(4)


def _add_image_text_table(doc, img_bytes: bytes, text_lines: list[tuple],
                           img_width_cm: float = 5.5, total_width_cm: float = 17.8):
    """
    1×2 table: left = EEG map image | right = styled text lines.
    text_lines: list of (text, bold, font_size, color_rgb)
    """
    txt_width = total_width_cm - img_width_cm
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c_img = tbl.cell(0, 0)
    c_txt = tbl.cell(0, 1)
    _set_cell_bg(c_img, _HEX["white"])
    _set_cell_bg(c_txt, _HEX["lteal"])
    _set_table_borders(tbl, "D0E8EC", "6")

    c_img.width = Cm(img_width_cm)
    c_txt.width = Cm(txt_width)
    c_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Insert image
    p_img = c_img.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_space_para(p_img)
    run_img = p_img.add_run()
    buf = io.BytesIO(img_bytes)
    run_img.add_picture(buf, width=Cm(img_width_cm - 0.3))

    # Insert text
    first = True
    for text, bold, fsize, color in text_lines:
        if first:
            p = c_txt.paragraphs[0]
            first = False
        else:
            p = c_txt.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(fsize)
        r.font.color.rgb = color
        r.font.name = "Calibri"

    return tbl


def _add_data_table(doc, headers: list, rows: list[list],
                    header_bg="teal", stripe_bg="lteal",
                    col_widths_cm: list = None):
    """Generic data table with teal header and striped rows."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "D0E8EC", "6")

    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, _HEX[header_bg])
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        bg = "lteal" if ri % 2 == 0 else "white"
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            _set_cell_bg(cell, _HEX[bg])
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(8)
            r.font.color.rgb = NAVY
            r.font.name = "Calibri"
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    if col_widths_cm:
        _col_widths(tbl, col_widths_cm)

    return tbl


# ── EEG image generator wrappers ───────────────────────────────────────────

def _tps_img(protocol: dict, condition_short: str) -> bytes:
    return generate_eeg_map(
        title=f"{condition_short} — {protocol['code']}",
        subtitle=protocol["symptom"],
        params_text=protocol["params_text"],
        tps_targets=protocol.get("targets", []),
        show_tps_legend=True,
        figsize=(3.5, 4.2),
        dpi=150,
    )


def _tdcs_img(protocol: dict, condition_short: str) -> bytes:
    return generate_eeg_map(
        title=f"{condition_short} — {protocol['code']}",
        subtitle=protocol["symptom"],
        params_text=protocol["params_text"],
        anodes=protocol.get("anodes", []),
        cathodes=protocol.get("cathodes", []),
        figsize=(3.5, 4.2),
        dpi=150,
    )


# ── Section builders ───────────────────────────────────────────────────────

def _build_title_page(doc, cond: dict, slug: str):
    today = datetime.date.today().strftime("%B %Y")

    _add_heading_band(doc, "SOZO BRAIN CENTER", "navy", WHITE, 18)
    _spacer(doc, 4)

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, _HEX["lteal"])
    _set_cell_borders(cell, "1A7A8A", "12")

    def _title_line(p, text, bold=False, size=12, color=NAVY, center=True):
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = "Calibri"
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    p0 = cell.paragraphs[0]
    _title_line(p0, "FELLOW ALL-IN-ONE PROTOCOL GUIDE", bold=True, size=18, color=NAVY)

    p1 = cell.add_paragraph()
    _title_line(p1, cond["display_name"].upper(), bold=True, size=22, color=TEAL)

    p2 = cell.add_paragraph()
    _title_line(p2, f"ICD-10: {cond['icd10']}  ·  Short code: {cond['short_name']}", size=10)

    p3 = cell.add_paragraph()
    _title_line(p3, "SOZO Neuromodulation Protocol Suite", size=11, color=DGRAY)

    p4 = cell.add_paragraph()
    _title_line(p4, f"Issued: {today}  ·  For qualified SOZO clinicians only", size=9, color=DGRAY)

    _spacer(doc, 6)

    # Protocol coverage summary
    n_tps   = len(cond.get("tps", []))
    n_tdcs  = len(cond.get("tdcs", []))
    n_plato = len(cond.get("plato", []))
    n_fen   = len(cond.get("phenotypes", []))
    _add_data_table(
        doc,
        headers=["Section", "Protocols Included"],
        rows=[
            ["TPS — Transcranial Pulse Stimulation",     f"T1 – T{n_tps}"],
            ["tDCS — Newronika HDCkit",                   f"C1 – C{n_tdcs}"],
            ["tDCS — PlatoScience Home",                  f"C1-PS – C{n_plato}-PS"],
            ["Multimodal SOZO Phenotypes",                f"F1 – F{n_fen}"],
            ["Safety & Adverse Event Management",         "Full grading table"],
        ],
        col_widths_cm=[12.0, 5.8],
    )


def _build_tps_section(doc, cond: dict):
    _page_break(doc)
    _add_heading_band(doc, "SECTION 1 — TRANSCRANIAL PULSE STIMULATION (TPS) PROTOCOLS", "navy", WHITE, 13)
    _spacer(doc, 4)

    warning = cond.get("tps_warning", TPS_WARNING_TEMPLATE.format(condition=cond["display_name"]))
    _add_notice_box(doc, warning, "lred", "WARNING")
    _spacer(doc, 3)
    _add_notice_box(doc, TPS_NOTE, "lteal", "NOTE")
    _spacer(doc, 6)

    for proto in cond.get("tps", []):
        _add_protocol_header(doc, proto["code"], proto["symptom"], proto.get("evidence", ""))
        img = _tps_img(proto, cond["short_name"])
        rationale = proto.get("rationale", "")
        text_lines = [
            (f"Target ROI: {', '.join(proto.get('targets', []))}", True, 9.5, TEAL),
            ("", False, 4, NAVY),
            (proto["params_text"], False, 9, NAVY),
            ("", False, 4, NAVY),
            ("Rationale:", True, 8.5, NAVY),
            (rationale, False, 8, DGRAY),
        ]
        _add_image_text_table(doc, img, text_lines)
        _spacer(doc, 8)


def _build_tdcs_hdckit_section(doc, cond: dict):
    _page_break(doc)
    _add_heading_band(doc, "SECTION 2 — tDCS PROTOCOLS — NEWRONIKA HDCkit (IN-CLINIC)", "navy", WHITE, 13)
    _spacer(doc, 4)
    _add_notice_box(doc, HDCKIT_NOTE, "lteal", "NOTE")
    _spacer(doc, 6)

    for proto in cond.get("tdcs", []):
        _add_protocol_header(doc, proto["code"], proto["symptom"])
        img = _tdcs_img(proto, cond["short_name"])
        anodes_str = ", ".join(proto.get("anodes", []))
        cathodes_str = ", ".join(proto.get("cathodes", []))
        rationale = proto.get("rationale", "")
        text_lines = [
            (f"Anode (+): {anodes_str}", True, 9.5, RGBColor(0xC0, 0x39, 0x2B)),
            (f"Cathode (−): {cathodes_str}", True, 9.5, RGBColor(0x2E, 0x6D, 0xA4)),
            ("", False, 4, NAVY),
            (proto["params_text"], False, 9, NAVY),
            ("", False, 4, NAVY),
            ("Rationale:", True, 8.5, NAVY),
            (rationale, False, 8, DGRAY),
        ]
        _add_image_text_table(doc, img, text_lines)
        _spacer(doc, 8)


def _build_plato_section(doc, cond: dict):
    _page_break(doc)
    _add_heading_band(doc, "SECTION 3 — tDCS PROTOCOLS — PLATOSCIENCE (HOME-BASED)", "navy", WHITE, 13)
    _spacer(doc, 4)

    # Device specs table
    ps = PLATOSCIENCE_DEVICE
    _add_data_table(doc, ps["headers"], ps["rows"], col_widths_cm=[6.5, 11.3])
    _spacer(doc, 4)
    _add_notice_box(doc, HOME_NOTE, "lteal", "NOTE")
    _spacer(doc, 6)

    for proto in cond.get("plato", []):
        _add_protocol_header(doc, proto["code"], proto["symptom"])
        img = _tdcs_img(proto, cond["short_name"])
        anodes_str = ", ".join(proto.get("anodes", []))
        cathodes_str = ", ".join(proto.get("cathodes", []))
        variant = proto.get("variant_note", "")
        text_lines = [
            (f"Anode (+): {anodes_str}", True, 9.5, RGBColor(0xC0, 0x39, 0x2B)),
            (f"Cathode (−): {cathodes_str}", True, 9.5, RGBColor(0x2E, 0x6D, 0xA4)),
            ("", False, 4, NAVY),
            (proto["params_text"], False, 9, NAVY),
            ("", False, 4, NAVY),
        ]
        if variant:
            text_lines += [("Variant Note:", True, 8.5, NAVY), (variant, False, 8, DGRAY)]
        _add_image_text_table(doc, img, text_lines)
        _spacer(doc, 8)


def _build_multimodal_section(doc, cond: dict):
    _page_break(doc)
    _add_heading_band(doc, "SECTION 4 — MULTIMODAL SOZO PHENOTYPE PROTOCOLS (F1–F9)", "navy", WHITE, 13)
    _spacer(doc, 6)

    for pheno in cond.get("phenotypes", []):
        # Phenotype header
        _add_heading_band(doc, f"{pheno['code']}  |  {pheno['name']}", "teal", WHITE, 11)
        _spacer(doc, 3)

        # SOZO Sequence table
        seq = pheno.get("sozo_seq", [])
        if seq:
            _add_data_table(
                doc,
                headers=["S — Stabilize", "O — Optimize", "Z — Zone Target", "O — Outcome"],
                rows=[seq[:4] + [""] * (4 - len(seq))],
                col_widths_cm=[4.45, 4.45, 4.45, 4.45],
            )
            _spacer(doc, 4)

        # TPS Combinations
        tps_combos = pheno.get("tps_combos", [])
        if tps_combos:
            _add_heading_band(doc, "TPS Combination Strategies", "mid", NAVY, 9)
            _add_data_table(
                doc,
                headers=["Combination", "Mechanism", "Timing", "Indication", "Evidence"],
                rows=tps_combos,
                col_widths_cm=[3.5, 4.5, 3.2, 3.5, 3.1],
            )
            _spacer(doc, 4)

        # Non-TPS Combinations
        non_combos = pheno.get("nontps_combos", [])
        if non_combos:
            _add_heading_band(doc, "Non-TPS Combination Strategies", "mid", NAVY, 9)
            _add_data_table(
                doc,
                headers=["Combination", "Indication", "Mechanism", "Timing", "Evidence"],
                rows=non_combos,
                col_widths_cm=[3.5, 3.5, 4.0, 3.2, 3.6],
            )
            _spacer(doc, 4)

        # Rationale box
        rat = pheno.get("rationale", "")
        if rat:
            _add_notice_box(doc, rat, "lteal", "Clinical Target:")
            _spacer(doc, 6)

        _hr(doc)
        _spacer(doc, 4)


def _build_safety_section(doc, cond: dict):
    _page_break(doc)
    _add_heading_band(doc, "SECTION 5 — SAFETY, ADVERSE EVENTS & CONTRAINDICATIONS", "navy", WHITE, 13)
    _spacer(doc, 6)

    # Side effects
    _add_heading_band(doc, "Common Side Effects & Management", "teal", WHITE, 10)
    _spacer(doc, 2)
    _add_data_table(
        doc,
        headers=["Side Effect", "Modality", "Frequency", "Management"],
        rows=SAFETY_SIDE_EFFECTS,
        col_widths_cm=[4.2, 3.8, 3.8, 6.0],
    )
    _spacer(doc, 6)

    # Adverse grading
    _add_heading_band(doc, "Adverse Event Grading Protocol", "teal", WHITE, 10)
    _spacer(doc, 2)
    _add_data_table(
        doc,
        headers=["Grade", "Definition", "Action"],
        rows=ADVERSE_GRADING,
        col_widths_cm=[4.5, 7.0, 6.3],
    )
    _spacer(doc, 6)

    # Contraindications
    _add_heading_band(doc, "Contraindications", "teal", WHITE, 10)
    _spacer(doc, 2)
    extra = cond.get("contraindications_extra", [])
    _add_data_table(
        doc,
        headers=["Contraindication", "Modality", "Type"],
        rows=CONTRAINDICATIONS_BASE + extra,
        col_widths_cm=[8.5, 5.5, 3.8],
    )
    _spacer(doc, 6)

    # Key outcome measures
    kom = cond.get("key_outcome_measures", "")
    if kom:
        _add_heading_band(doc, "Key Outcome Measures", "teal", WHITE, 10)
        _spacer(doc, 2)
        _add_notice_box(doc, kom, "lteal")


# ── Main document builder ──────────────────────────────────────────────────

def build_fellow_document(slug: str) -> Path:
    cond = CONDITIONS[slug]
    doc = Document()
    _set_margins(doc, 1.9, 1.9, 1.9, 1.9)

    # Set page size to Letter (8.5 × 11 in)
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)

    _build_title_page(doc, cond, slug)
    _build_tps_section(doc, cond)
    _build_tdcs_hdckit_section(doc, cond)
    _build_plato_section(doc, cond)
    _build_multimodal_section(doc, cond)
    _build_safety_section(doc, cond)

    # Save
    out_dir = ROOT / "outputs" / "documents" / slug / "fellow"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"FELLOW_ALL_IN_ONE_{slug.upper()}.docx"
    doc.save(str(out_path))
    return out_path


# ── CLI ────────────────────────────────────────────────────────────────────

def main():
    slugs = list(CONDITIONS.keys())
    if len(sys.argv) > 1:
        slugs = [s for s in sys.argv[1:] if s in CONDITIONS]
        if not slugs:
            print("Unknown slug(s). Available:", ", ".join(CONDITIONS.keys()))
            sys.exit(1)

    total = len(slugs)
    for i, slug in enumerate(slugs, 1):
        print(f"[{i}/{total}] Generating {slug} ...", end=" ", flush=True)
        try:
            path = build_fellow_document(slug)
            size_kb = path.stat().st_size // 1024
            print(f"OK  ->  {path.relative_to(ROOT)}  ({size_kb} KB)")
        except Exception as exc:
            import traceback
            print(f"FAILED — {exc}")
            traceback.print_exc()

    print("\nDone.")


if __name__ == "__main__":
    main()
