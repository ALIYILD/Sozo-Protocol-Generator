"""
Generate SOZO Brain Regions Reference DOCX.
Saves to: outputs/documents/shared/SOZO_Brain_Regions_Reference.docx
Run from project root: python scripts/generate_brain_regions_reference.py
"""
import json
import sys
from pathlib import Path
from collections import defaultdict

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x3A, 0x5C)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
DGRAY  = RGBColor(0x44, 0x44, 0x44)

BG = {
    "navy":       "1B3A5C",
    "blue":       "2E75B6",
    "white":      "FFFFFF",
    "alt_row":    "F0F4F8",
    "mid_blue":   "D6E4F0",
    "light_purple": "E6D6F0",
    "light_orange": "FFF0D6",
    "lgray":      "F5F5F5",
}

# Depth colour codes
DEPTH_BG = {
    "Cortical":    "D6E4F0",   # light blue
    "Subcortical": "E6D6F0",   # light purple
    "Deep":        "FFF0D6",   # light orange
}

# Modality depth classification
MODALITY_DEPTH = {
    "TPS":   "Cortical",
    "TMS":   "Cortical",
    "tDCS":  "Cortical",
    "tACS":  "Cortical",
    "PBM":   "Cortical",
    "PEMF":  "Cortical",
    "tRNS":  "Cortical",
    "LIFU":  "Deep subcortical",
    "DBS":   "Deep subcortical",
    "taVNS": "Peripheral/indirect",
    "CES":   "Peripheral/indirect",
}

# Canonical modality order for the summary section
MODALITY_ORDER = ["TPS", "TMS", "tDCS", "taVNS", "CES", "tACS", "PBM", "PEMF", "LIFU", "tRNS", "DBS"]


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), sz)
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _set_table_borders(table, color="CCCCCC", sz="4"):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell, color, sz)


def _no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


def _col_widths(table, widths_cm: list):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _set_margins(doc, top=2.0, bottom=2.0, left=2.0, right=2.0):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def _page_break(doc):
    doc.add_page_break()


def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    _no_space_para(p)


def _hr(doc, color="1B3A5C"):
    p = doc.add_paragraph()
    _no_space_para(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Text helpers ──────────────────────────────────────────────────────────────

def _add_run(para, text, bold=False, size=9, color=None, italic=False):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color
    if italic:
        r.font.italic = True
    return r


def _para(doc, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    _add_run(p, text, bold=bold, size=size, color=color or NAVY)
    return p


def _colored_para(doc, text, color=NAVY, bold=False, size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


# ── Band heading ──────────────────────────────────────────────────────────────

def _band(doc, text, bg="1B3A5C", fg=WHITE, size=13):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("  " + text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = fg
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    return tbl


# ── Section 1: Title banner ───────────────────────────────────────────────────

def _build_title(doc):
    # Main title banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "1B3A5C")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("SOZO Brain Center")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = WHITE
    r.font.name = "Calibri"

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run("Neuromodulation Target Brain Regions \u2014 Clinical Reference")
    r2.bold = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0xA8, 0xC8, 0xF0)
    r2.font.name = "Calibri"

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after  = Pt(10)
    r3 = p3.add_run("Evidence-Based Stimulation Target Atlas  |  April 2026")
    r3.bold = False
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    r3.font.name = "Calibri"

    _spacer(doc, 8)


# ── Section 2: Introduction ───────────────────────────────────────────────────

def _build_intro(doc, n_regions):
    _band(doc, "Introduction", bg="2E75B6", size=11)
    _spacer(doc, 4)

    intro = (
        f"This reference describes the {n_regions} key brain regions used as stimulation targets "
        "in SOZO neuromodulation protocols. Each entry details the region\u2019s anatomical "
        "location within the EEG 10\u201320 system (where applicable), its cortical depth "
        "(Cortical, Subcortical, or Deep), dominant hemisphere involvement, primary cognitive "
        "and physiological functions, and the clinical conditions most commonly addressed. "
        "The atlas also identifies which neuromodulation modalities \u2014 including TPS, TMS, "
        "tDCS, tACS, taVNS, CES, PBM, PEMF, LIFU, tRNS, and DBS \u2014 are capable of reaching "
        "each region based on published biophysical evidence. Use this document as a rapid "
        "clinical reference when selecting stimulation targets and matching modalities to "
        "patient presentations within the SOZO framework."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(intro)
    r.font.size = Pt(9.5)
    r.font.name = "Calibri"
    r.font.color.rgb = DGRAY
    _spacer(doc, 4)


# ── Section 3: Quick-Reference Table ─────────────────────────────────────────

def _build_quick_table(doc, regions):
    _band(doc, "Quick-Reference Table \u2014 All Brain Regions", bg="1B3A5C", size=11)
    _spacer(doc, 4)

    headers = ["Region", "Abbreviation", "EEG Position", "Depth", "Key Conditions", "Modalities"]
    n_cols = len(headers)
    n_rows = 1 + len(regions)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")

    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, "1B3A5C")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    # Data rows
    for ri, reg in enumerate(regions):
        drow = tbl.rows[ri + 1]
        row_bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        depth_bg = DEPTH_BG.get(reg.get("depth", ""), row_bg)

        vals = [
            reg.get("region", ""),
            reg.get("abbreviation", ""),
            reg.get("eeg_position", ""),
            reg.get("depth", ""),
            reg.get("key_conditions", ""),
            reg.get("modalities", reg.get("modalities_that_can_target", "")),
        ]

        for ci, val in enumerate(vals):
            cell = drow.cells[ci]
            if ci == 3:  # Depth column — colour coded
                _set_cell_bg(cell, depth_bg)
            else:
                _set_cell_bg(cell, row_bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(8)
            r.font.name = "Calibri"
            r.font.color.rgb = NAVY

    # Column widths (total ~17 cm usable)
    _col_widths(tbl, [4.2, 2.0, 2.4, 2.0, 4.4, 3.0])
    _spacer(doc, 8)


# ── Section 4: Detailed Region Cards ─────────────────────────────────────────

def _build_region_card(doc, reg):
    abbr = reg.get("abbreviation", "")
    region_name = reg.get("region", "")

    # Section heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f"{abbr}  \u2014  {region_name}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE
    r.font.name = "Calibri"

    # 2-column info table
    fields = [
        ("EEG Position",             reg.get("eeg_position", "")),
        ("Hemisphere",               reg.get("hemisphere", "")),
        ("Cortical Depth",           reg.get("depth", "")),
        ("Primary Functions",        reg.get("functions", "")),
        ("Key Clinical Conditions",  reg.get("key_conditions", "")),
        ("Modalities That Can Target", reg.get("modalities", reg.get("modalities_that_can_target", ""))),
        ("Clinical Notes",           reg.get("notes", "")),
    ]

    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")

    for i, (field, value) in enumerate(fields):
        cl = tbl.cell(i, 0)
        cv = tbl.cell(i, 1)

        _set_cell_bg(cl, "D6E4F0")
        _set_cell_bg(cv, "FFFFFF")

        # Field label
        pf = cl.paragraphs[0]
        pf.paragraph_format.space_before = Pt(2)
        pf.paragraph_format.space_after  = Pt(2)
        rf = pf.add_run(field)
        rf.bold = True
        rf.font.size = Pt(9)
        rf.font.name = "Calibri"
        rf.font.color.rgb = NAVY

        # Value
        pv = cv.paragraphs[0]
        pv.paragraph_format.space_before = Pt(2)
        pv.paragraph_format.space_after  = Pt(2)
        rv = pv.add_run(str(value))
        rv.font.size = Pt(9)
        rv.font.name = "Calibri"
        rv.font.color.rgb = DGRAY

    _col_widths(tbl, [4.5, 12.5])
    _spacer(doc, 6)


def _build_detail_cards(doc, regions):
    _page_break(doc)
    _band(doc, "Detailed Brain Region Cards", bg="1B3A5C", size=11)
    _spacer(doc, 6)

    for reg in regions:
        _build_region_card(doc, reg)


# ── Section 5: Modality Reach Summary ────────────────────────────────────────

def _build_modality_summary(doc, regions):
    _page_break(doc)
    _band(doc, "Modality Reach Summary", bg="1B3A5C", size=11)
    _spacer(doc, 4)

    # Introductory note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(
        "The table below summarises which brain regions each neuromodulation modality can "
        "physically reach, based on biophysical depth of penetration. "
        "Peripheral/indirect modalities influence central circuits via ascending pathways."
    )
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = DGRAY

    # Build modality → regions map
    modality_regions: dict = defaultdict(list)
    for reg in regions:
        mods_raw = reg.get("modalities", reg.get("modalities_that_can_target", ""))
        abbr = reg.get("abbreviation", reg.get("region", ""))
        for mod in [m.strip() for m in mods_raw.split(",") if m.strip()]:
            modality_regions[mod].append(abbr)

    # Build rows in canonical order, then any extras
    seen = set()
    ordered_mods = []
    for m in MODALITY_ORDER:
        if m in modality_regions:
            ordered_mods.append(m)
            seen.add(m)
    for m in sorted(modality_regions.keys()):
        if m not in seen:
            ordered_mods.append(m)

    headers = ["Modality", "Reachable Regions", "Max Depth"]
    n_rows = 1 + len(ordered_mods)
    tbl = doc.add_table(rows=n_rows, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl, "CCCCCC", "4")

    # Header
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _set_cell_bg(cell, "1B3A5C")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    # Data rows
    for ri, mod in enumerate(ordered_mods):
        drow = tbl.rows[ri + 1]
        row_bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
        depth = MODALITY_DEPTH.get(mod, "Cortical")
        reachable = ", ".join(sorted(modality_regions[mod]))

        vals = [mod, reachable, depth]
        for ci, val in enumerate(vals):
            cell = drow.cells[ci]
            _set_cell_bg(cell, row_bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(8.5)
            r.font.name = "Calibri"
            r.font.color.rgb = NAVY
            if ci == 0:
                r.bold = True

    _col_widths(tbl, [2.5, 11.5, 3.0])
    _spacer(doc, 8)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Load data
    data_path = ROOT / "scripts" / "excel_knowledge_dump.json"
    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    regions = data.get("brain_regions", [])
    n_regions = len(regions)

    # Output path
    out_dir = ROOT / "outputs" / "documents" / "shared"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "SOZO_Brain_Regions_Reference.docx"

    # Build document
    doc = Document()
    _set_margins(doc, top=2.0, bottom=2.0, left=2.0, right=2.0)

    _build_title(doc)
    _build_intro(doc, n_regions)
    _build_quick_table(doc, regions)
    _build_detail_cards(doc, regions)
    _build_modality_summary(doc, regions)

    doc.save(str(out_path))

    # Approximate page count: count table rows + loose paragraphs
    total_rows = sum(len(tbl.rows) for tbl in doc.tables)
    total_paras = len(doc.paragraphs)
    # ~3 table rows per "line", ~50 lines per page
    approx_pages = max(1, (total_rows * 3 + total_paras * 2) // 100)

    print(f"Brain Regions Reference generated: {out_path} \u2014 {n_regions} regions, ~{approx_pages} pages")


if __name__ == "__main__":
    main()
