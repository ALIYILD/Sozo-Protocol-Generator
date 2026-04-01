"""
Generate Partners Tier Psychological_Intake_PRS_Baseline_Partners DOCX files
Batch 2: ptsd, ocd, ms, asd, long_covid, tinnitus, insomnia (conditions 8-14)
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK = RGBColor(0x00, 0x00, 0x00)


def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None
    size = fr.font.size if fr else None
    try:
        color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except:
        color = None
    for r in para.runs:
        r.text = ""
    if fr:
        fr.text = replaced
        fr.font.bold = bold
        if size:
            fr.font.size = size
        if color:
            fr.font.color.rgb = color
    else:
        para.add_run(replaced)


def _para_set(para, new_text):
    _para_replace(para, "".join(r.text for r in para.runs), new_text)


def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text)
    run.font.bold = bold
    if white:
        run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type:
            run.font.color.rgb = C_BLACK


def _replace_table(table, rows):
    for ri in range(len(table.rows)):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells:
                _cell_write(cell, "")
            continue
        is_hdr = (ri == 0)
        for ci, cell in enumerate(row.cells):
            _cell_write(
                cell,
                str(rows[ri][ci]) if ci < len(rows[ri]) else "",
                bold=is_hdr,
                white=is_hdr,
            )


def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs):
            _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs):
                        _para_replace(para, old, new)


def build_prs(c):
    TEMPLATE = Path(
        r"C:/Users/yildi/OneDrive/Desktop/Parkinson D/Partners/Assessments/Psychological_Intake_PRS_Baseline_Partners.docx"
    )
    doc = Document(str(TEMPLATE))
    paras = doc.paragraphs
    tables = doc.tables
    for old, new in [
        ("Parkinson's Disease (PD)", c["full"]),
        ("Parkinson's Disease", c["full"]),
        ("Parkinson's", c["short"]),
        (" (PD)", f" ({c['abbr']})"),
        (" PD ", f" {c['abbr']} "),
        (" PD.", f" {c['abbr']}."),
        (" PD,", f" {c['abbr']},"),
        ("for PD", f"for {c['abbr']}"),
        ("PD Phenotype", f"{c['abbr']} Phenotype"),
        ("Hoehn & Yahr", "disease severity rating"),
        ("levodopa", c.get("med_name", "primary medication")),
    ]:
        _global_replace(doc, old, new)
    _para_set(paras[1], f"{c['short']} Psychological Intake & PRS Baseline")
    _para_set(paras[21], "B1. Primary Symptoms Domain")
    _replace_table(tables[3], c["primary_symptoms"])
    _cell_write(tables[5].rows[2].cells[1], c["phenotype_checkboxes"])
    slug = c["slug"]
    out = (
        Path("outputs/documents")
        / slug
        / "partners"
        / f"Psychological_Intake_PRS_Baseline_Partners_{slug}.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  [OK] {out}")
    return out


CONDITIONS = [
    {
        "slug": "ptsd",
        "short": "PTSD",
        "abbr": "PTSD",
        "full": "Post-Traumatic Stress Disorder (PTSD)",
        "med_name": "SSRI/prazosin",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Flashbacks / intrusive memories", "", "\u25a1", ""],
            ["Nightmares / trauma-related dreams", "", "\u25a1", ""],
            ["Hypervigilance / exaggerated startle", "", "\u25a1", ""],
            ["Avoidance of trauma reminders", "", "\u25a1", ""],
            ["Emotional numbing / detachment", "", "\u25a1", ""],
            ["Irritability / anger outbursts", "", "\u25a1", ""],
            ["Sleep disturbance", "", "\u25a1", ""],
            ["Concentration difficulties", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Re-experiencing (RE)  \u25a1 Avoidance (AV)  \u25a1 Hyperarousal (HY)  \u25a1 DI  \u25a1 CX  \u25a1 TB  \u25a1 DE",
    },
    {
        "slug": "ocd",
        "short": "OCD",
        "abbr": "OCD",
        "full": "Obsessive-Compulsive Disorder (OCD)",
        "med_name": "SSRI/clomipramine",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Obsession severity / intrusiveness", "", "\u25a1", ""],
            ["Time spent on compulsions (hrs/day)", "", "\u25a1", ""],
            ["Distress caused by obsessions", "", "\u25a1", ""],
            ["Ability to resist compulsions", "", "\u25a1", ""],
            ["Avoidance of triggering situations", "", "\u25a1", ""],
            ["Functional / occupational interference", "", "\u25a1", ""],
            ["Insight into OCD irrationality", "", "\u25a1", ""],
            ["Tic / motor symptom severity", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Contamination (CO)  \u25a1 Symmetry (SY)  \u25a1 Harm (HA)  \u25a1 HO  \u25a1 PO  \u25a1 TR  \u25a1 TI",
    },
    {
        "slug": "ms",
        "short": "MS",
        "abbr": "MS",
        "full": "Multiple Sclerosis (MS)",
        "med_name": "disease-modifying therapy",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Motor weakness / spasticity", "", "\u25a1", ""],
            ["Fatigue severity (MS-related)", "", "\u25a1", ""],
            ["Walking / gait impairment", "", "\u25a1", ""],
            ["Cognitive difficulties (processing speed)", "", "\u25a1", ""],
            ["Central neuropathic pain", "", "\u25a1", ""],
            ["Bladder / bowel symptoms", "", "\u25a1", ""],
            ["Depression / mood disturbance", "", "\u25a1", ""],
            ["Balance / coordination impairment", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Motor (MO)  \u25a1 Fatigue (FA)  \u25a1 Cognitive (CG)  \u25a1 AT  \u25a1 PA  \u25a1 DE  \u25a1 WA",
    },
    {
        "slug": "asd",
        "short": "ASD",
        "abbr": "ASD",
        "full": "Autism Spectrum Disorder (ASD)",
        "med_name": "atypical antipsychotic",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Social interaction difficulties", "", "\u25a1", ""],
            ["Communication challenges", "", "\u25a1", ""],
            ["Repetitive behaviours / rituals", "", "\u25a1", ""],
            ["Sensory hyper/hyposensitivities", "", "\u25a1", ""],
            ["Flexibility / transitions difficulty", "", "\u25a1", ""],
            ["Emotional regulation difficulty", "", "\u25a1", ""],
            ["Executive function challenges", "", "\u25a1", ""],
            ["Daily living skills impairment", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Social (SO)  \u25a1 Repetitive (RE)  \u25a1 Language (LA)  \u25a1 SE  \u25a1 EX  \u25a1 EM  \u25a1 AD",
    },
    {
        "slug": "long_covid",
        "short": "Long COVID",
        "abbr": "LC",
        "full": "Long COVID / Post-COVID Syndrome (LC)",
        "med_name": "symptomatic treatment",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Brain fog / cognitive difficulties", "", "\u25a1", ""],
            ["Post-exertional fatigue / PEM", "", "\u25a1", ""],
            ["Breathlessness / exercise intolerance", "", "\u25a1", ""],
            ["Sleep disruption", "", "\u25a1", ""],
            ["Headache", "", "\u25a1", ""],
            ["Mood / anxiety symptoms", "", "\u25a1", ""],
            ["Palpitations / chest symptoms", "", "\u25a1", ""],
            ["Pain / musculoskeletal symptoms", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Cognitive (CG)  \u25a1 Fatigue (FA)  \u25a1 Dysautonomia (DA)  \u25a1 NP  \u25a1 HA  \u25a1 EI  \u25a1 SL",
    },
    {
        "slug": "tinnitus",
        "short": "Tinnitus",
        "abbr": "TIN",
        "full": "Chronic Tinnitus (TIN)",
        "med_name": "tinnitus medication",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Tinnitus perceived loudness", "", "\u25a1", ""],
            ["Tinnitus-related distress / annoyance", "", "\u25a1", ""],
            ["Sleep disruption from tinnitus", "", "\u25a1", ""],
            ["Concentration difficulty", "", "\u25a1", ""],
            ["Emotional impact (anxiety / depression)", "", "\u25a1", ""],
            ["Hyperacusis / sound sensitivity", "", "\u25a1", ""],
            ["Avoidance of sound environments", "", "\u25a1", ""],
            ["Daily activity interference", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Tonal (TO)  \u25a1 Noise (NO)  \u25a1 Pulsatile (PU)  \u25a1 RE  \u25a1 SO  \u25a1 HY  \u25a1 DE",
    },
    {
        "slug": "insomnia",
        "short": "Insomnia",
        "abbr": "INS",
        "full": "Chronic Insomnia Disorder (INS)",
        "med_name": "hypnotic/sleep medication",
        "primary_symptoms": [
            ["Symptom", "Score (0–10)", "Relevant", "Notes"],
            ["Sleep-onset difficulty (SOL)", "", "\u25a1", ""],
            ["Night waking frequency (WASO)", "", "\u25a1", ""],
            ["Early morning awakening", "", "\u25a1", ""],
            ["Daytime fatigue / low energy", "", "\u25a1", ""],
            ["Daytime sleepiness", "", "\u25a1", ""],
            ["Cognitive impairment (daytime)", "", "\u25a1", ""],
            ["Mood effects of insomnia", "", "\u25a1", ""],
            ["Worry / anxiety about sleep", "", "\u25a1", ""],
        ],
        "phenotype_checkboxes": "\u25a1 Sleep-Onset (SO)  \u25a1 Maintenance (MA)  \u25a1 Early-Morning (EM)  \u25a1 HY  \u25a1 PS  \u25a1 PA  \u25a1 CI",
    },
]

if __name__ == "__main__":
    print("Generating Partners Tier — Psychological_Intake_PRS_Baseline — Batch 2")
    print("=" * 70)
    outputs = []
    for c in CONDITIONS:
        outputs.append(build_prs(c))
    print()
    print("Verifying output files:")
    print("-" * 70)
    all_ok = True
    for out in outputs:
        p = Path(out)
        if p.exists():
            size = p.stat().st_size
            print(f"  [PASS] {p.name:60s}  {size:>8,} bytes")
        else:
            print(f"  [FAIL] {p.name} — NOT FOUND")
            all_ok = False
    print()
    print(f"Result: {'ALL 7 FILES OK' if all_ok else 'SOME FILES MISSING'}")
