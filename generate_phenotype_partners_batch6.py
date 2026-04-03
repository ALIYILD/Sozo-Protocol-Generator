# DEPRECATED: This script is superseded by the canonical generation pipeline.
# Use instead: GenerationService.generate(condition="...", tier="...", doc_type="...")
# Or CLI: PYTHONPATH=src python -m sozo_generator.cli.main build condition --condition <slug> --tier <tier> --doc-type <type>
# See docs/MIGRATION_PLAN.md for details.

from pathlib import Path
from docx import Document
from docx.shared import RGBColor

_PROJECT_ROOT = Path(__file__).resolve().parent

C_WHITE = RGBColor(0xFF,0xFF,0xFF); C_BLACK = RGBColor(0x00,0x00,0x00)

def _para_replace(para, old, new):
    full = "".join(r.text for r in para.runs)
    if old not in full: return
    replaced = full.replace(old, new)
    fr = para.runs[0] if para.runs else None
    bold = fr.font.bold if fr else None; size = fr.font.size if fr else None
    try: color = fr.font.color.rgb if (fr and fr.font.color.type) else None
    except: color = None
    for r in para.runs: r.text = ""
    if fr:
        fr.text = replaced; fr.font.bold = bold
        if size: fr.font.size = size
        if color: fr.font.color.rgb = color
    else: para.add_run(replaced)

def _para_set(para, new_text): _para_replace(para, "".join(r.text for r in para.runs), new_text)

def _cell_write(cell, text, bold=False, white=False):
    for p in cell.paragraphs:
        for r in p.runs: r.text = ""
    para = cell.paragraphs[0]; run = para.runs[0] if para.runs else para.add_run()
    run.text = str(text); run.font.bold = bold
    if white: run.font.color.rgb = C_WHITE
    else:
        if run.font.color.type: run.font.color.rgb = C_BLACK

def _replace_table(table, rows):
    for ri in range(len(table.rows)):
        row = table.rows[ri]
        if ri >= len(rows):
            for cell in row.cells: _cell_write(cell, ""); continue
        is_hdr = (ri == 0)
        for ci, cell in enumerate(row.cells):
            _cell_write(cell, str(rows[ri][ci]) if ci < len(rows[ri]) else "", bold=is_hdr, white=is_hdr)

def _global_replace(doc, old, new):
    for para in doc.paragraphs:
        if old in "".join(r.text for r in para.runs): _para_replace(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in "".join(r.text for r in para.runs): _para_replace(para, old, new)

def build(c):
    TEMPLATE = _PROJECT_ROOT / "templates" / "gold_standard" / "Phenotype_Classification.docx"
    doc = Document(str(TEMPLATE)); paras = doc.paragraphs; tables = doc.tables
    for old, new in [
        ("Parkinson's Disease (PD)", c["full"]), ("Parkinson's Disease", c["full"]),
        ("Parkinson's", c["short"]), (" (PD)", f" ({c['abbr']})"),
        (" PD ", f" {c['abbr']} "), (" PD.", f" {c['abbr']}."),
        (" PD,", f" {c['abbr']},"), ("for PD", f"for {c['abbr']}"),
        ("H&Y Stage", c["severity_field"]), ("Disease Duration", c["duration_field"]),
        ("PD Phenotype", f"{c['abbr']} Phenotype"), ("Hoehn & Yahr", "disease severity rating"),
        ("levodopa", c.get("med_name","primary medication")),
    ]: _global_replace(doc, old, new)
    _para_set(paras[1], f"{c['short']} Phenotype Classification & Protocol Mapping")
    _para_set(paras[7], f"STEP 1 \u2014 {c['short']} Phenotype Determination")
    _para_set(paras[11], f"1.2 {c['short']} Phenotype Classification")
    _para_set(paras[14], f"ASSIGNED PHENOTYPE: {c['phenotype_codes']}")
    _para_set(paras[19], c["g1"]); _para_set(paras[21], c["g2"]); _para_set(paras[23], c["g3"])
    _para_set(paras[27], c["tps_ol"]); _para_set(paras[40], c["sel_opt"])
    _replace_table(tables[0],[["Field","Details"],["Patient Name",""],["File Number",""],[c["severity_field"],""],["Date",""],["Treating Doctor",""],[c["duration_field"],"___ years / ___ months"]])
    for ti, key in [(1,"dc"),(2,"ph"),(3,"nw"),(5,"t5"),(6,"t6"),(7,"t7"),(9,"t9"),(12,"tavns"),(13,"ces"),(14,"seq")]:
        _replace_table(tables[ti], c[key])
    _cell_write(tables[15].rows[1].cells[0], c["resp"])
    _cell_write(tables[15].rows[1].cells[1], "\u25a1 Completed | Score: ___ / baseline: ___")
    _cell_write(tables[18].rows[1].cells[0], f"{c['abbr']} Phenotype")
    out = Path("outputs/documents") / c["slug"] / "partners" / f"Phenotype_Classification_Partners_{c['slug']}.docx"
    out.parent.mkdir(parents=True, exist_ok=True); doc.save(str(out)); print(f"  [OK] {out}"); return out

ASD = {
    "slug":"asd","short":"ASD","abbr":"ASD","full":"Autism Spectrum Disorder (ASD)",
    "severity_field":"ADOS-2 / SRS-2 Severity","duration_field":"ASD Duration",
    "med_name":"atypical antipsychotic",
    "phenotype_codes":"\u25a1 SO  \u25a1 RE  \u25a1 LA  \u25a1 SE  \u25a1 EX  \u25a1 EM  \u25a1 AD",
    "g1":"3.1 Social / Language Phenotypes (SO / LA)","g2":"3.2 Repetitive / Sensory (RE / SE)",
    "g3":"3.3 Executive / Emotional / ADHD-Comorbid (EX / EM / AD)",
    "tps_ol":"TPS is OFF-LABEL for ASD. Requires Doctor authorisation and off-label consent.",
    "sel_opt":"SELECTED OPTION: \u25a1 A (Social/Language)  \u25a1 B (Repetitive/Sensory)  \u25a1 C (Executive)  \u25a1 D (Mixed)",
    "resp":"Reassess SOZO PRS + SRS-2 / VABS",
    "dc":[
        ["Data Point","Completed","Source"],
        ["ADOS-2 / ADI-R diagnostic assessment","\u25a1","Specialist Assessment"],
        ["Social communication profile (SRS-2, SCQ)","\u25a1","Rating Scales"],
        ["Repetitive behaviour severity (RBS-R, RBSER)","\u25a1","Rating Scales"],
        ["Language and communication level (VABS)","\u25a1","VABS / SLP Assessment"],
        ["Sensory processing profile (SPM, SSP)","\u25a1","Sensory Scales"],
        ["Executive function (BRIEF-A, ToL)","\u25a1","Neuropsychological"],
        ["ADHD comorbidity (ADHD-RS, Conners)","\u25a1","Rating Scales"],
        ["6-Network Bedside Assessment scores","\u25a1","Assessment Form"],
    ],
    "ph":[
        ["Dominant Features","Assign Phenotype","Code"],
        ["Social reciprocity deficits, reduced joint attention, peer relationship difficulty","Social (SO)","SO"],
        ["Restricted interests, stereotypies, insistence on sameness dominant","Repetitive (RE)","RE"],
        ["Language delay, communication difficulties, pragmatic deficits","Language (LA)","LA"],
        ["Sensory hyper/hyposensitivity, sensory seeking/avoidance dominant","Sensory (SE)","SE"],
        ["Executive dysfunction: planning, flexibility, initiation foregrounded","Executive (EX)","EX"],
        ["Emotional dysregulation, meltdowns, affect modulation difficulty","Emotional (EM)","EM"],
        ["ADHD comorbidity dominant: inattention/hyperactivity alongside ASD","ADHD-Comorbid (AD)","AD"],
    ],
    "nw":[
        ["Phenotype","Primary Network","Secondary Network","Tertiary"],
        ["SO","Social brain (R-TPJ/mPFC)","Limbic (amygdala)","CEN"],
        ["RE","Fronto-striatal (SMA)","CEN","SN"],
        ["LA","L-temporal (STG/IFG)","CEN (L-DLPFC)","DMN"],
        ["SE","SN (Insula)","SMN","Thalamic gating"],
        ["EX","CEN (L-DLPFC)","Fronto-striatal","DAN"],
        ["EM","SN (dACC/Insula)","Limbic (amygdala)","CEN"],
        ["AD","CEN (R-DLPFC)","SN","DMN"],
    ],
    "t5":[
        ["Goal","Anode","Cathode"],
        ["L-DLPFC social cognition enhancement (SO)","Left DLPFC (F3)","Right supraorbital"],
        ["R-TPJ social processing (SO/EM)","Right TPJ (P4)","Left supraorbital"],
        ["L-IFG language facilitation (LA)","Left IFG (F7)","Right supraorbital"],
    ],
    "t6":[
        ["Goal","Anode","Cathode"],
        ["Pre-SMA repetitive behaviour modulation (RE)","Left DLPFC","Pre-SMA cathodal"],
        ["Insula sensory gating (SE)","Right DLPFC cathodal","Left DLPFC anodal"],
    ],
    "t7":[
        ["Goal","Anode","Cathode"],
        ["R-DLPFC ADHD-executive (AD/EX)","Right DLPFC (F4)","Left supraorbital"],
        ["Emotional regulation DLPFC (EM)","Left DLPFC","Right DLPFC cathodal"],
    ],
    "t9":[
        ["Clinical Target","TPS Protocol","Phenotype"],
        ["L-DLPFC social-executive network","FT1","SO, EX"],
        ["R-TPJ social processing","FT3","SO, EM"],
        ["L-IFG/Broca language circuit","FT4","LA"],
        ["Pre-SMA repetitive behaviour","FT5","RE"],
        ["Insula sensory processing","FT7","SE"],
        ["Limbic emotional regulation","FT8","EM"],
        ["R-DLPFC ADHD circuit","FT9","AD"],
        ["Bilateral temporal-parietal","Multiple","LA / SO"],
        ["Multi-network distributed ASD","Multiple","MX / complex"],
    ],
    "tavns":[
        ["Domain","Details"],
        ["Indications","Emotional dysregulation, autonomic arousal, sleep disruption in ASD"],
        ["Protocol","20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
        ["Add taVNS?","\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "ces":[
        ["Domain","Details"],
        ["Indications","Anxiety comorbidity, insomnia, sensory overload, emotional dysregulation in ASD"],
        ["Protocol","20\u201360 min daily, Alpha-Stim AID, especially SE and EM phenotypes"],
        ["Add CES?","\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "seq":[
        ["Option","Sequence","Phenotype"],
        ["OPTION A \u2014 Social/Language","Wk 1\u20134: L-DLPFC + R-TPJ tDCS during social tasks (5\u00d7/wk) \u2192 Wk 2\u20136: TPS TPJ/DLPFC","SO / LA"],
        ["OPTION B \u2014 Repetitive/Sensory","Wk 1\u20134: Pre-SMA cathodal tDCS + Daily CES \u2192 Wk 3\u20136: TPS Pre-SMA/Insula","RE / SE"],
        ["OPTION C \u2014 Executive","Wk 1\u20132: CES/taVNS stabilization \u2192 Wk 2\u20136: R-DLPFC tDCS + TPS CEN","EX / AD"],
        ["OPTION D \u2014 Emotional/Mixed","Sequential: TPS (highest burden) \u2192 Maintenance L-DLPFC + R-TPJ tDCS","EM / MX"],
    ],
}

LONG_COVID = {
    "slug":"long_covid","short":"Long COVID","abbr":"LC",
    "full":"Long COVID / Post-COVID Syndrome (LC)",
    "severity_field":"Symptom Severity Score / DSQ","duration_field":"Post-COVID Duration",
    "med_name":"symptomatic treatment",
    "phenotype_codes":"\u25a1 CG  \u25a1 FA  \u25a1 DA  \u25a1 NP  \u25a1 HA  \u25a1 EI  \u25a1 SL",
    "g1":"3.1 Cognitive / Neuropsychiatric (CG / NP)","g2":"3.2 Fatigue / Exercise Intolerance (FA / EI)",
    "g3":"3.3 Dysautonomia / Headache / Sleep (DA / HA / SL)",
    "tps_ol":"TPS is OFF-LABEL for Long COVID. Requires Doctor authorisation and off-label consent.",
    "sel_opt":"SELECTED OPTION: \u25a1 A (Brain Fog/Cognitive)  \u25a1 B (Fatigue)  \u25a1 C (Dysautonomia)  \u25a1 D (Mixed)",
    "resp":"Reassess SOZO PRS + DSQ / MoCA / FSS",
    "dc":[
        ["Data Point","Completed","Source"],
        ["Post-COVID symptom duration and onset","\u25a1","Medical History"],
        ["Cognitive / brain fog assessment (MoCA, COGSTATE)","\u25a1","Neuropsychological"],
        ["Fatigue severity (FSS, MFIS, PEM assessment)","\u25a1","Fatigue Scales"],
        ["Dysautonomia screen (orthostatic HR/BP, COMPASS-31)","\u25a1","Autonomic Assessment"],
        ["Neuropsychiatric symptoms (PHQ-9, GAD-7, PCL-5)","\u25a1","Rating Scales"],
        ["Sleep quality (ISI, PSQI)","\u25a1","Sleep Scales"],
        ["Exercise tolerance / PEM diary","\u25a1","Activity Diary"],
        ["6-Network Bedside Assessment scores","\u25a1","Assessment Form"],
    ],
    "ph":[
        ["Dominant Features","Assign Phenotype","Code"],
        ["Brain fog: concentration, memory, processing speed impairment dominant","Cognitive (CG)","CG"],
        ["Post-exertional malaise, persistent fatigue, low stamina dominant","Fatigue (FA)","FA"],
        ["POTS, orthostatic intolerance, autonomic dysfunction dominant","Dysautonomia (DA)","DA"],
        ["Anxiety, depression, PTSD-like features post-COVID dominant","Neuropsychiatric (NP)","NP"],
        ["New-onset chronic headache or migraine post-COVID","Headache (HA)","HA"],
        ["Post-exertional intolerance, exercise limitation, PEM","Exercise Intolerance (EI)","EI"],
        ["Sleep disruption, hypersomnia or insomnia post-COVID","Sleep (SL)","SL"],
    ],
    "nw":[
        ["Phenotype","Primary Network","Secondary Network","Tertiary"],
        ["CG","CEN (bilateral DLPFC)","DMN","DAN"],
        ["FA","CEN (L-DLPFC)","SN","DMN"],
        ["DA","SN (Insula/autonomic cortex)","SMN","CEN"],
        ["NP","CEN (L-DLPFC)","Limbic","SN"],
        ["HA","SN (dACC/Insula)","Thalamic","CEN"],
        ["EI","SN + SMN","CEN","DMN"],
        ["SL","SN (arousal)","CEN (R-DLPFC cathodal)","DMN"],
    ],
    "t5":[
        ["Goal","Anode","Cathode"],
        ["Bilateral DLPFC brain fog (CG/FA)","Left DLPFC (F3)","Right supraorbital"],
        ["R-DLPFC cathodal arousal reduction (SL)","Right DLPFC cathodal","Left DLPFC anodal"],
        ["L-DLPFC neuropsychiatric (NP/FA)","Left DLPFC","Right supraorbital"],
    ],
    "t6":[
        ["Goal","Anode","Cathode"],
        ["SN/Insula dysautonomia modulation (DA)","Left DLPFC","Right DLPFC cathodal"],
        ["Thalamic headache modulation (HA)","Left DLPFC","Right DLPFC cathodal"],
    ],
    "t7":[
        ["Goal","Anode","Cathode"],
        ["Bilateral DLPFC fatigue/PEM (EI/FA)","Bilateral DLPFC","Extracephalic"],
        ["L-DLPFC antidepressant (NP-depression)","Left DLPFC (F3)","Right supraorbital"],
    ],
    "t9":[
        ["Clinical Target","TPS Protocol","Phenotype"],
        ["L-DLPFC brain fog / CEN","FT1","CG, FA, NP"],
        ["Bilateral DLPFC cognitive network","FT3","CG"],
        ["Insula/SN dysautonomia","FT4","DA"],
        ["Thalamic headache circuit","FT5","HA"],
        ["Limbic-prefrontal neuropsychiatric","FT7","NP"],
        ["R-DLPFC cathodal arousal (sleep)","FT8","SL"],
        ["Bilateral frontal fatigue","FT9","FA, EI"],
        ["SN + CEN combined","Multiple","DA / EI"],
        ["Multi-network Long COVID","Multiple","MX / severe"],
    ],
    "tavns":[
        ["Domain","Details"],
        ["Indications","Dysautonomia, fatigue, sleep disruption, neuropsychiatric \u2014 all LC phenotypes"],
        ["Protocol","20\u201330 min daily, 25 Hz, 250 \u00b5s, left tragus, 0.5 mA"],
        ["Add taVNS?","\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "ces":[
        ["Domain","Details"],
        ["Indications","Anxiety, insomnia, neuropsychiatric comorbidity \u2014 adjunct to all LC phenotypes"],
        ["Protocol","20\u201360 min daily, Alpha-Stim AID, especially NP, SL, FA phenotypes"],
        ["Add CES?","\u25a1 Yes  \u25a1 No | Rationale: __________"],
    ],
    "seq":[
        ["Option","Sequence","Phenotype"],
        ["OPTION A \u2014 Brain Fog","Wk 1\u20134: L-DLPFC tDCS (5\u00d7/wk, low intensity) \u2192 Wk 2\u20136: TPS DLPFC/CEN","CG / NP"],
        ["OPTION B \u2014 Fatigue","Wk 1\u20134: L-DLPFC tDCS + Daily taVNS \u2192 Wk 3\u20136: TPS frontal fatigue circuit","FA / EI"],
        ["OPTION C \u2014 Dysautonomia","Wk 1\u20132: CES/taVNS stabilization \u2192 Wk 2\u20136: SN tDCS + TPS Insula/autonomic","DA / HA"],
        ["OPTION D \u2014 Mixed","Sequential: TPS (highest burden) \u2192 Maintenance bilateral DLPFC tDCS + taVNS","SL / MX"],
    ],
}

if __name__ == "__main__":
    import os
    os.chdir(str(_PROJECT_ROOT))
    print("Building ASD...")
    asd_out = build(ASD)
    print("Building Long COVID...")
    lc_out = build(LONG_COVID)
    print("\nVerifying outputs:")
    for p in [asd_out, lc_out]:
        fp = Path(p)
        if fp.exists():
            print(f"  EXISTS  {fp}  ({fp.stat().st_size:,} bytes)")
        else:
            print(f"  MISSING {fp}")
